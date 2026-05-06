#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
except ModuleNotFoundError as exc:
    raise SystemExit(
        "python-docx is required. Install it with `python3 -m pip install python-docx` and rerun."
    ) from exc


LANGUAGE_ALIASES = {
    "zh": "zh-CN",
    "zh-cn": "zh-CN",
    "zh-hans": "zh-CN",
    "cn": "zh-CN",
    "chinese": "zh-CN",
    "中文": "zh-CN",
    "en": "en-US",
    "en-us": "en-US",
    "en-gb": "en-US",
    "english": "en-US",
    "英文": "en-US",
}


L10N = {
    "zh-CN": {
        "report_suffix": "漏洞报告.docx",
        "report_title": "{project_name} 库存在 {vuln_type} 漏洞报告",
        "attachment_notes_suffix": "_附件目录说明.md",
        "supplement_note_suffix": "_补充复现说明.md",
        "impact_versions": "影响版本",
        "vuln_description": "漏洞描述",
        "risk_assessment": "漏洞危险性评估",
        "analysis": "漏洞分析",
        "reproduction": "漏洞复现",
        "final_verdict": "最终判定：",
        "code_context": "关键代码上下文",
        "env_files": "验证环境关键文件",
        "repro_pending": "复现步骤待补充。",
        "related_artifacts": "相关文件/产物：",
        "commands": "执行命令：",
        "expected": "预期结果：",
        "observed": "实际结果：",
        "notes": "补充说明：",
        "results": "结果证据：",
        "attachment_notes_title": "# 附件目录说明",
        "report_for": "对应报告：{name}",
        "directory_structure": "## 目录结构",
        "report_body": "- `{name}`：漏洞确认报告正文。",
        "attachments_dir": "- `attachments/`：复现所需附件目录。",
        "attachment_usage": "## 附件用途说明",
        "supplement_title": "# {project_name} {vuln_type} 漏洞补充复现说明",
        "supplement_goal": "## 一、补充目的",
        "supplement_env": "## 二、复现环境",
        "supplement_shortest_path": "## 三、建议的最短复现路径",
        "supplement_evidence": "## 四、关键成功证据",
        "supplement_bundle": "## 五、补充材料说明",
        "supplement_conclusion": "## 六、结论",
        "supplement_docker_only": "本次补充材料默认以 Docker 或 Docker Compose 方式完成验证，不应回退到宿主机直接执行攻击流量。",
        "supplement_shortest_with_script": "建议优先使用打包目录根部的最小复现脚本：`{script}`",
        "supplement_shortest_without_script": "若未提供 bundle 根复现脚本，请按报告中的 Docker 复现步骤逐条执行。",
        "supplement_bundle_scripts": "补充材料中包含可直接录屏或复核的 bundle 根脚本：",
        "supplement_bundle_attachments": "复现依赖的附件文件保存在 `attachments/` 目录内。",
        "supplement_evidence_missing": "关键成功证据待补充。建议明确写出能够证明“漏洞复现成功”或“攻击成功”的输出片段。",
        "supplement_conclusion_default": "经 Docker 隔离环境复现，报告中所述问题可以稳定验证；建议在提交时同步附上最小复现脚本、附件目录说明及录屏视频。",
        "attachment_missing": "用途待补充。",
        "original_path": "- 原始路径：`{path}`",
        "purpose": "- 用途：{purpose}",
        "purpose_inline": "用途：{purpose}",
        "note_inline": "说明：{note}",
        "analysis_pending": "漏洞分析待补充。",
        "assessment_pending": "评估依据待补充。",
        "verdict_pending": "最终判定待补充。",
        "impact_package": "影响包",
        "impact_component": "影响组件",
        "impact_versions_label": "影响版本",
        "repo_url": "仓库链接",
        "cvss4_vector": "CVSS 4.0 向量",
        "cvss31_vector": "CVSS 3.1 向量",
        "base_score": "基础评分",
        "severity": "等级判定",
        "code_context_item": "{idx}. 代码上下文",
    },
    "en-US": {
        "report_suffix": "_report.docx",
        "report_title": "{project_name} {vuln_type} Vulnerability Report",
        "attachment_notes_suffix": "_attachment_index.md",
        "supplement_note_suffix": "_reproduction_note.md",
        "impact_versions": "Affected Versions",
        "vuln_description": "Vulnerability Description",
        "risk_assessment": "Risk Assessment",
        "analysis": "Vulnerability Analysis",
        "reproduction": "Reproduction",
        "final_verdict": "Final Verdict:",
        "code_context": "Key Code Context",
        "env_files": "Key Verification Environment Files",
        "repro_pending": "Reproduction steps are not provided yet.",
        "related_artifacts": "Related Files / Artifacts:",
        "commands": "Commands:",
        "expected": "Expected Result:",
        "observed": "Observed Result:",
        "notes": "Additional Notes:",
        "results": "Evidence:",
        "attachment_notes_title": "# Attachment Index",
        "report_for": "Report: {name}",
        "directory_structure": "## Directory Structure",
        "report_body": "- `{name}`: vulnerability report document.",
        "attachments_dir": "- `attachments/`: bundled files required for reproduction.",
        "attachment_usage": "## Attachment Usage",
        "supplement_title": "# {project_name} {vuln_type} Reproduction Supplement",
        "supplement_goal": "## 1. Purpose",
        "supplement_env": "## 2. Verification Environment",
        "supplement_shortest_path": "## 3. Recommended Shortest Reproduction Path",
        "supplement_evidence": "## 4. Key Success Evidence",
        "supplement_bundle": "## 5. Bundled Materials",
        "supplement_conclusion": "## 6. Conclusion",
        "supplement_docker_only": "This supplement is intended for Docker or Docker Compose based verification and must not fall back to direct host-side exploit traffic.",
        "supplement_shortest_with_script": "Prefer the bundle-root minimal reproduction helper script: `{script}`",
        "supplement_shortest_without_script": "If no bundle-root helper script is provided, follow the Docker reproduction steps in the report in the same order.",
        "supplement_bundle_scripts": "The bundle includes root-level scripts for direct screen recording or reviewer re-check:",
        "supplement_bundle_attachments": "Supporting reproduction artifacts are bundled under `attachments/`.",
        "supplement_evidence_missing": "Key success evidence is not populated yet. Add exact output lines proving successful reproduction or successful exploitation.",
        "supplement_conclusion_default": "Docker-based verification confirms that the reported issue can be reproduced reliably. Ship the minimal reproduction script, attachment index, and screen-recording proof together when submitting.",
        "attachment_missing": "Purpose is not provided yet.",
        "original_path": "- Original Path: `{path}`",
        "purpose": "- Purpose: {purpose}",
        "purpose_inline": "Purpose: {purpose}",
        "note_inline": "Note: {note}",
        "analysis_pending": "Analysis is not provided yet.",
        "assessment_pending": "Assessment rationale is not provided yet.",
        "verdict_pending": "Final verdict is not provided yet.",
        "impact_package": "Package",
        "impact_component": "Component",
        "impact_versions_label": "Affected Versions",
        "repo_url": "Repository URL",
        "cvss4_vector": "CVSS 4.0 Vector",
        "cvss31_vector": "CVSS 3.1 Vector",
        "base_score": "Base Score",
        "severity": "Severity",
        "code_context_item": "{idx}. Code Context",
    },
}


def normalize_language(value: Any, default: str = "zh-CN") -> str:
    text = str(value or "").strip()
    if not text:
        return default
    key = text.lower()
    return LANGUAGE_ALIASES.get(key, default if text not in L10N else text)


def language_field_candidates(base_key: str, language: str) -> list[str]:
    if language == "en-US":
        return [
            f"{base_key}_en_us",
            f"{base_key}_en",
            f"{base_key}_english",
        ]
    return [
        f"{base_key}_zh_cn",
        f"{base_key}_zh",
        f"{base_key}_cn",
        f"{base_key}_chinese",
    ]


def normalize_declared_language(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    return normalize_language(text, "")


def should_use_generic_field(language: str, declared_language: str = "") -> bool:
    if declared_language:
        return declared_language == language
    return language == "zh-CN"


def read_localized_string(
    data: dict[str, Any],
    base_key: str,
    language: str,
    *,
    declared_language: str = "",
) -> str:
    for key in language_field_candidates(base_key, language):
        value = str(data.get(key, "")).strip()
        if value:
            return value
    if should_use_generic_field(language, declared_language):
        return str(data.get(base_key, "")).strip()
    return ""


def read_localized_list(
    data: dict[str, Any],
    base_key: str,
    language: str,
    *,
    declared_language: str = "",
) -> list[str]:
    for key in language_field_candidates(base_key, language):
        if key in data:
            return ensure_list(data.get(key))
    if should_use_generic_field(language, declared_language):
        return ensure_list(data.get(base_key))
    return []


def localized_string(data: dict[str, Any], base_key: str, language: str) -> str:
    return read_localized_string(data, base_key, language, declared_language=infer_source_language(data))


def localized_list(data: dict[str, Any], base_key: str, language: str) -> list[str]:
    return read_localized_list(data, base_key, language, declared_language=infer_source_language(data))


def tr(language: str, key: str, **kwargs: Any) -> str:
    template = L10N[language][key]
    return template.format(**kwargs) if kwargs else template


def format_kv(language: str, label: str, value: Any) -> str:
    sep = ":" if language == "en-US" else "："
    return f"{label}{sep} {value}"


def ensure_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value).strip()
    return [text] if text else []


def ensure_mapping(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}


def ensure_relpath(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    if text.startswith("/"):
        return Path(text).name
    return text


def normalize_relative_to_project(value: Any, project_root: Path) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    path = Path(text).expanduser()
    if path.is_absolute():
        resolved_path = path.resolve()
        try:
            return resolved_path.relative_to(project_root.resolve()).as_posix()
        except ValueError:
            return resolved_path.name
    return Path(text).as_posix().lstrip("./")


def resolve_source_path(value: Any, project_root: Path) -> tuple[str, Path]:
    text = str(value or "").strip()
    if not text:
        return "", project_root
    source_path = Path(text).expanduser()
    if source_path.is_absolute():
        resolved = source_path.resolve()
        try:
            rel = resolved.relative_to(project_root.resolve()).as_posix()
        except ValueError:
            rel = resolved.name
        return rel, resolved
    rel = Path(text).as_posix().lstrip("./")
    return rel, (project_root / rel).resolve()


def slugify(value: str) -> str:
    value = re.sub(r'[\/:*?"<>|]+', "_", value.strip())
    value = re.sub(r"\s+", "_", value)
    return value.strip("_") or "confirmed_vulnerability"


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= char <= "\u9fff" for char in text)


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def choose_template(cli_path: str | None) -> Path | None:
    if cli_path:
        path = Path(cli_path).expanduser()
        if path.exists():
            return path
    script_root = Path(__file__).resolve().parent.parent
    default_paths = [
        script_root / "assets" / "confirmed-vuln-report-template.docx",
        script_root / "confirmed" / "confirmed-vuln-report-template.docx",
    ]
    return next((path for path in default_paths if path.exists()), None)


def infer_source_language(finding: dict[str, Any]) -> str:
    return normalize_declared_language(finding.get("report_language") or finding.get("output_language"))


def looks_like_short_vulnerability_name(value: Any) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    if "\n" in text or "\r" in text:
        return False
    lowered = text.lower()
    if "漏洞报告" in text or "vulnerability report" in lowered:
        return False
    if len(text) > 80:
        return False
    if not contains_cjk(text) and len(text.split()) > 12:
        return False
    return True


def fail_missing_vulnerability_name(finding: dict[str, Any], language: str) -> None:
    title = (
        read_localized_string(finding, "title", language, declared_language=infer_source_language(finding))
        or str(finding.get("title") or finding.get("title_zh") or finding.get("title_en") or "").strip()
    )
    hint = (
        "Confirmed findings must include vulnerability_name/vulnerability_name_zh/vulnerability_name_en "
        "or a short vuln_type/vuln_type_zh/vuln_type_en. Do not derive the bundle identity from a long title."
    )
    if title:
        raise SystemExit(f"{hint} Problem title: {title[:140]}")
    raise SystemExit(hint)


def localized_vuln_type(finding: dict[str, Any], language: str) -> str:
    declared_language = infer_source_language(finding)
    name = read_localized_string(finding, "vulnerability_name", language, declared_language=declared_language)
    if looks_like_short_vulnerability_name(name):
        return name
    if language == "en-US":
        for key in ("vulnerability_name_en",):
            value = str(finding.get(key) or "").strip()
            if looks_like_short_vulnerability_name(value) and not contains_cjk(value):
                return value
    else:
        for key in ("vulnerability_name", "vulnerability_name_zh"):
            value = str(finding.get(key) or "").strip()
            if looks_like_short_vulnerability_name(value) and contains_cjk(value):
                return value
    localized = read_localized_string(finding, "vuln_type", language, declared_language=declared_language)
    if looks_like_short_vulnerability_name(localized):
        return localized
    fallback = str(finding.get("vuln_type", "")).strip()
    if looks_like_short_vulnerability_name(fallback) and should_use_generic_field(language, declared_language):
        return fallback
    if looks_like_short_vulnerability_name(fallback) and language == "en-US" and not contains_cjk(fallback):
        return fallback
    if looks_like_short_vulnerability_name(fallback) and language == "zh-CN" and contains_cjk(fallback):
        return fallback
    fail_missing_vulnerability_name(finding, language)


def build_title(finding: dict[str, Any], language: str) -> str:
    declared_language = infer_source_language(finding)
    title = read_localized_string(finding, "title", language, declared_language=declared_language)
    if title:
        return title
    project_name = str(finding.get("project_name", "目标项目")).strip()
    vuln_type = localized_vuln_type(finding, language)
    return tr(language, "report_title", project_name=project_name, vuln_type=vuln_type)


def default_final_verdict(finding: dict[str, Any], language: str) -> list[str]:
    status = str(finding.get("verification_status") or "").strip()
    evidence = ensure_mapping(finding.get("verification_evidence"))
    if not status:
        status = str(evidence.get("verification_status") or "").strip()
    oracle = (
        localized_string(finding, "oracle", language)
        or str(finding.get("oracle") or evidence.get("oracle_token") or evidence.get("observed_observation") or "").strip()
    )
    remediation = (
        localized_string(finding, "remediation", language)
        or str(finding.get("remediation") or finding.get("fix") or finding.get("recommendation") or "").strip()
    )
    if language == "en-US":
        verdict = (
            "Final verdict: Docker evidence confirms this finding as a real vulnerability."
            if status == "confirmed_in_docker"
            else "Final verdict: this finding has been reviewed with the available verification evidence."
        )
        lines = [verdict]
        if oracle:
            lines.append(f"Success oracle: {oracle}")
        if remediation:
            lines.append(f"Remediation: {remediation}")
        return lines
    verdict = (
        "最终判定：Docker 证据确认该问题为真实漏洞。"
        if status == "confirmed_in_docker"
        else "最终判定：该问题已结合现有验证证据完成复核。"
    )
    lines = [verdict]
    if oracle:
        lines.append(f"成功判据：{oracle}")
    if remediation:
        lines.append(f"修复建议：{remediation}")
    return lines


def severity_cn_from_any(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return "中危"
    lowered = text.lower()
    if "critical" in lowered or "严重" in text:
        return "严重"
    if "high" in lowered or "高危" in text:
        return "高危"
    if "medium" in lowered or "中危" in text:
        return "中危"
    if "low" in lowered or "低危" in text:
        return "低危"
    return "中危"


def severity_en_from_any(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return "Medium"
    lowered = text.lower()
    if "critical" in lowered or "严重" in text:
        return "Critical"
    if "high" in lowered or "高危" in text:
        return "High"
    if "medium" in lowered or "中危" in text:
        return "Medium"
    if "low" in lowered or "低危" in text:
        return "Low"
    return "Medium"


def localized_severity_for_filename(finding: dict[str, Any], language: str) -> str:
    declared_language = infer_source_language(finding)
    if language == "en-US":
        explicit = read_localized_string(finding, "severity", language, declared_language=declared_language)
        if explicit:
            return severity_en_from_any(explicit)
        if should_use_generic_field(language, declared_language):
            for key in ("severity", "severity_en", "severity_label"):
                value = str(finding.get(key, "")).strip()
                if value:
                    return severity_en_from_any(value)
        return severity_en_from_any(finding.get("severity_cn"))
    explicit = read_localized_string(finding, "severity_cn", language, declared_language=declared_language)
    if explicit:
        return severity_cn_from_any(explicit)
    for key in ("severity_cn", "severity"):
        value = str(finding.get(key, "")).strip()
        if value:
            return severity_cn_from_any(value)
    return "中危"


ANALYSIS_MARKER_TRANSLATIONS = {
    "zh-CN": {
        "Location:": "位置：",
        "Entry / controllable input:": "入口/可控输入：",
        "Dangerous operation:": "危险函数/危险操作：",
        "Trigger path:": "触发路径：",
        "Root cause:": "根因：",
        "Why existing checks fail:": "现有校验为何失效：",
    },
    "en-US": {
        "位置：": "Location:",
        "入口/可控输入：": "Entry / controllable input:",
        "危险函数/危险操作：": "Dangerous operation:",
        "触发路径：": "Trigger path:",
        "根因：": "Root cause:",
        "现有校验为何失效：": "Why existing checks fail:",
    },
}


def localize_analysis_marker(text: str, language: str) -> str:
    stripped = str(text).strip()
    for source_marker, target_marker in ANALYSIS_MARKER_TRANSLATIONS[language].items():
        if stripped.startswith(source_marker):
            return f"{target_marker} {stripped[len(source_marker):].strip()}"
    return stripped


def build_filename(finding: dict[str, Any], language: str) -> str:
    declared_language = infer_source_language(finding)
    explicit = read_localized_string(finding, "filename", language, declared_language=declared_language)
    if explicit:
        return explicit if explicit.endswith(".docx") else f"{explicit}.docx"
    project_name = str(finding.get("project_name", "project")).strip()
    vuln_type = localized_vuln_type(finding, language)
    severity_label = localized_severity_for_filename(finding, language)
    if language == "en-US":
        return f"{slugify(project_name)}_{slugify(vuln_type)}_{slugify(severity_label)}_report.docx"
    return f"{slugify(project_name)}_{slugify(vuln_type)}_{slugify(severity_label)}漏洞报告.docx"


def build_attachment_notes_filename(docx_name: str, language: str) -> str:
    stem = Path(docx_name).stem
    return f"{stem}{tr(language, 'attachment_notes_suffix')}"


def build_supplement_note_filename(docx_name: str, language: str) -> str:
    stem = Path(docx_name).stem
    return f"{stem}{tr(language, 'supplement_note_suffix')}"


def build_bundle_dirname(docx_name: str) -> str:
    return Path(docx_name).stem


def add_paragraphs(doc: Document, paragraphs: Iterable[str], style: str = "Normal") -> None:
    for text in paragraphs:
        doc.add_paragraph(text, style=style)


def add_labeled_block(doc: Document, label: str, lines: list[str], style: str = "Normal") -> None:
    if not lines:
        return
    doc.add_paragraph(label, style=style)
    for line in lines:
        doc.add_paragraph(line, style=style)


def resolve_project_root(output_dir: Path, finding: dict[str, Any]) -> Path:
    explicit = str(finding.get("project_root_dir", "")).strip() or str(finding.get("project_root", "")).strip()
    if explicit:
        return Path(explicit).expanduser().resolve()
    config_path = locate_workspace_config(output_dir)
    if config_path and config_path.parent == output_dir.parent:
        return config_path.parent.parent.resolve()
    if output_dir.name == "confirmed" and output_dir.parent.name.startswith("security-research"):
        return output_dir.parent.parent.resolve()
    if output_dir.name == "confirmed":
        return output_dir.parent.resolve()
    return output_dir.resolve()


def locate_workspace_config(output_dir: Path) -> Path | None:
    candidates = [
        output_dir.parent / "asr-config.json",
        output_dir / "asr-config.json",
    ]
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate.resolve()
    return None


def read_workspace_output_language(output_dir: Path) -> str | None:
    config_path = locate_workspace_config(output_dir)
    if not config_path:
        return None
    try:
        data = json.loads(config_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(data, dict):
        return None
    language = normalize_language(data.get("output_language"), "")
    return language or None


def resolve_bundle_attachment_source(
    item: dict[str, Any], project_root: Path, workspace_dir: Path | None
) -> tuple[str, Path, bool]:
    raw_path = item.get("path")
    rel, source_path = resolve_source_path(raw_path, project_root)
    if source_path.exists() and source_path.is_file():
        return rel, source_path, True

    if not rel.startswith("attachments/") or workspace_dir is None:
        return rel, source_path, False

    desired_rel = rel[len("attachments/"):].lstrip("/")
    desired_name = Path(desired_rel).name
    candidates: list[Path] = []

    direct_candidate = workspace_dir / desired_rel
    if direct_candidate.exists() and direct_candidate.is_file():
        candidates.append(direct_candidate.resolve())

    if desired_rel.endswith(".json"):
        candidates.extend(sorted((workspace_dir / "evidence").glob(f"**/{desired_name}")))
    elif desired_rel.endswith(".sh"):
        candidates.extend(sorted((workspace_dir / "poc").glob(f"**/{desired_rel}")))
        if not candidates:
            candidates.extend(sorted(workspace_dir.glob(f"**/{desired_rel}")))

    seen: set[str] = set()
    unique: list[Path] = []
    for candidate in candidates:
        resolved = candidate.resolve()
        key = str(resolved)
        if key in seen or not resolved.exists() or not resolved.is_file():
            continue
        seen.add(key)
        unique.append(resolved)

    if len(unique) == 1:
        return desired_rel, unique[0], True
    return desired_rel, source_path, False


def collect_bundle_metadata(finding: dict[str, Any], project_root: Path, bundle_dir: Path) -> dict[str, str]:
    mapping: dict[str, str] = {}
    workspace_dir = bundle_dir.parent.parent if bundle_dir.parent.name == "confirmed" else None
    seen_targets: set[str] = set()
    source_to_bundle: dict[str, str] = {}

    def choose_target(rel: str, source_path: Path) -> str:
        base = Path(rel).name or source_path.name or "attachment"
        stem = Path(base).stem or "attachment"
        suffix = Path(base).suffix
        candidate = f"attachments/{base}"
        counter = 2
        while candidate in seen_targets:
            candidate = f"attachments/{stem}-{counter}{suffix}"
            counter += 1
        return candidate

    attachment_items = finding.get("attachments") if isinstance(finding.get("attachments"), list) else []
    for item in attachment_items:
        if not isinstance(item, dict):
            continue
        rel, source_path, resolved = resolve_bundle_attachment_source(item, project_root, workspace_dir)
        if not rel:
            continue
        source_key = str(source_path.resolve()) if resolved and source_path.exists() and source_path.is_file() else ""
        if source_key and source_key in source_to_bundle:
            bundled_rel = source_to_bundle[source_key]
            mapping[rel] = bundled_rel
            mapping[f"attachments/{rel.lstrip('/')}"] = bundled_rel
            mapping[str(item.get("path") or "")] = bundled_rel
            continue
        bundled_rel = choose_target(rel, source_path) if source_key else f"attachments/{rel.lstrip('/')}"
        if resolved and source_path.exists() and source_path.is_file():
            target_path = bundle_dir / bundled_rel
            target_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source_path, target_path)
            seen_targets.add(bundled_rel)
            source_to_bundle[str(source_path.resolve())] = bundled_rel
        mapping[rel] = bundled_rel
        mapping[f"attachments/{rel.lstrip('/')}"] = bundled_rel
        mapping[str(item.get("path") or "")] = bundled_rel

    env_items = finding.get("environment_files") if isinstance(finding.get("environment_files"), list) else []
    for item in env_items:
        if not isinstance(item, dict):
            continue
        rel, source_path = resolve_source_path(item.get("path"), project_root)
        if not rel or not source_path.exists() or not source_path.is_file():
            continue
        source_key = str(source_path.resolve())
        existing = source_to_bundle.get(source_key)
        if existing:
            mapping[rel] = existing
            mapping[source_key] = existing
            mapping[str(item.get("path") or "")] = existing
            continue
        bundled_rel = choose_target(rel, source_path)
        target_path = bundle_dir / bundled_rel
        target_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source_path, target_path)
        mapping[rel] = bundled_rel
        mapping[f"attachments/{rel.lstrip('/')}"] = bundled_rel
        mapping[source_key] = bundled_rel
        mapping[str(item.get("path") or "")] = bundled_rel
        source_to_bundle[source_key] = bundled_rel
    return mapping


def collect_bundle_root_artifacts(
    finding: dict[str, Any], project_root: Path, bundle_dir: Path, language: str, path_map: dict[str, str]
) -> list[dict[str, str]]:
    collected: list[dict[str, str]] = []
    seen_targets: set[str] = set()
    value = finding.get("bundle_root_artifacts")
    if not isinstance(value, list):
        return collected

    for item in value:
        if not isinstance(item, dict):
            continue
        generator_name = str(item.get("generator") or "").strip()
        input_path = str(item.get("path") or "").strip()
        rel = ""
        source_path = project_root
        if input_path:
            rel, source_path = resolve_source_path(input_path, project_root)
        output_name = str(item.get("output_name") or (Path(rel).name if rel else "")).strip()
        if not output_name:
            continue
        output_name = Path(output_name).as_posix().lstrip("./")
        if output_name == ".." or output_name.startswith("../"):
            continue
        if output_name in seen_targets:
            continue
        seen_targets.add(output_name)
        target_path = bundle_dir / output_name
        target_path.parent.mkdir(parents=True, exist_ok=True)
        should_generate = (
            generator_name in {"reviewer-recording-shell", "recording-shell"}
            or (not input_path and output_name.endswith(".sh") and output_name.startswith("run-"))
        )
        if should_generate:
            script_text = build_generated_recording_shell(finding, language, path_map, item)
            target_path.write_text(script_text, encoding="utf-8")
            target_path.chmod(0o755)
            original_path = (
                "generated:reviewer-recording-shell"
                if language == "en-US"
                else "生成产物：reviewer-recording-shell"
            )
        else:
            if not rel or not source_path.exists() or not source_path.is_file():
                continue
            shutil.copy2(source_path, target_path)
            original_path = rel
        collected.append(
            {
                "bundle_path": output_name,
                "original_path": original_path,
                "purpose": str(item.get("purpose") or "").strip(),
                "purpose_en": str(item.get("purpose_en") or "").strip(),
            }
        )
    return collected


def shell_quote(text: str) -> str:
    return "'" + str(text).replace("'", "'\"'\"'") + "'"


def shell_block(text: str, marker: str) -> list[str]:
    lines = [f"cat <<'{marker}'"]
    lines.extend(str(text).splitlines() or [""])
    lines.append(marker)
    return lines


def localized_script_step(step: dict[str, Any], index: int, language: str) -> str:
    title = read_localized_string(step, "title", language)
    if title:
        return title
    return f"Step {index}" if language == "en-US" else f"步骤 {index}"


def normalize_reviewer_detail(detail: str, language: str) -> str:
    text = str(detail)
    if language == "zh-CN":
        text = text.replace(
            "使用本地缓存的 node:20-alpine 镜像（无需网络拉取）",
            "优先使用本地已有的 node:20-alpine 镜像；如本地不存在，再按脚本提示拉取",
        )
        if "PoC 脚本位于 " in text and "/poc/ 目录下" in text:
            text = "PoC 脚本已随当前 bundle 打包在 `attachments/` 目录下"
    else:
        text = text.replace(
            "Use the locally cached node:20-alpine image (no network pull required)",
            "Prefer a locally available node:20-alpine image; pull it only if it is not already present",
        )
        if "PoC scripts are located under " in text and "/poc/" in text:
            text = "The PoC scripts are bundled under `attachments/` in this bundle"
    return text


def localized_script_details(step: dict[str, Any], language: str, path_map: dict[str, str]) -> list[str]:
    raw = localized_list(step, "details", language) or localized_list(step, "detail", language)
    return [normalize_reviewer_detail(replace_paths_in_text(item, path_map), language) for item in raw]


def localized_script_evidence(step: dict[str, Any], language: str) -> list[str]:
    return (
        localized_list(step, "results", language)
        or localized_list(step, "result", language)
        or localized_list(step, "observed", language)
    )


def rewrite_command_for_bundle(command: str, path_map: dict[str, str]) -> str:
    result = str(command)
    for original in sorted(path_map, key=len, reverse=True):
        bundled = path_map[original]
        if not original:
            continue
        if original.startswith("/"):
            result = result.replace(original, bundled)
            continue
        pattern = re.compile(rf"(?<!attachments/){re.escape(original)}")
        result = pattern.sub(bundled, result)
    return result


def generator_modes_for_artifact(item: dict[str, Any]) -> list[str]:
    options = ensure_mapping(item.get("generator_options"))
    requested = ensure_list(options.get("modes"))
    allowed = {"record", "quick", "record-dos", "quick-dos"}
    modes: list[str] = []
    for mode in requested:
        normalized = str(mode).strip()
        if normalized in allowed and normalized not in modes:
            modes.append(normalized)
    return modes or ["record", "quick"]


def generated_mode_profile(mode: str, language: str) -> dict[str, str]:
    profiles = {
        "zh-CN": {
            "record": {
                "intro": "录屏模式即将开始，如尚未开始录屏请先手动开启。",
                "done": "录屏流程结束，现在可以停止录屏。",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "quick": {
                "intro": "快速模式会保留简短停顿与高亮提示，方便审核查看。",
                "done": "快速复现流程结束。",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "record-dos": {
                "intro": "DoS 录屏模式即将开始，如尚未开始录屏请先手动开启。",
                "done": "DoS 录屏流程结束，现在可以停止录屏。",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "quick-dos": {
                "intro": "快速 DoS 模式会保留简短停顿与高亮提示，方便审核查看。",
                "done": "快速 DoS 复现流程结束。",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
        },
        "en-US": {
            "record": {
                "intro": "Recording mode starts now. Start screen recording first if needed.",
                "done": "Recording flow completed. You can stop screen recording now.",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "quick": {
                "intro": "Quick mode still keeps short pauses and highlighting around key evidence.",
                "done": "Quick reproduction flow completed.",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "record-dos": {
                "intro": "DoS recording mode starts now. Start screen recording first if needed.",
                "done": "DoS recording flow completed. You can stop screen recording now.",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
            "quick-dos": {
                "intro": "Quick DoS mode still keeps short pauses and highlighting around the final proof lines.",
                "done": "Quick DoS reproduction flow completed.",
                "usage": "[record|quick|record-dos|quick-dos]",
            },
        },
    }
    return profiles[language][mode]


def build_generated_recording_shell(
    finding: dict[str, Any], language: str, path_map: dict[str, str], artifact: dict[str, Any]
) -> str:
    project_name = str(finding.get("project_name", "target-project")).strip() or "target-project"
    vuln_type = localized_vuln_type(finding, language)
    title = build_title(finding, language)
    raw_steps = finding.get("reproduction")
    steps = [step for step in raw_steps if isinstance(step, dict)] if isinstance(raw_steps, list) else []
    evidence_lines = collect_balanced_reproduction_evidence_lines(finding, language, limit=4)
    supported_modes = generator_modes_for_artifact(artifact)
    usage_modes = "|".join(supported_modes)
    code_items = finding.get("code_context")
    first_code = code_items[0] if isinstance(code_items, list) and code_items and isinstance(code_items[0], dict) else {}
    code_location = ensure_relpath(first_code.get("location")) or (
        "Key code path is not provided yet." if language == "en-US" else "关键代码位置待补充。"
    )
    code_summary = localized_string(first_code, "summary", language)
    code_snippet = str(first_code.get("snippet") or "").strip()

    strings = {
        "zh-CN": {
            "banner": f"{project_name} {vuln_type} 复现录屏辅助脚本",
            "mode": "模式：",
            "runtime": "运行环境：",
            "docker_only": "该审核脚本仅支持 Docker，不提供宿主机回退执行。",
            "docker_unavailable": "Docker daemon 未启动，请先启动 Docker/OrbStack。",
            "code_hint_step": "展示关键代码提示",
            "code_hint_title": "关键代码位置：",
            "code_hint_summary": "代码摘要：",
            "step": "步骤",
            "details": "步骤说明：",
            "run_command": "执行命令：",
            "evidence_hint": "关键证据提示：",
            "final_focus": "请在录屏中让最终关键证据停留几秒。",
            "code_unavailable": "未提供可展示的代码片段，将直接进入复现步骤。",
            "review_title": title,
            "focus_location": "关键代码位置",
            "focus_risk": "风险路径",
            "focus_oracle": "观察点",
            "focus_risk_default": "请让审核人员看到外部可控输入如何进入危险操作或脆弱解析路径。",
            "focus_oracle_default": "请重点展示最终成功判据、崩溃判据或攻击效果判据。",
            "snippet_focus": "请重点看下方关键片段：外部输入、危险点、关键判据。",
            "snippet_truncated": "代码片段较长，当前仅展示前 24 行关键内容。",
            "evidence_summary_title": "证据汇总",
            "evidence_summary_result": "[结论] 请结合以上证据给出审核结论。",
        },
        "en-US": {
            "banner": f"{project_name} {vuln_type} recording helper",
            "mode": "Mode:",
            "runtime": "Runtime:",
            "docker_only": "This reviewer-facing helper only supports Docker and must not fall back to host execution.",
            "docker_unavailable": "Docker daemon is not running. Please start Docker first.",
            "code_hint_step": "Show key code hint",
            "code_hint_title": "Key code location:",
            "code_hint_summary": "Code summary:",
            "step": "Step",
            "details": "Step details:",
            "run_command": "Run command:",
            "evidence_hint": "Key evidence hint:",
            "final_focus": "Keep the final proof lines on screen for a few seconds.",
            "code_unavailable": "No code snippet is available yet; continuing directly to the reproduction steps.",
            "review_title": title,
            "focus_location": "Key path",
            "focus_risk": "Risk flow",
            "focus_oracle": "Watch for",
            "focus_risk_default": "Show how attacker-controlled input reaches the dangerous sink or fragile parsing path.",
            "focus_oracle_default": "Keep the final success oracle, crash oracle, or attack-effect oracle clearly visible.",
            "snippet_focus": "Focus on the snippet below: external input, dangerous operation, and final oracle.",
            "snippet_truncated": "The snippet is longer; only the first 24 key lines are shown here.",
            "evidence_summary_title": "Evidence Summary",
            "evidence_summary_result": "[Conclusion] Keep this evidence summary on screen before stopping the recording.",
        },
    }[language]

    snippet_lines = code_snippet.splitlines() if code_snippet else []
    snippet_preview = snippet_lines[:24]
    snippet_truncated = len(snippet_lines) > 24
    focus_risk_text = code_summary or strings["focus_risk_default"]
    focus_oracle_text = (
        evidence_lines[0]
        if evidence_lines
        else strings["focus_oracle_default"]
    )

    script_lines: list[str] = [
        "#!/bin/sh",
        "",
        "set -eu",
        "",
        'MODE="${1:-record}"',
        'RUNTIME="${2:-docker}"',
        'PAUSE_SHORT="${REVIEWER_PAUSE_SHORT:-2}"',
        'PAUSE_LONG="${REVIEWER_PAUSE_LONG:-4}"',
        'SCRIPT_DIR="$(CDPATH= cd -- "$(dirname -- "$0")" && pwd)"',
        'cd "$SCRIPT_DIR"',
        "",
        "if [ -t 1 ]; then",
        "    C_RESET=\"$(printf '\\033[0m')\"",
        "    C_BOLD=\"$(printf '\\033[1m')\"",
        "    C_BLUE=\"$(printf '\\033[34m')\"",
        "    C_MAGENTA=\"$(printf '\\033[35m')\"",
        "    C_RED=\"$(printf '\\033[31m')\"",
        "    C_GREEN=\"$(printf '\\033[32m')\"",
        "    C_YELLOW=\"$(printf '\\033[33m')\"",
        "    C_CYAN=\"$(printf '\\033[36m')\"",
        "    C_WHITE_ON_RED=\"$(printf '\\033[1;37;41m')\"",
        "    C_WHITE_ON_BLUE=\"$(printf '\\033[1;37;44m')\"",
        "    C_WHITE_ON_MAGENTA=\"$(printf '\\033[1;37;45m')\"",
        "    C_BLACK_ON_YELLOW=\"$(printf '\\033[1;30;43m')\"",
        "else",
        "    C_RESET=\"\"",
        "    C_BOLD=\"\"",
        "    C_BLUE=\"\"",
        "    C_MAGENTA=\"\"",
        "    C_RED=\"\"",
        "    C_GREEN=\"\"",
        "    C_YELLOW=\"\"",
        "    C_CYAN=\"\"",
        "    C_WHITE_ON_RED=\"\"",
        "    C_WHITE_ON_BLUE=\"\"",
        "    C_WHITE_ON_MAGENTA=\"\"",
        "    C_BLACK_ON_YELLOW=\"\"",
        "fi",
        "",
        "pause_step() {",
        "    sleep \"$1\"",
        "}",
        "",
        "print_banner() {",
        "    printf '\\n%s%s============================================================%s\\n' \"$C_CYAN\" \"$C_BOLD\" \"$C_RESET\"",
        "    printf '%s%s%s\\n' \"$C_CYAN\" \"$1\" \"$C_RESET\"",
        "    printf '%s%s============================================================%s\\n' \"$C_CYAN\" \"$C_BOLD\" \"$C_RESET\"",
        "}",
        "",
        "announce_step() {",
        "    printf '\\n%s%s[%s]%s %s\\n' \"$C_YELLOW\" \"$C_BOLD\" \"$1\" \"$C_RESET\" \"$2\"",
        "}",
        "",
        "focus_line() {",
        "    printf '%s%s %s %s\\n' \"$1\" \"$2\" \"$3\" \"$C_RESET\"",
        "}",
        "",
        "print_separator() {",
        "    printf '%s%s------------------------------------------------------------%s\\n' \"$C_MAGENTA\" \"$C_BOLD\" \"$C_RESET\"",
        "}",
        "",
        "highlight_success() {",
        "    printf '%s%s%s\\n' \"$C_GREEN\" \"$1\" \"$C_RESET\"",
        "}",
        "",
        "highlight_danger() {",
        "    printf '%s%s%s\\n' \"$C_RED\" \"$1\" \"$C_RESET\"",
        "}",
        "",
        "highlight_note() {",
        "    printf '%s%s%s\\n' \"$C_CYAN\" \"$1\" \"$C_RESET\"",
        "}",
        "",
        "docker_ready() {",
        "    docker info >/dev/null 2>&1",
        "}",
        "",
        "show_code_hint() {",
        f"    announce_step {shell_quote('0/' + str(len(steps) + 1))} {shell_quote(strings['code_hint_step'])}",
    ]
    if code_snippet:
        script_lines.append(
            f"    focus_line \"$C_WHITE_ON_BLUE\" {shell_quote(' ' + strings['focus_location'] + ' ')} {shell_quote(' ' + code_location + ' ')}"
        )
        script_lines.append(
            f"    focus_line \"$C_BLACK_ON_YELLOW\" {shell_quote(' ' + strings['focus_risk'] + ' ')} {shell_quote(' ' + focus_risk_text + ' ')}"
        )
        script_lines.append(
            f"    focus_line \"$C_WHITE_ON_MAGENTA\" {shell_quote(' ' + strings['focus_oracle'] + ' ')} {shell_quote(' ' + focus_oracle_text + ' ')}"
        )
        script_lines.append("    pause_step \"$PAUSE_SHORT\"")
        script_lines.append(f"    printf '%s%s %s%s\\n' \"$C_GREEN\" {shell_quote(strings['code_hint_title'])} {shell_quote(code_location)} \"$C_RESET\"")
        if code_summary:
            script_lines.append(f"    printf '%s%s %s%s\\n' \"$C_GREEN\" {shell_quote(strings['code_hint_summary'])} {shell_quote(code_summary)} \"$C_RESET\"")
        script_lines.append(f"    printf '%s%s%s%s\\n' \"$C_BLACK_ON_YELLOW\" \"$C_BOLD\" {shell_quote(strings['snippet_focus'])} \"$C_RESET\"")
        script_lines.append("    print_separator")
        script_lines.append("    cat <<'CODE_SNIPPET_EOF' | nl -ba | sed -n '1,24p'")
        script_lines.extend(snippet_preview or [""])
        script_lines.append("CODE_SNIPPET_EOF")
        script_lines.append("    print_separator")
        if snippet_truncated:
            script_lines.append(f"    highlight_note {shell_quote(strings['snippet_truncated'])}")
    else:
        script_lines.append(f"    printf '%s\\n' {shell_quote(strings['code_unavailable'])}")
    script_lines.extend([
        "    pause_step \"$PAUSE_LONG\"",
        "}",
        "",
        "show_evidence_summary() {",
        f"    printf '\\n%s%s==================== {strings['evidence_summary_title']} ====================%s\\n' \"$C_GREEN\" \"$C_BOLD\" \"$C_RESET\"",
    ])
    palette_vars = ["$C_WHITE_ON_BLUE", "$C_WHITE_ON_MAGENTA", "$C_WHITE_ON_RED", "$C_BLACK_ON_YELLOW"]
    for idx, line in enumerate(evidence_lines, start=1):
        color_var = palette_vars[(idx - 1) % len(palette_vars)]
        label = f"{strings['evidence_hint']} {idx}"
        script_lines.append(f"    focus_line \"{color_var}\" {shell_quote(' ' + label + ' ')} {shell_quote(' ' + line + ' ')}")
        script_lines.append("    pause_step 1")
    script_lines.extend([
        f"    printf '%s%s%s%s\\n' \"$C_GREEN\" \"$C_BOLD\" {shell_quote(strings['evidence_summary_result'])} \"$C_RESET\"",
        "    printf '%s%s====================================================%s\\n' \"$C_GREEN\" \"$C_BOLD\" \"$C_RESET\"",
        "}",
        "",
        "run_flow() {",
        "    show_code_hint",
    ])

    total_steps = len(steps)
    for idx, step in enumerate(steps, start=1):
        step_title = localized_script_step(step, idx, language)
        step_details = localized_script_details(step, language, path_map)
        commands = [rewrite_command_for_bundle(command, path_map) for command in ensure_list(step.get("commands") or step.get("command"))]
        evidence = localized_script_evidence(step, language)
        script_lines.append(f"    announce_step {shell_quote(f'{idx}/{max(total_steps, 1)}')} {shell_quote(step_title)}")
        for detail in step_details:
            script_lines.append(f"    printf '%s %s\\n' {shell_quote('-')} {shell_quote(detail)}")
        if step_details:
            script_lines.append("    pause_step \"$PAUSE_SHORT\"")
        for command in commands:
            script_lines.append(f"    printf '%s%s %s%s\\n' \"$C_CYAN\" {shell_quote(strings['run_command'])} {shell_quote(command)} \"$C_RESET\"")
            script_lines.append("    pause_step \"$PAUSE_SHORT\"")
            script_lines.append(f"    {command}")
            script_lines.append("    pause_step \"$PAUSE_SHORT\"")
        for line in evidence:
            script_lines.append(f"    highlight_success {shell_quote(strings['evidence_hint'] + ' ' + line)}")
        if evidence:
            script_lines.append("    pause_step \"$PAUSE_LONG\"")

    if evidence_lines:
        script_lines.append("    announce_step 'evidence' " + shell_quote(strings["final_focus"]))
        for line in evidence_lines:
            script_lines.append(f"    highlight_danger {shell_quote(line)}")
        script_lines.append("    pause_step \"$PAUSE_SHORT\"")
        script_lines.append("    show_evidence_summary")
        script_lines.append("    pause_step \"$PAUSE_LONG\"")
    script_lines.extend([
        "}",
        "",
        "main() {",
        "    if [ \"$RUNTIME\" != \"docker\" ]; then",
        f"        echo \"Usage: $0 [{usage_modes}] [docker]\" >&2",
        f"        echo {shell_quote(strings['docker_only'])} >&2",
        "        exit 1",
        "    fi",
        "",
        "    if ! docker_ready; then",
        f"        echo {shell_quote(strings['docker_unavailable'])} >&2",
        "        exit 1",
        "    fi",
        "",
        "    case \"$MODE\" in",
    ])
    for mode in supported_modes:
        profile = generated_mode_profile(mode, language)
        script_lines.append(f"        {mode})")
        if mode.startswith("quick"):
            script_lines.append("            PAUSE_SHORT=1")
            script_lines.append("            PAUSE_LONG=2")
        script_lines.append(f"            print_banner {shell_quote(strings['banner'])}")
        script_lines.append(f"            printf '%s %s\\n' {shell_quote(strings['mode'])} \"$MODE\"")
        script_lines.append(f"            printf '%s docker\\n' {shell_quote(strings['runtime'])}")
        if mode.startswith("record"):
            script_lines.append(f"            printf '%s\\n' {shell_quote(strings['review_title'])}")
        script_lines.append(f"            printf '%s\\n' {shell_quote(profile['intro'])}")
        if mode.startswith("record"):
            script_lines.append("            pause_step \"$PAUSE_LONG\"")
        else:
            script_lines.append("            pause_step \"$PAUSE_SHORT\"")
        script_lines.append("            run_flow")
        script_lines.append(f"            printf '%s\\n' {shell_quote(profile['done'])}")
        script_lines.append("            ;;")
    script_lines.extend([
        "        *)",
        f"            echo \"Usage: $0 [{usage_modes}] [docker]\" >&2",
        "            exit 1",
        "            ;;",
        "    esac",
        "}",
        "",
        "main \"$@\"",
        "",
    ])
    return "\n".join(script_lines)


def replace_paths_in_text(text: str, path_map: dict[str, str]) -> str:
    result = str(text)
    for original in sorted(path_map, key=len, reverse=True):
        bundled = path_map[original]
        if not original:
            continue
        if original.startswith("/"):
            result = result.replace(original, bundled)
            continue
        pattern = re.compile(rf"(?<!attachments/){re.escape(original)}")
        result = pattern.sub(bundled, result)
    result = result.replace("attachments/attachments/", "attachments/")
    return result


def bundled_path_for_value(value: Any, path_map: dict[str, str]) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    normalized = Path(text).as_posix().lstrip("./")
    if normalized.startswith("attachments/"):
        return path_map.get(normalized, path_map.get(text, normalized))
    return path_map.get(normalized, path_map.get(text, normalized))


def first_nonempty(*values: Any) -> str:
    for value in values:
        text = str(value or "").strip()
        if text:
            return text
    return ""


def first_docker_command(finding: dict[str, Any], fallback: str = "") -> str:
    raw_steps = finding.get("reproduction")
    if isinstance(raw_steps, list):
        for step in raw_steps:
            if not isinstance(step, dict):
                continue
            for command in ensure_list(step.get("commands") or step.get("command")):
                if "docker" in command.lower():
                    return command
    return fallback


def first_localized_step_value(finding: dict[str, Any], keys: tuple[str, ...], language: str) -> str:
    raw_steps = finding.get("reproduction")
    if not isinstance(raw_steps, list):
        return ""
    for step in reversed(raw_steps):
        if not isinstance(step, dict):
            continue
        for key in keys:
            values = localized_list(step, key, language)
            if values:
                return values[0]
    return ""


def collect_bundled_attachment_paths(path_map: dict[str, str]) -> list[str]:
    paths = sorted({value for value in path_map.values() if value.startswith("attachments/")})
    return paths


def infer_poc_path(finding: dict[str, Any], path_map: dict[str, str]) -> str:
    explicit = ensure_mapping(finding.get("verification_evidence")).get("poc_path")
    rel = bundled_path_for_value(explicit, path_map)
    if rel:
        return rel
    for value in collect_bundled_attachment_paths(path_map):
        lowered = value.lower()
        if "/poc/" in lowered or lowered.endswith(".py") or lowered.endswith(".js") or lowered.endswith(".sh"):
            return value
    attachments = collect_bundled_attachment_paths(path_map)
    return attachments[0] if attachments else ""


def write_verification_evidence(
    output_path: Path,
    finding: dict[str, Any],
    bundle_finding: dict[str, Any],
    path_map: dict[str, str],
    language: str,
) -> None:
    provided = ensure_mapping(finding.get("verification_evidence"))
    slug = first_nonempty(provided.get("finding_slug"), finding.get("slug"), bundle_finding.get("slug"), output_path.parent.name)
    status = first_nonempty(provided.get("verification_status"), finding.get("verification_status"), "confirmed_in_docker")
    if status != "confirmed_in_docker":
        raise SystemExit(
            "Refusing to render a confirmed bundle unless verification_status is confirmed_in_docker: "
            f"{status}"
        )

    evidence_files = []
    raw_evidence_files = provided.get("evidence_files")
    if isinstance(raw_evidence_files, list):
        for item in raw_evidence_files:
            rel = bundled_path_for_value(item, path_map)
            if rel and rel not in evidence_files:
                evidence_files.append(rel)
    for rel in collect_bundled_attachment_paths(path_map):
        if rel not in evidence_files:
            evidence_files.append(rel)

    poc_path = bundled_path_for_value(provided.get("poc_path"), path_map) or infer_poc_path(finding, path_map)
    if not poc_path:
        raise SystemExit(
            "verification_evidence.poc_path is required for confirmed bundles. "
            "Set findings[].verification_evidence.poc_path to a PoC file that is "
            "bundled under attachments/, or include a PoC attachment that the renderer can copy."
        )
    if poc_path and poc_path not in evidence_files:
        evidence_files.insert(0, poc_path)

    docker_command = first_nonempty(
        provided.get("docker_command"),
        finding.get("docker_command"),
        first_docker_command(bundle_finding),
    )
    docker_image = first_nonempty(
        provided.get("docker_image"),
        finding.get("docker_image"),
    )
    if not docker_image:
        raise SystemExit(
            "verification_evidence.docker_image is required for confirmed bundles. "
            "Use the concrete Docker image tag or Docker Compose service used during verification."
        )
    if docker_image == "project-specific Docker image or Docker Compose service":
        raise SystemExit(
            "verification_evidence.docker_image must name the concrete Docker image or Compose service, "
            "not the placeholder text."
        )
    expected = first_nonempty(
        provided.get("expected_observation"),
        finding.get("expected_observation"),
        first_localized_step_value(bundle_finding, ("expected",), language),
    )
    observed = first_nonempty(
        provided.get("observed_observation"),
        finding.get("observed_observation"),
        first_localized_step_value(bundle_finding, ("observed", "results", "result"), language),
    )
    oracle_token = first_nonempty(
        provided.get("oracle_token"),
        finding.get("oracle_token"),
        observed,
    )
    severity_result = first_nonempty(
        provided.get("severity_escalation_result"),
        finding.get("severity_escalation_result"),
        localized_list(ensure_mapping(bundle_finding.get("cvss")), "rationale", language)[0]
        if localized_list(ensure_mapping(bundle_finding.get("cvss")), "rationale", language)
        else "",
        "Severity escalation was attempted; the final severity reflects the strongest Docker-verified oracle.",
    )

    evidence = {
        "schema_version": 1,
        "finding_slug": slug,
        "verification_status": status,
        "docker_required": True,
        "docker_image": docker_image,
        "docker_command": docker_command,
        "poc_path": poc_path,
        "expected_observation": expected,
        "observed_observation": observed,
        "oracle_token": oracle_token,
        "evidence_files": evidence_files,
        "severity_escalation_attempted": bool(
            provided.get("severity_escalation_attempted", finding.get("severity_escalation_attempted", True))
        ),
        "severity_escalation_result": severity_result,
    }
    (output_path.parent / "verification-evidence.json").write_text(
        json.dumps(evidence, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def rewrite_value(value: Any, path_map: dict[str, str], project_root: Path) -> Any:
    if isinstance(value, str):
        return replace_paths_in_text(value, path_map)
    if isinstance(value, list):
        return [rewrite_value(item, path_map, project_root) for item in value]
    if isinstance(value, dict):
        rewritten = {}
        for key, item in value.items():
            if key == "location":
                rewritten[key] = normalize_relative_to_project(item, project_root)
            elif key == "path":
                rel = normalize_relative_to_project(item, project_root)
                rewritten[key] = path_map.get(rel, rel)
            else:
                rewritten[key] = rewrite_value(item, path_map, project_root)
        return rewritten
    return value


def prepare_finding_for_bundle(
    finding: dict[str, Any], path_map: dict[str, str], project_root: Path
) -> dict[str, Any]:
    rewritten = rewrite_value(finding, path_map, project_root)
    if isinstance(rewritten, dict):
        return rewritten
    return dict(finding)


def render_impact_paragraph(impact: dict[str, Any], language: str) -> str:
    lines = []
    affected_versions = localized_string(impact, "affected_versions", language)
    mapping = [
        (tr(language, "impact_package"), impact.get("package") or impact.get("project")),
        (tr(language, "impact_component"), ensure_relpath(impact.get("component"))),
        (tr(language, "impact_versions_label"), affected_versions or impact.get("version")),
        (tr(language, "repo_url"), impact.get("repo_url") or impact.get("repository")),
    ]
    for label, value in mapping:
        if value:
            lines.append(format_kv(language, label, value))
    lines.extend(localized_list(impact, "extra", language))
    return "\n".join(lines) if lines else ("Impact details are not provided yet." if language == "en-US" else "影响信息待补充。")


def render_cvss_paragraph(cvss: dict[str, Any], language: str) -> str:
    lines = []
    vector = str(cvss.get("vector") or "").strip()
    if vector.startswith("CVSS:3.1/"):
        vector_label = tr(language, "cvss31_vector")
    else:
        vector_label = tr(language, "cvss4_vector")
    mapping = [
        (vector_label, vector),
        (tr(language, "base_score"), cvss.get("score")),
        (
            tr(language, "severity"),
            severity_en_from_any(cvss.get("severity")) if language == "en-US" else severity_cn_from_any(cvss.get("severity")),
        ),
    ]
    for label, value in mapping:
        if value:
            lines.append(format_kv(language, label, value))
    return "\n".join(lines) if lines else ("CVSS details are not provided yet." if language == "en-US" else "CVSS 信息待补充。")


def validate_generated_docx(output_path: Path) -> None:
    if not output_path.exists():
        raise RuntimeError(f"DOCX was not created: {output_path}")
    if not zipfile.is_zipfile(output_path):
        raise RuntimeError(
            f"Generated file is not a valid DOCX/ZIP container: {output_path}. "
            "Do not save plain text or model replies with a .docx extension."
        )
    with zipfile.ZipFile(output_path) as zf:
        names = set(zf.namelist())
    required = {"[Content_Types].xml", "word/document.xml", "_rels/.rels"}
    missing = sorted(required - names)
    if missing:
        raise RuntimeError(
            f"Generated DOCX is missing required OOXML members {missing}: {output_path}"
        )


def render_code_context(doc: Document, finding: dict[str, Any], language: str) -> None:
    items = finding.get("code_context")
    if not isinstance(items, list) or not items:
        return
    doc.add_heading(tr(language, "code_context"), level=2)
    for idx, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        location = ensure_relpath(item.get("location"))
        summary = localized_string(item, "summary", language)
        snippet = str(item.get("snippet", "")).strip()
        explanation = localized_string(item, "explanation", language)
        heading = f"{idx}. {location}" if location else tr(language, "code_context_item", idx=idx)
        doc.add_paragraph(heading, style="List Bullet")
        if summary:
            doc.add_paragraph(summary)
        if snippet:
            for line in snippet.splitlines():
                doc.add_paragraph(line, style="Intense Quote")
        if explanation:
            doc.add_paragraph(explanation)


def render_environment_files(doc: Document, finding: dict[str, Any], language: str) -> None:
    items = finding.get("environment_files")
    if not isinstance(items, list) or not items:
        return
    doc.add_heading(tr(language, "env_files"), level=1)
    for idx, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        path = ensure_relpath(item.get("path"))
        purpose = localized_string(item, "purpose", language)
        snippet = str(item.get("snippet", "")).strip()
        note = localized_string(item, "note", language)
        doc.add_paragraph(f"{idx}. {path}", style="List Bullet")
        if purpose:
            doc.add_paragraph(tr(language, "purpose_inline", purpose=purpose))
        if note:
            doc.add_paragraph(tr(language, "note_inline", note=note))
        if snippet:
            for line in snippet.splitlines():
                doc.add_paragraph(line, style="Intense Quote")


def render_reproduction(doc: Document, reproduction_steps: list[dict[str, Any]], language: str) -> None:
    doc.add_heading(tr(language, "reproduction"), level=1)
    if not reproduction_steps:
        doc.add_paragraph(tr(language, "repro_pending"))
        return
    for index, step in enumerate(reproduction_steps, start=1):
        title = read_localized_string(step, "title", language)
        if not title:
            title = f"{index}. Reproduction Step" if language == "en-US" else f"{index}. 复现步骤"
        doc.add_heading(title, level=2)
        details = localized_list(step, "details", language) or localized_list(step, "detail", language)
        add_paragraphs(doc, details)
        add_labeled_block(doc, tr(language, "related_artifacts"), [ensure_relpath(x) for x in ensure_list(step.get("artifacts") or step.get("artifact"))])
        add_labeled_block(doc, tr(language, "commands"), ensure_list(step.get("commands") or step.get("command")), style="Intense Quote")
        add_labeled_block(doc, tr(language, "expected"), localized_list(step, "expected", language))
        add_labeled_block(doc, tr(language, "observed"), localized_list(step, "observed", language))
        add_labeled_block(doc, tr(language, "notes"), localized_list(step, "notes", language) or localized_list(step, "note", language))
        add_labeled_block(doc, tr(language, "results"), localized_list(step, "results", language) or localized_list(step, "result", language))


def collect_reproduction_evidence_lines(finding: dict[str, Any], language: str) -> list[str]:
    lines: list[str] = []
    for step_lines in collect_reproduction_evidence_groups(finding, language):
        lines.extend(step_lines)
    return lines


def collect_reproduction_evidence_groups(finding: dict[str, Any], language: str) -> list[list[str]]:
    grouped_lines: list[list[str]] = []
    raw_steps = finding.get("reproduction")
    if not isinstance(raw_steps, list):
        return grouped_lines
    for index, step in enumerate(raw_steps, start=1):
        if not isinstance(step, dict):
            continue
        title = read_localized_string(step, "title", language) or (
            f"{index}. Reproduction Step" if language == "en-US" else f"{index}. 复现步骤"
        )
        step_lines: list[str] = []
        seen_step_lines: set[str] = set()
        result_groups = localized_list(step, "results", language) or localized_list(step, "result", language)
        observed_group = localized_list(step, "observed", language)
        for group_name, group in (("result", result_groups), ("observed", observed_group)):
            for item in group:
                text = str(item).strip()
                if not text or text in seen_step_lines:
                    continue
                seen_step_lines.add(text)
                if group_name == "observed":
                    prefix = "Observed result: " if language == "en-US" else "实际结果："
                    if not text.startswith(prefix):
                        text = f"{prefix}{text}"
                step_lines.append(text)
        formatted_lines: list[str] = []
        for item in step_lines:
            text = str(item).strip()
            if text:
                formatted_lines.append(f"{title}: {text}")
        if formatted_lines:
            grouped_lines.append(formatted_lines)
    return grouped_lines


def collect_balanced_reproduction_evidence_lines(
    finding: dict[str, Any], language: str, limit: int = 4
) -> list[str]:
    if limit <= 0:
        return []
    groups = [prioritize_reviewer_evidence_group(group, language) for group in collect_reproduction_evidence_groups(finding, language) if group]
    if not groups:
        return []

    selected: list[str] = []
    for group in groups:
        if len(selected) >= limit:
            break
        selected.append(group.pop(0))

    while len(selected) < limit:
        made_progress = False
        for group in groups:
            if len(selected) >= limit:
                break
            if not group:
                continue
            selected.append(group.pop(0))
            made_progress = True
        if not made_progress:
            break

    return selected


def prioritize_reviewer_evidence_group(group: list[str], language: str) -> list[str]:
    generic_markers = (
        ["实际结果：已完成", "结果证据：已完成"]
        if language == "zh-CN"
        else ["Observed result: completed", "Observed result: done", "Evidence: completed"]
    )

    def score(text: str) -> tuple[int, int]:
        normalized = str(text).strip()
        is_generic = any(marker in normalized for marker in generic_markers)
        return (1 if is_generic else 0, len(normalized))

    return sorted(group, key=score)


def write_reproduction_supplement(
    output_path: Path,
    finding: dict[str, Any],
    language: str,
    path_map: dict[str, str],
    bundle_root_artifacts: list[dict[str, str]] | None = None,
) -> None:
    bundle_root_artifacts = bundle_root_artifacts or []
    project_name = str(finding.get("project_name", "target-project")).strip() or "target-project"
    vuln_type = localized_vuln_type(finding, language)
    lines = [
        tr(language, "supplement_title", project_name=project_name, vuln_type=vuln_type),
        "",
        tr(language, "supplement_goal"),
        "",
    ]
    description = localized_list(finding, "description", language)
    if description:
        lines.extend(description[:2])
    else:
        lines.append(
            "This supplement provides the shortest reproducible path and reviewer-facing success evidence for the confirmed vulnerability."
            if language == "en-US"
            else "本补充说明用于提供已确认漏洞的最短复现路径与面向审核人员的成功证据。"
        )
    lines.extend(["", tr(language, "supplement_env"), "", tr(language, "supplement_docker_only"), ""])
    lines.extend([tr(language, "supplement_shortest_path"), ""])
    scripts = [str(item.get("bundle_path") or "").strip() for item in bundle_root_artifacts if str(item.get("bundle_path") or "").strip()]
    if scripts:
        lines.append(tr(language, "supplement_shortest_with_script", script=scripts[0]))
    else:
        lines.append(tr(language, "supplement_shortest_without_script"))

    raw_steps = finding.get("reproduction")
    if isinstance(raw_steps, list):
        for step in raw_steps:
            if not isinstance(step, dict):
                continue
            title = read_localized_string(step, "title", language)
            details = [
                normalize_reviewer_detail(replace_paths_in_text(item, path_map), language)
                for item in (localized_list(step, "details", language) or localized_list(step, "detail", language))
            ]
            commands = [rewrite_command_for_bundle(command, path_map) for command in ensure_list(step.get("commands") or step.get("command"))]
            if title:
                lines.extend(["", title])
            for item in details:
                lines.append(item)
            for command in commands:
                lines.append(f"`{command}`")

    lines.extend(["", tr(language, "supplement_evidence"), ""])
    evidence_lines = collect_reproduction_evidence_lines(finding, language)
    if evidence_lines:
        lines.extend(f"- {line}" for line in evidence_lines)
    else:
        lines.append(tr(language, "supplement_evidence_missing"))

    lines.extend(["", tr(language, "supplement_bundle"), ""])
    if scripts:
        lines.append(tr(language, "supplement_bundle_scripts"))
        lines.extend(f"- `{script}`" for script in scripts)
    lines.append(tr(language, "supplement_bundle_attachments"))

    lines.extend(["", tr(language, "supplement_conclusion"), ""])
    final_verdict = localized_list(finding, "final_verdict", language)
    if final_verdict:
        lines.extend(final_verdict)
    else:
        lines.append(tr(language, "supplement_conclusion_default"))

    supplement_path = output_path.with_name(build_supplement_note_filename(output_path.name, language))
    supplement_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_attachment_notes(
    output_path: Path,
    finding: dict[str, Any],
    path_map: dict[str, str],
    project_root: Path,
    language: str,
    bundle_root_artifacts: list[dict[str, str]] | None = None,
) -> None:
    lines = [tr(language, "attachment_notes_title"), "", tr(language, "report_for", name=output_path.name), "", tr(language, "directory_structure"), ""]
    lines.append(tr(language, "report_body", name=output_path.name))
    if bundle_root_artifacts:
        for item in bundle_root_artifacts:
            bundled_path = str(item.get("bundle_path") or "").strip()
            if bundled_path:
                lines.append(
                    f"- `{bundled_path}`：{'最小复现脚本或辅助验证产物。' if language == 'zh-CN' else 'minimal reproduction helper or supporting verification artifact.'}"
                )
    if path_map:
        lines.append(tr(language, "attachments_dir"))
    lines.append("")
    lines.append(tr(language, "attachment_usage"))
    lines.append("")

    entries = []
    for key in ("attachments", "environment_files"):
        value = finding.get(key)
        if isinstance(value, list):
            entries.extend(value)

    seen: set[str] = set()
    index = 1
    for item in bundle_root_artifacts or []:
        bundled = str(item.get("bundle_path") or "").strip()
        original = str(item.get("original_path") or "").strip()
        if not bundled or not original or bundled in seen:
            continue
        seen.add(bundled)
        purpose = str(item.get("purpose_en") if language == "en-US" else item.get("purpose") or "").strip() or tr(language, "attachment_missing")
        lines.append(f"### {index}. `{bundled}`")
        lines.append(tr(language, "original_path", path=original))
        lines.append(tr(language, "purpose", purpose=purpose))
        lines.append("")
        index += 1

    for item in entries:
        if not isinstance(item, dict):
            continue
        original = normalize_relative_to_project(item.get("path"), project_root)
        bundled = path_map.get(original, original)
        if not original or bundled in seen:
            continue
        seen.add(bundled)
        purpose = localized_string(item, "purpose", language) or tr(language, "attachment_missing")
        lines.append(f"### {index}. `{bundled}`")
        lines.append(tr(language, "original_path", path=original))
        lines.append(tr(language, "purpose", purpose=purpose))
        lines.append("")
        index += 1

    notes_path = output_path.with_name(build_attachment_notes_filename(output_path.name, language))
    notes_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def render_analysis(doc: Document, finding: dict[str, Any], language: str) -> None:
    doc.add_heading(tr(language, "analysis"), level=1)
    analysis_items = localized_list(finding, "analysis", language)
    if analysis_items:
        for item in analysis_items:
            doc.add_paragraph(localize_analysis_marker(item, language), style="List Bullet")
    else:
        doc.add_paragraph(tr(language, "analysis_pending"), style="List Bullet")
    render_code_context(doc, finding, language)


def cvss4_vector_from_legacy(finding: dict[str, Any]) -> str:
    av_map = {"network": "N", "adjacent": "A", "local": "L", "physical": "P"}
    simple_map = {"none": "N", "low": "L", "high": "H", "required": "P"}
    attack_vector = av_map.get(str(finding.get("attack_vector", "")).strip().lower(), "L")
    attack_complexity = simple_map.get(str(finding.get("attack_complexity", "")).strip().lower(), "L")
    privileges_required = simple_map.get(str(finding.get("privileges_required", "")).strip().lower(), "N")
    user_interaction = simple_map.get(str(finding.get("user_interaction", "")).strip().lower(), "N")
    confidentiality = simple_map.get(str(finding.get("confidentiality_impact", "")).strip().lower(), "L")
    integrity = simple_map.get(str(finding.get("integrity_impact", "")).strip().lower(), "N")
    availability = simple_map.get(str(finding.get("availability_impact", "")).strip().lower(), "N")
    return (
        "CVSS:4.0/"
        f"AV:{attack_vector}/AC:{attack_complexity}/AT:N/PR:{privileges_required}/UI:{user_interaction}/"
        f"VC:{confidentiality}/VI:{integrity}/VA:{availability}/SC:N/SI:N/SA:N"
    )


def severity_cn_from_legacy(value: Any) -> str:
    return severity_cn_from_any(value)


def severity_label_from_score(score_text: str, language: str = "zh-CN") -> str:
    try:
        score = float(score_text)
    except ValueError:
        return "Medium" if language == "en-US" else "中危"
    if score >= 9:
        return "Critical" if language == "en-US" else "严重"
    if score >= 7:
        return "High" if language == "en-US" else "高危"
    if score >= 4:
        return "Medium" if language == "en-US" else "中危"
    return "Low" if language == "en-US" else "低危"


def transform_legacy_finding(finding: dict[str, Any], defaults: dict[str, Any] | None = None) -> dict[str, Any]:
    defaults = defaults or {}
    language = normalize_language(
        finding.get("report_language")
        or finding.get("output_language")
        or defaults.get("report_language")
        or defaults.get("output_language"),
        "zh-CN",
    )
    project_name = (
        str(
            finding.get("project_name")
            or finding.get("project")
            or finding.get("package_name")
            or defaults.get("project_name")
            or defaults.get("project")
            or defaults.get("package_name")
            or "target-project"
        ).strip().split("/")[-1]
        or "target-project"
    )
    raw_title = localized_string(finding, "title", language) or str(finding.get("title") or "").strip()
    vuln_type = (
        str(
            localized_string(finding, "vuln_type", language)
            or finding.get("vulnerability_type")
            or finding.get("type")
            or raw_title
            or "安全漏洞"
        ).strip()
        or "安全漏洞"
    )
    severity_cn = severity_cn_from_legacy(finding.get("severity") or defaults.get("severity"))
    location = str(finding.get("affected_component") or finding.get("vulnerable_file") or finding.get("file_path") or "").strip()
    vulnerable_code = str(finding.get("vulnerable_code") or "").strip()
    description = localized_list(finding, "description", language) or ensure_list(finding.get("description")) or [
        raw_title or (
            "This issue was verified in Docker, but the legacy findings.json entry does not provide a complete vulnerability description."
            if language == "en-US"
            else "该问题已通过 Docker 验证，但旧版 findings.json 未提供完整漏洞描述。"
        )
    ]
    impact_text = localized_string(finding, "impact", language) or str(finding.get("impact") or "").strip()
    remediation = localized_string(finding, "remediation", language) or str(finding.get("remediation") or "").strip()
    reproduction_steps = ensure_list(finding.get("reproduction_steps"))
    report_file = str(finding.get("report_file") or "").strip()
    filename = f"{Path(report_file).stem}.docx" if report_file else build_filename({
        "project_name": project_name,
        "vuln_type": vuln_type,
        "severity_cn": severity_cn,
    }, language)
    cvss_score = str(finding.get("cvss4_score") or finding.get("cvss_score") or "").strip() or ("TBD" if language == "en-US" else "待补充")
    raw_vector = str(finding.get("cvss4_vector") or finding.get("cvss_vector") or "").strip()
    if raw_vector.startswith("CVSS:4.0/") or raw_vector.startswith("CVSS:3.1/"):
        cvss_vector = raw_vector
    elif raw_vector:
        cvss_vector = f"CVSS:3.1/{raw_vector}"
    else:
        cvss_vector = cvss4_vector_from_legacy(finding)

    attachments = []
    for item in finding.get("attachments", []):
        if isinstance(item, str) and item.strip():
            attachments.append({
                "path": item.strip(),
                "purpose": "Reproduction or evidence attachment" if language == "en-US" else "复现或证据附件",
            })
        elif isinstance(item, dict):
            attachments.append(item)

    proof = finding.get("proof_of_concept")
    proof_items = proof if isinstance(proof, dict) else {}
    proof_titles = [
        localized_string(item, "title", language) or str(key)
        for key, item in proof_items.items()
        if isinstance(item, dict)
    ]
    preconditions = localized_list(finding, "preconditions", language) or ensure_list(finding.get("preconditions"))

    analysis = []
    if location:
        analysis.append(format_kv(language, "Location" if language == "en-US" else "位置", location))
    else:
        analysis.append("Location: the legacy findings.json entry does not provide an exact file location." if language == "en-US" else "位置：旧版 findings.json 未提供精确文件位置，需补充。")
    if language == "en-US":
        analysis.append(
            "Entry / controllable input: attacker-controlled input reaches the parser or vulnerable component through the confirmed PoC vectors"
            + (f" ({'; '.join(proof_titles)})." if proof_titles else ".")
        )
    else:
        analysis.append(
            "入口/可控输入：攻击者可控输入会通过已确认 PoC 向量进入解析器或漏洞组件"
            + (f"（{'; '.join(proof_titles)}）。" if proof_titles else "。")
        )
    if vulnerable_code:
        analysis.append(format_kv(language, "Dangerous operation" if language == "en-US" else "危险函数/危险操作", vulnerable_code))
    elif location:
        analysis.append(format_kv(language, "Dangerous operation" if language == "en-US" else "危险函数/危险操作", location))
    else:
        analysis.append("Dangerous operation: the legacy findings.json entry does not provide a complete sink description." if language == "en-US" else "危险函数/危险操作：旧版 findings.json 未提供完整危险操作描述，需结合源码补充。")
    analysis.append(
        "Trigger path: attacker-controlled input is parsed, normalized or transformed by the target library, then written into the result object or vulnerable sink without rejecting dangerous keys."
        if language == "en-US"
        else "触发路径：攻击者可控输入经目标库解析、规范化或转换后，在未拒绝危险键名/危险路径的情况下写入结果对象或危险操作点。"
    )
    analysis.append(
        (f"Root cause: {description[0]}" if language == "en-US" else f"根因：{description[0]}")
        if description else
        ("Root cause: the legacy findings.json entry does not provide a complete root-cause explanation." if language == "en-US" else "根因：旧版 findings.json 未提供完整根因说明，需结合源码补充。")
    )
    analysis.append(
        "Why existing checks fail: the confirmed PoC shows the current parsing and assignment path does not block the dangerous keys or values before they reach the sink."
        if language == "en-US"
        else "现有校验为何失效：已确认 PoC 表明，当前解析与赋值路径在危险键名或危险值到达 sink 前没有完成有效拦截。"
    )
    if impact_text:
        analysis.append(format_kv(language, "Practical impact / boundary" if language == "en-US" else "实际影响/边界", impact_text))
    if preconditions:
        analysis.append(
            ("Preconditions: " if language == "en-US" else "前置条件：") + "；".join(str(item) for item in preconditions)
        )

    code_context = []
    if location or vulnerable_code:
        code_context.append(
            {
                "location": location,
                "summary": raw_title or vuln_type,
                "snippet": vulnerable_code,
            }
        )

    reproduction = []
    if reproduction_steps:
        details = [
            (f"Step {idx}: {step}" if language == "en-US" else f"步骤 {idx}：{step}") for idx, step in enumerate(reproduction_steps, start=1)
        ]
        reproduction.append(
            {
                "title": "1. Reproduction Steps" if language == "en-US" else "1. 复现步骤",
                "details": details,
                "artifacts": [item["path"] for item in attachments if item.get("path")],
                "expected": ["Expected result: the issue can be reproduced in the Docker verification environment." if language == "en-US" else "预期结果：按上述步骤可在 Docker 验证环境中复现该问题。"],
                "observed": ["Observed result: this finding is marked as docker_verified=true." if language == "en-US" else "实际结果：该 finding 被标记为 docker_verified=true。"],
            }
        )
    elif proof_items:
        reproduction.append(
            {
                "title": "1. Environment Preparation" if language == "en-US" else "1. 环境准备",
                "details": [
                    "Use the Docker verification workspace and install the affected package version before running the PoC."
                    if language == "en-US"
                    else "使用 Docker 验证工作区，并在容器内安装受影响版本后执行 PoC。",
                    f"Package/version: {project_name} {finding.get('version_affected') or defaults.get('version_affected') or ''}".strip(),
                ],
                "expected": [
                    "Expected result: the Docker environment can execute the parser PoC."
                    if language == "en-US"
                    else "预期结果：Docker 环境能够执行解析器 PoC。"
                ],
                "observed": [
                    "Observed result: docker_verified=true in the structured finding."
                    if language == "en-US"
                    else "实际结果：结构化 finding 中记录 docker_verified=true。"
                ],
                "results": [
                    "Evidence: Docker verification completed for the confirmed finding."
                    if language == "en-US"
                    else "结果证据：该确认漏洞已完成 Docker 验证。"
                ],
            }
        )
        for idx, (key, item) in enumerate(proof_items.items(), start=2):
            if not isinstance(item, dict):
                continue
            title = localized_string(item, "title", language) or str(key)
            xml = str(item.get("xml") or "").strip()
            options = item.get("options")
            impact = localized_string(item, "impact", language) or str(item.get("impact") or item.get("note") or "").strip()
            details = [f"PoC vector: {title}" if language == "en-US" else f"PoC 向量：{title}"]
            if options:
                details.append(f"Options: {json.dumps(options, ensure_ascii=False)}" if language == "en-US" else f"解析选项：{json.dumps(options, ensure_ascii=False)}")
            if xml:
                details.append(f"Input XML: {xml}" if language == "en-US" else f"输入 XML：{xml}")
            reproduction.append(
                {
                    "title": f"{idx}. Verify {title}" if language == "en-US" else f"{idx}. 验证 {title}",
                    "details": details,
                    "expected": [
                        "Expected result: the dangerous property or behavior is observable in the parsed result."
                        if language == "en-US"
                        else "预期结果：解析结果中可观察到危险属性或危险行为。"
                    ],
                    "observed": [
                        f"Observed result: {impact}" if language == "en-US" else f"实际结果：{impact or title}"
                    ],
                    "results": [
                        f"Evidence: {title} was confirmed in Docker."
                        if language == "en-US"
                        else f"结果证据：{title} 已在 Docker 中确认。"
                    ],
                }
            )

    repo_url = str(defaults.get("repo_url") or defaults.get("repository") or defaults.get("repository_url") or finding.get("repository_url") or "").strip()
    affected_versions = str(
        finding.get("version_affected")
        or defaults.get("version_affected")
        or defaults.get("version")
        or finding.get("affected_versions")
        or ""
    ).strip()
    rationale = [
        (f"Assessment rationale: {impact_text}" if language == "en-US" else f"评估依据：{impact_text}") if impact_text else ("Assessment rationale is not provided yet." if language == "en-US" else "评估依据：请根据实际验证到的影响范围补充评分依据。")
    ]

    return {
        "project_name": project_name,
        "report_language": language,
        "project_root_dir": str(finding.get("project_root_dir", "")).strip() or str(finding.get("project_root", "")).strip(),
        "vuln_type": vuln_type,
        "severity_cn": severity_cn,
        "filename": filename,
        "title": raw_title or build_title({"project_name": project_name, "vuln_type": vuln_type}, language),
        "description": description,
        "impact": {
            "package": project_name,
            "component": location,
            "affected_versions": affected_versions,
            "repo_url": repo_url,
            "extra": [impact_text] if impact_text else [],
        },
        "cvss": {
            "vector": cvss_vector,
            "score": cvss_score,
            "severity": severity_label_from_score(cvss_score, language),
            "rationale": rationale,
        },
        "analysis": analysis,
        "code_context": code_context,
        "reproduction": reproduction,
        "final_verdict": [
            "Verified in Docker; this issue is confirmed as a real vulnerability." if language == "en-US" else "经 Docker 验证，该问题已被标记为确认漏洞。",
            remediation or ("Remediation: provide the minimal secure code fix based on the actual implementation." if language == "en-US" else "修复建议：请结合具体源码实现补充最小安全修复方案。"),
        ],
        "attachments": attachments,
        "docker_verified": bool(finding.get("docker_verified")),
    }


def load_findings(input_path: Path, cli_language: str | None = None) -> list[dict[str, Any]]:
    data = json.loads(input_path.read_text(encoding="utf-8"))
    if isinstance(data, dict) and isinstance(data.get("findings"), list):
        raw_findings = [item for item in data["findings"] if isinstance(item, dict)]
        defaults = {k: v for k, v in data.items() if k != "findings"}
        if cli_language:
            defaults["report_language"] = cli_language
            defaults["output_language"] = cli_language
        converted = []
        for item in raw_findings:
            has_standard_shape = (
                isinstance(item.get("cvss"), dict)
                and isinstance(item.get("impact"), dict)
                and isinstance(item.get("analysis"), list)
                and isinstance(item.get("reproduction"), list)
            )
            has_legacy_shape = any(
                key in item
                for key in (
                    "title_zh",
                    "description_zh",
                    "impact_zh",
                    "cvss_score",
                    "cvss_vector",
                    "affected_component",
                    "proof_of_concept",
                )
            )
            converted.append(transform_legacy_finding(item, defaults) if has_legacy_shape and not has_standard_shape else item)
        return converted
    if isinstance(data, dict):
        return [data]
    if isinstance(data, list):
        return [item for item in data if isinstance(item, dict)]
    raise ValueError("Input JSON must be an object or an array of objects.")


def validate_output_dir(output_dir: Path) -> None:
    if output_dir.name != "confirmed":
        raise SystemExit(
            "Confirmed vulnerability reports must be rendered into <repo>/<audit-workspace>/confirmed, "
            f"but got: {output_dir}"
        )


def resolve_finding_language(finding: dict[str, Any], cli_language: str | None) -> str:
    if cli_language:
        return cli_language
    return normalize_language(
        finding.get("report_language") or finding.get("output_language"),
        "zh-CN",
    )


def render_finding(
    finding: dict[str, Any],
    output_path: Path,
    template_path: Path | None,
    project_root: Path,
    language: str,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if finding.get("docker_verified") is False:
        raise SystemExit(
            f"Refusing to render unverified finding '{build_title(finding, language)}': docker_verified is false."
        )
    attachments_dir = output_path.parent / "attachments"
    if attachments_dir.exists():
        shutil.rmtree(attachments_dir)
    for stale in output_path.parent.iterdir():
        if not stale.is_file():
            continue
        if (
            stale.suffix in {".docx", ".md"}
            or (stale.suffix == ".sh" and stale.name.startswith("run-"))
            or stale.name == "verification-evidence.json"
        ):
            stale.unlink()
    path_map = collect_bundle_metadata(finding, project_root, output_path.parent)
    bundle_finding = prepare_finding_for_bundle(finding, path_map, project_root)
    bundle_root_artifacts = collect_bundle_root_artifacts(bundle_finding, project_root, output_path.parent, language, path_map)

    doc = Document(template_path) if template_path else Document()
    clear_document(doc)
    doc.add_paragraph(build_title(bundle_finding, language), style="Title")
    doc.add_heading(tr(language, "vuln_description"), level=1)
    add_paragraphs(doc, localized_list(bundle_finding, "description", language))
    doc.add_heading(tr(language, "impact_versions"), level=1)
    doc.add_paragraph(render_impact_paragraph(ensure_mapping(bundle_finding.get("impact")), language))
    doc.add_heading(tr(language, "risk_assessment"), level=1)
    cvss = ensure_mapping(bundle_finding.get("cvss"))
    doc.add_paragraph(render_cvss_paragraph(cvss, language))
    add_paragraphs(doc, localized_list(cvss, "rationale", language) or [tr(language, "assessment_pending")])
    render_analysis(doc, bundle_finding, language)
    render_environment_files(doc, bundle_finding, language)
    raw_steps = bundle_finding.get("reproduction")
    reproduction_steps = [step for step in raw_steps if isinstance(step, dict)] if isinstance(raw_steps, list) else []
    render_reproduction(doc, reproduction_steps, language)
    doc.add_heading(tr(language, "final_verdict"), level=2)
    add_paragraphs(doc, localized_list(bundle_finding, "final_verdict", language) or default_final_verdict(bundle_finding, language))
    doc.save(output_path)
    validate_generated_docx(output_path)
    write_attachment_notes(output_path, finding, path_map, project_root, language, bundle_root_artifacts)
    write_reproduction_supplement(output_path, bundle_finding, language, path_map, bundle_root_artifacts)
    write_verification_evidence(output_path, finding, bundle_finding, path_map, language)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render confirmed vulnerability DOCX reports from JSON.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--template-docx")
    parser.add_argument("--language")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    validate_output_dir(output_dir)
    template_path = choose_template(args.template_docx)
    cli_language = normalize_language(args.language, "zh-CN") if args.language else None
    effective_language = cli_language or read_workspace_output_language(output_dir) or "zh-CN"
    findings = load_findings(input_path, effective_language)
    if not findings:
        raise SystemExit("No findings found in input JSON.")
    project_root = resolve_project_root(output_dir, findings[0])
    for finding in findings:
        language = resolve_finding_language(finding, effective_language)
        docx_name = build_filename(finding, language)
        bundle_dir = output_dir / build_bundle_dirname(docx_name)
        render_finding(finding, bundle_dir / docx_name, template_path, project_root, language)
    print(f"Generated {len(findings)} report(s) in {output_dir}")


if __name__ == "__main__":
    main()
