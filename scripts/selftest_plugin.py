#!/usr/bin/env python3
from __future__ import annotations

import json
import hashlib
import importlib.util
import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from pathlib import Path
from unittest import mock
from xml.etree import ElementTree as ET


REPLAY_TRANSCRIPT_CORPUS_FILES = [
    "assets/fixtures/replay-transcript-corpus/manifest.json",
    "assets/fixtures/replay-transcript-corpus/positive-multi-command-transcript.log",
    "assets/fixtures/replay-transcript-corpus/positive-compose-service-transcript.log",
    "assets/fixtures/replay-transcript-corpus/positive-copied-with-provenance-transcript.log",
    "assets/fixtures/replay-transcript-corpus/negative-marker-only.log",
    "assets/fixtures/replay-transcript-corpus/negative-placeholder-only.log",
    "assets/fixtures/replay-transcript-corpus/negative-thin-explanatory.log",
    "assets/fixtures/replay-transcript-corpus/negative-oracle-missing.log",
    "assets/fixtures/replay-transcript-corpus/negative-copied-without-provenance.log",
]

AUDIT_STATE_PROTOCOL_R2_FIXTURE_FILES = [
    "assets/fixtures/audit-state-protocol-r2/fixture-manifest.json",
    "assets/fixtures/audit-state-protocol-r2/line-ending-cases.json",
    "assets/fixtures/audit-state-protocol-r2/valid-event-r2.json",
    "assets/fixtures/audit-state-protocol-r2/valid-state-r2.json",
    "assets/fixtures/audit-state-protocol-r2/valid-events-r2.jsonl",
    "assets/fixtures/audit-state-protocol-r2/valid-policy-events-r2.jsonl",
    "assets/fixtures/audit-state-protocol-r2/valid-new-event-r2.json",
    "assets/fixtures/audit-state-protocol-r2/legacy-event-r1.json",
    "assets/fixtures/audit-state-protocol-r2/legacy-state-r1.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-missing-seq.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-stage.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-status.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-reason-code.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-absolute-evidence.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-parent-evidence.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-file-uri-evidence.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-https-uri-evidence.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-unexpected-property.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-incomplete-transition-metadata.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-event-unknown-transition-kind.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-state-negative-revision.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-state-malformed-digest.json",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-zero-seq.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-duplicate-seq.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-nonmonotonic-seq.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-truncated.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-seq-gap.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-revision-chain.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-run-id-drift.jsonl",
    "assets/fixtures/audit-state-protocol-r2/invalid-events-middle-corruption.jsonl",
]

RECON_RESULT_FIXTURE_FILES = [
    "assets/fixtures/recon-result/README.md",
    "assets/fixtures/recon-result/manifest.json",
    "assets/fixtures/recon-result/service/repo/README.md",
    "assets/fixtures/recon-result/service/repo/src/app.py",
    "assets/fixtures/recon-result/service/repo/src/policy.py",
    "assets/fixtures/recon-result/service/workspace/zhulong-target.yaml",
    "assets/fixtures/recon-result/service/workspace/attack-surface.md",
    "assets/fixtures/recon-result/service/workspace/evidence/recon-notes.md",
    "assets/fixtures/recon-result/service/workspace/evidence/deployment.md",
    "assets/fixtures/recon-result/service/workspace/cases/complete-service.json",
    "assets/fixtures/recon-result/service/workspace/cases/partial-service.json",
    "assets/fixtures/recon-result/service/workspace/cases/blocked-service.json",
    "assets/fixtures/recon-result/library/repo/README.md",
    "assets/fixtures/recon-result/library/repo/src/parser.py",
    "assets/fixtures/recon-result/library/workspace/zhulong-target.yaml",
    "assets/fixtures/recon-result/library/workspace/attack-surface.md",
    "assets/fixtures/recon-result/library/workspace/evidence/recon-notes.md",
    "assets/fixtures/recon-result/library/workspace/evidence/consumer-boundary.md",
    "assets/fixtures/recon-result/library/workspace/cases/complete-library.json",
]

REQUIRED_FILES = [
    ".claude-plugin/plugin.json",
    ".codex-plugin/plugin.json",
    "AGENTS.md",
    "README.md",
    "assets/tool-registry.json",
    "assets/schemas/tool-registry.schema.json",
    "assets/context-catalog.json",
    "assets/schemas/context-catalog.schema.json",
    "assets/schemas/context-plan.schema.json",
    "assets/fixtures/context-planning/README.md",
    "assets/fixtures/context-planning/manifest.json",
    "assets/fixtures/tool-registry/README.md",
    "assets/fixtures/tool-registry/manifest.json",
    "assets/confirmed-vuln-report-template.docx",
    "assets/examples/confirmed-findings.example.json",
    "assets/examples/zhulong-target.example.yaml",
    "assets/examples/candidate.example.json",
    "assets/examples/verifier-verdict.example.json",
    "assets/references/false-positive-template.md",
    "assets/references/unverified-lead-template.md",
    "assets/references/final-summary-template.md",
    "assets/references/docker-resource-hygiene.md",
    "assets/references/docker-registry-fallbacks.example.json",
    "assets/references/bundle-rule-mapping.md",
    "assets/references/bundle-contract-template.json",
    "assets/references/bundle-generation-checklist.md",
    "assets/references/reviewer-readiness-validator-gates.md",
    "assets/references/p8-bundle-generation-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-metrics.json",
    "assets/references/p9-protocol-chain-real-workspace-dogfood-report.md",
    "assets/references/p9-protocol-chain-real-workspace-dogfood-metrics.json",
    "assets/references/variant-seed-template.md",
    "assets/schemas/bundle-contract.schema.json",
    "assets/schemas/variant-seed.schema.json",
    "assets/schemas/zhulong-target.schema.json",
    "assets/schemas/candidate.schema.json",
    "assets/schemas/candidate-identity-input.schema.json",
    "assets/schemas/candidate-dedup-inventory.schema.json",
    "assets/schemas/candidate-dedup-plan.schema.json",
    "assets/fixtures/candidate-identity/README.md",
    "assets/fixtures/candidate-identity/manifest.json",
    "assets/schemas/verifier-verdict.schema.json",
    "assets/schemas/recording-evidence.schema.json",
    "assets/schemas/audit-event.schema.json",
    "assets/schemas/stage-status.schema.json",
    "assets/schemas/recon-result.schema.json",
    "assets/schemas/triage-batch.schema.json",
    "assets/schemas/handoff-state.schema.json",
    "assets/schemas/workspace-checkpoint.schema.json",
    "assets/schemas/next-actions.schema.json",
    "assets/schemas/audit-timeline.schema.json",
    "assets/fixtures/next-actions/manifest.json",
    "assets/fixtures/audit-timeline/README.md",
    "assets/fixtures/audit-timeline/manifest.json",
    "assets/references/recon-result-template.json",
    "assets/fixtures/triage-batch/README.md",
    "assets/fixtures/triage-batch/manifest.json",
    "assets/fixtures/recording-evidence/README.md",
    "assets/fixtures/recording-evidence/manifest.template.json",
    "assets/fixtures/handoff-checkpoint/README.md",
    "assets/fixtures/handoff-checkpoint/manifest.json",
    "assets/references/java-web-audit-playbook.md",
    "assets/references/go-web-audit-playbook.md",
    "assets/references/nodejs-library-audit-playbook.md",
    "assets/references/nodejs-web-audit-playbook.md",
    "assets/references/php-swoole-audit-playbook.md",
    "assets/references/python-library-audit-playbook.md",
    "assets/references/python-web-audit-playbook.md",
    "assets/references/ssrf-checklist.md",
    "assets/references/path-traversal-checklist.md",
    "assets/references/prototype-pollution-checklist.md",
    "scripts/bootstrap_verification_workspace.sh",
    "scripts/asr_start.sh",
    "scripts/resolve_skill_root.sh",
    "scripts/zhulong_audit.sh",
    "scripts/prepare_target_repo.sh",
    "scripts/check_docker_gate.sh",
    "scripts/check_omc_runtime.sh",
    "scripts/check_sandbox_preflight.py",
    "scripts/check_security_tooling.sh",
    "scripts/run_initial_probes.sh",
    "scripts/run_verification_case.sh",
    "scripts/evidence_io.py",
    "scripts/manage_docker_resources.py",
    "scripts/workspace_state.py",
    "scripts/render_handoff_summary.py",
    "scripts/render_handoff_state.py",
    "scripts/validate_handoff_state.py",
    "scripts/create_workspace_checkpoint.py",
    "scripts/validate_workspace_checkpoint.py",
    "scripts/next_actions.py",
    "scripts/render_next_actions.py",
    "scripts/validate_next_actions.py",
    "scripts/audit_timeline.py",
    "scripts/render_audit_timeline.py",
    "scripts/validate_audit_timeline.py",
    "scripts/selftest_audit_timeline.py",
    "docs/runner-contracts/next-actions-contract-r1.md",
    "docs/runner-contracts/audit-timeline-r1.md",
    "docs/runner-contracts/candidate-identity-dedupe-r1.md",
    "scripts/assert_finalized_workspace.py",
    "scripts/audit_disposition.py",
    "scripts/blocked_verification.py",
    "scripts/refresh_workspace_helpers.sh",
    "scripts/sync_to_claude_skill.sh",
    "scripts/sync_to_codex_skill.sh",
    "scripts/write_audit_event.py",
    "scripts/audit_state_io.py",
    "scripts/audit_text_safety.py",
    "scripts/audit_transition_policy.py",
    "scripts/validate_audit_protocol.py",
    "scripts/recover_audit_state.py",
    "scripts/selftest_audit_state_protocol.py",
    "scripts/validate_workspace_state.py",
    "scripts/validate_target_contract.py",
    "scripts/validate_recon_result.py",
    "scripts/validate_triage_batch.py",
    "scripts/finalize_stage.py",
    "scripts/validate_candidate.py",
    "scripts/candidate_identity.py",
    "scripts/upgrade_candidate_identity.py",
    "scripts/candidate_dedup.py",
    "scripts/build_candidate_dedup_plan.py",
    "scripts/validate_candidate_dedup_plan.py",
    "scripts/validate_verifier_verdict.py",
    "scripts/validate_bundle_contract.py",
    "scripts/build_confirmed_bundle.py",
    "scripts/p8_dogfood_metrics.py",
    "scripts/verify_candidate.py",
    "scripts/plan_security_toolchain.py",
    "scripts/validate_tool_registry.py",
    "scripts/context_catalog.py",
    "scripts/validate_context_catalog.py",
    "scripts/plan_audit_context.py",
    "scripts/validate_context_plan.py",
    "scripts/render_confirmed_vuln_docx.py",
    "scripts/recording_identity.py",
    "scripts/auto_record_bundle.py",
    "scripts/validate_recording_evidence.py",
    "scripts/scaffold_bilingual_findings.py",
    "scripts/extract_variant_seed.py",
    "scripts/find_variant_candidates.py",
    "scripts/validate_report_bundle.py",
    "scripts/validate_all_report_bundles.py",
    "scripts/finalize_audit_workspace.py",
    "docs/runner-contracts/target-contract-r1.md",
    "docs/runner-contracts/finding-contract-r1.md",
    "docs/runner-contracts/independent-verifier-r1.md",
    "docs/runner-contracts/disposition-integration-r1.md",
    "docs/runner-contracts/candidate-identity-dedupe-r1.md",
    "docs/runner-contracts/contract-layer-r1-closure.md",
    "docs/runner-contracts/audit-state-protocol-r2.md",
    "docs/runner-contracts/recon-result-contract-r1.md",
    "docs/runner-contracts/triage-batch-contract-r1.md",
    "docs/runner-contracts/tool-effects-execution-boundaries-r1.md",
    "docs/runner-contracts/context-planning-r1.md",
    "assets/fixtures/audit-state-protocol-r2/fixture-manifest.json",
    "assets/fixtures/audit-state-protocol-r2/line-ending-cases.json",
    "assets/fixtures/contracts/confirmed_ssrf/zhulong-target.yaml",
    "assets/fixtures/contracts/confirmed_ssrf/candidate.json",
    "assets/fixtures/contracts/confirmed_ssrf/verifier-verdict.json",
    "assets/fixtures/contracts/confirmed_ssrf/expected-disposition.json",
    "assets/fixtures/contracts/false_positive_unreachable/zhulong-target.yaml",
    "assets/fixtures/contracts/false_positive_unreachable/candidate.json",
    "assets/fixtures/contracts/false_positive_unreachable/verifier-verdict.json",
    "assets/fixtures/contracts/false_positive_unreachable/expected-disposition.json",
    "assets/fixtures/contracts/unverified_oracle_weak/zhulong-target.yaml",
    "assets/fixtures/contracts/unverified_oracle_weak/candidate.json",
    "assets/fixtures/contracts/unverified_oracle_weak/verifier-verdict.json",
    "assets/fixtures/contracts/unverified_oracle_weak/expected-disposition.json",
    "assets/fixtures/contracts/blocked_manual_runtime/zhulong-target.yaml",
    "assets/fixtures/contracts/blocked_manual_runtime/candidate.json",
    "assets/fixtures/contracts/blocked_manual_runtime/verifier-verdict.json",
    "assets/fixtures/contracts/blocked_manual_runtime/expected-disposition.json",
    "assets/fixtures/p8-dogfood/README.md",
    "assets/fixtures/p8-dogfood/bad-contract.bundle-contract.json",
    "assets/fixtures/p8-dogfood/marker-only-replay-output.log",
    "assets/fixtures/p8-real-historical-bundle-dogfood/README.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/.contracts/historical-sample-01.bundle-contract.json",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/historical-sample-01/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/placeholder-replay-output.log",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-02/confirmed/historical-draft-bundle/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-03/confirmed/.contracts/historical-sample-03.bundle-contract.json",
    *REPLAY_TRANSCRIPT_CORPUS_FILES,
    *AUDIT_STATE_PROTOCOL_R2_FIXTURE_FILES,
    *RECON_RESULT_FIXTURE_FILES,
    "skills/zhulong/SKILL.md",
    "templates/claude-skill/SKILL.md",
]

INSTALLED_SKILL_REQUIRED_FILES = [
    "SKILL.md",
    "assets/schemas/next-actions.schema.json",
    "assets/schemas/audit-timeline.schema.json",
    "assets/fixtures/next-actions/manifest.json",
    "assets/fixtures/audit-timeline/README.md",
    "assets/fixtures/audit-timeline/manifest.json",
    "docs/runner-contracts/next-actions-contract-r1.md",
    "docs/runner-contracts/audit-timeline-r1.md",
    "scripts/next_actions.py",
    "scripts/render_next_actions.py",
    "scripts/validate_next_actions.py",
    "scripts/audit_timeline.py",
    "scripts/render_audit_timeline.py",
    "scripts/validate_audit_timeline.py",
    "scripts/selftest_audit_timeline.py",
    "README.plugin-package.md",
    "INSTALL.plugin-package.md",
    "assets/tool-registry.json",
    "assets/schemas/tool-registry.schema.json",
    "assets/context-catalog.json",
    "assets/schemas/context-catalog.schema.json",
    "assets/schemas/context-plan.schema.json",
    "assets/fixtures/context-planning/README.md",
    "assets/fixtures/context-planning/manifest.json",
    "assets/fixtures/tool-registry/README.md",
    "assets/fixtures/tool-registry/manifest.json",
    "assets/confirmed-vuln-report-template.docx",
    "assets/references/docker-resource-hygiene.md",
    "assets/references/docker-registry-fallbacks.example.json",
    "assets/references/bundle-rule-mapping.md",
    "assets/references/bundle-contract-template.json",
    "assets/references/bundle-generation-checklist.md",
    "assets/references/reviewer-readiness-validator-gates.md",
    "assets/references/p8-bundle-generation-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-metrics.json",
    "assets/references/p9-protocol-chain-real-workspace-dogfood-report.md",
    "assets/references/p9-protocol-chain-real-workspace-dogfood-metrics.json",
    "assets/references/variant-seed-template.md",
    "assets/schemas/bundle-contract.schema.json",
    "assets/schemas/variant-seed.schema.json",
    "assets/schemas/zhulong-target.schema.json",
    "assets/schemas/candidate.schema.json",
    "assets/schemas/candidate-identity-input.schema.json",
    "assets/schemas/candidate-dedup-inventory.schema.json",
    "assets/schemas/candidate-dedup-plan.schema.json",
    "assets/fixtures/candidate-identity/README.md",
    "assets/fixtures/candidate-identity/manifest.json",
    "assets/schemas/verifier-verdict.schema.json",
    "assets/schemas/recording-evidence.schema.json",
    "assets/schemas/audit-event.schema.json",
    "assets/schemas/stage-status.schema.json",
    "assets/schemas/recon-result.schema.json",
    "assets/schemas/triage-batch.schema.json",
    "assets/schemas/handoff-state.schema.json",
    "assets/schemas/workspace-checkpoint.schema.json",
    "assets/references/recon-result-template.json",
    "assets/fixtures/triage-batch/README.md",
    "assets/fixtures/triage-batch/manifest.json",
    "assets/fixtures/recording-evidence/README.md",
    "assets/fixtures/recording-evidence/manifest.template.json",
    "assets/fixtures/handoff-checkpoint/README.md",
    "assets/fixtures/handoff-checkpoint/manifest.json",
    "assets/examples/zhulong-target.example.yaml",
    "assets/examples/candidate.example.json",
    "assets/examples/verifier-verdict.example.json",
    "assets/references/nodejs-web-audit-playbook.md",
    "assets/references/php-swoole-audit-playbook.md",
    "assets/references/python-library-audit-playbook.md",
    "docs/CODEX_SKILL_ADAPTATION.md",
    "docs/INSTALL.md",
    "docs/USAGE.md",
    "scripts/asr_start.sh",
    "scripts/resolve_skill_root.sh",
    "scripts/zhulong_audit.sh",
    "scripts/bootstrap_verification_workspace.sh",
    "scripts/check_docker_gate.sh",
    "scripts/check_omc_runtime.sh",
    "scripts/check_sandbox_preflight.py",
    "scripts/check_security_tooling.sh",
    "scripts/run_initial_probes.sh",
    "scripts/run_verification_case.sh",
    "scripts/evidence_io.py",
    "scripts/manage_docker_resources.py",
    "scripts/workspace_state.py",
    "scripts/audit_state_io.py",
    "scripts/audit_text_safety.py",
    "scripts/audit_transition_policy.py",
    "scripts/validate_audit_protocol.py",
    "scripts/recover_audit_state.py",
    "scripts/selftest_audit_state_protocol.py",
    "scripts/render_confirmed_vuln_docx.py",
    "scripts/recording_identity.py",
    "scripts/auto_record_bundle.py",
    "scripts/validate_recording_evidence.py",
    "scripts/extract_variant_seed.py",
    "scripts/find_variant_candidates.py",
    "scripts/validate_report_bundle.py",
    "scripts/validate_target_contract.py",
    "scripts/validate_recon_result.py",
    "scripts/validate_tool_registry.py",
    "scripts/context_catalog.py",
    "scripts/validate_context_catalog.py",
    "scripts/plan_audit_context.py",
    "scripts/validate_context_plan.py",
    "scripts/validate_triage_batch.py",
    "scripts/finalize_stage.py",
    "scripts/validate_candidate.py",
    "scripts/candidate_identity.py",
    "scripts/upgrade_candidate_identity.py",
    "scripts/candidate_dedup.py",
    "scripts/build_candidate_dedup_plan.py",
    "scripts/validate_candidate_dedup_plan.py",
    "scripts/validate_verifier_verdict.py",
    "scripts/validate_bundle_contract.py",
    "scripts/build_confirmed_bundle.py",
    "scripts/p8_dogfood_metrics.py",
    "scripts/verify_candidate.py",
    "scripts/validate_all_report_bundles.py",
    "scripts/finalize_audit_workspace.py",
    "scripts/assert_finalized_workspace.py",
    "scripts/audit_disposition.py",
    "scripts/blocked_verification.py",
    "scripts/render_handoff_summary.py",
    "scripts/render_handoff_state.py",
    "scripts/validate_handoff_state.py",
    "scripts/create_workspace_checkpoint.py",
    "scripts/validate_workspace_checkpoint.py",
    "docs/runner-contracts/target-contract-r1.md",
    "docs/runner-contracts/finding-contract-r1.md",
    "docs/runner-contracts/independent-verifier-r1.md",
    "docs/runner-contracts/disposition-integration-r1.md",
    "docs/runner-contracts/candidate-identity-dedupe-r1.md",
    "docs/runner-contracts/contract-layer-r1-closure.md",
    "docs/runner-contracts/audit-state-protocol-r2.md",
    "docs/runner-contracts/recon-result-contract-r1.md",
    "docs/runner-contracts/triage-batch-contract-r1.md",
    "docs/runner-contracts/tool-effects-execution-boundaries-r1.md",
    "docs/runner-contracts/context-planning-r1.md",
    "assets/fixtures/audit-state-protocol-r2/fixture-manifest.json",
    "assets/fixtures/audit-state-protocol-r2/line-ending-cases.json",
    "assets/fixtures/contracts/confirmed_ssrf/zhulong-target.yaml",
    "assets/fixtures/contracts/confirmed_ssrf/candidate.json",
    "assets/fixtures/contracts/confirmed_ssrf/verifier-verdict.json",
    "assets/fixtures/contracts/confirmed_ssrf/expected-disposition.json",
    "assets/fixtures/contracts/false_positive_unreachable/zhulong-target.yaml",
    "assets/fixtures/contracts/false_positive_unreachable/candidate.json",
    "assets/fixtures/contracts/false_positive_unreachable/verifier-verdict.json",
    "assets/fixtures/contracts/false_positive_unreachable/expected-disposition.json",
    "assets/fixtures/contracts/unverified_oracle_weak/zhulong-target.yaml",
    "assets/fixtures/contracts/unverified_oracle_weak/candidate.json",
    "assets/fixtures/contracts/unverified_oracle_weak/verifier-verdict.json",
    "assets/fixtures/contracts/unverified_oracle_weak/expected-disposition.json",
    "assets/fixtures/contracts/blocked_manual_runtime/zhulong-target.yaml",
    "assets/fixtures/contracts/blocked_manual_runtime/candidate.json",
    "assets/fixtures/contracts/blocked_manual_runtime/verifier-verdict.json",
    "assets/fixtures/contracts/blocked_manual_runtime/expected-disposition.json",
    "assets/fixtures/p8-dogfood/README.md",
    "assets/fixtures/p8-dogfood/bad-contract.bundle-contract.json",
    "assets/fixtures/p8-dogfood/marker-only-replay-output.log",
    "assets/fixtures/p8-real-historical-bundle-dogfood/README.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/.contracts/historical-sample-01.bundle-contract.json",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/historical-sample-01/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/placeholder-replay-output.log",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-02/confirmed/historical-draft-bundle/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-03/confirmed/.contracts/historical-sample-03.bundle-contract.json",
    *REPLAY_TRANSCRIPT_CORPUS_FILES,
    *AUDIT_STATE_PROTOCOL_R2_FIXTURE_FILES,
    *RECON_RESULT_FIXTURE_FILES,
]

P8_RUNTIME_FILES = [
    "scripts/validate_bundle_contract.py",
    "scripts/build_confirmed_bundle.py",
    "scripts/p8_dogfood_metrics.py",
    "scripts/validate_report_bundle.py",
    "scripts/validate_all_report_bundles.py",
]

P8_REFERENCE_FILES = [
    "assets/schemas/bundle-contract.schema.json",
    "assets/references/bundle-rule-mapping.md",
    "assets/references/bundle-contract-template.json",
    "assets/references/bundle-generation-checklist.md",
    "assets/references/reviewer-readiness-validator-gates.md",
    "assets/references/p8-bundle-generation-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-metrics.json",
]

BUNDLE_RULE_MAPPING_REQUIRED_FIELDS = [
    "schema_version",
    "bundle",
    "bundle.slug",
    "bundle.language",
    "bundle.final_path",
    "bundle.one_vulnerability_only",
    "render",
    "render.source_findings_json",
    "render.finding_slug",
    "finding",
    "finding.project_name",
    "finding.vulnerability_name",
    "finding.bug_class",
    "finding.severity",
    "finding.attacker_condition",
    "finding.server_condition",
    "finding.security_impact",
    "docker_evidence",
    "docker_evidence.verification_status",
    "docker_evidence.docker_required",
    "docker_evidence.docker_command",
    "docker_evidence.oracle_token",
    "docker_evidence.expected_observation",
    "docker_evidence.observed_observation",
    "docker_evidence.severity_escalation_attempted",
    "entrypoint_evidence",
    "entrypoint_evidence.evidence_level",
    "entrypoint_evidence.attacker_controlled_entrypoint",
    "entrypoint_evidence.input_shape",
    "entrypoint_evidence.entrypoint_to_sink_path",
    "entrypoint_evidence.deterministic_impact_oracle",
    "entrypoint_evidence.replay_material",
    "replay",
    "replay.root_script.path",
    "replay.log.path",
    "replay.log.registration_targets",
    "direct_impact",
    "direct_impact.marker",
    "direct_impact.sync_targets",
    "files",
    "files.verification_evidence",
    "files.reviewer_evidence_index",
    "files.evidence_files",
    "files.attachments",
    "code_context",
    "code_context.entries",
    "source_binding",
    "source_binding.tested_ref",
    "source_binding.attacker_entrypoint",
    "source_binding.replay_observed_entrypoint",
    "source_binding.binding_mode",
    "source_binding.source_references",
    "fixture_provenance",
    "fixture_provenance.required",
    "fixture_provenance.replay_type",
    "fixture_provenance.synthetic_security_properties_present",
    "fixture_provenance.security_properties",
    "impact_claims",
    "deployment_prerequisites",
    "validity_review",
    "validity_review.validity_verdict",
    "validity_review.classification_decision",
    "validity_review.final_bug_class",
    "validity_review.final_severity",
    "validity_review.supported_impact_claim_ids",
    "validity_review.deployment_prerequisite_ids",
    "impact_tier",
    "impact_tier.bug_class",
    "impact_tier.ssrf.tier",
    "impact_tier.ssrf.claimed_exposures",
    "impact_tier.ssrf.stronger_impacts_not_claimed",
    "impact_tier.ssrf.artifact_backed_oracle",
    "variant_seed_readiness",
    "variant_seed_readiness.run_after_promote",
]

BUNDLE_RULE_MAPPING_COLUMNS = [
    "Contract field",
    "Readiness meaning",
    "Renderer / builder output",
    "Final validator or batch gate",
    "Evidence artifact",
    "Notes / non-claims",
]

BUNDLE_RULE_MAPPING_FORBIDDEN_CLAIMS = [
    "preflight proves a vulnerability",
    "preflight confirms a vulnerability",
    "contract proves a vulnerability",
    "contract confirms a vulnerability",
    "预检证明漏洞",
    "合同证明漏洞",
]

REVIEWER_READINESS_GATE_FAMILIES = [
    "SSRF Impact Overclaim",
    "Code Context Minimum Quality",
    "Replay Helper Pause Contract",
]

REVIEWER_READINESS_GATE_REQUIRED_TEXT = [
    "Reviewer-readiness gates improve",
    "They are final-bundle quality gates",
    "do not discover",
    "do not replace Docker evidence",
    "must only add stricter rejection",
    "False-positive boundary",
    "Accepted example",
    "Rejected example",
    "SSRF_IMPACT_OVERCLAIM",
    "CODE_CONTEXT_MINIMUM_QUALITY",
    "REPLAY_HELPER_PAUSE_CONTRACT",
    "ROOT_SCRIPT_CONTEXT_MISSING",
    "Any new reviewer-readiness validator gate",
    "positive and negative",
]

REVIEWER_READINESS_GATE_FORBIDDEN_CLAIMS = [
    "these gates prove vulnerabilities",
    "reviewer-readiness gates prove vulnerabilities",
    "these gates confirm vulnerabilities",
    "reviewer-readiness gates confirm vulnerabilities",
]

STABLE_CONTRACT_SEVERITIES = [
    "Critical",
    "High",
    "Medium",
    "Low",
    "Informational",
]

RECOMMENDED_BUG_CLASS_TEXT = [
    "recommended bug classes",
    "free text",
    "SSRF",
    "Path Traversal",
    "Prototype Pollution",
    "Command Injection",
    "Deserialization",
    "Authentication Bypass",
    "Authorization Bypass",
    "Information Disclosure",
    "Denial of Service",
]

P8_DOGFOOD_FILES = [
    "assets/fixtures/p8-dogfood/README.md",
    "assets/fixtures/p8-dogfood/bad-contract.bundle-contract.json",
    "assets/fixtures/p8-dogfood/marker-only-replay-output.log",
    "assets/references/p8-bundle-generation-dogfood-report.md",
    "scripts/p8_dogfood_metrics.py",
]

P8_REAL_HISTORICAL_DOGFOOD_FILES = [
    "assets/fixtures/p8-real-historical-bundle-dogfood/README.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/.contracts/historical-sample-01.bundle-contract.json",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/confirmed/historical-sample-01/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-01/placeholder-replay-output.log",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-02/confirmed/historical-draft-bundle/README.sanitized.md",
    "assets/fixtures/p8-real-historical-bundle-dogfood/workspaces/historical-sample-03/confirmed/.contracts/historical-sample-03.bundle-contract.json",
    "assets/references/p8-real-historical-bundle-dogfood-report.md",
    "assets/references/p8-real-historical-bundle-dogfood-metrics.json",
]

REPLAY_TRANSCRIPT_CORPUS_REQUIRED_IDS = {
    "positive-multi-command-transcript",
    "positive-compose-service-transcript",
    "positive-copied-with-provenance-transcript",
    "negative-marker-only",
    "negative-placeholder-only",
    "negative-thin-explanatory",
    "negative-oracle-missing",
    "negative-copied-without-provenance",
}

FORBIDDEN_INSTALLED_TOP_LEVEL = [
    "prompts",
    "zhulong-real-runs",
    "已提交",
    ".codex",
    ".claude",
    ".omc",
    "AGENTS.md",
]

PACKAGE_RESIDUE_SUFFIXES = (".hidden", ".bak", ".tmp", ".orig", ".rej", ".pyc")
PACKAGE_RESIDUE_NAMES = {"AGENTS.md", ".DS_Store"}
PACKAGE_RESIDUE_DIR_NAMES = {"__pycache__", ".omc"}


def run(command: list[str], cwd: Path) -> None:
    env = {
        **os.environ,
        "PYTHONDONTWRITEBYTECODE": "1",
        "PYTHONPYCACHEPREFIX": str(Path(tempfile.gettempdir()) / "zhulong-selftest-pycache"),
    }
    proc = subprocess.run(command, cwd=cwd, env=env, capture_output=True, text=True)
    if proc.returncode != 0:
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        raise SystemExit(f"FAILED: {' '.join(command)}\n{output}")


def run_capture(command: list[str], cwd: Path) -> str:
    proc = subprocess.run(command, cwd=cwd, capture_output=True, text=True)
    output = ((proc.stdout or "") + (proc.stderr or "")).strip()
    if proc.returncode != 0:
        raise SystemExit(f"FAILED: {' '.join(command)}\n{output}")
    return output


def require_command_output(command: list[str], cwd: Path, expected: str, label: str) -> None:
    output = run_capture(command, cwd)
    if output != expected:
        raise SystemExit(
            f"FAILED: unexpected output for {label}\n"
            f"Expected: {expected}\n"
            f"Actual: {output}"
        )


def run_capture_with_env(
    command: list[str],
    cwd: Path,
    env: dict[str, str],
    *,
    expected_returncode: int = 0,
) -> str:
    merged_env = {**os.environ, **env}
    proc = subprocess.run(command, cwd=cwd, env=merged_env, capture_output=True, text=True)
    output = ((proc.stdout or "") + (proc.stderr or "")).strip()
    if proc.returncode != expected_returncode:
        raise SystemExit(
            f"FAILED: {' '.join(command)}\n"
            f"Expected exit code {expected_returncode}, got {proc.returncode}\n{output}"
        )
    return output


def docx_text(docx_path: Path) -> list[str]:
    with zipfile.ZipFile(docx_path) as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: list[str] = []
    for para in root.findall(".//w:p", ns):
        text = "".join(node.text or "" for node in para.findall(".//w:t", ns)).strip()
        if text:
            lines.append(text)
    return lines


def rewrite_docx_paragraphs(docx_path: Path, replacer) -> None:
    from docx import Document

    doc = Document(docx_path)
    for paragraph in list(doc.paragraphs):
        replacement = replacer(paragraph.text.strip())
        if replacement is None:
            element = paragraph._element
            element.getparent().remove(element)
        elif replacement != paragraph.text:
            paragraph.text = replacement
    doc.save(docx_path)


def write_live_replay_log(bundle: Path, *, marker: str = "DIRECT_IMPACT_CONFIRMED", extra: str = "") -> None:
    replay_log = bundle / "attachments/evidence/replay-output.log"
    replay_log.parent.mkdir(parents=True, exist_ok=True)
    replay_log.write_text(
        "Zhulong reviewer replay log\n"
        "Generated at: 2026-06-16T00:00:00Z\n"
        "COMMAND: docker compose -f attachments/poc/docker-compose.selftest.yml up --abort-on-container-exit\n"
        "stdout: deterministic selftest replay completed\n"
        "success marker verified with grep -Fq\n"
        f"{marker}\n"
        f"{extra}".rstrip()
        + "\n",
        encoding="utf-8",
    )


def write_live_replay_logs(*bundles: Path, marker: str = "DIRECT_IMPACT_CONFIRMED", extra: str = "") -> None:
    for bundle in bundles:
        write_live_replay_log(bundle, marker=marker, extra=extra)




def run_with_env(command: list[str], cwd: Path, env: dict[str, str]) -> None:
    merged_env = {**os.environ, **env}
    proc = subprocess.run(command, cwd=cwd, env=merged_env, capture_output=True, text=True)
    if proc.returncode != 0:
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        raise SystemExit(f"FAILED: {' '.join(command)}\n{output}")


def run_expect_fail(command: list[str], cwd: Path, expected: str,
                   extra_env: dict[str, str] | None = None) -> None:
    env = {**os.environ, **extra_env} if extra_env else None
    proc = subprocess.run(command, cwd=cwd, env=env, capture_output=True, text=True)
    output = ((proc.stdout or "") + (proc.stderr or "")).strip()
    if proc.returncode == 0:
        raise SystemExit(f"FAILED: command unexpectedly succeeded: {' '.join(command)}")
    if expected not in output:
        raise SystemExit(
            f"FAILED: command did not fail with expected text: {expected}\n"
            f"Command: {' '.join(command)}\nOutput:\n{output}"
        )


def require_text(path: Path, needle: str, label: str) -> None:
    content = path.read_text(encoding="utf-8")
    if needle not in content:
        raise SystemExit(f"FAILED: missing expected text for {label}: {needle}")


def require_installed_package_hygiene(root: Path, label: str) -> None:
    issues: list[str] = []
    for path in root.rglob("*"):
        if path.name in PACKAGE_RESIDUE_DIR_NAMES and path.is_dir():
            issues.append(path.relative_to(root).as_posix() + "/")
        elif path.is_file() and (
            path.name in PACKAGE_RESIDUE_NAMES
            or path.name.endswith(PACKAGE_RESIDUE_SUFFIXES)
        ):
            issues.append(path.relative_to(root).as_posix())
    if issues:
        raise SystemExit(
            f"FAILED: {label} contains forbidden package residue: {sorted(issues)}"
        )


def forbid_text(path: Path, needle: str, label: str) -> None:
    content = path.read_text(encoding="utf-8")
    if needle in content:
        raise SystemExit(f"FAILED: forbidden text for {label}: {needle}")


def require_no_repo_text(plugin_root: Path, needle: str, label: str) -> None:
    checked_suffixes = {".md", ".py", ".sh", ".json"}
    for path in plugin_root.rglob("*"):
        if any(part in {".git", ".omc", "__pycache__"} for part in path.parts):
            continue
        if not path.is_file() or path.suffix not in checked_suffixes:
            continue
        if needle in path.read_text(encoding="utf-8", errors="ignore"):
            raise SystemExit(f"FAILED: forbidden repository text for {label}: {path}: {needle}")


def require_files(root: Path, rels: list[str], label: str) -> None:
    for rel in rels:
        if not (root / rel).exists():
            raise SystemExit(f"FAILED: missing {label} file: {root / rel}")


def exercise_tool_registry_contract(skill_root: Path) -> None:
    """Exercise the production registry validator and planner without tools or network."""
    validator = skill_root / "scripts/validate_tool_registry.py"
    planner = skill_root / "scripts/plan_security_toolchain.py"
    registry_path = skill_root / "assets/tool-registry.json"
    schema_path = skill_root / "assets/schemas/tool-registry.schema.json"
    manifest_path = skill_root / "assets/fixtures/tool-registry/manifest.json"
    require_files(
        skill_root,
        [
            "scripts/validate_tool_registry.py",
            "scripts/plan_security_toolchain.py",
            "assets/tool-registry.json",
            "assets/schemas/tool-registry.schema.json",
            "assets/fixtures/tool-registry/README.md",
            "assets/fixtures/tool-registry/manifest.json",
        ],
        "tool registry contract",
    )
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest.get("schema_version") != 1 or manifest.get("execution_policy") != "offline-production-validator-and-planner-only":
        raise SystemExit("FAILED: tool registry fixture manifest is not the expected offline contract")
    forbidden_execution = set(manifest.get("forbidden_execution", []))
    if not {"docker", "poc", "scanner", "network", "package_manager", "github", "llm"}.issubset(forbidden_execution):
        raise SystemExit("FAILED: tool registry fixture manifest is missing the offline execution exclusions")

    canonical = json.loads(registry_path.read_text(encoding="utf-8"))
    if canonical.get("schema_version") != 2:
        raise SystemExit("FAILED: Tool Registry schema_version must be numeric 2")

    def tool(document: dict, name: str) -> dict:
        for tier in document["tiers"]:
            for entry in tier["tools"]:
                if entry["name"] == name:
                    return entry
        raise SystemExit(f"FAILED: Tool Registry selftest could not find {name}")

    def invoke(registry: Path, *declared: str) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                sys.executable,
                str(validator),
                "--skill-root",
                str(skill_root),
                "--registry",
                str(registry),
                "--schema",
                str(schema_path),
                *declared,
                "--json",
            ],
            cwd=skill_root,
            capture_output=True,
            text=True,
        )

    def payload_for(proc: subprocess.CompletedProcess[str], label: str) -> dict:
        try:
            return json.loads(proc.stdout)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: {label} did not emit JSON: {proc.stdout!r} {proc.stderr!r}") from exc

    canonical_result = invoke(registry_path)
    if canonical_result.returncode != 0:
        raise SystemExit(f"FAILED: canonical Tool Registry rejected: {canonical_result.stdout}\n{canonical_result.stderr}")
    canonical_payload = payload_for(canonical_result, "canonical Tool Registry")
    if canonical_payload.get("authority") != "tool_metadata_only" or canonical_payload.get("tool_count", 0) < 30:
        raise SystemExit("FAILED: canonical Tool Registry returned unexpected metadata authority or tool count")

    for declared in (
        ("--tool", "source-inspection", "--stage", "recon", "--boundary", "host_read_only", "--effect", "source_read"),
        ("--tool", "semgrep", "--stage", "recon", "--boundary", "workspace_write", "--effect", "workspace_evidence_write"),
        ("--tool", "docker-verification-wrapper", "--stage", "verification", "--boundary", "docker_exec", "--effect", "target_code_execute"),
    ):
        result = invoke(registry_path, *declared)
        if result.returncode != 0 or not payload_for(result, "declared Tool Registry use").get("ok"):
            raise SystemExit(f"FAILED: declared Tool Registry use should be allowed: {declared}")

    with tempfile.TemporaryDirectory(prefix="zhulong-tool-registry-") as tempdir:
        temp_root = Path(tempdir)

        def clone() -> dict:
            return json.loads(json.dumps(canonical))

        def reject(label: str, document: dict, expected_code: str) -> None:
            sample = temp_root / f"{label}.json"
            sample.write_text(json.dumps(document, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            result = invoke(sample)
            payload = payload_for(result, label)
            if result.returncode == 0 or expected_code not in payload.get("issue_codes", []):
                raise SystemExit(f"FAILED: {label} did not reject with {expected_code}: {payload}")

        unknown_enum = clone()
        tool(unknown_enum, "semgrep")["effects"] = ["unknown_effect"]
        reject("unknown-enum", unknown_enum, "SCHEMA_INVALID")

        duplicate = clone()
        tool(duplicate, "gh")["name"] = "docker"
        reject("duplicate-tool", duplicate, "DUPLICATE_TOOL")

        prohibited_effect = clone()
        tool(prohibited_effect, "docker")["effects"] = ["source_read"]
        reject("prohibited-effect", prohibited_effect, "PROHIBITED_EFFECTS_FORBIDDEN")

        external_boundary_without_scope = clone()
        external_boundary_tool = tool(external_boundary_without_scope, "source-inspection")
        external_boundary_tool["execution_boundaries"] = ["host_read_only", "external_network"]
        reject("external-boundary-without-scope", external_boundary_without_scope, "NETWORK_SCOPE_BOUNDARY_CONFLICT")

        external_scope_without_boundary = clone()
        tool(external_scope_without_boundary, "source-inspection")["network_scope"] = "public_external"
        reject("external-scope-without-boundary", external_scope_without_boundary, "NETWORK_SCOPE_BOUNDARY_CONFLICT")

        active_dast_without_network = clone()
        active_dast_tool = tool(active_dast_without_network, "nuclei")
        active_dast_tool["execution_boundaries"] = ["workspace_write"]
        active_dast_tool["effects"] = ["workspace_evidence_write"]
        active_dast_tool["network_scope"] = "none"
        active_dast_tool["timeout_policy"] = "caller_required"
        active_dast_tool["evidence_outputs"] = [{"path_family": "evidence/dast/*.json", "artifact_type": "json"}]
        active_dast_tool["confirmation_authority"] = "candidate_only"
        active_dast_tool["controlled_wrapper"] = {
            "path": "scripts/run_initial_probes.sh",
            "contract_marker": "zhulong-tool-contract: initial-probes-v1",
        }
        active_dast_tool["planner_status"] = "wrapper_required"
        reject("active-dast-without-network", active_dast_without_network, "DAST_NETWORK_BOUNDARY_MISSING")

        missing_wrapper = clone()
        tool(missing_wrapper, "semgrep")["controlled_wrapper"] = None
        reject("missing-wrapper", missing_wrapper, "WRAPPER_REQUIRED")

        raw_authority = clone()
        tool(raw_authority, "docker")["confirmation_authority"] = "docker_oracle_material_only"
        reject("raw-docker-authority", raw_authority, "RAW_DOCKER_AUTHORITY_FORBIDDEN")

        dast_authority = clone()
        tool(dast_authority, "nuclei")["confirmation_authority"] = "docker_oracle_material_only"
        reject("raw-dast-authority", dast_authority, "SCANNER_AUTHORITY_FORBIDDEN")

        missing_evidence = clone()
        tool(missing_evidence, "semgrep")["evidence_outputs"] = []
        reject("missing-evidence", missing_evidence, "WORKSPACE_EVIDENCE_MISSING")

        timeout_marker = clone()
        timeout_marker_tool = tool(timeout_marker, "docker-verification-wrapper")
        timeout_marker_tool["controlled_wrapper"]["contract_marker"] = "zhulong-tool-contract: docker-verification-v1"
        reject("timeout-marker", timeout_marker, "TIMEOUT_CONTRACT_MISSING")

        sandbox_marker = clone()
        sandbox_marker_tool = tool(sandbox_marker, "docker-verification-wrapper")
        sandbox_marker_tool["controlled_wrapper"]["contract_marker"] = "zhulong-tool-contract: docker-verification-v1; timeout=mandatory"
        reject("sandbox-marker", sandbox_marker, "SANDBOX_CONTRACT_MISSING")

        for name, unsafe_path, expected_code in (
            ("wrapper-absolute", "/tmp/wrapper.sh", "WRAPPER_PATH_UNSAFE"),
            ("wrapper-uri", "file:///tmp/wrapper.sh", "WRAPPER_PATH_UNSAFE"),
            ("wrapper-traversal", "scripts/../run_initial_probes.sh", "WRAPPER_PATH_UNSAFE"),
            ("wrapper-backslash", "scripts\\run_initial_probes.sh", "WRAPPER_PATH_UNSAFE"),
            ("wrapper-directory", "scripts", "WRAPPER_TYPE_INVALID"),
        ):
            unsafe = clone()
            tool(unsafe, "semgrep")["controlled_wrapper"]["path"] = unsafe_path
            reject(name, unsafe, expected_code)

        unsafe_evidence = clone()
        tool(unsafe_evidence, "semgrep")["evidence_outputs"][0]["path_family"] = "evidence/../../outside.log"
        reject("evidence-traversal", unsafe_evidence, "EVIDENCE_PATH_UNSAFE")

        confirmed_role = clone()
        tool(confirmed_role, "source-inspection")["role"] = "confirmed text is not confirmation authority"
        sample = temp_root / "confirmed-role-text.json"
        sample.write_text(json.dumps(confirmed_role, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        role_result = invoke(sample)
        if role_result.returncode != 0 or not payload_for(role_result, "confirmed role text").get("ok"):
            raise SystemExit("FAILED: ordinary role text mentioning confirmed was treated as an authority")
        if tool(confirmed_role, "source-inspection")["confirmation_authority"] != "none":
            raise SystemExit("FAILED: source-inspection authority changed while testing role text")

        misuse = invoke(
            registry_path,
            "--tool", "source-inspection", "--stage", "verification", "--boundary", "docker_exec", "--effect", "target_code_execute",
        )
        misuse_payload = payload_for(misuse, "declared stage misuse")
        expected_misuse = {"TOOL_STAGE_FORBIDDEN", "TOOL_BOUNDARY_FORBIDDEN", "TOOL_EFFECT_FORBIDDEN"}
        if misuse.returncode == 0 or not expected_misuse.issubset(set(misuse_payload.get("issue_codes", []))):
            raise SystemExit(f"FAILED: declared stage misuse did not fail closed: {misuse_payload}")

        target = temp_root / "target"
        workspace = target / "security-research-tool-contract"
        target.mkdir()
        workspace.mkdir()
        (target / "package.json").write_text('{"name":"tool-contract-fixture"}\n', encoding="utf-8")
        (workspace / "asr-config.json").write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "workspace_root": workspace.name,
                    "workspace_created_at": "2026-07-21T00:00:00Z",
                    "confirmed_output_dir": f"{workspace.name}/confirmed",
                }
            ) + "\n",
            encoding="utf-8",
        )
        plan_result = subprocess.run(
            [sys.executable, str(planner), "--target-dir", str(target), "--workspace-dir", str(workspace), "--format", "json"],
            cwd=skill_root,
            capture_output=True,
            text=True,
        )
        if plan_result.returncode != 0:
            raise SystemExit(f"FAILED: Tool Registry planner rejected canonical input: {plan_result.stdout}\n{plan_result.stderr}")
        plan = payload_for(plan_result, "Tool Registry planner")
        catalog = {entry["name"]: entry for entry in plan.get("tool_catalog", [])}
        if catalog.get("source-inspection", {}).get("confirmation_authority") != "none":
            raise SystemExit("FAILED: source inspection received confirmation authority in planner metadata")
        if catalog.get("semgrep", {}).get("confirmation_authority") != "candidate_only":
            raise SystemExit("FAILED: scanner planner metadata is not candidate-only")
        if catalog.get("docker-verification-wrapper", {}).get("confirmation_authority") != "docker_oracle_material_only":
            raise SystemExit("FAILED: controlled Docker wrapper metadata lost oracle-material authority")
        hints = plan.get("command_hints", [])
        forbidden_hints = ("docker ", "nuclei", "ffuf", "sqlmap", "zap", "http://", "https://")
        if any(any(token in hint.lower() for token in forbidden_hints) or "run-initial-probes.sh" not in hint for hint in hints):
            raise SystemExit(f"FAILED: planner exposed a raw tool hint: {hints}")

        bad_registry = clone()
        bad_registry["schema_version"] = 99
        bad_path = temp_root / "planner-bad-registry.json"
        bad_path.write_text(json.dumps(bad_registry, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        bad_plan = subprocess.run(
            [sys.executable, str(planner), "--target-dir", str(target), "--workspace-dir", str(workspace), "--registry", str(bad_path), "--format", "json"],
            cwd=skill_root,
            capture_output=True,
            text=True,
        )
        if bad_plan.returncode == 0 or "Tool Registry validation failed closed" not in (bad_plan.stdout + bad_plan.stderr):
            raise SystemExit("FAILED: planner accepted an invalid registry override")

    print("TOOL REGISTRY SELFTEST PASSED: offline schema/authority/wrapper/planner matrix")


def exercise_context_planning_contract(skill_root: Path) -> None:
    """Exercise production advisory planning without opening references or running tools."""
    catalog = skill_root / "assets/context-catalog.json"
    catalog_validator = skill_root / "scripts/validate_context_catalog.py"
    planner = skill_root / "scripts/plan_audit_context.py"
    plan_validator = skill_root / "scripts/validate_context_plan.py"
    manifest = json.loads((skill_root / "assets/fixtures/context-planning/manifest.json").read_text(encoding="utf-8"))
    if manifest.get("execution_policy") != "offline-production-catalog-planner-validator-only":
        raise SystemExit("FAILED: context planning fixture is not offline-only")
    if not {"docker", "poc", "scanner", "network", "package_manager", "llm", "agent"}.issubset(set(manifest.get("forbidden_execution", []))):
        raise SystemExit("FAILED: context planning fixture omissions weaken offline boundary")
    run([sys.executable, str(catalog_validator), "--skill-root", str(skill_root), "--catalog", str(catalog), "--json"], skill_root)
    fixture_root = skill_root / "assets/fixtures/context-planning"
    cases = (("generic-recon", "recon", []), ("node-web-recon", "recon", []), ("python-library-candidate-generation", "candidate_generation", []), ("java-web-verification", "verification", []), ("ssrf-verification", "verification", ["ssrf"]))
    golden = manifest.get("golden_selection")
    if not isinstance(golden, dict) or set(golden) != {name for name, _, _ in cases}:
        raise SystemExit("FAILED: context planning golden selection fixture is incomplete")
    expected_negative_cases = manifest.get("negative_cases")
    if not isinstance(expected_negative_cases, list) or any(not isinstance(value, str) for value in expected_negative_cases) or len(expected_negative_cases) != len(set(expected_negative_cases)):
        raise SystemExit("FAILED: context planning negative-case manifest is not a unique label set")
    executed_negative_cases: set[str] = set()

    def record_negative(label: str) -> None:
        if label in executed_negative_cases:
            raise SystemExit(f"FAILED: context planning negative case ran twice: {label}")
        executed_negative_cases.add(label)

    with tempfile.TemporaryDirectory(prefix="zhulong-context-plan-") as tempdir:
        root = Path(tempdir)
        plans: dict[str, bytes] = {}
        for name, phase, bugs in cases:
            output = root / f"{name}.json"
            run([sys.executable, str(planner), "--target-dir", str(fixture_root / name), "--phase", phase, *sum((["--bug-class", item] for item in bugs), []), "--output", str(output), "--json"], skill_root)
            run([sys.executable, str(plan_validator), "--skill-root", str(skill_root), "--catalog", str(catalog), "--plan", str(output), "--json"], skill_root)
            plan = json.loads(output.read_text(encoding="utf-8")); plans[name] = output.read_bytes()
            if any(str(skill_root) in value for value in json.dumps(plan, ensure_ascii=False).split()):
                raise SystemExit("FAILED: context plan leaked a local absolute skill path")
            expected = golden[name]
            if not isinstance(expected, dict) or any(plan.get(group) != expected.get(group) for group in ("mandatory", "optional", "deferred")):
                raise SystemExit(f"FAILED: context planner selection drifted from static golden: {name}")
        repeat = root / "repeat.json"
        run_with_env([sys.executable, str(planner), "--target-dir", str(fixture_root / "node-web-recon"), "--phase", "recon", "--output", str(repeat), "--json"], skill_root, {"LC_ALL": "C", "TZ": "UTC"})
        if repeat.read_bytes() != plans["node-web-recon"]:
            raise SystemExit("FAILED: repeated context plan was not byte-identical")
        locale_plan = root / "locale.json"
        run_with_env([sys.executable, str(planner), "--target-dir", str(fixture_root / "node-web-recon"), "--phase", "recon", "--output", str(locale_plan), "--json"], skill_root, {"LC_ALL": "C", "TZ": "Asia/Shanghai"})
        if locale_plan.read_bytes() != plans["node-web-recon"]:
            raise SystemExit("FAILED: context plan changed with locale/timezone")
        record_negative("locale-timezone-determinism")
        canonical = json.loads(catalog.read_text(encoding="utf-8"))

        def expect_catalog_failure(label: str, value: dict, code: str, validation_root: Path = skill_root, record_case: bool = True) -> None:
            path = root / f"{label}.catalog.json"; path.write_text(json.dumps(value), encoding="utf-8")
            proc = subprocess.run([sys.executable, str(catalog_validator), "--skill-root", str(validation_root), "--catalog", str(path), "--json"], cwd=skill_root, capture_output=True, text=True)
            try:
                issue_codes = json.loads(proc.stdout).get("issue_codes", [])
            except json.JSONDecodeError as exc:
                raise SystemExit(f"FAILED: {label} catalog diagnostic was not JSON") from exc
            if proc.returncode == 0 or code not in issue_codes:
                raise SystemExit(f"FAILED: {label} did not reject {code}")
            if record_case:
                record_negative(label)

        def expect_plan_failure(label: str, value: dict, code: str) -> None:
            path = root / f"{label}.plan.json"; path.write_text(json.dumps(value), encoding="utf-8")
            proc = subprocess.run([sys.executable, str(plan_validator), "--skill-root", str(skill_root), "--catalog", str(catalog), "--plan", str(path), "--json"], cwd=skill_root, capture_output=True, text=True)
            try:
                issue_codes = json.loads(proc.stdout).get("issue_codes", [])
            except json.JSONDecodeError as exc:
                raise SystemExit(f"FAILED: {label} plan diagnostic was not JSON") from exc
            if proc.returncode == 0 or code not in issue_codes:
                raise SystemExit(f"FAILED: {label} did not reject {code}")
            record_negative(label)

        bad = json.loads(json.dumps(canonical)); bad["unknown"] = True; expect_catalog_failure("unknown-field", bad, "CONTEXT_CATALOG_SCHEMA_INVALID")
        bad = json.loads(json.dumps(canonical)); bad["modules"].append(dict(bad["modules"][0])); expect_catalog_failure("duplicate-module", bad, "CONTEXT_CATALOG_DUPLICATE_ID")
        for label, value in (("absolute-reference-path", "/tmp/context.md"), ("uri-reference-path", "file:///tmp/context.md"), ("traversal-reference-path", "assets/references/../context.md"), ("backslash-reference-path", "assets/references\\context.md")):
            bad = json.loads(json.dumps(canonical)); bad["modules"][0]["path"] = value; expect_catalog_failure(label, bad, "CONTEXT_REFERENCE_PATH_UNSAFE")
        bad = json.loads(json.dumps(canonical)); bad["modules"][0]["stacks"] = ["unknown"]; expect_catalog_failure("unknown-selector", bad, "CONTEXT_CATALOG_SCHEMA_INVALID")
        bad = json.loads(json.dumps(canonical)); bad["modules"][0]["phases"] = ["unknown"]; expect_catalog_failure("unknown-phase", bad, "CONTEXT_CATALOG_SCHEMA_INVALID")
        bad = json.loads(json.dumps(canonical)); bad["non_authority_statement"] = "This catalog does not grant read authority."; expect_catalog_failure("authority-drift", bad, "CONTEXT_CATALOG_AUTHORITY_INVALID", record_case=False)

        temporary_layout = root / "temporary-layout"
        shutil.copytree(skill_root / "assets", temporary_layout / "assets")
        reference_dir = temporary_layout / "assets/references"
        (reference_dir / "reference-symlink.md").symlink_to(reference_dir / "repo-preparation.md")
        (reference_dir / "reference-directory.md").mkdir()
        for label, reference, code in (("reference-symlink", "assets/references/reference-symlink.md", "CONTEXT_REFERENCE_SYMLINK"), ("missing-reference", "assets/references/missing-reference.md", "CONTEXT_REFERENCE_MISSING"), ("directory-reference", "assets/references/reference-directory.md", "CONTEXT_REFERENCE_TYPE_INVALID")):
            bad = json.loads(json.dumps(canonical)); bad["modules"][0]["path"] = reference; expect_catalog_failure(label, bad, code, temporary_layout)

        scope_cases = {
            "dogfood-reference": ["p8-bundle-generation-dogfood-report.md", "p8-real-historical-bundle-dogfood-report.md", "p8-real-historical-bundle-dogfood-metrics.json", "p9-protocol-chain-real-workspace-dogfood-report.md", "p9-protocol-chain-real-workspace-dogfood-metrics.json"],
            "template-reference": ["bundle-contract-template.json", "claude-code-invocation-template.md", "false-positive-template.md", "final-summary-template.md", "recon-result-template.json", "unverified-lead-template.md", "variant-seed-template.md"],
            "example-json-reference": ["docker-registry-fallbacks.example.json"],
        }
        for label, basenames in scope_cases.items():
            for basename in basenames:
                bad = json.loads(json.dumps(canonical)); bad["modules"][0]["path"] = f"assets/references/{basename}"
                path = root / f"{label}-{basename}.catalog.json"; path.write_text(json.dumps(bad), encoding="utf-8")
                proc = subprocess.run([sys.executable, str(catalog_validator), "--skill-root", str(skill_root), "--catalog", str(path), "--json"], cwd=skill_root, capture_output=True, text=True)
                payload = json.loads(proc.stdout)
                scope_issue = next((issue for issue in payload.get("issues", []) if issue.get("code") == "CONTEXT_REFERENCE_SCOPE_FORBIDDEN"), None)
                if proc.returncode == 0 or scope_issue is None or scope_issue.get("path") != "$.modules[0].path" or basename in scope_issue.get("message", ""):
                    raise SystemExit(f"FAILED: {label} accepted forbidden reference basename")
            record_negative(label)
        bad = json.loads(json.dumps(canonical)); bad["modules"][0]["path"] = "assets/references/dogfood-missing-report.md"
        path = root / "dogfood-scope-before-lstat.catalog.json"; path.write_text(json.dumps(bad), encoding="utf-8")
        proc = subprocess.run([sys.executable, str(catalog_validator), "--skill-root", str(skill_root), "--catalog", str(path), "--json"], cwd=skill_root, capture_output=True, text=True)
        if proc.returncode == 0 or "CONTEXT_REFERENCE_SCOPE_FORBIDDEN" not in json.loads(proc.stdout).get("issue_codes", []):
            raise SystemExit("FAILED: forbidden scope was not checked before reference lstat")
        for basename in ("attacker-container-pattern.md", "omc-runtime-stability.md", "output-language-and-path-contract.md"):
            bad = json.loads(json.dumps(canonical)); bad["modules"][0]["path"] = f"assets/references/{basename}"
            path = root / f"ordinary-{basename}.catalog.json"; path.write_text(json.dumps(bad), encoding="utf-8")
            proc = subprocess.run([sys.executable, str(catalog_validator), "--skill-root", str(skill_root), "--catalog", str(path), "--json"], cwd=skill_root, capture_output=True, text=True)
            if proc.returncode != 0 or "CONTEXT_REFERENCE_SCOPE_FORBIDDEN" in json.loads(proc.stdout).get("issue_codes", []):
                raise SystemExit("FAILED: ordinary unregistered reference was treated as forbidden scope")

        forged = json.loads(plans["node-web-recon"]); forged["optional"][0]["path"] = "assets/references/ssrf-checklist.md"; expect_plan_failure("forged-plan-module", forged, "CONTEXT_PLAN_SELECTION_INVALID")
        duplicate = json.loads(plans["node-web-recon"]); duplicate["mandatory"].append(duplicate["optional"][0]); expect_plan_failure("duplicate-plan-module", duplicate, "CONTEXT_PLAN_MODULE_DUPLICATE")
        missing = json.loads(plans["node-web-recon"]); missing["mandatory"].pop(); expect_plan_failure("missing-baseline", missing, "CONTEXT_PLAN_SELECTION_INVALID")
        reason_drift = json.loads(plans["node-web-recon"]); reason_drift["optional"][0]["reason_code"] = "WRONG_REASON"; expect_plan_failure("selector-reason-drift", reason_drift, "CONTEXT_PLAN_SELECTION_INVALID")
        noncanonical = json.loads(plans["node-web-recon"]); noncanonical["deferred"].reverse(); expect_plan_failure("noncanonical-order", noncanonical, "CONTEXT_PLAN_SELECTION_INVALID")
        catalog_drift = json.loads(plans["node-web-recon"]); catalog_drift["catalog"]["digest"] = "sha256:" + "0" * 64; expect_plan_failure("catalog-digest-drift", catalog_drift, "CONTEXT_PLAN_CATALOG_DRIFT")
        authority_drift = json.loads(plans["node-web-recon"]); authority_drift["authority"] = "loaded"; expect_plan_failure("authority-drift", authority_drift, "CONTEXT_PLAN_AUTHORITY_INVALID")
        unknown_claim = json.loads(plans["node-web-recon"]); unknown_claim["non_claims"].append("does prove a module was read"); expect_plan_failure("unknown-claim", unknown_claim, "CONTEXT_PLAN_NON_CLAIMS_INVALID")
        unknown_facts = json.loads(plans["node-web-recon"]); unknown_facts["input_facts"]["attack_surface_hints"].append("unrecognized-surface"); expect_plan_failure("unknown-plan-input-fact", unknown_facts, "CONTEXT_PLAN_INPUT_FACT_UNKNOWN")
        unsafe_output = root / "unsafe-output.json"; unsafe_output.symlink_to(root / "target")
        run_expect_fail([sys.executable, str(planner), "--target-dir", str(fixture_root / "generic-recon"), "--phase", "recon", "--output", str(unsafe_output)], skill_root, "output path must not be a symlink")
        record_negative("output-symlink")
    if set(expected_negative_cases) != executed_negative_cases:
        missing = sorted(set(expected_negative_cases) - executed_negative_cases)
        unexpected = sorted(executed_negative_cases - set(expected_negative_cases))
        raise SystemExit(f"FAILED: context planning manifest coverage mismatch: missing={missing}, unexpected={unexpected}")
    print("CONTEXT PLANNING SELFTEST PASSED: offline catalog/path/selection/determinism matrix")


def exercise_root_skill_kernel_contract(skill_root: Path) -> None:
    """Validate kernel -> phase reference -> production carrier relationships."""
    inventory_path = skill_root / "assets/root-skill-rule-inventory.json"
    inventory_schema = skill_root / "assets/schemas/root-skill-rule-inventory.schema.json"
    inventory_validator = skill_root / "scripts/validate_root_skill_rule_inventory.py"
    source_skill = skill_root / "skills/zhulong/SKILL.md"
    template_skill = skill_root / "templates/claude-skill/SKILL.md"
    if not source_skill.exists():
        source_skill = skill_root / "SKILL.md"
        template_skill = source_skill
    for path in (inventory_path, inventory_schema, inventory_validator):
        if not path.is_file() or path.is_symlink():
            raise SystemExit(f"FAILED: root Skill kernel contract file missing: {path.name}")
    run([
        sys.executable,
        str(inventory_validator),
        "--skill-root",
        str(skill_root),
        "--inventory",
        str(inventory_path),
        "--json",
    ], skill_root)

    skill_text = source_skill.read_text(encoding="utf-8")
    if source_skill.read_bytes() != template_skill.read_bytes():
        raise SystemExit("FAILED: source and template root Skills are not byte-identical")
    invariants = (
        "PoCs, exploit payloads, and verification traffic run only inside Docker",
        "Scanner, static, dependency, checklist, playbook, and LLM results are",
        "Confirmed requires a real attacker-controlled entrypoint",
        "Blocked verification is not `completed_no_confirmed_findings`",
        "Bind every claim to the exact tested source ref",
        "`rejected_unsafe_sandbox` never enters `confirmed/`",
        "Never use broad Docker prune",
        "Severity escalation and seeded variant discovery are separate required",
        "Each variant remains a candidate until its own Docker reproduction",
        "Final bundles use contract-first staging",
        "Only the canonical finalization gate and event establish completion",
        "Recording is an opt-in post-bundle gate",
        "Context plans, handoffs, checkpoints, and next-actions are advisory",
        "Confirmed bundles must not leak local absolute paths",
    )
    for invariant in invariants:
        if invariant not in skill_text:
            raise SystemExit(f"FAILED: root Skill kernel invariant missing: {invariant}")
    for obsolete_inventory_heading in ("## Installed Skill Runtime Contents", "## Standard Execution Order"):
        if obsolete_inventory_heading in skill_text:
            raise SystemExit("FAILED: root Skill still duplicates the full operational inventory")

    inventory = json.loads(inventory_path.read_text(encoding="utf-8"))
    catalog = json.loads((skill_root / "assets/context-catalog.json").read_text(encoding="utf-8"))
    catalog_paths = {item["path"] for item in catalog["modules"]}
    phase_paths = {
        "assets/references/audit-phase-intake-recon.md",
        "assets/references/audit-phase-candidate-triage.md",
        "assets/references/audit-phase-verification.md",
        "assets/references/audit-phase-variant-discovery.md",
        "assets/references/audit-phase-packaging-finalization.md",
        "assets/references/audit-phase-recording.md",
        "assets/references/audit-continuation-state.md",
    }
    if not phase_paths.issubset(catalog_paths):
        raise SystemExit("FAILED: one or more phase references are absent from the context catalog")
    for path in phase_paths:
        reference = skill_root / path
        if not reference.is_file() or reference.is_symlink():
            raise SystemExit(f"FAILED: phase reference is not a regular file: {path}")
        basename = reference.name
        if "dogfood" in basename or basename.endswith(("-template.md", "-template.json", ".example.json")):
            raise SystemExit(f"FAILED: phase reference triggers forbidden catalog scope: {basename}")
    catalog_modules = {item["path"]: item for item in catalog["modules"]}
    phases = {"intake", "recon", "candidate_generation", "verification", "severity_escalation", "packaging", "finalization", "variant_discovery"}
    for phase in phases:
        if not any(module["selection_policy"] == "baseline" and phase in module["phases"] and module["path"] in phase_paths for module in catalog["modules"]):
            raise SystemExit(f"FAILED: phase lacks a deterministic phase-reference baseline: {phase}")

    mutation_cases: list[tuple[str, callable]] = []

    def mutate(label: str, fn: callable) -> None:
        mutation_cases.append((label, fn))

    mutate("duplicate-id", lambda value: value["rules"].__setitem__(1, {**value["rules"][1], "rule_id": value["rules"][0]["rule_id"]}))
    mutate("empty-carriers", lambda value: value["rules"][0].__setitem__("carriers", []))
    mutate("unknown-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("type", "document"))
    mutate("absolute-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "/tmp/x"))
    mutate("uri-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "file:///tmp/x"))
    mutate("backslash-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "scripts\\x.py"))
    mutate("traversal-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "scripts/../x.py"))
    mutate("outside-allowlist", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "README.md"))
    mutate("missing-carrier", lambda value: value["rules"][0]["carriers"][0].__setitem__("path", "scripts/not-present.py"))
    mutate("wrong-symbol", lambda value: value["rules"][0]["carriers"][1].__setitem__("symbol", "NOT_A_REAL_SYMBOL"))
    mutate("production-symbol-missing", lambda value: value["rules"][0]["carriers"][1].pop("symbol"))
    mutate("retain-without-kernel", lambda value: value["rules"][0].__setitem__("carriers", value["rules"][0]["carriers"][1:]))
    mutate("retain-target-reference", lambda value: value["rules"][0]["target"].__setitem__("path", "assets/references/audit-phase-verification.md"))
    mutate("move-target-kernel", lambda value: value["rules"][14]["target"].__setitem__("path", "skills/zhulong/SKILL.md"))
    mutate("move-without-reference", lambda value: value["rules"][14].__setitem__("carriers", value["rules"][14]["carriers"][1:]))
    mutate("hard-move-reference-only", lambda value: (value["rules"][0].__setitem__("disposition", "move_to_reference"), value["rules"][0].__setitem__("target", {"path": "assets/references/audit-phase-verification.md", "section": "Working path"}), value["rules"][0].__setitem__("carriers", [{"type": "reference", "path": "assets/references/audit-phase-verification.md", "symbol": "Working path"}])))
    mutate("uncataloged-reference", lambda value: value["rules"][14]["target"].__setitem__("path", "assets/references/attacker-container-pattern.md"))
    mutate("unknown-rule-class", lambda value: value["rules"][0].__setitem__("rule_class", "advice"))
    mutate("unknown-disposition", lambda value: value["rules"][0].__setitem__("disposition", "drop"))
    mutate("unknown-rule-field", lambda value: value["rules"][0].__setitem__("authority", True))
    mutate("absolute-target", lambda value: value["rules"][14]["target"].__setitem__("path", "/tmp/ref.md"))
    mutate("empty-rules", lambda value: value.__setitem__("rules", []))

    with tempfile.TemporaryDirectory(prefix="zhulong-root-kernel-") as tempdir:
        temp_root = Path(tempdir)
        for label, fn in mutation_cases:
            mutated = json.loads(json.dumps(inventory))
            fn(mutated)
            mutation_path = temp_root / f"{label}.json"
            mutation_path.write_text(json.dumps(mutated), encoding="utf-8")
            proc = subprocess.run([
                sys.executable,
                str(inventory_validator),
                "--skill-root",
                str(skill_root),
                "--inventory",
                str(mutation_path),
                "--json",
            ], cwd=skill_root, capture_output=True, text=True)
            if proc.returncode == 0:
                raise SystemExit(f"FAILED: root Skill inventory mutation passed: {label}")
        removed = skill_text.replace(invariants[0], "removed invariant", 1)
        if removed == skill_text or invariants[0] in removed:
            raise SystemExit("FAILED: root invariant removal mutation was ineffective")
        layout = temp_root / "layout"
        required_paths = {"assets/context-catalog.json"}
        for rule in inventory["rules"]:
            required_paths.update(carrier["path"] for carrier in rule["carriers"])
        for relative in sorted(required_paths):
            source = skill_root / relative
            if relative == "skills/zhulong/SKILL.md" and not source.exists():
                source = skill_root / "SKILL.md"
            destination = layout / relative
            destination.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, destination)
        deleted_carrier = layout / "scripts/run_verification_case.sh"
        deleted_carrier.unlink()
        proc = subprocess.run([
            sys.executable,
            str(inventory_validator),
            "--skill-root",
            str(layout),
            "--inventory",
            str(inventory_path),
            "--json",
        ], cwd=skill_root, capture_output=True, text=True)
        if proc.returncode == 0:
            raise SystemExit("FAILED: deleted production carrier did not fail closed")
        shutil.copy2(skill_root / "scripts/run_verification_case.sh", deleted_carrier)
        layout_skill = layout / "skills/zhulong/SKILL.md"
        layout_skill.write_text(removed, encoding="utf-8")
        proc = subprocess.run([
            sys.executable,
            str(inventory_validator),
            "--skill-root",
            str(layout),
            "--inventory",
            str(inventory_path),
            "--json",
        ], cwd=skill_root, capture_output=True, text=True)
        if proc.returncode == 0:
            raise SystemExit("FAILED: removed root invariant did not fail closed")
    print(f"ROOT SKILL KERNEL SELFTEST PASSED: {len(mutation_cases) + 2} fail-closed mutations")


def exercise_workspace_tool_registry_snapshot(skill_root: Path, workspace: Path) -> None:
    expected_pairs = (
        (skill_root / "assets/tool-registry.json", workspace / "bin/tool-registry.json"),
        (skill_root / "assets/schemas/tool-registry.schema.json", workspace / "bin/tool-registry.schema.json"),
        (skill_root / "scripts/validate_tool_registry.py", workspace / "bin/validate_tool_registry.py"),
        (skill_root / "scripts/plan_security_toolchain.py", workspace / "bin/plan-security-toolchain.py"),
    )
    for source, copied in expected_pairs:
        if not copied.is_file() or source.read_bytes() != copied.read_bytes():
            raise SystemExit(f"FAILED: workspace Tool Registry snapshot drift: {source} != {copied}")
    run(
        [
            sys.executable,
            str(workspace / "bin/validate_tool_registry.py"),
            "--skill-root", str(workspace),
            "--registry", str(workspace / "bin/tool-registry.json"),
            "--schema", str(workspace / "bin/tool-registry.schema.json"),
            "--tool", "initial-probes-wrapper",
            "--stage", "recon",
            "--boundary", "workspace_write",
            "--effect", "workspace_evidence_write",
            "--json",
        ],
        workspace,
    )
    for adapter, marker in (
        (workspace / "scripts/run-initial-probes.sh", "zhulong-tool-contract: initial-probes-v1"),
        (workspace / "scripts/run-verification-case.sh", "zhulong-tool-contract: docker-verification-v1; timeout=mandatory; sandbox-preflight=mandatory"),
        (workspace / "scripts/check-sandbox-preflight.py", "zhulong-tool-contract: sandbox-preflight-v1"),
    ):
        require_text(adapter, marker, "workspace controlled-wrapper marker")
    outside_output = workspace.parent / "tool-registry-outside-output"
    rejected_output = subprocess.run(
        [
            "bash",
            str(workspace / "bin/run-initial-probes.sh"),
            "--repo-root", str(workspace.parent),
            "--workspace-dir", str(workspace),
            "--output-dir", str(outside_output),
        ],
        cwd=workspace,
        capture_output=True,
        text=True,
    )
    if rejected_output.returncode == 0 or "--output-dir must stay under" not in (rejected_output.stdout + rejected_output.stderr) or outside_output.exists():
        raise SystemExit("FAILED: initial-probe wrapper accepted or wrote an output path outside workspace evidence")


def load_validate_report_bundle_module(root: Path):
    module_path = root / "scripts/validate_report_bundle.py"
    spec = importlib.util.spec_from_file_location("zhulong_validate_report_bundle_selftest", module_path)
    if spec is None or spec.loader is None:
        raise SystemExit(f"FAILED: could not load validate_report_bundle.py from {module_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def exercise_replay_transcript_corpus(root: Path) -> None:
    corpus_dir = root / "assets/fixtures/replay-transcript-corpus"
    manifest_path = corpus_dir / "manifest.json"
    if not corpus_dir.is_dir():
        raise SystemExit(f"FAILED: missing replay transcript corpus directory: {corpus_dir}")
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise SystemExit(f"FAILED: replay transcript corpus manifest is invalid JSON: {exc}") from exc
    if manifest.get("schema_version") != 1:
        raise SystemExit("FAILED: replay transcript corpus manifest schema_version must be 1")
    samples = manifest.get("samples")
    if not isinstance(samples, list) or not samples:
        raise SystemExit("FAILED: replay transcript corpus manifest must contain samples")

    validator = load_validate_report_bundle_module(root)
    seen_ids: set[str] = set()
    positive_formats: set[str] = set()
    copied_with_provenance: dict[str, object] | None = None
    copied_without_provenance: dict[str, object] | None = None
    copied_without_provenance_path: Path | None = None

    for index, sample in enumerate(samples):
        if not isinstance(sample, dict):
            raise SystemExit(f"FAILED: replay transcript corpus samples[{index}] must be an object")
        sample_id = str(sample.get("id") or "").strip()
        if not sample_id:
            raise SystemExit(f"FAILED: replay transcript corpus samples[{index}].id must not be empty")
        if sample_id in seen_ids:
            raise SystemExit(f"FAILED: duplicate replay transcript corpus sample id: {sample_id}")
        seen_ids.add(sample_id)

        rel = str(sample.get("path") or "").strip()
        rel_path = Path(rel)
        if not rel or rel_path.is_absolute() or ".." in rel_path.parts or rel_path.suffix != ".log":
            raise SystemExit(f"FAILED: replay transcript corpus sample path must be corpus-relative .log: {sample_id}")
        sample_path = corpus_dir / rel_path
        if not sample_path.is_file():
            raise SystemExit(f"FAILED: missing replay transcript corpus sample: {sample_path}")

        text = sample_path.read_text(encoding="utf-8")
        classification = validator.classify_replay_transcript(text)
        expected_classification = sample.get("expected_classification")
        if classification.get("classification") != expected_classification:
            raise SystemExit(
                f"FAILED: replay transcript corpus classification mismatch for {sample_id}: "
                f"expected {expected_classification}, got {classification.get('classification')}"
            )
        issue_code = validator.REPLAY_CLASSIFICATION_ISSUE_CODES.get(classification.get("classification"))
        expected_issue_code = sample.get("expected_issue_code")
        if issue_code != expected_issue_code:
            raise SystemExit(
                f"FAILED: replay transcript corpus issue-code mismatch for {sample_id}: "
                f"expected {expected_issue_code}, got {issue_code}"
            )

        if expected_classification == "trusted_transcript":
            if "[command]" in text:
                positive_formats.add("bracket-command")
            if "docker compose" in text or "container logs:" in text:
                positive_formats.add("compose-service")
            if "Run command:" in text and "RAW OUTPUT:" in text:
                positive_formats.add("copied-run-command")
        if sample_id == "positive-copied-with-provenance-transcript":
            copied_with_provenance = sample
        if sample_id == "negative-copied-without-provenance":
            copied_without_provenance = sample
            copied_without_provenance_path = sample_path

    missing = sorted(REPLAY_TRANSCRIPT_CORPUS_REQUIRED_IDS - seen_ids)
    if missing:
        raise SystemExit(f"FAILED: replay transcript corpus missing required sample ids: {missing}")
    if len(positive_formats) < 3:
        raise SystemExit("FAILED: replay transcript corpus positive samples must use meaningfully different formats")
    if copied_with_provenance is None or copied_without_provenance is None or copied_without_provenance_path is None:
        raise SystemExit("FAILED: replay transcript corpus must include copied provenance boundary samples")

    with tempfile.TemporaryDirectory(prefix="zhulong-replay-corpus-") as tempdir:
        bundle = Path(tempdir)
        rel = "attachments/evidence/replay-output.log"
        log_path = bundle / rel
        log_path.parent.mkdir(parents=True, exist_ok=True)

        log_path.write_text((corpus_dir / str(copied_with_provenance["path"])).read_text(encoding="utf-8"), encoding="utf-8")
        manifest_payload = {
            "schema_version": 1,
            "replay_logs": [
                {
                    "path": rel,
                    "source_kind": copied_with_provenance.get("source_kind"),
                    "source_path": copied_with_provenance.get("source_path"),
                    "provenance": copied_with_provenance.get("provenance"),
                    "trust_classification": "trusted_transcript",
                }
            ],
        }
        (bundle / "bundle-build-manifest.json").write_text(
            json.dumps(manifest_payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        manifest_entries = validator.build_manifest_replay_entries(bundle)
        validator.validate_registered_replay_log_content(bundle, rel, manifest_entries.get(rel))

        log_path.write_text(copied_without_provenance_path.read_text(encoding="utf-8"), encoding="utf-8")
        manifest_payload["replay_logs"][0] = {
            "path": rel,
            "source_kind": copied_without_provenance.get("source_kind"),
            "trust_classification": "trusted_transcript",
        }
        (bundle / "bundle-build-manifest.json").write_text(
            json.dumps(manifest_payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        manifest_entries = validator.build_manifest_replay_entries(bundle)
        try:
            validator.validate_registered_replay_log_content(bundle, rel, manifest_entries.get(rel))
        except SystemExit as exc:
            if "lacks copied transcript provenance" not in str(exc):
                raise
        else:
            raise SystemExit("FAILED: copied replay transcript without provenance unexpectedly passed")

    docs = [
        root / "assets/references/reviewer-readiness-validator-gates.md",
        root / "assets/references/bundle-generation-checklist.md",
        root / "docs/RELEASE_CHECKLIST.md",
        root / "docs/WORKFLOW_DETAILS.md",
    ]
    for path in docs:
        require_text(path, "replay transcript corpus", f"replay transcript corpus docs pointer in {path.name}")
        require_text(path, "single rigid log format", f"no rigid replay format boundary in {path.name}")
    require_text(
        root / "docs/WORKFLOW_DETAILS.zh-CN.md",
        "复现记录样本集",
        "replay transcript corpus docs pointer in Chinese workflow",
    )
    require_text(
        root / "docs/WORKFLOW_DETAILS.zh-CN.md",
        "唯一且僵化的日志格式",
        "no rigid replay format boundary in Chinese workflow",
    )


def exercise_bundle_rule_mapping(root: Path) -> None:
    mapping_path = root / "assets/references/bundle-rule-mapping.md"
    if not mapping_path.is_file():
        raise SystemExit(f"FAILED: missing bundle rule mapping: {mapping_path}")
    content = mapping_path.read_text(encoding="utf-8")
    lowered = content.lower()
    normalized_lowered = re.sub(r"\s+", " ", lowered)
    for phrase in (
        "Contract preflight means `ready to render`",
        "generation readiness workflow gate",
        "does not prove a vulnerability",
        "Docker evidence",
        "final confirmed-bundle validation",
    ):
        if phrase.lower() not in normalized_lowered:
            raise SystemExit(f"FAILED: bundle rule mapping missing boundary wording: {phrase}")
    for field in BUNDLE_RULE_MAPPING_REQUIRED_FIELDS:
        if f"`{field}`" not in content:
            raise SystemExit(f"FAILED: bundle rule mapping missing contract field: {field}")
    for column in BUNDLE_RULE_MAPPING_COLUMNS:
        if column not in content:
            raise SystemExit(f"FAILED: bundle rule mapping missing column: {column}")
    for phrase in (
        "Any new contract field must be added to this mapping in the same change",
        "do not add it to the contract",
    ):
        if phrase not in content:
            raise SystemExit(f"FAILED: bundle rule mapping missing maintenance rule: {phrase}")
    for phrase in (
        "stable enum",
        "Critical",
        "Informational",
        "recommended bug classes",
        "free text",
        "Path Traversal",
        "Command Injection",
    ):
        if phrase not in content:
            raise SystemExit(f"FAILED: bundle rule mapping missing enum/recommended-values rationale: {phrase}")
    for phrase in BUNDLE_RULE_MAPPING_FORBIDDEN_CLAIMS:
        haystack = lowered if phrase.isascii() else content
        needle = phrase.lower() if phrase.isascii() else phrase
        if needle in haystack:
            raise SystemExit(f"FAILED: bundle rule mapping contains forbidden claim: {phrase}")


def exercise_reviewer_readiness_gate_classification(root: Path) -> None:
    reference = root / "assets/references/reviewer-readiness-validator-gates.md"
    for family in REVIEWER_READINESS_GATE_FAMILIES:
        require_text(reference, family, f"reviewer-readiness gate family {family}")
    for phrase in REVIEWER_READINESS_GATE_REQUIRED_TEXT:
        require_text(reference, phrase, f"reviewer-readiness gate classification text {phrase}")
    for phrase in REVIEWER_READINESS_GATE_FORBIDDEN_CLAIMS:
        forbid_text(reference, phrase, "reviewer-readiness gate non-confirmation boundary")

    validator = root / "scripts/validate_report_bundle.py"
    contract_validator = root / "scripts/validate_bundle_contract.py"
    release_checklist = root / "docs/RELEASE_CHECKLIST.md"

    require_text(contract_validator, "SSRF_IMPACT_OVERCLAIM", "SSRF contract issue code")
    require_text(validator, "SSRF_IMPACT_OVERCLAIM", "SSRF final validator collectable issue code")
    require_text(validator, "CODE_CONTEXT_MINIMUM_QUALITY", "code context minimum quality issue code")
    require_text(validator, "REPLAY_HELPER_PAUSE_CONTRACT", "replay helper pause contract issue code")
    require_text(validator, "REPLAY_HELPER_READINESS_PAUSE_SEPARATION", "replay helper readiness pause separation issue code")
    require_text(validator, "REPLAY_HELPER_ABSOLUTE_EVIDENCE_PATH", "replay helper bundle-relative evidence path issue code")
    require_text(validator, "ROOT_SCRIPT_CONTEXT_MISSING", "replay helper context issue code")
    require_text(validator, "report {heading} section contains placeholder-only code context", "code context placeholder rejection")
    require_text(validator, "must pause before proof command execution transitions", "replay pre-command pause rejection")
    require_text(validator, "must pause after proof command output transitions", "replay post-command pause rejection")
    require_text(validator, "Functional waits must use independent READY_/RETRY_/BACKOFF variables", "readiness pause separation rejection")
    require_text(validator, "bundle-relative replay log path", "absolute evidence path output rejection")

    # These fixtures are deterministic local file rewrites in the report-bundle
    # selftest section. They do not execute Docker, replay scripts, scanners,
    # package managers, network calls, or real target code.
    fixture_text = (root / "scripts/selftest_plugin.py").read_text(encoding="utf-8")
    fixture_pairs = {
        "SSRF gate positive": "good_ssrf_callback_bounded",
        "SSRF gate negative": "bad_ssrf_callback_overclaim",
        "Code context positive": "standard_bundle",
        "Code context negative": "bad_placeholder_code_context",
        "Replay pause positive": "standard_bundle",
        "Replay pause negative": "bad_missing_pre_command_pause",
        "Readiness pause separation positive": "READY_WAIT_SECONDS",
        "Readiness pause separation negative": "bad_readiness_pause_reuse",
        "Bundle-relative evidence path positive": "REPLAY_LOG_REL",
        "Bundle-relative evidence path negative": "bad_absolute_evidence_log_output",
    }
    for label, needle in fixture_pairs.items():
        if needle not in fixture_text:
            raise SystemExit(f"FAILED: missing reviewer-readiness fixture coverage for {label}: {needle}")

    require_text(
        release_checklist,
        "reviewer-readiness gate classification",
        "release checklist reviewer-readiness classification entry",
    )
    require_text(
        release_checklist,
        "positive and negative fixture coverage",
        "release checklist reviewer-readiness fixture entry",
    )


def exercise_p8_closure_contracts(root: Path) -> None:
    require_files(root, P8_RUNTIME_FILES, "P8 runtime")
    require_files(root, P8_REFERENCE_FILES, "P8 reference/schema")
    exercise_bundle_rule_mapping(root)
    exercise_reviewer_readiness_gate_classification(root)

    schema = json.loads((root / "assets/schemas/bundle-contract.schema.json").read_text(encoding="utf-8"))
    template = json.loads((root / "assets/references/bundle-contract-template.json").read_text(encoding="utf-8"))
    severity_schema = schema["properties"]["finding"]["properties"]["severity"]
    if severity_schema.get("enum") != STABLE_CONTRACT_SEVERITIES:
        raise SystemExit("FAILED: bundle contract schema must enforce stable finding.severity enum labels")
    if template.get("fixture_provenance") != {
        "required": False,
        "replay_type": "full_app",
        "synthetic_security_properties_present": False,
        "security_properties": [],
    }:
        raise SystemExit("FAILED: bundle contract template must retain the explicit full-app fixture security judgment")
    ssrf_template = template.get("impact_tier", {}).get("ssrf", {})
    if "artifact_backed_oracle" in ssrf_template:
        raise SystemExit("FAILED: callback-only bundle contract template must omit empty SSRF artifact_backed_oracle")

    if (root / "skills/zhulong/SKILL.md").exists():
        skill_path = root / "skills/zhulong/SKILL.md"
        template_path = root / "templates/claude-skill/SKILL.md"
        if skill_path.read_bytes() != template_path.read_bytes():
            raise SystemExit("FAILED: source skill and Claude skill template must be byte-identical")
    else:
        skill_path = root / "SKILL.md"

    require_text(skill_path, "## Confirmed bundle path", "root kernel confirmed bundle section")
    require_text(skill_path, "Final bundles use contract-first staging", "root kernel atomic bundle invariant")
    packaging_reference = root / "assets/references/audit-phase-packaging-finalization.md"
    require_text(packaging_reference, "confirmed/.staging/<slug>", "phase reference staging path")
    require_text(packaging_reference, "validate_all_report_bundles.py", "phase reference batch validation")
    require_text(packaging_reference, "finalization event", "phase reference canonical finalization")
    require_text(root / "scripts/build_confirmed_bundle.py", "atomic promote", "production bundle promotion carrier")
    require_text(root / "scripts/validate_report_bundle.py", "REPLAY_HELPER_ABSOLUTE_EVIDENCE_PATH", "production portability carrier")

    docs = [
        root / "docs/WORKFLOW_DETAILS.md",
        root / "docs/RELEASE_CHECKLIST.md",
        root / "assets/references/bundle-generation-checklist.md",
    ]
    for path in docs:
        require_text(path, "bundle contract", f"bundle short-path contract wording in {path.name}")
        require_text(path, "confirmed/.staging/<slug>", f"bundle staging path wording in {path.name}")
        require_text(path, "validate_all_report_bundles.py", f"bundle validate-all wording in {path.name}")
        require_text(path, "finalization", f"bundle finalization wording in {path.name}")
        require_text(path, "severity", f"bundle severity policy wording in {path.name}")
        require_text(path, "bug_class", f"bundle bug_class policy wording in {path.name}")

    chinese_workflow = root / "docs/WORKFLOW_DETAILS.zh-CN.md"
    require_text(chinese_workflow, "生成合同", "bundle short-path contract wording in Chinese workflow")
    require_text(chinese_workflow, "confirmed/.staging/<slug>", "bundle staging path wording in Chinese workflow")
    require_text(chinese_workflow, "validate_all_report_bundles.py", "bundle validate-all wording in Chinese workflow")
    require_text(chinese_workflow, "审计收尾", "bundle finalization wording in Chinese workflow")
    require_text(chinese_workflow, "finding.severity", "bundle severity policy wording in Chinese workflow")
    require_text(chinese_workflow, "bug_class", "bundle bug_class policy wording in Chinese workflow")

    for phrase in RECOMMENDED_BUG_CLASS_TEXT:
        require_text(
            root / "assets/references/bundle-generation-checklist.md",
            phrase,
            f"bundle checklist recommended bug_class wording {phrase}",
        )

    require_text(root / "docs/WORKFLOW_DETAILS.md", "does not prove a vulnerability", "bundle English preflight boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.md", "diagnostic mode", "bundle English all-errors diagnostic boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.md", "failed staging directory", "bundle English failed staging boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.md", "Marker-only replay", "bundle English marker-only replay boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.md", "stays candidate-only", "bundle English seeded variant boundary")

    require_text(root / "docs/WORKFLOW_DETAILS.zh-CN.md", "不能证明漏洞成立", "bundle Chinese preflight boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.zh-CN.md", "诊断模式", "bundle Chinese all-errors diagnostic boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.zh-CN.md", "不能称为已确认交付物", "bundle Chinese failed staging boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.zh-CN.md", "仅含标记的复现日志", "bundle Chinese marker-only replay boundary")
    require_text(root / "docs/WORKFLOW_DETAILS.zh-CN.md", "始终保持候选态", "bundle Chinese seeded variant boundary")

    forbidden_claims = [
        "contract preflight proves a vulnerability",
        "contract preflight confirms a vulnerability",
        "staging validation proves a vulnerability",
        "staging validation confirms a vulnerability",
        "all-errors proves a vulnerability",
        "all-errors confirms a vulnerability",
        "preflight 证明漏洞",
        "preflight 确认漏洞",
        "all-errors 证明漏洞",
        "all-errors 确认漏洞",
    ]
    for path in docs + [chinese_workflow, skill_path]:
        for phrase in forbidden_claims:
            forbid_text(path, phrase, f"P8 no diagnostic-confirmation claim in {path.name}")


def exercise_p8_dogfood_metrics(root: Path) -> None:
    require_files(root, P8_DOGFOOD_FILES, "P8 dogfood")
    report_path = root / "assets/references/p8-bundle-generation-dogfood-report.md"
    for heading in (
        "## Fixtures Used",
        "## Commands Represented",
        "## Old Retry-Loop Failure Mode",
        "## New P8 Flow Result",
        "## Metrics",
        "## Boundaries",
        "## P8 Closure State",
        "## Residual Risks",
    ):
        require_text(report_path, heading, f"P8 dogfood report heading {heading}")
    for phrase in (
        "does not execute Docker",
        "do not prove a real vulnerability",
        "replace Docker evidence",
        "Final bundle validation remains mandatory",
        "variant candidates as confirmed evidence",
    ):
        require_text(report_path, phrase, f"P8 dogfood report boundary {phrase}")

    with tempfile.TemporaryDirectory() as tempdir:
        temp_root = Path(tempdir)
        output = run_capture_with_env(
            [
                sys.executable,
                str(root / "scripts/p8_dogfood_metrics.py"),
                "--json",
            ],
            root,
            {"PYTHONPYCACHEPREFIX": str(temp_root / "pycache")},
        )
        metrics = json.loads(output)
        required = {
            "schema_version",
            "validator_invocation_count",
            "material_rewrite_count",
            "unique_error_count_per_invocation",
            "partial_confirmed_bundle_created",
            "manual_marker_patch_detected_or_required",
            "contract_preflight_caught_expected_issues",
            "staging_promote_required_for_final",
        }
        missing = sorted(required - set(metrics))
        if missing:
            raise SystemExit(f"FAILED: P8 dogfood metrics missing fields: {missing}")
        if metrics.get("schema_version") != 1:
            raise SystemExit("FAILED: P8 dogfood metrics schema_version mismatch")
        if metrics.get("validator_invocation_count", 0) <= 0:
            raise SystemExit("FAILED: P8 dogfood metrics did not record validator invocations")
        if metrics.get("material_rewrite_count") != 0:
            raise SystemExit("FAILED: P8 dogfood metrics should not require material rewrites")
        if metrics.get("partial_confirmed_bundle_created") is not False:
            raise SystemExit("FAILED: P8 dogfood staging failure created a final confirmed bundle")
        if metrics.get("manual_marker_patch_detected_or_required") is not False:
            raise SystemExit("FAILED: P8 dogfood required manual marker patching")
        if metrics.get("contract_preflight_caught_expected_issues") is not True:
            raise SystemExit("FAILED: P8 dogfood bad contract did not catch expected issues")
        if metrics.get("staging_promote_required_for_final") is not True:
            raise SystemExit("FAILED: P8 dogfood did not prove staging-before-final promotion")
        if max(metrics.get("unique_error_count_per_invocation") or [0]) < 3:
            raise SystemExit("FAILED: P8 dogfood did not return multiple errors in one invocation")

        cases = metrics.get("cases", {})
        bad_codes = set(cases.get("bad_contract", {}).get("issue_codes", []))
        for code in (
            "REPLAY_LOG_UNREGISTERED",
            "DIRECT_IMPACT_MARKER_DRIFT",
            "SSRF_IMPACT_OVERCLAIM",
            "FIXTURE_PROVENANCE_MISSING",
            "BUNDLE_PATH_ESCAPE",
            "FINAL_TARGET_EXISTS",
        ):
            if code not in bad_codes:
                raise SystemExit(f"FAILED: P8 dogfood bad contract missing issue code {code}")
        if cases.get("staging_build_failure", {}).get("final_bundle_created") is not False:
            raise SystemExit("FAILED: P8 dogfood staging failure promoted a final bundle")
        if cases.get("staging_build_failure", {}).get("failed_staging_preserved") is not True:
            raise SystemExit("FAILED: P8 dogfood did not preserve failed staging for diagnostics")
        marker_case = cases.get("marker_only_replay_log", {})
        if marker_case.get("rejected") is not True or "REPLAY_LOG_MARKER_ONLY" not in marker_case.get("issue_codes", []):
            raise SystemExit("FAILED: P8 dogfood marker-only replay log was not rejected")
        if marker_case.get("called_confirmed") is not False:
            raise SystemExit("FAILED: P8 dogfood marker-only case was described as confirmed")
        happy = cases.get("valid_contract_happy_path", {})
        if not (
            happy.get("contract_preflight_valid") is True
            and happy.get("promoted") is True
            and happy.get("batch_validation_passed") is True
        ):
            raise SystemExit("FAILED: P8 dogfood valid path did not reach promote and batch validation")
        comparison = metrics.get("comparison", {})
        if comparison.get("validator_invocation_count_delta", 0) <= 0:
            raise SystemExit("FAILED: P8 dogfood did not reduce validator invocation count")
        if comparison.get("material_rewrite_count_delta", 0) <= 0:
            raise SystemExit("FAILED: P8 dogfood did not reduce material rewrite count")

        generated_report = temp_root / "dogfood-report.md"
        generated_metrics = temp_root / "dogfood-metrics.json"
        run_with_env(
            [
                sys.executable,
                str(root / "scripts/p8_dogfood_metrics.py"),
                "--output-json",
                str(generated_metrics),
                "--output-report",
                str(generated_report),
            ],
            root,
            {"PYTHONPYCACHEPREFIX": str(temp_root / "pycache-output")},
        )
        generated = json.loads(generated_metrics.read_text(encoding="utf-8"))
        if generated.get("schema_version") != 1:
            raise SystemExit("FAILED: generated P8 dogfood JSON has wrong schema_version")
        require_text(generated_report, "## Old Retry-Loop Failure Mode", "generated P8 dogfood report")
        require_text(generated_report, "## Boundaries", "generated P8 dogfood report boundaries")


def exercise_p8_real_historical_dogfood(root: Path) -> None:
    require_files(root, P8_REAL_HISTORICAL_DOGFOOD_FILES, "P8 real historical dogfood")
    report_path = root / "assets/references/p8-real-historical-bundle-dogfood-report.md"
    metrics_path = root / "assets/references/p8-real-historical-bundle-dogfood-metrics.json"
    fixture_root = root / "assets/fixtures/p8-real-historical-bundle-dogfood"

    for heading in (
        "## Scope and Non-Claims",
        "## Sample Table",
        "## Metrics Table",
        "## Per-Sample Findings",
        "## Comparison With P8.6 Fixture Dogfood",
        "## Follow-Up",
    ):
        require_text(report_path, heading, f"P8 real historical dogfood report heading {heading}")
    for phrase in (
        "local historical dogfood",
        "did not execute Docker",
        "does not confirm new vulnerabilities",
        "not a production token-saving statistic",
        "P8.6 fixture measurement",
        "P8-post.4 historical dogfood measurement",
        "Production token savings: not measured and not claimed",
    ):
        require_text(report_path, phrase, f"P8 real historical dogfood boundary {phrase}")
    for phrase in (
        "Validator invocations",
        "Material rewrites during dogfood",
        "Partial confirmed bundle created",
        "historical-sample-01",
    ):
        require_text(report_path, phrase, f"P8 real historical dogfood metrics/sample text {phrase}")

    metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
    if metrics.get("schema_version") != 1:
        raise SystemExit("FAILED: P8 real historical dogfood metrics schema_version mismatch")
    if metrics.get("measurement_type") != "historical_dogfood":
        raise SystemExit("FAILED: P8 real historical dogfood metrics measurement_type mismatch")
    samples = metrics.get("samples")
    if not isinstance(samples, list) or not samples:
        raise SystemExit("FAILED: P8 real historical dogfood metrics must contain samples")
    sample_ids = {str(sample.get("sample_id")) for sample in samples if isinstance(sample, dict)}
    for expected in ("historical-sample-01", "historical-sample-02", "historical-sample-03"):
        if expected not in sample_ids:
            raise SystemExit(f"FAILED: P8 real historical dogfood metrics missing sample {expected}")
    if metrics.get("validator_invocation_count") != 4:
        raise SystemExit("FAILED: P8 real historical dogfood metrics validator invocation count mismatch")
    if metrics.get("material_rewrite_count") != 0:
        raise SystemExit("FAILED: P8 real historical dogfood should not rewrite materials")
    if metrics.get("partial_confirmed_bundle_count") != 1:
        raise SystemExit("FAILED: P8 real historical dogfood partial count mismatch")
    if metrics.get("measurement_type") == "fixture_dogfood":
        raise SystemExit("FAILED: P8 real historical dogfood metrics confused with fixture measurement")

    combined = "\n".join(
        path.read_text(encoding="utf-8", errors="ignore")
        for path in [report_path, metrics_path, *fixture_root.rglob("*")]
        if path.is_file()
    )
    forbidden_patterns = [
        r"/Users/",
        r"/home/",
        r"oss-vulnerability-" r"research",
        r"security-research-\d{8}",
        r"agent-studio",
        r"agent-runtime",
        r"nexent",
        r"spectral",
        r"linkding",
        r"yargs-parser",
    ]
    for pattern in forbidden_patterns:
        if re.search(pattern, combined, re.IGNORECASE):
            raise SystemExit(f"FAILED: P8 real historical dogfood leaked forbidden text pattern: {pattern}")
    forbidden_claims = [
        "confirmed a new vulnerability",
        "new vulnerability was confirmed",
        "production token savings were",
        "saved production tokens",
        "token-saving percentage",
    ]
    lowered = combined.lower()
    for phrase in forbidden_claims:
        if phrase in lowered:
            raise SystemExit(f"FAILED: P8 real historical dogfood contains forbidden claim: {phrase}")

    def run_json(command: list[str], *, expected_returncode: int) -> dict[str, object]:
        proc = subprocess.run(command, cwd=root, capture_output=True, text=True)
        if proc.returncode != expected_returncode:
            output = ((proc.stdout or "") + (proc.stderr or "")).strip()
            raise SystemExit(
                f"FAILED: P8 real historical dogfood command returned {proc.returncode}, "
                f"expected {expected_returncode}: {' '.join(command)}\n{output}"
            )
        try:
            return json.loads(proc.stdout)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: P8 real historical dogfood command did not emit JSON: {command}") from exc

    sample01 = fixture_root / "workspaces/historical-sample-01"
    sample01_payload = run_json(
        [
            sys.executable,
            str(root / "scripts/validate_bundle_contract.py"),
            "--workspace-dir",
            str(sample01),
            "--repo-root",
            str(sample01.parent),
            "--contract",
            str(sample01 / "confirmed/.contracts/historical-sample-01.bundle-contract.json"),
            "--all-errors",
            "--json",
        ],
        expected_returncode=1,
    )
    sample01_codes = {str(issue.get("code")) for issue in sample01_payload.get("issues", []) if isinstance(issue, dict)}
    for code in (
        "FINAL_TARGET_EXISTS",
        "REPLAY_LOG_UNREGISTERED",
        "DIRECT_IMPACT_MARKER_DRIFT",
        "CODE_CONTEXT_TOO_THIN",
        "SSRF_IMPACT_OVERCLAIM",
    ):
        if code not in sample01_codes:
            raise SystemExit(f"FAILED: P8 real historical sample 01 missing issue code {code}")

    validator = load_validate_report_bundle_module(root)
    placeholder = (sample01 / "placeholder-replay-output.log").read_text(encoding="utf-8")
    classification = validator.classify_replay_transcript(placeholder)
    issue_code = validator.REPLAY_CLASSIFICATION_ISSUE_CODES.get(classification.get("classification"))
    if classification.get("classification") != "placeholder_log" or issue_code != "REPLAY_LOG_PLACEHOLDER":
        raise SystemExit("FAILED: P8 real historical placeholder replay log was not rejected")

    sample02 = fixture_root / "workspaces/historical-sample-02"
    sample02_payload = run_json(
        [
            sys.executable,
            str(root / "scripts/validate_all_report_bundles.py"),
            "--confirmed-dir",
            str(sample02 / "confirmed"),
            "--language",
            "zh-CN",
            "--json",
        ],
        expected_returncode=1,
    )
    summary = sample02_payload.get("summary")
    if not isinstance(summary, dict) or summary.get("partial_confirmed_bundle") != 1:
        raise SystemExit("FAILED: P8 real historical sample 02 did not classify partial bundle")

    sample03 = fixture_root / "workspaces/historical-sample-03"
    sample03_payload = run_json(
        [
            sys.executable,
            str(root / "scripts/validate_bundle_contract.py"),
            "--workspace-dir",
            str(sample03),
            "--repo-root",
            str(sample03.parent),
            "--contract",
            str(sample03 / "confirmed/.contracts/historical-sample-03.bundle-contract.json"),
            "--all-errors",
            "--json",
        ],
        expected_returncode=1,
    )
    sample03_codes = {str(issue.get("code")) for issue in sample03_payload.get("issues", []) if isinstance(issue, dict)}
    if "REPLAY_LOG_UNREGISTERED" not in sample03_codes:
        raise SystemExit(f"FAILED: P8 real historical sample 03 lost replay registration rejection: {sorted(sample03_codes)}")


def exercise_p9_protocol_chain_real_workspace_dogfood(root: Path) -> None:
    report_path = root / "assets/references/p9-protocol-chain-real-workspace-dogfood-report.md"
    metrics_path = root / "assets/references/p9-protocol-chain-real-workspace-dogfood-metrics.json"
    require_files(root, [str(report_path.relative_to(root)), str(metrics_path.relative_to(root))], "P9.12.1 protocol-chain dogfood")

    report_text = report_path.read_text(encoding="utf-8")
    required_headings = (
        "## 结论",
        "## 边界与 Non-Claims",
        "## 历史 R1 样本保留",
        "## 真实 R2 Pilot：目标、Bootstrap 与协议链",
        "## Real-Copy CAS",
        "## Real-Copy Rebuild",
        "## Fresh-Context Agent B",
        "## Deterministic R2 Fixture Regression",
        "## Original Workspace Immutability",
        "## Metrics 与资格计算",
        "## 最终判定",
    )
    for heading in required_headings:
        if heading not in report_text:
            raise SystemExit(f"FAILED: P9.12.1 report heading missing: {heading}")
    for phrase in (
        "本次 P9.12.1 结果为 **passed**",
        "旧实现报告与旧 cross-audit 文档仍作为历史记录保留",
        "不把它们计入真实 R2 acceptance",
        "STATE_REVISION_CONFLICT",
        "journal before/after SHA-256 相同且 byte-identical",
        "correct 为 9，unknown 为 0，incorrect 为 0",
        "file mutation violations 为 0",
        "fixture regression 继续单独计量，永远不计入 real-workspace acceptance",
        "closure eligibility 为 `eligible_for_next_phase`",
    ):
        if phrase not in report_text:
            raise SystemExit(f"FAILED: P9.12.1 report result phrase missing: {phrase}")

    metrics = json.loads(metrics_path.read_text(encoding="utf-8"))

    def is_digest(value: object) -> bool:
        return isinstance(value, str) and bool(re.fullmatch(r"[0-9a-f]{64}", value))

    def evaluate_payload(doc: dict) -> bool:
        if (
            doc.get("schema_version") != 1
            or doc.get("measurement_type") != "real_workspace_protocol_dogfood"
            or doc.get("result") != "passed"
            or doc.get("blocker_statement") is not None
            or doc.get("closure_eligibility") != "eligible_for_next_phase"
        ):
            return False

        samples = doc.get("historical_r1_samples")
        expected_categories = {
            "sample-no-confirmed": "no_confirmed",
            "sample-blocked-verification": "blocked_verification",
            "sample-validated-bundle": "validated_confirmed_bundle",
        }
        if not isinstance(samples, list) or len(samples) != 3:
            return False
        actual_categories = {
            str(item.get("sample_id")): str(item.get("category"))
            for item in samples
            if isinstance(item, dict)
        }
        if actual_categories != expected_categories or any(
            not isinstance(item, dict)
            or item.get("protocol_mode") != "legacy_r1"
            or item.get("selection_gate_passed") is not True
            or item.get("counts_toward_real_r2_acceptance") is not False
            for item in samples
        ):
            return False

        aggregate = doc.get("aggregate")
        if not isinstance(aggregate, dict):
            return False
        historical_aggregate = {
            "historical_sample_count": len(samples),
            "historical_validator_invocation_count": sum(int(item.get("validator_invocation_count", 0)) for item in samples),
            "historical_derived_artifact_write_count": sum(int(item.get("derived_artifact_write_count", 0)) for item in samples),
            "historical_sanitization_count": sum(int(item.get("sanitization_count", 0)) for item in samples),
            "historical_contradiction_count": sum(int(item.get("contradiction_count", 0)) for item in samples),
        }
        if any(aggregate.get(key) != value for key, value in historical_aggregate.items()):
            return False

        pilot = doc.get("real_r2_pilot")
        contract = pilot.get("target_contract") if isinstance(pilot, dict) else None
        bootstrap = pilot.get("bootstrap") if isinstance(pilot, dict) else None
        derived = pilot.get("derived_artifacts") if isinstance(pilot, dict) else None
        if not all(isinstance(value, dict) for value in (pilot, contract, bootstrap, derived)):
            return False
        if not (
            pilot.get("origin") == "production_bootstrap_on_real_repository_copy"
            and pilot.get("protocol_mode") == "r2"
            and isinstance(pilot.get("tested_ref_digest"), str)
            and bool(re.fullmatch(r"[0-9a-f]{40}", pilot.get("tested_ref_digest")))
            and contract.get("repo_root") == "."
            and contract.get("tested_ref_matches_copy_head") is True
            and contract.get("validator_result") == "passed"
            and bootstrap.get("production_bootstrap_invoked") is True
            and bootstrap.get("initial_event_written") is True
            and bootstrap.get("observation_event_written") is True
            and bootstrap.get("event_count") == 2
            and bootstrap.get("state_revision") == 2
            and bootstrap.get("scope_is_empty") is True
            and bootstrap.get("authority_fact_count") == 0
            and bootstrap.get("verification_executed") is False
            and bootstrap.get("docker_resources_created") == 0
            and bootstrap.get("symlink_count") == 0
            and is_digest(bootstrap.get("baseline_manifest_sha256"))
            and derived.get("handoff_integrity") == "valid"
            and derived.get("tested_ref_verified") is True
            and derived.get("no_authority_facts") is True
            and pilot.get("original_copy_mutation_violations") == 0
            and pilot.get("counts_toward_real_r2_acceptance") is True
        ):
            return False

        cas = doc.get("real_copy_cas")
        if not isinstance(cas, dict) or not (
            cas.get("measurement_type") == "real_r2_copy_cas_conflict"
            and cas.get("origin") == "real_r2_pilot_copy"
            and cas.get("protocol_mode") == "r2"
            and cas.get("writer_process_count") == 2
            and cas.get("pids_distinct") is True
            and cas.get("success_count") == 1
            and cas.get("conflict_count") == 1
            and cas.get("conflict_code") == "STATE_REVISION_CONFLICT"
            and cas.get("loser_journal_committed") is False
            and cas.get("loser_state_view_updated") is False
            and cas.get("loser_authority_mutation_count") == 0
            and cas.get("authority_fact_delta") == 0
            and cas.get("counts_toward_real_r2_acceptance") is True
        ):
            return False

        rebuild = doc.get("real_copy_rebuild")
        if not isinstance(rebuild, dict) or not (
            rebuild.get("measurement_type") == "real_r2_copy_state_rebuild"
            and rebuild.get("origin") == "real_r2_pilot_copy"
            and rebuild.get("protocol_mode") == "r2"
            and rebuild.get("corruption_detected_by_read_only_check") is True
            and rebuild.get("check_result") == "drift_detected"
            and rebuild.get("check_issue_codes") == ["STATE_REVISION_MISMATCH"]
            and rebuild.get("check_was_read_only") is True
            and rebuild.get("apply_mode") == "explicit_digest_cas"
            and rebuild.get("apply_success") is True
            and is_digest(rebuild.get("journal_before_sha256"))
            and rebuild.get("journal_before_sha256") == rebuild.get("journal_after_sha256")
            and rebuild.get("journal_byte_identical") is True
            and is_digest(rebuild.get("rebuilt_state_sha256"))
            and rebuild.get("rebuilt_state_sha256") == rebuild.get("pilot_state_sha256")
            and rebuild.get("state_canonical_json_equivalent") is True
            and rebuild.get("journal_mutation_violations") == 0
            and rebuild.get("authority_fact_delta") == 0
            and rebuild.get("counts_toward_real_r2_acceptance") is True
        ):
            return False

        fresh = doc.get("fresh_context")
        receipt = fresh.get("receipt") if isinstance(fresh, dict) else None
        comparison = fresh.get("comparison") if isinstance(fresh, dict) else None
        agent_b = fresh.get("agent_b_observation") if isinstance(fresh, dict) else None
        agent_a = fresh.get("agent_a_expected") if isinstance(fresh, dict) else None
        if not all(isinstance(value, dict) for value in (fresh, receipt, comparison, agent_b, agent_a)):
            return False
        if not (
            fresh.get("method") == "platform_subagent_opaque_tool_identity"
            and fresh.get("mechanism") == "platform_subagent"
            and fresh.get("tool_version") == "multi_agent_v1"
            and fresh.get("fork_context") is False
            and fresh.get("attempt_count") == 1
            and is_digest(receipt.get("child_identity_sha256"))
            and receipt.get("parent_identity_sha256") is None
            and receipt.get("parent_identity_exposed") is False
            and receipt.get("parent_child_distinct") is True
            and receipt.get("pid_distinct") is None
            and receipt.get("resume") is False
            and is_digest(receipt.get("input_sha256"))
            and is_digest(receipt.get("allowed_input_digest"))
            and is_digest(receipt.get("observation_sha256"))
            and receipt.get("timestamps_present") is True
            and receipt.get("exit_status") == "completed"
            and receipt.get("raw_chat_saved") is False
            and receipt.get("hidden_reasoning_saved") is False
            and comparison.get("key_field_count") == 9
            and comparison.get("correct_field_count") == 9
            and comparison.get("unknown_field_count") == 0
            and comparison.get("incorrect_field_count") == 0
            and comparison.get("critical_incorrect_count") == 0
            and comparison.get("result") == "passed"
            and fresh.get("workspace_mutation_violations") == 0
            and fresh.get("counts_toward_real_r2_acceptance") is True
        ):
            return False
        observation_keys = {
            "protocol_mode",
            "target_contract_valid",
            "tested_ref_verified",
            "event_chain_valid",
            "handoff_state_valid",
            "checkpoint_valid",
            "next_actions_valid",
            "scope_is_empty",
            "authority_facts_absent",
        }
        if set(agent_b) != observation_keys or set(agent_a) != observation_keys or agent_b != agent_a:
            return False
        if agent_b.get("protocol_mode") != "r2" or any(
            agent_b.get(key) is not True for key in observation_keys - {"protocol_mode"}
        ):
            return False

        original = doc.get("original_workspaces")
        if not isinstance(original, dict) or not (
            original.get("manifest_root_count") == 4
            and original.get("before_after_file_manifest_equal") is True
            and original.get("before_after_symlink_manifest_equal") is True
            and original.get("mutation_violations") == 0
            and original.get("symlink_mutation_violations") == 0
        ):
            return False

        fixture = doc.get("deterministic_fixture_regression")
        if not isinstance(fixture, dict) or not (
            fixture.get("measurement_type") == "deterministic_r2_fixture_regression"
            and fixture.get("counts_toward_real_workspace_acceptance") is False
            and fixture.get("concurrent_writer_success_count") == 1
            and fixture.get("concurrent_writer_conflict_count") == 1
            and fixture.get("recovery_check_count") == 1
            and fixture.get("recovery_apply_count") == 1
            and fixture.get("journal_mutation_violations") == 0
            and fixture.get("rebuilt_state_valid") is True
        ):
            return False

        acceptance = doc.get("real_workspace_acceptance")
        if not isinstance(acceptance, dict):
            return False
        cas_satisfied = (
            cas.get("success_count") == 1
            and cas.get("conflict_count") == 1
            and cas.get("loser_journal_committed") is False
            and cas.get("pids_distinct") is True
        )
        rebuild_satisfied = (
            rebuild.get("apply_success") is True
            and rebuild.get("journal_byte_identical") is True
            and rebuild.get("state_canonical_json_equivalent") is True
            and rebuild.get("journal_mutation_violations") == 0
        )
        fresh_satisfied = (
            comparison.get("correct_field_count") == comparison.get("key_field_count")
            and comparison.get("unknown_field_count") == 0
            and comparison.get("incorrect_field_count") == 0
            and receipt.get("parent_child_distinct") is True
            and receipt.get("resume") is False
        )
        expected_acceptance = {
            "real_r2_workspace_count": int(pilot.get("counts_toward_real_r2_acceptance") is True),
            "cas_conflict_on_real_copy_satisfied": cas_satisfied,
            "state_rebuild_on_real_copy_satisfied": rebuild_satisfied,
            "fresh_context_isolation_satisfied": fresh_satisfied,
            "original_workspace_mutation_violations": original.get("mutation_violations"),
            "full_regression_passed": doc.get("p9_1_p9_11_regression_result") == "passed",
        }
        expected_acceptance["eligible_for_next_phase"] = (
            expected_acceptance["real_r2_workspace_count"] > 0
            and expected_acceptance["cas_conflict_on_real_copy_satisfied"] is True
            and expected_acceptance["state_rebuild_on_real_copy_satisfied"] is True
            and expected_acceptance["fresh_context_isolation_satisfied"] is True
            and expected_acceptance["original_workspace_mutation_violations"] == 0
            and expected_acceptance["full_regression_passed"] is True
        )
        if any(acceptance.get(key) != value for key, value in expected_acceptance.items()):
            return False
        if doc.get("p9_1_p9_11_regression_result") != "passed":
            return False
        if not {
            "does not measure token use",
            "does not prove or reconfirm a vulnerability",
            "fixture results do not satisfy real-workspace CAS or rebuild acceptance",
            "phase closure eligibility does not authorize execution of out-of-scope workloads",
        }.issubset(set(doc.get("non_claims") or [])):
            return False
        return True

    if not evaluate_payload(metrics):
        raise SystemExit("FAILED: P9.12.1 metrics are not a positive ledger-derived result")
    aggregate = metrics["aggregate"]
    derived_aggregate = {
        "real_r2_pilot_count": int(metrics["real_r2_pilot"]["counts_toward_real_r2_acceptance"] is True),
        "real_copy_cas_success_count": metrics["real_copy_cas"]["success_count"],
        "real_copy_cas_conflict_count": metrics["real_copy_cas"]["conflict_count"],
        "real_copy_rebuild_apply_count": int(metrics["real_copy_rebuild"]["apply_success"] is True),
        "fresh_context_attempt_count": metrics["fresh_context"]["attempt_count"],
        "fresh_context_correct_field_count": metrics["fresh_context"]["comparison"]["correct_field_count"],
        "fresh_context_unknown_field_count": metrics["fresh_context"]["comparison"]["unknown_field_count"],
        "fresh_context_incorrect_field_count": metrics["fresh_context"]["comparison"]["incorrect_field_count"],
        "fixture_count": 1,
        "original_workspace_mutation_violations": metrics["original_workspaces"]["mutation_violations"],
        "p9_1_p9_11_regression_passed": metrics["p9_1_p9_11_regression_result"] == "passed",
    }
    if any(aggregate.get(key) != value for key, value in derived_aggregate.items()):
        raise SystemExit("FAILED: P9.12.1 aggregate is not derived from detailed ledgers")

    # Fail-closed mutation checks stay in memory: this exercise is static-only.
    mutations = (
        ("pilot synthetic origin", ("real_r2_pilot", "origin"), "synthetic_mini_repo"),
        ("pilot fixture origin", ("real_r2_pilot", "origin"), "fixture"),
        ("CAS fixture origin", ("real_copy_cas", "origin"), "fixture"),
        ("rebuild synthetic origin", ("real_copy_rebuild", "origin"), "synthetic_mini_repo"),
        ("CAS measurement type", ("real_copy_cas", "measurement_type"), "synthetic_cas_conflict"),
        ("rebuild measurement type", ("real_copy_rebuild", "measurement_type"), "fixture_state_rebuild"),
        ("CAS loser journal flag", ("real_copy_cas", "loser_journal_committed"), True),
        ("rebuild journal identity", ("real_copy_rebuild", "journal_byte_identical"), False),
        ("fixture acceptance flag", ("deterministic_fixture_regression", "counts_toward_real_workspace_acceptance"), True),
        ("historical protocol", ("historical_r1_samples", 0, "protocol_mode"), "r2"),
        ("fresh comparison", ("fresh_context", "comparison", "incorrect_field_count"), 1),
    )
    for label, path, value in mutations:
        mutated = json.loads(json.dumps(metrics))
        cursor = mutated
        for component in path[:-1]:
            cursor = cursor[component]
        cursor[path[-1]] = value
        if evaluate_payload(mutated):
            raise SystemExit(f"FAILED: P9.12.1 fail-closed mutation accepted: {label}")

    # This checker is deliberately static: it must not launch workload or agent processes.
    import ast

    source = Path(__file__).read_text(encoding="utf-8")
    start = source.index("def exercise_p9_protocol_chain_real_workspace_dogfood")
    end = source.index("\ndef valid_target_contract_yaml", start) + 1
    function_tree = ast.parse(source[start:end])
    disallowed_calls = {
        "subprocess.run",
        "subprocess.Popen",
        "subprocess.check_call",
        "subprocess.check_output",
        "os.system",
        "socket.socket",
        "requests.get",
        "urllib.request.urlopen",
        "spawn_agent",
        "multi_agent_v1__spawn_agent",
    }
    for node in ast.walk(function_tree):
        if not isinstance(node, ast.Call):
            continue
        if isinstance(node.func, ast.Attribute) and isinstance(node.func.value, ast.Name):
            call_name = f"{node.func.value.id}.{node.func.attr}"
        elif isinstance(node.func, ast.Name):
            call_name = node.func.id
        else:
            call_name = ""
        if call_name in disallowed_calls:
            raise SystemExit(f"FAILED: P9.12.1 static checker invoked disallowed call: {call_name}")

    public_text = report_text + "\n" + metrics_path.read_text(encoding="utf-8")
    forbidden_patterns = (
        r"/Users/",
        r"/home/",
        r"security-research-\d{8}",
        r"app-platform",
        r"nexent",
        r"atomcode",
        r"fit-framework",
        r"agent-store",
        r"BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY",
        r"sk-[A-Za-z0-9]{16,}",
        r"(?:^|[^a-z])(?:session|child|parent)[ _-]?id\s*[:=]\s*[0-9a-f]{8}-[0-9a-f-]{27,}",
        r"docker (?:system|builder|buildx) prune",
        r"kill -(?:9|TERM|KILL)",
    )
    for pattern in forbidden_patterns:
        if re.search(pattern, public_text, re.IGNORECASE):
            raise SystemExit(f"FAILED: P9.12.1 public artifacts leaked forbidden pattern: {pattern}")
def valid_target_contract_yaml(*, runtime_type: str = "docker-compose", entrypoints: str | None = None) -> str:
    scope_entrypoints = entrypoints
    if scope_entrypoints is None:
        scope_entrypoints = """    - id: "import-url"
      kind: "http"
      route: "POST /api/import"
      auth: "low-privileged-user"
      attacker_control:
        - "url"
"""
    if runtime_type == "manual-blocked":
        return f"""schema_version: 1
target:
  name: "selftest-service"
  repo_root: "."
  tested_ref: "local-state"
  language_hint:
    - "python"
runtime:
  type: "manual-blocked"
verify:
  mode: "manual-required"
scope:
  entrypoints:
{scope_entrypoints if scope_entrypoints.strip() else "    []"}
  trust_boundaries:
    - "external request -> application boundary"
  in_scope_bug_classes:
    - "SSRF"
  out_of_scope:
    - "host shell access"
"""
    if runtime_type == "docker":
        runtime_block = """runtime:
  type: "docker"
  healthcheck:
    command: "curl -fsS http://127.0.0.1:8080/health"
    timeout_seconds: 30
build:
  command: "docker build -t zhulong-target-selftest ."
  network_required: false
start:
  command: "docker run --name zhulong-target-selftest -d -p 127.0.0.1:8080:8080 zhulong-target-selftest"
  readiness:
    command: "curl -fsS http://127.0.0.1:8080/health"
    timeout_seconds: 60
"""
    else:
        runtime_block = """runtime:
  type: "docker-compose"
  compose_file: "docker-compose.zhulong.yml"
  service: "app"
  healthcheck:
    command: "curl -fsS http://127.0.0.1:8080/health"
    timeout_seconds: 30
build:
  command: "docker compose -f docker-compose.zhulong.yml build"
  network_required: true
start:
  command: "docker compose -f docker-compose.zhulong.yml up -d"
  readiness:
    command: "curl -fsS http://127.0.0.1:8080/health"
    timeout_seconds: 60
"""
    cleanup = (
        "docker rm -f zhulong-target-selftest"
        if runtime_type == "docker"
        else "docker compose -f docker-compose.zhulong.yml down -v"
    )
    return f"""schema_version: 1
target:
  name: "selftest-service"
  repo_root: "."
  tested_ref: "local-state"
  language_hint:
    - "python"
    - "node"
{runtime_block}verify:
  mode: "fresh-container"
  allowed_network: "local-only"
  success_oracles:
    - type: "http_response"
    - type: "log_pattern"
  cleanup:
    command: "{cleanup}"
scope:
  entrypoints:
{scope_entrypoints if scope_entrypoints.strip() else "    []"}
  trust_boundaries:
    - "external request -> application boundary"
  in_scope_bug_classes:
    - "SSRF"
  out_of_scope:
    - "host shell access"
"""


def exercise_target_contract_validator(plugin_root: Path) -> None:
    validator = plugin_root / "scripts/validate_target_contract.py"
    example = plugin_root / "assets/examples/zhulong-target.example.yaml"
    schema = json.loads((plugin_root / "assets/schemas/zhulong-target.schema.json").read_text(encoding="utf-8"))
    if schema.get("title") != "Zhulong Target Contract R1":
        raise SystemExit("FAILED: zhulong-target schema title mismatch")
    if str(schema.get("$id", "")).startswith("https://zhulong.local/"):
        raise SystemExit("FAILED: zhulong-target schema must not use a placeholder $id URI")
    if '"additionalProperties": true' in json.dumps(schema, sort_keys=True):
        raise SystemExit("FAILED: zhulong-target schema must not leave known contract objects open to arbitrary properties")

    output = run_capture([sys.executable, str(validator), str(example)], plugin_root)
    if "OK: zhulong-target valid" not in output or "runtime_type=docker-compose" not in output:
        raise SystemExit(f"FAILED: example target contract did not validate as docker-compose:\n{output}")

    with tempfile.TemporaryDirectory() as tempdir:
        tmp = Path(tempdir)

        def write_case(name: str, content: str) -> Path:
            path = tmp / f"{name}.yaml"
            path.write_text(content, encoding="utf-8")
            return path

        compose_path = write_case("valid-compose", valid_target_contract_yaml())
        docker_path = write_case("valid-docker", valid_target_contract_yaml(runtime_type="docker"))
        manual_path = write_case("valid-manual", valid_target_contract_yaml(runtime_type="manual-blocked"))
        empty_entrypoints_path = write_case(
            "empty-entrypoints",
            valid_target_contract_yaml(entrypoints="    []"),
        )

        if "runtime_type=docker-compose" not in run_capture([sys.executable, str(validator), str(compose_path)], plugin_root):
            raise SystemExit("FAILED: valid docker-compose target contract did not pass")
        if "runtime_type=docker" not in run_capture([sys.executable, str(validator), str(docker_path)], plugin_root):
            raise SystemExit("FAILED: valid docker target contract did not pass")
        manual_output = run_capture([sys.executable, str(validator), str(manual_path)], plugin_root)
        if "runtime_type=manual-blocked" not in manual_output or "non_confirmable=true" not in manual_output:
            raise SystemExit(f"FAILED: manual-blocked target contract did not report non-confirmable:\n{manual_output}")
        empty_output = run_capture([sys.executable, str(validator), str(empty_entrypoints_path)], plugin_root)
        if "recon_incomplete=true" not in empty_output:
            raise SystemExit(f"FAILED: empty entrypoints did not mark recon incomplete:\n{empty_output}")

        missing_field = valid_target_contract_yaml().replace('  name: "selftest-service"\n', "")
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("missing-field", missing_field))],
            plugin_root,
            "missing required string: $.target.name",
        )

        invalid_runtime = valid_target_contract_yaml().replace('type: "docker-compose"', 'type: "podman"')
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("invalid-runtime", invalid_runtime))],
            plugin_root,
            "$.runtime.type must be one of",
        )

        no_compose_file = valid_target_contract_yaml().replace('  compose_file: "docker-compose.zhulong.yml"\n', "")
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("missing-compose-file", no_compose_file))],
            plugin_root,
            "missing required string: $.runtime.compose_file",
        )
        no_service = valid_target_contract_yaml().replace('  service: "app"\n', "")
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("missing-compose-service", no_service))],
            plugin_root,
            "missing required string: $.runtime.service",
        )

        absolute_path = valid_target_contract_yaml().replace('repo_root: "."', 'repo_root: "/Users/example/service"')
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("absolute-path", absolute_path))],
            plugin_root,
            "operator-local absolute path",
        )

        for index, unsafe_ref in enumerate(("/Users/alice/source", "（/Users/alice/secret.txt）", "`/Users/alice/secret.txt`", "ghp_SELFTEST_TOKEN", "file:///private/ref", "sha256:\nunsafe"), start=1):
            unsafe_tested_ref = valid_target_contract_yaml().replace('tested_ref: "local-state"', f"tested_ref: {json.dumps(unsafe_ref)}")
            run_expect_fail(
                [sys.executable, str(validator), str(write_case(f"unsafe-tested-ref-{index}", unsafe_tested_ref))],
                plugin_root,
                "forbidden source-identity material",
            )

        traversal_path = valid_target_contract_yaml().replace(
            'compose_file: "docker-compose.zhulong.yml"',
            'compose_file: "../docker-compose.yml"',
        )
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("parent-traversal", traversal_path))],
            plugin_root,
            "parent path traversal",
        )

        broad_prune = valid_target_contract_yaml().replace(
            "docker compose -f docker-compose.zhulong.yml down -v",
            "docker " + "system " + "prune -af",
        )
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("broad-prune", broad_prune))],
            plugin_root,
            "broad Docker prune",
        )

        dangerous_kill = valid_target_contract_yaml().replace(
            "docker compose -f docker-compose.zhulong.yml down -v",
            "kill " + "-9 1234",
        )
        run_expect_fail(
            [sys.executable, str(validator), str(write_case("dangerous-kill", dangerous_kill))],
            plugin_root,
            "dangerous PID kill",
        )

        unsafe_cases = {
            "privileged": "docker run --privileged example",
            "host-network": "docker run --network host example",
            "docker-socket": "docker run -v ./docker.sock:/sock example",
            "credential-mount": "docker run -v .npmrc:/workspace/.npmrc example",
        }
        for name, command in unsafe_cases.items():
            content = valid_target_contract_yaml(runtime_type="docker").replace(
                "docker run --name zhulong-target-selftest -d -p 127.0.0.1:8080:8080 zhulong-target-selftest",
                command,
            )
            run_expect_fail(
                [sys.executable, str(validator), str(write_case(f"unsafe-{name}", content))],
                plugin_root,
                "must not request privileged",
            )


def exercise_recon_result_contract(plugin_root: Path) -> None:
    fixture_root = plugin_root / "assets/fixtures/recon-result"
    manifest_path = fixture_root / "manifest.json"
    schema_path = plugin_root / "assets/schemas/recon-result.schema.json"
    template_path = plugin_root / "assets/references/recon-result-template.json"
    validator = plugin_root / "scripts/validate_recon_result.py"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    schema = json.loads(schema_path.read_text(encoding="utf-8"))
    template = json.loads(template_path.read_text(encoding="utf-8"))

    if schema.get("title") != "Zhulong Recon Coverage Contract R1":
        raise SystemExit("FAILED: Recon result schema title mismatch")
    if schema.get("$schema") != "https://json-schema.org/draft/2020-12/schema":
        raise SystemExit("FAILED: Recon result schema must declare Draft 2020-12")
    if not str(schema.get("$id", "")).startswith("https://github.com/"):
        raise SystemExit("FAILED: Recon result schema must use a repository-backed $id")

    def assert_closed_objects(value, path: str = "$") -> None:
        if isinstance(value, dict):
            if value.get("type") == "object" and value.get("additionalProperties") is not False:
                raise SystemExit(f"FAILED: Recon result schema leaves object open: {path}")
            for key, child in value.items():
                assert_closed_objects(child, f"{path}.{key}")
        elif isinstance(value, list):
            for index, child in enumerate(value):
                assert_closed_objects(child, f"{path}[{index}]")

    assert_closed_objects(schema)
    for field in (
        "schema_version",
        "recon_id",
        "target_binding",
        "attack_surface_binding",
        "coverage",
        "coverage_gaps",
        "unresolved_blockers",
        "focus_refs",
    ):
        if field not in template or template[field] is None:
            raise SystemExit(f"FAILED: Recon result template lacks representative field: {field}")
    if template.get("status") != "partial":
        raise SystemExit("FAILED: Recon result template must demonstrate partial coverage")
    if manifest.get("schema_version") != 1:
        raise SystemExit("FAILED: Recon result fixture manifest schema_version mismatch")

    help_proc = subprocess.run(
        [sys.executable, str(validator), "--help"],
        cwd=plugin_root,
        env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
        capture_output=True,
        text=True,
    )
    help_output = (help_proc.stdout or "") + (help_proc.stderr or "")
    if help_proc.returncode != 0 or "--recon-result" not in help_output or "--json" not in help_output:
        raise SystemExit(f"FAILED: Recon result validator CLI help is incomplete:\n{help_output}")

    def clone(value):
        return json.loads(json.dumps(value))

    def path_parts(expression: str) -> list[str | int]:
        parts: list[str | int] = []
        for component in expression.split("."):
            if not component:
                raise SystemExit(f"FAILED: malformed fixture mutation path: {expression}")
            if component.isdigit():
                parts.append(int(component))
            else:
                parts.append(component)
        return parts

    def get_at(value, expression: str):
        current = value
        for part in path_parts(expression):
            current = current[part]
        return current

    def parent_at(value, expression: str):
        parts = path_parts(expression)
        if not parts:
            raise SystemExit(f"FAILED: empty fixture mutation path: {expression}")
        current = value
        for part in parts[:-1]:
            current = current[part]
        return current, parts[-1]

    def set_at(value, expression: str, replacement) -> None:
        parent, leaf = parent_at(value, expression)
        if isinstance(parent, dict) and isinstance(leaf, str):
            parent[leaf] = clone(replacement)
        elif isinstance(parent, list) and isinstance(leaf, int):
            parent[leaf] = clone(replacement)
        else:
            raise SystemExit(f"FAILED: invalid fixture mutation target: {expression}")

    def drop_at(value, expression: str) -> None:
        parent, leaf = parent_at(value, expression)
        if isinstance(parent, dict) and isinstance(leaf, str):
            del parent[leaf]
        elif isinstance(parent, list) and isinstance(leaf, int):
            parent.pop(leaf)
        else:
            raise SystemExit(f"FAILED: invalid fixture drop target: {expression}")

    def snapshot_tree(root: Path) -> dict[str, bytes]:
        snapshot: dict[str, bytes] = {}
        for base, directories, files in os.walk(root, followlinks=False):
            for name in [*directories, *files]:
                path = Path(base) / name
                rel = path.relative_to(root).as_posix()
                if path.is_symlink():
                    snapshot[rel] = b"SYMLINK:" + os.readlink(path).encode("utf-8")
                elif path.is_file():
                    snapshot[rel] = path.read_bytes()
        return snapshot

    positive_cases = manifest.get("positive_cases", [])
    positive_by_id: dict[str, dict] = {}
    positive_case_meta: dict[str, dict] = {}
    for case in positive_cases:
        case_id = case["id"]
        scenario_root = fixture_root / case["scenario"]
        result_path = scenario_root / case["result"]
        positive_by_id[case_id] = json.loads(result_path.read_text(encoding="utf-8"))
        positive_case_meta[case_id] = case

    def run_validator_case(
        repo_root: Path,
        workspace_dir: Path,
        result_path: Path,
        *,
        expect_ok: bool,
        expected_codes: list[str],
        label: str,
    ) -> dict:
        before_repo = snapshot_tree(repo_root)
        before_workspace = snapshot_tree(workspace_dir)
        with tempfile.TemporaryDirectory(prefix="zhulong-recon-cache-") as cache_dir:
            env = {
                **os.environ,
                "PYTHONDONTWRITEBYTECODE": "1",
                "PYTHONPYCACHEPREFIX": cache_dir,
            }
            proc = subprocess.run(
                [
                    sys.executable,
                    str(validator),
                    "--repo-root",
                    str(repo_root),
                    "--workspace-dir",
                    str(workspace_dir),
                    "--recon-result",
                    str(result_path),
                    "--json",
                ],
                cwd=plugin_root,
                env=env,
                capture_output=True,
                text=True,
            )
        output = (proc.stdout or "") + (proc.stderr or "")
        for local_path in (repo_root.resolve(), workspace_dir.resolve()):
            if str(local_path) in output:
                raise SystemExit(f"FAILED: Recon validator leaked a local path for {label}: {output}")
        try:
            payload = json.loads(proc.stdout)
        except (TypeError, json.JSONDecodeError) as exc:
            raise SystemExit(f"FAILED: Recon validator did not return JSON for {label}: {output}") from exc
        if proc.returncode == 0 and not expect_ok:
            raise SystemExit(f"FAILED: negative Recon fixture unexpectedly passed: {label}\n{output}")
        if proc.returncode != 0 and expect_ok:
            raise SystemExit(f"FAILED: positive Recon fixture failed: {label}\n{output}")
        if bool(payload.get("ok")) != expect_ok:
            raise SystemExit(f"FAILED: Recon validator ok flag mismatch for {label}: {payload}")
        codes = set(payload.get("issue_codes", []))
        missing_codes = [code for code in expected_codes if code not in codes]
        if missing_codes:
            raise SystemExit(
                f"FAILED: Recon validator missed expected issue code(s) for {label}: {missing_codes}\n{output}"
            )
        after_repo = snapshot_tree(repo_root)
        after_workspace = snapshot_tree(workspace_dir)
        if before_repo != after_repo or before_workspace != after_workspace:
            raise SystemExit(f"FAILED: Recon validator mutated source or workspace material for {label}")
        return payload

    for case in positive_cases:
        scenario_root = fixture_root / case["scenario"]
        run_validator_case(
            scenario_root / "repo",
            scenario_root / "workspace",
            scenario_root / case["result"],
            expect_ok=True,
            expected_codes=[],
            label=case["id"],
        )

    for case in manifest.get("negative_cases", []):
        base_id = case["base"]
        if base_id not in positive_by_id:
            raise SystemExit(f"FAILED: negative Recon fixture references unknown base case: {base_id}")
        base_meta = positive_case_meta[base_id]
        with tempfile.TemporaryDirectory(prefix=f"zhulong-recon-{case['id']}-") as tempdir:
            temp_root = Path(tempdir)
            repo_root = temp_root / "repo"
            workspace_dir = temp_root / "workspace"
            scenario_root = fixture_root / base_meta["scenario"]
            shutil.copytree(scenario_root / "repo", repo_root)
            shutil.copytree(scenario_root / "workspace", workspace_dir)
            data = clone(positive_by_id[base_id])
            symlink_kind: str | None = None
            for mutation in case.get("mutations", []):
                kind = mutation["kind"]
                if kind == "drop":
                    drop_at(data, mutation["path"])
                elif kind == "set":
                    set_at(data, mutation["path"], mutation.get("value"))
                elif kind == "copy":
                    source_case = mutation["from_case"]
                    set_at(data, mutation["path"], get_at(positive_by_id[source_case], mutation["from_path"]))
                elif kind in {
                    "symlink_result",
                    "symlink_target",
                    "symlink_attack_surface",
                    "symlink_source",
                    "symlink_evidence",
                }:
                    symlink_kind = kind
                elif kind == "natural_language_only":
                    for category in (
                        "technology_stack",
                        "public_entrypoints",
                        "trust_boundaries",
                        "high_risk_sinks",
                        "security_policy_explanations",
                        "default_deployment_assumptions",
                        "priority_areas",
                        "deferred_areas",
                    ):
                        data[category] = []
                        data["coverage"][category]["item_ids"] = []
                        data["coverage"][category]["reason"] = "未发现问题"
                    data["focus_refs"] = []
                else:
                    raise SystemExit(f"FAILED: unknown Recon fixture mutation kind: {kind}")

            outside_dir = temp_root / "outside"
            outside_dir.mkdir()

            def make_escape_symlink(path: Path, name: str, content: bytes) -> None:
                outside_path = outside_dir / name
                outside_path.write_bytes(content)
                if path.exists() or path.is_symlink():
                    path.unlink()
                path.symlink_to(outside_path)

            result_path = workspace_dir / "recon-result.json"
            if symlink_kind == "symlink_result":
                make_escape_symlink(result_path, "result.json", json.dumps(data).encode("utf-8"))
            else:
                result_path.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
                if symlink_kind == "symlink_target":
                    target_path = workspace_dir / "zhulong-target.yaml"
                    make_escape_symlink(target_path, "target.yaml", target_path.read_bytes())
                elif symlink_kind == "symlink_attack_surface":
                    surface_path = workspace_dir / "attack-surface.md"
                    make_escape_symlink(surface_path, "attack-surface.md", surface_path.read_bytes())
                elif symlink_kind == "symlink_source":
                    source_path = repo_root / "src/app.py"
                    make_escape_symlink(source_path, "app.py", source_path.read_bytes())
                elif symlink_kind == "symlink_evidence":
                    evidence_path = workspace_dir / "evidence/recon-notes.md"
                    make_escape_symlink(evidence_path, "recon-notes.md", evidence_path.read_bytes())

            run_validator_case(
                repo_root,
                workspace_dir,
                result_path,
                expect_ok=False,
                expected_codes=case["expected_codes"],
                label=case["id"],
            )

    with tempfile.TemporaryDirectory(prefix="zhulong-recon-partial-blocker-") as tempdir:
        temp_root = Path(tempdir)
        repo_root = temp_root / "repo"
        workspace_dir = temp_root / "workspace"
        scenario_root = fixture_root / positive_case_meta["complete-service"]["scenario"]
        shutil.copytree(scenario_root / "repo", repo_root)
        shutil.copytree(scenario_root / "workspace", workspace_dir)
        data = clone(positive_by_id["complete-service"])
        data["status"] = "partial"
        data["unresolved_blockers"] = clone(positive_by_id["blocked-service"]["unresolved_blockers"])
        result_path = workspace_dir / "recon-result.json"
        result_path.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
        run_validator_case(
            repo_root,
            workspace_dir,
            result_path,
            expect_ok=True,
            expected_codes=[],
            label="partial-with-blocker-only",
        )

    print("RECON RESULT SELFTEST PASSED: schema/CLI/status/reference/read-only matrix")


def valid_candidate_contract(overrides: dict | None = None) -> dict:
    candidate = {
        "schema_version": 1,
        "candidate_id": "CAND-0001",
        "title": "SSRF through import URL fetch",
        "bug_class": "SSRF",
        "status": "candidate",
        "target_ref": {
            "target_config": "zhulong-target.yaml",
            "tested_ref": "git-sha-or-local-state",
        },
        "entrypoint": {
            "id": "import-url",
            "kind": "http",
            "route": "POST /api/import",
        },
        "attacker_model": {
            "required_auth": "low_privileged_user",
            "attacker_controls": ["url"],
            "environment_assumptions": [
                "import endpoint enabled",
                "backend can make outbound HTTP requests",
            ],
        },
        "claim": {
            "source": "user-controlled import URL",
            "sink": "server-side HTTP fetch",
            "missing_constraint": "private IP or metadata address denylist is absent or bypassable",
            "impact": "internal service probing or metadata fetch",
        },
        "poc": {
            "kind": "script",
            "path": "poc/reproduce.sh",
            "expected_oracle": {
                "type": "callback_observed",
                "description": "attacker-controlled callback server receives request from target container",
            },
        },
        "evidence": {
            "static_locations": [
                {
                    "path": "src/importer.py",
                    "start_line": 42,
                    "end_line": 67,
                    "reason": "URL is fetched server-side without network range validation",
                }
            ],
            "dynamic_evidence": [],
        },
        "finder": {
            "source": "agent",
            "created_at": "2026-06-18T00:00:00Z",
        },
    }
    if overrides:
        candidate.update(overrides)
    return candidate


def valid_verifier_verdict(overrides: dict | None = None) -> dict:
    verdict = {
        "schema_version": 1,
        "candidate_id": "CAND-0001",
        "verdict": "confirmed_in_docker",
        "verification_status": "confirmed_in_docker",
        "evidence_level": "entrypoint_reproduced",
        "target_ref": {
            "target_config": "zhulong-target.yaml",
            "tested_ref": "git-sha-or-local-state",
        },
        "environment": {
            "fresh_container": True,
            "runtime_type": "docker-compose",
            "host_network": False,
            "privileged": False,
            "docker_socket_mounted": False,
            "credential_paths_mounted": False,
            "egress_policy": "local-only",
        },
        "commands": [
            {
                "name": "poc",
                "command": "bash candidates/CAND-0001/poc/reproduce.sh",
                "exit_code": 0,
            }
        ],
        "oracle_result": {
            "type": "callback_observed",
            "success": True,
            "summary": "Callback server received request from target container.",
        },
        "attacker_entrypoint": {
            "id": "import-url",
            "kind": "http",
            "route": "POST /api/import",
            "input_shape": "JSON body field url controlled by a low-privilege user.",
            "entrypoint_to_sink_path": "POST /api/import parses url and reaches the server-side HTTP fetch helper.",
            "deterministic_impact_oracle": "Callback server receives request from target container.",
        },
        "replay_material": {
            "description": "Reviewer-facing verifier replay log.",
            "path": "verifier/CAND-0001/logs/poc.log",
        },
        "disposition_recommendation": "confirmed_in_docker",
        "negative_checks": [
            {
                "check": "auth prerequisite matches candidate",
                "passed": True,
            }
        ],
        "artifacts": ["verifier/CAND-0001/logs/poc.log"],
        "verified_at": "2026-06-18T00:00:00Z",
    }
    if overrides:
        verdict.update(overrides)
    return verdict


def json_clone(value: dict) -> dict:
    return json.loads(json.dumps(value))


def write_json_fixture(path: Path, value: dict) -> Path:
    path.write_text(json.dumps(value, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
    return path


def valid_bundle_contract(overrides: dict | None = None) -> dict:
    contract = {
        "schema_version": 1,
        "bundle": {
            "slug": "example-ssrf-callback",
            "language": "zh-CN",
            "final_path": "confirmed/example-ssrf-callback",
            "one_vulnerability_only": True,
            "fail_if_final_path_exists": True,
        },
        "render": {
            "source_findings_json": "confirmed/findings.json",
            "finding_slug": "example-ssrf-callback",
        },
        "finding": {
            "project_name": "example-project",
            "vulnerability_name": "SSRF callback reachability",
            "bug_class": "SSRF",
            "severity": "Medium",
            "attacker_condition": "Authenticated low-privilege user controls the import URL.",
            "server_condition": "Default import feature is enabled and outbound requests are allowed.",
            "security_impact": (
                "Docker evidence confirms callback reachability only; response, configuration, "
                "credential, and sensitive-data exposure are not claimed."
            ),
        },
        "docker_evidence": {
            "verification_status": "confirmed_in_docker",
            "docker_required": True,
            "docker_command": "docker compose -f attachments/poc/docker-compose.zhulong.yml up --abort-on-container-exit",
            "oracle_token": "DIRECT_IMPACT_CONFIRMED",
            "expected_observation": "Listener receives one request from the target container.",
            "observed_observation": "Replay log contains DIRECT_IMPACT_CONFIRMED.",
            "severity_escalation_attempted": True,
        },
        "entrypoint_evidence": {
            "evidence_level": "entrypoint_reproduced",
            "attacker_controlled_entrypoint": "POST /api/import",
            "input_shape": "JSON body field url controlled by a low-privilege user.",
            "entrypoint_to_sink_path": "POST /api/import parses url and reaches the server-side HTTP fetch helper.",
            "deterministic_impact_oracle": "Callback listener receives a request and replay log contains DIRECT_IMPACT_CONFIRMED.",
            "replay_material": {
                "description": "Bundle-root helper and replay log reproduce the attacker-entrypoint path.",
                "path": "attachments/evidence/replay-output.log",
            },
        },
        "replay": {
            "root_script": {"path": "run-example-ssrf-callback-recording.sh"},
            "log": {
                "path": "attachments/evidence/replay-output.log",
                "registration_targets": ["files.evidence_files", "files.reviewer_evidence_index"],
            },
        },
        "direct_impact": {
            "marker": "DIRECT_IMPACT_CONFIRMED",
            "sync_targets": [
                {"target": "replay.root_script", "marker": "DIRECT_IMPACT_CONFIRMED"},
                {"target": "replay.log", "marker": "DIRECT_IMPACT_CONFIRMED"},
                {"target": "files.verification_evidence", "marker": "DIRECT_IMPACT_CONFIRMED"},
                {"target": "files.reviewer_evidence_index", "marker": "DIRECT_IMPACT_CONFIRMED"},
                {"target": "reviewer_material", "marker": "DIRECT_IMPACT_CONFIRMED"},
            ],
        },
        "files": {
            "verification_evidence": "verification-evidence.json",
            "reviewer_evidence_index": "attachments/reviewer-evidence-index.json",
            "evidence_files": [
                "attachments/evidence/replay-output.log",
                "attachments/poc/reproduce.py",
            ],
            "attachments": [
                "attachments/poc/docker-compose.zhulong.yml",
                "attachments/poc/reproduce.py",
                "attachments/evidence/replay-output.log",
                "attachments/reviewer-evidence-index.json",
            ],
        },
        "code_context": {
            "entries": [
                {
                    "source_path": "src/importer.py",
                    "line_range": "42-67",
                    "input_to_sink_chain": "The user-controlled import URL reaches the server-side HTTP client.",
                    "missing_guard": "No resolved-IP or private-range validation runs before the fetch.",
                    "verified_impact_boundary": "Docker evidence proves callback reachability only.",
                }
            ]
        },
        "source_binding": {
            "tested_ref": "SELFTEST_REF",
            "attacker_entrypoint": "POST /api/import",
            "replay_observed_entrypoint": "POST /api/import",
            "binding_mode": "exact",
            "target_entrypoint_id": "import-url",
            "source_defined_entrypoint": "POST /api/import",
            "materials": {
                "target_config": "target/zhulong-target.yaml",
                "verifier_verdict": "verifier/verifier-verdict.json",
            },
            "source_references": [
                {
                    "id": "SRC-ENTRY",
                    "role": "entrypoint",
                    "path": "src/importer.py",
                    "start_line": 1,
                    "end_line": 1,
                    "hash_kind": "snippet",
                    "sha256": "0" * 64,
                    "exact_token": "POST /api/import",
                },
                {
                    "id": "SRC-SINK",
                    "role": "sink",
                    "path": "src/importer.py",
                    "start_line": 2,
                    "end_line": 2,
                    "hash_kind": "snippet",
                    "sha256": "0" * 64,
                    "exact_token": "http_fetch(url)",
                },
                {
                    "id": "SRC-CONFIG",
                    "role": "prerequisite",
                    "path": "src/importer.py",
                    "start_line": 3,
                    "end_line": 3,
                    "hash_kind": "snippet",
                    "sha256": "0" * 64,
                    "exact_token": "OUTBOUND_REQUESTS_ENABLED = True",
                },
            ],
        },
        "fixture_provenance": {
            "required": False,
            "replay_type": "full_app",
            "synthetic_security_properties_present": False,
            "security_properties": [],
        },
        "deployment_prerequisites": [
            {
                "id": "DEP-OUTBOUND",
                "description": "The tested deployment enables outbound requests for the import feature.",
                "source_reference_ids": ["SRC-CONFIG"],
                "reviewer_material_required": True,
            }
        ],
        "impact_claims": [
            {
                "id": "IMPACT-CALLBACK",
                "category": "network_reachability",
                "statement": "The tested attacker entrypoint causes one server-side outbound callback.",
                "deterministic_oracle": {
                    "token": "DIRECT_IMPACT_CONFIRMED",
                    "evidence_path": "attachments/evidence/replay-output.log",
                },
                "source_bound_prerequisite_ids": ["SRC-ENTRY", "SRC-SINK"],
                "depends_on_security_property_ids": [],
                "verified_deployment_prerequisite_ids": ["DEP-OUTBOUND"],
                "supported_bug_classes": ["SSRF"],
                "severity_ceiling": "Medium",
                "unsupported_stronger_impacts": ["Response content and sensitive data exposure are not proven."],
            }
        ],
        "validity_review": {
            "validity_verdict": "conditionally_confirmed",
            "classification_decision": "unchanged",
            "original_bug_class": "SSRF",
            "original_severity": "Medium",
            "final_bug_class": "SSRF",
            "final_cwe": "CWE-918",
            "final_severity": "Medium",
            "supported_impact_claim_ids": ["IMPACT-CALLBACK"],
            "deployment_prerequisite_ids": ["DEP-OUTBOUND"],
            "rationale": "Source and replay bind the same attacker entrypoint and prove callback reachability only.",
            "stronger_impacts_not_claimed": ["Response content and sensitive data exposure are not proven."],
        },
        "impact_tier": {
            "bug_class": "SSRF",
            "ssrf": {
                "tier": "callback_reachability",
                "claimed_exposures": ["callback_reachability"],
                "stronger_impacts_not_claimed": [
                    "response_content_exposure",
                    "configuration_exposure",
                    "credential_exposure",
                    "sensitive_data_exposure",
                ],
            },
        },
        "variant_seed_readiness": {
            "run_after_promote": True,
        },
    }
    if overrides:
        for key, value in overrides.items():
            contract[key] = value
    return contract


def prepare_source_bound_fixture(repo_root: Path, workspace: Path, contract: dict) -> dict:
    source_path = repo_root / "src/importer.py"
    source_path.parent.mkdir(parents=True, exist_ok=True)
    source_text = (
        'ENTRYPOINT = "POST /api/import"\n'
        "result = http_fetch(url)\n"
        "OUTBOUND_REQUESTS_ENABLED = True\n"
    )
    source_path.write_text(source_text, encoding="utf-8")
    if not (repo_root / ".git").exists():
        run(["git", "init", "-q", str(repo_root)], repo_root.parent)
        run(["git", "-C", str(repo_root), "add", "src/importer.py"], repo_root)
        run(
            [
                "git", "-C", str(repo_root),
                "-c", "user.name=Zhulong Selftest",
                "-c", "user.email=selftest@example.invalid",
                "commit", "-q", "-m", "source-bound selftest fixture",
            ],
            repo_root,
        )
    tested_ref = run_capture(["git", "-C", str(repo_root), "rev-parse", "HEAD"], repo_root).strip()
    binding = contract["source_binding"]
    binding["tested_ref"] = tested_ref
    lines = source_text.splitlines(keepends=True)
    for ref in binding["source_references"]:
        start = int(ref["start_line"])
        end = int(ref["end_line"])
        snippet = "".join(lines[start - 1:end]).encode("utf-8")
        ref["sha256"] = hashlib.sha256(snippet).hexdigest()
    target_path = workspace / binding["materials"]["target_config"]
    target_path.parent.mkdir(parents=True, exist_ok=True)
    target_path.write_text(
        "schema_version: 1\n"
        "target:\n"
        "  name: example-project\n"
        "  repo_root: .\n"
        f"  tested_ref: {tested_ref}\n"
        "  language_hint: [python]\n"
        "runtime:\n"
        "  type: docker\n"
        "verify:\n"
        "  mode: local\n"
        "scope:\n"
        "  entrypoints:\n"
        "    - id: import-url\n"
        "      kind: http\n"
        "      route: POST /api/import\n"
        "  trust_boundaries: []\n"
        "  in_scope_bug_classes: [SSRF]\n"
        "  out_of_scope: []\n",
        encoding="utf-8",
    )
    verdict_path = workspace / binding["materials"]["verifier_verdict"]
    verdict_path.parent.mkdir(parents=True, exist_ok=True)
    verdict_path.write_text(
        json.dumps(
            {
                "target_ref": {"target_config": binding["materials"]["target_config"], "tested_ref": tested_ref},
                "attacker_entrypoint": {"route": "POST /api/import"},
            },
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )
    return contract


def run_bundle_contract_json(
    plugin_root: Path,
    workspace: Path,
    contract_path: Path,
    *,
    expected_returncode: int,
    repo_root: Path | None = None,
) -> dict:
    validator = plugin_root / "scripts/validate_bundle_contract.py"
    output = run_capture_with_env(
        [
            sys.executable,
            str(validator),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_root or workspace.parent),
            "--contract",
            str(contract_path),
            "--all-errors",
            "--json",
        ],
        plugin_root,
        {"PYTHONPYCACHEPREFIX": str(workspace / "pycache")},
        expected_returncode=expected_returncode,
    )
    return json.loads(output)


def require_bundle_issue(payload: dict, code: str) -> None:
    codes = {issue.get("code") for issue in payload.get("issues", [])}
    if code not in codes:
        raise SystemExit(f"FAILED: expected bundle contract issue code {code}; got {sorted(codes)}")


def run_report_bundle_all_errors_json(
    plugin_root: Path,
    bundle: Path,
    *,
    language: str = "zh-CN",
    expected_returncode: int = 0,
    output_errors: Path | None = None,
) -> dict:
    command = [
        sys.executable,
        str(plugin_root / "scripts/validate_report_bundle.py"),
        "--bundle-dir",
        str(bundle),
        "--language",
        language,
        "--all-errors",
        "--json",
    ]
    if output_errors is not None:
        command.extend(["--output-errors", str(output_errors)])
    merged_env = {**os.environ, "PYTHONPYCACHEPREFIX": str(bundle.parent.parent / ".pycache-selftest")}
    proc = subprocess.run(
        command,
        cwd=plugin_root,
        env=merged_env,
        capture_output=True,
        text=True,
    )
    stdout = (proc.stdout or "").strip()
    stderr = (proc.stderr or "").strip()
    if proc.returncode != expected_returncode:
        raise SystemExit(
            f"FAILED: {' '.join(command)}\n"
            f"Expected exit code {expected_returncode}, got {proc.returncode}\n"
            f"stdout:\n{stdout}\nstderr:\n{stderr}"
        )
    if output_errors is not None:
        if not output_errors.is_file():
            raise SystemExit("FAILED: --output-errors did not write a JSON file")
        return json.loads(output_errors.read_text(encoding="utf-8"))
    try:
        return json.loads(stdout)
    except json.JSONDecodeError as exc:
        raise SystemExit(
            "FAILED: --all-errors --json stdout was not JSON\n"
            f"stdout prefix={stdout[:500]!r}\n"
            f"stderr prefix={stderr[:500]!r}"
        ) from exc


def require_report_issue(payload: dict, code: str) -> None:
    codes = {issue.get("code") for issue in payload.get("issues", [])}
    if code not in codes:
        raise SystemExit(
            f"FAILED: expected report bundle issue code {code}; got {sorted(codes)}\n"
            + json.dumps(payload.get("issues", []), ensure_ascii=False, indent=2)
        )


def exercise_bundle_contract_validator(plugin_root: Path) -> None:
    validator = plugin_root / "scripts/validate_bundle_contract.py"
    template = plugin_root / "assets/references/bundle-contract-template.json"
    schema = json.loads((plugin_root / "assets/schemas/bundle-contract.schema.json").read_text(encoding="utf-8"))
    if schema.get("title") != "Zhulong Bundle Contract R1":
        raise SystemExit("FAILED: bundle contract schema title mismatch")
    if str(schema.get("$id", "")).startswith("https://zhulong.local/"):
        raise SystemExit("FAILED: bundle contract schema must not use a placeholder $id URI")
    if '"additionalProperties": true' in json.dumps(schema, sort_keys=True):
        raise SystemExit("FAILED: bundle contract schema must not leave known contract objects open to arbitrary properties")
    json.loads(template.read_text(encoding="utf-8"))

    with tempfile.TemporaryDirectory() as tempdir:
        root = Path(tempdir)

        def new_workspace(name: str) -> Path:
            workspace = root / name
            (workspace / "confirmed" / ".contracts").mkdir(parents=True)
            return workspace

        def write_contract(workspace: Path, name: str, contract: dict) -> Path:
            prepare_source_bound_fixture(root, workspace, contract)
            return write_json_fixture(workspace / "confirmed" / ".contracts" / f"{name}.bundle-contract.json", contract)

        good_workspace = new_workspace("good")
        good_contract = write_contract(good_workspace, "good", valid_bundle_contract())
        good_payload = run_bundle_contract_json(plugin_root, good_workspace, good_contract, expected_returncode=0)
        if not good_payload.get("valid") or good_payload.get("issues"):
            raise SystemExit(f"FAILED: valid minimal bundle contract did not pass:\n{good_payload}")

        callback_workspace = new_workspace("callback")
        callback_contract = write_contract(callback_workspace, "callback", valid_bundle_contract())
        callback_payload = run_bundle_contract_json(plugin_root, callback_workspace, callback_contract, expected_returncode=0)
        if not callback_payload.get("valid"):
            raise SystemExit("FAILED: bounded SSRF callback/reachability contract did not pass")

        entrypoint_mismatch = json_clone(valid_bundle_contract())
        entrypoint_mismatch["source_binding"]["replay_observed_entrypoint"] = "POST /api/other"
        workspace = new_workspace("source-entrypoint-mismatch")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", entrypoint_mismatch), expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_ENTRYPOINT_MISMATCH")

        composed_mismatch = json_clone(valid_bundle_contract())
        composed_mismatch["source_binding"].update(
            {
                "binding_mode": "composed",
                "component_joiner": " ",
                "resolved_entrypoint": "POST /api/missing",
                "components": [
                    {"id": "METHOD", "value": "POST", "source_reference_ids": ["SRC-ENTRY"]},
                    {"id": "ROUTE", "value": "/api/missing", "source_reference_ids": ["SRC-ENTRY"]},
                ],
            }
        )
        composed_mismatch["source_binding"].pop("source_defined_entrypoint", None)
        workspace = new_workspace("composed-mismatch")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", composed_mismatch), expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_REF_MISMATCH")
        require_bundle_issue(payload, "SOURCE_ENTRYPOINT_MISMATCH")

        bad_hash = json_clone(valid_bundle_contract())
        workspace = new_workspace("source-hash-mismatch")
        bad_hash_path = write_contract(workspace, "bad", bad_hash)
        bad_hash_doc = json.loads(bad_hash_path.read_text(encoding="utf-8"))
        bad_hash_doc["source_binding"]["source_references"][0]["sha256"] = "f" * 64
        write_json_fixture(bad_hash_path, bad_hash_doc)
        payload = run_bundle_contract_json(plugin_root, workspace, bad_hash_path, expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_FILE_MISMATCH")

        dirty_source = json_clone(valid_bundle_contract())
        workspace = new_workspace("source-differs-from-tested-ref")
        dirty_source_path = write_contract(workspace, "bad", dirty_source)
        with (root / "src/importer.py").open("a", encoding="utf-8") as handle:
            handle.write("# uncommitted source change outside every hashed snippet\n")
        payload = run_bundle_contract_json(plugin_root, workspace, dirty_source_path, expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_FILE_MISMATCH")

        bad_line = json_clone(valid_bundle_contract())
        bad_line["source_binding"]["source_references"][0]["end_line"] = 999
        workspace = new_workspace("source-line-mismatch")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", bad_line), expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_REF_MISMATCH")

        symlink_escape = json_clone(valid_bundle_contract())
        symlink_escape["source_binding"]["source_references"][0]["path"] = "src/escape.py"
        workspace = new_workspace("source-symlink-escape")
        symlink_contract = write_contract(workspace, "bad", symlink_escape)
        outside_source = root.parent / f"{root.name}-outside.py"
        outside_source.write_text('ENTRYPOINT = "POST /api/import"\n', encoding="utf-8")
        (root / "src/escape.py").symlink_to(outside_source)
        payload = run_bundle_contract_json(plugin_root, workspace, symlink_contract, expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_FILE_MISMATCH")
        (root / "src/escape.py").unlink()
        outside_source.unlink()

        bad_ref = json_clone(valid_bundle_contract())
        workspace = new_workspace("tested-ref-mismatch")
        bad_ref_path = write_contract(workspace, "bad", bad_ref)
        bad_ref_doc = json.loads(bad_ref_path.read_text(encoding="utf-8"))
        bad_ref_doc["source_binding"]["tested_ref"] = "0" * 40
        write_json_fixture(bad_ref_path, bad_ref_doc)
        payload = run_bundle_contract_json(plugin_root, workspace, bad_ref_path, expected_returncode=1)
        require_bundle_issue(payload, "SOURCE_REF_MISMATCH")

        synthetic_privilege = json_clone(valid_bundle_contract())
        synthetic_privilege["fixture_provenance"].update(
            {
                "synthetic_security_properties_present": True,
                "security_properties": [
                    {
                        "id": "PROP-ADMIN",
                        "origin": "synthetic",
                        "use": "impact_prerequisite",
                        "meaning": "The fixture creates an administrator identity.",
                        "source_reference_ids": [],
                        "cannot_support_impact_claim_ids": [],
                    }
                ],
            }
        )
        synthetic_privilege["impact_claims"][0]["depends_on_security_property_ids"] = ["PROP-ADMIN"]
        workspace = new_workspace("synthetic-privilege")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", synthetic_privilege), expected_returncode=1)
        require_bundle_issue(payload, "SYNTHETIC_PROPERTY_SUPPORTS_IMPACT")

        oracle_drift = json_clone(valid_bundle_contract())
        oracle_drift["impact_claims"][0]["deterministic_oracle"]["token"] = "UNREGISTERED_ORACLE"
        workspace = new_workspace("impact-oracle-drift")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", oracle_drift), expected_returncode=1)
        require_bundle_issue(payload, "IMPACT_CLAIM_UNSUPPORTED")

        severity_overclaim = json_clone(valid_bundle_contract())
        severity_overclaim["finding"]["severity"] = "High"
        severity_overclaim["validity_review"]["final_severity"] = "High"
        severity_overclaim["validity_review"]["original_severity"] = "High"
        workspace = new_workspace("severity-overclaim")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", severity_overclaim), expected_returncode=1)
        require_bundle_issue(payload, "SEVERITY_EVIDENCE_MISMATCH")

        not_valid = json_clone(valid_bundle_contract())
        not_valid["validity_review"]["validity_verdict"] = "not_valid"
        workspace = new_workspace("not-valid")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", not_valid), expected_returncode=1)
        require_bundle_issue(payload, "VALIDITY_VERDICT_NOT_PROMOTABLE")

        downgraded = json_clone(valid_bundle_contract())
        downgraded["validity_review"]["classification_decision"] = "downgraded"
        downgraded["validity_review"]["original_severity"] = "High"
        workspace = new_workspace("downgraded")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "good", downgraded), expected_returncode=0)
        if not payload.get("valid"):
            raise SystemExit("FAILED: evidence-bounded downgraded contract did not pass")

        reclassified = json_clone(valid_bundle_contract())
        reclassified["validity_review"]["classification_decision"] = "reclassified"
        reclassified["validity_review"]["original_bug_class"] = "Unbounded Request Forwarding"
        workspace = new_workspace("reclassified")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "good", reclassified), expected_returncode=0)
        if not payload.get("valid"):
            raise SystemExit("FAILED: evidence-bounded reclassified contract did not pass")

        fixture_workspace = new_workspace("fixture")
        fixture_doc = json_clone(valid_bundle_contract())
        fixture_doc["fixture_provenance"] = {
            "required": True,
            "replay_type": "minimal_fixture",
            "upstream_sources": ["src/importer.py"],
            "preserved_behavior": "The fixture preserves URL parsing and the HTTP fetch sink.",
            "sufficiency_reason": "The vulnerable source-to-sink boundary is independent of omitted UI code.",
            "consumer_boundary": "The consuming application passes attacker-controlled URLs to the library API.",
            "non_claims": ["No response content exposure is claimed."],
            "synthetic_security_properties_present": True,
            "security_properties": [
                {
                    "id": "PROP-ORACLE",
                    "origin": "synthetic",
                    "use": "oracle_only",
                    "meaning": "A marker identifies the deterministic callback oracle.",
                    "source_reference_ids": [],
                    "cannot_support_impact_claim_ids": ["IMPACT-CALLBACK"],
                }
            ],
        }
        fixture_contract = write_contract(fixture_workspace, "fixture", fixture_doc)
        fixture_payload = run_bundle_contract_json(plugin_root, fixture_workspace, fixture_contract, expected_returncode=0)
        if not fixture_payload.get("valid"):
            raise SystemExit("FAILED: fixture replay contract with provenance did not pass")

        missing_registration = json_clone(valid_bundle_contract())
        missing_registration["replay"]["log"]["registration_targets"] = ["files.reviewer_evidence_index"]
        missing_registration["files"]["reviewer_evidence_index"] = ""
        workspace = new_workspace("missing-registration")
        payload = run_bundle_contract_json(
            plugin_root,
            workspace,
            write_contract(workspace, "bad", missing_registration),
            expected_returncode=1,
        )
        require_bundle_issue(payload, "REPLAY_LOG_UNREGISTERED")

        marker_drift = json_clone(valid_bundle_contract())
        marker_drift["direct_impact"]["sync_targets"][1]["marker"] = "DIFFERENT_MARKER"
        workspace = new_workspace("marker-drift")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", marker_drift), expected_returncode=1)
        require_bundle_issue(payload, "DIRECT_IMPACT_MARKER_DRIFT")

        missing_provenance = json_clone(valid_bundle_contract())
        missing_provenance["fixture_provenance"] = {
            "required": True,
            "replay_type": "minimal_fixture",
            "upstream_sources": [],
        }
        workspace = new_workspace("missing-provenance")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", missing_provenance), expected_returncode=1)
        require_bundle_issue(payload, "FIXTURE_PROVENANCE_MISSING")

        ssrf_overclaim = json_clone(valid_bundle_contract())
        ssrf_overclaim["impact_tier"]["ssrf"]["claimed_exposures"] = [
            "callback_reachability",
            "response_content_exposure",
        ]
        workspace = new_workspace("ssrf-overclaim")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", ssrf_overclaim), expected_returncode=1)
        require_bundle_issue(payload, "SSRF_IMPACT_OVERCLAIM")

        path_escape = json_clone(valid_bundle_contract())
        path_escape["files"]["evidence_files"][0] = "../outside/replay-output.log"
        workspace = new_workspace("path-escape")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", path_escape), expected_returncode=1)
        require_bundle_issue(payload, "BUNDLE_PATH_ESCAPE")

        code_context_thin = json_clone(valid_bundle_contract())
        code_context_thin["code_context"]["entries"] = [{"source_path": "src/importer.py"}]
        workspace = new_workspace("thin-code-context")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", code_context_thin), expected_returncode=1)
        require_bundle_issue(payload, "CODE_CONTEXT_TOO_THIN")

        code_level_only = json_clone(valid_bundle_contract())
        code_level_only["entrypoint_evidence"]["evidence_level"] = "code_level_reproduced"
        workspace = new_workspace("code-level-only")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", code_level_only), expected_returncode=1)
        require_bundle_issue(payload, "CODE_LEVEL_ONLY_NOT_BUNDLE_READY")

        blocked_entrypoint = json_clone(valid_bundle_contract())
        blocked_entrypoint["entrypoint_evidence"]["evidence_level"] = "blocked_entrypoint_verification"
        workspace = new_workspace("blocked-entrypoint")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", blocked_entrypoint), expected_returncode=1)
        require_bundle_issue(payload, "ENTRYPOINT_EVIDENCE_MISSING")

        missing_entrypoint_path = json_clone(valid_bundle_contract())
        missing_entrypoint_path["entrypoint_evidence"]["entrypoint_to_sink_path"] = ""
        workspace = new_workspace("missing-entrypoint-path")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", missing_entrypoint_path), expected_returncode=1)
        require_bundle_issue(payload, "ENTRYPOINT_TO_SINK_PATH_MISSING")

        missing_impact_oracle = json_clone(valid_bundle_contract())
        missing_impact_oracle["entrypoint_evidence"]["deterministic_impact_oracle"] = ""
        workspace = new_workspace("missing-impact-oracle")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", missing_impact_oracle), expected_returncode=1)
        require_bundle_issue(payload, "IMPACT_ORACLE_MISSING")

        missing_replay_material = json_clone(valid_bundle_contract())
        missing_replay_material["entrypoint_evidence"]["replay_material"] = {"description": ""}
        workspace = new_workspace("missing-replay-material")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", missing_replay_material), expected_returncode=1)
        require_bundle_issue(payload, "REPLAY_MATERIAL_MISSING")

        invalid_severity = json_clone(valid_bundle_contract())
        invalid_severity["finding"]["severity"] = "高危"
        workspace = new_workspace("invalid-severity")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", invalid_severity), expected_returncode=1)
        require_bundle_issue(payload, "SEVERITY_ENUM_INVALID")

        open_bug_class = json_clone(valid_bundle_contract())
        open_bug_class["finding"]["bug_class"] = "unsafe native plugin loading"
        open_bug_class["impact_tier"] = {"bug_class": "unsafe native plugin loading"}
        open_bug_class["impact_claims"][0]["supported_bug_classes"] = ["unsafe native plugin loading"]
        open_bug_class["validity_review"]["original_bug_class"] = "unsafe native plugin loading"
        open_bug_class["validity_review"]["final_bug_class"] = "unsafe native plugin loading"
        workspace = new_workspace("open-bug-class")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "open", open_bug_class), expected_returncode=0)
        if not payload.get("valid"):
            raise SystemExit(f"FAILED: free-text bug_class bundle contract did not pass:\n{payload}")

        existing_target = json_clone(valid_bundle_contract())
        workspace = new_workspace("existing-target")
        (workspace / "confirmed" / existing_target["bundle"]["slug"]).mkdir(parents=True)
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", existing_target), expected_returncode=1)
        require_bundle_issue(payload, "FINAL_TARGET_EXISTS")

        outside_final = json_clone(valid_bundle_contract())
        outside_final["bundle"]["final_path"] = "confirmed/other-slug"
        workspace = new_workspace("outside-final")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", outside_final), expected_returncode=1)
        require_bundle_issue(payload, "BUNDLE_PATH_ESCAPE")

        multi_bad = json_clone(valid_bundle_contract())
        multi_bad["docker_evidence"]["verification_status"] = "failed_timeout"
        multi_bad["replay"]["log"]["registration_targets"] = []
        multi_bad["direct_impact"]["sync_targets"][0]["marker"] = "DRIFTED"
        multi_bad["code_context"]["entries"] = [{}]
        workspace = new_workspace("multi-bad")
        payload = run_bundle_contract_json(plugin_root, workspace, write_contract(workspace, "bad", multi_bad), expected_returncode=1)
        if len(payload.get("issues", [])) < 3:
            raise SystemExit(f"FAILED: bad bundle contract did not return 3+ JSON issues:\n{payload}")
        require_bundle_issue(payload, "DOCKER_STATUS_NOT_CONFIRMED")
        require_bundle_issue(payload, "REPLAY_LOG_UNREGISTERED")
        require_bundle_issue(payload, "DIRECT_IMPACT_MARKER_DRIFT")
        require_bundle_issue(payload, "CODE_CONTEXT_TOO_THIN")

    help_output = run_capture([sys.executable, str(validator), "--help"], plugin_root)
    if "--repo-root" not in help_output:
        raise SystemExit("FAILED: bundle contract validator help is missing --repo-root")


def build_wrapper_source_finding(
    plugin_root: Path,
    repo_dir: Path,
    workspace: Path,
    *,
    slug: str = "demo-app_Path_Traversal_高危漏洞报告",
) -> str:
    (repo_dir / "docker").mkdir(parents=True, exist_ok=True)
    (repo_dir / "poc").mkdir(parents=True, exist_ok=True)
    (repo_dir / "evidence").mkdir(parents=True, exist_ok=True)
    (repo_dir / "docker/docker-compose.attacker.yml").write_text(
        "services:\n  attacker:\n    image: alpine:3.20\n",
        encoding="utf-8",
    )
    (repo_dir / "poc/path_traversal.py").write_text("print('root:x:0:0:')\n", encoding="utf-8")
    (repo_dir / "evidence/replay-output.log").write_text(
        "Zhulong reviewer replay log\n"
        "Generated at: 2026-06-24T00:00:00Z\n"
        "COMMAND: docker compose -f attachments/docker/docker-compose.attacker.yml up --abort-on-container-exit\n"
        "stdout: deterministic wrapper selftest replay completed\n"
        "success marker verified with grep -Fq\n"
        "root:x:0:0:\n"
        "DIRECT_IMPACT_CONFIRMED\n",
        encoding="utf-8",
    )

    finding = json.loads((plugin_root / "assets/examples/confirmed-findings.example.json").read_text(encoding="utf-8"))[0]
    finding["filename"] = f"{slug}.docx"
    finding["slug"] = "demo-app-path-traversal"
    finding["project_root_dir"] = "."
    finding.setdefault("verification_evidence", {})["finding_slug"] = finding["slug"]
    evidence_files = finding["verification_evidence"].setdefault("evidence_files", [])
    if "attachments/evidence/replay-output.log" not in evidence_files:
        evidence_files.append("attachments/evidence/replay-output.log")
    finding.setdefault("attachments", []).append(
        {
            "path": "evidence/replay-output.log",
            "purpose": "真实 replay 日志，证明 wrapper 自测的 direct-impact 标记来自已存在附件。",
        }
    )
    (workspace / "confirmed").mkdir(parents=True, exist_ok=True)
    (workspace / "confirmed/findings.json").write_text(
        json.dumps({"findings": [finding]}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return slug


def build_wrapper_contract(workspace: Path, slug: str, *, finding_slug: str = "demo-app-path-traversal") -> Path:
    contract = valid_bundle_contract()
    contract["bundle"]["slug"] = slug
    contract["bundle"]["language"] = "zh-CN"
    contract["bundle"]["final_path"] = f"confirmed/{slug}"
    contract["render"] = {
        "source_findings_json": "confirmed/findings.json",
        "finding_slug": finding_slug,
    }
    contract["finding"] = {
        "project_name": "demo-app",
        "vulnerability_name": "目录遍历",
        "bug_class": "Path Traversal",
        "severity": "High",
        "attacker_condition": "远程攻击者控制下载路径参数。",
        "server_condition": "服务端启用受影响下载接口并从本地文件系统读取文件。",
        "security_impact": "Docker 证据证明攻击者可读取容器内敏感文件内容。",
    }
    contract["impact_tier"]["bug_class"] = "Path Traversal"
    contract["impact_tier"].pop("ssrf", None)
    contract["impact_claims"][0].update(
        {
            "category": "unauthorized_read",
            "statement": "The tested entrypoint reads a file outside the intended directory boundary.",
            "supported_bug_classes": ["Path Traversal"],
            "severity_ceiling": "High",
            "unsupported_stronger_impacts": ["Arbitrary file write and code execution are not proven."],
        }
    )
    contract["validity_review"].update(
        {
            "original_bug_class": "Path Traversal",
            "original_severity": "High",
            "final_bug_class": "Path Traversal",
            "final_cwe": "CWE-22",
            "final_severity": "High",
            "rationale": "Source-bound entrypoint and sink evidence support only the verified file-read boundary.",
            "stronger_impacts_not_claimed": ["Arbitrary file write and code execution are not proven."],
            "cvss": {
                "version": "4.0",
                "vector": "CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:N/VA:N/SC:N/SI:N/SA:N",
                "score": 8.7,
            },
        }
    )
    prepare_source_bound_fixture(workspace.parent, workspace, contract)
    contract_path = workspace / "confirmed/.contracts" / f"{slug}.bundle-contract.json"
    contract_path.parent.mkdir(parents=True, exist_ok=True)
    return write_json_fixture(contract_path, contract)


def new_build_wrapper_workspace(root: Path, name: str) -> tuple[Path, Path]:
    repo_dir = root / name / "repo"
    workspace = repo_dir / "security-research-wrapper"
    (workspace / "confirmed/.contracts").mkdir(parents=True, exist_ok=True)
    (workspace / "asr-config.json").write_text(
        json.dumps(
            {
                "workspace_root": workspace.name,
                "workspace_created_at": "2026-06-24T00:00:00Z",
                "confirmed_output_dir": f"{workspace.name}/confirmed",
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )
    return repo_dir, workspace


def exercise_build_confirmed_bundle_wrapper(plugin_root: Path) -> None:
    wrapper = plugin_root / "scripts/build_confirmed_bundle.py"
    renderer = plugin_root / "scripts/render_confirmed_vuln_docx.py"
    with tempfile.TemporaryDirectory() as tempdir:
        root = Path(tempdir)

        repo_dir, workspace = new_build_wrapper_workspace(root, "positive")
        slug = build_wrapper_source_finding(plugin_root, repo_dir, workspace)
        contract = build_wrapper_contract(workspace, slug)
        run([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_dir),
            "--contract",
            str(contract),
            "--language",
            "zh-CN",
        ], plugin_root)
        final_bundle = workspace / "confirmed" / slug
        if not final_bundle.is_dir():
            raise SystemExit("FAILED: build_confirmed_bundle.py did not promote the final bundle")
        manifest_path = final_bundle / "bundle-build-manifest.json"
        if not manifest_path.is_file():
            raise SystemExit("FAILED: promoted bundle is missing bundle-build-manifest.json")
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        for key in ("contract_sha256", "tested_ref", "source_binding_sha256"):
            if not str(manifest.get(key) or "").strip():
                raise SystemExit(f"FAILED: source-bound build manifest missing {key}")
        for required_name in ("findings.json", "validity-review.json", "verification-evidence.json"):
            if not (final_bundle / required_name).is_file():
                raise SystemExit(f"FAILED: promoted source-bound bundle is missing {required_name}")
        reviewer_index = json.loads((final_bundle / "attachments/reviewer-evidence-index.json").read_text(encoding="utf-8"))
        for key in ("source_binding", "fixture_provenance", "impact_claims", "deployment_prerequisites", "validity_review"):
            if key not in reviewer_index:
                raise SystemExit(f"FAILED: reviewer evidence index missing source-bound field {key}")
        replay_logs = manifest.get("replay_logs")
        if not isinstance(replay_logs, list) or not replay_logs:
            raise SystemExit("FAILED: bundle-build-manifest.json must include replay_logs provenance")
        replay_entry = replay_logs[0]
        for key in ("path", "source_kind", "trust_classification", "sha256", "source_path", "provenance"):
            if not str(replay_entry.get(key) or "").strip():
                raise SystemExit(f"FAILED: replay log manifest entry missing {key}")
        if replay_entry.get("path") != "attachments/evidence/replay-output.log":
            raise SystemExit("FAILED: replay log manifest path mismatch")
        if replay_entry.get("source_kind") != "copied_successful_transcript":
            raise SystemExit("FAILED: replay log manifest source_kind mismatch")
        if replay_entry.get("trust_classification") != "trusted_transcript":
            raise SystemExit("FAILED: replay log manifest did not classify valid transcript as trusted")
        for value in replay_entry.values():
            if isinstance(value, str) and (value.startswith("/") or "/Users/" in value or "\\Users\\" in value):
                raise SystemExit("FAILED: replay log manifest contains a local absolute path")
        if (workspace / "confirmed/.staging" / slug).exists():
            raise SystemExit("FAILED: staging bundle remained visible after promote")

        docx_path = next(final_bundle.glob("*.docx"))
        original_docx = docx_path.read_bytes()
        prerequisite_text = "The tested deployment enables outbound requests for the import feature."
        rewrite_docx_paragraphs(docx_path, lambda text: None if prerequisite_text in text else text)
        run_expect_fail(
            [
                sys.executable,
                str(plugin_root / "scripts/validate_report_bundle.py"),
                "--bundle-dir",
                str(final_bundle),
                "--language",
                "zh-CN",
            ],
            plugin_root,
            "DOCX omits final validity/classification/severity/condition material",
        )
        docx_path.write_bytes(original_docx)

        direct_repo, direct_workspace = new_build_wrapper_workspace(root, "direct-renderer")
        direct_slug = build_wrapper_source_finding(plugin_root, direct_repo, direct_workspace)
        run([
            sys.executable,
            str(renderer),
            "--input",
            str(direct_workspace / "confirmed/findings.json"),
            "--output-dir",
            str(direct_workspace / "confirmed"),
            "--language",
            "zh-CN",
        ], direct_repo)
        if not (direct_workspace / "confirmed" / direct_slug).is_dir():
            raise SystemExit("FAILED: old direct renderer command no longer creates the bundle")

        invalid_repo, invalid_workspace = new_build_wrapper_workspace(root, "invalid-contract")
        invalid_slug = build_wrapper_source_finding(plugin_root, invalid_repo, invalid_workspace)
        invalid_contract_data = valid_bundle_contract()
        invalid_contract_data["bundle"]["slug"] = invalid_slug
        invalid_contract_data["bundle"]["final_path"] = f"confirmed/{invalid_slug}"
        invalid_contract_data["render"] = {
            "source_findings_json": "confirmed/findings.json",
            "finding_slug": "demo-app-path-traversal",
        }
        invalid_contract_data["docker_evidence"]["verification_status"] = "failed_timeout"
        invalid_contract = write_json_fixture(
            invalid_workspace / "confirmed/.contracts/invalid.bundle-contract.json",
            invalid_contract_data,
        )
        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(invalid_workspace),
            "--repo-root",
            str(invalid_repo),
            "--contract",
            str(invalid_contract),
            "--language",
            "zh-CN",
        ], plugin_root, "bundle contract preflight failed")
        if (invalid_workspace / "confirmed" / invalid_slug).exists():
            raise SystemExit("FAILED: invalid contract created a final bundle")

        invalid_validity_repo, invalid_validity_workspace = new_build_wrapper_workspace(root, "not-valid-contract")
        invalid_validity_slug = build_wrapper_source_finding(plugin_root, invalid_validity_repo, invalid_validity_workspace)
        invalid_validity_contract = build_wrapper_contract(invalid_validity_workspace, invalid_validity_slug)
        invalid_validity_data = json.loads(invalid_validity_contract.read_text(encoding="utf-8"))
        invalid_validity_data["validity_review"]["validity_verdict"] = "withdrawn"
        write_json_fixture(invalid_validity_contract, invalid_validity_data)
        run_expect_fail(
            [
                sys.executable,
                str(wrapper),
                "--workspace-dir",
                str(invalid_validity_workspace),
                "--repo-root",
                str(invalid_validity_repo),
                "--contract",
                str(invalid_validity_contract),
                "--language",
                "zh-CN",
            ],
            plugin_root,
            "VALIDITY_VERDICT_NOT_PROMOTABLE",
        )
        if (invalid_validity_workspace / "confirmed" / invalid_validity_slug).exists():
            raise SystemExit("FAILED: withdrawn validity contract created a final bundle")

        multi_repo, multi_workspace = new_build_wrapper_workspace(root, "multi-finding")
        multi_slug = build_wrapper_source_finding(plugin_root, multi_repo, multi_workspace)
        multi_data = json.loads((multi_workspace / "confirmed/findings.json").read_text(encoding="utf-8"))
        multi_data["findings"].append(json.loads(json.dumps(multi_data["findings"][0])))
        (multi_workspace / "confirmed/findings.json").write_text(
            json.dumps(multi_data, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        multi_contract = build_wrapper_contract(multi_workspace, multi_slug)
        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(multi_workspace),
            "--repo-root",
            str(multi_repo),
            "--contract",
            str(multi_contract),
            "--language",
            "zh-CN",
        ], plugin_root, "selected 2 findings")
        if (multi_workspace / "confirmed" / multi_slug).exists():
            raise SystemExit("FAILED: multi-finding source promoted a final bundle")

        validation_repo, validation_workspace = new_build_wrapper_workspace(root, "validation-failure")
        validation_slug = build_wrapper_source_finding(plugin_root, validation_repo, validation_workspace)
        (validation_repo / "evidence/replay-output.log").write_text(
            "Zhulong reviewer replay log placeholder.\n"
            "Run the bundle-root replay script to refresh this file with live reviewer output.\n"
            "Replay contract direct-impact marker: DIRECT_IMPACT_CONFIRMED\n",
            encoding="utf-8",
        )
        validation_contract = build_wrapper_contract(validation_workspace, validation_slug)
        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(validation_workspace),
            "--repo-root",
            str(validation_repo),
            "--contract",
            str(validation_contract),
            "--language",
            "zh-CN",
            "--keep-failed-staging",
        ], plugin_root, "validate_report_bundle.py")
        if (validation_workspace / "confirmed" / validation_slug).exists():
            raise SystemExit("FAILED: staging validation failure promoted a final bundle")
        if not (validation_workspace / "confirmed/.staging" / validation_slug).is_dir():
            raise SystemExit("FAILED: --keep-failed-staging did not preserve failed staging")

        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_dir),
            "--contract",
            str(contract),
            "--language",
            "zh-CN",
        ], plugin_root, "FINAL_TARGET_EXISTS")

        run([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_dir),
            "--contract",
            str(contract),
            "--language",
            "zh-CN",
            "--replace-existing-validated-bundle",
        ], plugin_root)
        trash_entries = sorted((workspace / "confirmed/.staging/.trash").glob(f"{slug}-*"))
        if not trash_entries:
            raise SystemExit("FAILED: replacement did not move existing final bundle to .staging/.trash/")
        if not (workspace / "confirmed" / slug / "bundle-build-manifest.json").is_file():
            raise SystemExit("FAILED: replacement did not promote the new valid bundle")

        stale_repo, stale_workspace = new_build_wrapper_workspace(root, "stale-staging")
        stale_slug = build_wrapper_source_finding(plugin_root, stale_repo, stale_workspace)
        stale_contract = build_wrapper_contract(stale_workspace, stale_slug)
        (stale_workspace / "confirmed/.staging" / stale_slug).mkdir(parents=True)
        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(stale_workspace),
            "--repo-root",
            str(stale_repo),
            "--contract",
            str(stale_contract),
            "--language",
            "zh-CN",
        ], plugin_root, "staging target already exists")
        if (stale_workspace / "confirmed" / stale_slug).exists():
            raise SystemExit("FAILED: pre-existing staging directory promoted a final bundle")

        escape_repo, escape_workspace = new_build_wrapper_workspace(root, "escape")
        escape_slug = build_wrapper_source_finding(plugin_root, escape_repo, escape_workspace)
        escape_contract_data = valid_bundle_contract()
        escape_contract_data["bundle"]["slug"] = escape_slug
        escape_contract_data["bundle"]["final_path"] = "confirmed/other-slug"
        escape_contract_data["render"] = {
            "source_findings_json": "confirmed/findings.json",
            "finding_slug": "demo-app-path-traversal",
        }
        escape_contract = write_json_fixture(
            escape_workspace / "confirmed/.contracts/escape.bundle-contract.json",
            escape_contract_data,
        )
        run_expect_fail([
            sys.executable,
            str(wrapper),
            "--workspace-dir",
            str(escape_workspace),
            "--repo-root",
            str(escape_repo),
            "--contract",
            str(escape_contract),
            "--language",
            "zh-CN",
        ], plugin_root, "bundle contract preflight failed")
        if (escape_workspace / "confirmed" / escape_slug).exists():
            raise SystemExit("FAILED: final path escape promoted a final bundle")


def exercise_finding_contract_validators(plugin_root: Path) -> None:
    candidate_validator = plugin_root / "scripts/validate_candidate.py"
    verdict_validator = plugin_root / "scripts/validate_verifier_verdict.py"
    candidate_example = plugin_root / "assets/examples/candidate.example.json"
    verdict_example = plugin_root / "assets/examples/verifier-verdict.example.json"

    candidate_schema = json.loads((plugin_root / "assets/schemas/candidate.schema.json").read_text(encoding="utf-8"))
    if candidate_schema.get("title") != "Zhulong Candidate Finding Contract R1/R2":
        raise SystemExit("FAILED: candidate schema title mismatch")
    if str(candidate_schema.get("$id", "")).startswith("https://zhulong.local/"):
        raise SystemExit("FAILED: candidate schema must not use a placeholder $id URI")
    if '"additionalProperties": true' in json.dumps(candidate_schema, sort_keys=True):
        raise SystemExit("FAILED: candidate schema must not leave known contract objects open to arbitrary properties")

    verdict_schema = json.loads((plugin_root / "assets/schemas/verifier-verdict.schema.json").read_text(encoding="utf-8"))
    if verdict_schema.get("title") != "Zhulong Verifier Verdict Contract R1":
        raise SystemExit("FAILED: verifier verdict schema title mismatch")
    if str(verdict_schema.get("$id", "")).startswith("https://zhulong.local/"):
        raise SystemExit("FAILED: verifier verdict schema must not use a placeholder $id URI")
    if '"additionalProperties": true' in json.dumps(verdict_schema, sort_keys=True):
        raise SystemExit("FAILED: verifier verdict schema must not leave known contract objects open to arbitrary properties")

    candidate_output = run_capture([sys.executable, str(candidate_validator), str(candidate_example)], plugin_root)
    if "OK: candidate valid" not in candidate_output or "status=candidate" not in candidate_output:
        raise SystemExit(f"FAILED: candidate example did not validate:\n{candidate_output}")
    verdict_output = run_capture([
        sys.executable,
        str(verdict_validator),
        "--candidate",
        str(candidate_example),
        str(verdict_example),
    ], plugin_root)
    if "OK: verifier-verdict valid" not in verdict_output or "verdict=confirmed_in_docker" not in verdict_output:
        raise SystemExit(f"FAILED: verifier verdict example did not validate:\n{verdict_output}")

    with tempfile.TemporaryDirectory() as tempdir:
        tmp = Path(tempdir)

        good_candidate = write_json_fixture(tmp / "candidate-good.json", valid_candidate_contract())
        run([sys.executable, str(candidate_validator), str(good_candidate)], plugin_root)

        bad_status = json_clone(valid_candidate_contract())
        bad_status["status"] = "confirmed_in_docker"
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-confirmed-status.json", bad_status))],
            plugin_root,
            "$.status must be exactly candidate",
        )

        missing_claim_source = json_clone(valid_candidate_contract())
        del missing_claim_source["claim"]["source"]
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-missing-claim-source.json", missing_claim_source))],
            plugin_root,
            "missing required string: $.claim.source",
        )

        missing_claim_sink = json_clone(valid_candidate_contract())
        del missing_claim_sink["claim"]["sink"]
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-missing-claim-sink.json", missing_claim_sink))],
            plugin_root,
            "missing required string: $.claim.sink",
        )

        missing_poc_oracle_type = json_clone(valid_candidate_contract())
        del missing_poc_oracle_type["poc"]["expected_oracle"]["type"]
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-missing-poc-oracle.json", missing_poc_oracle_type))],
            plugin_root,
            "missing required string: $.poc.expected_oracle.type",
        )

        absolute_path = json_clone(valid_candidate_contract())
        absolute_path["evidence"]["static_locations"][0]["path"] = "/Users/example/project/src/importer.py"
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-absolute-path.json", absolute_path))],
            plugin_root,
            "operator-local absolute path",
        )

        for index, unsafe_ref in enumerate(("/Users/alice/source", "ghp_SELFTEST_TOKEN", "sha256:\nunsafe"), start=1):
            unsafe_candidate_ref = json_clone(valid_candidate_contract())
            unsafe_candidate_ref["target_ref"]["tested_ref"] = unsafe_ref
            run_expect_fail(
                [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / f"candidate-unsafe-tested-ref-{index}.json", unsafe_candidate_ref))],
                plugin_root,
                "forbidden source-identity material",
            )

        traversal_path = json_clone(valid_candidate_contract())
        traversal_path["poc"]["path"] = "../poc/reproduce.sh"
        run_expect_fail(
            [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / "candidate-parent-traversal.json", traversal_path))],
            plugin_root,
            "parent path traversal",
        )

        unsafe_cases = {
            "broad-prune": ("poc", "path", "docker " + "system " + "prune -af"),
            "dangerous-kill": ("poc", "path", "kill " + "-9 1234"),
            "unsafe-runtime": ("poc", "path", "docker run --privileged -v docker.sock:/sock example"),
        }
        for name, (section, key, value) in unsafe_cases.items():
            bad_candidate = json_clone(valid_candidate_contract())
            bad_candidate[section][key] = value
            expected = "broad Docker prune" if name == "broad-prune" else "dangerous PID kill" if name == "dangerous-kill" else "must not request privileged"
            run_expect_fail(
                [sys.executable, str(candidate_validator), str(write_json_fixture(tmp / f"candidate-{name}.json", bad_candidate))],
                plugin_root,
                expected,
            )

        good_verdict = write_json_fixture(tmp / "verdict-good.json", valid_verifier_verdict())
        run([sys.executable, str(verdict_validator), "--candidate", str(good_candidate), str(good_verdict)], plugin_root)

        unsafe_verdict_ref = json_clone(valid_verifier_verdict())
        unsafe_verdict_ref["target_ref"]["tested_ref"] = "/private/tmp/secret-ref"
        run_expect_fail(
            [sys.executable, str(verdict_validator), str(write_json_fixture(tmp / "verdict-unsafe-tested-ref.json", unsafe_verdict_ref))],
            plugin_root,
            "forbidden source-identity material",
        )

        confirmed_failure_cases = [
            ("fresh-container-false", ("environment", "fresh_container"), False, "fresh_container=true"),
            ("host-network-true", ("environment", "host_network"), True, "host_network=false"),
            ("privileged-true", ("environment", "privileged"), True, "privileged=false"),
            ("docker-socket-mounted", ("environment", "docker_socket_mounted"), True, "docker_socket_mounted=false"),
            ("credential-paths-mounted", ("environment", "credential_paths_mounted"), True, "credential_paths_mounted=false"),
            ("oracle-success-false", ("oracle_result", "success"), False, "oracle_result.success=true"),
            ("empty-artifacts", ("artifacts",), [], "non-empty artifacts"),
            ("code-level-only", ("evidence_level",), "code_level_reproduced", "code_level_reproduced evidence is supporting evidence only"),
            ("blocked-entrypoint", ("evidence_level",), "blocked_entrypoint_verification", "blocked_entrypoint_verification cannot produce confirmed_in_docker"),
            ("missing-input-shape", ("attacker_entrypoint", "input_shape"), "", "$.attacker_entrypoint.input_shape"),
            ("missing-entrypoint-path", ("attacker_entrypoint", "entrypoint_to_sink_path"), "", "$.attacker_entrypoint.entrypoint_to_sink_path"),
            ("missing-impact-oracle", ("attacker_entrypoint", "deterministic_impact_oracle"), "", "$.attacker_entrypoint.deterministic_impact_oracle"),
            ("missing-replay-path", ("replay_material", "path"), "", "$.replay_material requires path or generation_command"),
        ]
        for name, path_keys, value, expected in confirmed_failure_cases:
            bad_verdict = json_clone(valid_verifier_verdict())
            if len(path_keys) == 1:
                bad_verdict[path_keys[0]] = value
            else:
                bad_verdict[path_keys[0]][path_keys[1]] = value
            run_expect_fail(
                [sys.executable, str(verdict_validator), str(write_json_fixture(tmp / f"verdict-{name}.json", bad_verdict))],
                plugin_root,
                expected,
            )

        for verdict_name in ["blocked", "false_positive", "unverified"]:
            non_confirmed = json_clone(valid_verifier_verdict())
            non_confirmed["verdict"] = verdict_name
            non_confirmed["verification_status"] = verdict_name
            non_confirmed["disposition_recommendation"] = verdict_name
            non_confirmed["commands"] = []
            non_confirmed["artifacts"] = []
            non_confirmed["evidence_level"] = "blocked_entrypoint_verification" if verdict_name == "blocked" else "code_level_reproduced"
            non_confirmed["oracle_result"]["success"] = False
            non_confirmed["oracle_result"]["summary"] = f"{verdict_name} decision has a clear selftest reason."
            non_confirmed["reason"] = f"{verdict_name} selftest reason"
            run([sys.executable, str(verdict_validator), str(write_json_fixture(tmp / f"verdict-{verdict_name}.json", non_confirmed))], plugin_root)

        mismatch = json_clone(valid_verifier_verdict())
        mismatch["candidate_id"] = "CAND-9999"
        run_expect_fail(
            [
                sys.executable,
                str(verdict_validator),
                "--candidate",
                str(good_candidate),
                str(write_json_fixture(tmp / "verdict-candidate-id-mismatch.json", mismatch)),
            ],
            plugin_root,
            "candidate_id mismatch",
        )


def exercise_independent_verifier(plugin_root: Path) -> None:
    verifier = plugin_root / "scripts/verify_candidate.py"
    verdict_validator = plugin_root / "scripts/validate_verifier_verdict.py"
    doc_path = plugin_root / "docs/runner-contracts/independent-verifier-r1.md"
    require_text(doc_path, "finder cannot self-certify", "independent verifier finder/verifier separation")
    require_text(doc_path, "does not read `finder-notes.md`", "independent verifier ignores finder notes")
    require_text(doc_path, "ZC-004 handles disposition promotion", "independent verifier ZC-004 boundary")
    require_text(doc_path, "does not create a confirmed bundle", "independent verifier bundle boundary")

    with tempfile.TemporaryDirectory(prefix="zhulong-verifier-selftest-") as tempdir:
        tmp = Path(tempdir)
        workspace = tmp / "security-research-selftest"
        workspace.mkdir(parents=True)
        fakebin = tmp / "fakebin"
        fakebin.mkdir()
        docker_marker = tmp / "docker-called"
        fake_docker = fakebin / "docker"
        fake_docker.write_text(f"#!/usr/bin/env bash\nprintf called > {docker_marker}\nexit 77\n", encoding="utf-8")
        fake_docker.chmod(0o755)
        env = {"PATH": f"{fakebin}{os.pathsep}{os.environ.get('PATH', '')}"}

        def write_target(name: str, content: str) -> Path:
            path = workspace / name
            path.write_text(content, encoding="utf-8")
            return path

        def candidate_doc(*, tested_ref: str = "local-state", oracle: str = "callback_observed") -> dict:
            candidate = valid_candidate_contract(
                {
                    "target_ref": {
                        "target_config": "zhulong-target.yaml",
                        "tested_ref": tested_ref,
                    }
                }
            )
            candidate["poc"]["expected_oracle"]["type"] = oracle
            return candidate

        def write_candidate(name: str, candidate: dict) -> Path:
            cand_dir = workspace / "candidates" / candidate.get("candidate_id", "CAND-0001")
            cand_dir.mkdir(parents=True, exist_ok=True)
            (cand_dir / "finder-notes.md").write_text(
                "Finder note says confirmed, but verifier must ignore this wording.\n",
                encoding="utf-8",
            )
            return write_json_fixture(cand_dir / name, candidate)

        def verdict_path(candidate_id: str = "CAND-0001", suffix: str = "verifier-verdict.json") -> Path:
            return workspace / "verifier" / candidate_id / suffix

        def run_verify(
            target: Path,
            candidate: Path,
            out: Path,
            *,
            run_id: str,
            expected_returncode: int = 1,
            extra: list[str] | None = None,
        ) -> str:
            command = [
                sys.executable,
                str(verifier),
                "--target-config",
                str(target),
                "--candidate",
                str(candidate),
                "--workspace",
                str(workspace),
                "--out",
                str(out),
                "--run-id",
                run_id,
                "--dry-run",
                "--no-execute",
            ]
            if extra:
                command.extend(extra)
            return run_capture_with_env(command, plugin_root, env, expected_returncode=expected_returncode)

        def load_json(path: Path) -> dict:
            return json.loads(path.read_text(encoding="utf-8"))

        def assert_valid_verdict(candidate: Path, out: Path) -> dict:
            run([sys.executable, str(verdict_validator), "--candidate", str(candidate), str(out)], plugin_root)
            return load_json(out)

        target = write_target("zhulong-target.yaml", valid_target_contract_yaml(runtime_type="docker"))
        candidate = write_candidate("candidate.json", candidate_doc())

        out_unverified = verdict_path(suffix="unverified.json")
        output = run_verify(target, candidate, out_unverified, run_id="no-execute")
        if "verdict=unverified" not in output:
            raise SystemExit(f"FAILED: no-execute verifier did not report unverified:\n{output}")
        unverified_doc = assert_valid_verdict(candidate, out_unverified)
        if unverified_doc["verdict"] != "unverified" or unverified_doc["oracle_result"]["success"]:
            raise SystemExit("FAILED: no-execute verifier did not produce unverified oracle mismatch")
        if not (workspace / "verifier/CAND-0001/runs/no-execute/verifier.log").exists():
            raise SystemExit("FAILED: verifier run directory was not created under workspace")

        out_allow_execute = verdict_path(suffix="allow-execute-blocked.json")
        allow_execute_output = run_verify(
            target,
            candidate,
            out_allow_execute,
            run_id="allow-execute-blocked",
            extra=["--allow-execute"],
        )
        if "verdict=blocked" not in allow_execute_output:
            raise SystemExit("FAILED: R1 --allow-execute did not remain explicitly blocked")
        allow_execute_doc = assert_valid_verdict(candidate, out_allow_execute)
        if (
            allow_execute_doc["verdict"] != "blocked"
            or "Docker execution is not implemented in R1 verifier" not in allow_execute_doc.get("reason", "")
        ):
            raise SystemExit("FAILED: R1 --allow-execute boundary changed")
        if docker_marker.exists():
            raise SystemExit("FAILED: R1 --allow-execute invoked Docker")

        manual_target = write_target("manual-target.yaml", valid_target_contract_yaml(runtime_type="manual-blocked"))
        manual_candidate_doc = candidate_doc()
        manual_candidate_doc["target_ref"]["target_config"] = "manual-target.yaml"
        manual_candidate = write_candidate("candidate-manual.json", manual_candidate_doc)
        out_manual = verdict_path(suffix="manual-blocked.json")
        run_verify(manual_target, manual_candidate, out_manual, run_id="manual-blocked")
        manual_doc = assert_valid_verdict(manual_candidate, out_manual)
        if manual_doc["verdict"] != "blocked" or "manual-blocked" not in manual_doc.get("reason", ""):
            raise SystemExit("FAILED: manual-blocked target did not produce blocked verdict")

        unsupported_candidate_doc = candidate_doc(oracle="unsupported_fixture_oracle")
        unsupported_candidate = write_candidate("candidate-unsupported.json", unsupported_candidate_doc)
        out_unsupported = verdict_path(suffix="unsupported.json")
        run_verify(target, unsupported_candidate, out_unsupported, run_id="unsupported-oracle")
        unsupported_doc = assert_valid_verdict(unsupported_candidate, out_unsupported)
        if unsupported_doc["verdict"] != "blocked" or "unsupported oracle type" not in unsupported_doc.get("reason", ""):
            raise SystemExit("FAILED: unsupported oracle did not block verification")

        out_confirmed = verdict_path(suffix="fixture-confirmed.json")
        confirmed_output = run_verify(
            target,
            candidate,
            out_confirmed,
            run_id="fixture-confirmed",
            expected_returncode=1,
            extra=["--dry-run-result", "confirmed_in_docker"],
        )
        if "verdict=blocked" not in confirmed_output:
            raise SystemExit(f"FAILED: fixture confirmed verifier did not report blocked entrypoint verification:\n{confirmed_output}")
        confirmed_doc = assert_valid_verdict(candidate, out_confirmed)
        if confirmed_doc["verdict"] != "blocked" or confirmed_doc.get("evidence_level") != "blocked_entrypoint_verification":
            raise SystemExit("FAILED: fixture confirmed dry-run did not stay blocked_entrypoint_verification")
        if "SIMULATED dry-run fixture" not in confirmed_doc["oracle_result"]["summary"]:
            raise SystemExit("FAILED: fixture blocked verdict is not clearly marked simulated")
        if not confirmed_doc["artifacts"]:
            raise SystemExit("FAILED: fixture blocked verdict must preserve fixture artifacts as supporting evidence")

        invalid_target = write_target(
            "invalid-target.yaml",
            valid_target_contract_yaml(runtime_type="docker").replace('  name: "selftest-service"\n', ""),
        )
        invalid_target_candidate_doc = candidate_doc()
        invalid_target_candidate_doc["target_ref"]["target_config"] = "invalid-target.yaml"
        invalid_target_candidate = write_candidate("candidate-invalid-target.json", invalid_target_candidate_doc)
        out_invalid_target = verdict_path(suffix="invalid-target.json")
        run_verify(invalid_target, invalid_target_candidate, out_invalid_target, run_id="invalid-target")
        invalid_target_doc = assert_valid_verdict(invalid_target_candidate, out_invalid_target)
        if invalid_target_doc["verdict"] != "blocked" or "invalid target" not in invalid_target_doc.get("reason", ""):
            raise SystemExit("FAILED: invalid target did not produce blocked verifier verdict")

        invalid_candidate_doc = candidate_doc()
        invalid_candidate_doc["unexpected_field"] = "candidate remains invalid before confirmation"
        invalid_candidate = write_candidate("candidate-invalid.json", invalid_candidate_doc)
        out_invalid_candidate = verdict_path(suffix="invalid-candidate.json")
        invalid_candidate_output = run_verify(
            target,
            invalid_candidate,
            out_invalid_candidate,
            run_id="invalid-candidate",
        )
        if "generated verifier verdict failed validation" not in invalid_candidate_output:
            raise SystemExit("FAILED: invalid candidate did not fail before verifier success")
        invalid_candidate_doc_out = load_json(out_invalid_candidate)
        if invalid_candidate_doc_out["verdict"] == "confirmed_in_docker":
            raise SystemExit("FAILED: invalid candidate produced confirmed verdict")

        mismatch_candidate_doc = candidate_doc(tested_ref="different-ref")
        mismatch_candidate = write_candidate("candidate-mismatch.json", mismatch_candidate_doc)
        out_mismatch = verdict_path(suffix="target-ref-mismatch.json")
        run_verify(target, mismatch_candidate, out_mismatch, run_id="target-ref-mismatch")
        mismatch_doc = assert_valid_verdict(mismatch_candidate, out_mismatch)
        if mismatch_doc["verdict"] != "blocked" or "target_ref.tested_ref" not in mismatch_doc.get("reason", ""):
            raise SystemExit("FAILED: target_ref mismatch did not block verification")

        unsafe_target = write_target(
            "unsafe-target.yaml",
            valid_target_contract_yaml(runtime_type="docker").replace(
                "docker run --name zhulong-target-selftest -d -p 127.0.0.1:8080:8080 zhulong-target-selftest",
                "docker run --privileged example",
            ),
        )
        unsafe_candidate_doc = candidate_doc()
        unsafe_candidate_doc["target_ref"]["target_config"] = "unsafe-target.yaml"
        unsafe_candidate = write_candidate("candidate-unsafe.json", unsafe_candidate_doc)
        out_unsafe = verdict_path(suffix="unsafe-runtime.json")
        run_verify(unsafe_target, unsafe_candidate, out_unsafe, run_id="unsafe-runtime")
        unsafe_doc = assert_valid_verdict(unsafe_candidate, out_unsafe)
        if unsafe_doc["verdict"] != "blocked" or "invalid target" not in unsafe_doc.get("reason", ""):
            raise SystemExit("FAILED: unsafe runtime text did not block verification")

        out_outside = tmp / "outside-verdict.json"
        outside_output = run_verify(target, candidate, out_outside, run_id="outside-output")
        if "verifier verdict output must stay under the workspace" not in outside_output:
            raise SystemExit("FAILED: verifier allowed output path outside workspace")

        if docker_marker.exists():
            raise SystemExit("FAILED: independent verifier selftest invoked Docker")


def exercise_disposition_integration(plugin_root: Path) -> None:
    disposition = plugin_root / "scripts/audit_disposition.py"
    doc_path = plugin_root / "docs/runner-contracts/disposition-integration-r1.md"
    require_text(doc_path, "Candidate-only records stay `status=candidate`", "disposition candidate-only guard")
    require_text(doc_path, "finder-notes.md` is human context only", "disposition finder notes boundary")
    require_text(doc_path, "ZC-004 does not generate confirmed bundles", "disposition bundle boundary")
    require_text(doc_path, "validate_report_bundle.py pass", "disposition future confirmed bundle gate")

    with tempfile.TemporaryDirectory(prefix="zhulong-disposition-selftest-") as tempdir:
        tmp = Path(tempdir)

        def make_workspace(name: str = "security-research-selftest") -> Path:
            workspace = tmp / name
            workspace.mkdir(parents=True, exist_ok=True)
            return workspace

        def write_candidate(workspace: Path, candidate: dict | None = None) -> Path:
            doc = candidate or valid_candidate_contract()
            cand_dir = workspace / "candidates" / doc.get("candidate_id", "CAND-0001")
            cand_dir.mkdir(parents=True, exist_ok=True)
            (cand_dir / "finder-notes.md").write_text(
                "Finder notes say confirmed_in_docker, but notes are not verifier evidence.\n",
                encoding="utf-8",
            )
            return write_json_fixture(cand_dir / "candidate.json", doc)

        def verdict_doc(status: str, overrides: dict | None = None) -> dict:
            doc = json_clone(valid_verifier_verdict())
            doc["verdict"] = status
            doc["verification_status"] = status
            doc["disposition_recommendation"] = status
            if status != "confirmed_in_docker":
                doc["environment"]["fresh_container"] = False
                doc["commands"] = []
                doc["artifacts"] = []
                doc["evidence_level"] = "blocked_entrypoint_verification" if status == "blocked" else "code_level_reproduced"
                doc["oracle_result"]["success"] = False
                doc["oracle_result"]["summary"] = f"{status} selftest disposition reason."
                doc["reason"] = f"{status} selftest disposition reason"
            if overrides:
                doc.update(overrides)
            return doc

        def write_verdict(workspace: Path, status: str, overrides: dict | None = None) -> Path:
            path = workspace / "verifier/CAND-0001/verifier-verdict.json"
            path.parent.mkdir(parents=True, exist_ok=True)
            return write_json_fixture(path, verdict_doc(status, overrides))

        def load_ledger(workspace: Path) -> dict:
            return json.loads((workspace / "audit-disposition.json").read_text(encoding="utf-8"))

        def candidate_record(workspace: Path) -> dict:
            ledger = load_ledger(workspace)
            records = ledger.get("candidate_dispositions", [])
            if len(records) != 1:
                raise SystemExit(f"FAILED: expected exactly one candidate disposition record: {ledger}")
            return records[0]

        def assert_status(workspace: Path, expected: str) -> dict:
            record = candidate_record(workspace)
            if record.get("status") != expected:
                raise SystemExit(f"FAILED: expected disposition status={expected}, got {record}")
            return record

        def run_update(workspace: Path, candidate: Path, verdict: Path, *, expected_returncode: int = 0) -> str:
            return run_capture_with_env(
                [
                    sys.executable,
                    str(disposition),
                    "--workspace",
                    str(workspace),
                    "--candidate",
                    str(candidate.relative_to(workspace)),
                    "--verdict",
                    str(verdict.relative_to(workspace)),
                    "--update-from-verdict",
                ],
                plugin_root,
                {},
                expected_returncode=expected_returncode,
            )

        candidate_only = make_workspace("candidate-only")
        write_candidate(candidate_only)
        run([sys.executable, str(disposition), "--workspace-dir", str(candidate_only), "--write"], plugin_root)
        candidate_only_record = assert_status(candidate_only, "candidate")
        if candidate_only_record.get("source") != "candidate-contract" or "verdict_path" in candidate_only_record:
            raise SystemExit(f"FAILED: candidate-only disposition should not contain verifier source: {candidate_only_record}")

        notes_guard = make_workspace("finder-notes-guard")
        write_candidate(notes_guard)
        run([sys.executable, str(disposition), "--workspace-dir", str(notes_guard), "--write"], plugin_root)
        assert_status(notes_guard, "candidate")

        for status in ["confirmed_in_docker", "false_positive", "unverified", "blocked"]:
            workspace = make_workspace(f"verdict-{status}")
            candidate_path = write_candidate(workspace)
            verdict_path = write_verdict(workspace, status)
            output = run_update(workspace, candidate_path, verdict_path)
            if "AUDIT DISPOSITION OK" not in output:
                raise SystemExit(f"FAILED: update-from-verdict did not validate ledger for {status}:\n{output}")
            record = assert_status(workspace, status)
            if record.get("source") != "verifier-verdict":
                raise SystemExit(f"FAILED: verdict disposition source drifted: {record}")
            if record.get("candidate_path") != "candidates/CAND-0001/candidate.json":
                raise SystemExit(f"FAILED: candidate path should be workspace-relative: {record}")
            if record.get("verdict_path") != "verifier/CAND-0001/verifier-verdict.json":
                raise SystemExit(f"FAILED: verdict path should be workspace-relative: {record}")
            if status == "confirmed_in_docker" and not record.get("oracle_summary"):
                raise SystemExit("FAILED: confirmed disposition must carry oracle summary")

        invalid_workspace = make_workspace("invalid-verdict")
        invalid_candidate = write_candidate(invalid_workspace)
        run([sys.executable, str(disposition), "--workspace-dir", str(invalid_workspace), "--write"], plugin_root)
        before_invalid = (invalid_workspace / "audit-disposition.json").read_text(encoding="utf-8")
        invalid_doc = verdict_doc("confirmed_in_docker")
        invalid_doc["environment"]["fresh_container"] = False
        invalid_verdict = invalid_workspace / "verifier/CAND-0001/verifier-verdict.json"
        invalid_verdict.parent.mkdir(parents=True, exist_ok=True)
        write_json_fixture(invalid_verdict, invalid_doc)
        run_update(invalid_workspace, invalid_candidate, invalid_verdict, expected_returncode=1)
        after_invalid = (invalid_workspace / "audit-disposition.json").read_text(encoding="utf-8")
        if after_invalid != before_invalid:
            raise SystemExit("FAILED: invalid verifier verdict changed audit-disposition.json")
        assert_status(invalid_workspace, "candidate")

        mismatch_workspace = make_workspace("mismatch-verdict")
        mismatch_candidate = write_candidate(mismatch_workspace)
        run([sys.executable, str(disposition), "--workspace-dir", str(mismatch_workspace), "--write"], plugin_root)
        before_mismatch = (mismatch_workspace / "audit-disposition.json").read_text(encoding="utf-8")
        mismatch_doc = verdict_doc("confirmed_in_docker")
        mismatch_doc["candidate_id"] = "CAND-9999"
        mismatch_verdict = mismatch_workspace / "verifier/CAND-0001/verifier-verdict.json"
        mismatch_verdict.parent.mkdir(parents=True, exist_ok=True)
        write_json_fixture(mismatch_verdict, mismatch_doc)
        run_update(mismatch_workspace, mismatch_candidate, mismatch_verdict, expected_returncode=1)
        after_mismatch = (mismatch_workspace / "audit-disposition.json").read_text(encoding="utf-8")
        if after_mismatch != before_mismatch:
            raise SystemExit("FAILED: candidate/verdict mismatch changed audit-disposition.json")
        assert_status(mismatch_workspace, "candidate")


def exercise_contract_fixture_chain(plugin_root: Path) -> None:
    target_validator = plugin_root / "scripts/validate_target_contract.py"
    candidate_validator = plugin_root / "scripts/validate_candidate.py"
    verdict_validator = plugin_root / "scripts/validate_verifier_verdict.py"
    disposition = plugin_root / "scripts/audit_disposition.py"
    fixture_root = plugin_root / "assets/fixtures/contracts"
    doc_path = plugin_root / "docs/runner-contracts/contract-layer-r1-closure.md"

    require_text(doc_path, "ZC-005 locks the cross-step fixture chain", "contract closure ZC-005 summary")
    require_text(doc_path, "not an autonomous runner", "contract closure runner boundary")
    require_text(doc_path, "Confirmed bundle validation remains separate", "contract closure bundle boundary")
    require_text(doc_path, "No Docker, PoC, replay, scanner, network, GitHub, package registry, or LLM", "contract closure selftest boundary")

    expected_statuses = {
        "confirmed_ssrf": "confirmed_in_docker",
        "false_positive_unreachable": "false_positive",
        "unverified_oracle_weak": "unverified",
        "blocked_manual_runtime": "blocked",
    }

    with tempfile.TemporaryDirectory(prefix="zhulong-contract-chain-selftest-") as tempdir:
        tmp = Path(tempdir)
        fakebin = tmp / "fakebin"
        fakebin.mkdir()
        docker_marker = tmp / "docker-called"
        fake_docker = fakebin / "docker"
        fake_docker.write_text(f"#!/usr/bin/env bash\nprintf called > {docker_marker}\nexit 77\n", encoding="utf-8")
        fake_docker.chmod(0o755)
        env = {"PATH": f"{fakebin}{os.pathsep}{os.environ.get('PATH', '')}"}

        candidate_only_checked = False
        finder_notes_checked = False

        for fixture_name, expected_status in expected_statuses.items():
            fixture_dir = fixture_root / fixture_name
            target_fixture = fixture_dir / "zhulong-target.yaml"
            candidate_fixture = fixture_dir / "candidate.json"
            verdict_fixture = fixture_dir / "verifier-verdict.json"
            expected_fixture = fixture_dir / "expected-disposition.json"
            for path in [target_fixture, candidate_fixture, verdict_fixture, expected_fixture]:
                if not path.exists():
                    raise SystemExit(f"FAILED: missing contract fixture file: {path}")

            run_capture_with_env([sys.executable, str(target_validator), str(target_fixture)], plugin_root, env)
            run_capture_with_env([sys.executable, str(candidate_validator), str(candidate_fixture)], plugin_root, env)
            run_capture_with_env(
                [sys.executable, str(verdict_validator), "--candidate", str(candidate_fixture), str(verdict_fixture)],
                plugin_root,
                env,
            )

            candidate_doc = json.loads(candidate_fixture.read_text(encoding="utf-8"))
            expected_doc = json.loads(expected_fixture.read_text(encoding="utf-8"))
            candidate_id = candidate_doc["candidate_id"]

            workspace = tmp / f"workspace-{fixture_name}"
            workspace.mkdir()
            shutil.copy2(target_fixture, workspace / "zhulong-target.yaml")
            candidate_path = workspace / "candidates" / candidate_id / "candidate.json"
            verdict_path = workspace / "verifier" / candidate_id / "verifier-verdict.json"
            candidate_path.parent.mkdir(parents=True)
            verdict_path.parent.mkdir(parents=True)
            shutil.copy2(candidate_fixture, candidate_path)
            shutil.copy2(verdict_fixture, verdict_path)
            (candidate_path.parent / "finder-notes.md").write_text(
                "Finder notes claim confirmed_in_docker, but fixture status must come only from verifier-verdict.json.\n",
                encoding="utf-8",
            )

            if not candidate_only_checked:
                candidate_only_workspace = tmp / "candidate-only-guard"
                candidate_only_workspace.mkdir()
                shutil.copy2(target_fixture, candidate_only_workspace / "zhulong-target.yaml")
                candidate_only_path = candidate_only_workspace / "candidates" / candidate_id / "candidate.json"
                candidate_only_path.parent.mkdir(parents=True)
                shutil.copy2(candidate_fixture, candidate_only_path)
                (candidate_only_path.parent / "finder-notes.md").write_text(
                    "Finder note says confirmed_in_docker while verifier verdict is withheld.\n",
                    encoding="utf-8",
                )
                run_capture_with_env(
                    [sys.executable, str(disposition), "--workspace-dir", str(candidate_only_workspace), "--write"],
                    plugin_root,
                    env,
                )
                candidate_only_ledger = json.loads(
                    (candidate_only_workspace / "audit-disposition.json").read_text(encoding="utf-8")
                )
                records = candidate_only_ledger.get("candidate_dispositions", [])
                if len(records) != 1 or records[0].get("status") != "candidate" or records[0].get("verdict_path"):
                    raise SystemExit(f"FAILED: candidate-only fixture promoted without verdict: {candidate_only_ledger}")
                candidate_only_checked = True

            output = run_capture_with_env(
                [
                    sys.executable,
                    str(disposition),
                    "--workspace",
                    str(workspace),
                    "--candidate",
                    str(candidate_path.relative_to(workspace)),
                    "--verdict",
                    str(verdict_path.relative_to(workspace)),
                    "--update-from-verdict",
                ],
                plugin_root,
                env,
            )
            if "AUDIT DISPOSITION OK" not in output:
                raise SystemExit(f"FAILED: contract fixture disposition update did not pass for {fixture_name}:\n{output}")
            ledger = json.loads((workspace / "audit-disposition.json").read_text(encoding="utf-8"))
            records = ledger.get("candidate_dispositions", [])
            if len(records) != 1:
                raise SystemExit(f"FAILED: expected one candidate disposition record for {fixture_name}: {ledger}")
            record = records[0]
            if record.get("status") != expected_status or record.get("status") != expected_doc.get("status"):
                raise SystemExit(f"FAILED: fixture {fixture_name} status mismatch: {record}")
            for key in ["source", "candidate_path", "verdict_path", "target_ref", "claim", "oracle_summary"]:
                if record.get(key) != expected_doc.get(key):
                    raise SystemExit(
                        f"FAILED: fixture {fixture_name} disposition field {key} mismatch\n"
                        f"Expected: {expected_doc.get(key)!r}\n"
                        f"Actual: {record.get(key)!r}"
                    )
            if fixture_name == "false_positive_unreachable" and record.get("status") == expected_status:
                finder_notes_checked = True

        if not candidate_only_checked:
            raise SystemExit("FAILED: candidate-only no-verdict guard did not run")
        if not finder_notes_checked:
            raise SystemExit("FAILED: finder-notes status guard did not run")
        if docker_marker.exists():
            raise SystemExit("FAILED: contract fixture chain invoked Docker")


def valid_variant_seed_card() -> dict:
    return {
        "schema_version": 1,
        "seed_id": "seed-confirmed-ssrf-import-url",
        "confirmed_bundle_path": "confirmed/ssrf-import-url",
        "bug_class": "SSRF",
        "root_cause": "Import flow trusted a URL before enforcing private-network deny rules.",
        "source_pattern": "Authenticated attacker controls the import URL in the HTTP request body.",
        "propagation_pattern": "Request body URL reaches the server-side fetch helper without host revalidation.",
        "sink_pattern": "Server-side HTTP fetch/open-url sink can reach internal network targets.",
        "missing_constraint_pattern": "Missing canonicalization and private-address denylist immediately before the sink.",
        "trigger_condition": "Import feature enabled; low-privilege authenticated user can submit imports.",
        "docker_success_oracle": (
            "Docker Compose replay observed the callback and verification-evidence.json records "
            "verification_status=confirmed_in_docker."
        ),
        "search_scope": {
            "repository": "same-target-repository",
            "default": "exclude generated outputs and confirmed bundles",
        },
        "negative_filters": [
            "tests/",
            "docs/",
            "examples/",
            "fixtures/",
            "call sites with canonical host validation before fetch",
        ],
    }


def write_variant_seed_card(path: Path, overrides: dict | None = None, *, drop: str | None = None) -> None:
    card = valid_variant_seed_card()
    if overrides:
        card.update(overrides)
    if drop:
        card.pop(drop, None)
    path.write_text(json.dumps(card, ensure_ascii=False, sort_keys=True) + "\n", encoding="utf-8")


def exercise_variant_seed_card_validation(plugin_root: Path, workspace: Path) -> None:
    validator = plugin_root / "scripts/validate_report_bundle.py"
    seeds_dir = workspace / "variant-seed-fixtures"
    seeds_dir.mkdir(parents=True, exist_ok=True)

    good = seeds_dir / "good-seed.json"
    write_variant_seed_card(good)
    run([sys.executable, str(validator), "--variant-seed-card", str(good)], plugin_root)

    draft_unknown = seeds_dir / "draft-unknown-seed.json"
    write_variant_seed_card(draft_unknown, {"root_cause": "unknown"})
    run([sys.executable, str(validator), "--variant-seed-card", str(draft_unknown), "--variant-seed-draft"], plugin_root)

    bad_cases = [
        ("empty-root-cause.json", {"root_cause": ""}, None, "field 'root_cause' must be a non-empty string"),
        (
            "source-lacks-attacker-control.json",
            {"source_pattern": "import_url parameter"},
            None,
            "source_pattern must describe attacker-controlled or untrusted input",
        ),
        ("missing-sink.json", None, "sink_pattern", "missing required field(s): sink_pattern"),
        (
            "empty-oracle.json",
            {"docker_success_oracle": ""},
            None,
            "field 'docker_success_oracle' must be a non-empty string",
        ),
        (
            "unknown-final.json",
            {"docker_success_oracle": "unknown"},
            None,
            "must not use unknown for required root-cause/source/sink/oracle",
        ),
        (
            "absolute-bundle-path.json",
            {"confirmed_bundle_path": str(Path.cwd() / "confirmed" / "seed")},
            None,
            "confirmed_bundle_path must be bundle-relative or workspace-relative",
        ),
    ]
    for filename, overrides, drop, expected in bad_cases:
        path = seeds_dir / filename
        write_variant_seed_card(path, overrides, drop=drop)
        run_expect_fail([sys.executable, str(validator), "--variant-seed-card", str(path)], plugin_root, expected)


def write_extractor_fixture_bundle(
    bundle: Path,
    *,
    verification_status: str = "confirmed_in_docker",
    include_verification: bool = True,
    complete_finding: bool = True,
    local_path_text: str = "",
    finding_payload: dict | None = None,
) -> None:
    bundle.mkdir(parents=True, exist_ok=True)
    if include_verification:
        (bundle / "verification-evidence.json").write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "finding_slug": bundle.name,
                    "verification_status": verification_status,
                    "docker_required": True,
                    "docker_image": "zhulong-selftest:local",
                    "docker_command": "docker compose run verifier",
                    "poc_path": "attachments/poc.js",
                    "expected_observation": "SSRF callback marker is observed",
                    "observed_observation": "SSRF callback marker is observed in Docker replay",
                    "oracle_token": "SSRF_CALLBACK_CONFIRMED",
                    "evidence_files": ["attachments/evidence/replay.log"],
                    "severity_escalation_attempted": True,
                    "severity_escalation_result": "no_higher_impact_found",
                },
                ensure_ascii=False,
                sort_keys=True,
            ),
            encoding="utf-8",
        )
    if finding_payload is not None:
        (bundle / "findings.json").write_text(
            json.dumps(finding_payload, ensure_ascii=False, sort_keys=True),
            encoding="utf-8",
        )
    elif complete_finding:
        (bundle / "findings.json").write_text(
            json.dumps(
                {
                    "finding_slug": bundle.name,
                    "bug_class": "SSRF",
                    "root_cause": (
                        "Import flow trusts the submitted URL before canonical private-network checks. "
                        + local_path_text
                    ).strip(),
                    "source_pattern": "Authenticated attacker controls the import URL in the HTTP request body.",
                    "propagation_pattern": "Request body URL reaches the server-side import fetch helper unchanged.",
                    "sink_pattern": "Server-side HTTP fetch sink can reach internal network services.",
                    "missing_constraint_pattern": "Missing canonicalization and private-address denylist immediately before the sink.",
                    "trigger_condition": "Import feature enabled; low-privilege authenticated user can submit imports.",
                    "search_scope": "all repositories on GitHub",
                },
                ensure_ascii=False,
                sort_keys=True,
            ),
            encoding="utf-8",
        )


def historical_finding_fixture(*, local_path_text: str = "") -> dict:
    suffix = (" " + local_path_text) if local_path_text else ""
    return {
        "slug": "historical-config-path-traversal",
        "bug_class": "Path Traversal / Code Execution",
        "title": "historical config path traversal",
        "analysis": [
            "入口/可控输入：configPath = argv[configKey]，来自用户完全控制的命令行参数。",
            "危险函数/危险操作：mixin.resolve(mixin.cwd(), configPath) 将攻击者提供的路径解析为绝对路径后直接传递给 mixin.require()，后者调用 require(path) 或 readFileSync(path)。",
            "触发路径：--config=../../tmp/payload -> argv['config'] 获取路径 -> path.resolve(cwd(), path) 解析 -> require(resolvedPath) 加载并执行 JS / readFileSync 读取文件 -> 内容合并至 argv。",
            "根因：setConfig 未对 configPath 进行路径规范化校验，也未限制可加载的文件类型或路径模式。" + suffix,
            "缺失校验：未确保规范化后的路径仍在用户预期目录范围内。",
            "触发条件：攻击者可以控制命令行参数并启用 config 加载功能。",
        ],
        "analysis_en": [
            "Entry / controllable input: configPath = argv[configKey], sourced from fully attacker-controlled command-line arguments.",
            "Dangerous operation: mixin.resolve(mixin.cwd(), configPath) resolves the attacker-supplied path and passes it directly to mixin.require(), which calls require(path) or readFileSync(path).",
            "Root cause: setConfig does not perform path normalization validation and does not restrict loadable file types or path patterns.",
        ],
        "code_context": [
            {
                "summary": "setConfig 从 argv 中读取 configKey 对应的路径值，通过 mixin.resolve 解析后传递给 mixin.require，期间未对路径进行任何校验。",
                "summary_en": "setConfig reads the configKey path from argv, resolves it via mixin.resolve, and passes it to mixin.require without any path validation.",
            },
            {
                "summary": "mixin.require 在 CommonJS 下调用 require(path) 可执行任意 JS 代码；在 ESM 下调用 readFileSync 可读取任意文件。",
                "summary_en": "mixin.require calls require(path) in CommonJS (enables arbitrary code execution); uses readFileSync in ESM (enables arbitrary file read).",
            },
        ],
    }


def exercise_extract_variant_seed(plugin_root: Path, workspace: Path) -> None:
    extractor = plugin_root / "scripts/extract_variant_seed.py"
    validator = plugin_root / "scripts/validate_report_bundle.py"
    fixture_root = workspace / "variant-extractor-fixtures"
    fixture_workspace = fixture_root / "audit-workspace"
    confirmed_dir = fixture_workspace / "confirmed"
    output_dir = fixture_workspace / "evidence/variant-analysis"
    output_dir.mkdir(parents=True, exist_ok=True)

    positive_bundle = confirmed_dir / "ssrf-import-url"
    fixture_local_path = "/" + "Users" + "/" + "fixture" + "/private/repo"
    write_extractor_fixture_bundle(positive_bundle, local_path_text=fixture_local_path)
    final_output = output_dir / "positive-seed.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(positive_bundle),
            "--output",
            str(final_output),
        ],
        plugin_root,
    )
    run([sys.executable, str(validator), "--variant-seed-card", str(final_output)], plugin_root)
    records = [json.loads(line) for line in final_output.read_text(encoding="utf-8").splitlines() if line.strip()]
    if len(records) != 1:
        raise SystemExit("FAILED: extractor should produce exactly one final seed card")
    card = records[0]
    if card.get("confirmed_bundle_path") != "confirmed/ssrf-import-url":
        raise SystemExit(f"FAILED: extractor emitted non-relative confirmed_bundle_path: {card}")
    if card.get("search_scope", {}).get("repository") != "same-target-repository":
        raise SystemExit(f"FAILED: extractor did not emit structured same-repository scope: {card}")
    if "all repositories" in json.dumps(card, ensure_ascii=False):
        raise SystemExit("FAILED: extractor preserved broad input search_scope into final output")
    if fixture_local_path in json.dumps(card, ensure_ascii=False):
        raise SystemExit("FAILED: extractor emitted local absolute path text")
    if "<local-absolute-path>" not in card.get("root_cause", ""):
        raise SystemExit("FAILED: extractor did not sanitize local absolute path text")

    historical_bundle = confirmed_dir / "historical-config-path-traversal"
    write_extractor_fixture_bundle(
        historical_bundle,
        finding_payload=historical_finding_fixture(local_path_text=fixture_local_path),
    )
    historical_output = output_dir / "historical-seeds.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(historical_bundle),
            "--output",
            str(historical_output),
        ],
        plugin_root,
    )
    run([sys.executable, str(validator), "--variant-seed-card", str(historical_output)], plugin_root)
    historical_records = [
        json.loads(line) for line in historical_output.read_text(encoding="utf-8").splitlines() if line.strip()
    ]
    historical_card = historical_records[0]
    if historical_card.get("root_cause") == "unknown" or historical_card.get("source_pattern") == "unknown":
        raise SystemExit(f"FAILED: historical natural-language analysis did not fill root/source: {historical_card}")
    if historical_card.get("sink_pattern") == "unknown" or "Sink/API" not in historical_card.get("sink_pattern", ""):
        raise SystemExit(f"FAILED: historical natural-language analysis did not fill safe sink hint: {historical_card}")
    historical_text = json.dumps(historical_card, ensure_ascii=False)
    if fixture_local_path in historical_text or "<local-absolute-path>" not in historical_text:
        raise SystemExit("FAILED: historical natural-language extraction did not sanitize local absolute path text")

    explicit_precedence_bundle = confirmed_dir / "explicit-precedence"
    explicit_payload = historical_finding_fixture()
    explicit_payload.update(
        {
            "root_cause": "Explicit root cause keeps the structured field before natural-language hints.",
            "source_pattern": "Authenticated attacker controls the explicit HTTP request body source.",
            "sink_pattern": "Explicit filesystem file read sink remains authoritative.",
        }
    )
    write_extractor_fixture_bundle(explicit_precedence_bundle, finding_payload=explicit_payload)
    explicit_output = output_dir / "explicit-precedence.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(explicit_precedence_bundle),
            "--output",
            str(explicit_output),
        ],
        plugin_root,
    )
    explicit_card = json.loads(explicit_output.read_text(encoding="utf-8").splitlines()[0])
    if explicit_card.get("root_cause") != explicit_payload["root_cause"]:
        raise SystemExit("FAILED: natural-language hints overwrote explicit root_cause")
    if explicit_card.get("source_pattern") != explicit_payload["source_pattern"]:
        raise SystemExit("FAILED: natural-language hints overwrote explicit source_pattern")
    if explicit_card.get("sink_pattern") != explicit_payload["sink_pattern"]:
        raise SystemExit("FAILED: natural-language hints overwrote explicit sink_pattern")

    code_context_sink_bundle = confirmed_dir / "code-context-sink"
    write_extractor_fixture_bundle(
        code_context_sink_bundle,
        finding_payload={
            "slug": "code-context-sink",
            "bug_class": "Path Traversal",
            "root_cause": "Path normalization validation is missing before loading a user-selected file.",
            "source_pattern": "Attacker-controlled CLI argument supplies the config path.",
            "analysis": [
                "触发路径：attacker CLI argument -> configPath -> resolve -> loader helper。",
                "缺失校验：missing canonical directory-boundary validation before the loader helper。",
            ],
            "code_context": [
                {
                    "summary_en": (
                        "loader helper calls require(path) and readFileSync(path), enabling dangerous "
                        "filesystem file read and code execution behavior."
                    )
                }
            ],
        },
    )
    code_context_output = output_dir / "code-context-sink.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(code_context_sink_bundle),
            "--output",
            str(code_context_output),
        ],
        plugin_root,
    )
    run([sys.executable, str(validator), "--variant-seed-card", str(code_context_output)], plugin_root)
    code_context_card = json.loads(code_context_output.read_text(encoding="utf-8").splitlines()[0])
    if "filesystem file read" not in code_context_card.get("sink_pattern", ""):
        raise SystemExit("FAILED: code_context summary_en did not fill sink hint")
    if code_context_card.get("propagation_pattern") == "unknown":
        raise SystemExit("FAILED: analysis trigger path did not fill propagation hint")

    missing_verification = confirmed_dir / "missing-verification"
    write_extractor_fixture_bundle(missing_verification, include_verification=False)
    run_expect_fail(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(missing_verification),
            "--output",
            str(output_dir / "missing.jsonl"),
        ],
        plugin_root,
        "missing verification-evidence.json",
    )

    blocked_bundle = confirmed_dir / "blocked-verification"
    write_extractor_fixture_bundle(blocked_bundle, verification_status="blocked_docker_unavailable")
    run_expect_fail(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(blocked_bundle),
            "--output",
            str(output_dir / "blocked.jsonl"),
        ],
        plugin_root,
        "verification_status must be confirmed_in_docker",
    )

    outside_bundle = fixture_root / "outside-bundle"
    write_extractor_fixture_bundle(outside_bundle)
    run_expect_fail(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(outside_bundle),
            "--output",
            str(output_dir / "outside.jsonl"),
        ],
        plugin_root,
        "bundle-dir must be inside workspace confirmed/ directory",
    )

    incomplete_bundle = confirmed_dir / "incomplete-source-sink"
    write_extractor_fixture_bundle(incomplete_bundle, complete_finding=False)
    incomplete_final = output_dir / "incomplete-final.jsonl"
    incomplete_draft = output_dir / "drafts.jsonl"
    incomplete_note = output_dir / "seed-incomplete-source-sink.md"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(incomplete_bundle),
            "--output",
            str(incomplete_final),
            "--allow-draft",
            "--draft-output",
            str(incomplete_draft),
            "--seed-note-output",
            str(incomplete_note),
        ],
        plugin_root,
    )
    if incomplete_final.exists():
        raise SystemExit("FAILED: incomplete extraction wrote an invalid final output")
    if not incomplete_draft.exists() or not incomplete_note.exists():
        raise SystemExit("FAILED: incomplete extraction did not write draft artifacts")
    run([sys.executable, str(validator), "--variant-seed-card", str(incomplete_draft), "--variant-seed-draft"], plugin_root)
    require_text(incomplete_note, "no variant can be confirmed from this note", "extractor draft note guardrail")
    require_text(incomplete_note, "Possible Unmapped Hints", "extractor draft note unmapped-hints section")

    bug_class_only_bundle = confirmed_dir / "bug-class-only-natural-language"
    write_extractor_fixture_bundle(
        bug_class_only_bundle,
        finding_payload={
            "slug": "bug-class-only-natural-language",
            "bug_class": "Path Traversal",
            "title": "Possible path traversal in config loading",
            "analysis": ["漏洞类型：Path Traversal", "Title: config path traversal"],
        },
    )
    bug_class_draft = output_dir / "bug-class-only-draft.jsonl"
    bug_class_note = output_dir / "seed-bug-class-only.md"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(bug_class_only_bundle),
            "--output",
            str(output_dir / "bug-class-only-final.jsonl"),
            "--allow-draft",
            "--draft-output",
            str(bug_class_draft),
            "--seed-note-output",
            str(bug_class_note),
        ],
        plugin_root,
    )
    bug_class_card = json.loads(bug_class_draft.read_text(encoding="utf-8").splitlines()[0])
    if bug_class_card.get("root_cause") != "unknown" or bug_class_card.get("source_pattern") != "unknown":
        raise SystemExit("FAILED: bug class/title-only analysis fabricated final seed fields")
    require_text(bug_class_note, "findings.json:analysis", "extractor bug-class-only unmapped source")

    bare_source_bundle = confirmed_dir / "bare-variable-source"
    bare_source_payload = historical_finding_fixture()
    bare_source_payload["analysis"] = [
        "根因：loader does not perform path normalization validation before loading the file.",
        "入口/可控输入：configPath",
        "危险函数/危险操作：dangerous filesystem file read sink via readFileSync(path).",
    ]
    bare_source_payload.pop("analysis_en", None)
    bare_source_payload.pop("code_context", None)
    write_extractor_fixture_bundle(bare_source_bundle, finding_payload=bare_source_payload)
    bare_source_draft = output_dir / "bare-source-draft.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(bare_source_bundle),
            "--output",
            str(output_dir / "bare-source-final.jsonl"),
            "--allow-draft",
            "--draft-output",
            str(bare_source_draft),
            "--seed-note-output",
            str(output_dir / "seed-bare-source.md"),
        ],
        plugin_root,
    )
    bare_source_card = json.loads(bare_source_draft.read_text(encoding="utf-8").splitlines()[0])
    if bare_source_card.get("source_pattern") != "unknown":
        raise SystemExit("FAILED: bare variable name filled source_pattern")

    bare_sink_bundle = confirmed_dir / "bare-file-path-sink"
    bare_sink_payload = historical_finding_fixture()
    bare_sink_payload["analysis"] = [
        "根因：loader does not perform path normalization validation before loading the file.",
        "入口/可控输入：attacker-controlled CLI argument supplies configPath.",
        "危险函数/危险操作：lib/yargs-parser.ts:655-692",
    ]
    bare_sink_payload.pop("analysis_en", None)
    bare_sink_payload.pop("code_context", None)
    write_extractor_fixture_bundle(bare_sink_bundle, finding_payload=bare_sink_payload)
    bare_sink_draft = output_dir / "bare-sink-draft.jsonl"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(bare_sink_bundle),
            "--output",
            str(output_dir / "bare-sink-final.jsonl"),
            "--allow-draft",
            "--draft-output",
            str(bare_sink_draft),
            "--seed-note-output",
            str(output_dir / "seed-bare-sink.md"),
        ],
        plugin_root,
    )
    bare_sink_card = json.loads(bare_sink_draft.read_text(encoding="utf-8").splitlines()[0])
    if bare_sink_card.get("sink_pattern") != "unknown":
        raise SystemExit("FAILED: bare file path filled sink_pattern")

    ambiguous_bundle = confirmed_dir / "ambiguous-natural-language"
    write_extractor_fixture_bundle(
        ambiguous_bundle,
        finding_payload={
            "slug": "ambiguous-natural-language",
            "bug_class": "Path Traversal",
            "analysis": [
                (
                    "Root cause / source / sink: attacker-controlled request path reaches dangerous "
                    "filesystem file read sink because validation is missing."
                )
            ],
        },
    )
    ambiguous_draft = output_dir / "ambiguous-draft.jsonl"
    ambiguous_note = output_dir / "seed-ambiguous.md"
    run(
        [
            sys.executable,
            str(extractor),
            "--workspace-dir",
            str(fixture_workspace),
            "--bundle-dir",
            str(ambiguous_bundle),
            "--output",
            str(output_dir / "ambiguous-final.jsonl"),
            "--allow-draft",
            "--draft-output",
            str(ambiguous_draft),
            "--seed-note-output",
            str(ambiguous_note),
        ],
        plugin_root,
    )
    ambiguous_card = json.loads(ambiguous_draft.read_text(encoding="utf-8").splitlines()[0])
    if ambiguous_card.get("root_cause") != "unknown" or ambiguous_card.get("source_pattern") != "unknown":
        raise SystemExit("FAILED: ambiguous multi-field hint fabricated root/source fields")
    require_text(ambiguous_note, "Root cause / source / sink", "extractor ambiguous unmapped hint")


VARIANT_CANDIDATE_REQUIRED_FIELDS = {
    "schema_version",
    "candidate_id",
    "variant_of",
    "bug_class",
    "file",
    "entry",
    "source_match",
    "sink_match",
    "root_cause_similarity",
    "negative_evidence",
    "rank",
    "score",
    "status",
    "recommended_next_step",
    "evidence_basis",
}


def valid_variant_candidate_record(overrides: dict | None = None, *, drop: str | None = None) -> dict:
    record = {
        "schema_version": 1,
        "candidate_id": "candidate-0123456789abcdef",
        "variant_of": "seed-confirmed-ssrf-import-url",
        "bug_class": "SSRF",
        "file": "src/routes/import.js",
        "entry": "route POST /import",
        "source_match": {
            "family": "source",
            "keyword": "req.body",
            "line": 4,
            "snippet": "const importUrl = req.body.url;",
        },
        "sink_match": {
            "family": "http-fetch",
            "keyword": "fetch(",
            "line": 5,
            "snippet": "const response = await fetch(importUrl);",
        },
        "root_cause_similarity": [
            "same sink family: http-fetch",
            "attacker-controlled source indicator present",
        ],
        "negative_evidence": [],
        "rank": 1,
        "score": 7,
        "status": "candidate",
        "recommended_next_step": (
            "Candidate only. Run independent Docker or Docker Compose verification "
            "and confirmed-bundle validation before any confirmation decision."
        ),
        "evidence_basis": ["source and sink signals matched in the same repository file"],
    }
    if overrides:
        record.update(overrides)
    if drop:
        record.pop(drop, None)
    return record


def write_variant_candidates_jsonl(path: Path, records: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "".join(json.dumps(record, ensure_ascii=False, sort_keys=True) + "\n" for record in records),
        encoding="utf-8",
    )


def write_finalization_variant_artifacts(workspace: Path) -> None:
    variant_dir = workspace / "evidence" / "variant-analysis"
    variant_dir.mkdir(parents=True, exist_ok=True)
    confirmed_bundles = sorted(
        path for path in (workspace / "confirmed").iterdir()
        if path.is_dir() and not path.name.startswith(".")
    )
    if not confirmed_bundles:
        raise SystemExit("FAILED: cannot write finalization seed fixture without a confirmed bundle")
    write_variant_seed_card(
        variant_dir / "seeds.jsonl",
        {"confirmed_bundle_path": f"confirmed/{confirmed_bundles[0].name}"},
    )
    write_variant_candidates_jsonl(
        variant_dir / "variant-candidates.jsonl",
        [valid_variant_candidate_record()],
    )


def exercise_variant_candidate_validation(plugin_root: Path, workspace: Path) -> None:
    validator = plugin_root / "scripts/validate_report_bundle.py"
    candidates_dir = workspace / "variant-candidate-validation-fixtures"
    candidates_dir.mkdir(parents=True, exist_ok=True)

    good_jsonl = candidates_dir / "variant-candidates.jsonl"
    write_variant_candidates_jsonl(good_jsonl, [valid_variant_candidate_record()])
    run([sys.executable, str(validator), "--variant-candidates", str(good_jsonl)], plugin_root)

    good_array = candidates_dir / "variant-candidates.json"
    good_array.write_text(
        json.dumps([valid_variant_candidate_record({"candidate_id": "candidate-fedcba9876543210"})], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    run([sys.executable, str(validator), "--variant-candidates", str(good_array)], plugin_root)

    bad_cases = [
        ("missing-required-field.jsonl", None, "sink_match", "missing required field(s): sink_match"),
        ("status-confirmed-in-docker.jsonl", {"status": "confirmed_in_docker"}, None, "status must be exactly candidate"),
        ("status-confirmed.jsonl", {"status": "confirmed"}, None, "status must be exactly candidate"),
        (
            "next-step-missing-bundle-validation.jsonl",
            {"recommended_next_step": "Run Docker verification later."},
            None,
            "recommended_next_step must require",
        ),
        (
            "absolute-file.jsonl",
            {"file": "/" + "Users" + "/" + "fixture" + "/repo/src/routes/import.js"},
            None,
            "operator-local absolute path",
        ),
        ("traversal-file.jsonl", {"file": "../src/routes/import.js"}, None, "path traversal"),
        ("confirmed-file.jsonl", {"file": "confirmed/ssrf-import-url/poc.js"}, None, "under confirmed/"),
        ("variant-analysis-file.jsonl", {"file": "evidence/variant-analysis/variant-candidates.jsonl"}, None, "evidence/variant-analysis"),
        ("zh-confirmation-text.jsonl", {"evidence_basis": ["漏洞已确认"]}, None, "forbidden confirmation/report wording"),
        ("en-confirmation-text.jsonl", {"evidence_basis": ["VULNERABILITY CONFIRMED"]}, None, "forbidden confirmation/report wording"),
        (
            "local-path-text.jsonl",
            {"evidence_basis": ["/" + "Users" + "/" + "fixture" + "/private/evidence.log"]},
            None,
            "operator-local absolute path",
        ),
        ("docx-generation.jsonl", {"evidence_basis": ["generate DOCX report for reviewer"]}, None, "forbidden confirmation/report wording"),
        ("confirmed-destination.jsonl", {"evidence_basis": ["write candidate to confirmed/output"]}, None, "forbidden confirmation/report wording"),
        (
            "similarity-confirms.jsonl",
            {"root_cause_similarity": ["similarity to seed is sufficient to confirm this issue"]},
            None,
            "seed similarity, score, or ranking is sufficient",
        ),
        (
            "ranking-docker-proof.jsonl",
            {"evidence_basis": ["candidate ranking is Docker proof for this issue"]},
            None,
            "seed similarity, score, or ranking is sufficient",
        ),
    ]
    for filename, overrides, drop, expected in bad_cases:
        path = candidates_dir / filename
        write_variant_candidates_jsonl(path, [valid_variant_candidate_record(overrides, drop=drop)])
        run_expect_fail([sys.executable, str(validator), "--variant-candidates", str(path)], plugin_root, expected)


def write_variant_candidate_fixture_sources(repo_root: Path, temp_root: Path) -> None:
    routes_dir = repo_root / "src/routes"
    routes_dir.mkdir(parents=True, exist_ok=True)
    (routes_dir / "import.js").write_text(
        "import express from 'express';\n"
        "const router = express.Router();\n"
        "router.post('/import', async (req, res) => {\n"
        "  const importUrl = req.body.url;\n"
        "  const response = await fetch(importUrl);\n"
        "  res.send(await response.text());\n"
        "});\n",
        encoding="utf-8",
    )
    tests_dir = repo_root / "tests"
    tests_dir.mkdir(parents=True, exist_ok=True)
    (tests_dir / "import.test.js").write_text(
        "test('fixture candidate stays out of ranked source output', async () => {\n"
        "  const importUrl = req.body.url;\n"
        "  await fetch(importUrl);\n"
        "});\n",
        encoding="utf-8",
    )
    (repo_root / "src/sink-only.js").write_text(
        "export async function pingInternalService() {\n"
        "  return fetch('https://example.invalid/status');\n"
        "}\n",
        encoding="utf-8",
    )
    (repo_root / "src/source-only.js").write_text(
        "export function readImportUrl(req) {\n"
        "  return req.body.url;\n"
        "}\n",
        encoding="utf-8",
    )
    outside_dir = temp_root / "outside-variant-source"
    outside_dir.mkdir(parents=True, exist_ok=True)
    (outside_dir / "escape.js").write_text(
        "export async function escaped(req) {\n"
        "  return fetch(req.body.url);\n"
        "}\n",
        encoding="utf-8",
    )
    symlink_path = repo_root / "linked-outside"
    try:
        symlink_path.symlink_to(outside_dir, target_is_directory=True)
    except OSError:
        pass


def assert_no_forbidden_candidate_language(record: dict) -> None:
    forbidden = ("confirmed_in_docker", "vulnerability confirmed", "漏洞已确认", "已确认")
    for key, value in record.items():
        if key == "variant_of":
            continue
        text = json.dumps(value, ensure_ascii=False).lower()
        if "confirmed" in text:
            raise SystemExit(f"FAILED: candidate field {key} contains final confirmation language: {record}")
        for needle in forbidden:
            if needle.lower() in text:
                raise SystemExit(f"FAILED: candidate field {key} contains forbidden confirmation wording: {record}")


def exercise_find_variant_candidates(plugin_root: Path, workspace: Path, seed_bundle_source: Path) -> None:
    finder = plugin_root / "scripts/find_variant_candidates.py"
    fixture_root = workspace / "variant-candidate-fixtures"
    repo_root = fixture_root / "repo"
    temp_root = fixture_root
    fixture_workspace = repo_root / "security-research-variant"
    fixture_workspace.mkdir(parents=True, exist_ok=True)
    (fixture_workspace / "asr-config.json").write_text(
        json.dumps(
            {
                "workspace_root": fixture_workspace.name,
                "workspace_created_at": "2026-06-20T00:00:00Z",
                "confirmed_output_dir": f"{fixture_workspace.name}/confirmed",
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )
    write_variant_candidate_fixture_sources(repo_root, temp_root)

    seed_bundle = fixture_workspace / "confirmed" / seed_bundle_source.name
    seed_bundle.parent.mkdir(parents=True, exist_ok=True)
    if seed_bundle.exists():
        shutil.rmtree(seed_bundle)
    shutil.copytree(seed_bundle_source, seed_bundle)
    variant_dir = fixture_workspace / "evidence/variant-analysis"
    variant_dir.mkdir(parents=True, exist_ok=True)
    seed_card = variant_dir / "seeds.jsonl"
    write_variant_seed_card(seed_card, {"confirmed_bundle_path": f"confirmed/{seed_bundle.name}"})
    output = variant_dir / "variant-candidates.jsonl"
    command = [
        sys.executable,
        str(finder),
        "--repo-root",
        str(repo_root),
        "--workspace-dir",
        str(fixture_workspace),
        "--seed-card",
        str(seed_card),
        "--seed-id",
        "seed-confirmed-ssrf-import-url",
        "--output",
        str(output),
        "--limit",
        "20",
        "--language",
        "nodejs",
    ]
    run(command, plugin_root)
    records = [json.loads(line) for line in output.read_text(encoding="utf-8").splitlines() if line.strip()]
    if not records:
        raise SystemExit("FAILED: variant candidate finder produced no candidates for a valid fixture")
    first_text = output.read_text(encoding="utf-8")
    run_expect_fail(command, plugin_root, "refusing to overwrite existing file without --force")
    run([*command, "--force"], plugin_root)
    if output.read_text(encoding="utf-8") != first_text:
        raise SystemExit("FAILED: variant candidate finder output is not deterministic across repeated runs")

    files = [record.get("file") for record in records]
    if files[0] != "src/routes/import.js":
        raise SystemExit(f"FAILED: expected real source candidate to rank first, got: {files}")
    forbidden_files = {"tests/import.test.js", "src/sink-only.js", "src/source-only.js", "linked-outside/escape.js"}
    if any(path in forbidden_files for path in files):
        raise SystemExit(f"FAILED: variant finder emitted excluded or source/sink-incomplete file: {files}")

    for index, record in enumerate(records, start=1):
        if set(record) != VARIANT_CANDIDATE_REQUIRED_FIELDS:
            raise SystemExit(f"FAILED: variant candidate output fields drifted: {record}")
        if record.get("schema_version") != 1:
            raise SystemExit(f"FAILED: candidate schema_version must be 1: {record}")
        if record.get("status") != "candidate":
            raise SystemExit(f"FAILED: variant candidate status must stay candidate: {record}")
        if record.get("variant_of") != "seed-confirmed-ssrf-import-url":
            raise SystemExit(f"FAILED: variant_of must equal the selected seed id: {record}")
        if Path(str(record.get("file", ""))).is_absolute():
            raise SystemExit(f"FAILED: candidate file must be repo-relative: {record}")
        if record.get("rank") != index or not isinstance(record.get("rank"), int) or record["rank"] < 1:
            raise SystemExit(f"FAILED: candidate rank must be a positive sorted integer: {record}")
        if not isinstance(record.get("score"), int) or record["score"] <= 0:
            raise SystemExit(f"FAILED: candidate score must be positive for emitted fixtures: {record}")
        next_step = str(record.get("recommended_next_step") or "")
        if "Docker" not in next_step or "verification" not in next_step or "before any confirmation" not in next_step:
            raise SystemExit(f"FAILED: candidate next step must require Docker verification before confirmation: {record}")
        assert_no_forbidden_candidate_language(record)

    missing_seed = variant_dir / "missing-seeds.jsonl"
    run_expect_fail(
        [
            sys.executable,
            str(finder),
            "--repo-root",
            str(repo_root),
            "--workspace-dir",
            str(fixture_workspace),
            "--seed-card",
            str(missing_seed),
            "--seed-id",
            "seed-confirmed-ssrf-import-url",
            "--output",
            str(variant_dir / "missing-output.jsonl"),
        ],
        plugin_root,
        "seed-card file does not exist",
    )

    invalid_seed = variant_dir / "invalid-seed.jsonl"
    write_variant_seed_card(invalid_seed, {"source_pattern": "import_url parameter"})
    run_expect_fail(
        [
            sys.executable,
            str(finder),
            "--repo-root",
            str(repo_root),
            "--workspace-dir",
            str(fixture_workspace),
            "--seed-card",
            str(invalid_seed),
            "--seed-id",
            "seed-confirmed-ssrf-import-url",
            "--output",
            str(variant_dir / "invalid-output.jsonl"),
        ],
        plugin_root,
        "source_pattern must describe attacker-controlled or untrusted input",
    )

    other_scope_seed = variant_dir / "other-scope-seed.jsonl"
    write_variant_seed_card(
        other_scope_seed,
        {
            "confirmed_bundle_path": f"confirmed/{seed_bundle.name}",
            "search_scope": {"repository": "different-target-repository"},
        },
    )
    run_expect_fail(
        [
            sys.executable,
            str(finder),
            "--repo-root",
            str(repo_root),
            "--workspace-dir",
            str(fixture_workspace),
            "--seed-card",
            str(other_scope_seed),
            "--seed-id",
            "seed-confirmed-ssrf-import-url",
            "--output",
            str(variant_dir / "other-scope-output.jsonl"),
        ],
        plugin_root,
        "search_scope.repository must be same-target-repository",
    )

    outside_bundle_seed = variant_dir / "outside-bundle-seed.jsonl"
    write_variant_seed_card(outside_bundle_seed, {"confirmed_bundle_path": "../confirmed/ssrf-import-url"})
    run_expect_fail(
        [
            sys.executable,
            str(finder),
            "--repo-root",
            str(repo_root),
            "--workspace-dir",
            str(fixture_workspace),
            "--seed-card",
            str(outside_bundle_seed),
            "--seed-id",
            "seed-confirmed-ssrf-import-url",
            "--output",
            str(variant_dir / "outside-bundle-output.jsonl"),
        ],
        plugin_root,
        "confirmed_bundle_path must be a normalized workspace-relative confirmed/<bundle> path",
    )

    outside_workspace = temp_root / "outside-audit-workspace"
    outside_workspace.mkdir(parents=True, exist_ok=True)
    run_expect_fail(
        [
            sys.executable,
            str(finder),
            "--repo-root",
            str(repo_root),
            "--workspace-dir",
            str(outside_workspace),
            "--seed-card",
            str(seed_card),
            "--seed-id",
            "seed-confirmed-ssrf-import-url",
            "--output",
            str(variant_dir / "outside-workspace-output.jsonl"),
        ],
        plugin_root,
        "workspace-dir must be inside repo-root",
    )


def exercise_variant_seed_confirmed_bundle_gate(plugin_root: Path, workspace: Path, valid_bundle: Path) -> None:
    validator = plugin_root / "scripts/validate_report_bundle.py"
    gate_dir = workspace / "evidence" / "variant-analysis"
    gate_dir.mkdir(parents=True, exist_ok=True)
    seed_path = gate_dir / "seeds.jsonl"

    def validate_seed(overrides: dict | None = None, *, expected: str | None = None, draft: bool = False) -> None:
        write_variant_seed_card(
            seed_path,
            {"confirmed_bundle_path": f"confirmed/{valid_bundle.name}", **(overrides or {})},
        )
        command = [
            sys.executable,
            str(validator),
            "--workspace-dir",
            str(workspace),
            "--variant-seed-card",
            str(seed_path),
        ]
        if draft:
            command.append("--variant-seed-draft")
        if expected is None:
            run(command, plugin_root)
        else:
            run_expect_fail(command, plugin_root, expected)

    validate_seed()

    bad_paths = [
        ("candidate-findings.md#C02", "markdown row or anchor"),
        ("manual-notes/seed.md", "candidate/manual/evidence-only material"),
        ("evidence/docker/ssrf-import-url", "candidate/manual/evidence-only material"),
        ("docker/ssrf-import-url", "candidate/manual/evidence-only material"),
        ("evidence/variant-analysis/manual-seed.jsonl", "candidate/manual/evidence-only material"),
    ]
    for bad_path, expected in bad_paths:
        validate_seed({"confirmed_bundle_path": bad_path}, expected=expected)

    zero_workspace = workspace / "issue17-zero-confirmed-workspace"
    zero_seed = zero_workspace / "evidence" / "variant-analysis" / "seeds.jsonl"
    zero_seed.parent.mkdir(parents=True, exist_ok=True)
    write_variant_seed_card(zero_seed, {"confirmed_bundle_path": f"confirmed/{valid_bundle.name}"})
    run_expect_fail(
        [
            sys.executable,
            str(validator),
            "--workspace-dir",
            str(zero_workspace),
            "--variant-seed-card",
            str(zero_seed),
        ],
        plugin_root,
        "confirmed bundle directory does not exist",
    )

    partial_bundle = workspace / "confirmed" / "issue17-partial-bundle"
    if partial_bundle.exists():
        shutil.rmtree(partial_bundle)
    partial_bundle.mkdir(parents=True)
    (partial_bundle / "verification-evidence.json").write_text(
        json.dumps({"verification_status": "confirmed_in_docker"}, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    validate_seed(
        {"confirmed_bundle_path": "confirmed/issue17-partial-bundle"},
        expected="must reference a validated confirmed bundle",
    )

    validation_failed_bundle = workspace / "confirmed" / "issue17-validation-failed-bundle"
    if validation_failed_bundle.exists():
        shutil.rmtree(validation_failed_bundle)
    shutil.copytree(valid_bundle, validation_failed_bundle)
    verification_path = validation_failed_bundle / "verification-evidence.json"
    verification = json.loads(verification_path.read_text(encoding="utf-8"))
    verification["verification_status"] = "blocked_docker_unavailable"
    verification_path.write_text(json.dumps(verification, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    validate_seed(
        {"confirmed_bundle_path": "confirmed/issue17-validation-failed-bundle"},
        expected="must reference a validated confirmed bundle",
    )

    validate_seed({"root_cause": "unknown"}, expected="must not use unknown")
    validate_seed({"root_cause": "unknown"}, draft=True)
    for transient in (partial_bundle, validation_failed_bundle):
        if transient.exists():
            shutil.rmtree(transient)


def iter_json_strings(value) -> list[str]:
    if isinstance(value, str):
        return [value]
    if isinstance(value, dict):
        strings: list[str] = []
        for nested in value.values():
            strings.extend(iter_json_strings(nested))
        return strings
    if isinstance(value, list):
        strings = []
        for nested in value:
            strings.extend(iter_json_strings(nested))
        return strings
    return []


def is_empty_manifest_component(value) -> bool:
    if value is None or value is False or value == "":
        return True
    if isinstance(value, (list, dict)) and not value:
        return True
    return False


def require_relative_manifest_path(
    plugin_root: Path,
    manifest: dict,
    field: str,
    *,
    must_be_file: bool = False,
) -> Path:
    value = manifest.get(field)
    if not isinstance(value, str) or not value.strip():
        raise SystemExit(f"FAILED: Claude plugin manifest missing relative path field: {field}")
    token = value.strip()
    if "://" in token:
        raise SystemExit(f"FAILED: Claude plugin manifest path must be relative, not URL: {field}={token}")
    if Path(token).is_absolute() or token.startswith("~"):
        raise SystemExit(f"FAILED: Claude plugin manifest path must be relative: {field}={token}")
    candidate = (plugin_root / token).resolve()
    try:
        candidate.relative_to(plugin_root.resolve())
    except ValueError as exc:
        raise SystemExit(f"FAILED: Claude plugin manifest path escapes package: {field}={token}") from exc
    if must_be_file and not candidate.is_file():
        raise SystemExit(f"FAILED: Claude plugin manifest file path does not exist: {field}={token}")
    if not must_be_file and not candidate.exists():
        raise SystemExit(f"FAILED: Claude plugin manifest path does not exist: {field}={token}")
    return candidate


def validate_claude_plugin_manifest(plugin_root: Path) -> None:
    manifest_path = plugin_root / ".claude-plugin/plugin.json"
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise SystemExit(f"FAILED: invalid Claude plugin manifest JSON: {manifest_path}: {exc}") from exc

    if not isinstance(manifest, dict):
        raise SystemExit("FAILED: Claude plugin manifest must be a JSON object")
    if manifest.get("name") != "zhulong":
        raise SystemExit("FAILED: Claude plugin manifest name must be zhulong")
    for field in ("version", "displayName", "description"):
        if not isinstance(manifest.get(field), str) or not manifest[field].strip():
            raise SystemExit(f"FAILED: Claude plugin manifest missing metadata field: {field}")
    description = manifest["description"].lower()
    if "docker" not in description or "audit" not in description:
        raise SystemExit("FAILED: Claude plugin manifest description must describe Docker audit workflow")

    for field in ("skills", "scripts", "assets"):
        require_relative_manifest_path(plugin_root, manifest, field)
    if not (plugin_root / "skills/zhulong/SKILL.md").is_file():
        raise SystemExit("FAILED: Claude plugin package missing skills/zhulong/SKILL.md")
    if not (plugin_root / "scripts").is_dir():
        raise SystemExit("FAILED: Claude plugin package missing scripts/")
    if not (plugin_root / "assets").is_dir():
        raise SystemExit("FAILED: Claude plugin package missing assets/")

    runtime = manifest.get("runtime")
    if runtime is not None:
        if not isinstance(runtime, dict):
            raise SystemExit("FAILED: Claude plugin manifest runtime must be metadata object")
        if runtime.get("entrypoint"):
            require_relative_manifest_path(
                plugin_root,
                runtime,
                "entrypoint",
                must_be_file=True,
            )

    for text in iter_json_strings(manifest):
        if text.startswith("file://") or text.startswith("/"):
            raise SystemExit(f"FAILED: Claude plugin manifest contains an absolute local path: {text}")
        if re.match(r"^[A-Za-z]:[\\/]", text):
            raise SystemExit(f"FAILED: Claude plugin manifest contains a Windows absolute path: {text}")
        operator_local_path = "/" + "Users" + "/" + "localuser"
        if operator_local_path in text:
            raise SystemExit("FAILED: Claude plugin manifest contains operator-local absolute path")

    forbidden_component_fields = (
        "hooks",
        "mcpServers",
        "mcp_servers",
        "apps",
        "agents",
        "commands",
        "services",
        "platformServices",
        "backgroundServices",
        "daemons",
    )
    for field in forbidden_component_fields:
        if field in manifest and not is_empty_manifest_component(manifest[field]):
            raise SystemExit(
                "FAILED: Claude plugin manifest must not declare runtime component "
                f"{field!r}; Zhulong remains skill-and-scripts only"
            )


def exercise_agents_shim(plugin_root: Path) -> None:
    shim_path = plugin_root / "AGENTS.md"
    if not shim_path.is_file():
        raise SystemExit("FAILED: missing repo-root AGENTS.md shim")

    content = shim_path.read_text(encoding="utf-8")
    normalized_content = re.sub(r"\s+", " ", content)
    lines = content.splitlines()
    if len(lines) > 80:
        raise SystemExit(f"FAILED: repo-root AGENTS.md must stay a short shim, got {len(lines)} lines")

    required_texts = [
        "# Zhulong Agent Shim",
        "lightweight instruction shim",
        "docs/AGENTS.md",
        "CONTRIBUTING.md",
        "docs/RELEASE_CHECKLIST.md",
        "Treat this plugin source tree as canonical",
        "Installed Claude and Codex skill directories are generated runtime copies",
        "When the user asks for repository-level security audit",
        "Docker-based PoC reproduction",
        "confirmed vulnerability bundles",
        "seeded variant discovery",
        "use `$zhulong`",
        "Do not duplicate the full `$zhulong` skill contract here.",
        "Scanner, static-analysis, LLM, and dependency findings remain candidates until",
        "attacker-entrypoint Docker reproduction, source-bound validity checks, and",
        "confirmed-bundle validation support them.",
        "Do not execute PoC or exploit verification directly on the host.",
        "Docker / Docker Compose verification flow",
        "Confirmed findings must live only in validated Zhulong confirmed bundles.",
        "Final recording is optional",
        "ordinary confirmed status does not imply recording readiness.",
        "Do not use broad Docker prune or PID kill behavior.",
        "skills/zhulong/SKILL.md",
        "assets/references/",
        "docs/WORKFLOW_DETAILS.md",
        "docs/WORKFLOW_DETAILS.zh-CN.md",
        "docs/CODEX_SKILL_ADAPTATION.md",
        "deterministic validators/selftests",
    ]
    for needle in required_texts:
        normalized_needle = re.sub(r"\s+", " ", needle)
        if needle not in content and normalized_needle not in normalized_content:
            raise SystemExit(f"FAILED: repo-root AGENTS.md missing required shim text: {needle}")

    forbidden_skill_contract_headings = [
        "## Installed Skill Runtime Contents",
        "## Plugin-Owned Hard Constraints",
        "## Standard Execution Order",
    ]
    for heading in forbidden_skill_contract_headings:
        if heading in content:
            raise SystemExit(f"FAILED: repo-root AGENTS.md duplicated skill contract heading: {heading}")

    forbidden_command_patterns = [
        r"docker\s+(system|builder|buildx)\s+prune",
        r"\b(builder|system|buildx)\s+prune\b",
        r"kill\s+-(TERM|9|KILL)",
        r"\bSIG" + r"KILL\b",
        r"cleanup-suspect-pid\s+.*--apply",
    ]
    for pattern in forbidden_command_patterns:
        if re.search(pattern, content, flags=re.IGNORECASE):
            raise SystemExit(f"FAILED: repo-root AGENTS.md contains forbidden command pattern: {pattern}")

    platform_terms = [
        "mcp server",
        "daemon",
        "dashboard",
        "database",
        "rag",
        "vector store",
        "hooks",
        "rules",
        "marketplace",
        "docker socket",
    ]
    requirement_words = r"(?:require|requires|required|requiring|mandatory|must)"
    for term in platform_terms:
        escaped = re.escape(term)
        patterns = [
            rf"\b{requirement_words}\b[^\n.]*\b{escaped}\b",
            rf"\b{escaped}\b[^\n.]*\b{requirement_words}\b",
        ]
        for pattern in patterns:
            if re.search(pattern, content, flags=re.IGNORECASE):
                raise SystemExit(
                    "FAILED: repo-root AGENTS.md must not make platform/service "
                    f"claims mandatory: {term}"
                )


def exercise_p7_wording_closure(plugin_root: Path) -> None:
    skill_path = plugin_root / "skills/zhulong/SKILL.md"
    template_path = plugin_root / "templates/claude-skill/SKILL.md"
    skill_text = skill_path.read_text(encoding="utf-8")
    template_text = template_path.read_text(encoding="utf-8")
    if skill_text != template_text:
        raise SystemExit("FAILED: skills/zhulong/SKILL.md and template SKILL.md must stay identical")

    for path in (skill_path, template_path):
        forbid_text(path, "Use this Claude Code skill when", "P7.5 local-agent-neutral skill opening")
        require_text(path, "Use this local-agent skill when", "P7.5 local-agent-neutral skill opening")

    repo_prep = plugin_root / "assets/references/repo-preparation.md"
    forbid_text(
        repo_prep,
        '$HOME/.claude/skills/zhulong/scripts/asr_start.sh',
        "repo preparation primary launcher must be platform-neutral",
    )
    require_text(
        repo_prep,
        "bash <skill-root>/scripts/zhulong_audit.sh --source <local-path-or-repo-url>",
        "repo preparation platform-neutral launcher",
    )

    invocation_template = plugin_root / "assets/references/claude-code-invocation-template.md"
    forbid_text(
        invocation_template,
        "bash scripts/asr_start.sh --source <repo-or-url>",
        "manual fallback must use platform-neutral launcher",
    )
    require_text(
        invocation_template,
        "bash scripts/zhulong_audit.sh --source <repo-or-url>",
        "manual fallback platform-neutral launcher",
    )

    future_only_phrases = [
        "P7.1 defines only the layout and sync contract",
        "P7.2 must add Codex",
        "Codex 适配在 P7.1 只定义布局和同步契约",
        "P7.2 需要实现 Codex",
        "A future phase may add first-class repo-scoped install guidance",
    ]
    for rel in [
        "docs/CODEX_SKILL_ADAPTATION.md",
        "docs/WORKFLOW_DETAILS.md",
        "docs/WORKFLOW_DETAILS.zh-CN.md",
        "docs/INSTALL.md",
        "docs/USAGE.md",
        "docs/USAGE.zh-CN.md",
        "docs/RELEASE_CHECKLIST.md",
        "README.md",
        "README.zh-CN.md",
    ]:
        path = plugin_root / rel
        for phrase in future_only_phrases:
            forbid_text(path, phrase, f"P7 completed-status wording in {rel}")

    require_text(
        plugin_root / "docs/CODEX_SKILL_ADAPTATION.md",
        "Current Codex support status",
        "Codex support current status",
    )
    require_text(
        plugin_root / "docs/CODEX_SKILL_ADAPTATION.md",
        "Source, Claude installed, and Codex installed regression checks are part of",
        "Codex support release validation status",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.md",
        "Codex user-level skill support is also available",
        "workflow details Codex installed runtime support",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.zh-CN.md",
        "Codex 用户级 Skill 也已支持",
        "Chinese workflow details Codex installed runtime support",
    )
    require_text(
        plugin_root / "README.md",
        "Codex user-level skill support with installed selftest, platform-neutral launcher",
        "README Codex completed support",
    )
    forbid_text(
        plugin_root / "README.md",
        "including Codex and Cursor",
        "README planned Codex support must be completed",
    )
    require_text(
        plugin_root / "README.zh-CN.md",
        "Codex 用户级 Skill 支持、安装目录自检、平台无关启动入口",
        "Chinese README Codex completed support",
    )
    forbid_text(
        plugin_root / "README.zh-CN.md",
        "例如 Codex 和 Cursor",
        "Chinese README planned Codex support must be completed",
    )


def require_probe_record(
    summary_path: Path,
    output_dir: Path,
    probe_name: str,
    expected_status: str,
    expected_exit_code: int | None,
    reason_snippet: str,
    forbidden_reason_snippet: str = "",
) -> dict:
    summary_data = json.loads(summary_path.read_text(encoding="utf-8"))
    probes = summary_data.get("probes") or []
    probe = next((item for item in probes if item.get("name") == probe_name), None)
    if probe is None:
        raise SystemExit(f"FAILED: missing probe record for {probe_name}")
    if probe.get("status") != expected_status:
        raise SystemExit(
            f"FAILED: {probe_name} status mismatch: "
            f"expected {expected_status}, got {probe.get('status')}"
        )
    if probe.get("exit_code") != expected_exit_code:
        raise SystemExit(
            f"FAILED: {probe_name} exit_code mismatch: "
            f"expected {expected_exit_code}, got {probe.get('exit_code')}"
        )
    reason = str(probe.get("reason") or "")
    if reason_snippet not in reason:
        raise SystemExit(f"FAILED: {probe_name} reason missing expected text: {reason_snippet}")
    if forbidden_reason_snippet and forbidden_reason_snippet in reason:
        raise SystemExit(f"FAILED: {probe_name} reason contains forbidden text: {forbidden_reason_snippet}")

    status_path = output_dir / f"{probe_name}.status"
    if not status_path.exists():
        raise SystemExit(f"FAILED: missing probe status file for {probe_name}: {status_path}")
    status_lines = status_path.read_text(encoding="utf-8").splitlines()
    if not status_lines or status_lines[0] != expected_status:
        raise SystemExit(f"FAILED: {probe_name}.status does not match summary status")
    status_exit = None
    for line in status_lines[1:]:
        if line.startswith("exit_code="):
            status_exit = int(line.split("=", 1)[1])
            break
    if expected_exit_code is not None and status_exit != expected_exit_code:
        raise SystemExit(f"FAILED: {probe_name}.status exit_code does not match summary exit_code")
    return probe


RUNTIME_STATUS_FIELDS = {
    "checked_at",
    "recommended_mode",
    "teams_enabled",
    "suspect_teammate_pids",
    "suspect_teammate_processes",
    "stale_swarm_sockets",
    "live_swarm_sockets",
    "ignored_current_session_teammate_pids",
    "ignored_current_session_teammate_processes",
    "cleanup_actions",
    "attempt_history",
    "heartbeat_seen",
    "resume_step",
    "unresolved_review_only",
    "clean",
}


def require_runtime_status_shape(status: dict, label: str) -> None:
    missing = sorted(RUNTIME_STATUS_FIELDS - set(status))
    if missing:
        raise SystemExit(f"FAILED: runtime hygiene status missing fields for {label}: {missing}")
    if status.get("recommended_mode") not in {"native_team_ready", "cleanup_needed", "single_agent_only"}:
        raise SystemExit(f"FAILED: invalid runtime hygiene mode for {label}: {status.get('recommended_mode')}")
    for key in (
        "suspect_teammate_pids",
        "suspect_teammate_processes",
        "stale_swarm_sockets",
        "live_swarm_sockets",
        "ignored_current_session_teammate_pids",
        "ignored_current_session_teammate_processes",
        "cleanup_actions",
        "attempt_history",
        "unresolved_review_only",
    ):
        if not isinstance(status.get(key), list):
            raise SystemExit(f"FAILED: runtime hygiene status field must be a list for {label}: {key}")


def run_omc_runtime_mock(
    script_path: Path,
    workspace: Path,
    plugin_root: Path,
    *,
    args: list[str] | None = None,
    env: dict[str, str] | None = None,
    expected_returncode: int = 0,
) -> dict:
    command = [
        "bash",
        str(script_path),
        "--workspace-dir",
        str(workspace),
        "--json",
    ]
    if args:
        command.extend(args)
    output = run_capture_with_env(
        command,
        plugin_root,
        {
            "ZHULONG_OMC_MOCK_TEAMS_ENABLED": "1",
            **(env or {}),
        },
        expected_returncode=expected_returncode,
    )
    try:
        status = json.loads(output)
    except json.JSONDecodeError as exc:
        raise SystemExit(f"FAILED: OMC runtime helper did not emit JSON for {script_path}: {output}") from exc
    require_runtime_status_shape(status, script_path.name)
    status_path = workspace / "runtime/runtime-hygiene-status.json"
    if not status_path.exists():
        raise SystemExit(f"FAILED: OMC runtime helper did not write status file: {status_path}")
    stored = json.loads(status_path.read_text(encoding="utf-8"))
    require_runtime_status_shape(stored, f"stored {script_path.name}")
    return status


def exercise_omc_runtime_hygiene(script_path: Path, workspace: Path, plugin_root: Path) -> None:
    def assert_no_teammate_signal_actions(status: dict, label: str) -> None:
        forbidden = {
            item.get("status")
            for item in status.get("cleanup_actions", [])
            if item.get("kind") == "cleanup_suspect_pid"
        }
        if forbidden & {"term_sent", "terminated"}:
            raise SystemExit(f"FAILED: OMC teammate PID cleanup recorded signal action for {label}: {status}")

    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        env={"ZHULONG_OMC_MOCK_TEAMMATE_RECORDS": "12345|111|222|333|ttys123|S+|claude --teammate-mode tmux audit"},
    )
    if status.get("recommended_mode") != "cleanup_needed" or status.get("suspect_teammate_pids") != ["12345"]:
        raise SystemExit(f"FAILED: OMC runtime helper did not report exact suspect PID: {status}")
    process = (status.get("suspect_teammate_processes") or [{}])[0]
    if process.get("tty") != "ttys123" or process.get("active_session_uncertain") is not True:
        raise SystemExit(f"FAILED: OMC runtime helper did not preserve process metadata: {status}")
    if not any(item.get("kind") == "suspect_teammate_pid" for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC runtime helper must record suspect teammate PIDs as review-only")
    assert_no_teammate_signal_actions(status, "initial suspect report")
    handoff = workspace / "handoff-summary.md"
    require_text(handoff, "## OMC Runtime Hygiene", "handoff runtime hygiene section")
    require_text(handoff, "`12345`", "handoff suspect PID")
    require_text(handoff, "review-only", "handoff review-only teammate PID guidance")
    require_text(handoff, "ttys123", "handoff suspect PID process metadata")
    forbidden_apply = "--" + "cleanup-suspect-pid 12345 --" + "apply"
    forbid_text(handoff, forbidden_apply, "handoff must not recommend teammate PID cleanup apply")

    kill_log = workspace / "runtime/mock-kill-dry-run.log"
    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-suspect-pid", "12346"],
        env={
            "ZHULONG_OMC_MOCK_TEAMMATE_RECORDS": "12346|claude --teammate-mode tmux audit",
            "ZHULONG_OMC_MOCK_PID_EXISTS": "12346",
            "ZHULONG_OMC_MOCK_KILL_LOG": str(kill_log),
        },
    )
    if kill_log.exists() and kill_log.read_text(encoding="utf-8").strip():
        raise SystemExit("FAILED: OMC exact PID dry-run sent a signal")
    if status.get("cleanup_actions"):
        raise SystemExit(f"FAILED: OMC teammate PID review-only dry-run must not create cleanup actions: {status}")
    if not any(item.get("kind") == "suspect_teammate_pid" and "review-only" in item.get("reason", "") for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC exact PID dry-run must be unresolved review-only")
    assert_no_teammate_signal_actions(status, "review-only dry-run")

    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-suspect-pid", "12347", "--apply"],
        env={
            "ZHULONG_OMC_MOCK_PID_EXISTS": "12347",
            "ZHULONG_OMC_MOCK_CMDLINE_RECORDS": "12347|python unrelated.py",
        },
        expected_returncode=1,
    )
    if not any(item.get("status") == "refused" and "command line" in item.get("reason", "") for item in status.get("cleanup_actions", [])):
        raise SystemExit("FAILED: OMC exact PID cleanup must refuse command-line mismatch")
    assert_no_teammate_signal_actions(status, "command-line mismatch")

    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-suspect-pid", "12348", "--apply"],
        env={
            "ZHULONG_OMC_MOCK_TEAMMATE_RECORDS": "12348|claude --teammate-mode tmux audit",
            "ZHULONG_OMC_MOCK_PID_EXISTS": "12348",
            "ZHULONG_OMC_MOCK_LIVE_SOCKETS": "/private/tmp/tmux-501/claude-swarm-live",
        },
        expected_returncode=1,
    )
    if not any(item.get("kind") == "live_swarm_socket" for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC exact PID cleanup must refuse when a live swarm socket exists")
    assert_no_teammate_signal_actions(status, "live swarm socket")

    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-suspect-pid", "12349", "--apply"],
        env={
            "ZHULONG_OMC_MOCK_TEAMMATE_RECORDS": "12349|claude --teammate-mode tmux audit",
            "ZHULONG_OMC_MOCK_PID_EXISTS": "12349",
            "ZHULONG_OMC_MOCK_CURRENT_SESSION_TEAMMATE_PIDS": "12349",
        },
        expected_returncode=1,
    )
    if status.get("ignored_current_session_teammate_pids") != ["12349"]:
        raise SystemExit("FAILED: OMC runtime helper must record ignored current-session teammate PIDs")
    ignored_process = (status.get("ignored_current_session_teammate_processes") or [{}])[0]
    if ignored_process.get("active_session_uncertain") is not False:
        raise SystemExit(f"FAILED: current-session teammate metadata must be marked certain/protected: {status}")
    if not any(item.get("kind") == "current_session_teammate_pid" for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC exact PID cleanup must refuse current-session teammate PIDs")
    assert_no_teammate_signal_actions(status, "current-session PID")

    kill_log = workspace / "runtime/mock-kill-apply.log"
    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-suspect-pid", "12350", "--apply"],
        env={
            "ZHULONG_OMC_MOCK_TEAMMATE_RECORDS": "12350|claude --teammate-mode tmux audit",
            "ZHULONG_OMC_MOCK_PID_EXISTS": "12350",
            "ZHULONG_OMC_MOCK_KILL_LOG": str(kill_log),
        },
        expected_returncode=1,
    )
    if kill_log.exists() and kill_log.read_text(encoding="utf-8").strip():
        raise SystemExit("FAILED: OMC exact PID cleanup with --apply must not send a signal")
    if not any("refused --apply" in item.get("reason", "") for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC exact PID cleanup with --apply must be refused review-only")
    assert_no_teammate_signal_actions(status, "apply refused teammate PID")

    socket_log = workspace / "runtime/mock-socket-cleanup.log"
    status = run_omc_runtime_mock(
        script_path,
        workspace,
        plugin_root,
        args=["--cleanup-stale"],
        env={
            "ZHULONG_OMC_MOCK_STALE_SOCKETS": "/tmp/claude-swarm-stale\n/tmp/not-omc.sock",
            "ZHULONG_OMC_MOCK_SOCKET_CLEANUP_LOG": str(socket_log),
        },
    )
    socket_text = socket_log.read_text(encoding="utf-8")
    if "REMOVE_SOCKET /tmp/claude-swarm-stale" not in socket_text or "not-omc.sock" in socket_text:
        raise SystemExit("FAILED: OMC stale socket cleanup must remove only stale claude-swarm sockets")
    if not any(item.get("kind") == "stale_swarm_socket" and "non-claude-swarm" in item.get("reason", "") for item in status.get("unresolved_review_only", [])):
        raise SystemExit("FAILED: OMC stale socket cleanup must review-only non claude-swarm mock sockets")

    run_expect_fail(
        ["bash", str(script_path), "--force-kill-suspect-teammates"],
        plugin_root,
        "Refusing deprecated --force-kill-suspect-teammates",
    )


def run_sandbox_preflight(
    script_path: Path,
    workspace: Path,
    plugin_root: Path,
    args: list[str],
    *,
    expected_returncode: int,
) -> dict:
    status_path = workspace / "runtime/sandbox-preflight-status.json"
    status_before = status_path.read_bytes() if status_path.exists() else None
    output = run_capture_with_env(
        [
            sys.executable,
            str(script_path),
            "--workspace-dir",
            str(workspace),
            "--case-id",
            "sandbox-selftest",
            "--json",
            *args,
        ],
        plugin_root,
        {},
        expected_returncode=expected_returncode,
    )
    try:
        status = json.loads(output)
    except json.JSONDecodeError as exc:
        raise SystemExit(f"FAILED: sandbox preflight did not emit JSON: {output}") from exc
    for key in ("checked_at", "ok", "status", "findings", "labels", "resume_step", "review_only"):
        if key not in status:
            raise SystemExit(f"FAILED: sandbox preflight missing status field: {key}")
    if expected_returncode == 0 and not status_path.exists():
        raise SystemExit("FAILED: sandbox preflight did not write runtime/sandbox-preflight-status.json")
    if expected_returncode != 0:
        status_after = status_path.read_bytes() if status_path.exists() else None
        if status_after != status_before:
            raise SystemExit("FAILED: rejected sandbox preflight wrote runtime status before Docker execution")
    return status


def require_sandbox_rejection(status: dict, pattern: str, label: str) -> None:
    if status.get("status") != "rejected_unsafe_sandbox" or status.get("ok") is not False:
        raise SystemExit(f"FAILED: sandbox preflight should reject {label}: {status}")
    if not any(item.get("pattern") == pattern for item in status.get("findings", [])):
        raise SystemExit(f"FAILED: sandbox preflight missing pattern {pattern} for {label}: {status}")
    if not status.get("review_only"):
        raise SystemExit(f"FAILED: rejected sandbox status must be review-only for {label}")


def exercise_sandbox_preflight(script_path: Path, workspace: Path, plugin_root: Path) -> None:
    fixtures = workspace / "sandbox-preflight-fixtures"
    fixtures.mkdir(parents=True, exist_ok=True)

    compose_privileged = fixtures / "privileged.yml"
    compose_privileged.write_text("services:\n  app:\n    image: alpine\n    privileged: true\n", encoding="utf-8")
    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(compose_privileged)], expected_returncode=1),
        "privileged_true",
        "Compose privileged:true",
    )

    compose_host_network = fixtures / "host-network.yml"
    compose_host_network.write_text("services:\n  app:\n    image: alpine\n    network_mode: host\n", encoding="utf-8")
    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(compose_host_network)], expected_returncode=1),
        "network_mode_host",
        "Compose network_mode:host",
    )

    compose_host_pid = fixtures / "host-pid.yml"
    compose_host_pid.write_text("services:\n  app:\n    image: alpine\n    pid: host\n", encoding="utf-8")
    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(compose_host_pid)], expected_returncode=1),
        "pid_host",
        "Compose pid:host",
    )

    compose_sock = fixtures / "docker-sock.yml"
    compose_sock.write_text(
        "services:\n  app:\n    image: alpine\n    volumes:\n      - /var/run/docker.sock:/var/run/docker.sock\n",
        encoding="utf-8",
    )
    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(compose_sock)], expected_returncode=1),
        "docker_socket_mount",
        "Compose Docker socket mount",
    )

    compose_root = fixtures / "root-mount.yml"
    compose_root.write_text("services:\n  app:\n    image: alpine\n    volumes:\n      - /:/host:ro\n", encoding="utf-8")
    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(compose_root)], expected_returncode=1),
        "host_root_mount",
        "Compose host root mount",
    )

    script_sock_root = fixtures / "unsafe-run.sh"
    script_sock_root.write_text(
        "#!/usr/bin/env bash\n"
        "docker run --rm -v /var/run/docker.sock:/var/run/docker.sock -v /:/host alpine true\n",
        encoding="utf-8",
    )
    status = run_sandbox_preflight(script_path, workspace, plugin_root, ["--shell-script", str(script_sock_root)], expected_returncode=1)
    require_sandbox_rejection(status, "docker_socket_mount", "script Docker socket mount")
    require_sandbox_rejection(status, "host_root_mount", "script host root mount")

    require_sandbox_rejection(
        run_sandbox_preflight(script_path, workspace, plugin_root, ["--docker-run-arg=--privileged"], expected_returncode=1),
        "docker_run_privileged",
        "docker run --privileged",
    )
    require_sandbox_rejection(
        run_sandbox_preflight(
            script_path,
            workspace,
            plugin_root,
            ["--docker-run-arg=--network", "--docker-run-arg=host"],
            expected_returncode=1,
        ),
        "docker_run_network_host",
        "docker run --network host",
    )
    require_sandbox_rejection(
        run_sandbox_preflight(
            script_path,
            workspace,
            plugin_root,
            ["--docker-run-arg=--pid=host"],
            expected_returncode=1,
        ),
        "docker_run_pid_host",
        "docker run --pid host",
    )

    safe_compose = fixtures / "safe-attacker.yml"
    safe_compose.write_text(
        "services:\n"
        "  attacker:\n"
        "    image: alpine:3.20\n"
        "    privileged: false\n"
        "    labels:\n"
        "      org.zhulong.managed: \"true\"\n"
        "    cap_drop:\n"
        "      - ALL\n"
        "    security_opt:\n"
        "      - no-new-privileges:true\n"
        "    network_mode: bridge\n"
        "    volumes:\n"
        "      - type: bind\n"
        "        source: ./poc\n"
        "        target: /workspace/poc\n"
        "        read_only: true\n",
        encoding="utf-8",
    )
    status = run_sandbox_preflight(script_path, workspace, plugin_root, ["--compose-file", str(safe_compose)], expected_returncode=0)
    if status.get("status") != "passed" or status.get("findings"):
        raise SystemExit(f"FAILED: safe Zhulong attacker compose should pass sandbox preflight: {status}")


def exercise_runner_sandbox_rejection(run_script: Path, workspace: Path, plugin_root: Path) -> None:
    writer = plugin_root / "scripts/write_audit_event.py"

    def write_transition(name: str, stage: str, transition_kind: str) -> None:
        run(
            [
                sys.executable,
                str(writer),
                "--workspace-dir",
                str(workspace),
                "--event",
                name,
                "--stage",
                stage,
                "--status",
                "running",
                "--transition-kind",
                transition_kind,
                "--message",
                "Seed the bounded sandbox-rejection policy path.",
                "--accept-current-revision",
            ],
            plugin_root,
        )

    state_path = workspace / "stage-status.json"
    if not state_path.exists():
        write_transition("sandbox_policy_started", "intake", "start")
    state = json.loads(state_path.read_text(encoding="utf-8"))
    if state.get("stage") != "intake" or state.get("status") != "running":
        raise SystemExit("FAILED: sandbox-rejection selftest requires an intake/running workspace")
    for name, stage in [
        ("sandbox_recon_started", "recon"),
        ("sandbox_candidates_started", "candidate_generation"),
        ("sandbox_triage_started", "triage"),
        ("sandbox_verification_started", "verification"),
    ]:
        write_transition(name, stage, "advance")

    fakebin = workspace / "fakebin"
    fakebin.mkdir(parents=True, exist_ok=True)
    docker_log = workspace / "runtime/docker-called.log"
    fake_docker = fakebin / "docker"
    fake_docker.write_text(
        "#!/usr/bin/env bash\n"
        "echo docker-called \"$@\" >> \"$ZHULONG_DOCKER_CALL_LOG\"\n"
        "exit 0\n",
        encoding="utf-8",
    )
    fake_docker.chmod(0o755)

    status_path = workspace / "runtime/sandbox-preflight-status.json"
    status_before = status_path.read_bytes() if status_path.exists() else None

    output = run_capture_with_env(
        [
            "bash",
            str(run_script),
            "--workspace-dir",
            str(workspace),
            "--case-id",
            "unsafe-host-network",
            "--mode",
            "docker-run",
            "--image",
            "alpine:3.20",
            "--timeout-seconds",
            "30",
            "--allow-exit-zero-oracle",
            "--network",
            "host",
            "--",
            "true",
        ],
        plugin_root,
        {
            "PATH": f"{fakebin}{os.pathsep}{os.environ.get('PATH', '')}",
            "ZHULONG_DOCKER_CALL_LOG": str(docker_log),
        },
        expected_returncode=1,
    )
    if "rejected_unsafe_sandbox" not in output:
        raise SystemExit(f"FAILED: runner did not surface rejected_unsafe_sandbox:\n{output}")
    if docker_log.exists():
        raise SystemExit("FAILED: run_verification_case.sh called docker after sandbox preflight rejection")
    result_path = workspace / "evidence/unsafe-host-network/verification-result.json"
    if result_path.exists():
        raise SystemExit("FAILED: rejected sandbox preflight created wrapper evidence before Docker execution")
    status_after = status_path.read_bytes() if status_path.exists() else None
    if status_after != status_before:
        raise SystemExit("FAILED: rejected runner changed the prior runtime sandbox status")


def exercise_verification_wrapper_state_boundary(plugin_root: Path, temp_root: Path) -> None:
    wrapper = plugin_root / "scripts/run_verification_case.sh"
    writer = plugin_root / "scripts/write_audit_event.py"
    matrix_root = temp_root / "verification-wrapper-state-boundary"
    matrix_root.mkdir(parents=True, exist_ok=True)

    def authority_fingerprint(workspace: Path) -> dict[str, object]:
        result: dict[str, object] = {}
        for name in ("audit-events.jsonl", "stage-status.json"):
            path = workspace / name
            if path.is_symlink():
                result[name] = {
                    "kind": "symlink",
                    "target": os.readlink(path),
                    "target_sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                }
            elif path.exists():
                result[name] = {
                    "kind": "regular",
                    "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                }
            else:
                result[name] = {"kind": "missing"}
        return result

    def write_transition(
        workspace: Path,
        event: str,
        stage: str,
        transition_kind: str,
        *,
        status: str = "running",
        extra: list[str] | None = None,
    ) -> None:
        command = [
            sys.executable,
            str(writer),
            "--workspace-dir",
            str(workspace),
            "--event",
            event,
            "--stage",
            stage,
            "--status",
            status,
            "--transition-kind",
            transition_kind,
            "--message",
            "Seed the verification-wrapper state-boundary selftest.",
            "--accept-current-revision",
        ]
        if extra:
            command.extend(extra)
        run(command, plugin_root)

    def make_workspace(name: str, stage: str | None, status: str = "running") -> Path:
        workspace = matrix_root / name
        workspace.mkdir(parents=True, exist_ok=True)
        (workspace / "asr-config.json").write_text('{"schema_version":1}\n', encoding="utf-8")
        (workspace / "audit-log.md").write_text("# Audit Log\n", encoding="utf-8")
        if stage is None:
            return workspace
        event_prefix = name.replace("-", "_")
        write_transition(workspace, f"{event_prefix}_intake", "intake", "start")
        ordered = ["recon", "candidate_generation", "triage", "verification"]
        for next_stage in ordered:
            if ordered.index(next_stage) > ordered.index(stage):
                break
            write_transition(workspace, f"{event_prefix}_{next_stage}", next_stage, "advance")
        if status == "completed":
            write_transition(workspace, f"{event_prefix}_completed", stage, "complete", status="completed")
        elif status == "blocked":
            write_transition(
                workspace,
                f"{event_prefix}_blocked",
                stage,
                "block",
                status="blocked",
                extra=[
                    "--blocker",
                    "A verification prerequisite is unavailable.",
                    "--resume-step",
                    "Resolve the prerequisite and explicitly retry this verification case.",
                ],
            )
        return workspace

    def install_fake_docker(workspace: Path) -> tuple[Path, Path]:
        fakebin = workspace / "fakebin"
        fakebin.mkdir(exist_ok=True)
        call_log = workspace / "docker-cli.log"
        poc_log = workspace / "poc-command.log"
        fake_docker = fakebin / "docker"
        fake_docker.write_text(
            "#!/usr/bin/env bash\n"
            "printf '%s\\n' \"$*\" >> \"$ZHULONG_DOCKER_CALL_LOG\"\n"
            "if [[ \"$1\" == \"info\" && \"$ZHULONG_STUB_BEHAVIOR\" == \"drift_before\" ]]; then\n"
            "  python3 \"$ZHULONG_WRITER\" --workspace-dir \"$ZHULONG_TEST_WORKSPACE\" "
            "--event concurrent_stage_drift --stage severity_escalation --status running "
            "--transition-kind advance --from-stage verification --from-status running "
            "--message 'Concurrent writer moved the stage.' --accept-current-revision >/dev/null\n"
            "fi\n"
            "if [[ \"$1\" == \"info\" && \"$ZHULONG_STUB_BEHAVIOR\" == \"writer_missing\" ]]; then\n"
            "  mv \"$ZHULONG_TEST_WRITER\" \"$ZHULONG_TEST_WRITER.hidden\"\n"
            "fi\n"
            "if [[ \"$1\" == \"info\" && \"$ZHULONG_STUB_BEHAVIOR\" == \"blocked\" ]]; then exit 1; fi\n"
            "if [[ \"$1\" == \"run\" ]]; then\n"
            "  printf 'poc-command\\n' >> \"$ZHULONG_POC_COMMAND_LOG\"\n"
            "  evidence=''\n"
            "  for arg in \"$@\"; do\n"
            "    case \"$arg\" in\n"
            "      type=bind,source=*,target=/workspace/evidence*) evidence=\"${arg#type=bind,source=}\"; evidence=\"${evidence%%,target=/workspace/evidence*}\" ;;\n"
            "    esac\n"
            "  done\n"
            "  case \"$ZHULONG_STUB_BEHAVIOR\" in\n"
            "    stdout_symlink) rm -f \"$evidence/stdout.log\"; ln -s \"$ZHULONG_STUB_MARKER\" \"$evidence/stdout.log\" ;;\n"
            "    stdout_hardlink) rm -f \"$evidence/stdout.log\"; ln \"$ZHULONG_STUB_MARKER\" \"$evidence/stdout.log\" ;;\n"
            "    stdout_fifo) rm -f \"$evidence/stdout.log\"; mkfifo \"$evidence/stdout.log\" ;;\n"
            "    stdout_dir) rm -f \"$evidence/stdout.log\"; mkdir \"$evidence/stdout.log\" ;;\n"
            "    result_symlink) rm -f \"$evidence/verification-result.json\"; ln -s ../../stage-status.json \"$evidence/verification-result.json\" ;;\n"
            "    ancestor_symlink) mv \"$evidence\" \"$evidence.real\"; ln -s \"$(basename \"$evidence.real\")\" \"$evidence\" ;;\n"
            "    running_replace) (sleep 0.05; rm -f \"$evidence/stdout.log\"; ln -s \"$ZHULONG_STUB_MARKER\" \"$evidence/stdout.log\") & ;;\n"
            "  esac\n"
            "  if [[ \"$ZHULONG_STUB_BEHAVIOR\" == \"drift_after\" ]]; then\n"
            "    python3 \"$ZHULONG_WRITER\" --workspace-dir \"$ZHULONG_TEST_WORKSPACE\" "
            "--event concurrent_postrun_drift --stage severity_escalation --status running "
            "--transition-kind advance --from-stage verification --from-status running "
            "--message 'Concurrent writer moved the stage after Docker.' --accept-current-revision >/dev/null\n"
            "  fi\n"
            "  if [[ \"$ZHULONG_STUB_BEHAVIOR\" == \"not_reproduced\" ]]; then printf 'no-match\\n'; exit 0; fi\n"
            "  if [[ \"$ZHULONG_STUB_BEHAVIOR\" == \"timeout\" ]]; then exit 124; fi\n"
            "  printf 'ZHULONG_STUB_ORACLE\\n'\n"
            "fi\n"
            "exit 0\n",
            encoding="utf-8",
        )
        fake_docker.chmod(0o755)
        return call_log, poc_log

    def invoke(
        workspace: Path,
        case_id: str,
        *,
        behavior: str = "confirmed",
        run_script: Path = wrapper,
        test_writer: Path | None = None,
        evidence_dir: Path | None = None,
    ) -> tuple[subprocess.CompletedProcess[str], dict[str, object], list[str], list[str]]:
        call_log, poc_log = install_fake_docker(workspace)
        fakebin = workspace / "fakebin"
        marker = workspace / "stub-external-marker.txt"
        marker.write_text("marker-before\n", encoding="utf-8")
        environment = {
            **os.environ,
            "PATH": f"{fakebin}{os.pathsep}{os.environ.get('PATH', '')}",
            "ZHULONG_DOCKER_CALL_LOG": str(call_log),
            "ZHULONG_POC_COMMAND_LOG": str(poc_log),
            "ZHULONG_STUB_BEHAVIOR": behavior,
            "ZHULONG_WRITER": str(writer),
            "ZHULONG_TEST_WORKSPACE": str(workspace),
            "ZHULONG_TEST_WRITER": str(test_writer or writer),
            "ZHULONG_STUB_MARKER": str(marker),
        }
        command = [
            "bash",
            str(run_script),
            "--workspace-dir",
            str(workspace),
            "--case-id",
            case_id,
            "--mode",
            "docker-run",
            "--image",
            "stub:local",
            "--timeout-seconds",
            "10",
            "--expected-oracle",
            "ZHULONG_STUB_ORACLE",
        ]
        if evidence_dir is not None:
            command.extend(["--evidence-dir", str(evidence_dir)])
        command.extend(["--", "true"])
        proc = subprocess.run(
            command,
            cwd=plugin_root,
            env=environment,
            capture_output=True,
            text=True,
        )
        result_path = workspace / "evidence" / case_id / "verification-result.json"
        result = json.loads(result_path.read_text(encoding="utf-8")) if result_path.exists() else {}
        docker_calls = call_log.read_text(encoding="utf-8").splitlines() if call_log.exists() else []
        poc_calls = poc_log.read_text(encoding="utf-8").splitlines() if poc_log.exists() else []
        return proc, result, docker_calls, poc_calls

    def event_names(workspace: Path) -> list[str]:
        return [
            json.loads(line)["event_name"]
            for line in (workspace / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]

    for name, stage, status in (
        ("wrong-candidate", "candidate_generation", "running"),
        ("wrong-triage-running", "triage", "running"),
        ("wrong-triage-completed", "triage", "completed"),
        ("wrong-verification-completed", "verification", "completed"),
    ):
        workspace = make_workspace(name, stage, status)
        before = authority_fingerprint(workspace)
        proc, result, docker_calls, poc_calls = invoke(workspace, name)
        if proc.returncode == 0 or result.get("status") != "blocked_state_precondition":
            raise SystemExit(f"FAILED: wrong-stage case did not fail closed: {name}: {proc.stdout}{proc.stderr}")
        if result.get("code") != "VERIFICATION_STATE_PRECONDITION_FAILED":
            raise SystemExit(f"FAILED: wrong-stage case did not use stable precondition code: {name}")
        if result.get("docker_invoked") or result.get("poc_command_invoked") or result.get("oracle_matched"):
            raise SystemExit(f"FAILED: wrong-stage case claimed Docker or oracle execution: {name}")
        if docker_calls or poc_calls:
            raise SystemExit(f"FAILED: wrong-stage case invoked stubbed Docker: {name}: {docker_calls}")
        if authority_fingerprint(workspace) != before:
            raise SystemExit(f"FAILED: wrong-stage case mutated journal/state: {name}")

    confirmed_workspace = make_workspace("valid-confirmed", "verification")
    proc, result, docker_calls, poc_calls = invoke(confirmed_workspace, "valid-confirmed")
    if proc.returncode != 0 or result.get("status") != "confirmed_in_docker":
        raise SystemExit(f"FAILED: verification/running confirmed path failed: {proc.stdout}{proc.stderr}")
    if event_names(confirmed_workspace)[-2:] != ["verification_case_started", "verification_case_completed"]:
        raise SystemExit("FAILED: confirmed path did not commit start then result observations")
    if len(poc_calls) != 1 or not any(call.startswith("run ") for call in docker_calls):
        raise SystemExit("FAILED: confirmed path did not invoke exactly one PoC container command")

    for attack in ("stdout_symlink", "stdout_hardlink", "stdout_fifo", "stdout_dir", "result_symlink", "ancestor_symlink", "running_replace"):
        attack_workspace = make_workspace(f"attack-{attack}", "verification")
        before = authority_fingerprint(attack_workspace)
        before_events = event_names(attack_workspace)
        proc, result, _docker_calls, _poc_calls = invoke(attack_workspace, f"attack-{attack}", behavior=attack)
        if proc.returncode == 0:
            raise SystemExit(f"FAILED: malicious Docker stub unexpectedly confirmed {attack}")
        if result and result.get("status") == "confirmed_in_docker":
            raise SystemExit(f"FAILED: malicious Docker stub influenced confirmed result for {attack}: {result}")
        after_events = event_names(attack_workspace)
        if after_events != before_events and after_events != before_events + ["verification_case_started"]:
            raise SystemExit(f"FAILED: malicious Docker stub appended an unauthorized authority event for {attack}: {after_events}")
        if json.loads((attack_workspace / "stage-status.json").read_text(encoding="utf-8")).get("status") != "running":
            raise SystemExit(f"FAILED: malicious Docker stub changed workflow authority for {attack}")
        marker = attack_workspace / "stub-external-marker.txt"
        if marker.read_text(encoding="utf-8") != "marker-before\n":
            raise SystemExit(f"FAILED: malicious Docker stub changed external marker for {attack}")

    rejected_workspace = make_workspace("valid-rejected", "verification")
    proc, result, _, poc_calls = invoke(rejected_workspace, "valid-rejected", behavior="not_reproduced")
    if proc.returncode == 0 or result.get("status") != "rejected_not_reproducible":
        raise SystemExit("FAILED: not-reproduced path did not preserve the Docker evidence status")
    if event_names(rejected_workspace)[-2:] != ["verification_case_started", "verification_case_rejected"]:
        raise SystemExit("FAILED: not-reproduced path did not use same-stage observations")
    if len(poc_calls) != 1:
        raise SystemExit("FAILED: not-reproduced path did not invoke exactly one PoC command")

    blocked_workspace = make_workspace("valid-blocked", "verification")
    proc, result, docker_calls, poc_calls = invoke(blocked_workspace, "valid-blocked", behavior="blocked")
    state = json.loads((blocked_workspace / "stage-status.json").read_text(encoding="utf-8"))
    last_event = json.loads((blocked_workspace / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()[-1])
    if proc.returncode == 0 or result.get("status") != "blocked_docker_unavailable":
        raise SystemExit("FAILED: blocked verification result did not fail closed")
    if last_event.get("transition_kind") != "block" or state.get("stage") != "verification" or state.get("status") != "blocked":
        raise SystemExit("FAILED: blocked result did not use verification/running -> verification/blocked")
    if len(docker_calls) != 1 or poc_calls:
        raise SystemExit("FAILED: Docker-unavailable result crossed the PoC boundary")

    resume_workspace = make_workspace("valid-resume", "verification", "blocked")
    proc, result, _, poc_calls = invoke(resume_workspace, "valid-resume")
    if proc.returncode != 0 or result.get("status") != "confirmed_in_docker":
        raise SystemExit("FAILED: explicit verification/blocked retry did not resume and complete")
    if event_names(resume_workspace)[-3:] != [
        "verification_case_resumed",
        "verification_case_started",
        "verification_case_completed",
    ]:
        raise SystemExit("FAILED: explicit blocked retry event sequence is incorrect")
    if len(poc_calls) != 1:
        raise SystemExit("FAILED: explicit blocked retry did not invoke exactly one PoC command")

    invalid_workspaces: list[tuple[str, Path]] = []
    missing_state = make_workspace("missing-state", "verification")
    (missing_state / "stage-status.json").unlink()
    invalid_workspaces.append(("missing-state", missing_state))

    stale_state = make_workspace("stale-state", "verification")
    stale_doc = json.loads((stale_state / "stage-status.json").read_text(encoding="utf-8"))
    stale_doc["last_event_name"] = "stale_view"
    (stale_state / "stage-status.json").write_text(json.dumps(stale_doc, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    invalid_workspaces.append(("stale-state", stale_state))

    corrupt_journal = make_workspace("corrupt-journal", "verification")
    with (corrupt_journal / "audit-events.jsonl").open("a", encoding="utf-8") as stream:
        stream.write("{not-json}\n")
    invalid_workspaces.append(("corrupt-journal", corrupt_journal))

    journal_symlink = make_workspace("journal-symlink", "verification")
    journal_target = journal_symlink / "journal-target"
    (journal_symlink / "audit-events.jsonl").rename(journal_target)
    (journal_symlink / "audit-events.jsonl").symlink_to(journal_target.name)
    invalid_workspaces.append(("journal-symlink", journal_symlink))

    state_symlink = make_workspace("state-symlink", "verification")
    state_target = state_symlink / "state-target"
    (state_symlink / "stage-status.json").rename(state_target)
    (state_symlink / "stage-status.json").symlink_to(state_target.name)
    invalid_workspaces.append(("state-symlink", state_symlink))

    protocol_mismatch = make_workspace("protocol-mismatch", "verification")
    legacy_state = {
        "schema_version": 1,
        "plugin": "zhulong",
        "plugin_version": "selftest",
        "stage": "candidate_verifying",
        "status": "running",
        "last_event_at": "2026-07-31T00:00:00Z",
        "blocker": None,
        "resume_step": None,
        "workspace": "<workspace>",
        "target_repo": "<repo>",
        "last_event": "legacy",
        "last_message": "legacy",
    }
    (protocol_mismatch / "stage-status.json").write_text(
        json.dumps(legacy_state, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    invalid_workspaces.append(("protocol-mismatch", protocol_mismatch))

    for name, workspace in invalid_workspaces:
        before = authority_fingerprint(workspace)
        proc, result, docker_calls, poc_calls = invoke(workspace, name)
        if proc.returncode == 0 or result.get("status") != "blocked_state_precondition":
            raise SystemExit(f"FAILED: invalid R2 authority case did not reject: {name}")
        if docker_calls or poc_calls:
            raise SystemExit(f"FAILED: invalid R2 authority case invoked Docker: {name}")
        if authority_fingerprint(workspace) != before:
            raise SystemExit(f"FAILED: invalid R2 authority case mutated authority bytes: {name}")

    drift_workspace = make_workspace("concurrent-drift-before", "verification")
    proc, result, docker_calls, poc_calls = invoke(drift_workspace, "concurrent-drift-before", behavior="drift_before")
    if proc.returncode == 0 or result.get("status") != "blocked_authority_event_commit":
        raise SystemExit("FAILED: concurrent pre-start stage drift did not fail closed")
    if poc_calls or any(call.startswith("run ") for call in docker_calls):
        raise SystemExit("FAILED: concurrent pre-start stage drift allowed the PoC command")
    if "verification_case_started" in event_names(drift_workspace):
        raise SystemExit("FAILED: concurrent pre-start stage drift appended a start event")

    isolated = matrix_root / "isolated-layout"
    shutil.copytree(plugin_root / "scripts", isolated / "scripts")
    shutil.copytree(plugin_root / "assets", isolated / "assets")
    writer_failure = make_workspace("start-writer-failure", "verification")
    isolated_writer = isolated / "scripts/write_audit_event.py"
    proc, result, docker_calls, poc_calls = invoke(
        writer_failure,
        "start-writer-failure",
        behavior="writer_missing",
        run_script=isolated / "scripts/run_verification_case.sh",
        test_writer=isolated_writer,
    )
    if proc.returncode == 0 or result.get("status") != "blocked_authority_event_commit":
        raise SystemExit(
            "FAILED: start-event writer failure did not fail closed: "
            f"exit={proc.returncode} result={result} stdout={proc.stdout} stderr={proc.stderr}"
        )
    if poc_calls or any(call.startswith("run ") for call in docker_calls):
        raise SystemExit("FAILED: start-event writer failure allowed the PoC command")

    post_drift = make_workspace("concurrent-drift-after", "verification")
    proc, result, _, poc_calls = invoke(post_drift, "concurrent-drift-after", behavior="drift_after")
    if proc.returncode == 0 or result.get("wrapper_status") != "blocked_authority_event_commit":
        raise SystemExit("FAILED: post-run authority commit failure was reported as complete")
    if result.get("status") != "confirmed_in_docker" or not result.get("oracle_matched") or len(poc_calls) != 1:
        raise SystemExit("FAILED: post-run authority failure did not preserve the Docker evidence truth")
    if result.get("authority_event_committed") is not False:
        raise SystemExit("FAILED: post-run authority failure claimed an event commit")

    legacy_workspace = make_workspace("legacy-r1", None)
    legacy_event = {
        "ts": "2026-07-31T00:00:00Z",
        "event": "legacy_seed",
        "stage": "candidate_verifying",
        "status": "running",
        "message": "legacy baseline",
        "details": {},
    }
    (legacy_workspace / "audit-events.jsonl").write_text(json.dumps(legacy_event, sort_keys=True) + "\n", encoding="utf-8")
    (legacy_workspace / "stage-status.json").write_text(
        json.dumps(legacy_state, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    legacy_before = authority_fingerprint(legacy_workspace)
    proc, result, docker_calls, poc_calls = invoke(legacy_workspace, "legacy-r1")
    if proc.returncode != 0 or result.get("status") != "confirmed_in_docker":
        raise SystemExit("FAILED: R1 wrapper compatibility changed")
    if authority_fingerprint(legacy_workspace) == legacy_before or len(docker_calls) != 3 or len(poc_calls) != 1:
        raise SystemExit("FAILED: R1 wrapper baseline mutation/execution behavior changed")

    no_state = make_workspace("no-state", None)
    no_state_before = authority_fingerprint(no_state)
    proc, result, docker_calls, poc_calls = invoke(no_state, "no-state")
    if proc.returncode == 0 or result.get("status") != "blocked_state_precondition":
        raise SystemExit("FAILED: no-state wrapper did not fail closed before execution")
    if (
        result.get("code") != "VERIFICATION_STATE_PRECONDITION_FAILED"
        or result.get("state_issue_code") != "AUTHORITATIVE_STATE_MISSING"
    ):
        raise SystemExit("FAILED: no-state path did not expose the stable missing-authority precondition")
    if authority_fingerprint(no_state) != no_state_before:
        raise SystemExit("FAILED: no-state path silently created or upgraded audit authority")
    if docker_calls or poc_calls:
        raise SystemExit("FAILED: no-state wrapper crossed the Docker or PoC boundary")

    for index, invalid_case_id in enumerate((".hidden", "..", "bad case", "case/slash", r"case\\slash"), start=1):
        invalid_case_workspace = make_workspace(f"invalid-case-id-{index}", None)
        before = authority_fingerprint(invalid_case_workspace)
        proc, _result, docker_calls, poc_calls = invoke(invalid_case_workspace, invalid_case_id)
        if proc.returncode != 2:
            raise SystemExit(f"FAILED: unsafe case ID was not rejected before execution: {invalid_case_id!r}")
        if docker_calls or poc_calls:
            raise SystemExit(f"FAILED: unsafe case ID crossed the Docker or PoC boundary: {invalid_case_id!r}")
        if authority_fingerprint(invalid_case_workspace) != before:
            raise SystemExit(f"FAILED: unsafe case ID mutated authority bytes: {invalid_case_id!r}")
        if (invalid_case_workspace / "evidence").exists():
            raise SystemExit(f"FAILED: unsafe case ID created an evidence directory: {invalid_case_id!r}")

    custom_evidence_workspace = make_workspace("custom-evidence-dir", None)
    custom_evidence_before = authority_fingerprint(custom_evidence_workspace)
    proc, _result, docker_calls, poc_calls = invoke(
        custom_evidence_workspace,
        "valid-case",
        evidence_dir=custom_evidence_workspace / "evidence" / "other-case",
    )
    if proc.returncode != 2 or docker_calls or poc_calls:
        raise SystemExit("FAILED: custom evidence directory outside the case boundary was accepted")
    if authority_fingerprint(custom_evidence_workspace) != custom_evidence_before:
        raise SystemExit("FAILED: rejected custom evidence directory mutated authority bytes")

    external_evidence = matrix_root / "external-evidence-target"
    external_evidence.mkdir(parents=True, exist_ok=True)
    symlink_evidence_workspace = make_workspace("symlink-evidence-dir", None)
    (symlink_evidence_workspace / "evidence").symlink_to(external_evidence, target_is_directory=True)
    symlink_before = authority_fingerprint(symlink_evidence_workspace)
    proc, _result, docker_calls, poc_calls = invoke(symlink_evidence_workspace, "valid-case")
    if proc.returncode != 2 or docker_calls or poc_calls:
        raise SystemExit("FAILED: symlink evidence ancestor was accepted")
    if authority_fingerprint(symlink_evidence_workspace) != symlink_before:
        raise SystemExit("FAILED: rejected symlink evidence directory mutated authority bytes")

    for workspace in matrix_root.iterdir():
        if not workspace.is_dir():
            continue
        forbidden = [
            workspace / "audit-disposition.json",
            workspace / "confirmed",
            workspace / "finalization-result.json",
            workspace / "verifier",
        ]
        if any(path.exists() for path in forbidden):
            raise SystemExit(f"FAILED: wrapper state-boundary selftest created authority side effects: {workspace.name}")

    print("VERIFICATION WRAPPER STATE BOUNDARY SELFTEST PASSED")


def exercise_sandbox_ledger_guard(workspace: Path, plugin_root: Path) -> None:
    sys.path.insert(0, str(plugin_root / "scripts"))
    from audit_disposition import synthesize_disposition_ledger  # type: ignore

    unverified_path = workspace / "unverified-leads.md"
    original = unverified_path.read_text(encoding="utf-8", errors="ignore") if unverified_path.exists() else ""
    try:
        unverified_path.write_text(
            "| Lead ID | Suspected Weakness | Evidence So Far | Missing Evidence | Docker Confirmation Status | Safe Resume Step | High-Confidence-Unverified? | Material blocker? | Default runtime scope? | Why completion is still safe? |\n"
            "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |\n"
            "| U-SANDBOX | SSRF candidate | PoC needs unsafe Docker flags | safe sandbox rewrite | rejected_unsafe_sandbox | rewrite verification container without host/privileged/docker.sock/root mount | Yes | Yes | default runtime | unsafe sandbox rejection is blocked/unverified, not confirmed |\n",
            encoding="utf-8",
        )
        ledger = synthesize_disposition_ledger(workspace, merge_existing=False)
    finally:
        unverified_path.write_text(original, encoding="utf-8")

    matches = [
        item for item in ledger.get("items", [])
        if "sandbox" in str(item.get("id", "")).lower() or "sandbox" in str(item.get("title", "")).lower()
    ]
    if not matches:
        raise SystemExit("FAILED: audit disposition ledger did not capture rejected_unsafe_sandbox lead")
    if any(item.get("state") == "confirmed" or item.get("confirmed_bundle_path") for item in matches):
        raise SystemExit(f"FAILED: rejected_unsafe_sandbox entered confirmed ledger state: {matches}")
    if not any(item.get("state") in {"blocked", "unverified"} for item in matches):
        raise SystemExit(f"FAILED: rejected_unsafe_sandbox must stay blocked/unverified: {matches}")


def exercise_structured_blocker_cli(plugin_root: Path) -> None:
    """Exercise structured blocker precedence and same-identity recovery through a subprocess."""
    with tempfile.TemporaryDirectory(prefix="zhulong-structured-blocker-") as tempdir:
        workspace = Path(tempdir) / "workspace"
        result_path = workspace / "evidence/C1/verification-result.json"
        result_path.parent.mkdir(parents=True)
        blocked = {
            "schema_version": 1,
            "case_id": "C1",
            "status": "blocked_missing_image",
            "classification_reason": "fixture blocker",
            "authority_event_committed": True,
        }
        result_path.write_text(json.dumps(blocked, sort_keys=True) + "\n", encoding="utf-8")

        probe = (
            "import json,sys; sys.path.insert(0, sys.argv[2]); "
            "from blocked_verification import detect_blocked_verification; "
            "print(json.dumps(detect_blocked_verification(__import__('pathlib').Path(sys.argv[1])), sort_keys=True))"
        )
        blocked_proc = subprocess.run(
            [sys.executable, "-c", probe, str(workspace), str(plugin_root / "scripts")],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        if blocked_proc.returncode != 0:
            raise SystemExit(f"FAILED: structured blocker subprocess failed: {blocked_proc.stderr}")
        blocked_payload = json.loads(blocked_proc.stdout)
        if blocked_payload.get("blocked") is not True or blocked_payload.get("structured_fact_count", 0) < 1:
            raise SystemExit(f"FAILED: structured blocked_missing_image was not authoritative: {blocked_payload}")

        resolved = dict(blocked)
        resolved.update({"status": "confirmed_in_docker", "oracle_matched": True})
        result_path.write_text(json.dumps(resolved, sort_keys=True) + "\n", encoding="utf-8")
        resolved_proc = subprocess.run(
            [sys.executable, "-c", probe, str(workspace), str(plugin_root / "scripts")],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        if resolved_proc.returncode != 0:
            raise SystemExit(f"FAILED: structured blocker recovery subprocess failed: {resolved_proc.stderr}")
        resolved_payload = json.loads(resolved_proc.stdout)
        if resolved_payload.get("blocked") is not False or "case:C1" not in resolved_payload.get("resolved_identities", []):
            raise SystemExit(f"FAILED: same-identity confirmed result did not clear blocker: {resolved_payload}")
    print("STRUCTURED BLOCKER SELFTEST PASSED: blocked facts first, same-identity recovery only")


def exercise_recording_evidence_gate(plugin_root: Path) -> None:
    """Exercise the public offline recording identity/media/transaction gate."""

    try:
        from PIL import Image, ImageDraw
    except ImportError as exc:
        raise SystemExit(f"FAILED: recording selftest requires Pillow: {exc}") from exc

    validator_path = plugin_root / "scripts/validate_recording_evidence.py"
    auto_recorder_path = plugin_root / "scripts/auto_record_bundle.py"
    if str(plugin_root / "scripts") not in sys.path:
        sys.path.insert(0, str(plugin_root / "scripts"))
    required = [
        "scripts/recording_identity.py",
        "scripts/validate_recording_evidence.py",
        "scripts/auto_record_bundle.py",
        "assets/schemas/recording-evidence.schema.json",
        "assets/fixtures/recording-evidence/README.md",
        "assets/fixtures/recording-evidence/manifest.template.json",
    ]
    require_files(plugin_root, required, "recording evidence gate")
    if (plugin_root / "skills/zhulong/SKILL.md").is_file():
        if (plugin_root / "skills/zhulong/SKILL.md").read_bytes() != (plugin_root / "templates/claude-skill/SKILL.md").read_bytes():
            raise SystemExit("FAILED: Claude and Codex public skill source files drifted")

    validator_spec = importlib.util.spec_from_file_location("zhulong_recording_validator_selftest", validator_path)
    if validator_spec is None or validator_spec.loader is None:
        raise SystemExit("FAILED: cannot load recording validator")
    validator_module = importlib.util.module_from_spec(validator_spec)
    sys.modules[validator_spec.name] = validator_module
    validator_spec.loader.exec_module(validator_module)

    def finding() -> dict[str, object]:
        return {
            "slug": "example-finding",
            "project_name": "example-project",
            "vuln_type": "Example finding",
            "source_binding": {
                "tested_ref": "v0.0.0-test",
                "source_bound_ref": "v0.0.0-test",
                "entrypoint": "example trigger context",
            },
            "attacker_condition": "example trigger context",
            "code_context": [{
                "location": "src/example.py:1",
                "summary": "example input",
                "explanation": "example operation",
            }],
            "verification_evidence": {
                "finding_slug": "example-finding",
                "oracle_token": "EXAMPLE_ORACLE_CONFIRMED",
                "direct_impact_marker": "DIRECT_IMPACT_CONFIRMED",
            },
        }

    code_context = "location=src/example.py:1;source=example input;sink=example operation"
    canonical = {
        "software_name": "example-project",
        "tested_ref": "v0.0.0-test",
        "tested_ref_kind": "version",
        "finding_slug": "example-finding",
        "direct_impact_marker": "DIRECT_IMPACT_CONFIRMED",
        "oracle_marker": "EXAMPLE_ORACLE_CONFIRMED",
        "code_context_identity": code_context,
        "trigger_context_identity": "example trigger context",
    }
    screenshot_paths = [
        "attachments/evidence/screenshots/01-target-identity.png",
        "attachments/evidence/screenshots/02-code-or-trigger-context.png",
        "attachments/evidence/screenshots/03-final-impact.png",
    ]
    stages = ["identity", "code_or_trigger_context", "final_impact"]
    timestamps = [0.5, 1.5, 2.5]
    stage_markers = [
        "example-project v0.0.0-test",
        code_context,
        "DIRECT_IMPACT_CONFIRMED",
    ]

    def fixture(root: Path) -> tuple[Path, Path]:
        bundle = root / "example-finding"
        checkpoint_dir = root / "checkpoint-images"
        (bundle / "attachments/evidence/screenshots").mkdir(parents=True)
        checkpoint_dir.mkdir()
        source_finding = finding()
        (bundle / "findings.json").write_text(json.dumps({"findings": [source_finding]}, indent=2) + "\n", encoding="utf-8")
        (bundle / "validity-review.json").write_text(json.dumps({
            "schema_version": 1,
            "finding_slug": "example-finding",
            "project_name": "example-project",
            "source_binding": source_finding["source_binding"],
            "trigger_context": "example trigger context",
            "code_context": source_finding["code_context"],
            "oracle_token": "EXAMPLE_ORACLE_CONFIRMED",
            "direct_impact_marker": "DIRECT_IMPACT_CONFIRMED",
        }, indent=2) + "\n", encoding="utf-8")
        proof = bundle / "attachments/evidence/example-proof.txt"
        proof.write_text("EXAMPLE_ORACLE_CONFIRMED\n", encoding="utf-8")
        (bundle / "attachments/evidence/replay-output.log").write_text(
            "[command] example local fixture replay\nEXAMPLE_ORACLE_CONFIRMED\nDIRECT_IMPACT_CONFIRMED\n",
            encoding="utf-8",
        )
        helper = bundle / "run-example-finding-recording.sh"
        helper.write_text(
            "#!/bin/sh\nrecording_checkpoint() { :; }\n"
            "ZHULONG_RECORDING_STAGE_DIR=\"\"\n"
            "ZHULONG_RECORDING_STAGE_ACK_DIR=\"\"\n"
            "ZHULONG_RECORDING_OWNER_MARKER=\"\"\n",
            encoding="utf-8",
        )
        helper.chmod(0o755)
        frames = []
        frame_text = [
            ["example-project", "v0.0.0-test", "IDENTITY_CONFIRMED"],
            ["src/example.py:1", "example trigger context", "CODE_OR_TRIGGER_CONTEXT_CONFIRMED"],
            ["DIRECT_IMPACT_CONFIRMED", "EXAMPLE_ORACLE_CONFIRMED", "FINAL_IMPACT_CONFIRMED"],
        ]
        for index, lines in enumerate(frame_text):
            image = Image.new("RGB", (320, 180), [(38, 54, 84), (39, 76, 68), (92, 54, 54)][index])
            draw = ImageDraw.Draw(image)
            for row, line in enumerate(lines):
                draw.text((18, 24 + row * 42), line, fill=(245, 245, 245))
            frames.append(image)
        video = bundle / "attachments/evidence/final-recording.gif"
        frames[0].save(video, save_all=True, append_images=frames[1:], duration=[1000, 1000, 1000], loop=0)
        for stage, timestamp in zip(stages, timestamps):
            path = checkpoint_dir / f"{stage}.png"
            validator_module._extract_frame(video, timestamp, path)

        verification = {
            "schema_version": 1,
            "finding_slug": "example-finding",
            "verification_status": "confirmed_in_docker",
            "docker_required": True,
            "docker_image": "example-image:v0.0.0-test",
            "docker_command": "docker run example-image:v0.0.0-test",
            "poc_path": proof.relative_to(bundle).as_posix(),
            "expected_observation": "EXAMPLE_ORACLE_CONFIRMED",
            "observed_observation": "EXAMPLE_ORACLE_CONFIRMED",
            "oracle_token": "EXAMPLE_ORACLE_CONFIRMED",
            "direct_impact_marker": "DIRECT_IMPACT_CONFIRMED",
            "evidence_files": [],
            "severity_escalation_attempted": True,
            "severity_escalation_result": "Sanitized fixture only.",
        }
        stage_items = []
        screenshot_items = []
        for sequence, (stage, timestamp, marker) in enumerate(zip(stages, timestamps, stage_markers), start=1):
            source_path = checkpoint_dir / f"{stage}.png"
            source_hash = hashlib.sha256(source_path.read_bytes()).hexdigest()
            screenshot_path = bundle / screenshot_paths[sequence - 1]
            shutil.copy2(source_path, screenshot_path)
            frame_hash = hashlib.sha256(source_path.read_bytes()).hexdigest()
            observations = {
                "identity": ["example-project", "v0.0.0-test", marker],
                "code_or_trigger_context": [code_context, "example trigger context", marker],
                "final_impact": ["DIRECT_IMPACT_CONFIRMED", "EXAMPLE_ORACLE_CONFIRMED", marker],
            }[stage]
            stage_items.append({
                "stage": stage,
                "sequence": sequence,
                "event_timestamp": float(sequence),
                "video_timestamp": timestamp,
                "hold_start": max(0.0, timestamp - 0.4),
                "hold_end": min(3.0, timestamp + 0.4),
                "expected_marker": marker,
                "source_name": "fixture source",
                "source_window_identity": "fixture-terminal-window",
                "canonical_identity": {
                    "software_name": canonical["software_name"],
                    "tested_ref": canonical["tested_ref"],
                    "finding_slug": canonical["finding_slug"],
                    "code_context_identity": canonical["code_context_identity"],
                    "trigger_context_identity": canonical["trigger_context_identity"],
                },
                "source_checkpoint": {"name": source_path.name, "sha256": source_hash, "width": 320, "height": 180},
                "frame": {
                    "sha256": frame_hash,
                    "width": 320,
                    "height": 180,
                    "perceptual_similarity": 1.0,
                    "recording_time_observations": observations,
                },
            })
            screenshot_items.append({
                "stage": stage,
                "path": screenshot_paths[sequence - 1],
                "sha256": hashlib.sha256(screenshot_path.read_bytes()).hexdigest(),
                "size": screenshot_path.stat().st_size,
                "width": 320,
                "height": 180,
                "video_timestamp": timestamp,
                "source_frame_sha256": frame_hash,
            })
        inventory_rel = "attachments/recording-screenshot-inventory.md"
        (bundle / inventory_rel).write_text("\n".join(f"- `{path}`" for path in screenshot_paths) + "\n", encoding="utf-8")
        verification["evidence_files"] = [
            proof.relative_to(bundle).as_posix(),
            "attachments/evidence/replay-output.log",
            "attachments/evidence/final-recording.gif",
            *screenshot_paths,
            "recording-evidence.json",
            inventory_rel,
        ]
        (bundle / "verification-evidence.json").write_text(json.dumps(verification, indent=2) + "\n", encoding="utf-8")
        (bundle / "attachments/reviewer-evidence-index.json").write_text(json.dumps({
            "schema_version": 1,
            "replay_command": "./run-example-finding-recording.sh quick docker",
            "evidence_artifacts": [{"path": path} for path in verification["evidence_files"]],
            "oracle_tokens": ["EXAMPLE_ORACLE_CONFIRMED", "DIRECT_IMPACT_CONFIRMED"],
        }, indent=2) + "\n", encoding="utf-8")
        manifest = {
            "schema_version": 1,
            "recording_status": "staging",
            "canonical_identity": canonical,
            "video": {
                "path": "attachments/evidence/final-recording.gif",
                "sha256": hashlib.sha256(video.read_bytes()).hexdigest(),
                "size": video.stat().st_size,
                "duration_seconds": 3.0,
                "width": 320,
                "height": 180,
            },
            "replay": {"script_path": helper.name, "script_sha256": hashlib.sha256(helper.read_bytes()).hexdigest(), "exit_code": 0},
            "obs": {
                "source_name": "fixture source",
                "source_kind": "fixture_media",
                "window_identity": "fixture-terminal-window",
                "window_title": "Zhulong fixture terminal",
                "window_stable": True,
            },
            "stages": stage_items,
            "screenshots": screenshot_items,
            "registrations": {
                "verification_evidence_path": "verification-evidence.json",
                "reviewer_index_path": "attachments/reviewer-evidence-index.json",
                "attachment_inventory_path": inventory_rel,
                "screenshot_paths": screenshot_paths,
            },
            "archive": {
                "status": "not_ready",
                "archive_name": "example-finding.zip",
                "testzip": None,
                "required_entries": ["recording-evidence.json", *screenshot_paths, "attachments/evidence/final-recording.gif"],
                "recording_ready": False,
                "submission_ready": False,
            },
            "transaction": {
                "owner": "zhulong-recording",
                "owner_marker": ".zhulong-recording-transaction.json",
                "promotion_status": "staging",
                "rollback_safe": True,
                "full_recording_time_validated": False,
            },
        }
        (bundle / "recording-evidence.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
        return bundle, checkpoint_dir

    def copy_case(root: Path, base_bundle: Path, base_checkpoint: Path, name: str) -> tuple[Path, Path]:
        bundle = root / name
        checkpoint = root / f"{name}-checkpoints"
        shutil.copytree(base_bundle, bundle)
        shutil.copytree(base_checkpoint, checkpoint)
        return bundle, checkpoint

    def error_code(bundle: Path, checkpoint: Path, archive: Path | None = None) -> str:
        command = [sys.executable, str(validator_path), "--bundle-dir", str(bundle), "--checkpoint-dir", str(checkpoint), "--json"]
        if archive is not None:
            command.extend(["--archive", str(archive), "--archive-root", bundle.name])
        proc = subprocess.run(command, cwd=plugin_root, capture_output=True, text=True)
        if proc.returncode == 0:
            raise SystemExit(f"FAILED: recording negative case unexpectedly passed: {bundle.name}")
        output = proc.stdout.strip().splitlines()
        if not output:
            raise SystemExit(f"FAILED: recording validator emitted no JSON for {bundle.name}: {proc.stderr}")
        return str(json.loads(output[-1]).get("error_code") or "")

    with tempfile.TemporaryDirectory(prefix="zhulong-recording-selftest-") as tempdir:
        root = Path(tempdir)
        base_bundle, base_checkpoint = fixture(root)

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-01")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["obs"]["window_stable"] = False
        data["stages"][0]["frame"]["recording_time_observations"] = ["UNRELATED_NONBLACK_WINDOW"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_WRONG_WINDOW":
            raise SystemExit("FAILED: case 1 wrong-window gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-01a")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][0]["frame"]["recording_time_observations"] = ["UNRELATED_NONBLACK_WINDOW"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_IDENTITY_FRAME_MISSING":
            raise SystemExit("FAILED: case 1a identity-specific error must precede the generic marker error")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-01b")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][0]["frame"]["recording_time_observations"] = ["example-project", "v0.0.0-test"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_VIDEO_CONTENT_UNVERIFIED":
            raise SystemExit("FAILED: case 1b identity marker absence must retain the generic content error")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-02")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][0]["video_timestamp"] = 3.2
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_VIDEO_CONTENT_UNVERIFIED":
            raise SystemExit("FAILED: case 2 video-start gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-03")
        data = json.loads((case / "recording-evidence.json").read_text())
        source = checkpoint / "final_impact.png"
        shutil.copy2(checkpoint / "identity.png", source)
        data["stages"][2]["source_checkpoint"]["sha256"] = hashlib.sha256(source.read_bytes()).hexdigest()
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        case_03_error = error_code(case, checkpoint)
        if case_03_error != "RECORDING_STAGE_FRAME_MISMATCH":
            raise SystemExit(f"FAILED: case 3 source/frame mismatch gate: {case_03_error}")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-04")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][1]["source_window_identity"] = "changed-window"
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_WRONG_WINDOW":
            raise SystemExit("FAILED: case 4 source/window stability gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-05")
        (case / screenshot_paths[2]).unlink()
        case_05_error = error_code(case, checkpoint)
        if case_05_error != "RECORDING_SCREENSHOT_MISSING":
            raise SystemExit(f"FAILED: case 5 missing screenshot gate: {case_05_error}")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-06")
        shutil.copy2(case / screenshot_paths[0], case / screenshot_paths[1])
        data = json.loads((case / "recording-evidence.json").read_text())
        data["screenshots"][1]["sha256"] = hashlib.sha256((case / screenshot_paths[1]).read_bytes()).hexdigest()
        data["screenshots"][1]["size"] = (case / screenshot_paths[1]).stat().st_size
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_SCREENSHOT_DUPLICATE":
            raise SystemExit("FAILED: case 6 duplicate screenshot gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-07")
        verification = json.loads((case / "verification-evidence.json").read_text())
        verification["evidence_files"] = [item for item in verification["evidence_files"] if "03-final-impact" not in item]
        (case / "verification-evidence.json").write_text(json.dumps(verification, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_SCREENSHOT_UNREGISTERED":
            raise SystemExit("FAILED: case 7 screenshot registration gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-08")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["video"]["sha256"] = "0" * 64
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_HASH_MISMATCH":
            raise SystemExit("FAILED: case 8 hash tamper gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-08a")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][1]["expected_marker"] = "CONTEXT_FRAME"
        data["stages"][1]["frame"]["recording_time_observations"] = ["CONTEXT_FRAME"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_STAGE_FRAME_MISMATCH":
            raise SystemExit("FAILED: case 8a code/context-specific error changed")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-09")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][2]["expected_marker"] = "FINAL_IMPACT_FRAME"
        data["stages"][2]["frame"]["recording_time_observations"] = ["EXAMPLE_ORACLE_CONFIRMED", "FINAL_IMPACT_FRAME"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_IMPACT_FRAME_MISSING":
            raise SystemExit("FAILED: case 9 direct-impact gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-10")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["replay"]["exit_code"] = 7
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        if error_code(case, checkpoint) != "RECORDING_REPLAY_FAILED":
            raise SystemExit("FAILED: case 10 replay exit gate")

        case, checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-11")
        archive = root / "case-11.zip"
        with zipfile.ZipFile(archive, "w") as handle:
            handle.write(case / "recording-evidence.json", f"{case.name}/recording-evidence.json")
        if error_code(case, checkpoint, archive) != "RECORDING_ARCHIVE_INCOMPLETE":
            raise SystemExit("FAILED: case 11 incomplete archive gate")

        case, _checkpoint = copy_case(root, base_bundle, base_checkpoint, "case-11a")
        data = json.loads((case / "recording-evidence.json").read_text())
        data["stages"][0]["frame"]["recording_time_observations"] = ["forged identity claim"]
        (case / "recording-evidence.json").write_text(json.dumps(data, indent=2) + "\n")
        finalize_without_checkpoint = subprocess.run(
            [sys.executable, str(validator_path), "--bundle-dir", str(case), "--finalize", "--json"],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        if finalize_without_checkpoint.returncode == 0 or "RECORDING_VIDEO_CONTENT_UNVERIFIED" not in finalize_without_checkpoint.stdout:
            raise SystemExit("FAILED: case 11a forged recording-time observations finalized without live checkpoints")
        persisted = json.loads((case / "recording-evidence.json").read_text())
        if persisted["recording_status"] != "staging" or persisted["transaction"]["full_recording_time_validated"] is not False:
            raise SystemExit("FAILED: case 11a missing checkpoints changed promotion authority")

        auto_spec = importlib.util.spec_from_file_location("zhulong_auto_recording_selftest", auto_recorder_path)
        if auto_spec is None or auto_spec.loader is None:
            raise SystemExit("FAILED: cannot load public auto recorder")
        auto_module = importlib.util.module_from_spec(auto_spec)
        sys.modules[auto_spec.name] = auto_module
        auto_spec.loader.exec_module(auto_module)

        retained_dir = root / "retained-unpromoted"
        parsed = auto_module.parse_args(["example-bundle", "--keep-unpromoted-archive", str(retained_dir)])
        if parsed.keep_unpromoted_archive != retained_dir or parsed.zip_on_fail:
            raise SystemExit("FAILED: recording diagnostic archive CLI did not parse the explicit destination")
        legacy = auto_module.parse_args(["example-bundle", "--zip-on-fail"])
        if not legacy.zip_on_fail or legacy.keep_unpromoted_archive is not None:
            raise SystemExit("FAILED: deprecated --zip-on-fail compatibility parsing drifted")
        legacy_probe = subprocess.run(
            [sys.executable, str(auto_recorder_path), str(root / "missing-bundle"), "--zip-on-fail"],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        if legacy_probe.returncode == 0 or "deprecated" not in legacy_probe.stderr or (root / "missing-bundle.zip").exists():
            raise SystemExit("FAILED: deprecated --zip-on-fail did not warn and remain non-producing")
        staged_diagnostic_zip = root / "verified-unpromoted.zip"
        with zipfile.ZipFile(staged_diagnostic_zip, "w") as handle:
            handle.writestr("example-finding/evidence.txt", "verified bytes\n")
        final_diagnostic_zip = root / "example-finding.zip"
        final_diagnostic_zip.write_bytes(b"original-final-zip-bytes\n")
        retained = auto_module.retain_unpromoted_archive(staged_diagnostic_zip, retained_dir, base_bundle)
        if retained.parent != retained_dir.resolve() or retained.name == final_diagnostic_zip.name:
            raise SystemExit("FAILED: retained archive path is not explicitly diagnostic")
        if retained.read_bytes() != staged_diagnostic_zip.read_bytes() or final_diagnostic_zip.read_bytes() != b"original-final-zip-bytes\n":
            raise SystemExit("FAILED: diagnostic archive retention changed archive bytes or the final path")
        try:
            auto_module.retain_unpromoted_archive(staged_diagnostic_zip, retained_dir, base_bundle)
        except RuntimeError:
            pass
        else:
            raise SystemExit("FAILED: diagnostic archive retention overwrote an existing diagnostic copy")
        inside_bundle_dir = base_bundle / "forbidden-diagnostics"
        try:
            auto_module.retain_unpromoted_archive(staged_diagnostic_zip, inside_bundle_dir, base_bundle)
        except RuntimeError:
            pass
        else:
            raise SystemExit("FAILED: diagnostic archive retention accepted a directory inside the final bundle")
        shutil.rmtree(inside_bundle_dir, ignore_errors=True)

        final_dir = root / "transaction-final"
        final_dir.mkdir()
        (final_dir / "original.txt").write_text("original bytes\n", encoding="utf-8")
        final_zip = root / "transaction-final.zip"
        with zipfile.ZipFile(final_zip, "w") as handle:
            handle.writestr("transaction-final/original.txt", "original bytes\n")
        before_dir = hashlib.sha256((final_dir / "original.txt").read_bytes()).hexdigest()
        before_zip = hashlib.sha256(final_zip.read_bytes()).hexdigest()
        stage_dir = root / "transaction-stage"
        stage_dir.mkdir()
        (stage_dir / ".zhulong-recording-transaction.json").write_text("{}\n", encoding="utf-8")
        (stage_dir / "recording-evidence.json").write_text(
            json.dumps({"recording_status": "staging", "transaction": {"full_recording_time_validated": False}}) + "\n",
            encoding="utf-8",
        )
        try:
            auto_module.transactional_promote(stage_dir, final_dir, root / "missing-stage.zip", final_zip)
        except RuntimeError as exc:
            if "full recording-time validation" not in str(exc):
                raise
            pass
        else:
            raise SystemExit("FAILED: case 12 unvalidated transaction unexpectedly succeeded")
        (stage_dir / "recording-evidence.json").write_text(
            json.dumps({"recording_status": "passed", "transaction": {"full_recording_time_validated": True}}) + "\n",
            encoding="utf-8",
        )
        try:
            auto_module.transactional_promote(stage_dir, final_dir, root / "missing-stage.zip", final_zip)
        except (RuntimeError, OSError):
            pass
        else:
            raise SystemExit("FAILED: case 12 interrupted transaction unexpectedly succeeded")
        if hashlib.sha256((final_dir / "original.txt").read_bytes()).hexdigest() != before_dir or hashlib.sha256(final_zip.read_bytes()).hexdigest() != before_zip:
            raise SystemExit("FAILED: case 12 rollback changed original bytes")

        positive = run_capture([sys.executable, str(validator_path), "--bundle-dir", str(base_bundle), "--checkpoint-dir", str(base_checkpoint), "--finalize", "--json"], plugin_root)
        positive_result = json.loads(positive)
        if (
            positive_result.get("status") != "passed"
            or positive_result.get("validation_mode") != "full_recording_time"
            or positive_result.get("live_checkpoint_proof_recomputed") is not True
            or positive_result.get("recording_time_observations_authority") != "non_authoritative_consistency_claims"
        ):
            raise SystemExit("FAILED: case 13 sanitized positive fixture")
        finalized_manifest = json.loads((base_bundle / "recording-evidence.json").read_text())
        if finalized_manifest["transaction"]["full_recording_time_validated"] is not True:
            raise SystemExit("FAILED: case 13 full validation did not persist promotion authority")
        artifact_only = json.loads(
            run_capture([sys.executable, str(validator_path), "--bundle-dir", str(base_bundle), "--json"], plugin_root)
        )
        if (
            artifact_only.get("validation_mode") != "artifact_only"
            or artifact_only.get("live_checkpoint_proof_recomputed") is not False
            or artifact_only.get("recording_time_observations_authority") != "non_authoritative_consistency_claims"
        ):
            raise SystemExit("FAILED: case 13 artifact-only revalidation overstated recording-time proof")
        positive_archive = root / "case-13-positive.zip"
        with zipfile.ZipFile(positive_archive, "w", compression=zipfile.ZIP_DEFLATED) as handle:
            for path in sorted(base_bundle.rglob("*")):
                if path.is_file() and not path.is_symlink():
                    handle.write(path, f"{base_bundle.name}/{path.relative_to(base_bundle)}")
        positive_archive_result = run_capture(
            [
                sys.executable,
                str(validator_path),
                "--bundle-dir",
                str(base_bundle),
                "--checkpoint-dir",
                str(base_checkpoint),
                "--archive",
                str(positive_archive),
                "--archive-root",
                base_bundle.name,
                "--json",
            ],
            plugin_root,
        )
        if json.loads(positive_archive_result).get("status") != "passed":
            raise SystemExit("FAILED: case 13 positive archive fixture")

        render_spec = importlib.util.spec_from_file_location("zhulong_render_recording_selftest", plugin_root / "scripts/render_confirmed_vuln_docx.py")
        if render_spec is None or render_spec.loader is None:
            raise SystemExit("FAILED: cannot load renderer for normal replay case")
        render_module = importlib.util.module_from_spec(render_spec)
        sys.modules[render_spec.name] = render_module
        render_spec.loader.exec_module(render_module)
        helper_text = render_module.build_generated_recording_shell(finding(), "en-US", {}, {"generator_options": {"modes": ["quick"]}})
        helper_path = root / "normal-replay-helper.sh"
        helper_path.write_text(helper_text.rsplit("main \"$@\"", 1)[0] + "recording_checkpoint identity 'example-project v0.0.0-test'\n", encoding="utf-8")
        proc = subprocess.run(["sh", "-c", f'. "{helper_path}"'], cwd=plugin_root, capture_output=True, text=True, timeout=3)
        if proc.returncode != 0:
            raise SystemExit(f"FAILED: case 14 normal replay waited/failed: {proc.stdout}{proc.stderr}")

        helper_prefix = helper_text.rsplit("main \"$@\"", 1)[0]

        def run_real_checkpoint_handshake(label: str, writer: object, expect_success: bool) -> None:
            session_root = root / f"checkpoint-{label}"
            events = session_root / "events"
            acknowledgements = session_root / "acks"
            events.mkdir(parents=True)
            acknowledgements.mkdir()
            owner = session_root / "owner.json"
            auto_module.write_json_atomic(owner, {"owner": "zhulong-recording"})
            generated_helper = session_root / "generated-helper.sh"
            generated_helper.write_text(
                helper_prefix
                + "recording_checkpoint identity 'example-project v0.0.0-test'\n"
                + "printf '%s\\n' '__ZHULONG_CHECKPOINT_SENTINEL__'\n",
                encoding="utf-8",
            )
            environment = dict(os.environ)
            environment.update(
                {
                    "ZHULONG_RECORDING_PROTOCOL_VERSION": "1",
                    "ZHULONG_RECORDING_ROOT": str(session_root),
                    "ZHULONG_RECORDING_STAGE_DIR": str(events),
                    "ZHULONG_RECORDING_STAGE_ACK_DIR": str(acknowledgements),
                    "ZHULONG_RECORDING_OWNER_MARKER": str(owner),
                    "ZHULONG_RECORDING_ACK_TIMEOUT_SECONDS": "1",
                    "ZHULONG_RECORDING_ACK_POLL_SECONDS": "0.02",
                }
            )
            checkpoint_proc = subprocess.Popen(
                ["sh", str(generated_helper)],
                cwd=plugin_root,
                env=environment,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            event_path = events / "1-identity.event.json"
            deadline = time.monotonic() + 2.0
            while not event_path.is_file():
                if checkpoint_proc.poll() is not None:
                    stdout, stderr = checkpoint_proc.communicate()
                    raise SystemExit(f"FAILED: checkpoint {label} exited before emitting an event: {stdout}{stderr}")
                if time.monotonic() >= deadline:
                    checkpoint_proc.kill()
                    stdout, stderr = checkpoint_proc.communicate()
                    raise SystemExit(f"FAILED: checkpoint {label} did not emit an event: {stdout}{stderr}")
                time.sleep(0.02)
            event = json.loads(event_path.read_text(encoding="utf-8"))
            acknowledgement = {
                "protocol_version": 1,
                "status": "ack",
                "stage": event["stage"],
                "sequence": event["sequence"],
                "event_timestamp": event["event_timestamp"],
                "expected_marker": event["expected_marker"],
            }
            ack_path = acknowledgements / f"{event['sequence']}-{event['stage']}.ack.json"
            assert callable(writer)
            writer(ack_path, acknowledgement, session_root)
            try:
                stdout, stderr = checkpoint_proc.communicate(timeout=4)
            except subprocess.TimeoutExpired:
                checkpoint_proc.kill()
                stdout, stderr = checkpoint_proc.communicate()
                raise SystemExit(f"FAILED: checkpoint {label} did not reach its short timeout: {stdout}{stderr}")
            sentinel_seen = "__ZHULONG_CHECKPOINT_SENTINEL__" in stdout
            if expect_success:
                if checkpoint_proc.returncode != 0 or not sentinel_seen:
                    raise SystemExit(f"FAILED: checkpoint {label} rejected a semantically valid ack: {stdout}{stderr}")
            elif checkpoint_proc.returncode == 0 or sentinel_seen:
                raise SystemExit(f"FAILED: checkpoint {label} accepted an invalid ack: {stdout}{stderr}")

        def pretty_ack(path: Path, acknowledgement: dict[str, object], _session_root: Path) -> None:
            auto_module.write_json_atomic(path, acknowledgement)

        def compact_ack(path: Path, acknowledgement: dict[str, object], _session_root: Path) -> None:
            path.write_text(json.dumps(acknowledgement, separators=(",", ":")) + "\n", encoding="utf-8")

        def reordered_ack(path: Path, acknowledgement: dict[str, object], _session_root: Path) -> None:
            reordered = {key: acknowledgement[key] for key in reversed(tuple(acknowledgement))}
            path.write_text("\n  " + json.dumps(reordered, indent=4) + "\n", encoding="utf-8")

        def invalid_ack(field: str, value: object) -> object:
            def writer(path: Path, acknowledgement: dict[str, object], _session_root: Path) -> None:
                modified = dict(acknowledgement)
                modified[field] = value
                auto_module.write_json_atomic(path, modified)

            return writer

        def malformed_ack(path: Path, _acknowledgement: dict[str, object], _session_root: Path) -> None:
            path.write_text('{"status": "ack",', encoding="utf-8")

        def symlink_ack(path: Path, acknowledgement: dict[str, object], session_root: Path) -> None:
            target = session_root / "outside-ack.json"
            auto_module.write_json_atomic(target, acknowledgement)
            path.symlink_to(target)

        run_real_checkpoint_handshake("pretty", pretty_ack, True)
        run_real_checkpoint_handshake("compact", compact_ack, True)
        run_real_checkpoint_handshake("reordered", reordered_ack, True)
        run_real_checkpoint_handshake("wrong-status", invalid_ack("status", "ignored"), False)
        run_real_checkpoint_handshake("wrong-stage", invalid_ack("stage", "final_impact"), False)
        run_real_checkpoint_handshake("wrong-sequence", invalid_ack("sequence", 2), False)
        run_real_checkpoint_handshake("string-sequence", invalid_ack("sequence", "1"), False)
        run_real_checkpoint_handshake("wrong-marker", invalid_ack("expected_marker", "unexpected marker"), False)
        run_real_checkpoint_handshake("malformed", malformed_ack, False)
        run_real_checkpoint_handshake("symlink", symlink_ack, False)

        layout_files = ["SKILL.md"] if not (plugin_root / "skills/zhulong/SKILL.md").is_file() else ["skills/zhulong/SKILL.md", "templates/claude-skill/SKILL.md"]
        for rel in required + layout_files:
            if not (plugin_root / rel).is_file():
                raise SystemExit(f"FAILED: case 15 public layout missing {rel}")
        print("RECORDING SELFTEST PASSED: identity/media/screenshot/transaction and semantic checkpoint gates")


def exercise_audit_state_protocol_r2(plugin_root: Path) -> None:
    """Exercise the R2 protocol as a read-only, schema-shaped contract."""

    validator = plugin_root / "scripts/validate_audit_protocol.py"
    fixture_root = plugin_root / "assets/fixtures/audit-state-protocol-r2"
    schemas = {
        "audit-event.schema.json": (
            "https://github.com/Torchbearer127/zhulong/blob/main/assets/schemas/audit-event.schema.json"
        ),
        "stage-status.schema.json": (
            "https://github.com/Torchbearer127/zhulong/blob/main/assets/schemas/stage-status.schema.json"
        ),
    }

    def assert_strict_object_boundaries(value: object, label: str) -> None:
        if isinstance(value, dict):
            if value.get("type") == "object" and value.get("additionalProperties") is not False:
                raise SystemExit(
                    "FAILED: R2 schema object does not reject unknown properties: " + label
                )
            for key, child in value.items():
                assert_strict_object_boundaries(child, f"{label}.{key}")
        elif isinstance(value, list):
            for index, child in enumerate(value):
                assert_strict_object_boundaries(child, f"{label}[{index}]")

    loaded_schemas: dict[str, dict[str, object]] = {}
    for filename, expected_id in schemas.items():
        schema_path = plugin_root / "assets/schemas" / filename
        try:
            schema = json.loads(schema_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            raise SystemExit(f"FAILED: R2 schema is not valid JSON: {schema_path}: {exc}") from exc
        if not isinstance(schema, dict):
            raise SystemExit(f"FAILED: R2 schema root is not an object: {schema_path}")
        if schema.get("$schema") != "https://json-schema.org/draft/2020-12/schema":
            raise SystemExit(f"FAILED: R2 schema has the wrong draft declaration: {schema_path}")
        if schema.get("$id") != expected_id:
            raise SystemExit(f"FAILED: R2 schema has the wrong canonical GitHub $id: {schema_path}")
        assert_strict_object_boundaries(schema, filename)
        loaded_schemas[filename] = schema

    expected_stages = [
        "intake",
        "recon",
        "candidate_generation",
        "triage",
        "verification",
        "severity_escalation",
        "variant_discovery",
        "packaging",
        "finalization",
        "recording",
    ]
    expected_statuses = ["running", "paused", "blocked", "completed"]
    expected_transition_kinds = [
        "start",
        "observe",
        "advance",
        "pause",
        "block",
        "resume",
        "skip",
        "return",
        "reopen",
        "complete",
    ]
    expected_event_types = ["stage_transition", "state_observation", "checkpoint", "recovery", "recording"]
    expected_reason_codes = [
        "normal_progress",
        "operator_request",
        "prerequisite_missing",
        "policy_or_safety_block",
        "verification_blocked",
        "validation_failed",
        "external_dependency",
        "manual_review_required",
        "interrupted",
        "recovery_requested",
        "scope_change",
        "not_applicable",
    ]
    required_event_fields = {
        "schema_version",
        "seq",
        "run_id",
        "ts",
        "stage",
        "event_type",
        "event_name",
        "from_status",
        "to_status",
        "reason_code",
        "subjects",
        "evidence_refs",
        "next_actions",
        "expected_state_revision",
        "details",
    }
    required_state_fields = {
        "schema_version",
        "plugin",
        "plugin_version",
        "run_id",
        "state_revision",
        "last_event_seq",
        "event_log_digest",
        "stage",
        "status",
        "last_event_at",
        "last_event_type",
        "last_event_name",
        "blocker",
        "resume_step",
    }
    event_schema = loaded_schemas["audit-event.schema.json"]
    state_schema = loaded_schemas["stage-status.schema.json"]
    if not required_event_fields.issubset(set(event_schema.get("required", []))):
        raise SystemExit("FAILED: R2 event schema is missing required protocol fields")
    if not required_state_fields.issubset(set(state_schema.get("required", []))):
        raise SystemExit("FAILED: R2 state schema is missing required protocol fields")
    for schema_name, schema in loaded_schemas.items():
        definitions = schema.get("$defs")
        if not isinstance(definitions, dict):
            raise SystemExit(f"FAILED: R2 schema has no definitions: {schema_name}")
        if definitions.get("stage", {}).get("enum") != expected_stages:
            raise SystemExit(f"FAILED: R2 schema stage enum drifted: {schema_name}")
        if definitions.get("status", {}).get("enum") != expected_statuses:
            raise SystemExit(f"FAILED: R2 schema status enum drifted: {schema_name}")
    event_definitions = event_schema["$defs"]
    if not isinstance(event_definitions, dict):
        raise SystemExit("FAILED: R2 event schema definitions are malformed")
    if event_definitions.get("event_type", {}).get("enum") != expected_event_types:
        raise SystemExit("FAILED: R2 event schema event_type enum drifted")
    if event_definitions.get("reason_code", {}).get("enum") != expected_reason_codes:
        raise SystemExit("FAILED: R2 event schema reason_code enum drifted")
    if event_definitions.get("transition_kind", {}).get("enum") != expected_transition_kinds:
        raise SystemExit("FAILED: R2 event schema transition_kind enum drifted")

    protocol_doc = plugin_root / "docs/runner-contracts/audit-state-protocol-r2.md"
    require_text(protocol_doc, "audit-events.jsonl", "R2 audit journal authority wording")
    require_text(protocol_doc, "P9.2", "R2 deferred lock/CAS wording")
    require_text(protocol_doc, "P9.3", "R2 deferred transition graph wording")
    require_text(protocol_doc, "P9.4", "R2 deferred rebuild/migration wording")
    for workflow_doc in [
        plugin_root / "docs/WORKFLOW_DETAILS.md",
        plugin_root / "docs/WORKFLOW_DETAILS.zh-CN.md",
    ]:
        if workflow_doc.exists():
            require_text(
                workflow_doc,
                "runner-contracts/audit-state-protocol-r2.md",
                f"R2 workflow link in {workflow_doc.name}",
            )

    def assert_success(
        flag: str,
        path: Path,
        *,
        expected_mode: str,
        expected_record_count: int,
        expected_transition_policy: str | None = None,
    ) -> None:
        output = run_capture(
            [sys.executable, str(validator), flag, str(path), "--json"],
            plugin_root,
        )
        try:
            payload = json.loads(output)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: R2 validator did not emit JSON for {path}: {output}") from exc
        if payload.get("ok") is not True:
            raise SystemExit(f"FAILED: R2 validator reported unsuccessful fixture: {path}: {payload}")
        if payload.get("mode") != expected_mode:
            raise SystemExit(
                f"FAILED: R2 validator mode mismatch for {path}: "
                f"expected {expected_mode}, got {payload.get('mode')}"
            )
        if payload.get("record_count") != expected_record_count:
            raise SystemExit(
                f"FAILED: R2 validator record count mismatch for {path}: "
                f"expected {expected_record_count}, got {payload.get('record_count')}"
            )
        if (
            expected_transition_policy is not None
            and payload.get("transition_policy") != expected_transition_policy
        ):
            raise SystemExit(
                f"FAILED: R2 validator transition policy mismatch for {path}: "
                f"expected {expected_transition_policy}, got {payload.get('transition_policy')}"
            )

    def assert_failure(
        flag: str,
        path: Path,
        *,
        expected_code: str,
        expected_line: int | None = None,
    ) -> None:
        proc = subprocess.run(
            [sys.executable, str(validator), flag, str(path), "--json"],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        if proc.returncode == 0:
            raise SystemExit(f"FAILED: R2 invalid fixture unexpectedly passed: {path}")
        try:
            payload = json.loads(output)
        except json.JSONDecodeError as exc:
            raise SystemExit(
                f"FAILED: R2 validator error was not deterministic JSON for {path}: {output}"
            ) from exc
        if payload.get("code") != expected_code:
            raise SystemExit(
                f"FAILED: R2 validator error code mismatch for {path}: "
                f"expected {expected_code}, got {payload.get('code')}"
            )
        if expected_line is not None and payload.get("line") != expected_line:
            raise SystemExit(
                f"FAILED: R2 validator line mismatch for {path}: "
                f"expected {expected_line}, got {payload.get('line')}"
            )

    fixture_snapshot = {
        path.relative_to(fixture_root).as_posix(): path.read_bytes()
        for path in sorted(fixture_root.iterdir())
        if path.is_file()
    }
    if set(fixture_snapshot) != {
        Path(rel).name for rel in AUDIT_STATE_PROTOCOL_R2_FIXTURE_FILES
    }:
        raise SystemExit("FAILED: R2 fixture inventory differs from the selftest inventory")

    valid_event = json.loads((fixture_root / "valid-event-r2.json").read_text(encoding="utf-8"))
    valid_state = json.loads((fixture_root / "valid-state-r2.json").read_text(encoding="utf-8"))
    for event_field, state_field in [
        ("run_id", "run_id"),
        ("seq", "last_event_seq"),
        ("ts", "last_event_at"),
        ("event_type", "last_event_type"),
        ("event_name", "last_event_name"),
    ]:
        if valid_event.get(event_field) != valid_state.get(state_field):
            raise SystemExit(
                "FAILED: positive R2 state fixture does not match its event fixture: "
                f"{event_field} != {state_field}"
            )

    assert_success(
        "--event",
        fixture_root / "valid-event-r2.json",
        expected_mode="r2",
        expected_record_count=1,
    )
    assert_success(
        "--event",
        fixture_root / "valid-new-event-r2.json",
        expected_mode="r2",
        expected_record_count=1,
        expected_transition_policy="transition_policy_v1",
    )
    assert_success(
        "--state",
        fixture_root / "valid-state-r2.json",
        expected_mode="r2",
        expected_record_count=1,
    )
    assert_success(
        "--events-jsonl",
        fixture_root / "valid-events-r2.jsonl",
        expected_mode="r2",
        expected_record_count=2,
        expected_transition_policy="pre_policy_r2",
    )
    assert_success(
        "--events-jsonl",
        fixture_root / "valid-policy-events-r2.jsonl",
        expected_mode="r2",
        expected_record_count=4,
        expected_transition_policy="transition_policy_v1",
    )
    assert_success(
        "--event",
        fixture_root / "legacy-event-r1.json",
        expected_mode="legacy_r1",
        expected_record_count=1,
    )
    assert_success(
        "--state",
        fixture_root / "legacy-state-r1.json",
        expected_mode="legacy_r1",
        expected_record_count=1,
    )

    for flag, filename, code, line in [
        ("--event", "invalid-event-missing-seq.json", "MISSING_REQUIRED_FIELD", None),
        ("--event", "invalid-event-stage.json", "INVALID_STAGE", None),
        ("--event", "invalid-event-status.json", "INVALID_STATUS", None),
        ("--event", "invalid-event-reason-code.json", "INVALID_REASON_CODE", None),
        ("--event", "invalid-event-absolute-evidence.json", "INVALID_EVIDENCE_REF", None),
        ("--event", "invalid-event-parent-evidence.json", "INVALID_EVIDENCE_REF", None),
        ("--event", "invalid-event-file-uri-evidence.json", "INVALID_EVIDENCE_REF", None),
        ("--event", "invalid-event-https-uri-evidence.json", "INVALID_EVIDENCE_REF", None),
        ("--event", "invalid-event-unexpected-property.json", "UNEXPECTED_PROPERTY", None),
        ("--event", "invalid-event-incomplete-transition-metadata.json", "TRANSITION_METADATA_INCOMPLETE", None),
        ("--event", "invalid-event-unknown-transition-kind.json", "INVALID_TRANSITION_KIND", None),
        ("--state", "invalid-state-negative-revision.json", "INVALID_STATE_REVISION", None),
        ("--state", "invalid-state-malformed-digest.json", "INVALID_EVENT_LOG_DIGEST", None),
        ("--events-jsonl", "invalid-events-zero-seq.jsonl", "INVALID_SEQ", 1),
        ("--events-jsonl", "invalid-events-duplicate-seq.jsonl", "DUPLICATE_SEQ", 2),
        ("--events-jsonl", "invalid-events-nonmonotonic-seq.jsonl", "NON_MONOTONIC_SEQ", 2),
        ("--events-jsonl", "invalid-events-truncated.jsonl", "MALFORMED_JSON", 2),
        ("--events-jsonl", "invalid-events-seq-gap.jsonl", "JOURNAL_SEQ_GAP", 2),
        ("--events-jsonl", "invalid-events-revision-chain.jsonl", "JOURNAL_REVISION_CHAIN_MISMATCH", 1),
        ("--events-jsonl", "invalid-events-run-id-drift.jsonl", "JOURNAL_RUN_ID_DRIFT", 2),
        ("--events-jsonl", "invalid-events-middle-corruption.jsonl", "MALFORMED_JSON", 2),
    ]:
        assert_failure(
            flag,
            fixture_root / filename,
            expected_code=code,
            expected_line=line,
        )

    with tempfile.TemporaryDirectory(prefix="zhulong-audit-state-r2-") as temp_dir:
        temp_root = Path(temp_dir)
        unsupported_event = json.loads((fixture_root / "valid-event-r2.json").read_text(encoding="utf-8"))
        unsupported_event["schema_version"] = 9
        unsupported_event_path = temp_root / "unsupported-event.json"
        unsupported_event_path.write_text(
            json.dumps(unsupported_event, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        assert_failure(
            "--event",
            unsupported_event_path,
            expected_code="SCHEMA_VERSION_UNSUPPORTED",
        )

        root_array_path = temp_root / "root-array.json"
        root_array_path.write_text("[]\n", encoding="utf-8")
        assert_failure("--event", root_array_path, expected_code="ROOT_NOT_OBJECT")

        paused_state = json.loads((fixture_root / "valid-state-r2.json").read_text(encoding="utf-8"))
        paused_state["status"] = "paused"
        paused_state_path = temp_root / "paused-without-blocker.json"
        paused_state_path.write_text(
            json.dumps(paused_state, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        assert_failure("--state", paused_state_path, expected_code="MISSING_BLOCKER")

        legacy_journal = temp_root / "legacy-events.jsonl"
        legacy_record = json.loads((fixture_root / "legacy-event-r1.json").read_text(encoding="utf-8"))
        legacy_journal.write_text(
            json.dumps(legacy_record, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        assert_success(
            "--events-jsonl",
            legacy_journal,
            expected_mode="legacy_r1",
            expected_record_count=1,
        )
        mixed_journal = temp_root / "mixed-events.jsonl"
        r2_record = json.loads((fixture_root / "valid-event-r2.json").read_text(encoding="utf-8"))
        mixed_journal.write_text(
            json.dumps(legacy_record, ensure_ascii=False, sort_keys=True)
            + "\n"
            + json.dumps(r2_record, ensure_ascii=False, sort_keys=True)
            + "\n",
            encoding="utf-8",
        )
        assert_failure(
            "--events-jsonl",
            mixed_journal,
            expected_code="MIXED_JOURNAL_MODE",
            expected_line=2,
        )

        finalization_event = json.loads((fixture_root / "valid-event-r2.json").read_text(encoding="utf-8"))
        finalization_event.update(
            {
                "stage": "finalization",
                "event_type": "stage_transition",
                "event_name": "finalization_succeeded",
                "from_status": "running",
                "to_status": "completed",
                "expected_state_revision": 9,
                "next_actions": [],
                "details": {"summary": "Shape-only finalization fixture; no completion is asserted."},
            }
        )
        no_side_effect_workspace = temp_root / "no-side-effect-workspace"
        no_side_effect_workspace.mkdir()
        finalization_event_path = no_side_effect_workspace / "shape-valid-finalization-event-r2.json"
        finalization_event_path.write_text(
            json.dumps(finalization_event, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        assert_success(
            "--event",
            finalization_event_path,
            expected_mode="r2",
            expected_record_count=1,
        )
        for generated in [
            "audit-events.jsonl",
            "audit-disposition.json",
            "stage-status.json",
            "confirmed",
            "handoff-summary.md",
            "SUMMARY.md",
            "finalization-summary.md",
        ]:
            if (no_side_effect_workspace / generated).exists():
                raise SystemExit(
                    "FAILED: R2 read-only validator created a completion artifact: " + generated
                )

        r1_repo = temp_root / "legacy-r1-repo"
        r1_workspace = r1_repo / "security-research-r1"
        r1_workspace.mkdir(parents=True)
        r1_state = json.loads((fixture_root / "legacy-state-r1.json").read_text(encoding="utf-8"))
        r1_state.update(
            {
                "stage": "completed",
                "status": "completed",
                "workspace": r1_workspace.name,
                "target_repo": ".",
                "last_event": "finalization_succeeded",
                "last_message": "Legacy finalization claim fixture.",
            }
        )
        r1_event = json.loads((fixture_root / "legacy-event-r1.json").read_text(encoding="utf-8"))
        r1_event.update(
            {
                "event": "finalization_succeeded",
                "stage": "finalization",
                "status": "ok",
                "message": "Legacy finalization claim fixture.",
                "details": {
                    "result": "completed_with_confirmed_bundles",
                    "docker_clean": True,
                },
            }
        )
        (r1_workspace / "stage-status.json").write_text(
            json.dumps(r1_state, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        (r1_workspace / "audit-events.jsonl").write_text(
            json.dumps(r1_event, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        run_expect_fail(
            [
                sys.executable,
                str(plugin_root / "scripts/validate_workspace_state.py"),
                "--workspace-dir",
                str(r1_workspace),
                "--repo-root",
                str(r1_repo),
                "--skip-latest-check",
            ],
            plugin_root,
            "validated_confirmed_bundle_count=0",
        )
        assertion = subprocess.run(
            [
                sys.executable,
                str(plugin_root / "scripts/assert_finalized_workspace.py"),
                "--workspace-dir",
                str(r1_workspace),
                "--json",
            ],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        assertion_output = ((assertion.stdout or "") + (assertion.stderr or "")).strip()
        if assertion.returncode == 0:
            raise SystemExit("FAILED: fake legacy R1 completion workspace unexpectedly passed finalization")
        try:
            assertion_payload = json.loads(assertion_output)
        except json.JSONDecodeError as exc:
            raise SystemExit(
                "FAILED: finalization assertion did not emit JSON for fake R1 workspace: "
                + assertion_output
            ) from exc
        assertion_errors = assertion_payload.get("errors")
        if not isinstance(assertion_errors, list):
            raise SystemExit("FAILED: fake R1 finalization assertion did not return an error list")
        combined_errors = "\n".join(str(error) for error in assertion_errors)
        if "missing finalization_succeeded event" in combined_errors:
            raise SystemExit("FAILED: finalization assertion did not recognize the legacy R1 success event")
        if "validated_confirmed_bundle_count=0" not in combined_errors:
            raise SystemExit("FAILED: fake R1 completion did not fail the confirmed bundle gate")
        if "audit-disposition.json" not in combined_errors:
            raise SystemExit("FAILED: fake R1 completion did not fail the disposition ledger gate")
        if "docker-cleanliness-status.json is missing" not in combined_errors:
            raise SystemExit("FAILED: fake R1 completion did not fail the Docker cleanliness evidence gate")

    fixture_after = {
        path.relative_to(fixture_root).as_posix(): path.read_bytes()
        for path in sorted(fixture_root.iterdir())
        if path.is_file()
    }
    if fixture_after != fixture_snapshot:
        raise SystemExit("FAILED: R2 validator or selftest mutated protocol fixtures")
    print("AUDIT STATE PROTOCOL R2 SELFTEST PASSED: strict/R1/read-only/sequence/finalization gates")


def exercise_audit_state_protocol_closure(plugin_root: Path) -> None:
    """Keep the P9.5 runner independently runnable and impossible to omit here."""
    runner = plugin_root / "scripts/selftest_audit_state_protocol.py"
    output = run_capture([sys.executable, str(runner), "--json"], plugin_root)
    try:
        summary = json.loads(output)
    except json.JSONDecodeError as exc:
        raise SystemExit(f"FAILED: P9.5 state protocol runner did not emit JSON: {output}") from exc
    if summary.get("ok") is not True:
        raise SystemExit(f"FAILED: P9.5 state protocol runner failed: {summary}")
    if summary.get("runner_version") != "p9.5-r2" or summary.get("manifest_schema_version") != 2 or not summary.get("manifest_digest"):
        raise SystemExit("FAILED: P9.5 state protocol runner summary lacks deterministic identity")
    if set(summary.get("fixture_ids", [])) != {
        "blocked-verification", "completed-no-confirmed", "damaged-journal-middle", "damaged-journal-tail",
        "legacy-r1", "modern-r2-running", "partial-commit-rebuild", "recording-entered-completed",
        "recording-optional-skipped", "state-only-fake-completed", "validated-confirmed-bundle", "variant-candidate-returned",
    }:
        raise SystemExit("FAILED: P9.5 fixture inventory drifted")
    if set(summary.get("concurrency", {})) != {
        "current_revision_24", "explicit_cas_12", "independent_workspaces_12x12", "invalid_transition_vs_valid_writer",
        "held_lock_timeout_and_reuse", "lock_owner_exit_and_reuse", "writer_vs_rebuild_apply", "two_rebuild_applies", "separate_r1_r2_writes",
    }:
        raise SystemExit("FAILED: P9.5 concurrency matrix drifted")
    if [item.get("case_id") for item in summary.get("fault_cases", [])] != [
        "before_journal_open", "partial_journal_bytes", "after_journal_fsync", "state_temp_partial", "before_replace",
        "after_replace_before_dirsync", "append_failure", "state_replace_failure", "recovery_replace_failure", "lock_owner_exit",
    ] or not all(item.get("asserted") is True for item in summary.get("fault_cases", [])):
        raise SystemExit("FAILED: P9.5 hard-exit matrix drifted")
    metrics = summary.get("metrics", {})
    if not isinstance(metrics, dict) or any(value != 0 for value in metrics.values()):
        raise SystemExit("FAILED: P9.5 state protocol runner reported an immutability violation")


def exercise_handoff_checkpoint_contract(plugin_root: Path) -> None:
    """Exercise the production handoff/checkpoint CLIs with sanitized temp workspaces."""
    fixture_root = plugin_root / "assets/fixtures/handoff-checkpoint"
    manifest = json.loads((fixture_root / "manifest.json").read_text(encoding="utf-8"))
    required_positive = {
        "r2-running-candidate-blocker", "paused-blocked-recovery-material",
        "completed-no-confirmed", "completed-with-validated-bundle",
        "same-revision-checkpoint-idempotence", "legal-historical-checkpoint",
    }
    required_negative = {
        "malformed-journal-or-state", "stale-handoff-revision", "forged-confirmed-count",
        "tested-ref-or-artifact-digest-drift", "checkpoint-filename-mismatch",
        "same-revision-conflicting-bytes", "absolute-uri-traversal-backslash-path",
        "symlink-input-output-parent", "concurrent-event-append", "atomic-temp-fsync-replace-fault",
        "advisory-notes-cannot-confirm", "docker-evidence-without-validated-bundle",
        "unvalidated-recording-manifest", "missing-or-invalid-validator",
        "completion-claim-without-validated-bundle", "validated-bundle-removed-after-completion",
        "protocol-aware-completion-source-attribution",
    }
    if not required_positive.issubset(set(manifest.get("positive_cases", []))) or not required_negative.issubset(set(manifest.get("negative_cases", []))):
        raise SystemExit("FAILED: handoff/checkpoint fixture manifest is incomplete")

    render_cli = plugin_root / "scripts/render_handoff_state.py"
    validate_cli = plugin_root / "scripts/validate_handoff_state.py"
    create_cli = plugin_root / "scripts/create_workspace_checkpoint.py"
    checkpoint_cli = plugin_root / "scripts/validate_workspace_checkpoint.py"
    event_writer = plugin_root / "scripts/write_audit_event.py"
    require_files(plugin_root, [
        "assets/schemas/handoff-state.schema.json",
        "assets/schemas/workspace-checkpoint.schema.json",
        "scripts/render_handoff_state.py",
        "scripts/validate_handoff_state.py",
        "scripts/create_workspace_checkpoint.py",
        "scripts/validate_workspace_checkpoint.py",
    ], "handoff/checkpoint contract")

    def new_workspace(root: Path, *, stage: str = "intake", status: str = "running") -> Path:
        workspace = root / "workspace"
        workspace.mkdir()
        command = [
            sys.executable, str(event_writer), "--workspace-dir", str(workspace),
            "--plugin-version", "selftest", "--event", "workspace_created",
            "--stage", stage, "--status", status, "--transition-kind", "start",
            "--event-status", "ok", "--message", "sanitized fixture", "--accept-current-revision",
        ]
        run(command, plugin_root)
        return workspace

    def append_event(
        workspace: Path,
        event_name: str,
        *,
        stage: str,
        status: str,
        transition_kind: str,
        blocker: str | None = None,
        resume_step: str | None = None,
        details: dict[str, Any] | None = None,
    ) -> None:
        command = [
            sys.executable, str(event_writer), "--workspace-dir", str(workspace),
            "--plugin-version", "selftest", "--event", event_name,
            "--stage", stage, "--status", status, "--transition-kind", transition_kind,
            "--event-status", "ok", "--message", "sanitized fixture", "--accept-current-revision",
        ]
        if blocker is not None:
            command.extend(["--blocker", blocker])
        if resume_step is not None:
            command.extend(["--resume-step", resume_step])
        if details is not None:
            command.extend(["--details-json", json.dumps(details, ensure_ascii=False, sort_keys=True)])
        run(command, plugin_root)

    def render(workspace: Path, *, expected: int = 0, env: dict[str, str] | None = None) -> dict[str, Any] | None:
        command = [sys.executable, str(render_cli), "--workspace-dir", str(workspace), "--repo-root", str(workspace.parent), "--json"]
        output = run_capture_with_env(command, plugin_root, env or {}, expected_returncode=expected)
        if expected:
            return None
        try:
            return json.loads(output)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: handoff renderer did not emit JSON: {output}") from exc

    def validate(workspace: Path, *, expected: int = 0, checkpoint: str | None = None) -> dict[str, Any] | None:
        if checkpoint is None:
            command = [sys.executable, str(validate_cli), "--workspace-dir", str(workspace), "--repo-root", str(workspace.parent), "--json"]
        else:
            command = [sys.executable, str(checkpoint_cli), "--workspace-dir", str(workspace), "--repo-root", str(workspace.parent), "--checkpoint", checkpoint, "--json"]
        output = run_capture_with_env(command, plugin_root, {}, expected_returncode=expected)
        if expected:
            return None
        try:
            return json.loads(output)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: handoff/checkpoint validator did not emit JSON: {output}") from exc

    def completion_fixture(workspace: Path, result: Any) -> None:
        for event_name, stage in (
            ("completion_recon", "recon"),
            ("completion_candidates", "candidate_generation"),
            ("completion_triage", "triage"),
            ("completion_verification", "verification"),
            ("completion_finalization", "finalization"),
        ):
            append_event(workspace, event_name, stage=stage, status="running", transition_kind="advance")
        append_event(
            workspace,
            "finalization_succeeded",
            stage="finalization",
            status="completed",
            transition_kind="complete",
            details={"result": result},
        )

    def completion_issue_codes(state: dict[str, Any]) -> set[str]:
        return {
            str(item.get("code"))
            for item in state.get("integrity", {}).get("issues", [])
            if isinstance(item, dict) and str(item.get("code") or "").startswith("COMPLETION_")
        }

    def completion_issues_by_code(state: dict[str, Any]) -> dict[str, dict[str, Any]]:
        return {
            str(item.get("code")): item
            for item in state.get("integrity", {}).get("issues", [])
            if isinstance(item, dict) and str(item.get("code") or "").startswith("COMPLETION_")
        }

    def assert_authority_unchanged(workspace: Path, before: dict[str, bytes | None], label: str) -> None:
        for relative, original in before.items():
            path = workspace / relative
            current = path.read_bytes() if path.exists() else None
            if current != original:
                raise SystemExit(f"FAILED: {label} validator changed authoritative {relative}")

    from workspace_state import _completion_result_from_authority, _completion_source_from_protocol

    if _completion_result_from_authority(
        {"result": "completed_with_confirmed_bundles"}, [], protocol_mode="legacy_r1"
    ) != "completed_with_confirmed_bundles":
        raise SystemExit("FAILED: R1 top-level completion result compatibility regressed")
    if _completion_source_from_protocol("legacy_r1") != (
        "stage-status.json", "legacy stage-status.json completion result"
    ) or _completion_source_from_protocol("r2") != (
        "audit-events.jsonl", "finalization_succeeded event completion result"
    ):
        raise SystemExit("FAILED: protocol-aware completion source mapping drifted")
    if _completion_result_from_authority(
        {},
        [
            {"event": "result_observed", "details": {"result": "completed_with_confirmed_bundles"}},
            {"event": "finalization_succeeded", "details": {"result": "completed_no_confirmed_findings"}},
        ],
        protocol_mode="r2",
    ) != "completed_no_confirmed_findings":
        raise SystemExit("FAILED: R2 completion result did not use the canonical finalization event")
    if _completion_result_from_authority(
        {},
        [
            {"event": "finalization_succeeded", "details": {"result": "completed_with_confirmed_bundles"}},
            {"event": "finalization_failed", "details": {}},
        ],
        protocol_mode="r2",
    ):
        raise SystemExit("FAILED: a later finalization failure did not supersede the prior completion claim")

    with tempfile.TemporaryDirectory(prefix="zhulong-handoff-r1-source-selftest-") as r1_temp:
        r1_root = Path(r1_temp)
        r1_workspace = r1_root / "workspace"
        r1_workspace.mkdir()
        write_json_fixture(r1_workspace / "stage-status.json", {
            "schema_version": 1,
            "plugin": "zhulong",
            "plugin_version": "selftest",
            "stage": "completed",
            "status": "completed",
            "last_event_at": "2026-07-22T00:00:00Z",
            "blocker": None,
            "resume_step": None,
            "workspace": "workspace",
            "target_repo": ".",
            "last_event": "finalization_succeeded",
            "last_message": "sanitized fixture",
            "result": "completed_with_confirmed_bundles",
        })
        (r1_workspace / "audit-events.jsonl").write_text(
            json.dumps({
                "ts": "2026-07-22T00:00:00Z",
                "event": "finalization_succeeded",
                "stage": "completed",
                "status": "completed",
                "message": "sanitized fixture",
                "details": {},
            }, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n",
            encoding="utf-8",
        )
        r1_state = render(r1_workspace).get("state", {})
        r1_claim_issue = completion_issues_by_code(r1_state).get("COMPLETION_CLAIM_UNSUPPORTED", {})
        if (
            r1_claim_issue.get("path") != "stage-status.json"
            or "legacy stage-status.json completion result" not in str(r1_claim_issue.get("message") or "")
            or "finalization_succeeded event" in str(r1_claim_issue.get("message") or "")
        ):
            raise SystemExit("FAILED: R1 completion diagnostic no longer identifies legacy state authority")
        r1_workspace_output = run_capture_with_env([
            sys.executable, str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir", str(r1_workspace), "--repo-root", str(r1_root), "--skip-latest-check",
        ], plugin_root, {}, expected_returncode=1)
        if (
            "legacy stage-status.json completion result" not in r1_workspace_output
            or "finalization_succeeded event completion result" in r1_workspace_output
        ):
            raise SystemExit("FAILED: R1 workspace completion diagnostic was attributed to R2 event metadata")

    with tempfile.TemporaryDirectory(prefix="zhulong-handoff-selftest-") as temp:
        root = Path(temp)
        workspace = new_workspace(root)
        candidate_path = workspace / "candidates/CAND-0001/candidate.json"
        candidate_path.parent.mkdir(parents=True)
        write_json_fixture(candidate_path, valid_candidate_contract())
        verdict_path = workspace / "verifier/CAND-0001/verifier-verdict.json"
        verdict_path.parent.mkdir(parents=True)
        write_json_fixture(verdict_path, valid_verifier_verdict())
        (workspace / "agent-notes.md").write_text("Confirmed bundles: 99\n", encoding="utf-8")
        first = render(workspace)
        state = first.get("state") if isinstance(first, dict) else None
        if not isinstance(state, dict) or state.get("counts", {}).get("candidates") != 1 or state.get("counts", {}).get("verdicts") != 1 or state.get("counts", {}).get("validated_confirmed_bundles") != 0 or state.get("integrity", {}).get("overall") != "blocked":
            raise SystemExit("FAILED: running candidate handoff did not remain conservatively blocked")
        if state.get("advisory_notes", {}).get("status") != "advisory":
            raise SystemExit("FAILED: agent-notes.md was not recorded as advisory")
        if validate(workspace).get("ok") is not True:
            raise SystemExit("FAILED: current handoff did not validate")
        original_verdict = verdict_path.read_bytes()
        drifted_verdict = json.loads(original_verdict.decode("utf-8"))
        drifted_verdict["target_ref"]["tested_ref"] = "different-structured-ref"
        verdict_path.write_text(json.dumps(drifted_verdict, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
        validate(workspace, expected=1)
        verdict_path.write_bytes(original_verdict)
        original_candidate = candidate_path.read_bytes()
        drifted_candidate = json.loads(original_candidate.decode("utf-8"))
        drifted_candidate["title"] = "changed after handoff"
        candidate_path.write_text(json.dumps(drifted_candidate, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
        validate(workspace, expected=1)
        candidate_path.write_bytes(original_candidate)
        render(workspace)

        disposition_root = root / "candidate-disposition"
        disposition_root.mkdir()
        disposition_workspace = new_workspace(disposition_root)
        disposition_candidate = disposition_workspace / "candidates/CAND-0001/candidate.json"
        disposition_candidate.parent.mkdir(parents=True)
        write_json_fixture(disposition_candidate, valid_candidate_contract())
        run([sys.executable, str(plugin_root / "scripts/audit_disposition.py"), "--workspace-dir", str(disposition_workspace), "--write"], plugin_root)
        disposition_state = render(disposition_workspace).get("state", {})
        if disposition_state.get("counts", {}).get("dispositions") != 1 or disposition_state.get("disposition", {}).get("status") != "valid":
            raise SystemExit("FAILED: disposition ledger IDs/count were not derived from the existing validator")

        blocked_root = root / "paused-blocked"
        blocked_root.mkdir()
        blocked_workspace = new_workspace(blocked_root)
        append_event(
            blocked_workspace,
            "selftest_blocked",
            stage="intake",
            status="blocked",
            transition_kind="block",
            blocker="Docker gate requires operator review",
            resume_step="check-docker-gate",
        )
        blocked_state = render(blocked_workspace).get("state", {})
        if blocked_state.get("status") != "blocked" or not blocked_state.get("blocker", {}).get("active") or blocked_state.get("resume", {}).get("entrypoint") != "resume_stage":
            raise SystemExit("FAILED: blocked workspace did not expose a safe recovery handoff")

        completed_root = root / "completed-no-confirmed"
        completed_root.mkdir()
        completed_workspace = new_workspace(completed_root)
        for event_name, stage in (
            ("selftest_recon", "recon"),
            ("selftest_candidates", "candidate_generation"),
            ("selftest_triage", "triage"),
            ("selftest_verification", "verification"),
            ("selftest_finalization", "finalization"),
        ):
            append_event(completed_workspace, event_name, stage=stage, status="running", transition_kind="advance")
        append_event(
            completed_workspace,
            "finalization_succeeded",
            stage="finalization",
            status="completed",
            transition_kind="complete",
            details={"result": "completed_no_confirmed_findings"},
        )
        completed_state = render(completed_workspace).get("state", {})
        if completed_state.get("finalization", {}).get("status") != "completed" or completed_state.get("counts", {}).get("validated_confirmed_bundles") != 0:
            raise SystemExit("FAILED: completed-no-confirmed handoff did not preserve conservative bundle facts")
        if not {"COMPLETION_DISPOSITION_UNVERIFIABLE", "COMPLETION_DOCKER_GATE_UNSATISFIED"}.issubset(completion_issue_codes(completed_state)):
            raise SystemExit("FAILED: R2 completed_no_confirmed_findings metadata was not recognized conservatively")
        completed_issues = completion_issues_by_code(completed_state)
        if (
            completed_issues["COMPLETION_DISPOSITION_UNVERIFIABLE"].get("path") != "audit-disposition.json"
            or completed_issues["COMPLETION_DOCKER_GATE_UNSATISFIED"].get("path") != "docker/docker-cleanliness-status.json"
        ):
            raise SystemExit("FAILED: disposition or Docker completion issue source path drifted")

        completion_claim_root = root / "completion-claim-without-validated-bundle"
        completion_claim_root.mkdir()
        completion_claim_workspace = new_workspace(completion_claim_root)
        completion_fixture(completion_claim_workspace, "completed_with_confirmed_bundles")
        completion_claim_state = render(completion_claim_workspace).get("state", {})
        completion_codes = completion_issue_codes(completion_claim_state)
        if completion_claim_state.get("integrity", {}).get("overall") != "blocked" or "COMPLETION_CLAIM_UNSUPPORTED" not in completion_codes:
            raise SystemExit("FAILED: R2 completion claim with zero validated bundles did not fail closed")
        completion_claim_issue = completion_issues_by_code(completion_claim_state).get("COMPLETION_CLAIM_UNSUPPORTED", {})
        if (
            completion_claim_issue.get("path") != "audit-events.jsonl"
            or "finalization_succeeded event completion result" not in str(completion_claim_issue.get("message") or "")
        ):
            raise SystemExit("FAILED: R2 completion claim diagnostic did not identify its finalization event authority")
        authority_before = {
            relative: (completion_claim_workspace / relative).read_bytes() if (completion_claim_workspace / relative).exists() else None
            for relative in ("audit-events.jsonl", "stage-status.json", "audit-disposition.json")
        }
        handoff_failure = json.loads(run_capture_with_env([
            sys.executable, str(validate_cli), "--workspace-dir", str(completion_claim_workspace),
            "--repo-root", str(completion_claim_root), "--json",
        ], plugin_root, {}, expected_returncode=1))
        if handoff_failure.get("classification") == "current" or "COMPLETION_CLAIM_UNSUPPORTED" not in handoff_failure.get("issue_codes", []):
            raise SystemExit("FAILED: handoff validator accepted an unsupported completion claim as current")
        completion_workspace_output = run_capture_with_env([
            sys.executable, str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir", str(completion_claim_workspace), "--repo-root", str(completion_claim_root), "--skip-latest-check",
        ], plugin_root, {}, expected_returncode=1)
        if (
            "finalization_succeeded event completion result" not in completion_workspace_output
            or "validated_confirmed_bundle_count=0" not in completion_workspace_output
            or "stage-status.json declares" in completion_workspace_output
        ):
            raise SystemExit("FAILED: R2 workspace completion diagnostic has inaccurate source attribution")
        assert_authority_unchanged(completion_claim_workspace, authority_before, "completion-claim")

        injected_result_root = root / "non-finalization-result-metadata"
        injected_result_root.mkdir()
        injected_result_workspace = new_workspace(injected_result_root)
        append_event(
            injected_result_workspace,
            "completion_result_observation",
            stage="intake",
            status="running",
            transition_kind="observe",
            details={"result": "completed_with_confirmed_bundles"},
        )
        if completion_issue_codes(render(injected_result_workspace).get("state", {})):
            raise SystemExit("FAILED: non-finalization result metadata created a completion claim")

        for label, result in (("unknown", "unexpected_completion_value"), ("non-string", True)):
            result_root = root / f"completion-result-{label}"
            result_root.mkdir()
            result_workspace = new_workspace(result_root)
            completion_fixture(result_workspace, result)
            if completion_issue_codes(render(result_workspace).get("state", {})):
                raise SystemExit("FAILED: unknown or non-string R2 result was mistaken for a valid completion result")

        bundle_root = root / "completed-with-validated-bundle"
        bundle_root.mkdir()
        bundle_workspace = new_workspace(bundle_root)
        (bundle_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": bundle_workspace.name,
                "workspace_created_at": "2026-07-21T00:00:00Z",
                "confirmed_output_dir": f"{bundle_workspace.name}/confirmed",
            }, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        bundle_slug = build_wrapper_source_finding(plugin_root, bundle_root, bundle_workspace)
        bundle_contract = build_wrapper_contract(bundle_workspace, bundle_slug)
        run([
            sys.executable, str(plugin_root / "scripts/build_confirmed_bundle.py"),
            "--workspace-dir", str(bundle_workspace), "--repo-root", str(bundle_root),
            "--contract", str(bundle_contract), "--language", "zh-CN",
        ], plugin_root)
        write_finalization_variant_artifacts(bundle_workspace)
        (bundle_workspace / "docker").mkdir(parents=True, exist_ok=True)
        (bundle_workspace / "docker/docker-cleanliness-status.json").write_text(
            json.dumps({"clean": True, "strict": True}, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        (bundle_workspace / "docker/docker-resource-baseline.json").write_text(
            json.dumps({"baseline": "selftest"}, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        for event_name, stage in (
            ("bundle_recon", "recon"),
            ("bundle_candidates", "candidate_generation"),
            ("bundle_triage", "triage"),
            ("bundle_verification", "verification"),
            ("bundle_finalization", "finalization"),
        ):
            append_event(bundle_workspace, event_name, stage=stage, status="running", transition_kind="advance")
        append_event(
            bundle_workspace,
            "finalization_succeeded",
            stage="finalization",
            status="completed",
            transition_kind="complete",
            details={"result": "completed_with_confirmed_bundles", "docker_clean": True},
        )
        bundle_state = render(bundle_workspace).get("state", {})
        if bundle_state.get("counts", {}).get("validated_confirmed_bundles") != 1 or bundle_state.get("variant_analysis", {}).get("status") != "completed":
            raise SystemExit("FAILED: completed-with-validated-bundle handoff did not retain bundle and seeded-variant gates")
        variant_candidates = bundle_workspace / "evidence/variant-analysis/variant-candidates.jsonl"
        variant_candidates_bytes = variant_candidates.read_bytes()
        variant_candidates.unlink()
        variant_blocked_state = render(bundle_workspace).get("state", {})
        if variant_blocked_state.get("integrity", {}).get("overall") != "blocked" or "COMPLETION_VARIANT_GATE_UNSATISFIED" not in completion_issue_codes(variant_blocked_state):
            raise SystemExit("FAILED: confirmed completion without formal variant analysis did not fail closed")
        variant_issue = completion_issues_by_code(variant_blocked_state).get("COMPLETION_VARIANT_GATE_UNSATISFIED", {})
        if (
            variant_issue.get("path") != "audit-events.jsonl"
            or "finalization_succeeded event completion result" not in str(variant_issue.get("message") or "")
        ):
            raise SystemExit("FAILED: R2 variant completion diagnostic did not identify its finalization event authority")
        variant_candidates.write_bytes(variant_candidates_bytes)
        render(bundle_workspace)
        bundle_dir = bundle_workspace / "confirmed" / bundle_slug
        if not bundle_dir.is_dir():
            raise SystemExit("FAILED: completion regression fixture has no validated bundle to remove")
        shutil.rmtree(bundle_dir)
        removed_bundle_state = render(bundle_workspace).get("state", {})
        if removed_bundle_state.get("integrity", {}).get("overall") != "blocked" or "COMPLETION_CLAIM_UNSUPPORTED" not in completion_issue_codes(removed_bundle_state):
            raise SystemExit("FAILED: removing a finalized confirmed bundle did not block completion integrity")
        removed_authority_before = {
            relative: (bundle_workspace / relative).read_bytes() if (bundle_workspace / relative).exists() else None
            for relative in ("audit-events.jsonl", "stage-status.json", "audit-disposition.json")
        }
        removed_handoff_failure = json.loads(run_capture_with_env([
            sys.executable, str(validate_cli), "--workspace-dir", str(bundle_workspace),
            "--repo-root", str(bundle_root), "--json",
        ], plugin_root, {}, expected_returncode=1))
        if removed_handoff_failure.get("classification") == "current" or "COMPLETION_CLAIM_UNSUPPORTED" not in removed_handoff_failure.get("issue_codes", []):
            raise SystemExit("FAILED: handoff validator accepted a post-finalization bundle removal as current")
        run_expect_fail([
            sys.executable, str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir", str(bundle_workspace), "--repo-root", str(bundle_root), "--skip-latest-check",
        ], plugin_root, "validated_confirmed_bundle_count=0")
        assert_authority_unchanged(bundle_workspace, removed_authority_before, "removed-bundle")

        recording_root = root / "unvalidated-recording"
        recording_root.mkdir()
        recording_workspace = new_workspace(recording_root)
        (recording_workspace / "recording-evidence.json").write_text(
            json.dumps({"recording_status": "passed"}, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        recording_state = render(recording_workspace).get("state", {})
        if recording_state.get("recording", {}).get("status") != "manifest_passed_not_revalidated" or recording_state.get("counts", {}).get("recording_manifests") != 1:
            raise SystemExit("FAILED: unvalidated recording manifest was promoted to an authoritative pass")

        invalid_validator_root = root / "invalid-validator"
        invalid_validator_root.mkdir()
        invalid_validator_workspace = new_workspace(invalid_validator_root)
        (invalid_validator_workspace / "recon-result.json").write_text("{}\n", encoding="utf-8")
        (invalid_validator_workspace / "bin/validate-recon-result.py").mkdir(parents=True)
        invalid_validator_state = render(invalid_validator_workspace).get("state", {})
        invalid_codes = {item.get("code") for item in invalid_validator_state.get("integrity", {}).get("issues", []) if isinstance(item, dict)}
        if "STRUCTURED_RESULT_INVALID" not in invalid_codes:
            raise SystemExit("FAILED: invalid validator JSON/executable did not fail closed")

        checkpoint_output = json.loads(run_capture([
            sys.executable, str(create_cli), "--workspace-dir", str(workspace), "--repo-root", str(root), "--json",
        ], plugin_root))
        if checkpoint_output.get("idempotent") is not False:
            raise SystemExit("FAILED: first checkpoint creation was not a publish")
        checkpoint_path = str(checkpoint_output["path"])
        second_checkpoint = json.loads(run_capture([
            sys.executable, str(create_cli), "--workspace-dir", str(workspace), "--repo-root", str(root), "--json",
        ], plugin_root))
        if second_checkpoint.get("idempotent") is not True or second_checkpoint.get("sha256") != checkpoint_output.get("sha256"):
            raise SystemExit("FAILED: same-revision checkpoint was not byte-idempotent")
        if validate(workspace, checkpoint=checkpoint_path).get("classification") != "current":
            raise SystemExit("FAILED: current checkpoint did not validate")

        # Docker evidence and notes remain non-confirming material.
        (workspace / "evidence/docker").mkdir(parents=True)
        (workspace / "evidence/docker/verification-evidence.json").write_text("{}\n", encoding="utf-8")
        docker_state = render(workspace).get("state", {})
        if docker_state.get("counts", {}).get("validated_confirmed_bundles") != 0 or docker_state.get("counts", {}).get("docker_evidence_only") < 1:
            raise SystemExit("FAILED: Docker evidence-only state was not represented conservatively")

        # Forged derived counts and unsafe checkpoint values fail read-only validation.
        handoff_path = workspace / "handoff-state.json"
        handoff_bytes = handoff_path.read_bytes()
        forged = json.loads(handoff_bytes.decode("utf-8"))
        forged["counts"]["validated_confirmed_bundles"] = 99
        handoff_path.write_text(json.dumps(forged, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
        validate(workspace, expected=1)
        handoff_path.write_bytes(handoff_bytes)
        render(workspace)

        original_checkpoint = (workspace / checkpoint_path).read_bytes()
        forged_checkpoint = json.loads(original_checkpoint.decode("utf-8"))
        for unsafe_value in ("/tmp/escape", "../escape", "https://example.invalid", "a\\b", "~/.ssh/id_rsa"):
            forged_checkpoint["resume"]["parameters"] = [{"name": "artifact", "value": unsafe_value}]
            (workspace / checkpoint_path).write_text(json.dumps(forged_checkpoint, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
            validate(workspace, expected=1, checkpoint=checkpoint_path)
        (workspace / checkpoint_path).write_bytes(original_checkpoint)

        # Same-revision conflicting bytes never overwrite the original checkpoint.
        (workspace / checkpoint_path).write_bytes(original_checkpoint + b"tampered\n")
        run_expect_fail([
            sys.executable, str(create_cli), "--workspace-dir", str(workspace), "--repo-root", str(root), "--json",
        ], plugin_root, "CHECKPOINT_CONFLICTING_BYTES")
        (workspace / checkpoint_path).write_bytes(original_checkpoint)

        # Event append makes the old handoff stale; after regeneration the old
        # checkpoint is a legal historical snapshot when non-volatile inputs agree.
        run([
            sys.executable, str(event_writer), "--workspace-dir", str(workspace), "--plugin-version", "selftest",
            "--event", "handoff_observe", "--stage", "current", "--status", "current",
            "--transition-kind", "observe", "--event-status", "ok", "--message", "next", "--accept-current-revision",
        ], plugin_root)
        validate(workspace, expected=1)
        render(workspace)
        historical = validate(workspace, checkpoint=checkpoint_path)
        if historical.get("classification") != "valid_historical":
            raise SystemExit(f"FAILED: legal historical checkpoint was not distinguished: {historical}")
        historical_bytes = (workspace / checkpoint_path).read_bytes()
        historical_doc = json.loads(historical_bytes.decode("utf-8"))
        historical_doc["event_digest"] = "sha256:" + "0" * 64
        (workspace / checkpoint_path).write_text(json.dumps(historical_doc, ensure_ascii=False, sort_keys=True, indent=2) + "\n", encoding="utf-8")
        validate(workspace, expected=1, checkpoint=checkpoint_path)
        (workspace / checkpoint_path).write_bytes(historical_bytes)
        leading_zero_checkpoint = workspace / "checkpoints/01.json"
        leading_zero_checkpoint.write_bytes(historical_bytes)
        validate(workspace, expected=1, checkpoint="checkpoints/01.json")
        leading_zero_checkpoint.unlink()

        # Symlink input and atomic fault injection must fail without replacing
        # the previous handoff. The selftest does not touch Docker or network.
        candidate_link = workspace / "candidate.json"
        candidate_link.symlink_to(workspace / "stage-status.json")
        render(workspace, expected=1)
        candidate_link.unlink()
        before_fault = handoff_path.read_bytes()
        render(workspace, expected=1, env={"ZHULONG_TEST_FAIL_HANDOFF_WRITE": "1"})
        if handoff_path.read_bytes() != before_fault:
            raise SystemExit("FAILED: handoff temp-write fault changed the published bytes")
        render(workspace, expected=1, env={"ZHULONG_TEST_FAIL_HANDOFF_REPLACE": "1"})
        if handoff_path.read_bytes() != before_fault:
            raise SystemExit("FAILED: handoff replace fault changed the published bytes")

        symlink_output = root / "symlink-output"
        symlink_output.mkdir()
        symlink_workspace = new_workspace(symlink_output)
        render(symlink_workspace)
        outside_handoff = root / "outside-handoff.json"
        outside_handoff.write_text("outside\n", encoding="utf-8")
        (symlink_workspace / "handoff-state.json").unlink()
        (symlink_workspace / "handoff-state.json").symlink_to(outside_handoff)
        render(symlink_workspace, expected=1)
        (symlink_workspace / "handoff-state.json").unlink()
        (symlink_workspace / "handoff-state.json").write_bytes(before_fault)

        checkpoint_fault_root = root / "checkpoint-fault"
        checkpoint_fault_root.mkdir()
        checkpoint_fault_workspace = new_workspace(checkpoint_fault_root)
        render(checkpoint_fault_workspace)
        checkpoint_fault_command = [
            sys.executable, str(create_cli), "--workspace-dir", str(checkpoint_fault_workspace),
            "--repo-root", str(checkpoint_fault_root), "--json",
        ]
        run_capture_with_env(checkpoint_fault_command, plugin_root, {"ZHULONG_TEST_FAIL_CHECKPOINT_REPLACE": "1"}, expected_returncode=1)
        if (checkpoint_fault_workspace / "checkpoints").exists():
            published = list((checkpoint_fault_workspace / "checkpoints").glob("*.json"))
            if published:
                raise SystemExit("FAILED: checkpoint replace fault left a published checkpoint")

        checkpoint_parent_root = root / "checkpoint-parent-symlink"
        checkpoint_parent_root.mkdir()
        checkpoint_parent_workspace = new_workspace(checkpoint_parent_root)
        render(checkpoint_parent_workspace)
        outside_checkpoints = root / "outside-checkpoints"
        outside_checkpoints.mkdir()
        (checkpoint_parent_workspace / "checkpoints").symlink_to(outside_checkpoints, target_is_directory=True)
        run_capture_with_env([
            sys.executable, str(create_cli), "--workspace-dir", str(checkpoint_parent_workspace),
            "--repo-root", str(checkpoint_parent_root), "--json",
        ], plugin_root, {}, expected_returncode=1)
        if list(outside_checkpoints.iterdir()):
            raise SystemExit("FAILED: checkpoint parent symlink received a published file")

        concurrent_root = root / "concurrent-append"
        concurrent_root.mkdir()
        concurrent_workspace = new_workspace(concurrent_root)
        render(concurrent_workspace)
        concurrent_journal = concurrent_workspace / "audit-events.jsonl"
        concurrent_handoff = concurrent_workspace / "handoff-state.json"
        journal_before = concurrent_journal.read_bytes()
        handoff_before = concurrent_handoff.read_bytes()
        pause_marker = root / "handoff-paused.marker"
        concurrent_proc = subprocess.Popen(
            [sys.executable, str(render_cli), "--workspace-dir", str(concurrent_workspace), "--repo-root", str(concurrent_root), "--json"],
            cwd=plugin_root,
            env={**os.environ, "ZHULONG_TEST_HANDOFF_PAUSE": "1", "ZHULONG_TEST_HANDOFF_PAUSE_MARKER": str(pause_marker)},
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        for _ in range(100):
            if pause_marker.exists():
                break
            time.sleep(0.01)
        else:
            concurrent_proc.kill()
            concurrent_proc.communicate()
            raise SystemExit("FAILED: handoff concurrency selftest did not reach its pause seam")
        with concurrent_journal.open("ab") as stream:
            stream.write(journal_before.splitlines()[0] + b"\n")
            stream.flush()
            os.fsync(stream.fileno())
        concurrent_stdout, concurrent_stderr = concurrent_proc.communicate(timeout=5)
        concurrent_output = (concurrent_stdout or "") + (concurrent_stderr or "")
        if concurrent_proc.returncode == 0 or "CONCURRENT_STATE_CHANGED" not in concurrent_output:
            raise SystemExit(f"FAILED: concurrent journal append was not rejected: {concurrent_output}")
        concurrent_journal.write_bytes(journal_before)
        if concurrent_handoff.read_bytes() != handoff_before:
            raise SystemExit("FAILED: concurrent journal append changed the previous handoff bytes")

        malformed_root = root / "malformed"
        malformed_root.mkdir()
        malformed_workspace = new_workspace(malformed_root)
        with (malformed_workspace / "audit-events.jsonl").open("ab") as stream:
            stream.write(b"{malformed\n")
        render(malformed_workspace, expected=1)

    print("HANDOFF/CHECKPOINT SELFTEST PASSED: derived-only, conservative, atomic, historical, symlink-safe")


def exercise_next_actions_contract(plugin_root: Path) -> None:
    """Run production next-action CLIs on a sanitized, authority-valid workspace."""
    fixture = plugin_root / "assets/fixtures/next-actions/manifest.json"
    manifest = json.loads(fixture.read_text(encoding="utf-8"))
    required_positive = {"handoff-stale-only", "candidate-verdict-missing", "docker-oracle-unproven", "no-action", "stable-multiple-actions"}
    required_negative = {"zero-bundle-variant-forgery", "markdown-note-injection", "unsafe-entrypoint", "unsafe-path", "symlink-input-output-parent", "schema-only-validator", "r1-masquerading-as-r2"}
    if not required_positive.issubset(set(manifest.get("positive_cases", []))) or not required_negative.issubset(set(manifest.get("negative_cases", []))):
        raise SystemExit("FAILED: next-actions fixture manifest is incomplete")
    require_files(plugin_root, [
        "assets/schemas/next-actions.schema.json", "assets/fixtures/next-actions/manifest.json",
        "docs/runner-contracts/next-actions-contract-r1.md", "scripts/next_actions.py",
        "scripts/render_next_actions.py", "scripts/validate_next_actions.py",
    ], "next-actions contract")
    with tempfile.TemporaryDirectory(prefix="zhulong-next-actions-") as tempdir:
        root = Path(tempdir)
        source = plugin_root / "assets/fixtures/recon-result/service"
        repo, workspace = root / "repo", root / "workspace"
        shutil.copytree(source / "repo", repo)
        shutil.copytree(source / "workspace", workspace)
        writer = plugin_root / "scripts/write_audit_event.py"
        run([sys.executable, str(writer), "--workspace-dir", str(workspace), "--plugin-version", "selftest", "--event", "next_actions_fixture", "--stage", "intake", "--status", "running", "--transition-kind", "start", "--event-status", "ok", "--message", "sanitized fixture", "--accept-current-revision"], plugin_root)
        handoff = plugin_root / "scripts/render_handoff_state.py"
        render = plugin_root / "scripts/render_next_actions.py"
        validate = plugin_root / "scripts/validate_next_actions.py"
        run([sys.executable, str(handoff), "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json"], plugin_root)
        authority_before = {(workspace / name): (workspace / name).read_bytes() for name in ("audit-events.jsonl", "stage-status.json", "handoff-state.json")}
        first = run_capture([sys.executable, str(render), "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json"], plugin_root)
        published = (workspace / "next-actions.json").read_bytes()
        second = run_capture([sys.executable, str(render), "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json"], plugin_root)
        if published != (workspace / "next-actions.json").read_bytes() or json.loads(first).get("state") != json.loads(second).get("state"):
            raise SystemExit("FAILED: next-actions generation is not deterministic")
        if any(path.read_bytes() != raw for path, raw in authority_before.items()):
            raise SystemExit("FAILED: next-actions generator changed authority inputs")
        validator_before = (workspace / "next-actions.json").read_bytes()
        run([sys.executable, str(validate), "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json"], plugin_root)
        if (workspace / "next-actions.json").read_bytes() != validator_before:
            raise SystemExit("FAILED: next-actions validator wrote its input")
        run([sys.executable, str(writer), "--workspace-dir", str(workspace), "--plugin-version", "selftest", "--event", "next_actions_stale", "--stage", "intake", "--status", "running", "--transition-kind", "observe", "--event-status", "ok", "--message", "sanitized stale fixture", "--accept-current-revision"], plugin_root)
        stale = json.loads(run_capture([sys.executable, str(render), "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json"], plugin_root))["state"]
        if stale.get("classification") != "action_required" or [item.get("blocking_code") for item in stale.get("actions", [])] != ["HANDOFF_STALE"]:
            raise SystemExit("FAILED: stale handoff did not produce only HANDOFF_STALE")
    print("NEXT-ACTIONS SELFTEST PASSED: derived-only, deterministic, stale-safe, read-only validation")


def exercise_audit_state_writer(plugin_root: Path) -> None:
    """Exercise P9.2 commits with real subprocesses and internal-only faults."""
    writer = plugin_root / "scripts/write_audit_event.py"
    fixture_root = plugin_root / "assets/fixtures/audit-state-protocol-r2"

    def invoke(
        workspace: Path,
        name: str,
        *,
        status: str = "running",
        stage: str = "intake",
        transition_kind: str = "observe",
        revision: int | None = None,
        current: bool = True,
        extra: list[str] | None = None,
        include_enhanced_material: bool = True,
    ) -> tuple[int, dict[str, object], str]:
        command = [
            sys.executable, str(writer), "--workspace-dir", str(workspace),
            "--event", name, "--stage", stage, "--status", status,
            "--transition-kind", transition_kind,
            "--message", f"Selftest event {name}.", "--json",
        ]
        if revision is not None:
            command.extend(["--expected-state-revision", str(revision)])
        elif current:
            command.append("--accept-current-revision")
        if status in {"blocked", "paused"}:
            command.extend(["--blocker", "selftest blocker", "--resume-step", "resume selftest"])
        if transition_kind in {"resume", "skip", "return", "reopen"} and include_enhanced_material:
            reason_code = {
                "resume": "recovery_requested",
                "skip": "not_applicable",
                "return": "validation_failed",
                "reopen": "validation_failed",
            }[transition_kind]
            command.extend([
                "--reason-code", reason_code,
                "--subject", "run:selftest",
                "--evidence-ref", "evidence/selftest/transition.json",
                "--next-action-json",
                json.dumps(
                    {
                        "action_id": "continue-selftest-transition",
                        "action_type": "review",
                        "subject_ids": ["run:selftest"],
                        "summary": "Continue the bounded selftest transition.",
                        "evidence_refs": ["evidence/selftest/transition.json"],
                    },
                    sort_keys=True,
                ),
                "--details-json",
                json.dumps(
                    {
                        "summary": f"Selftest event {name}.",
                        "reason_detail": "The selftest supplies the required transition evidence.",
                    },
                    sort_keys=True,
                ),
            ])
        if extra:
            command.extend(extra)
        proc = subprocess.run(command, cwd=plugin_root, capture_output=True, text=True)
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        try:
            payload = json.loads((proc.stdout or "").strip())
        except json.JSONDecodeError:
            payload = {}
        return proc.returncode, payload, output

    with tempfile.TemporaryDirectory(prefix="zhulong-audit-writer-") as temp_dir:
        root = Path(temp_dir)
        workspace = root / "r2"
        workspace.mkdir()
        code, first, output = invoke(workspace, "first_event", transition_kind="start")
        if code != 0 or first.get("mode") != "r2" or first.get("seq") != 1 or first.get("state_revision") != 1:
            raise SystemExit("FAILED: first R2 writer commit did not create seq=1/revision=1: " + output)
        code, blocked, output = invoke(workspace, "paused_event", status="paused", transition_kind="pause")
        if code != 0 or blocked.get("state_revision") != 2:
            raise SystemExit("FAILED: blocked R2 writer commit failed: " + output)
        paused_state = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
        if not paused_state.get("blocker") or not paused_state.get("resume_step"):
            raise SystemExit("FAILED: blocked R2 event did not materialize blocker/resume_step")
        code, running, output = invoke(workspace, "running_event", transition_kind="resume")
        running_state = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
        if code != 0 or running_state.get("blocker") is not None or running_state.get("resume_step") is not None:
            raise SystemExit("FAILED: running R2 event did not clear blocker/resume_step: " + output)
        digest = "sha256:" + hashlib.sha256((workspace / "audit-events.jsonl").read_bytes()).hexdigest()
        if running_state.get("event_log_digest") != digest:
            raise SystemExit("FAILED: R2 state digest does not match exact committed journal bytes")
        stale = root / "stale"
        shutil.copytree(workspace, stale)
        stale_journal_before = (stale / "audit-events.jsonl").read_bytes()
        (stale / "stage-status.json").unlink()
        code, payload, output = invoke(stale, "stale_view", transition_kind="observe")
        if code == 0 or payload.get("code") != "STATE_VIEW_MISSING" or (stale / "audit-events.jsonl").read_bytes() != stale_journal_before:
            raise SystemExit("FAILED: missing committed R2 state view did not fail closed: " + output)

        concurrent = root / "concurrent"
        concurrent.mkdir()
        if invoke(concurrent, "concurrent_start", transition_kind="start")[0] != 0:
            raise SystemExit("FAILED: could not seed concurrent R2 workspace")
        processes = [
            subprocess.Popen([
                sys.executable, str(writer), "--workspace-dir", str(concurrent),
                "--event", f"concurrent_{index}", "--stage", "intake", "--status", "running",
                "--transition-kind", "observe",
                "--message", f"Concurrent writer {index}.", "--accept-current-revision", "--json",
            ], cwd=plugin_root, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            for index in range(7)
        ]
        for process in processes:
            stdout, stderr = process.communicate(timeout=15)
            if process.returncode != 0:
                raise SystemExit("FAILED: concurrent current-revision writer failed: " + (stdout + stderr))
        concurrent_events = [json.loads(line) for line in (concurrent / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()]
        concurrent_state = json.loads((concurrent / "stage-status.json").read_text(encoding="utf-8"))
        if [event.get("seq") for event in concurrent_events] != list(range(1, 9)):
            raise SystemExit("FAILED: concurrent R2 writers did not produce contiguous unique seq values")
        if concurrent_state.get("state_revision") != 8:
            raise SystemExit("FAILED: concurrent R2 writers did not produce final revision 8")

        cas = root / "cas"
        cas.mkdir()
        if invoke(cas, "seed", transition_kind="start")[0] != 0:
            raise SystemExit("FAILED: could not seed CAS workspace")
        cas_commands = [
            [sys.executable, str(writer), "--workspace-dir", str(cas), "--event", f"cas_{index}",
             "--stage", "intake", "--status", "running", "--message", "CAS selftest.",
             "--transition-kind", "observe", "--expected-state-revision", "1", "--json"]
            for index in range(2)
        ]
        cas_processes = [subprocess.Popen(command, cwd=plugin_root, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True) for command in cas_commands]
        cas_results = [process.communicate(timeout=15) + (process.returncode,) for process in cas_processes]
        successes = sum(1 for stdout, stderr, code in cas_results if code == 0)
        conflicts = sum(1 for stdout, stderr, code in cas_results if code != 0 and "STATE_REVISION_CONFLICT" in stdout + stderr)
        if successes != 1 or conflicts != 1:
            raise SystemExit("FAILED: equal-revision CAS writers did not yield one success and one conflict")
        if (
            len((cas / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()) != 2
            or json.loads((cas / "stage-status.json").read_text(encoding="utf-8")).get("state_revision") != 2
        ):
            raise SystemExit("FAILED: rejected equal-revision CAS writer changed journal or state")

        held = root / "held-lock"
        held.mkdir()
        if invoke(held, "seed", transition_kind="start")[0] != 0:
            raise SystemExit("FAILED: could not seed lock timeout workspace")
        lock_holder = subprocess.Popen([
            sys.executable, "-c",
            "import fcntl,pathlib,time; p=pathlib.Path(__import__('sys').argv[1]) / '.audit-state.lock'; "
            "f=p.open('a+'); fcntl.flock(f, fcntl.LOCK_EX); time.sleep(1.0)", str(held),
        ])
        time.sleep(0.15)
        journal_before = (held / "audit-events.jsonl").read_bytes()
        code, payload, output = invoke(
            held,
            "timeout",
            transition_kind="observe",
            extra=["--lock-timeout-seconds", "0.05"],
        )
        lock_holder.wait(timeout=5)
        if code == 0 or payload.get("code") != "LOCK_TIMEOUT" or (held / "audit-events.jsonl").read_bytes() != journal_before:
            raise SystemExit("FAILED: held lock did not cause bounded zero-write LOCK_TIMEOUT: " + output)
        if not (held / ".audit-state.lock").is_file() or invoke(held, "after_lock", transition_kind="observe")[0] != 0:
            raise SystemExit("FAILED: persistent lock file was not reusable after release")

        legacy = root / "legacy"
        legacy.mkdir()
        legacy_event = json.loads((fixture_root / "legacy-event-r1.json").read_text(encoding="utf-8"))
        (legacy / "audit-events.jsonl").write_text(
            json.dumps(legacy_event, ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        shutil.copy2(fixture_root / "legacy-state-r1.json", legacy / "stage-status.json")
        code, payload, output = invoke(
            legacy,
            "legacy_append",
            extra=["--protocol-mode", "legacy-r1"],
        )
        if code != 0 or payload.get("mode") != "legacy_r1" or payload.get("cas_mode") != "unavailable":
            raise SystemExit("FAILED: existing R1 workspace was not preserved as legacy R1: " + output)
        if '"schema_version":2' in (legacy / "audit-events.jsonl").read_text(encoding="utf-8"):
            raise SystemExit("FAILED: legacy R1 journal was mixed with an R2 event")

        legacy_sensitive = root / "legacy-sensitive"
        shutil.copytree(legacy, legacy_sensitive)
        legacy_journal_before = (legacy_sensitive / "audit-events.jsonl").read_bytes()
        legacy_state_before = (legacy_sensitive / "stage-status.json").read_bytes()
        code, payload, output = invoke(
            legacy_sensitive,
            "legacy_sensitive",
            extra=["--protocol-mode", "legacy-r1", "--message", "Review /etc/passwd only."],
        )
        if code == 0 or payload.get("code") != "EVENT_SENSITIVE_TEXT_FORBIDDEN":
            raise SystemExit("FAILED: sensitive R1 message was not rejected: " + output)
        if (legacy_sensitive / "audit-events.jsonl").read_bytes() != legacy_journal_before or (legacy_sensitive / "stage-status.json").read_bytes() != legacy_state_before:
            raise SystemExit("FAILED: rejected sensitive R1 message changed journal/state bytes")

        unsafe = root / "unsafe"
        unsafe.mkdir()
        (unsafe / "target.jsonl").write_text("\n", encoding="utf-8")
        (unsafe / "audit-events.jsonl").symlink_to(unsafe / "target.jsonl")
        code, payload, output = invoke(unsafe, "unsafe_path", transition_kind="start")
        if code == 0 or payload.get("code") != "JOURNAL_PATH_UNSAFE":
            raise SystemExit("FAILED: symlink journal path was not rejected: " + output)
        unsafe_lock = root / "unsafe-lock"
        unsafe_lock.mkdir()
        (unsafe_lock / "lock-target").write_text("x", encoding="utf-8")
        (unsafe_lock / ".audit-state.lock").symlink_to(unsafe_lock / "lock-target")
        code, payload, output = invoke(unsafe_lock, "unsafe_lock", transition_kind="start")
        if code == 0 or payload.get("code") != "LOCK_PATH_UNSAFE":
            raise SystemExit("FAILED: symlink lock path was not rejected: " + output)

        invalid = root / "invalid"
        invalid.mkdir()
        code, payload, output = invoke(
            invalid,
            "bad_evidence",
            transition_kind="start",
            extra=["--evidence-ref", "https://example.invalid/evidence"],
        )
        if code == 0 or payload.get("code") != "EVENT_VALIDATION_FAILED" or (invalid / "audit-events.jsonl").exists():
            raise SystemExit("FAILED: invalid evidence ref did not fail before journal/state writes: " + output)
        code, payload, output = invoke(
            invalid,
            "bad_details",
            transition_kind="start",
            extra=["--details-json", '{"nested":{"value":true}}'],
        )
        if code == 0 or payload.get("code") != "EVENT_VALIDATION_FAILED" or (invalid / "audit-events.jsonl").exists():
            raise SystemExit("FAILED: nested details did not fail before journal/state writes: " + output)

        sys.path.insert(0, str(plugin_root / "scripts"))
        try:
            import audit_state_io as state_io
        finally:
            sys.path.pop(0)
        fault = root / "fault"
        fault.mkdir()
        request = {
            "accept_current_revision": True, "expected_state_revision": None, "run_id": "",
            "timestamp": "2026-07-18T00:00:00Z", "stage": "intake", "to_status": "running",
            "event_type": "checkpoint", "event_name": "fault_event", "reason_code": "normal_progress",
            "subjects": [], "evidence_refs": [], "next_actions": [],
            "details": {"summary": "Fault injection selftest."}, "blocker": "", "resume_step": "",
            "transition_kind": "start", "expected_from_stage": "", "expected_from_status": "",
            "plugin_version": "selftest",
            "legacy_event": {}, "legacy_state": {},
        }
        with mock.patch.object(state_io, "_atomic_replace_state", side_effect=state_io.AuditStateError("STATE_VIEW_REPLACE_FAILED", "injected")):
            try:
                state_io.commit_event(fault, mode_policy="r2", lock_timeout_seconds=1, request=request)
            except state_io.AuditStateError as exc:
                if exc.code != "STATE_VIEW_REPLACE_FAILED" or not exc.fields.get("journal_committed"):
                    raise SystemExit("FAILED: state replacement fault did not retain committed-journal semantics")
            else:
                raise SystemExit("FAILED: injected state replacement fault unexpectedly succeeded")
        if len((fault / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()) != 1 or (fault / "stage-status.json").exists():
            raise SystemExit("FAILED: state replacement fault did not retain exactly one journal event and no partial state")
        try:
            state_io.commit_event(fault, mode_policy="r2", lock_timeout_seconds=1, request=request)
        except state_io.AuditStateError as exc:
            if exc.code != "STATE_VIEW_MISSING":
                raise SystemExit("FAILED: post-commit stale state did not fail closed")
        else:
            raise SystemExit("FAILED: post-commit stale state allowed a duplicate event")

        temp_write_fault = root / "temp-write-fault"
        temp_write_fault.mkdir()
        with mock.patch.object(state_io, "_write_state_temp", side_effect=state_io.AuditStateError("STATE_VIEW_WRITE_FAILED", "injected")):
            try:
                state_io.commit_event(temp_write_fault, mode_policy="r2", lock_timeout_seconds=1, request=request)
            except state_io.AuditStateError as exc:
                if exc.code != "STATE_VIEW_WRITE_FAILED" or not exc.fields.get("journal_committed"):
                    raise SystemExit("FAILED: state temporary-write fault did not retain committed-journal semantics")
            else:
                raise SystemExit("FAILED: injected state temporary-write fault unexpectedly succeeded")
        if len((temp_write_fault / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()) != 1 or (temp_write_fault / "stage-status.json").exists():
            raise SystemExit("FAILED: state temporary-write fault left an invalid state view")

        append_fault = root / "append-fault"
        append_fault.mkdir()
        with mock.patch.object(state_io, "_safe_append_fsync", side_effect=state_io.AuditStateError("JOURNAL_APPEND_FAILED", "injected")):
            try:
                state_io.commit_event(append_fault, mode_policy="r2", lock_timeout_seconds=1, request=request)
            except state_io.AuditStateError as exc:
                if exc.code != "JOURNAL_APPEND_FAILED":
                    raise SystemExit("FAILED: injected append error changed failure semantics")
            else:
                raise SystemExit("FAILED: injected journal append failure unexpectedly succeeded")
        if (append_fault / "audit-events.jsonl").exists() or (append_fault / "stage-status.json").exists():
            raise SystemExit("FAILED: injected journal append failure changed state or journal bytes")

    print("AUDIT STATE WRITER SELFTEST PASSED: lock/CAS/durability/R1/path-safety")


def exercise_audit_state_recovery(plugin_root: Path) -> None:
    recovery = plugin_root / "scripts/recover_audit_state.py"
    writer = plugin_root / "scripts/write_audit_event.py"

    def call(workspace: Path, *args: str) -> tuple[int, dict[str, Any]]:
        proc = subprocess.run(
            [sys.executable, str(recovery), "--workspace-dir", str(workspace), *args, "--json"],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        raw = (proc.stdout or proc.stderr).strip()
        try:
            payload = json.loads(raw)
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: recovery CLI did not emit JSON: {raw}") from exc
        return proc.returncode, payload

    with tempfile.TemporaryDirectory(prefix="zhulong-state-recovery-") as tempdir:
        root = Path(tempdir)
        current = root / "current-r2"
        current.mkdir()
        run([
            sys.executable,
            str(writer),
            "--workspace-dir", str(current),
            "--protocol-mode", "r2",
            "--accept-current-revision",
            "--plugin-version", "0.4.0-recovery-fixture",
            "--event", "intake_started",
            "--stage", "intake",
            "--status", "running",
            "--transition-kind", "start",
            "--message", "Recovery fixture started.",
            "--json",
        ], plugin_root)
        journal_before = (current / "audit-events.jsonl").read_bytes()
        state_before = (current / "stage-status.json").read_bytes()
        code, checked = call(current, "--check")
        if code != 0 or checked.get("rebuildability") != "complete_from_journal" or checked.get("drift"):
            raise SystemExit("FAILED: valid new R2 recovery check was not an exact journal-only match")
        journal_digest = str(checked["journal"]["digest"])
        state_digest = str(checked["state"]["digest"])

        (current / "stage-status.json").unlink()
        code, missing = call(current, "--check")
        if code == 0 or missing.get("rebuildability") != "complete_from_journal":
            raise SystemExit("FAILED: missing state was not reported as rebuildable without a write")
        if (current / "stage-status.json").exists() or (current / "audit-events.jsonl").read_bytes() != journal_before:
            raise SystemExit("FAILED: recovery dry-run modified journal or state")
        code, applied = call(
            current,
            "--apply",
            "--expected-journal-digest", journal_digest,
            "--expect-state-missing",
        )
        if code != 0 or not applied.get("applied") or (current / "stage-status.json").read_bytes() != state_before:
            raise SystemExit("FAILED: missing state was not rebuilt byte-exactly")

        tampered = json.loads(state_before)
        tampered["stage"] = "recon"
        tampered_bytes = (json.dumps(tampered, ensure_ascii=False, sort_keys=True, indent=2) + "\n").encode("utf-8")
        (current / "stage-status.json").write_bytes(tampered_bytes)
        code, stale = call(current, "--check")
        if code == 0 or not any(item.get("code") == "STATE_STAGE_MISMATCH" for item in stale.get("drift", [])):
            raise SystemExit("FAILED: field-level state drift was not diagnosed")
        code, conflict = call(
            current,
            "--apply",
            "--expected-journal-digest", journal_digest,
            "--expected-state-digest", state_digest,
        )
        if code == 0 or conflict.get("code") != "STATE_DIGEST_CONFLICT":
            raise SystemExit("FAILED: state digest CAS conflict was not enforced")
        if (current / "stage-status.json").read_bytes() != tampered_bytes or (current / "audit-events.jsonl").read_bytes() != journal_before:
            raise SystemExit("FAILED: failed CAS apply modified journal or state")
        tampered_digest = "sha256:" + hashlib.sha256(tampered_bytes).hexdigest()
        code, replaced = call(
            current,
            "--apply",
            "--expected-journal-digest", journal_digest,
            "--expected-state-digest", tampered_digest,
        )
        if code != 0 or not replaced.get("applied") or (current / "stage-status.json").read_bytes() != state_before:
            raise SystemExit("FAILED: matching state digest did not authorize exact replacement")

        historical = root / "anchored-historical-r2"
        historical.mkdir()
        old_journal = (plugin_root / "assets/fixtures/audit-state-protocol-r2/valid-events-r2.jsonl").read_bytes()
        old_events = [json.loads(line) for line in old_journal.decode("utf-8").splitlines() if line.strip()]
        latest = old_events[-1]
        old_state = {
            "schema_version": 2,
            "plugin": "zhulong",
            "plugin_version": "0.4.0-anchored",
            "run_id": latest["run_id"],
            "state_revision": len(old_events),
            "last_event_seq": latest["seq"],
            "event_log_digest": "sha256:" + hashlib.sha256(old_journal).hexdigest(),
            "stage": latest["stage"],
            "status": latest["to_status"],
            "last_event_at": latest["ts"],
            "last_event_type": latest["event_type"],
            "last_event_name": latest["event_name"],
            "blocker": None,
            "resume_step": None,
        }
        (historical / "audit-events.jsonl").write_bytes(old_journal)
        historical_state_bytes = (json.dumps(old_state, ensure_ascii=False, sort_keys=True, indent=2) + "\n").encode("utf-8")
        (historical / "stage-status.json").write_bytes(historical_state_bytes)
        code, anchored = call(historical, "--check")
        if code != 0 or anchored.get("rebuildability") != "complete_with_anchored_legacy_metadata":
            raise SystemExit("FAILED: exact historical state anchor was not accepted conservatively")
        first_line = old_journal.splitlines(keepends=True)[0]
        first = old_events[0]
        stale_anchor = {
            **old_state,
            "run_id": first["run_id"],
            "state_revision": 1,
            "last_event_seq": 1,
            "event_log_digest": "sha256:" + hashlib.sha256(first_line).hexdigest(),
            "stage": first["stage"],
            "status": first["to_status"],
            "last_event_at": first["ts"],
            "last_event_type": first["event_type"],
            "last_event_name": first["event_name"],
        }
        (historical / "stage-status.json").write_text(
            json.dumps(stale_anchor, ensure_ascii=False, sort_keys=True, indent=2) + "\n",
            encoding="utf-8",
        )
        code, stale_anchored = call(historical, "--check")
        if code == 0 or stale_anchored.get("rebuildability") != "complete_with_anchored_legacy_metadata":
            raise SystemExit("FAILED: stale state did not supply metadata only through its exact journal-prefix anchor")
        (historical / "stage-status.json").unlink()
        code, unavailable = call(historical, "--check")
        if code == 0 or unavailable.get("rebuildability") != "blocked_missing_metadata":
            raise SystemExit("FAILED: missing historical plugin-version provenance did not fail closed")

        damaged = root / "damaged"
        damaged.mkdir()
        (damaged / "audit-events.jsonl").write_bytes(journal_before[:-1])
        code, no_newline = call(damaged, "--check")
        if code == 0 or no_newline["journal"]["classification"] != "journal_final_newline_missing":
            raise SystemExit("FAILED: valid final event without newline was not distinguished")
        (damaged / "audit-events.jsonl").write_bytes(journal_before + b'{"schema_version":2')
        code, tail = call(damaged, "--check")
        if code == 0 or tail["journal"]["classification"] != "journal_tail_incomplete":
            raise SystemExit("FAILED: incomplete final tail was not distinguished")
        (damaged / "audit-events.jsonl").write_bytes(journal_before + b"{bad}\n" + journal_before)
        code, middle = call(damaged, "--check")
        if code == 0 or middle["journal"]["classification"] != "journal_middle_corruption":
            raise SystemExit("FAILED: middle corruption was not distinguished")

        legacy = root / "legacy-r1"
        legacy.mkdir()
        legacy_event = json.loads((plugin_root / "assets/fixtures/audit-state-protocol-r2/legacy-event-r1.json").read_text(encoding="utf-8"))
        legacy_state = (plugin_root / "assets/fixtures/audit-state-protocol-r2/legacy-state-r1.json").read_bytes()
        (legacy / "audit-events.jsonl").write_text(json.dumps(legacy_event, sort_keys=True) + "\n", encoding="utf-8")
        (legacy / "stage-status.json").write_bytes(legacy_state)
        legacy_before = {path.name: path.read_bytes() for path in legacy.iterdir()}
        code, preflight = call(legacy, "--migration-preflight")
        if code != 0 or not preflight["r1_migration_preflight"].get("available"):
            raise SystemExit("FAILED: valid R1 migration preflight was unavailable")
        if legacy_before != {path.name: path.read_bytes() for path in legacy.iterdir()}:
            raise SystemExit("FAILED: R1 migration preflight wrote source files")

    print("AUDIT STATE RECOVERY SELFTEST PASSED: bytes/diagnostics/provenance/CAS/R1")


def exercise_audit_transition_policy(plugin_root: Path) -> None:
    """Exercise P9.3 policy intent at the real writer and validator boundary."""
    writer = plugin_root / "scripts/write_audit_event.py"
    validator = plugin_root / "scripts/validate_audit_protocol.py"
    default = object()

    def snapshot(workspace: Path) -> dict[str, bytes | None]:
        return {
            name: (workspace / name).read_bytes() if (workspace / name).exists() else None
            for name in ("audit-events.jsonl", "stage-status.json", ".audit-state.lock")
        }

    def transition_action(name: str) -> dict[str, object]:
        evidence_ref = "evidence/selftest/transition-context.json"
        return {
            "action_id": f"continue-{name}",
            "action_type": "review",
            "subject_ids": ["run:transition-selftest"],
            "summary": "Continue the bounded transition-policy selftest.",
            "evidence_refs": [evidence_ref],
        }

    def invoke(
        workspace: Path,
        name: str,
        *,
        stage: str,
        status: str,
        transition_kind: str,
        reason_code: str | object = default,
        reason_detail: str | None | object = default,
        subjects: list[str] | object = default,
        evidence_refs: list[str] | object = default,
        next_actions: list[dict[str, object]] | object = default,
        blocker: str = "",
        resume_step: str = "",
        include_block_context: bool = True,
        expected_from_stage: str = "",
        expected_from_status: str = "",
    ) -> tuple[int, dict[str, object], str]:
        command = [
            sys.executable,
            str(writer),
            "--workspace-dir",
            str(workspace),
            "--event",
            name,
            "--stage",
            stage,
            "--status",
            status,
            "--transition-kind",
            transition_kind,
            "--message",
            f"Transition-policy selftest event {name}.",
            "--accept-current-revision",
            "--json",
        ]
        if expected_from_stage:
            command.extend(["--from-stage", expected_from_stage])
        if expected_from_status:
            command.extend(["--from-status", expected_from_status])
        if status in {"paused", "blocked"} and include_block_context:
            command.extend([
                "--blocker",
                blocker or "selftest transition blocker",
                "--resume-step",
                resume_step or "Review evidence/selftest/transition-context.json before resuming.",
            ])

        enhanced = transition_kind in {"resume", "skip", "return", "reopen"}
        if enhanced:
            if reason_code is default:
                reason_code = {
                    "resume": "recovery_requested",
                    "skip": "not_applicable",
                    "return": "validation_failed",
                    "reopen": "validation_failed",
                }[transition_kind]
            if reason_detail is default:
                reason_detail = "The selftest records the evidence-based reason for this non-default transition."
            if subjects is default:
                subjects = ["run:transition-selftest"]
            if evidence_refs is default:
                evidence_refs = ["evidence/selftest/transition-context.json"]
            if next_actions is default:
                next_actions = [transition_action(name)]
        else:
            if reason_code is default:
                reason_code = None
            if reason_detail is default:
                reason_detail = None
            if subjects is default:
                subjects = []
            if evidence_refs is default:
                evidence_refs = []
            if next_actions is default:
                next_actions = []

        if isinstance(reason_code, str) and reason_code:
            command.extend(["--reason-code", reason_code])
        if isinstance(subjects, list):
            for subject in subjects:
                command.extend(["--subject", subject])
        if isinstance(evidence_refs, list):
            for evidence_ref in evidence_refs:
                command.extend(["--evidence-ref", evidence_ref])
        if isinstance(next_actions, list):
            for action in next_actions:
                command.extend(["--next-action-json", json.dumps(action, sort_keys=True)])
        details: dict[str, str] = {"summary": f"Transition-policy selftest event {name}."}
        if isinstance(reason_detail, str):
            details["reason_detail"] = reason_detail
        command.extend(["--details-json", json.dumps(details, sort_keys=True)])
        proc = subprocess.run(command, cwd=plugin_root, capture_output=True, text=True)
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        try:
            payload = json.loads((proc.stdout or "").strip())
        except json.JSONDecodeError:
            payload = {}
        return proc.returncode, payload, output

    def require_ok(result: tuple[int, dict[str, object], str], label: str) -> dict[str, object]:
        code, payload, output = result
        if code != 0 or payload.get("ok") is not True:
            raise SystemExit(f"FAILED: transition policy positive case {label}: {output}")
        return payload

    def require_rejection(
        workspace: Path,
        label: str,
        expected_code: str,
        **kwargs: object,
    ) -> None:
        before = snapshot(workspace)
        code, payload, output = invoke(workspace, label, **kwargs)  # type: ignore[arg-type]
        if code == 0 or payload.get("code") != expected_code:
            raise SystemExit(
                f"FAILED: transition policy rejection {label} expected {expected_code}: {output}"
            )
        if snapshot(workspace) != before:
            raise SystemExit(f"FAILED: rejected transition {label} changed journal, state, or lock bytes")

    def seed_to_verification(workspace: Path) -> None:
        require_ok(
            invoke(
                workspace,
                "policy_started",
                stage="intake",
                status="running",
                transition_kind="start",
            ),
            "start",
        )
        for name, stage in [
            ("policy_recon_started", "recon"),
            ("policy_candidates_started", "candidate_generation"),
            ("policy_triage_started", "triage"),
            ("policy_verification_started", "verification"),
        ]:
            require_ok(
                invoke(
                    workspace,
                    name,
                    stage=stage,
                    status="running",
                    transition_kind="advance",
                ),
                name,
            )

    def prepared_enhanced_workspace(
        root: Path,
        transition_kind: str,
        label: str,
    ) -> tuple[Path, str, str]:
        workspace = root / f"enhanced-{transition_kind}-{label}"
        workspace.mkdir()
        if transition_kind == "resume":
            require_ok(
                invoke(
                    workspace,
                    "resume_started",
                    stage="intake",
                    status="running",
                    transition_kind="start",
                ),
                "resume start",
            )
            require_ok(
                invoke(
                    workspace,
                    "resume_blocked",
                    stage="intake",
                    status="blocked",
                    transition_kind="block",
                ),
                "resume block",
            )
            return workspace, "intake", "running"
        if transition_kind == "skip":
            seed_to_verification(workspace)
            return workspace, "severity_escalation", "completed"
        if transition_kind == "return":
            seed_to_verification(workspace)
            return workspace, "triage", "running"
        if transition_kind == "reopen":
            require_ok(
                invoke(
                    workspace,
                    "reopen_started",
                    stage="intake",
                    status="running",
                    transition_kind="start",
                ),
                "reopen start",
            )
            require_ok(
                invoke(
                    workspace,
                    "reopen_completed",
                    stage="intake",
                    status="completed",
                    transition_kind="complete",
                ),
                "reopen complete",
            )
            return workspace, "intake", "running"
        raise AssertionError(f"unknown enhanced kind: {transition_kind}")

    with tempfile.TemporaryDirectory(prefix="zhulong-transition-policy-") as temp_dir:
        root = Path(temp_dir)
        positive = root / "positive"
        positive.mkdir()
        first = require_ok(
            invoke(
                positive,
                "intake_started",
                stage="intake",
                status="running",
                transition_kind="start",
            ),
            "first start",
        )
        if first.get("seq") != 1:
            raise SystemExit("FAILED: transition policy first start did not allocate seq=1")
        require_ok(
            invoke(
                positive,
                "intake_observed",
                stage="intake",
                status="running",
                transition_kind="observe",
                expected_from_stage="intake",
                expected_from_status="running",
            ),
            "same-stage observe",
        )
        require_ok(
            invoke(
                positive,
                "intake_observed_at_locked_current",
                stage="current",
                status="current",
                transition_kind="observe",
            ),
            "locked current observe",
        )
        require_ok(
            invoke(
                positive,
                "intake_paused",
                stage="intake",
                status="paused",
                transition_kind="pause",
            ),
            "pause",
        )
        require_ok(
            invoke(
                positive,
                "intake_resumed_from_pause",
                stage="intake",
                status="running",
                transition_kind="resume",
            ),
            "resume paused",
        )
        require_ok(
            invoke(
                positive,
                "intake_blocked",
                stage="intake",
                status="blocked",
                transition_kind="block",
            ),
            "block",
        )
        require_ok(
            invoke(
                positive,
                "intake_resumed_from_block",
                stage="intake",
                status="running",
                transition_kind="resume",
            ),
            "resume blocked",
        )
        require_ok(
            invoke(
                positive,
                "intake_completed",
                stage="intake",
                status="completed",
                transition_kind="complete",
            ),
            "complete",
        )
        require_ok(
            invoke(
                positive,
                "intake_reopened",
                stage="intake",
                status="running",
                transition_kind="reopen",
            ),
            "reopen",
        )
        for name, stage in [
            ("recon_started", "recon"),
            ("candidates_started", "candidate_generation"),
            ("triage_started", "triage"),
            ("verification_started", "verification"),
        ]:
            require_ok(
                invoke(positive, name, stage=stage, status="running", transition_kind="advance"),
                f"advance {stage}",
            )
        require_ok(
            invoke(
                positive,
                "severity_skipped",
                stage="severity_escalation",
                status="completed",
                transition_kind="skip",
            ),
            "optional skip",
        )
        require_ok(
            invoke(
                positive,
                "variant_started",
                stage="variant_discovery",
                status="running",
                transition_kind="advance",
            ),
            "advance variant discovery",
        )
        require_ok(
            invoke(
                positive,
                "variant_returned_for_verification",
                stage="verification",
                status="running",
                transition_kind="return",
            ),
            "variant return to verification",
        )
        require_ok(
            invoke(
                positive,
                "verification_returned_to_triage",
                stage="triage",
                status="running",
                transition_kind="return",
            ),
            "verification return to triage",
        )
        require_ok(
            invoke(
                positive,
                "verification_restarted",
                stage="verification",
                status="running",
                transition_kind="advance",
            ),
            "reverification advance",
        )
        require_ok(
            invoke(
                positive,
                "packaging_started",
                stage="packaging",
                status="running",
                transition_kind="advance",
            ),
            "advance packaging",
        )
        require_ok(
            invoke(
                positive,
                "finalization_started",
                stage="finalization",
                status="running",
                transition_kind="advance",
            ),
            "advance finalization",
        )
        require_ok(
            invoke(
                positive,
                "finalization_completed",
                stage="finalization",
                status="completed",
                transition_kind="complete",
            ),
            "complete finalization state",
        )
        positive_events = [
            json.loads(line)
            for line in (positive / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        if any(event.get("stage") == "recording" for event in positive_events):
            raise SystemExit("FAILED: omitted recording was represented without a recording claim")
        require_ok(
            invoke(
                positive,
                "recording_skipped_without_archive_claim",
                stage="recording",
                status="completed",
                transition_kind="skip",
            ),
            "optional recording skip after completed finalization",
        )
        if (positive / "recordings").exists() or (positive / "recording-evidence.json").exists():
            raise SystemExit("FAILED: recording skip created a recording artifact or submission claim")
        positive_events = [
            json.loads(line)
            for line in (positive / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()
        ]
        for event in positive_events:
            if set(("from_stage", "transition_kind", "transition_policy_version", "blocker", "resume_step")) - set(event):
                raise SystemExit("FAILED: new P9.3 event omitted complete transition metadata")
        policy_output = run_capture(
            [sys.executable, str(validator), "--events-jsonl", str(positive / "audit-events.jsonl"), "--json"],
            plugin_root,
        )
        if json.loads(policy_output).get("transition_policy") != "transition_policy_v1":
            raise SystemExit("FAILED: policy journal did not validate as transition_policy_v1")

        for transition_kind in ("resume", "skip", "return", "reopen"):
            for label, expected_code, kwargs in [
                ("missing_reason_detail", "TRANSITION_REASON_DETAIL_REQUIRED", {"reason_detail": None}),
                ("missing_subject", "TRANSITION_SUBJECT_REQUIRED", {"subjects": []}),
                ("missing_evidence", "TRANSITION_EVIDENCE_REQUIRED", {"evidence_refs": []}),
                ("missing_next_action", "TRANSITION_NEXT_ACTION_REQUIRED", {"next_actions": []}),
            ]:
                workspace, stage, status = prepared_enhanced_workspace(root, transition_kind, label)
                require_rejection(
                    workspace,
                    f"{transition_kind}_{label}",
                    expected_code,
                    stage=stage,
                    status=status,
                    transition_kind=transition_kind,
                    **kwargs,
                )

        paused = root / "paused"
        paused.mkdir()
        require_ok(invoke(paused, "paused_start", stage="intake", status="running", transition_kind="start"), "paused start")
        require_ok(invoke(paused, "paused_pause", stage="intake", status="paused", transition_kind="pause"), "paused state")
        require_rejection(
            paused,
            "paused_observe_running",
            "OBSERVE_STATE_CHANGED",
            stage="intake",
            status="running",
            transition_kind="observe",
        )
        blocked = root / "blocked"
        blocked.mkdir()
        require_ok(invoke(blocked, "blocked_start", stage="intake", status="running", transition_kind="start"), "blocked start")
        require_ok(invoke(blocked, "blocked_block", stage="intake", status="blocked", transition_kind="block"), "blocked state")
        require_rejection(
            blocked,
            "blocked_observe_running",
            "OBSERVE_STATE_CHANGED",
            stage="intake",
            status="running",
            transition_kind="observe",
        )
        require_rejection(
            blocked,
            "blocked_resume_with_unsafe_evidence",
            "EVENT_VALIDATION_FAILED",
            stage="intake",
            status="running",
            transition_kind="resume",
            evidence_refs=["https://example.invalid/evidence"],
        )

        completed = root / "completed"
        completed.mkdir()
        require_ok(invoke(completed, "completed_start", stage="intake", status="running", transition_kind="start"), "completed start")
        require_ok(invoke(completed, "completed_complete", stage="intake", status="completed", transition_kind="complete"), "completed state")
        require_rejection(
            completed,
            "completed_observe_running",
            "OBSERVE_STATE_CHANGED",
            stage="intake",
            status="running",
            transition_kind="observe",
        )
        require_rejection(
            completed,
            "completed_resume",
            "RESUME_TRANSITION_INVALID",
            stage="intake",
            status="running",
            transition_kind="resume",
        )
        running = root / "running"
        running.mkdir()
        require_ok(invoke(running, "running_start", stage="intake", status="running", transition_kind="start"), "running start")
        require_rejection(
            running,
            "pause_without_context",
            "EVENT_VALIDATION_FAILED",
            stage="intake",
            status="paused",
            transition_kind="pause",
            include_block_context=False,
        )
        require_rejection(
            running,
            "block_without_context",
            "EVENT_VALIDATION_FAILED",
            stage="intake",
            status="blocked",
            transition_kind="block",
            include_block_context=False,
        )
        require_rejection(
            running,
            "running_resume",
            "RESUME_TRANSITION_INVALID",
            stage="intake",
            status="running",
            transition_kind="resume",
        )
        require_rejection(
            running,
            "running_reopen",
            "REOPEN_TRANSITION_INVALID",
            stage="intake",
            status="running",
            transition_kind="reopen",
        )
        require_rejection(
            running,
            "observe_changes_stage",
            "OBSERVE_STATE_CHANGED",
            stage="recon",
            status="running",
            transition_kind="observe",
        )
        require_rejection(
            running,
            "source_stage_mismatch",
            "SOURCE_STAGE_MISMATCH",
            stage="intake",
            status="running",
            transition_kind="observe",
            expected_from_stage="recon",
        )
        require_rejection(
            running,
            "source_status_mismatch",
            "SOURCE_STATUS_MISMATCH",
            stage="intake",
            status="running",
            transition_kind="observe",
            expected_from_status="blocked",
        )
        require_rejection(
            running,
            "unknown_transition_kind",
            "INVALID_TRANSITION_KIND",
            stage="intake",
            status="running",
            transition_kind="teleport",
        )

        backward = root / "backward"
        backward.mkdir()
        seed_to_verification(backward)
        require_rejection(
            backward,
            "advance_backward",
            "ADVANCE_STAGE_EDGE_INVALID",
            stage="triage",
            status="running",
            transition_kind="advance",
        )
        require_rejection(
            backward,
            "return_forward",
            "RETURN_STAGE_EDGE_INVALID",
            stage="packaging",
            status="running",
            transition_kind="return",
        )
        require_rejection(
            backward,
            "skip_mandatory_packaging",
            "SKIP_TRANSITION_INVALID",
            stage="packaging",
            status="completed",
            transition_kind="skip",
        )

        promotion = root / "promotion-boundary"
        promotion.mkdir()
        (promotion / "asr-config.json").write_text('{"schema_version":1}\n', encoding="utf-8")
        (promotion / "candidate-findings.md").write_text(
            "Scanner and LLM notes say confirmed_in_docker, but no authoritative verdict exists.\n",
            encoding="utf-8",
        )
        (promotion / "evidence").mkdir()
        (promotion / "evidence/docker-only.json").write_text(
            '{"verification_status":"confirmed_in_docker","authority":"docker-only"}\n',
            encoding="utf-8",
        )
        seed_to_verification(promotion)
        require_ok(
            invoke(
                promotion,
                "packaging_observed_after_scanner_note",
                stage="packaging",
                status="running",
                transition_kind="advance",
            ),
            "syntactic verification to packaging transition",
        )
        require_ok(
            invoke(
                promotion,
                "fake_finalization_started",
                stage="finalization",
                status="running",
                transition_kind="advance",
            ),
            "fake finalization transition",
        )
        require_ok(
            invoke(
                promotion,
                "fake_finalization_completed",
                stage="finalization",
                status="completed",
                transition_kind="complete",
            ),
            "fake finalization complete",
        )
        require_ok(
            invoke(
                promotion,
                "fake_recording_started",
                stage="recording",
                status="running",
                transition_kind="advance",
            ),
            "fake recording transition",
        )
        for generated in (
            "confirmed",
            "audit-disposition.json",
            "handoff-summary.md",
            "SUMMARY.md",
            "recordings",
            "recording-evidence.json",
        ):
            if (promotion / generated).exists():
                raise SystemExit(
                    "FAILED: workflow transition created a promotion or completion artifact: " + generated
                )
        assertion = subprocess.run(
            [
                sys.executable,
                str(plugin_root / "scripts/assert_finalized_workspace.py"),
                "--workspace-dir",
                str(promotion),
                "--json",
            ],
            cwd=plugin_root,
            capture_output=True,
            text=True,
        )
        if assertion.returncode == 0:
            raise SystemExit("FAILED: fake transition-only finalization passed substantive finalization assertion")

        prepolicy = root / "prepolicy-prefix"
        prepolicy.mkdir()
        legacy_prefix = (
            plugin_root / "assets/fixtures/audit-state-protocol-r2/valid-events-r2.jsonl"
        ).read_bytes()
        prefix_events = [
            json.loads(line)
            for line in legacy_prefix.decode("utf-8").splitlines()
            if line.strip()
        ]
        latest = prefix_events[-1]
        (prepolicy / "audit-events.jsonl").write_bytes(legacy_prefix)
        (prepolicy / "stage-status.json").write_text(
            json.dumps(
                {
                    "schema_version": 2,
                    "plugin": "zhulong",
                    "plugin_version": "fixture",
                    "run_id": latest["run_id"],
                    "state_revision": len(prefix_events),
                    "last_event_seq": latest["seq"],
                    "event_log_digest": "sha256:" + hashlib.sha256(legacy_prefix).hexdigest(),
                    "stage": latest["stage"],
                    "status": latest["to_status"],
                    "last_event_at": latest["ts"],
                    "last_event_type": latest["event_type"],
                    "last_event_name": latest["event_name"],
                    "blocker": None,
                    "resume_step": None,
                },
                sort_keys=True,
                indent=2,
            )
            + "\n",
            encoding="utf-8",
        )
        require_ok(
            invoke(
                prepolicy,
                "prepolicy_recon_observed",
                stage="recon",
                status="running",
                transition_kind="observe",
            ),
            "pre-policy R2 continuation",
        )
        prefix_output = run_capture(
            [sys.executable, str(validator), "--events-jsonl", str(prepolicy / "audit-events.jsonl"), "--json"],
            plugin_root,
        )
        if json.loads(prefix_output).get("transition_policy") != "pre_policy_r2_prefix_then_transition_policy_v1":
            raise SystemExit("FAILED: pre-P9.3 R2 prefix was not visibly classified at the policy boundary")

    print("AUDIT TRANSITION POLICY SELFTEST PASSED: FSM/compatibility/rejection/promotion boundaries")


def exercise_triage_batch_and_stage_finalization(plugin_root: Path) -> None:
    """Exercise the production P9 triage/finalizer paths without runtime execution."""
    validator = plugin_root / "scripts/validate_triage_batch.py"
    finalizer = plugin_root / "scripts/finalize_stage.py"
    writer = plugin_root / "scripts/write_audit_event.py"
    fixture_root = plugin_root / "assets/fixtures/triage-batch"
    required = [
        "assets/schemas/triage-batch.schema.json",
        "docs/runner-contracts/triage-batch-contract-r1.md",
        "assets/fixtures/triage-batch/README.md",
        "assets/fixtures/triage-batch/manifest.json",
        "scripts/validate_triage_batch.py",
        "scripts/finalize_stage.py",
    ]
    require_files(plugin_root, required, "triage batch/finalizer")
    manifest = json.loads((fixture_root / "manifest.json").read_text(encoding="utf-8"))
    for fixture_id in (
        "complete-five-advisories", "partial-unprocessed-candidate", "blocked-batch-recovery",
        "duplicate-cycle", "candidate-id-path-swap", "wrong-result-digest", "duplicate-finalization",
    ):
        if fixture_id not in set(manifest.get("positive_cases", [])) | set(manifest.get("negative_cases", [])):
            raise SystemExit(f"FAILED: triage fixture manifest omitted {fixture_id}")

    def clone(value: object) -> object:
        return json.loads(json.dumps(value))

    def digest(path: Path) -> str:
        return "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()

    def invoke_json(command: list[str], *, expected_returncode: int, label: str) -> dict:
        proc = subprocess.run(command, cwd=plugin_root, capture_output=True, text=True)
        output = ((proc.stdout or "") + (proc.stderr or "")).strip()
        if proc.returncode != expected_returncode:
            raise SystemExit(f"FAILED: {label} returned {proc.returncode}, expected {expected_returncode}: {output}")
        try:
            return json.loads((proc.stdout or "").strip())
        except json.JSONDecodeError as exc:
            raise SystemExit(f"FAILED: {label} did not emit JSON: {output}") from exc

    def decision(candidate_id: str, recommendation: str, **extra: object) -> dict:
        value: dict[str, object] = {
            "candidate_id": candidate_id,
            "recommendation": recommendation,
            "reason_code": "TRIAGE_REVIEW_REQUIRED",
            "evidence": ["Sanitized candidate-only evidence was reviewed."],
            "next_action": "Keep this advisory record for the next manual review step.",
        }
        value.update(extra)
        return value

    def materialize(root: Path) -> tuple[Path, Path, dict, list[str]]:
        repo = root / "repo"
        workspace = root / "workspace"
        recon_fixture = plugin_root / "assets/fixtures/recon-result/service"
        shutil.copytree(recon_fixture / "repo", repo)
        shutil.copytree(recon_fixture / "workspace", workspace)
        shutil.copy2(workspace / "cases/complete-service.json", workspace / "recon-result.json")
        contract_fixture = plugin_root / "assets/fixtures/contracts/confirmed_ssrf"
        base_candidate = json.loads((contract_fixture / "candidate.json").read_text(encoding="utf-8"))
        candidate_ids: list[str] = []
        inventory: list[dict] = []
        for index in range(1, 6):
            candidate_id = f"CAND-TRIAGE-{index}"
            candidate_ids.append(candidate_id)
            candidate = clone(base_candidate)
            assert isinstance(candidate, dict)
            candidate["candidate_id"] = candidate_id
            candidate["title"] = f"Sanitized triage candidate {index}"
            candidate["finder"]["created_at"] = f"2026-06-{18 + index:02d}T00:00:00Z"
            candidate["target_ref"] = {
                "target_config": "zhulong-target.yaml",
                "tested_ref": "fixture-service-ref-001",
            }
            path = workspace / "candidates" / candidate_id / "candidate.json"
            path.parent.mkdir(parents=True)
            path.write_text(json.dumps(candidate, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            inventory.append({"path": path.relative_to(workspace).as_posix(), "sha256": digest(path), "candidate_id": candidate_id})
        batch = {
            "schema_version": 1,
            "batch_id": "TRIAGE-SANITIZED-1",
            "status": "complete",
            "target_binding": {
                "target_contract_path": "zhulong-target.yaml",
                "target_contract_sha256": digest(workspace / "zhulong-target.yaml"),
                "tested_ref": "fixture-service-ref-001",
            },
            "recon_binding": {
                "path": "recon-result.json",
                "sha256": digest(workspace / "recon-result.json"),
                "recon_id": "RECON-SERVICE-001",
            },
            "candidate_inventory": inventory,
            "decisions": [
                decision(candidate_ids[0], "recommend_verification", verification_reason="The entrypoint-to-sink claim merits independent review.", docker_applicability="applicable", required_evidence=["Attacker entrypoint and deterministic oracle evidence."], verification_order=1),
                decision(candidate_ids[1], "unverified", missing_evidence=["A deterministic entrypoint observation is missing."], next_action="Collect the missing candidate-only evidence before requesting verification."),
                decision(candidate_ids[2], "false_positive", counterevidence=["The documented boundary rejects this candidate-only path."], next_action="Retain the counterevidence for reviewer inspection."),
                decision(candidate_ids[3], "duplicate", duplicate_of_candidate_id=candidate_ids[2], next_action="Review the canonical candidate before any independent verification."),
                decision(candidate_ids[4], "blocked", blocker_code="FIXTURE_RUNTIME_DEPENDENCY", recovery_condition="The sanitized prerequisite is available.", resume_action="Resume candidate-only review after the prerequisite is recorded.", next_action="Resume only after the recorded recovery condition."),
            ],
            "unprocessed_candidates": [],
            "batch_gaps": [],
            "batch_blockers": [],
            "next_actions": [{"action_id": "ACTION-TRIAGE-REVIEW", "summary": "Review this advisory batch without changing a disposition."}],
        }
        return repo, workspace, batch, candidate_ids

    def validate_batch(repo: Path, workspace: Path, batch: dict, *, expected_ok: bool, expected_codes: set[str], label: str) -> dict:
        path = workspace / "triage-batch.json"
        path.write_text(json.dumps(batch, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        payload = invoke_json([
            sys.executable, str(validator), "--repo-root", str(repo), "--workspace-dir", str(workspace),
            "--triage-batch", "triage-batch.json", "--json",
        ], expected_returncode=0 if expected_ok else 1, label=label)
        if bool(payload.get("ok")) != expected_ok:
            raise SystemExit(f"FAILED: {label} success mismatch: {payload}")
        actual_codes = set(payload.get("issue_codes", []))
        if not expected_codes.issubset(actual_codes):
            raise SystemExit(f"FAILED: {label} codes {actual_codes} do not include {expected_codes}")
        if (workspace / "audit-disposition.json").exists() or (workspace / "audit-events.jsonl").exists():
            raise SystemExit(f"FAILED: read-only triage validator wrote downstream authority material for {label}")
        return payload

    with tempfile.TemporaryDirectory(prefix="zhulong-triage-batch-") as tempdir:
        root = Path(tempdir)
        repo, workspace, complete, candidate_ids = materialize(root)
        validate_batch(repo, workspace, complete, expected_ok=True, expected_codes=set(), label="complete five advisories")

        invalid_recon_digest = clone(complete)
        assert isinstance(invalid_recon_digest, dict)
        invalid_recon_digest["recon_binding"]["sha256"] = "sha256:" + "0" * 64
        validate_batch(repo, workspace, invalid_recon_digest, expected_ok=False, expected_codes={"DIGEST_MISMATCH_RECON"}, label="recon digest drift")

        invalid_recon_id = clone(complete)
        assert isinstance(invalid_recon_id, dict)
        invalid_recon_id["recon_binding"]["recon_id"] = "RECON-SERVICE-OTHER"
        validate_batch(repo, workspace, invalid_recon_id, expected_ok=False, expected_codes={"RECON_ID_MISMATCH"}, label="recon id mismatch")

        original_recon = (workspace / "recon-result.json").read_bytes()
        invalid_recon = json.loads(original_recon.decode("utf-8"))
        invalid_recon["status"] = "invalid"
        (workspace / "recon-result.json").write_text(json.dumps(invalid_recon, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        invalid_recon_contract = clone(complete)
        assert isinstance(invalid_recon_contract, dict)
        invalid_recon_contract["recon_binding"]["sha256"] = digest(workspace / "recon-result.json")
        validate_batch(repo, workspace, invalid_recon_contract, expected_ok=False, expected_codes={"RECON_BINDING_INVALID"}, label="recon production validator rejection")
        (workspace / "recon-result.json").write_bytes(original_recon)

        escaped_recon = clone(complete)
        assert isinstance(escaped_recon, dict)
        escaped_recon["recon_binding"]["path"] = "../recon-result.json"
        validate_batch(repo, workspace, escaped_recon, expected_ok=False, expected_codes={"PATH_UNSAFE"}, label="recon path escape")

        missing_recon = clone(complete)
        assert isinstance(missing_recon, dict)
        missing_recon["recon_binding"]["path"] = "missing-recon-result.json"
        validate_batch(repo, workspace, missing_recon, expected_ok=False, expected_codes={"FILE_MISSING"}, label="recon missing")

        original_target = (workspace / "zhulong-target.yaml").read_bytes()
        moved_recon = workspace / "recon-result.real.json"
        (workspace / "recon-result.json").rename(moved_recon)
        (workspace / "recon-result.json").symlink_to(moved_recon.name)
        try:
            symlink_recon = clone(complete)
            assert isinstance(symlink_recon, dict)
            validate_batch(repo, workspace, symlink_recon, expected_ok=False, expected_codes={"SYMLINK_ESCAPE"}, label="recon symlink")
        finally:
            (workspace / "recon-result.json").unlink()
            moved_recon.rename(workspace / "recon-result.json")

        swapped_candidate = clone(complete)
        assert isinstance(swapped_candidate, dict)
        swapped_candidate["candidate_inventory"][0]["path"] = "recon-result.json"
        swapped_candidate["candidate_inventory"][0]["sha256"] = digest(workspace / "recon-result.json")
        validate_batch(repo, workspace, swapped_candidate, expected_ok=False, expected_codes={"CANDIDATE_INVALID"}, label="candidate recon binding swap")

        swapped_target = clone(complete)
        assert isinstance(swapped_target, dict)
        swapped_target["target_binding"]["target_contract_path"] = "recon-result.json"
        swapped_target["target_binding"]["target_contract_sha256"] = digest(workspace / "recon-result.json")
        validate_batch(repo, workspace, swapped_target, expected_ok=False, expected_codes={"TARGET_CONTRACT_INVALID"}, label="target recon binding swap")
        if (workspace / "zhulong-target.yaml").read_bytes() != original_target:
            raise SystemExit("FAILED: triage mutation matrix changed the target contract")

        partial = clone(complete)
        assert isinstance(partial, dict)
        partial["status"] = "partial"
        partial["decisions"] = partial["decisions"][:-1]
        partial["unprocessed_candidates"] = [{"candidate_id": candidate_ids[4], "reason_code": "EVIDENCE_PENDING", "evidence": ["Sanitized collection has not completed."], "next_action": "Collect the explicit missing evidence."}]
        validate_batch(repo, workspace, partial, expected_ok=True, expected_codes=set(), label="partial batch")

        blocked = clone(complete)
        assert isinstance(blocked, dict)
        blocked["status"] = "blocked"
        blocked["batch_blockers"] = [{"code": "BATCH_FIXTURE_BLOCKER", "affected_candidate_ids": [candidate_ids[4]], "evidence": ["A sanitized batch-level prerequisite is absent."], "recovery_condition": "The prerequisite is restored.", "resume_action": "Resume triage after restoration."}]
        validate_batch(repo, workspace, blocked, expected_ok=True, expected_codes=set(), label="blocked batch")

        cases: list[tuple[str, dict, set[str]]] = []
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["decisions"][3]["duplicate_of_candidate_id"] = candidate_ids[3]; cases.append(("duplicate self", invalid, {"DUPLICATE_SELF_REFERENCE"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["decisions"][2]["recommendation"] = "duplicate"; invalid["decisions"][2].pop("counterevidence"); invalid["decisions"][2]["duplicate_of_candidate_id"] = candidate_ids[3]; cases.append(("duplicate cycle", invalid, {"DUPLICATE_CYCLE"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["decisions"][3]["duplicate_of_candidate_id"] = "CAND-TRIAGE-MISSING"; cases.append(("duplicate missing", invalid, {"DUPLICATE_TARGET_UNKNOWN"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["candidate_inventory"][0]["sha256"] = "sha256:" + "0" * 64; cases.append(("candidate digest drift", invalid, {"DIGEST_MISMATCH_CANDIDATE"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["candidate_inventory"][0]["candidate_id"] = candidate_ids[1]; cases.append(("candidate id path swap", invalid, {"DUPLICATE_CANDIDATE_ID", "CANDIDATE_ID_PATH_SWAP"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["candidate_inventory"][0]["path"] = "../candidate.json"; cases.append(("candidate path escape", invalid, {"PATH_UNSAFE"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["decisions"] = invalid["decisions"][:-1]; cases.append(("omitted decision", invalid, {"CANDIDATE_DECISION_OMITTED", "BATCH_COMPLETENESS_INVALID"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["decisions"].append(clone(invalid["decisions"][0])); cases.append(("duplicate decision", invalid, {"DUPLICATE_DECISION"}))
        invalid = clone(partial); assert isinstance(invalid, dict); invalid["decisions"].append(clone(complete["decisions"][4])); cases.append(("processed unprocessed overlap", invalid, {"PROCESSED_UNPROCESSED_OVERLAP"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["disposition"] = "confirmed"; cases.append(("forbidden disposition", invalid, {"FORBIDDEN_AUTHORITY_FIELD", "SCHEMA_INVALID"}))
        invalid = clone(complete); assert isinstance(invalid, dict); invalid["promotion"] = "confirmed"; cases.append(("forbidden promotion", invalid, {"FORBIDDEN_AUTHORITY_FIELD", "SCHEMA_INVALID"}))
        for label, bad_batch, codes in cases:
            validate_batch(repo, workspace, bad_batch, expected_ok=False, expected_codes=codes, label=label)

        validate_batch(repo, workspace, complete, expected_ok=True, expected_codes=set(), label="restored complete triage")
        for event_name, stage, transition in (
            ("triage_handoff_intake", "intake", "start"),
            ("triage_handoff_recon", "recon", "advance"),
            ("triage_handoff_candidates", "candidate_generation", "advance"),
            ("triage_handoff_started", "triage", "advance"),
        ):
            invoke_json([
                sys.executable, str(writer), "--workspace-dir", str(workspace),
                "--event", event_name, "--stage", stage, "--status", "running",
                "--transition-kind", transition, "--message", "Seed a portable handoff triage fixture.",
                "--accept-current-revision", "--json",
            ], expected_returncode=0, label=f"handoff {event_name}")
        handoff = invoke_json([
            sys.executable, str(plugin_root / "scripts/render_handoff_state.py"),
            "--workspace-dir", str(workspace), "--repo-root", str(repo), "--json",
        ], expected_returncode=0, label="handoff triage validator parity")
        handoff_state = handoff.get("state", {})
        triage_state = handoff_state.get("triage", {}) if isinstance(handoff_state, dict) else {}
        issue_codes = {
            str(item.get("code"))
            for item in handoff_state.get("integrity", {}).get("issues", [])
            if isinstance(item, dict)
        } if isinstance(handoff_state, dict) else set()
        if triage_state.get("status") != "complete" or "STRUCTURED_RESULT_INVALID" in issue_codes:
            raise SystemExit(
                "FAILED: handoff triage status diverged from the production triage CLI "
                f"status={triage_state.get('status')} issues={sorted(issue_codes)}"
            )
        help_output = run_capture([sys.executable, str(validator), "--help"], plugin_root)
        if "--recon-result" in help_output:
            raise SystemExit("FAILED: triage CLI unexpectedly exposes a second Recon input")

    def seed_running_stage(workspace: Path, stage: str) -> int:
        stages = ["intake", "recon", "candidate_generation", "triage"]
        target_index = stages.index(stage)
        for index, name in enumerate(stages[:target_index + 1]):
            transition = "start" if index == 0 else "advance"
            payload = invoke_json([
                sys.executable, str(writer), "--workspace-dir", str(workspace), "--event", f"p9_seed_{name}",
                "--stage", name, "--status", "running", "--transition-kind", transition,
                "--message", "Seed a deterministic P9 stage-finalizer fixture.", "--accept-current-revision", "--json",
            ], expected_returncode=0, label=f"seed {stage}")
            if payload.get("state_revision") != index + 1:
                raise SystemExit(f"FAILED: seed revision mismatch for {stage}: {payload}")
        return target_index + 1

    def finalize(repo: Path, workspace: Path, stage: str, result_name: str, revision: int, *, expected_returncode: int, label: str, expected_digest: str | None = None) -> dict:
        result_path = workspace / result_name
        return invoke_json([
            sys.executable, str(finalizer), "--workspace-dir", str(workspace), "--repo-root", str(repo),
            "--stage", stage, "--result", result_name, "--expected-result-sha256", expected_digest or digest(result_path),
            "--expected-state-revision", str(revision), "--json",
        ], expected_returncode=expected_returncode, label=label)

    recon_fixture = plugin_root / "assets/fixtures/recon-result/service"
    for result_name, expected_status in [("complete-service.json", "completed"), ("partial-service.json", "paused"), ("blocked-service.json", "blocked")]:
        with tempfile.TemporaryDirectory(prefix="zhulong-finalize-recon-") as tempdir:
            root = Path(tempdir)
            repo = root / "repo"; workspace = root / "workspace"
            shutil.copytree(recon_fixture / "repo", repo)
            shutil.copytree(recon_fixture / "workspace", workspace)
            shutil.copy2(workspace / "cases" / result_name, workspace / "recon-result.json")
            revision = seed_running_stage(workspace, "recon")
            payload = finalize(repo, workspace, "recon", "recon-result.json", revision, expected_returncode=0, label=f"recon {result_name}")
            if payload.get("journal_committed") is not True or payload.get("state_view_updated") is not True:
                raise SystemExit(f"FAILED: recon finalization did not report a committed state: {payload}")
            state = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
            if state.get("stage") != "recon" or state.get("status") != expected_status:
                raise SystemExit(f"FAILED: recon finalization state mismatch: {state}")
            if (workspace / "audit-disposition.json").exists():
                raise SystemExit("FAILED: stage finalizer wrote audit-disposition.json")

    for status_name, batch_status in [("complete", "complete"), ("partial", "partial"), ("blocked", "blocked")]:
        with tempfile.TemporaryDirectory(prefix="zhulong-finalize-triage-") as tempdir:
            root = Path(tempdir)
            repo, workspace, complete, candidate_ids = materialize(root)
            batch = {"complete": complete, "partial": clone(complete), "blocked": clone(complete)}[status_name]
            assert isinstance(batch, dict)
            if batch_status == "partial":
                batch["status"] = "partial"; batch["decisions"] = batch["decisions"][:-1]
                batch["unprocessed_candidates"] = [{"candidate_id": candidate_ids[4], "reason_code": "EVIDENCE_PENDING", "evidence": ["Pending."], "next_action": "Collect evidence."}]
            if batch_status == "blocked":
                batch["status"] = "blocked"; batch["batch_blockers"] = [{"code": "BATCH_FIXTURE_BLOCKER", "affected_candidate_ids": [candidate_ids[4]], "evidence": ["Blocked."], "recovery_condition": "Restored.", "resume_action": "Resume triage."}]
            result_path = workspace / "triage-batch.json"; result_path.write_text(json.dumps(batch, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            revision = seed_running_stage(workspace, "triage")
            journal_before = (workspace / "audit-events.jsonl").read_bytes()
            state_before = (workspace / "stage-status.json").read_bytes()
            wrong = finalize(repo, workspace, "triage", "triage-batch.json", revision, expected_returncode=1, label="wrong triage digest", expected_digest="sha256:" + "f" * 64)
            if wrong.get("code") != "RESULT_DIGEST_CONFLICT" or (workspace / "audit-events.jsonl").read_bytes() != journal_before or (workspace / "stage-status.json").read_bytes() != state_before:
                raise SystemExit("FAILED: failed digest CAS changed audit state")
            wrong_type = finalize(repo, workspace, "recon", "triage-batch.json", revision, expected_returncode=1, label="wrong stage/result type")
            if wrong_type.get("code") != "CONTRACT_INVALID" or (workspace / "audit-events.jsonl").read_bytes() != journal_before or (workspace / "stage-status.json").read_bytes() != state_before:
                raise SystemExit("FAILED: wrong stage/result type changed audit state")
            stale = finalize(repo, workspace, "triage", "triage-batch.json", revision + 1, expected_returncode=1, label="stale triage revision")
            if stale.get("code") != "STATE_REVISION_CONFLICT" or (workspace / "audit-events.jsonl").read_bytes() != journal_before or (workspace / "stage-status.json").read_bytes() != state_before:
                raise SystemExit("FAILED: stale revision changed audit state")
            payload = finalize(repo, workspace, "triage", "triage-batch.json", revision, expected_returncode=0, label=f"triage {status_name}")
            events_after = (workspace / "audit-events.jsonl").read_bytes()
            duplicate = finalize(repo, workspace, "triage", "triage-batch.json", revision, expected_returncode=1, label="duplicate triage finalization")
            if duplicate.get("code") not in {"STATE_REVISION_CONFLICT", "CURRENT_STAGE_MISMATCH"} or (workspace / "audit-events.jsonl").read_bytes() != events_after:
                raise SystemExit(f"FAILED: duplicate finalization appended a second terminal event: {duplicate}")
            event = json.loads(events_after.decode("utf-8").splitlines()[-1])
            forbidden = {"confirmed", "severity", "cvss", "verdict", "disposition", "bundle_ready", "promotion"}
            if forbidden & set(json.dumps(event).lower().replace("-", "_").split('"')):
                raise SystemExit("FAILED: triage event contains forbidden downstream authority wording")
            expected_state = {"complete": "completed", "partial": "paused", "blocked": "blocked"}[status_name]
            if json.loads((workspace / "stage-status.json").read_text(encoding="utf-8")).get("status") != expected_state:
                raise SystemExit("FAILED: triage finalizer mapping mismatch")

    print("TRIAGE BATCH / STAGE FINALIZATION SELFTEST PASSED: advisory/binding/CAS/R2 matrix")


def selftest_installed_skill(skill_root: Path) -> None:
    for rel in INSTALLED_SKILL_REQUIRED_FILES:
        path = skill_root / rel
        if not path.exists():
            raise SystemExit(f"FAILED: missing required installed skill file: {path}")
    for rel in FORBIDDEN_INSTALLED_TOP_LEVEL:
        path = skill_root / rel
        if path.exists():
            raise SystemExit(f"FAILED: installed skill contains forbidden top-level material: {path}")
    timeline_bundle_fixture = (
        skill_root
        / "assets/fixtures/audit-timeline/completed-confirmed/workspace"
    )
    if timeline_bundle_fixture.exists():
        raise SystemExit(
            "FAILED: installed skill contains a materialized audit-timeline confirmed bundle; "
            "the production builder fixture must stay temporary"
        )

    run([sys.executable, "-m", "py_compile",
         str(skill_root / "scripts/plan_security_toolchain.py"),
         str(skill_root / "scripts/validate_tool_registry.py"),
         str(skill_root / "scripts/workspace_state.py"),
         str(skill_root / "scripts/render_handoff_summary.py"),
         str(skill_root / "scripts/render_handoff_state.py"),
         str(skill_root / "scripts/validate_handoff_state.py"),
         str(skill_root / "scripts/create_workspace_checkpoint.py"),
         str(skill_root / "scripts/validate_workspace_checkpoint.py"),
         str(skill_root / "scripts/next_actions.py"),
         str(skill_root / "scripts/render_next_actions.py"),
         str(skill_root / "scripts/validate_next_actions.py"),
         str(skill_root / "scripts/audit_timeline.py"),
         str(skill_root / "scripts/render_audit_timeline.py"),
         str(skill_root / "scripts/validate_audit_timeline.py"),
         str(skill_root / "scripts/selftest_audit_timeline.py"),
         str(skill_root / "scripts/assert_finalized_workspace.py"),
         str(skill_root / "scripts/audit_disposition.py"),
         str(skill_root / "scripts/blocked_verification.py"),
         str(skill_root / "scripts/audit_state_io.py"),
         str(skill_root / "scripts/audit_text_safety.py"),
         str(skill_root / "scripts/audit_transition_policy.py"),
         str(skill_root / "scripts/write_audit_event.py"),
         str(skill_root / "scripts/validate_audit_protocol.py"),
         str(skill_root / "scripts/recover_audit_state.py"),
         str(skill_root / "scripts/validate_workspace_state.py"),
         str(skill_root / "scripts/validate_target_contract.py"),
         str(skill_root / "scripts/validate_recon_result.py"),
         str(skill_root / "scripts/validate_triage_batch.py"),
         str(skill_root / "scripts/finalize_stage.py"),
         str(skill_root / "scripts/validate_candidate.py"),
         str(skill_root / "scripts/validate_verifier_verdict.py"),
         str(skill_root / "scripts/validate_bundle_contract.py"),
         str(skill_root / "scripts/build_confirmed_bundle.py"),
         str(skill_root / "scripts/p8_dogfood_metrics.py"),
         str(skill_root / "scripts/verify_candidate.py"),
         str(skill_root / "scripts/check_sandbox_preflight.py"),
         str(skill_root / "scripts/evidence_io.py"),
         str(skill_root / "scripts/manage_docker_resources.py"),
         str(skill_root / "scripts/render_confirmed_vuln_docx.py"),
         str(skill_root / "scripts/recording_identity.py"),
         str(skill_root / "scripts/auto_record_bundle.py"),
         str(skill_root / "scripts/validate_recording_evidence.py"),
         str(skill_root / "scripts/extract_variant_seed.py"),
         str(skill_root / "scripts/find_variant_candidates.py"),
         str(skill_root / "scripts/validate_report_bundle.py"),
         str(skill_root / "scripts/validate_all_report_bundles.py"),
         str(skill_root / "scripts/finalize_audit_workspace.py")], skill_root)
    exercise_target_contract_validator(skill_root)
    exercise_tool_registry_contract(skill_root)
    exercise_context_planning_contract(skill_root)
    exercise_root_skill_kernel_contract(skill_root)
    exercise_recon_result_contract(skill_root)
    exercise_triage_batch_and_stage_finalization(skill_root)
    exercise_finding_contract_validators(skill_root)
    exercise_bundle_contract_validator(skill_root)
    exercise_build_confirmed_bundle_wrapper(skill_root)
    exercise_p8_closure_contracts(skill_root)
    exercise_replay_transcript_corpus(skill_root)
    exercise_p8_dogfood_metrics(skill_root)
    exercise_p8_real_historical_dogfood(skill_root)
    exercise_p9_protocol_chain_real_workspace_dogfood(skill_root)
    exercise_independent_verifier(skill_root)
    exercise_disposition_integration(skill_root)
    exercise_contract_fixture_chain(skill_root)
    exercise_audit_state_protocol_r2(skill_root)
    exercise_audit_state_protocol_closure(skill_root)
    exercise_handoff_checkpoint_contract(skill_root)
    exercise_next_actions_contract(skill_root)
    run([sys.executable, str(skill_root / "scripts/selftest_audit_timeline.py")], skill_root)
    exercise_audit_state_writer(skill_root)
    exercise_audit_state_recovery(skill_root)
    exercise_audit_transition_policy(skill_root)
    exercise_recording_evidence_gate(skill_root)

    for script in [
        "scripts/bootstrap_verification_workspace.sh",
        "scripts/asr_start.sh",
        "scripts/resolve_skill_root.sh",
        "scripts/zhulong_audit.sh",
        "scripts/prepare_target_repo.sh",
        "scripts/check_docker_gate.sh",
        "scripts/check_omc_runtime.sh",
        "scripts/check_security_tooling.sh",
        "scripts/run_initial_probes.sh",
        "scripts/run_verification_case.sh",
    ]:
        run(["bash", "-n", str(skill_root / script)], skill_root)

    expected_root = str(skill_root.resolve())
    require_command_output(
        ["bash", str(skill_root / "scripts/resolve_skill_root.sh")],
        skill_root,
        expected_root,
        "installed skill root resolver",
    )
    require_command_output(
        ["bash", str(skill_root / "scripts/zhulong_audit.sh"), "--print-skill-root"],
        skill_root,
        expected_root,
        "installed skill launcher root print",
    )

    forbid_text(
        skill_root / "SKILL.md",
        "Use this Claude Code skill when",
        "installed skill local-agent-neutral opening",
    )
    forbid_text(
        skill_root / "scripts/find_variant_candidates.py",
        "import subprocess",
        "installed variant candidate finder must not shell out",
    )
    forbid_text(
        skill_root / "scripts/find_variant_candidates.py",
        "subprocess.run",
        "installed variant candidate finder must not run subprocesses",
    )
    require_text(
        skill_root / "assets/references/unverified-lead-template.md",
        "Material blocker?",
        "installed skill unverified lead materiality template",
    )
    require_text(
        skill_root / "assets/references/false-positive-template.md",
        "expected_behavior",
        "installed skill false-positive expected behavior reason code",
    )
    require_text(
        skill_root / "assets/references/false-positive-template.md",
        "outside_security_boundary",
        "installed skill false-positive outside boundary reason code",
    )
    require_text(
        skill_root / "assets/references/false-positive-template.md",
        "requires_non_default_admin_trust",
        "installed skill false-positive admin trust reason code",
    )
    require_text(
        skill_root / "assets/references/unverified-lead-template.md",
        "Security policy / scope checked",
        "installed skill unverified security policy check field",
    )

    operator_local_path = "/" + "Users" + "/" + "localuser"
    require_no_repo_text(skill_root, operator_local_path, "operator-local absolute path")
    stale_asr_name = "autonomous-security" + "-researcher"
    require_no_repo_text(skill_root, stale_asr_name, "stale ASR naming")

    with tempfile.TemporaryDirectory(prefix="zhulong-installed-omc-selftest-") as tempdir:
        repo_dir = Path(tempdir) / "repo"
        repo_dir.mkdir(parents=True, exist_ok=True)
        run([
            "bash",
            str(skill_root / "scripts/bootstrap_verification_workspace.sh"),
            "--target-dir",
            str(repo_dir),
            "--workspace-name",
            "security-research-installed-omc",
        ], skill_root)
        exercise_omc_runtime_hygiene(
            skill_root / "scripts/check_omc_runtime.sh",
            repo_dir / "security-research-installed-omc",
            skill_root,
        )
        installed_workspace = repo_dir / "security-research-installed-omc"
        exercise_workspace_tool_registry_snapshot(skill_root, installed_workspace)
        exercise_sandbox_preflight(
            skill_root / "scripts/check_sandbox_preflight.py",
            installed_workspace,
            skill_root,
        )
        exercise_runner_sandbox_rejection(
            skill_root / "scripts/run_verification_case.sh",
            installed_workspace,
            skill_root,
        )
        exercise_verification_wrapper_state_boundary(skill_root, Path(tempdir))
        exercise_sandbox_ledger_guard(installed_workspace, skill_root)
        exercise_structured_blocker_cli(skill_root)

    print(f"SELFTEST PASSED: installed skill layout {skill_root}")


def exercise_candidate_identity_dedup(plugin_root: Path) -> None:
    validator = plugin_root / "scripts/validate_candidate.py"
    upgrader = plugin_root / "scripts/upgrade_candidate_identity.py"
    builder = plugin_root / "scripts/build_candidate_dedup_plan.py"
    plan_validator = plugin_root / "scripts/validate_candidate_dedup_plan.py"
    manifest = json.loads(
        (plugin_root / "assets/fixtures/candidate-identity/manifest.json").read_text(encoding="utf-8")
    )
    manifest_cases = manifest.get("positive_cases", []) + manifest.get("negative_cases", [])
    if not manifest_cases or any(not isinstance(case, dict) for case in manifest_cases):
        raise SystemExit("FAILED: candidate identity manifest cases must be structured objects")
    declared_case_ids = {str(case.get("id") or "") for case in manifest_cases}
    if "" in declared_case_ids or len(declared_case_ids) != len(manifest_cases):
        raise SystemExit("FAILED: candidate identity manifest case IDs must be non-empty and unique")
    for case in manifest_cases:
        if not isinstance(case.get("base"), str) or not isinstance(case.get("mutation"), str) or not isinstance(case.get("expected"), dict):
            raise SystemExit(f"FAILED: candidate identity manifest case is incomplete: {case.get('id')}")
    executed_case_ids: set[str] = set()

    def covered(case_id: str) -> None:
        if case_id not in declared_case_ids:
            raise SystemExit(f"FAILED: candidate identity selftest executed undeclared case: {case_id}")
        if case_id in executed_case_ids:
            raise SystemExit(f"FAILED: candidate identity selftest executed a case twice: {case_id}")
        executed_case_ids.add(case_id)

    def sha(path: Path) -> str:
        return "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()

    with tempfile.TemporaryDirectory(prefix="zhulong-candidate-identity-") as tempdir:
        root = Path(tempdir); repo = root / "repo"; workspace = root / "workspace"
        (repo / "src").mkdir(parents=True); (workspace / "evidence").mkdir(parents=True)
        (workspace / "candidates/CAND-0001").mkdir(parents=True)
        (workspace / "candidates/CAND-0002").mkdir(parents=True)
        (workspace / "verifier/CAND-0001").mkdir(parents=True)
        (repo / "src/importer.py").write_text("def import_url(value):\n    return fetch(value)\n", encoding="utf-8")
        for name in ("agent-a.json", "agent-b.json", "manual-c.json"):
            (workspace / "evidence" / name).write_text(json.dumps({"source": name}, sort_keys=True) + "\n", encoding="utf-8")

        candidate_a = valid_candidate_contract()
        path_a_r1 = write_json_fixture(workspace / "candidates/CAND-0001/candidate-r1.json", candidate_a)
        legacy = json.loads(run_capture([sys.executable, str(validator), str(path_a_r1), "--json"], plugin_root))
        if legacy.get("protocol_mode") != "legacy_r1":
            raise SystemExit("FAILED: legal R1 candidate did not expose legacy_r1 protocol mode")
        covered("legacy-r1-visible-mode")

        def identity_input(artifact: str, **overrides: object) -> dict:
            value = {
                "schema_version": 1, "target_commit": "git-sha-or-local-state", "trust_boundary_id": "tenant-api",
                "sink_family": "http_request", "root_cause_family": "missing_validation",
                "primary_source_path": "src/importer.py",
                "provenance": [{"source_kind": "agent", "source_id": artifact.removesuffix(".json"),
                    "artifact_path": f"evidence/{artifact}", "artifact_sha256": sha(workspace / "evidence" / artifact),
                    "observed_at": "2026-07-22T00:00:00Z", "producer": {"name": "fixture-agent", "version": "1"}}],
            }
            value.update(overrides); return value

        input_a = write_json_fixture(workspace / "identity-a.json", identity_input("agent-a.json"))
        path_a = workspace / "candidates/CAND-0001/candidate-r2.json"
        run([sys.executable, str(upgrader), "--candidate", str(path_a_r1), "--repo-root", str(repo), "--identity-input", str(input_a), "--output", str(path_a), "--json"], plugin_root)
        covered("explicit-r1-to-r2-upgrade")
        first_a = path_a.read_bytes()
        run([sys.executable, str(upgrader), "--candidate", str(path_a_r1), "--repo-root", str(repo), "--identity-input", str(input_a), "--output", str(path_a), "--json"], plugin_root)
        if path_a.read_bytes() != first_a:
            raise SystemExit("FAILED: same-byte Candidate R2 upgrade retry was not idempotent")
        covered("same-output-idempotence")
        checked_a = json.loads(run_capture([sys.executable, str(validator), str(path_a), "--repo-root", str(repo), "--json"], plugin_root))
        if checked_a.get("protocol_mode") != "r2" or not str(checked_a.get("fingerprint", "")).startswith("sha256:"):
            raise SystemExit("FAILED: upgraded Candidate R2 did not expose a recomputed fingerprint")
        upgraded_a = json.loads(path_a.read_text(encoding="utf-8"))
        if upgraded_a["relationships"]["legacy_id_mapping"][0]["legacy_candidate_id"] != "CAND-0001":
            raise SystemExit("FAILED: explicit upgrade did not preserve the legacy candidate ID")
        covered("legacy-id-preserved")

        verdict_doc = valid_verifier_verdict()
        verdict_doc["candidate_binding"] = {"protocol_mode": "r2", "candidate_sha256": sha(path_a), "fingerprint": checked_a["fingerprint"]}
        verdict_path = write_json_fixture(workspace / "verifier/CAND-0001/verifier-verdict.json", verdict_doc)
        run([sys.executable, str(plugin_root / "scripts/validate_verifier_verdict.py"), "--candidate", str(path_a), str(verdict_path)], plugin_root)
        covered("verdict-r2-digest-fingerprint-binding")
        run(
            [
                sys.executable,
                str(plugin_root / "scripts/audit_disposition.py"),
                "--workspace-dir",
                str(workspace),
                "--candidate",
                str(path_a.relative_to(workspace)),
                "--verdict",
                str(verdict_path.relative_to(workspace)),
                "--update-from-verdict",
                "--json",
            ],
            plugin_root,
        )
        ledger = json.loads((workspace / "audit-disposition.json").read_text(encoding="utf-8"))
        record = ledger["candidate_dispositions"][0]
        if record.get("candidate_sha256") != sha(path_a) or record.get("candidate_fingerprint") != checked_a["fingerprint"]:
            raise SystemExit("FAILED: disposition omitted Candidate R2 digest/fingerprint binding")
        drifted_a = json_clone(upgraded_a); drifted_a["title"] = "Digest drift without identity-component drift"
        write_json_fixture(path_a, drifted_a)
        run_expect_fail(
            [sys.executable, str(plugin_root / "scripts/audit_disposition.py"), "--workspace-dir", str(workspace), "--json"],
            plugin_root,
            "title does not match candidate.json",
        )
        covered("disposition-r2-drift-binding")
        path_a.write_bytes(first_a)

        candidate_b = json_clone(candidate_a); candidate_b["candidate_id"] = "CAND-0002"
        candidate_b["title"] = "Same claim with a producer-specific title"
        candidate_b["entrypoint"] = {"id": "IMPORT-URL", "kind": "HTTP", "route": "post //api/import/"}
        candidate_b["evidence"]["static_locations"][0]["start_line"] = 7
        candidate_b["evidence"]["static_locations"][0]["end_line"] = 9
        path_b_r1 = write_json_fixture(workspace / "candidates/CAND-0002/candidate-r1.json", candidate_b)
        legacy_b = json.loads(run_capture([sys.executable, str(validator), str(path_b_r1), "--json"], plugin_root))
        input_b = write_json_fixture(workspace / "identity-b.json", identity_input("agent-b.json"))
        path_b = workspace / "candidates/CAND-0002/candidate-r2.json"
        run([sys.executable, str(upgrader), "--candidate", str(path_b_r1), "--repo-root", str(repo), "--identity-input", str(input_b), "--output", str(path_b), "--json"], plugin_root)
        checked_b = json.loads(run_capture([sys.executable, str(validator), str(path_b), "--repo-root", str(repo), "--json"], plugin_root))
        if checked_a["fingerprint"] != checked_b["fingerprint"]:
            raise SystemExit("FAILED: equivalent entrypoint/title/line/provenance variants changed the fingerprint")
        covered("semantic-entrypoint-normalization")
        covered("nonidentity-fields-independent")

        inventory_path = workspace / "candidate-inventory.json"
        def item(path: Path, checked: dict) -> dict:
            document = json.loads(path.read_text(encoding="utf-8"))
            return {"path": path.relative_to(workspace).as_posix(), "sha256": sha(path), "candidate_id": document["candidate_id"],
                "fingerprint": checked.get("fingerprint"), "target_tested_ref": document["target_ref"]["tested_ref"]}
        inventory = {"schema_version": 1, "inventory_id": "INVENTORY-SELFTEST", "candidates": [item(path_a, checked_a), item(path_b, checked_b)]}
        write_json_fixture(inventory_path, inventory)
        plan_one = workspace / "dedup-plan-one.json"
        run([sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace), "--inventory", str(inventory_path), "--output", str(plan_one), "--json"], plugin_root)
        run([sys.executable, str(plan_validator), "--repo-root", str(repo), "--workspace-dir", str(workspace), "--plan", str(plan_one), "--json"], plugin_root)
        plan = json.loads(plan_one.read_text(encoding="utf-8"))
        if plan["classifications"][0]["classification"] != "exact_duplicate" or plan["exact_groups"][0]["canonical_candidate_id"] != "CAND-0001" or len(plan["exact_groups"][0]["merged_provenance"]) != 2:
            raise SystemExit("FAILED: exact duplicate canonical winner or provenance union is incorrect")
        covered("exact-duplicate-canonical-winner")
        covered("complete-provenance-union-preview")
        inventory["candidates"].reverse(); write_json_fixture(inventory_path, inventory)
        plan_two = workspace / "dedup-plan-two.json"
        run([sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace), "--inventory", str(inventory_path), "--output", str(plan_two), "--json"], plugin_root)
        if plan_one.read_bytes() != plan_two.read_bytes():
            raise SystemExit("FAILED: inventory order changed deterministic plan bytes")

        mixed = {"schema_version": 1, "inventory_id": "INVENTORY-MIXED", "candidates": [item(path_b_r1, legacy_b), item(path_a, checked_a)]}
        write_json_fixture(inventory_path, mixed); mixed_plan = workspace / "dedup-plan-mixed.json"
        run([sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace), "--inventory", str(inventory_path), "--output", str(mixed_plan), "--json"], plugin_root)
        if json.loads(mixed_plan.read_text(encoding="utf-8"))["classifications"][0]["classification"] != "review_required":
            raise SystemExit("FAILED: mixed R1/R2 inventory was not classified conservatively")
        covered("mixed-r1-r2-conservative")

        # Repeat the same semantic inventory under environment and filesystem
        # variation. The plan binds the inventory path, so reuse that path and
        # vary only order, mtimes, locale, timezone, and hash seed.
        inventory["candidates"].reverse()
        write_json_fixture(inventory_path, inventory)
        os.utime(path_a, (1, 1)); os.utime(path_b, (2, 2))
        plan_env = workspace / "dedup-plan-env.json"
        run_capture_with_env(
            [sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace),
             "--inventory", str(inventory_path), "--output", str(plan_env), "--json"],
            plugin_root,
            {"LC_ALL": "C", "TZ": "Pacific/Kiritimati", "PYTHONHASHSEED": "731"},
        )
        if plan_one.read_bytes() != plan_env.read_bytes():
            raise SystemExit("FAILED: locale/timezone/hash-seed/mtime changed deterministic plan bytes")
        covered("environment-and-order-determinism")

        (repo / "src/other.py").write_text("def other(value):\n    return value\n", encoding="utf-8")

        def upgrade_variant(
            label: str,
            candidate_id: str,
            *,
            candidate_changes: dict[str, object] | None = None,
            identity_changes: dict[str, object] | None = None,
        ) -> tuple[Path, dict, dict]:
            candidate_doc = json_clone(candidate_a)
            candidate_doc["candidate_id"] = candidate_id
            for key, value in (candidate_changes or {}).items():
                candidate_doc[key] = value
            case_dir = workspace / "cases" / label
            case_dir.mkdir(parents=True, exist_ok=True)
            r1_path = write_json_fixture(case_dir / "candidate-r1.json", candidate_doc)
            identity_doc = identity_input("manual-c.json")
            identity_doc.update(identity_changes or {})
            identity_path = write_json_fixture(case_dir / "identity-input.json", identity_doc)
            r2_path = case_dir / "candidate-r2.json"
            run([sys.executable, str(upgrader), "--candidate", str(r1_path), "--repo-root", str(repo),
                 "--identity-input", str(identity_path), "--output", str(r2_path), "--json"], plugin_root)
            checked = json.loads(run_capture(
                [sys.executable, str(validator), str(r2_path), "--repo-root", str(repo), "--json"], plugin_root
            ))
            return r2_path, checked, json.loads(r2_path.read_text(encoding="utf-8"))

        def plan_for(label: str, members: list[tuple[Path, dict]]) -> tuple[Path, dict]:
            local_inventory = {
                "schema_version": 1,
                "inventory_id": "INVENTORY-" + label.upper().replace("_", "-"),
                "candidates": [item(path, checked) for path, checked in members],
            }
            local_inventory_path = write_json_fixture(workspace / f"inventory-{label}.json", local_inventory)
            output = workspace / f"dedup-plan-{label}.json"
            run([sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace),
                 "--inventory", str(local_inventory_path), "--output", str(output), "--json"], plugin_root)
            return output, json.loads(output.read_text(encoding="utf-8"))

        def expect_classification(label: str, member: tuple[Path, dict], expected: str, case_id: str) -> None:
            _output, local_plan = plan_for(label, [(path_a, checked_a), member])
            actual = local_plan["classifications"][0]["classification"]
            if actual != expected:
                raise SystemExit(f"FAILED: {case_id} expected {expected}, got {actual}")
            covered(case_id)

        commit_candidate = json_clone(candidate_a)
        commit_candidate["target_ref"] = {**candidate_a["target_ref"], "tested_ref": "different-tested-ref"}
        commit_path, commit_checked, _commit_doc = upgrade_variant(
            "commit-diff", "CAND-COMMIT", candidate_changes={"target_ref": commit_candidate["target_ref"]},
            identity_changes={"target_commit": "different-tested-ref"},
        )
        expect_classification("commit-diff", (commit_path, commit_checked), "distinct", "target-commit-diff")

        entry_path, entry_checked, _entry_doc = upgrade_variant(
            "entry-diff", "CAND-ENTRY",
            candidate_changes={"entrypoint": {"id": "other-entry", "kind": "http", "route": "POST /api/other"}},
        )
        expect_classification("entry-diff", (entry_path, entry_checked), "review_required", "entrypoint-diff")

        boundary_path, boundary_checked, _boundary_doc = upgrade_variant(
            "boundary-diff", "CAND-BOUNDARY", identity_changes={"trust_boundary_id": "admin-api"},
        )
        expect_classification("boundary-diff", (boundary_path, boundary_checked), "review_required", "trust-boundary-diff")

        sink_path, sink_checked, _sink_doc = upgrade_variant(
            "sink-diff", "CAND-SINK", identity_changes={"sink_family": "file_write"},
        )
        expect_classification("sink-diff", (sink_path, sink_checked), "review_required", "sink-family-diff")

        cause_path, cause_checked, _cause_doc = upgrade_variant(
            "cause-diff", "CAND-CAUSE", identity_changes={"root_cause_family": "insufficient_validation"},
        )
        expect_classification("cause-diff", (cause_path, cause_checked), "review_required", "root-cause-diff")
        covered("r2-partial-review-required")
        covered("fuzzy-no-auto-merge")

        source_path, source_checked, _source_doc = upgrade_variant(
            "source-diff", "CAND-SOURCE", identity_changes={"primary_source_path": "src/other.py"},
        )
        expect_classification("source-diff", (source_path, source_checked), "review_required", "primary-source-path-diff")

        distinct_path, distinct_checked, distinct_doc = upgrade_variant(
            "structured-distinct", "CAND-DISTINCT",
            candidate_changes={"entrypoint": {"id": "different-entry", "kind": "cli", "route": "import other"}},
            identity_changes={"trust_boundary_id": "local-cli", "sink_family": "file_write",
                              "root_cause_family": "authorization_missing", "primary_source_path": "src/other.py"},
        )
        expect_classification("structured-distinct", (distinct_path, distinct_checked), "distinct", "structured-distinct")
        distinct_doc["title"] = candidate_a["title"]
        distinct_doc["bug_class"] = candidate_a["bug_class"]
        distinct_doc["claim"] = json_clone(candidate_a["claim"])
        write_json_fixture(distinct_path, distinct_doc)
        distinct_checked = json.loads(run_capture(
            [sys.executable, str(validator), str(distinct_path), "--repo-root", str(repo), "--json"], plugin_root
        ))
        expect_classification("text-only", (distinct_path, distinct_checked), "distinct", "text-only-no-merge")

        def checked_for(path: Path) -> dict:
            return json.loads(run_capture(
                [sys.executable, str(validator), str(path), "--repo-root", str(repo), "--json"], plugin_root
            ))

        def relationship_ref(path: Path, checked: dict, *, digest: str | None = None) -> dict:
            document = json.loads(path.read_text(encoding="utf-8"))
            return {"candidate_id": document["candidate_id"], "fingerprint": checked["fingerprint"],
                    "path": path.relative_to(workspace).as_posix(), "sha256": digest or sha(path)}

        def write_relationship_copy(label: str, candidate_id: str, source: dict | None = None) -> tuple[Path, dict]:
            document = json_clone(source or upgraded_a)
            document["candidate_id"] = candidate_id
            document["relationships"] = {
                "duplicate_of": None,
                "legacy_id_mapping": [{"legacy_candidate_id": candidate_id, "current_candidate_id": candidate_id}],
                "merged_from": [],
            }
            case_dir = workspace / "relationship-cases" / label
            case_dir.mkdir(parents=True, exist_ok=True)
            path = write_json_fixture(case_dir / "candidate.json", document)
            return path, checked_for(path)

        def builder_expect_fail(label: str, members: list[tuple[Path, dict]], expected: str, case_id: str) -> None:
            local_inventory = {"schema_version": 1, "inventory_id": "INVENTORY-REL-" + label.upper(),
                               "candidates": [item(path, checked) for path, checked in members]}
            local_inventory_path = write_json_fixture(workspace / f"inventory-rel-{label}.json", local_inventory)
            run_expect_fail(
                [sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace),
                 "--inventory", str(local_inventory_path), "--output", str(workspace / f"plan-rel-{label}.json"), "--json"],
                plugin_root, expected,
            )
            covered(case_id)

        unknown_path, _unknown_checked = write_relationship_copy("unknown", "CAND-REL-UNKNOWN")
        unknown_doc = json.loads(unknown_path.read_text(encoding="utf-8"))
        unknown_doc["relationships"]["duplicate_of"] = {
            "candidate_id": "CAND-MISSING", "fingerprint": checked_a["fingerprint"],
            "path": "candidates/CAND-MISSING/candidate.json", "sha256": "sha256:" + "1" * 64,
        }
        write_json_fixture(unknown_path, unknown_doc)
        builder_expect_fail("unknown", [(unknown_path, checked_for(unknown_path)), (path_b, checked_b)],
                            "outside the explicit inventory", "unknown-relationship-target")

        cycle_members: list[tuple[Path, dict]] = []
        for suffix in ("A", "B", "C"):
            cycle_members.append(write_relationship_copy(f"cycle-{suffix.lower()}", f"CAND-CYCLE-{suffix}"))
        for index, (cycle_path, cycle_checked) in enumerate(cycle_members):
            next_path, next_checked = cycle_members[(index + 1) % len(cycle_members)]
            cycle_doc = json.loads(cycle_path.read_text(encoding="utf-8"))
            cycle_doc["relationships"]["duplicate_of"] = relationship_ref(next_path, next_checked)
            write_json_fixture(cycle_path, cycle_doc)
        cycle_members = [(path, checked_for(path)) for path, _checked in cycle_members]
        builder_expect_fail("cycle", cycle_members, "relationship cycle", "duplicate-cycle")

        bidi_a, bidi_a_checked = write_relationship_copy("bidi-a", "CAND-BIDI-A")
        bidi_b, bidi_b_checked = write_relationship_copy("bidi-b", "CAND-BIDI-B")
        bidi_a_doc = json.loads(bidi_a.read_text(encoding="utf-8")); bidi_b_doc = json.loads(bidi_b.read_text(encoding="utf-8"))
        bidi_a_doc["relationships"]["duplicate_of"] = relationship_ref(bidi_b, bidi_b_checked)
        bidi_b_doc["relationships"]["duplicate_of"] = relationship_ref(bidi_a, bidi_a_checked)
        write_json_fixture(bidi_a, bidi_a_doc); write_json_fixture(bidi_b, bidi_b_doc)
        builder_expect_fail("bidi", [(bidi_a, checked_for(bidi_a)), (bidi_b, checked_for(bidi_b))],
                            "bidirectional", "bidirectional-duplicate")

        chain_c, chain_c_checked = write_relationship_copy("chain-c", "CAND-CHAIN-C")
        chain_b, _chain_b_checked = write_relationship_copy("chain-b", "CAND-CHAIN-B")
        chain_b_doc = json.loads(chain_b.read_text(encoding="utf-8"))
        chain_b_doc["relationships"]["duplicate_of"] = relationship_ref(chain_c, chain_c_checked)
        write_json_fixture(chain_b, chain_b_doc); chain_b_checked = checked_for(chain_b)
        chain_a, _chain_a_checked = write_relationship_copy("chain-a", "CAND-CHAIN-A")
        chain_a_doc = json.loads(chain_a.read_text(encoding="utf-8"))
        chain_a_doc["relationships"]["duplicate_of"] = relationship_ref(chain_b, chain_b_checked)
        write_json_fixture(chain_a, chain_a_doc)
        builder_expect_fail("chain", [(chain_a, checked_for(chain_a)), (chain_b, chain_b_checked), (chain_c, chain_c_checked)],
                            "must not itself be subordinate", "duplicate-chain-canonical-subordinate")

        loss_path, _loss_checked = write_relationship_copy("provenance-loss", "CAND-PROVENANCE-LOSS")
        loss_doc = json.loads(loss_path.read_text(encoding="utf-8"))
        loss_doc["relationships"]["merged_from"] = [relationship_ref(path_b, checked_b)]
        write_json_fixture(loss_path, loss_doc)
        builder_expect_fail("provenance", [(loss_path, checked_for(loss_path)), (path_b, checked_b)],
                            "provenance union is incomplete", "provenance-loss")

        binding_path, _binding_checked = write_relationship_copy("binding-drift", "CAND-BINDING-DRIFT")
        binding_doc = json.loads(binding_path.read_text(encoding="utf-8"))
        binding_doc["relationships"]["duplicate_of"] = relationship_ref(path_b, checked_b, digest="sha256:" + "2" * 64)
        write_json_fixture(binding_path, binding_doc)
        builder_expect_fail("binding", [(binding_path, checked_for(binding_path)), (path_b, checked_b)],
                            "binding drift", "relationship-binding-drift")

        forged = json_clone(upgraded_a); forged["identity"]["fingerprint"] = "sha256:" + "0" * 64
        run_expect_fail([sys.executable, str(validator), str(write_json_fixture(workspace / "forged.json", forged)), "--json"], plugin_root, "fingerprint")
        covered("forged-fingerprint")
        unknown = json_clone(upgraded_a); unknown["schema_version"] = 99
        run_expect_fail([sys.executable, str(validator), str(write_json_fixture(workspace / "unknown-version.json", unknown)), "--json"], plugin_root, "schema_version")
        covered("unknown-version")
        self_ref = json_clone(upgraded_a); self_ref["relationships"]["duplicate_of"] = {"candidate_id": "CAND-0001", "fingerprint": checked_a["fingerprint"], "path": path_a.relative_to(workspace).as_posix(), "sha256": sha(path_a)}
        run_expect_fail([sys.executable, str(validator), str(write_json_fixture(workspace / "self-ref.json", self_ref)), "--json"], plugin_root, "self-reference")
        covered("self-reference")

        silent_r2 = json_clone(candidate_a)
        silent_r2.update({"identity": upgraded_a["identity"], "provenance": upgraded_a["provenance"],
                          "relationships": upgraded_a["relationships"]})
        run_expect_fail([sys.executable, str(validator), str(write_json_fixture(workspace / "silent-r2.json", silent_r2)), "--json"],
                        plugin_root, "unsupported field")
        covered("legacy-silent-r2")

        confirmed_like = json_clone(upgraded_a); confirmed_like["title"] = "confirmed vulnerability from candidate metadata"
        run_expect_fail([sys.executable, str(validator), str(write_json_fixture(workspace / "confirmed-like.json", confirmed_like)), "--json"],
                        plugin_root, "confirmed-like")
        covered("confirmed-authority-injection")

        def upgrade_must_reject(label: str, identity_doc: dict) -> None:
            (workspace / "rejected-inputs").mkdir(parents=True, exist_ok=True)
            input_path = write_json_fixture(workspace / "rejected-inputs" / f"{label}.json", identity_doc)
            proc = subprocess.run(
                [sys.executable, str(upgrader), "--candidate", str(path_a_r1), "--repo-root", str(repo),
                 "--identity-input", str(input_path), "--output", str(workspace / "rejected-outputs" / f"{label}.json"), "--json"],
                cwd=plugin_root, capture_output=True, text=True,
            )
            if proc.returncode == 0:
                raise SystemExit(f"FAILED: unsafe/nonportable identity input unexpectedly succeeded: {label}")

        for label, unsafe_path in (
            ("absolute", "/private/fixture.py"),
            ("uri", "file://fixture.py"),
            ("traversal", "../src/importer.py"),
            ("backslash", "src\\importer.py"),
            ("tilde", "~/src/importer.py"),
        ):
            upgrade_must_reject(label, identity_input("agent-a.json", primary_source_path=unsafe_path))
        outside_source = root / "outside.py"; outside_source.write_text("outside\n", encoding="utf-8")
        symlink_source = repo / "src/link.py"; symlink_source.symlink_to(outside_source)
        upgrade_must_reject("symlink", identity_input("agent-a.json", primary_source_path="src/link.py"))
        covered("unsafe-path-matrix")

        nonportable_inputs = [
            identity_input("agent-a.json", primary_source_path="/Users/operator/private.py"),
            identity_input("agent-a.json"),
            identity_input("agent-a.json"),
            identity_input("agent-a.json"),
        ]
        nonportable_inputs[1]["provenance"][0]["source_id"] = "system_prompt"
        nonportable_inputs[2]["provenance"][0]["producer"]["version"] = "hidden reasoning"
        nonportable_inputs[3]["provenance"][0]["producer"]["version"] = "token=fixture-value"
        for index, identity_doc in enumerate(nonportable_inputs):
            upgrade_must_reject(f"nonportable-{index}", identity_doc)
        covered("nonportable-input-matrix")

        run_expect_fail([sys.executable, str(upgrader), "--candidate", str(path_a_r1), "--repo-root", str(repo), "--identity-input", str(input_a), "--output", str(path_a_r1), "--json"], plugin_root, "in-place")
        covered("in-place-output")
        conflict = workspace / "conflict.json"; conflict.write_text("{}\n", encoding="utf-8")
        run_expect_fail([sys.executable, str(upgrader), "--candidate", str(path_a_r1), "--repo-root", str(repo), "--identity-input", str(input_a), "--output", str(conflict), "--json"], plugin_root, "different bytes")
        covered("existing-output-conflict")

        drift_inventory = {"schema_version": 1, "inventory_id": "INVENTORY-DRIFT", "candidates": [item(path_a, checked_a), item(path_b, checked_b)]}
        drift_inventory["candidates"][0]["sha256"] = "sha256:" + "3" * 64
        drift_inventory_path = write_json_fixture(workspace / "inventory-drift.json", drift_inventory)
        run_expect_fail(
            [sys.executable, str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace),
             "--inventory", str(drift_inventory_path), "--output", str(workspace / "plan-drift.json"), "--json"],
            plugin_root, "candidate digest drift",
        )
        covered("candidate-input-drift")

        # Deterministically fault the second production load used immediately
        # before publication. This exercises the CLI main path without a flaky
        # timing race and proves that no output is published after input drift.
        scripts_dir = str(plugin_root / "scripts")
        inserted_scripts_dir = scripts_dir not in sys.path
        if inserted_scripts_dir:
            sys.path.insert(0, scripts_dir)
        try:
            spec = importlib.util.spec_from_file_location(
                "zhulong_build_candidate_dedup_plan_drift_selftest", builder
            )
            if spec is None or spec.loader is None:
                raise SystemExit("FAILED: could not load production dedup builder for drift recheck")
            builder_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(builder_module)
            stable_inventory, stable_candidates = builder_module.load_inventory(repo, workspace, inventory_path)
            drifted_inventory = dict(stable_inventory)
            drifted_inventory["_sha256"] = "sha256:" + "4" * 64
            drift_output = workspace / "plan-concurrent-drift.json"
            argv = [str(builder), "--repo-root", str(repo), "--workspace-dir", str(workspace),
                    "--inventory", str(inventory_path), "--output", str(drift_output), "--json"]
            with mock.patch.object(
                builder_module, "load_inventory",
                side_effect=[(stable_inventory, stable_candidates), (drifted_inventory, stable_candidates)],
            ), mock.patch.object(builder_module, "atomic_publish") as publish, mock.patch.object(sys, "argv", argv), \
                    mock.patch.object(sys, "stdout", io.StringIO()):
                try:
                    builder_module.main()
                except SystemExit as exc:
                    if exc.code != 1:
                        raise SystemExit("FAILED: concurrent input drift did not fail closed") from exc
                else:
                    raise SystemExit("FAILED: concurrent input drift unexpectedly succeeded")
                publish.assert_not_called()
            if drift_output.exists():
                raise SystemExit("FAILED: concurrent input drift published an output")
        finally:
            if inserted_scripts_dir:
                sys.path.pop(0)
        covered("concurrent-input-drift-recheck")

        tampered = json.loads(plan_one.read_text(encoding="utf-8")); tampered["classifications"][0]["classification"] = "distinct"
        tampered_path = write_json_fixture(workspace / "tampered-plan.json", tampered)
        run_expect_fail([sys.executable, str(plan_validator), "--repo-root", str(repo), "--workspace-dir", str(workspace), "--plan", str(tampered_path), "--json"], plugin_root, "stale, forged")
        covered("plan-recompute-forgery")

        (workspace / "confirmed").mkdir(parents=True, exist_ok=True)
        authority_paths = [
            write_json_fixture(workspace / "audit-events.jsonl", {"sentinel": "journal"}),
            write_json_fixture(workspace / "stage-status.json", {"sentinel": "state"}),
            write_json_fixture(workspace / "recording-evidence.json", {"sentinel": "recording"}),
            write_json_fixture(workspace / "confirmed/sentinel.json", {"sentinel": "bundle"}),
            write_json_fixture(workspace / "evidence/protected-sentinel.json", {"sentinel": "evidence"}),
            path_a,
            verdict_path,
            workspace / "audit-disposition.json",
        ]
        authority_before = {path: path.read_bytes() for path in authority_paths}
        run([sys.executable, str(validator), str(path_a), "--repo-root", str(repo), "--json"], plugin_root)
        run([sys.executable, str(plan_validator), "--repo-root", str(repo), "--workspace-dir", str(workspace),
             "--plan", str(plan_one), "--json"], plugin_root)
        for path, before in authority_before.items():
            if path.read_bytes() != before:
                raise SystemExit(f"FAILED: candidate identity tooling modified authority artifact: {path}")
        covered("authority-artifact-immutability")

        missing_cases = sorted(declared_case_ids - executed_case_ids)
        extra_cases = sorted(executed_case_ids - declared_case_ids)
        if missing_cases or extra_cases:
            raise SystemExit(
                "FAILED: candidate identity manifest/selftest coverage drift; "
                f"missing={missing_cases} extra={extra_cases}"
            )

    print("CANDIDATE IDENTITY/DEDUP SELFTEST PASSED: R1/R2, deterministic, advisory-only, drift-safe")


def main() -> None:
    plugin_root = Path(__file__).resolve().parent.parent

    if (plugin_root / "SKILL.md").exists() and not (plugin_root / ".codex-plugin/plugin.json").exists():
        selftest_installed_skill(plugin_root)
        return

    for rel in REQUIRED_FILES:
        path = plugin_root / rel
        if not path.exists():
            raise SystemExit(f"FAILED: missing required plugin file: {path}")

    exercise_agents_shim(plugin_root)
    exercise_p7_wording_closure(plugin_root)
    exercise_p8_closure_contracts(plugin_root)
    exercise_replay_transcript_corpus(plugin_root)
    exercise_p8_dogfood_metrics(plugin_root)
    exercise_p8_real_historical_dogfood(plugin_root)
    exercise_p9_protocol_chain_real_workspace_dogfood(plugin_root)
    exercise_target_contract_validator(plugin_root)
    exercise_tool_registry_contract(plugin_root)
    exercise_context_planning_contract(plugin_root)
    exercise_root_skill_kernel_contract(plugin_root)
    exercise_recon_result_contract(plugin_root)
    exercise_triage_batch_and_stage_finalization(plugin_root)
    exercise_finding_contract_validators(plugin_root)
    exercise_candidate_identity_dedup(plugin_root)
    exercise_bundle_contract_validator(plugin_root)
    exercise_build_confirmed_bundle_wrapper(plugin_root)
    exercise_independent_verifier(plugin_root)
    exercise_disposition_integration(plugin_root)
    exercise_contract_fixture_chain(plugin_root)
    exercise_audit_state_protocol_r2(plugin_root)
    exercise_audit_state_protocol_closure(plugin_root)
    exercise_handoff_checkpoint_contract(plugin_root)
    exercise_next_actions_contract(plugin_root)
    run([sys.executable, str(plugin_root / "scripts/selftest_audit_timeline.py")], plugin_root)
    exercise_audit_state_writer(plugin_root)
    exercise_audit_state_recovery(plugin_root)
    exercise_audit_transition_policy(plugin_root)
    exercise_recording_evidence_gate(plugin_root)
    exercise_structured_blocker_cli(plugin_root)

    plugin_json = json.loads((plugin_root / ".codex-plugin/plugin.json").read_text(encoding="utf-8"))
    if plugin_json.get("name") != "zhulong":
        raise SystemExit("FAILED: plugin.json name mismatch")
    validate_claude_plugin_manifest(plugin_root)

    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "must never be written under `confirmed/`",
        "false-positive template confirmed-output guardrail",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "Docker verification status",
        "false-positive template Docker status field",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "expected_behavior",
        "false-positive template expected behavior reason code",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "outside_security_boundary",
        "false-positive template outside boundary reason code",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "requires_non_default_admin_trust",
        "false-positive template admin trust reason code",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "default_config_not_vulnerable",
        "false-positive template default config reason code",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "insufficient_attacker_condition",
        "false-positive template attacker-condition reason code",
    )
    require_text(
        plugin_root / "assets/references/false-positive-template.md",
        "insufficient_security_impact",
        "false-positive template security-impact reason code",
    )
    require_text(
        plugin_root / "assets/references/unverified-lead-template.md",
        "Safe resume step",
        "unverified lead template resume field",
    )
    require_text(
        plugin_root / "assets/references/unverified-lead-template.md",
        "high-confidence-unverified/",
        "unverified lead template high-confidence guardrail",
    )
    require_text(
        plugin_root / "assets/references/unverified-lead-template.md",
        "Why completion is still safe?",
        "unverified lead template materiality rationale field",
    )
    require_text(
        plugin_root / "assets/references/unverified-lead-template.md",
        "Confirmed-output guardrail",
        "unverified lead template confirmed-output guardrail field",
    )
    require_text(
        plugin_root / "assets/references/unverified-lead-template.md",
        "Security policy / scope checked",
        "unverified lead template security policy check field",
    )
    require_text(
        plugin_root / "assets/references/final-summary-template.md",
        "false positives / non-security defects",
        "final summary false-positive section",
    )
    require_text(
        plugin_root / "assets/references/final-summary-template.md",
        "high-confidence-but-not-Docker-confirmed leads",
        "final summary high-confidence unverified section",
    )
    require_text(
        plugin_root / "assets/references/final-summary-template.md",
        "<audit-workspace>/SUMMARY.md",
        "final summary stable workspace summary requirement",
    )
    require_text(
        plugin_root / "assets/references/final-summary-template.md",
        "Docker-confirmed but bundle incomplete",
        "final summary partial bundle section",
    )
    require_text(
        plugin_root / "assets/references/confirmed-vuln-docx-format.md",
        "Claude Code DOCX Editing Rule",
        "confirmed-vuln-docx-format docx workflow section",
    )
    require_text(
        plugin_root / "assets/references/confirmed-vuln-docx-format.md",
        "Verification Evidence JSON",
        "confirmed-vuln-docx-format verification evidence schema",
    )
    require_text(
        plugin_root / "assets/references/confirmed-vuln-docx-format.md",
        "partial confirmed bundle",
        "confirmed-vuln-docx-format partial bundle guardrail",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "partial_confirmed_bundle",
        "handoff renderer partial bundle classification hint",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "Finalization integrity",
        "handoff renderer finalization integrity hint",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "OMC Runtime Hygiene",
        "handoff renderer OMC runtime hygiene section",
    )
    require_text(
        plugin_root / "scripts/check_omc_runtime.sh",
        "--cleanup-suspect-pid",
        "OMC runtime helper exact PID cleanup flag",
    )
    require_text(
        plugin_root / "scripts/check_omc_runtime.sh",
        "--force-kill-suspect-teammates",
        "OMC runtime helper deprecated broad cleanup refusal",
    )
    require_text(
        plugin_root / "scripts/asr_start.sh",
        "--prompt-runtime-pid-review",
        "asr launcher optional runtime PID review prompt flag",
    )
    require_text(
        plugin_root / "scripts/assert_finalized_workspace.py",
        "FINALIZATION INTEGRITY FAILED",
        "finalization integrity checker failure heading",
    )
    require_text(
        plugin_root / "scripts/audit_disposition.py",
        "state=confirmed requires confirmed_bundle_path",
        "audit disposition confirmed bundle gate",
    )
    require_text(
        plugin_root / "scripts/blocked_verification.py",
        "Docker Hub pull rate limit blocked runtime verification",
        "blocked verification Docker Hub recovery guidance",
    )
    require_text(
        plugin_root / "scripts/finalize_audit_workspace.py",
        "Blocked Docker/runtime verification prevents completed_no_confirmed_findings",
        "finalization blocks blocked verification no-confirmed success",
    )
    require_text(
        plugin_root / "scripts/finalize_audit_workspace.py",
        "zhulong_completion_summary_placeholder",
        "finalization writes stable summary placeholder",
    )
    require_text(
        plugin_root / "scripts/manage_docker_resources.py",
        "BuildKit cache blocker",
        "Docker strict BuildKit blocker messaging",
    )
    require_text(
        plugin_root / "assets/references/docker-resource-hygiene.md",
        "cannot be auto-deleted safely",
        "Docker hygiene BuildKit review-only blocker guidance",
    )
    require_text(
        plugin_root / "assets/references/docker-resource-hygiene.md",
        "Registry Fallback Guidance",
        "Docker hygiene registry fallback guidance",
    )
    require_text(
        plugin_root / "assets/references/document-output-stability.md",
        "Claude Code built-in `Documents` skill",
        "document-output-stability Documents skill rule",
    )
    require_text(
        plugin_root / "assets/references/java-web-audit-playbook.md",
        "source-to-sink",
        "Java Web playbook source-to-sink contract",
    )
    require_text(
        plugin_root / "assets/references/java-web-audit-playbook.md",
        "Minimum entry inventory fields",
        "Java Web playbook entry inventory contract",
    )
    require_text(
        plugin_root / "assets/references/java-web-audit-playbook.md",
        "Current Verification Status",
        "Java Web playbook verification status field",
    )
    require_text(
        plugin_root / "assets/references/go-web-audit-playbook.md",
        "source-to-sink",
        "Go Web playbook source-to-sink contract",
    )
    require_text(
        plugin_root / "assets/references/go-web-audit-playbook.md",
        "Minimum entry inventory fields",
        "Go Web playbook entry inventory contract",
    )
    require_text(
        plugin_root / "assets/references/go-web-audit-playbook.md",
        "Current Verification Status",
        "Go Web playbook verification status field",
    )
    node_library_playbook = plugin_root / "assets/references/nodejs-library-audit-playbook.md"
    for expected in (
        "Fast Model",
        "Minimum library inventory fields",
        "Public API / CLI",
        "Input Shape",
        "Caller-Controlled Options",
        "Transformation Path",
        "High-Risk Sink",
        "Consumer Impact Assumption",
        "Current Verification Status",
        "Source-To-Sink Tracing Guidance",
        "Docker-Only Verification Reminders",
        "Package metadata, API matches",
        "cannot confirm a vulnerability by",
        "verification_status=confirmed_in_docker",
        "OWASP Prototype Pollution Prevention Cheat Sheet",
    ):
        require_text(
            node_library_playbook,
            expected,
            f"Node.js Library playbook required text {expected}",
        )
    python_library_playbook = plugin_root / "assets/references/python-library-audit-playbook.md"
    for expected in (
        "Python Library / Framework Audit Playbook",
        "Minimum Python library inventory fields",
        "Public API / Hook",
        "Caller-Controlled Options",
        "Consumer Impact Assumption",
        "Source-To-Sink Tracing Guidance",
        "Docker-Only Verification Reminders",
        "Flask, Werkzeug, Jinja, Click",
        "Do not force a route / method / handler table",
        "cannot confirm a vulnerability by themselves",
        "verification_status=confirmed_in_docker",
    ):
        require_text(
            python_library_playbook,
            expected,
            f"Python Library playbook required text {expected}",
        )
    php_swoole_playbook = plugin_root / "assets/references/php-swoole-audit-playbook.md"
    for expected in (
        "PHP / Swoole Web Audit Playbook",
        "Minimum entry inventory fields",
        "Route / Command / Worker",
        "HTTP-exposed controllers from CLI-only",
        "curl_exec",
        "GraphQL",
        "Docker-Only Verification Reminders",
        "verification_status=confirmed_in_docker",
    ):
        require_text(
            php_swoole_playbook,
            expected,
            f"PHP/Swoole playbook required text {expected}",
        )
    registry_fallback = plugin_root / "assets/references/docker-registry-fallbacks.example.json"
    registry_data = json.loads(registry_fallback.read_text(encoding="utf-8"))
    if registry_data.get("policy", {}).get("configurable_not_hardcoded") is not True:
        raise SystemExit("FAILED: registry fallback example must be configurable, not hardcoded")
    for expected_field in (
        "original_image_ref",
        "attempted_image_ref",
        "registry_source",
        "final_digest",
        "failure_reason",
    ):
        if expected_field not in json.dumps(registry_data, ensure_ascii=False):
            raise SystemExit(f"FAILED: registry fallback example missing field: {expected_field}")
    for playbook in (
        "nodejs-web-audit-playbook.md",
        "python-web-audit-playbook.md",
    ):
        path = plugin_root / "assets/references" / playbook
        for expected in (
            "Fast Model",
            "Minimum entry inventory fields",
            "Route / Endpoint",
            "Method",
            "Handler / Controller",
            "Authentication Requirement",
            "Input Source",
            "Downstream Sink / Service",
            "Current Verification Status",
            "Source-To-Sink Tracing Guidance",
            "Docker-Only Verification Reminders",
            "Reference Sources",
            "cannot confirm a vulnerability by themselves",
            "not exhaustive and must not narrow exploration",
            "Do not generate DOCX reports from playbook hypotheses alone",
            "verification_status=confirmed_in_docker",
            "OWASP Web Security Testing Guide",
        ):
            require_text(
                path,
                expected,
                f"{playbook} required text {expected}",
            )
    require_text(
        plugin_root / "assets/references/python-web-audit-playbook.md",
        "Django REST framework viewsets",
        "Python Web playbook DRF authoritative source",
    )
    require_text(
        plugin_root / "assets/references/recommended-security-tooling.md",
        "Do not paste raw scanner logs into `attack-surface.md`",
        "recommended tooling attack-surface log-dump guardrail",
    )
    for checklist in (
        "ssrf-checklist.md",
        "path-traversal-checklist.md",
        "prototype-pollution-checklist.md",
    ):
        path = plugin_root / "assets/references" / checklist
        for section in (
            "Scope and When To Use It",
            "Common Sources",
            "High-Risk Sinks",
            "Source-To-Sink Tracing Hints",
            "Docker-Only Verification Ideas",
            "Severity-Escalation Evidence To Seek",
            "Common False Positives",
            "Confirmed-Only Routing Reminder",
        ):
            require_text(
                path,
                section,
                f"{checklist} section {section}",
            )
        require_text(
            path,
            "cannot confirm a vulnerability",
            f"{checklist} reasoning-only guardrail",
        )
        require_text(
            path,
            "Do not generate DOCX reports from this checklist alone",
            f"{checklist} no-DOCX guardrail",
        )
        require_text(
            path,
            "verification_status=confirmed_in_docker",
            f"{checklist} Docker-confirmed-only guardrail",
        )
    require_text(
        plugin_root / "assets/references/attacker-container-pattern.md",
        "Verification Runner Contract",
        "attacker container verification runner contract section",
    )
    require_text(
        plugin_root / "assets/references/attacker-container-pattern.md",
        "failed_timeout",
        "attacker container runner timeout label",
    )
    require_text(
        plugin_root / "assets/references/attacker-container-pattern.md",
        "resource limits are managed by the Compose files",
        "attacker container compose resource limit note",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "STABLE_LABELS=\"blocked_state_precondition blocked_authority_event_commit blocked_docker_unavailable blocked_missing_image failed_timeout failed_resource_limit rejected_unsafe_sandbox rejected_not_reproducible confirmed_in_docker\"",
        "verification runner stable labels",
    )
    require_text(
        plugin_root / "scripts/run_initial_probes.sh",
        "PROBE_STATUS_LABELS=\"ran_ok skipped_tool_missing skipped_no_package_sources failed_nonfatal failed_fatal\"",
        "initial probes stable status labels",
    )
    require_text(
        plugin_root / "scripts/run_initial_probes.sh",
        "initial-probes-summary.json",
        "initial probes structured summary filename",
    )
    require_text(
        plugin_root / "scripts/run_initial_probes.sh",
        "No package sources found",
        "initial probes OSV no package sources classifier",
    )
    require_text(
        plugin_root / "scripts/run_initial_probes.sh",
        "--report-format json",
        "initial probes gitleaks JSON report mode",
    )
    require_text(
        plugin_root / "scripts/run_initial_probes.sh",
        "Full Secret and Match values are omitted",
        "initial probes gitleaks secret redaction contract",
    )
    require_text(
        plugin_root / "assets/references/python-web-audit-playbook.md",
        "Werkzeug Debugger / Gunicorn Verification Hint",
        "Python Web playbook Werkzeug debugger section",
    )
    require_text(
        plugin_root / "assets/references/python-web-audit-playbook.md",
        "WEB_CONCURRENCY=1",
        "Python Web playbook Gunicorn single-worker verification hint",
    )
    require_text(
        plugin_root / "assets/references/python-web-audit-playbook.md",
        "Never recommend enabling Flask/Werkzeug debugger",
        "Python Web playbook no production debugger guardrail",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "Heavy Logs To Avoid Unless Needed",
        "handoff renderer heavy-log avoidance heading",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "Confirmed-Only Routing Guardrails",
        "handoff renderer confirmed-only heading",
    )
    require_text(
        plugin_root / "scripts/render_handoff_summary.py",
        "Do not copy raw scanner logs into this handoff",
        "handoff renderer raw-log dump warning",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "--timeout-seconds is required and must be positive",
        "verification runner mandatory timeout contract",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "--memory \"$MEMORY_LIMIT\"",
        "verification runner memory limit",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "--cpus \"$CPU_LIMIT\"",
        "verification runner CPU limit",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "--pids-limit \"$PIDS_LIMIT\"",
        "verification runner pids limit",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "--network \"$NETWORK\"",
        "verification runner explicit network",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "managed_by_compose_file",
        "verification runner compose resource limit reporting",
    )
    forbid_text(
        plugin_root / "scripts/run_verification_case.sh",
        "stable_status_labels",
        "verification runner static labels in result json",
    )
    require_text(
        plugin_root / "scripts/run_verification_case.sh",
        "no host fallback is provided",
        "verification runner no-host-fallback contract",
    )
    forbid_text(
        plugin_root / "scripts/run_verification_case.sh",
        "may execute PoC logic directly on the host",
        "verification runner positive host execution wording",
    )
    operator_local_path = "/" + "Users" + "/" + "localuser"
    require_no_repo_text(plugin_root, operator_local_path, "operator-local absolute path")
    stale_asr_name = "autonomous-security" + "-researcher"
    require_no_repo_text(plugin_root, stale_asr_name, "stale ASR naming")
    require_text(
        plugin_root / "assets/references/claude-code-invocation-template.md",
        "severity-escalation pass",
        "Claude invocation template severity escalation contract",
    )
    require_text(
        plugin_root / "assets/references/claude-code-invocation-template.md",
        "Do not execute `web_search`",
        "Claude invocation template web lookup shell-safety contract",
    )
    require_text(
        plugin_root / "assets/references/claude-code-invocation-template.md",
        "Do not produce a thin report",
        "Claude invocation template report-depth contract",
    )
    require_text(
        plugin_root / "assets/references/claude-code-invocation-template.md",
        "exactly one vulnerability",
        "Claude invocation template one-finding-per-bundle contract",
    )
    require_text(
        plugin_root / "assets/references/variant-seed-template.md",
        "A seed card is auxiliary evidence only",
        "variant seed template auxiliary-evidence boundary",
    )
    require_text(
        plugin_root / "assets/references/variant-seed-template.md",
        "independent Docker or Docker Compose reproduction and confirmed",
        "variant seed template Docker reproduction boundary",
    )
    require_text(
        plugin_root / "assets/references/variant-seed-template.md",
        "records must stay `status=candidate`",
        "variant seed template candidate-output boundary",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.md",
        "`scripts/extract_variant_seed.py` is an offline helper",
        "workflow details seed-card extractor contract",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.md",
        "`scripts/find_variant_candidates.py` reads one final Variant Seed Card",
        "workflow details candidate finder contract",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.zh-CN.md",
        "从已有的确认漏洞包中提取同类漏洞种子卡",
        "Chinese workflow details seed-card extractor contract",
    )
    require_text(
        plugin_root / "docs/WORKFLOW_DETAILS.zh-CN.md",
        "执行 `scripts/find_variant_candidates.py`",
        "Chinese workflow details candidate finder contract",
    )
    forbid_text(
        plugin_root / "scripts/find_variant_candidates.py",
        "import subprocess",
        "variant candidate finder must not shell out",
    )
    forbid_text(
        plugin_root / "scripts/find_variant_candidates.py",
        "subprocess.run",
        "variant candidate finder must not run subprocesses",
    )
    canonical_prompt = plugin_root / "assets" / "references" / "claude-code-invocation-template.md"
    root_prompt = plugin_root.parent.parent / "claude-code-zhulong-prompt-template.md"
    if root_prompt.exists() and canonical_prompt.read_text(encoding="utf-8") != root_prompt.read_text(encoding="utf-8"):
        raise SystemExit(
            "FAILED: root prompt template is out of sync with the canonical plugin invocation template. "
            "Run scripts/sync_to_claude_skill.sh --sync-root-prompt-template or resync the repository prompt copy."
        )

    run([sys.executable, "-m", "py_compile",
         str(plugin_root / "scripts/plan_security_toolchain.py"),
         str(plugin_root / "scripts/validate_tool_registry.py"),
         str(plugin_root / "scripts/workspace_state.py"),
         str(plugin_root / "scripts/render_handoff_summary.py"),
         str(plugin_root / "scripts/next_actions.py"),
         str(plugin_root / "scripts/render_next_actions.py"),
         str(plugin_root / "scripts/validate_next_actions.py"),
         str(plugin_root / "scripts/audit_timeline.py"),
         str(plugin_root / "scripts/render_audit_timeline.py"),
         str(plugin_root / "scripts/validate_audit_timeline.py"),
         str(plugin_root / "scripts/selftest_audit_timeline.py"),
         str(plugin_root / "scripts/assert_finalized_workspace.py"),
         str(plugin_root / "scripts/audit_disposition.py"),
         str(plugin_root / "scripts/blocked_verification.py"),
         str(plugin_root / "scripts/audit_state_io.py"),
         str(plugin_root / "scripts/audit_text_safety.py"),
         str(plugin_root / "scripts/audit_transition_policy.py"),
         str(plugin_root / "scripts/write_audit_event.py"),
         str(plugin_root / "scripts/validate_audit_protocol.py"),
         str(plugin_root / "scripts/validate_workspace_state.py"),
         str(plugin_root / "scripts/validate_target_contract.py"),
         str(plugin_root / "scripts/validate_candidate.py"),
         str(plugin_root / "scripts/validate_verifier_verdict.py"),
         str(plugin_root / "scripts/validate_bundle_contract.py"),
         str(plugin_root / "scripts/build_confirmed_bundle.py"),
         str(plugin_root / "scripts/p8_dogfood_metrics.py"),
         str(plugin_root / "scripts/verify_candidate.py"),
         str(plugin_root / "scripts/check_sandbox_preflight.py"),
         str(plugin_root / "scripts/manage_docker_resources.py"),
         str(plugin_root / "scripts/render_confirmed_vuln_docx.py"),
         str(plugin_root / "scripts/scaffold_bilingual_findings.py"),
         str(plugin_root / "scripts/extract_variant_seed.py"),
         str(plugin_root / "scripts/find_variant_candidates.py"),
         str(plugin_root / "scripts/validate_report_bundle.py"),
         str(plugin_root / "scripts/validate_all_report_bundles.py"),
         str(plugin_root / "scripts/finalize_audit_workspace.py")], plugin_root)

    run(["bash", "-n", str(plugin_root / "scripts/bootstrap_verification_workspace.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/asr_start.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/resolve_skill_root.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/zhulong_audit.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/prepare_target_repo.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/check_docker_gate.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/check_omc_runtime.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/check_security_tooling.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/run_initial_probes.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/run_verification_case.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/refresh_workspace_helpers.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/sync_to_claude_skill.sh")], plugin_root)
    run(["bash", "-n", str(plugin_root / "scripts/sync_to_codex_skill.sh")], plugin_root)
    expected_plugin_root = str(plugin_root.resolve())
    require_command_output(
        ["bash", str(plugin_root / "scripts/resolve_skill_root.sh")],
        plugin_root,
        expected_plugin_root,
        "source skill root resolver",
    )
    require_command_output(
        ["bash", str(plugin_root / "scripts/zhulong_audit.sh"), "--print-skill-root"],
        plugin_root,
        expected_plugin_root,
        "source launcher root print",
    )
    run(["bash", str(plugin_root / "scripts/run_verification_case.sh"), "--help"], plugin_root)
    run([sys.executable, str(plugin_root / "scripts/manage_docker_resources.py"), "--help"], plugin_root)
    run([sys.executable, str(plugin_root / "scripts/render_handoff_summary.py"), "--help"], plugin_root)
    run([sys.executable, str(plugin_root / "scripts/assert_finalized_workspace.py"), "--help"], plugin_root)
    run([sys.executable, str(plugin_root / "scripts/finalize_audit_workspace.py"), "--help"], plugin_root)
    require_text(
        plugin_root / "scripts/manage_docker_resources.py",
        '"image", "ls", "-a", "--no-trunc"',
        "Docker cleanup helper snapshots dangling images",
    )

    with tempfile.TemporaryDirectory(prefix="asr-plugin-selftest-") as tempdir:
        repo_dir = Path(tempdir) / "repo"
        repo_dir.mkdir(parents=True, exist_ok=True)
        for unsafe_name in ("", ".", "..", ".hidden", "/tmp/escape", "../escape", "case/name", r"case\\name", "name with space", "~", "http://example.invalid", "%2e%2e"):
            unsafe_target = Path(tempdir) / ("unsafe-" + str(len(list(Path(tempdir).iterdir()))))
            unsafe_target.mkdir()
            before = sorted(item.name for item in unsafe_target.iterdir())
            proc = subprocess.run(
                [
                    "bash",
                    str(plugin_root / "scripts/bootstrap_verification_workspace.sh"),
                    "--target-dir",
                    str(unsafe_target),
                    "--workspace-name",
                    unsafe_name,
                ],
                cwd=plugin_root,
                capture_output=True,
                text=True,
            )
            if proc.returncode == 0 or sorted(item.name for item in unsafe_target.iterdir()) != before:
                raise SystemExit(f"FAILED: unsafe bootstrap workspace name crossed a side-effect boundary: {unsafe_name!r}")
        workspace_name = "security-research-selftest"
        run([
            "bash",
            str(plugin_root / "scripts/bootstrap_verification_workspace.sh"),
            "--target-dir",
            str(repo_dir),
            "--workspace-name",
            workspace_name,
        ], plugin_root)
        workspace = repo_dir / workspace_name
        if not (workspace / "bin/check_security_tooling.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing check_security_tooling.sh")
        if not (workspace / "bin/check-docker-gate.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing check-docker-gate.sh")
        if not (workspace / "bin/run-initial-probes.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing run-initial-probes.sh")
        if not (workspace / "bin/run-verification-case.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing run-verification-case.sh")
        if not (workspace / "bin/check-sandbox-preflight.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing check-sandbox-preflight.py")
        if not (workspace / "bin/evidence_io.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing evidence_io.py")
        if not (workspace / "bin/manage-docker-resources.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing manage-docker-resources.py")
        if not (workspace / "bin/render-handoff-summary.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing render-handoff-summary.py")
        if not (workspace / "bin/workspace_state.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing workspace_state.py")
        if not (workspace / "bin/audit_text_safety.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing audit_text_safety.py")
        for next_actions_helper in [
            "next_actions.py", "render-next-actions.py", "validate-next-actions.py",
            "next-actions.schema.json",
        ]:
            if not (workspace / "bin" / next_actions_helper).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing {next_actions_helper}")
        for next_actions_adapter in ["render-next-actions.py", "validate-next-actions.py"]:
            if not (workspace / "scripts" / next_actions_adapter).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing scripts/{next_actions_adapter}")
        for timeline_helper in [
            "audit_timeline.py", "render-audit-timeline.py",
            "validate-audit-timeline.py", "audit-timeline.schema.json",
        ]:
            if not (workspace / "bin" / timeline_helper).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing {timeline_helper}")
        for timeline_adapter in ["render-audit-timeline.py", "validate-audit-timeline.py"]:
            if not (workspace / "scripts" / timeline_adapter).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing scripts/{timeline_adapter}")
        if not (workspace / "bin/assert-finalized-workspace.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing assert-finalized-workspace.py")
        if not (workspace / "bin/blocked_verification.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing blocked_verification.py")
        for contract_helper in [
            "validate_target_contract.py",
            "validate_recon_result.py",
            "validate-recon-result.py",
            "validate-triage-batch.py",
            "validate_candidate.py",
            "validate_verifier_verdict.py",
            "verify_candidate.py",
        ]:
            if not (workspace / "bin" / contract_helper).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing {contract_helper}")
        for contract_schema in ["audit-timeline.schema.json", "recon-result.schema.json", "triage-batch.schema.json"]:
            if not (workspace / "assets" / "schemas" / contract_schema).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing assets/schemas/{contract_schema}")
        for contract_adapter in ["validate-recon-result.py", "validate-triage-batch.py"]:
            if not (workspace / "scripts" / contract_adapter).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing scripts/{contract_adapter}")
        for variant_helper in [
            "extract_variant_seed.py",
            "find_variant_candidates.py",
        ]:
            if not (workspace / "bin" / variant_helper).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing {variant_helper}")
            if not (workspace / "scripts" / variant_helper).exists():
                raise SystemExit(f"FAILED: bootstrapped workspace is missing scripts/{variant_helper}")
        exercise_variant_seed_card_validation(plugin_root, workspace)
        exercise_extract_variant_seed(plugin_root, workspace)
        exercise_variant_candidate_validation(plugin_root, workspace)
        if not (workspace / "bin/audit_disposition.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing audit_disposition.py")
        if not (workspace / "scripts/render-handoff-summary.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing scripts/render-handoff-summary.py")
        if not (workspace / "scripts/assert-finalized-workspace.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing scripts/assert-finalized-workspace.py")
        if not (workspace / "handoff-summary.md").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing handoff-summary.md")
        require_text(
            workspace / "handoff-summary.md",
            "<!-- schema_version: 1 -->",
            "bootstrapped handoff schema version marker",
        )
        require_text(
            workspace / "handoff-summary.md",
            "It is not a vulnerability report",
            "bootstrapped handoff non-report disclaimer",
        )
        for heading in (
            "Target and Workspace",
            "Current Stage / Status",
            "Recommended First Reads",
            "Context-Slimming Rules",
            "Attack-Surface Highlights",
            "Initial Probe Summary",
            "Blocked Verification Status",
            "Audit Disposition Ledger",
            "Candidate Findings",
            "False Positives / Non-Security Defects",
            "Unverified Leads",
            "Confirmed Bundle Pointers",
            "Heavy Logs To Avoid Unless Needed",
            "Next Safe Steps",
            "Confirmed-Only Routing Guardrails",
        ):
            require_text(
                workspace / "handoff-summary.md",
                heading,
                f"bootstrapped handoff heading {heading}",
            )
        require_text(
            workspace / "handoff-summary.md",
            "Read lightweight files first",
            "bootstrapped handoff context-slimming rule",
        )
        require_text(
            workspace / "handoff-summary.md",
            "Avoid default-reading full raw logs",
            "bootstrapped handoff raw-log avoidance rule",
        )
        require_text(
            workspace / "handoff-summary.md",
            "Confirmed vulnerabilities belong only under `confirmed/<one-folder-per-vulnerability>/`",
            "bootstrapped handoff confirmed-only guardrail",
        )
        require_text(
            workspace / "handoff-summary.md",
            "Finalization integrity: `not_finalized`",
            "bootstrapped handoff finalization integrity state",
        )

        docker_baseline = workspace / "docker" / "baseline-fixture.json"
        docker_current = workspace / "docker" / "current-fixture.json"
        docker_baseline.write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "captured_at": "2026-04-28T00:00:00Z",
                    "docker_available": True,
                    "images": [{"id": "sha256:base", "repository": "node", "tag": "20-alpine"}],
                    "volumes": [{"name": "existing-volume", "driver": "local"}],
                    "networks": [{"id": "net0", "name": "bridge", "driver": "bridge"}],
                    "containers": [{"id": "container0", "name": "existing", "state": "exited"}],
                    "build_cache": [{"id": "cache0", "reclaimable": True, "size": "1MB"}],
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        docker_current.write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "captured_at": "2026-04-28T00:10:00Z",
                    "docker_available": True,
                    "images": [
                        {"id": "sha256:base", "repository": "node", "tag": "20-alpine"},
                        {
                            "id": "sha256:new",
                            "repository": "target-app",
                            "tag": "latest",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": workspace_name,
                            },
                        },
                        {
                            "id": "sha256:foreign",
                            "repository": "other-app",
                            "tag": "latest",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": "security-research-other",
                            },
                        },
                        {
                            "id": "sha256:projectonly",
                            "repository": "project-only-app",
                            "tag": "latest",
                            "labels": {"com.zhulong.project": workspace_name},
                        },
                        {
                            "id": "sha256:legacy",
                            "repository": "starlette-verify",
                            "tag": "zhulong",
                            "labels": {"com.zhulong.workspace": workspace_name},
                        },
                        {
                            "id": "sha256:compose",
                            "repository": "<none>",
                            "tag": "<none>",
                            "labels": {"com.docker.compose.project": "zhulong-test-compose"},
                        },
                        {"id": "sha256:compose-pulled", "repository": "mysql", "tag": "5.7"},
                        {"id": "sha256:unlabeled", "repository": "parallel-app", "tag": "latest"},
                    ],
                    "volumes": [
                        {"name": "existing-volume", "driver": "local"},
                        {
                            "name": "target-created-volume",
                            "driver": "local",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": workspace_name,
                            },
                        },
                        {
                            "name": "target-compose-volume",
                            "driver": "local",
                            "labels": {"com.docker.compose.project": "zhulong-test-compose"},
                        },
                        {
                            "name": "legacy-labeled-volume",
                            "driver": "local",
                            "labels": {"com.zhulong.workspace": workspace_name},
                        },
                        {
                            "name": "project-only-volume",
                            "driver": "local",
                            "labels": {"com.zhulong.project": workspace_name},
                        },
                        {"name": "parallel-created-volume", "driver": "local"},
                    ],
                    "networks": [
                        {"id": "net0", "name": "bridge", "driver": "bridge"},
                        {
                            "id": "net1",
                            "name": "target-created-network",
                            "driver": "bridge",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": workspace_name,
                            },
                        },
                        {
                            "id": "net3",
                            "name": "target-compose-network",
                            "driver": "bridge",
                            "labels": {"com.docker.compose.project": "zhulong-test-compose"},
                        },
                        {
                            "id": "net4",
                            "name": "legacy-labeled-network",
                            "driver": "bridge",
                            "labels": {"com.zhulong.workspace": workspace_name},
                        },
                        {
                            "id": "net5",
                            "name": "project-only-network",
                            "driver": "bridge",
                            "labels": {"com.zhulong.project": workspace_name},
                        },
                        {"id": "net2", "name": "parallel-created-network", "driver": "bridge"},
                    ],
                    "containers": [
                        {"id": "container0", "name": "existing", "state": "exited"},
                        {
                            "id": "container1",
                            "name": "target-stopped",
                            "state": "exited",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": workspace_name,
                            },
                        },
                        {
                            "id": "container2",
                            "name": "target-running",
                            "state": "running",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": workspace_name,
                            },
                        },
                        {
                            "id": "container6",
                            "name": "legacy-labeled-stopped",
                            "state": "exited",
                            "labels": {"com.zhulong.workspace": workspace_name},
                        },
                        {
                            "id": "container5",
                            "name": "target-compose-stopped",
                            "state": "exited",
                            "labels": {"com.docker.compose.project": "zhulong-test-compose"},
                        },
                        {
                            "id": "container3",
                            "name": "other-zhulong-stopped",
                            "state": "exited",
                            "labels": {
                                "org.zhulong.managed": "true",
                                "org.zhulong.workspace": "security-research-other",
                            },
                        },
                        {
                            "id": "container7",
                            "name": "project-only-stopped",
                            "state": "exited",
                            "labels": {"com.zhulong.project": workspace_name},
                        },
                        {"id": "container4", "name": "parallel-unlabeled-stopped", "state": "exited"},
                    ],
                    "build_cache": [
                        {"id": "cache0", "reclaimable": True, "size": "1MB"},
                        {"id": "cache1", "reclaimable": True, "size": "2MB"},
                        {"id": "cache2", "reclaimable": False, "size": "3MB"},
                    ],
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--show-created",
        ], plugin_root)
        cleanup_plan = json.loads((workspace / "docker" / "docker-cleanup-plan.json").read_text(encoding="utf-8"))
        if cleanup_plan.get("safety_policy", {}).get("uses_docker_prune") is not False:
            raise SystemExit("FAILED: Docker cleanup helper must not use broad prune semantics")
        if cleanup_plan.get("safety_policy", {}).get("delete_unowned_resources") is not False:
            raise SystemExit("FAILED: Docker cleanup helper must not delete unowned resources")
        planned_images = {item.get("id") for item in cleanup_plan.get("images", [])}
        if planned_images != {"sha256:new", "sha256:legacy"}:
            raise SystemExit(f"FAILED: Docker cleanup image plan should only include owned new image: {planned_images}")
        planned_volumes = {item.get("name") for item in cleanup_plan.get("volumes", [])}
        if planned_volumes != {"target-created-volume", "legacy-labeled-volume"}:
            raise SystemExit(f"FAILED: Docker cleanup volume plan should only include owned new volume: {planned_volumes}")
        planned_networks = {item.get("name") for item in cleanup_plan.get("networks", [])}
        if planned_networks != {"target-created-network", "legacy-labeled-network"}:
            raise SystemExit(f"FAILED: Docker cleanup network plan should only include owned new non-default network: {planned_networks}")
        running_skipped = {item.get("name") for item in cleanup_plan.get("containers", {}).get("running_owned_skipped", [])}
        if running_skipped != {"target-running"}:
            raise SystemExit("FAILED: Docker cleanup helper must skip running containers by default")
        planned_containers = {item.get("name") for item in cleanup_plan.get("containers", {}).get("stopped_owned", [])}
        if planned_containers != {"target-stopped", "legacy-labeled-stopped"}:
            raise SystemExit(f"FAILED: Docker cleanup should only remove owned stopped containers: {planned_containers}")
        skipped_containers = {item.get("name") for item in cleanup_plan.get("containers", {}).get("unattributed_new_skipped", [])}
        if skipped_containers != {"other-zhulong-stopped", "parallel-unlabeled-stopped", "project-only-stopped", "target-compose-stopped"}:
            raise SystemExit(f"FAILED: Docker cleanup must skip foreign/unlabeled containers: {skipped_containers}")
        skipped_images = {item.get("id") for item in cleanup_plan.get("unattributed_new_skipped", {}).get("images", [])}
        if skipped_images != {"sha256:foreign", "sha256:projectonly", "sha256:unlabeled", "sha256:compose", "sha256:compose-pulled"}:
            raise SystemExit(f"FAILED: Docker cleanup must skip foreign/unlabeled images: {skipped_images}")
        skipped_volumes = {item.get("name") for item in cleanup_plan.get("unattributed_new_skipped", {}).get("volumes", [])}
        if skipped_volumes != {"parallel-created-volume", "project-only-volume", "target-compose-volume"}:
            raise SystemExit(f"FAILED: Docker cleanup must skip unlabeled volumes: {skipped_volumes}")
        skipped_networks = {item.get("name") for item in cleanup_plan.get("unattributed_new_skipped", {}).get("networks", [])}
        if skipped_networks != {"parallel-created-network", "project-only-network", "target-compose-network"}:
            raise SystemExit(f"FAILED: Docker cleanup must skip unlabeled networks: {skipped_networks}")
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--show-created",
            "--adopt-compose-project",
            "zhulong-*",
            "--adopt-image-ref",
            "*",
            "--adopt-network-name",
            "parallel-*",
            "--adopt-volume-name",
            "parallel-*",
            "--adopt-build-cache",
            "--adopt-build-cache-id",
            "cache*",
        ], plugin_root)
        cleanup_plan = json.loads((workspace / "docker" / "docker-cleanup-plan.json").read_text(encoding="utf-8"))
        planned_images = {item.get("id") for item in cleanup_plan.get("images", [])}
        planned_volumes = {item.get("name") for item in cleanup_plan.get("volumes", [])}
        planned_networks = {item.get("name") for item in cleanup_plan.get("networks", [])}
        planned_build_cache = {item.get("id") for item in cleanup_plan.get("build_cache", {}).get("adopted_reclaimable", [])}
        if (
            planned_images != {"sha256:new", "sha256:legacy"}
            or planned_volumes != {"target-created-volume", "legacy-labeled-volume"}
            or planned_networks != {"target-created-network", "legacy-labeled-network"}
            or planned_build_cache
        ):
            raise SystemExit("FAILED: Docker adoption flags must use exact literal matches, not wildcard/prefix semantics")
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--show-created",
            "--adopt-compose-project",
            "zhulong-test-compose",
            "--adopt-image-ref",
            "node:20-alpine",
            "--adopt-image-ref",
            "mysql:5.7",
            "--adopt-build-cache",
            "--adopt-build-cache-id",
            "cache1",
            "--adopt-volume-name",
            "existing-volume",
            "--adopt-network-name",
            "parallel-created-network",
            "--adopt-volume-name",
            "parallel-created-volume",
        ], plugin_root)
        cleanup_plan = json.loads((workspace / "docker" / "docker-cleanup-plan.json").read_text(encoding="utf-8"))
        planned_images = {item.get("id") for item in cleanup_plan.get("images", [])}
        if planned_images != {"sha256:new", "sha256:legacy", "sha256:compose", "sha256:compose-pulled"}:
            raise SystemExit(f"FAILED: adopted compose/image resources should enter cleanup image plan: {planned_images}")
        planned_volumes = {item.get("name") for item in cleanup_plan.get("volumes", [])}
        if planned_volumes != {"target-created-volume", "legacy-labeled-volume", "target-compose-volume", "parallel-created-volume"}:
            raise SystemExit(f"FAILED: adopted compose/exact resources should enter cleanup volume plan: {planned_volumes}")
        planned_networks = {item.get("name") for item in cleanup_plan.get("networks", [])}
        if planned_networks != {"target-created-network", "legacy-labeled-network", "target-compose-network", "parallel-created-network"}:
            raise SystemExit(f"FAILED: adopted compose/exact resources should enter cleanup network plan: {planned_networks}")
        planned_containers = {item.get("name") for item in cleanup_plan.get("containers", {}).get("stopped_owned", [])}
        if planned_containers != {"target-stopped", "legacy-labeled-stopped", "target-compose-stopped"}:
            raise SystemExit(f"FAILED: adopted compose containers should enter cleanup plan: {planned_containers}")
        planned_build_cache = {item.get("id") for item in cleanup_plan.get("build_cache", {}).get("adopted_reclaimable", [])}
        if planned_build_cache != {"cache1"}:
            raise SystemExit(f"FAILED: adopted build cache should enter cleanup plan: {planned_build_cache}")
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--show-created",
            "--adopt-build-cache",
        ], plugin_root)
        cleanup_plan = json.loads((workspace / "docker" / "docker-cleanup-plan.json").read_text(encoding="utf-8"))
        if cleanup_plan.get("build_cache", {}).get("adopted_reclaimable"):
            raise SystemExit("FAILED: BuildKit cache adoption must require explicit cache IDs")
        skipped_build_cache = {item.get("id") for item in cleanup_plan.get("build_cache", {}).get("unattributed_new_skipped", [])}
        if skipped_build_cache != {"cache1"}:
            raise SystemExit(f"FAILED: unattributed BuildKit cache should remain review-only without exact IDs: {skipped_build_cache}")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--show-created",
            "--adopt-build-cache-id",
            "cache1",
        ], plugin_root, "--adopt-build-cache-id requires --adopt-build-cache")

        def docker_overwrite_fixture(
            name: str,
            *,
            images: list[dict[str, object]] | None = None,
            volumes: list[dict[str, object]] | None = None,
            networks: list[dict[str, object]] | None = None,
            containers: list[dict[str, object]] | None = None,
            build_cache: list[dict[str, object]] | None = None,
        ) -> Path:
            path = workspace / "docker" / f"{name}.json"
            path.write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "captured_at": "2026-04-28T00:15:00Z",
                        "docker_available": True,
                        "images": [
                            {"id": "sha256:base", "repository": "node", "tag": "20-alpine"},
                            *(images or []),
                        ],
                        "volumes": [
                            {"name": "existing-volume", "driver": "local"},
                            *(volumes or []),
                        ],
                        "networks": [
                            {"id": "net0", "name": "bridge", "driver": "bridge"},
                            *(networks or []),
                        ],
                        "containers": [
                            {"id": "container0", "name": "existing", "state": "exited"},
                            *(containers or []),
                        ],
                        "build_cache": [
                            {"id": "cache0", "reclaimable": True, "size": "1MB"},
                            *(build_cache or []),
                        ],
                    },
                    indent=2,
                ),
                encoding="utf-8",
            )
            return path

        overwrite_owned_image = docker_overwrite_fixture(
            "current-overwrite-owned-image",
            images=[
                {
                    "id": "sha256:owned-overwrite",
                    "repository": "owned-overwrite",
                    "tag": "latest",
                    "labels": {
                        "org.zhulong.managed": "true",
                        "org.zhulong.workspace": workspace_name,
                    },
                }
            ],
        )
        overwrite_unlabeled_image = docker_overwrite_fixture(
            "current-overwrite-unlabeled-image",
            images=[{"id": "sha256:unlabeled-overwrite", "repository": "unlabeled-overwrite", "tag": "latest"}],
        )
        overwrite_unlabeled_network_volume = docker_overwrite_fixture(
            "current-overwrite-unlabeled-network-volume",
            volumes=[{"name": "unlabeled-overwrite-volume", "driver": "local"}],
            networks=[{"id": "net-overwrite", "name": "unlabeled-overwrite-network", "driver": "bridge"}],
        )
        overwrite_build_cache = docker_overwrite_fixture(
            "current-overwrite-build-cache",
            build_cache=[{"id": "cache-overwrite", "reclaimable": True, "size": "2MB"}],
        )
        for overwrite_current in (
            overwrite_owned_image,
            overwrite_unlabeled_image,
            overwrite_unlabeled_network_volume,
            overwrite_build_cache,
        ):
            run_expect_fail([
                sys.executable,
                str(plugin_root / "scripts/manage_docker_resources.py"),
                "--workspace-dir",
                str(workspace),
                "--baseline-file",
                str(docker_baseline),
                "--current-file",
                str(overwrite_current),
                "--capture-baseline",
                "--force-overwrite-baseline",
            ], plugin_root, "hide Docker residue from strict cleanliness checks")
        cleanup_plan = json.loads((workspace / "docker" / "docker-cleanup-plan.json").read_text(encoding="utf-8"))
        skipped_build_cache = {
            item.get("id")
            for item in cleanup_plan.get("build_cache", {}).get("unattributed_new_skipped", [])
        }
        if skipped_build_cache != {"cache-overwrite"}:
            raise SystemExit(f"FAILED: baseline overwrite refusal must record BuildKit cache residue: {skipped_build_cache}")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--capture-baseline",
            "--force-overwrite-baseline",
        ], plugin_root, "Refusing to overwrite Docker baseline while post-baseline resources remain")
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--cleanup-created",
        ], plugin_root)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_current),
            "--verify-clean",
        ], plugin_root, "owned Docker resources remain")
        cleanliness_status = json.loads((workspace / "docker" / "docker-cleanliness-status.json").read_text(encoding="utf-8"))
        if cleanliness_status.get("clean") is not False:
            raise SystemExit("FAILED: Docker verify-clean must fail when owned resources remain")
        docker_clean_current = workspace / "docker" / "current-clean-fixture.json"
        docker_clean_current.write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "captured_at": "2026-04-28T00:20:00Z",
                    "docker_available": True,
                    "images": [
                        {"id": "sha256:base", "repository": "node", "tag": "20-alpine"},
                        {"id": "sha256:foreign", "repository": "other-app", "tag": "latest"},
                    ],
                    "volumes": [
                        {"name": "existing-volume", "driver": "local"},
                        {"name": "parallel-created-volume", "driver": "local"},
                    ],
                    "networks": [
                        {"id": "net0", "name": "bridge", "driver": "bridge"},
                        {"id": "net2", "name": "parallel-created-network", "driver": "bridge"},
                    ],
                    "containers": [
                        {"id": "container0", "name": "existing", "state": "exited"},
                        {"id": "container4", "name": "parallel-unlabeled-stopped", "state": "exited"},
                    ],
                    "build_cache": [
                        {"id": "cache0", "reclaimable": True, "size": "1MB"},
                        {"id": "cache3", "reclaimable": True, "size": "2MB"},
                    ],
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_clean_current),
            "--verify-clean",
        ], plugin_root)
        cleanliness_status = json.loads((workspace / "docker" / "docker-cleanliness-status.json").read_text(encoding="utf-8"))
        if cleanliness_status.get("clean") is not True:
            raise SystemExit("FAILED: Docker verify-clean must pass when no current-workspace owned resources remain")
        strict_blocker_proc = subprocess.run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_clean_current),
            "--verify-clean",
            "--strict",
        ], cwd=plugin_root, capture_output=True, text=True)
        strict_blocker_output = (strict_blocker_proc.stdout or "") + (strict_blocker_proc.stderr or "")
        if strict_blocker_proc.returncode == 0:
            raise SystemExit("FAILED: Docker strict verify-clean unexpectedly passed on post-baseline unattributed resources")
        for expected in (
            "unattributed Docker resources remain",
            "BuildKit cache blocker",
            "review-only and cannot be auto-deleted safely",
            "must remain blocked",
            "must not manually mark the audit completed",
            "--adopt-build-cache --adopt-build-cache-id <cache-id>",
        ):
            if expected not in strict_blocker_output:
                raise SystemExit(f"FAILED: Docker strict BuildKit blocker output missing: {expected}")
        cleanliness_status = json.loads((workspace / "docker" / "docker-cleanliness-status.json").read_text(encoding="utf-8"))
        if cleanliness_status.get("clean") is not False or cleanliness_status.get("strict") is not True:
            raise SystemExit("FAILED: Docker strict verify-clean must fail on post-baseline unattributed resources")
        docker_strict_clean_current = workspace / "docker" / "current-strict-clean-fixture.json"
        docker_strict_clean_current.write_text(
            json.dumps(
                {
                    "schema_version": 1,
                    "captured_at": "2026-04-28T00:30:00Z",
                    "docker_available": True,
                    "images": [{"id": "sha256:base", "repository": "node", "tag": "20-alpine"}],
                    "volumes": [{"name": "existing-volume", "driver": "local"}],
                    "networks": [{"id": "net0", "name": "bridge", "driver": "bridge"}],
                    "containers": [{"id": "container0", "name": "existing", "state": "exited"}],
                    "build_cache": [{"id": "cache0", "reclaimable": True, "size": "1MB"}],
                },
                indent=2,
            ),
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_baseline),
            "--current-file",
            str(docker_strict_clean_current),
            "--verify-clean",
            "--strict",
        ], plugin_root)
        cleanliness_status = json.loads((workspace / "docker" / "docker-cleanliness-status.json").read_text(encoding="utf-8"))
        if cleanliness_status.get("clean") is not True or cleanliness_status.get("strict") is not True:
            raise SystemExit("FAILED: Docker strict verify-clean must pass when the Docker state matches the baseline")
        docker_overwrite_success_baseline = workspace / "docker" / "baseline-overwrite-success.json"
        docker_overwrite_success_baseline.write_text(docker_baseline.read_text(encoding="utf-8"), encoding="utf-8")
        run([
            sys.executable,
            str(plugin_root / "scripts/manage_docker_resources.py"),
            "--workspace-dir",
            str(workspace),
            "--baseline-file",
            str(docker_overwrite_success_baseline),
            "--current-file",
            str(docker_strict_clean_current),
            "--capture-baseline",
            "--force-overwrite-baseline",
        ], plugin_root)
        overwritten_baseline = json.loads(docker_overwrite_success_baseline.read_text(encoding="utf-8"))
        if overwritten_baseline.get("captured_at") != "2026-04-28T00:30:00Z":
            raise SystemExit("FAILED: force-overwrite baseline should succeed only after current state matches baseline residue-free")
        require_text(
            workspace / "handoff-summary.md",
            "Do not generate DOCX reports from handoff content",
            "bootstrapped handoff no-DOCX guardrail",
        )
        require_text(
            workspace / "handoff-summary.md",
            "Do not copy raw scanner logs into this handoff",
            "bootstrapped handoff raw-log dump warning",
        )
        if not (workspace / "scripts/run-verification-case.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing scripts/run-verification-case.sh")
        if not (workspace / "bin/asr-start.sh").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing asr-start.sh")
        fake_bin = Path(tempdir) / "fake-bin"
        fake_bin.mkdir(parents=True, exist_ok=True)
        fake_osv = fake_bin / "osv-scanner"
        fake_osv.write_text(
            "#!/usr/bin/env bash\n"
            "echo 'No package sources found'\n"
            "exit 128\n",
            encoding="utf-8",
        )
        fake_osv.chmod(0o755)
        fake_gitleaks = fake_bin / "gitleaks"
        fake_gitleaks.write_text(
            "#!/usr/bin/env bash\n"
            "set -euo pipefail\n"
            "report_path=''\n"
            "while [[ $# -gt 0 ]]; do\n"
            "  case \"$1\" in\n"
            "    --report-path)\n"
            "      report_path=\"${2:-}\"\n"
            "      shift 2\n"
            "      ;;\n"
            "    *)\n"
            "      shift\n"
            "      ;;\n"
            "  esac\n"
            "done\n"
            "[[ -n \"$report_path\" ]] || exit 64\n"
            "python3 - <<'PY' \"$report_path\"\n"
            "import json, sys\n"
            "items = [\n"
            "    {\n"
            "        'RuleID': 'generic-api-key',\n"
            "        'Description': 'Generic API Key',\n"
            "        'File': 'config/example.env',\n"
            "        'StartLine': 3,\n"
            "        'Commit': 'abcdef1234567890',\n"
            "        'Secret': 'sk_live_SUPER_SECRET_VALUE_123456',\n"
            "        'Match': 'API_KEY=sk_live_SUPER_SECRET_VALUE_123456',\n"
            "    },\n"
            "    {\n"
            "        'RuleID': 'private-key',\n"
            "        'Description': 'Private Key',\n"
            "        'File': 'tests/fixtures/key.pem',\n"
            "        'StartLine': 1,\n"
            "        'Secret': '-----BEGIN PRIVATE KEY-----FAKESECRET-----END PRIVATE KEY-----',\n"
            "    },\n"
            "]\n"
            "open(sys.argv[1], 'w', encoding='utf-8').write(json.dumps(items))\n"
            "PY\n"
            "echo 'leaks found: 2'\n"
            "exit 1\n",
            encoding="utf-8",
        )
        fake_gitleaks.chmod(0o755)
        fake_trivy = fake_bin / "trivy"
        fake_trivy.write_text(
            "#!/usr/bin/env bash\n"
            "echo 'simulated scanner failure for nonfatal classification' >&2\n"
            "exit 2\n",
            encoding="utf-8",
        )
        fake_trivy.chmod(0o755)
        fake_grype = fake_bin / "grype"
        fake_grype.write_text(
            "#!/usr/bin/env bash\n"
            "echo 'simulated command execution failure for fatal classification' >&2\n"
            "exit 127\n",
            encoding="utf-8",
        )
        fake_grype.chmod(0o755)
        probe_env = {
            **dict(),
            "PATH": f"{fake_bin}:/usr/bin:/bin:/usr/sbin:/sbin",
            "HOME": str(Path.home()),
        }
        run_with_env([
            "/bin/bash",
            str(workspace / "bin/run-initial-probes.sh"),
            "--repo-root",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root, probe_env)
        initial_summary = workspace / "evidence/initial-probes/initial-probes-summary.json"
        if not initial_summary.exists():
            raise SystemExit("FAILED: run_initial_probes did not write initial-probes-summary.json")
        summary_data = json.loads(initial_summary.read_text(encoding="utf-8"))
        for field in ("schema_version", "generated_at", "repo_root", "workspace_dir", "output_dir", "probes"):
            if field not in summary_data:
                raise SystemExit(f"FAILED: initial-probes-summary.json is missing {field}")
        labels = set(summary_data.get("stable_status_labels") or [])
        expected_labels = {"ran_ok", "skipped_tool_missing", "skipped_no_package_sources", "failed_nonfatal", "failed_fatal"}
        if labels != expected_labels:
            raise SystemExit(f"FAILED: initial probe stable labels mismatch: {sorted(labels)}")
        probes = summary_data.get("probes") or []
        if not isinstance(probes, list) or not probes:
            raise SystemExit("FAILED: initial-probes-summary.json probes must be a non-empty list")
        by_name = {probe.get("name"): probe for probe in probes}
        require_probe_record(
            initial_summary,
            workspace / "evidence/initial-probes",
            "osv-scanner",
            "skipped_no_package_sources",
            128,
            "no supported package lockfile",
            "exited non-zero",
        )
        if by_name.get("semgrep", {}).get("status") != "skipped_tool_missing":
            raise SystemExit("FAILED: missing semgrep was not classified as skipped_tool_missing")
        gitleaks_probe = by_name.get("gitleaks", {})
        if gitleaks_probe.get("status") != "failed_nonfatal":
            raise SystemExit("FAILED: gitleaks leak-found exit was not classified as failed_nonfatal")
        if gitleaks_probe.get("exit_code") != 1:
            raise SystemExit("FAILED: gitleaks leak-found exit code was not preserved")
        gitleaks_summary = gitleaks_probe.get("summary") or {}
        if gitleaks_summary.get("finding_count") != 2:
            raise SystemExit("FAILED: gitleaks summary did not preserve finding_count")
        samples = gitleaks_summary.get("sample_findings") or []
        if not (2 <= len(samples) <= 5):
            raise SystemExit("FAILED: gitleaks summary samples were not captured")
        if samples[0].get("rule_id") != "generic-api-key" or samples[0].get("file") != "config/example.env":
            raise SystemExit("FAILED: gitleaks summary did not include actionable metadata")
        summary_text = json.dumps(gitleaks_probe, ensure_ascii=False)
        for forbidden_secret in (
            "sk_live_SUPER_SECRET_VALUE_123456",
            "API_KEY=sk_live_SUPER_SECRET_VALUE_123456",
            "-----BEGIN PRIVATE KEY-----FAKESECRET-----END PRIVATE KEY-----",
        ):
            if forbidden_secret in summary_text:
                raise SystemExit("FAILED: gitleaks summary copied a secret-like value verbatim")
        if "secret_sha256_12" not in summary_text or "secret_redacted" not in summary_text:
            raise SystemExit("FAILED: gitleaks summary should include only redacted/hash secret hints")
        for field in ("top_rule_ids", "path_category_counts", "top_rule_path_categories"):
            if field not in gitleaks_summary:
                raise SystemExit(f"FAILED: gitleaks summary missing aggregation field: {field}")
        if str(gitleaks_summary.get("raw_log_path", "")).startswith("/"):
            raise SystemExit("FAILED: gitleaks raw_log_path should be relative")
        if not (workspace / "evidence/initial-probes/gitleaks.log").exists():
            raise SystemExit("FAILED: gitleaks raw log was not preserved")
        if not (workspace / "evidence/initial-probes/gitleaks.json").exists():
            raise SystemExit("FAILED: gitleaks JSON report was not preserved")
        if by_name.get("syft", {}).get("status") != "skipped_tool_missing":
            raise SystemExit("FAILED: missing syft was not classified as skipped_tool_missing")
        if by_name.get("trivy", {}).get("status") != "failed_nonfatal":
            raise SystemExit("FAILED: non-zero trivy was not classified as failed_nonfatal")
        if by_name.get("trivy", {}).get("exit_code") != 2:
            raise SystemExit("FAILED: failed_nonfatal trivy exit code was not preserved")
        if by_name.get("grype", {}).get("status") != "failed_fatal":
            raise SystemExit("FAILED: exit-127 grype was not classified as failed_fatal")
        if by_name.get("grype", {}).get("exit_code") != 127:
            raise SystemExit("FAILED: failed_fatal grype exit code was not preserved")
        if by_name.get("semgrep", {}).get("command") != "(not executed)":
            raise SystemExit("FAILED: skipped semgrep should use a descriptive command placeholder")
        for probe in probes:
            for field in ("name", "status", "command", "exit_code", "log_path", "reason", "next_action"):
                if field not in probe:
                    raise SystemExit(f"FAILED: initial probe record missing {field}: {probe}")
            for path_field in ("log_path",):
                value = str(probe.get(path_field) or "")
                if value.startswith("/"):
                    raise SystemExit(f"FAILED: initial probe {path_field} should be relative: {value}")
        if str(summary_data.get("repo_root")).startswith("/"):
            raise SystemExit("FAILED: initial-probes-summary.json repo_root should not leak an absolute path")
        if str(summary_data.get("workspace_dir")).startswith("/"):
            raise SystemExit("FAILED: initial-probes-summary.json workspace_dir should not leak an absolute path")
        if str(summary_data.get("output_dir")).startswith("/"):
            raise SystemExit("FAILED: initial-probes-summary.json output_dir should not leak an absolute path")
        fake_gitleaks.write_text(
            "#!/usr/bin/env bash\n"
            "set -euo pipefail\n"
            "report_path=''\n"
            "while [[ $# -gt 0 ]]; do\n"
            "  case \"$1\" in --report-path) report_path=\"${2:-}\"; shift 2 ;; *) shift ;; esac\n"
            "done\n"
            "[[ -n \"$report_path\" ]] || exit 64\n"
            "python3 - <<'PY' \"$report_path\"\n"
            "import json, sys\n"
            "items=[]\n"
            "paths=['tests/fixtures/a.env','docs/example.md','examples/demo.env','app/config/specs/open-api.json','src/Service.php','fixtures/key.txt']\n"
            "rules=['generic-api-key','jwt','private-key']\n"
            "for i in range(36):\n"
            "    secret=f'SECRET_VALUE_{i:04d}_DO_NOT_COPY'\n"
            "    items.append({'RuleID': rules[i % len(rules)], 'Description': 'Synthetic secret', 'File': paths[i % len(paths)], 'StartLine': i + 1, 'Commit': f'commit{i % 4}', 'Secret': secret, 'Match': 'TOKEN=' + secret})\n"
            "open(sys.argv[1], 'w', encoding='utf-8').write(json.dumps(items))\n"
            "PY\n"
            "echo 'leaks found: 36'\n"
            "exit 1\n",
            encoding="utf-8",
        )
        fake_gitleaks.chmod(0o755)
        large_gitleaks_output = workspace / "evidence/initial-probes-gitleaks-large"
        run_with_env([
            "/bin/bash",
            str(workspace / "bin/run-initial-probes.sh"),
            "--repo-root",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
            "--output-dir",
            str(large_gitleaks_output),
        ], plugin_root, probe_env)
        large_probe = require_probe_record(
            large_gitleaks_output / "initial-probes-summary.json",
            large_gitleaks_output,
            "gitleaks",
            "failed_nonfatal",
            1,
            "gitleaks exited non-zero",
        )
        large_summary = large_probe.get("summary") or {}
        if large_summary.get("finding_count") != 36:
            raise SystemExit("FAILED: large gitleaks summary did not preserve finding_count")
        if len(large_summary.get("sample_findings") or []) > 5:
            raise SystemExit("FAILED: large gitleaks summary sample_findings exceeded cap")
        for field in ("top_rule_ids", "path_category_counts", "top_rule_path_categories", "top_commits"):
            if not large_summary.get(field):
                raise SystemExit(f"FAILED: large gitleaks summary missing populated aggregation: {field}")
        large_summary_text = json.dumps(large_summary, ensure_ascii=False)
        if "SECRET_VALUE_" in large_summary_text or "TOKEN=SECRET" in large_summary_text:
            raise SystemExit("FAILED: large gitleaks summary copied secret-like values verbatim")
        fake_osv.write_text(
            "#!/usr/bin/env bash\n"
            "echo 'OSV scan completed successfully with no vulnerable packages'\n"
            "exit 0\n",
            encoding="utf-8",
        )
        fake_osv.chmod(0o755)
        osv_ok_output = workspace / "evidence/initial-probes-osv-ok"
        run_with_env([
            "/bin/bash",
            str(workspace / "bin/run-initial-probes.sh"),
            "--repo-root",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
            "--output-dir",
            str(osv_ok_output),
        ], plugin_root, probe_env)
        require_probe_record(
            osv_ok_output / "initial-probes-summary.json",
            osv_ok_output,
            "osv-scanner",
            "ran_ok",
            0,
            "completed with exit code 0",
            "exited non-zero",
        )
        fake_osv.write_text(
            "#!/usr/bin/env bash\n"
            "echo 'simulated unexpected OSV failure' >&2\n"
            "exit 42\n",
            encoding="utf-8",
        )
        fake_osv.chmod(0o755)
        osv_failure_output = workspace / "evidence/initial-probes-osv-failure"
        run_with_env([
            "/bin/bash",
            str(workspace / "bin/run-initial-probes.sh"),
            "--repo-root",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
            "--output-dir",
            str(osv_failure_output),
        ], plugin_root, probe_env)
        require_probe_record(
            osv_failure_output / "initial-probes-summary.json",
            osv_failure_output,
            "osv-scanner",
            "failed_nonfatal",
            42,
            "exited non-zero for a reason other than no package sources",
        )
        run([
            sys.executable,
            str(workspace / "bin/render-handoff-summary.py"),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            workspace / "handoff-summary.md",
            "osv-scanner: skipped_no_package_sources",
            "rendered handoff initial probe status",
        )
        require_text(
            workspace / "handoff-summary.md",
            "semgrep: skipped_tool_missing",
            "rendered handoff missing tool status",
        )
        require_text(
            workspace / "bin/run-verification-case.sh",
            "failed_resource_limit",
            "bootstrapped verification runner stable labels",
        )
        require_text(
            workspace / "bin/run-verification-case.sh",
            "Verification command timed out. Re-analyze service readiness",
            "bootstrapped verification runner timeout guidance",
        )
        if not (workspace / "unverified-leads.md").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing unverified-leads.md")
        if not (workspace / "attack-surface.md").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing attack-surface.md")
        for heading in (
            "Repository / Stack Summary",
            "External Entry Points",
            "Trusted and Untrusted Input Sources / Trust Boundaries",
            "Auth / Session / Permission Boundaries",
            "High-Risk Sinks",
            "Source-to-Sink Hypotheses",
            "Docker Verification Status",
            "Confirmed / False-Positive / Unverified Routing Reminder",
            "Next Safe Audit Steps",
        ):
            require_text(
                workspace / "attack-surface.md",
                heading,
                f"bootstrapped attack-surface heading {heading}",
            )
        require_text(
            workspace / "attack-surface.md",
            "not a vulnerability report, not raw scanner output",
            "bootstrapped attack-surface non-report guardrail",
        )
        require_text(
            workspace / "candidate-findings.md",
            "Source-to-Sink Hypothesis",
            "bootstrapped candidate findings stable columns",
        )
        require_text(
            workspace / "false-positives.md",
            "False Positives and Non-Security Defects",
            "bootstrapped false positives stable heading",
        )
        require_text(
            workspace / "unverified-leads.md",
            "High-Confidence-Unverified?",
            "bootstrapped unverified leads stable columns",
        )
        require_text(
            workspace / "unverified-leads.md",
            "Why completion is still safe?",
            "bootstrapped unverified leads materiality columns",
        )
        if not (workspace / "stage-status.json").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing stage-status.json")
        if not (workspace / "audit-events.jsonl").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing audit-events.jsonl")
        if not (workspace / "bin/write-audit-event.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing write-audit-event.py")
        if not (workspace / "bin/audit_transition_policy.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing audit_transition_policy.py")
        if not (workspace / "bin/validate-workspace-state.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing validate-workspace-state.py")
        if not (workspace / "bin/plan-security-toolchain.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing plan-security-toolchain.py")
        exercise_workspace_tool_registry_snapshot(plugin_root, workspace)
        if not (workspace / "bin/render-confirmed-vuln-docx.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing render-confirmed-vuln-docx.py")
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir",
            str(workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        sandbox_workspace = Path(tempdir) / "sandbox-preflight-workspace"
        sandbox_workspace.mkdir(parents=True, exist_ok=True)
        (sandbox_workspace / "asr-config.json").write_text('{"schema_version":1}\n', encoding="utf-8")
        (sandbox_workspace / "audit-log.md").write_text("# Audit Log\n", encoding="utf-8")
        exercise_sandbox_preflight(
            plugin_root / "scripts/check_sandbox_preflight.py",
            sandbox_workspace,
            plugin_root,
        )
        exercise_sandbox_preflight(
            workspace / "bin/check-sandbox-preflight.py",
            sandbox_workspace,
            plugin_root,
        )
        exercise_runner_sandbox_rejection(
            plugin_root / "scripts/run_verification_case.sh",
            sandbox_workspace,
            plugin_root,
        )
        exercise_verification_wrapper_state_boundary(plugin_root, Path(tempdir))
        exercise_sandbox_ledger_guard(sandbox_workspace, plugin_root)
        run([
            "bash",
            str(workspace / "bin/check_omc_runtime.sh"),
            "--workspace-dir",
            str(workspace),
            "--json",
        ], plugin_root)
        exercise_omc_runtime_hygiene(
            plugin_root / "scripts/check_omc_runtime.sh",
            workspace,
            plugin_root,
        )
        exercise_omc_runtime_hygiene(
            workspace / "bin/check_omc_runtime.sh",
            workspace,
            plugin_root,
        )
        run([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        (repo_dir / "pom.xml").write_text(
            "<project><modelVersion>4.0.0</modelVersion><groupId>selftest</groupId><artifactId>demo</artifactId><version>1</version></project>\n",
            encoding="utf-8",
        )
        java_controller = repo_dir / "src/main/java/example/DemoController.java"
        java_controller.parent.mkdir(parents=True, exist_ok=True)
        java_controller.write_text(
            "@RestController\nclass DemoController {\n  @GetMapping(\"/demo\") String demo(@RequestParam String name) { return name; }\n}\n",
            encoding="utf-8",
        )
        planner_output = run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        for expected in (
            "attack_surface_guidance:",
            "Java Web: inventory Spring/JAX-RS/Servlet routes",
            "Minimum entry inventory fields: route or endpoint, method, handler/controller",
            "current verification status",
        ):
            if expected not in planner_output:
                raise SystemExit(f"FAILED: planner output missing attack-surface guidance text: {expected}")
        go_router = repo_dir / "cmd/server/main.go"
        go_router.parent.mkdir(parents=True, exist_ok=True)
        (repo_dir / "go.mod").write_text(
            "module selftest\n\ngo 1.22\n",
            encoding="utf-8",
        )
        go_router.write_text(
            "package main\n\nimport \"net/http\"\n\nfunc main() {\n  http.HandleFunc(\"/demo\", func(w http.ResponseWriter, r *http.Request) {})\n}\n",
            encoding="utf-8",
        )
        mixed_planner_output = run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        if mixed_planner_output.count("Minimum entry inventory fields: route or endpoint, method, handler/controller") != 1:
            raise SystemExit("FAILED: planner output duplicated minimum entry inventory fields for mixed Java/Go workspace")
        (repo_dir / "package.json").write_text(
            '{"name":"selftest","version":"1.0.0","dependencies":{"express":"^4.18.0","fastify":"^4.0.0","next":"^14.0.0","lodash":"^4.17.21"}}\n',
            encoding="utf-8",
        )
        node_route = repo_dir / "routes/proxy.js"
        node_route.parent.mkdir(parents=True, exist_ok=True)
        node_route.write_text(
            "const fs = require('fs');\n"
            "const path = require('path');\n"
            "const express = require('express');\n"
            "const fastify = require('fastify')();\n"
            "const _ = require('lodash');\n"
            "const app = express();\n"
            "app.get('/proxy', proxy);\n"
            "fastify.post('/upload', async function route(request, reply) { return reply.send({ok: true}); });\n"
            "export default function handler(req, res) { return res.json({ok: true}); }\n"
            "async function proxy(req) {\n"
            "  await fetch(req.query.url);\n"
            "  fs.readFileSync(path.join('/srv/files', req.query.filename));\n"
            "  _.merge({}, JSON.parse('{\"__proto__\":{\"polluted\":true}}'));\n"
            "}\n",
            encoding="utf-8",
        )
        checklist_output = run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        for expected in (
            "local_knowledge_checklists:",
            "assets/references/ssrf-checklist.md",
            "assets/references/path-traversal-checklist.md",
            "assets/references/prototype-pollution-checklist.md",
        ):
            if expected not in checklist_output:
                raise SystemExit(f"FAILED: planner output missing checklist recommendation: {expected}")
        for expected in (
            "assets/references/nodejs-web-audit-playbook.md",
            "Node.js Web: inventory Express/Koa/Fastify/Next.js routes",
        ):
            if expected not in checklist_output:
                raise SystemExit(f"FAILED: planner output missing Node.js Web playbook recommendation: {expected}")
        library_repo = Path(tempdir) / "node-library-repo"
        library_repo.mkdir(parents=True, exist_ok=True)
        (library_repo / "package.json").write_text(
            '{"name":"selftest-library","version":"1.0.0","main":"lib/index.js","exports":"./lib/index.js"}\n',
            encoding="utf-8",
        )
        (library_repo / "lib").mkdir(parents=True, exist_ok=True)
        (library_repo / "lib" / "index.js").write_text(
            "exports.parse = function parse(input, options = {}) { return {input, options}; };\n",
            encoding="utf-8",
        )
        (library_repo / "api").mkdir(parents=True, exist_ok=True)
        (library_repo / "api" / "README.md").write_text(
            "# Public API documentation\n\nThis directory documents exported library APIs; it is not an HTTP API.\n",
            encoding="utf-8",
        )
        stale_workspace_api = library_repo / "security-research-20250101-000000/api/app.py"
        stale_workspace_api.parent.mkdir(parents=True, exist_ok=True)
        stale_workspace_api.write_text(
            "from flask import Flask\napp = Flask(__name__)\n@app.route('/stale')\ndef stale(): return 'stale'\n",
            encoding="utf-8",
        )
        library_plan = json.loads(run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(library_repo),
            "--workspace-dir",
            str(workspace),
            "--format",
            "json",
        ], plugin_root))
        library_hints = set(library_plan["attack_surface_hints"])
        if "node-library" not in library_hints:
            raise SystemExit("FAILED: planner did not classify pure Node.js package as node-library")
        for unexpected in ("http-api", "node-web", "python-web"):
            if unexpected in library_hints:
                raise SystemExit(f"FAILED: planner treated pure Node.js package as {unexpected}")
        if "assets/references/nodejs-library-audit-playbook.md" not in library_plan["specialized_playbooks"]:
            raise SystemExit("FAILED: planner did not recommend Node.js Library playbook")
        guidance_text = "\n".join(library_plan["attack_surface_guidance"])
        for expected in (
            "Node.js Library: inventory exported APIs",
            "Minimum library inventory fields: public API or CLI",
            "distinguish library-local behavior from application-level impact",
        ):
            if expected not in guidance_text:
                raise SystemExit(f"FAILED: planner output missing Node.js Library guidance: {expected}")
        (repo_dir / "requirements.txt").write_text(
            "flask\nfastapi\ndjango\n",
            encoding="utf-8",
        )
        python_app = repo_dir / "api/app.py"
        python_app.parent.mkdir(parents=True, exist_ok=True)
        python_app.write_text(
            "from flask import Flask, request\n"
            "from fastapi import FastAPI, UploadFile\n"
            "from django.urls import path\n\n"
            "app = Flask(__name__)\n"
            "api = FastAPI()\n\n"
            "@app.route('/download')\n"
            "def download():\n"
            "    return request.args.get('file', '')\n\n"
            "@api.post('/upload')\n"
            "async def upload(file: UploadFile):\n"
            "    return {'name': file.filename}\n\n"
            "urlpatterns = [path('demo/', lambda request: None)]\n",
            encoding="utf-8",
        )
        python_planner_output = run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(repo_dir),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        for expected in (
            "assets/references/python-web-audit-playbook.md",
            "Python Web: inventory Flask/Django/FastAPI/Starlette routes",
        ):
            if expected not in python_planner_output:
                raise SystemExit(f"FAILED: planner output missing Python Web playbook recommendation: {expected}")
        python_library_repo = Path(tempdir) / "python-library-repo"
        python_library_repo.mkdir(parents=True, exist_ok=True)
        (python_library_repo / "pyproject.toml").write_text(
            "[project]\nname = \"werkzeug-style-selftest\"\nversion = \"1.0.0\"\n",
            encoding="utf-8",
        )
        package_dir = python_library_repo / "src/werkzeug_style_selftest"
        package_dir.mkdir(parents=True, exist_ok=True)
        (package_dir / "__init__.py").write_text(
            "def parse_path(user_value, options=None):\n"
            "    return user_value\n",
            encoding="utf-8",
        )
        python_library_plan = json.loads(run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(python_library_repo),
            "--workspace-dir",
            str(workspace),
            "--format",
            "json",
        ], plugin_root))
        if "python-library" not in python_library_plan["attack_surface_hints"]:
            raise SystemExit("FAILED: planner did not classify pure Python package as python-library")
        if "http-api" in python_library_plan["attack_surface_hints"]:
            raise SystemExit("FAILED: planner forced a web route model on a pure Python library")
        if "assets/references/python-library-audit-playbook.md" not in python_library_plan["specialized_playbooks"]:
            raise SystemExit("FAILED: planner did not recommend Python Library playbook")
        python_library_guidance = "\n".join(python_library_plan["attack_surface_guidance"])
        for expected in (
            "Python Library: inventory public APIs",
            "Minimum Python library inventory fields: public API or hook",
            "do not force a route/method/handler table",
        ):
            if expected not in python_library_guidance:
                raise SystemExit(f"FAILED: planner output missing Python Library guidance: {expected}")
        appwrite_like_repo = Path(tempdir) / "appwrite-like-repo"
        appwrite_like_repo.mkdir(parents=True, exist_ok=True)
        (appwrite_like_repo / "composer.json").write_text(
            json.dumps({
                "name": "selftest/appwrite-like",
                "require": {
                    "php": "^8.3",
                    "ext-swoole": "*",
                    "utopia-php/framework": "^0.0.0",
                },
            }),
            encoding="utf-8",
        )
        php_worker = appwrite_like_repo / "src/Appwrite/Platform/Workers/Webhooks.php"
        php_worker.parent.mkdir(parents=True, exist_ok=True)
        php_worker.write_text(
            "<?php\n"
            "namespace Appwrite\\Platform\\Workers;\n"
            "use Swoole\\Runtime;\n"
            "final class Webhooks { public function execute($url) { $ch = curl_init($url); return curl_exec($ch); } }\n",
            encoding="utf-8",
        )
        (appwrite_like_repo / "frontend").mkdir()
        (appwrite_like_repo / "frontend/package-lock.json").write_text('{"lockfileVersion":3}\n', encoding="utf-8")
        (appwrite_like_repo / "tests/resources").mkdir(parents=True)
        (appwrite_like_repo / "tests/resources/package-lock.json").write_text('{"lockfileVersion":3}\n', encoding="utf-8")
        (appwrite_like_repo / "docker-compose.yml").write_text(
            "services:\n  appwrite:\n    image: appwrite/appwrite:dev\n",
            encoding="utf-8",
        )
        appwrite_plan = json.loads(run_capture([
            sys.executable,
            str(workspace / "bin/plan-security-toolchain.py"),
            "--target-dir",
            str(appwrite_like_repo),
            "--workspace-dir",
            str(workspace),
            "--format",
            "json",
        ], plugin_root))
        appwrite_hints = set(appwrite_plan["attack_surface_hints"])
        for expected in ("php-web", "php-swoole", "docker-compose"):
            if expected not in appwrite_hints:
                raise SystemExit(f"FAILED: Appwrite-like planner missing {expected}")
        playbooks = appwrite_plan["specialized_playbooks"]
        if not playbooks or playbooks[0] != "assets/references/php-swoole-audit-playbook.md":
            raise SystemExit(f"FAILED: Appwrite-like planner should lead with PHP/Swoole playbook: {playbooks}")
        if "assets/references/nodejs-web-audit-playbook.md" in playbooks:
            raise SystemExit("FAILED: Appwrite-like planner should not lead with Node Web from frontend/test lockfiles")
        appwrite_guidance = "\n".join(appwrite_plan["attack_surface_guidance"])
        for expected in (
            "PHP/Swoole: inventory Utopia routes/controllers",
            "frontend/test package-lock files as secondary",
        ):
            if expected not in appwrite_guidance:
                raise SystemExit(f"FAILED: Appwrite-like planner output missing PHP/Swoole guidance: {expected}")
        run_with_env([
            "bash",
            str(plugin_root / "scripts/refresh_workspace_helpers.sh"),
            "--workspace",
            str(workspace),
        ], plugin_root, {
            "SKILL_DIR": str(plugin_root),
            "HOME": str(Path.home()),
            "PATH": "/usr/bin:/bin:/usr/sbin:/sbin",
        })
        (repo_dir / "docker").mkdir(parents=True, exist_ok=True)
        (repo_dir / "poc").mkdir(parents=True, exist_ok=True)
        (repo_dir / "docker" / "docker-compose.attacker.yml").write_text(
            "services:\n  attacker:\n    image: alpine:3.20\n",
            encoding="utf-8",
        )
        (repo_dir / "poc" / "path_traversal.py").write_text(
            "print('demo poc')\\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(plugin_root / "assets/examples/confirmed-findings.example.json"),
            "--output-dir",
            str(workspace / "confirmed"),
            "--language",
            "zh-CN",
        ], plugin_root)
        run([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(plugin_root / "assets/examples/confirmed-findings.example.json"),
            "--output-dir",
            str(workspace / "confirmed"),
            "--language",
            "en-US",
        ], plugin_root)
        rendered_bundles = sorted(
            [
                path for path in (workspace / "confirmed").iterdir()
                if path.is_dir() and not path.name.startswith(".")
            ],
            key=lambda path: path.name,
        )
        if len(rendered_bundles) < 2:
            raise SystemExit("FAILED: bilingual example confirmed bundles were not rendered during selftest")
        zh_bundle = next((path for path in rendered_bundles if "漏洞报告" in path.name), None)
        en_bundle = next((path for path in rendered_bundles if path.name.endswith("_report")), None)
        if zh_bundle is None:
            raise SystemExit("FAILED: zh-CN confirmed bundle was not rendered during selftest")
        if en_bundle is None:
            raise SystemExit("FAILED: en-US confirmed bundle was not rendered during selftest")
        if not (zh_bundle / "verification-evidence.json").exists():
            raise SystemExit("FAILED: zh-CN confirmed bundle is missing verification-evidence.json")
        if not (en_bundle / "verification-evidence.json").exists():
            raise SystemExit("FAILED: en-US confirmed bundle is missing verification-evidence.json")
        write_live_replay_logs(zh_bundle, en_bundle)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(zh_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        standard_fixture_poc = repo_dir / "poc/jwt-forge-poc.py"
        standard_fixture_poc.parent.mkdir(parents=True, exist_ok=True)
        standard_fixture_poc.write_text("print('forged token accepted')\n", encoding="utf-8")
        standard_fixture_evidence = repo_dir / "poc/forged-token-response.json"
        standard_fixture_evidence.write_text('{"ok":true,"user":{"id":1}}\n', encoding="utf-8")
        standard_fixture = workspace / "standard-vulnerability-name-finding.json"
        standard_fixture.write_text(json.dumps({
            "project_name": "gothinkster/node-express-realworld-example-app",
            "vulnerability_id": "SELFTEST-001",
            "vulnerability_name": "硬编码 JWT 密钥导致身份认证绕过",
            "vulnerability_name_en": "Hardcoded JWT Secret Leading to Authentication Bypass",
            "severity": "critical",
            "severity_cn": "严重",
            "cwe": "CWE-798: Use of Hardcoded Credentials",
            "description": [
                "默认配置缺失 JWT_SECRET 时，应用回退到公开硬编码密钥，攻击者可伪造认证 token。"
            ],
            "impact": {
                "package": "gothinkster/node-express-realworld-example-app",
                "component": "src/app/routes/auth/auth.ts",
                "affected_versions": "default configuration",
                "repo_url": "https://github.com/gothinkster/node-express-realworld-example-app",
            },
            "cvss": {
                "vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:C/C:H/I:H/A:L",
                "score": "9.3",
                "severity": "Critical",
                "rationale": ["评估依据：攻击者可伪造任意用户 token，并完成未授权读写操作。"],
            },
            "analysis": [
                "位置：src/app/routes/auth/auth.ts 使用 process.env.JWT_SECRET || 'superSecret'。",
                "入口/可控输入：攻击者提交自签名 JWT token，请求受保护 API。",
                "危险函数/危险操作：express-jwt 使用公开默认密钥验证 HS256 token。",
                "触发路径：缺失 JWT_SECRET -> 默认密钥生效 -> 攻击者签发 token -> API 接受认证。",
                "根因：认证密钥存在硬编码回退值。",
                "现有校验为何失效：启动流程没有强制要求安全 JWT_SECRET。",
            ],
            "code_context": [
                {
                    "location": "src/app/routes/auth/auth.ts:18-26",
                    "summary": "认证中间件在缺失 JWT_SECRET 时回退到公开字符串 superSecret。",
                    "snippet": "const secret = process.env.JWT_SECRET || 'superSecret';\napp.use(jwt({ secret, algorithms: ['HS256'] }));",
                    "explanation": (
                        "攻击者可控输入是 Authorization 请求头中的自签名 JWT；该 token 通过认证中间件传播到 "
                        "express-jwt 校验 sink。缺失 guard 是启动阶段没有强制校验 JWT_SECRET，也没有拒绝默认密钥；"
                        "相邻的算法限制只限定 HS256，并不足以阻断公开密钥签名。Docker 已验证影响边界为身份认证绕过和未授权资料读取，不声称宿主机代码执行。"
                    ),
                }
            ],
            "reproduction": [
                {
                    "title": "1. 伪造 token 并访问受保护接口",
                    "details": [
                        "在 Docker Compose 环境中不设置 JWT_SECRET，确认应用按默认配置启动。",
                        "使用公开硬编码密钥 superSecret 构造 HS256 JWT，载荷中写入 user.id=1。",
                        "将伪造 token 放入 Authorization: Token <token> 请求头访问受保护接口。",
                    ],
                    "commands": [
                        "python3 poc/jwt-forge-poc.py",
                        "curl -s http://localhost:3000/api/user -H 'Authorization: Token <FORGED_TOKEN>'",
                    ],
                    "expected": ["预期结果：伪造 token 被服务端接受。"],
                    "observed": ["实际结果：HTTP 200 返回用户资料。"],
                    "results": [
                        "结果证据：forged-token-response.json 显示认证绕过成功。",
                        "结果证据：响应中包含 user.id=1，且没有返回 401 未授权错误。",
                    ],
                }
            ],
            "verification_status": "confirmed_in_docker",
            "verification_evidence": {
                "docker_image": "selftest-realworld-api",
                "docker_command": "docker compose up -d",
                "poc_path": "poc/jwt-forge-poc.py",
                "evidence_files": ["poc/forged-token-response.json"],
                "expected_observation": "预期结果：伪造 token 被服务端接受。",
                "observed_observation": "实际结果：HTTP 200 返回用户资料。",
                "oracle_token": "认证绕过成功",
                "severity_escalation_attempted": True,
                "severity_escalation_result": "Critical impact confirmed in Docker.",
            },
            "attachments": [
                {"path": "poc/jwt-forge-poc.py", "purpose": "JWT 伪造 PoC"},
                {"path": "poc/forged-token-response.json", "purpose": "认证绕过响应证据"},
            ],
            "bundle_root_artifacts": [
                {
                    "generator": "reviewer-recording-shell",
                    "output_name": "run-selftest-jwt-recording.sh",
                    "purpose": "审核复现脚本",
                    "generator_options": {"modes": ["quick"]},
                }
            ],
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        run([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(standard_fixture),
            "--output-dir",
            str(workspace / "confirmed"),
            "--language",
            "zh-CN",
        ], plugin_root)
        standard_bundle = next(
            (
                path for path in (workspace / "confirmed").iterdir()
                if path.is_dir() and "硬编码" in path.name and "安全漏洞" not in path.name
            ),
            None,
        )
        if standard_bundle is None:
            raise SystemExit("FAILED: standard vulnerability_name fixture did not render a finding-specific bundle name")
        standard_docx = next(standard_bundle.glob("*.docx"))
        write_live_replay_log(standard_bundle)
        nested_compose = standard_bundle / "attachments/poc/docker-compose.selftest.yml"
        nested_compose.parent.mkdir(parents=True, exist_ok=True)
        nested_compose.write_text("services:\n  app:\n    image: alpine:3.20\n", encoding="utf-8")
        if "硬编码" not in standard_docx.name or "安全漏洞" in standard_docx.name:
            raise SystemExit("FAILED: standard vulnerability_name fixture rendered a generic DOCX filename")
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(standard_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        exercise_variant_seed_confirmed_bundle_gate(plugin_root, workspace, standard_bundle)
        exercise_find_variant_candidates(plugin_root, workspace, standard_bundle)
        standard_lines = docx_text(standard_docx)
        if not standard_lines or "硬编码 JWT 密钥导致身份认证绕过" not in standard_lines[0]:
            raise SystemExit("FAILED: standard vulnerability_name fixture rendered a generic DOCX title")
        if "最终判定待补充" in "\n".join(standard_lines):
            raise SystemExit("FAILED: standard vulnerability_name fixture left final verdict placeholder text")
        with zipfile.ZipFile(standard_docx) as archive:
            standard_docx_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        if "Courier New" not in standard_docx_xml or 'w:sz w:val="17"' not in standard_docx_xml:
            raise SystemExit("FAILED: generated DOCX code context must use compact 8-9 pt monospace snippets")
        if "IntenseQuote" in standard_docx_xml:
            raise SystemExit("FAILED: generated DOCX code context must not use oversized quote/display code blocks")
        standard_text = "\n".join(standard_lines)
        for label in ("攻击者条件", "服务端条件", "安全影响"):
            if label not in standard_text:
                raise SystemExit(f"FAILED: zh-CN confirmed report is missing quality-gate label: {label}")
        if "实际场景中的危害与利用方式" not in standard_text:
            raise SystemExit("FAILED: zh-CN confirmed report is missing real-world exploitability section")
        en_docx = next(en_bundle.glob("*.docx"))
        en_text = "\n".join(docx_text(en_docx))
        for label in ("Attacker Condition", "Server Condition", "Security Impact"):
            if label not in en_text:
                raise SystemExit(f"FAILED: en-US confirmed report is missing quality-gate label: {label}")
        if "Real-World Exploitability" not in en_text:
            raise SystemExit("FAILED: en-US confirmed report is missing real-world exploitability section")
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(en_bundle),
            "--language",
            "en-US",
        ], plugin_root)
        en_scripts = sorted(en_bundle.glob("run-*.sh"))
        if en_scripts:
            en_script_text = en_scripts[0].read_text(encoding="utf-8")
            for expected in ("print_target_identity", "Software and Version", "Tested Software", "Tested Version / Branch"):
                if expected not in en_script_text:
                    raise SystemExit(f"FAILED: generated en-US recording script is missing target identity marker: {expected}")
        standard_script = standard_bundle / "run-selftest-jwt-recording.sh"
        standard_script_text = standard_script.read_text(encoding="utf-8")
        if "announce_step '代码'" not in standard_script_text or re.search(
            r"(?:announce_step\s+['\"]0/\d+|\[0/\d+\]|Step\s+0/\d+|步骤\s*0/\d+)",
            standard_script_text,
        ):
            raise SystemExit("FAILED: generated zh-CN recording script must use [代码], not 0/N, for code hints")
        for expected in (
            "print_target_identity",
            "软件名称和版本号",
            "测试软件名称",
            "测试版本/分支",
            "gothinkster/node-express-realworld-example-app",
            "default configuration",
            "REPLAY_LOG",
            "REPLAY_LOG_REL=\"attachments/evidence/replay-output.log\"",
            "replay-output.log",
            "READY_WAIT_SECONDS=\"${ZHULONG_READY_WAIT_SECONDS:-1}\"",
            "READY_RETRY_COUNT=\"${ZHULONG_READY_RETRY_COUNT:-30}\"",
            "ready_sleep()",
            "run_logged_command",
            "verify_success_marker",
            "SUCCESS_MARKER",
            "DIRECT_IMPACT_MARKER",
            "DIRECT_IMPACT_CONFIRMED",
            "record_direct_impact_marker",
            "> \"$command_output\" 2>&1",
            "cat \"$command_output\" >> \"$REPLAY_LOG\"",
            "grep -Fq -- \"$marker\" \"$REPLAY_LOG\"",
            "show_evidence_summary",
            "代码上下文屏",
            "show_vulnerability_analysis",
            "代码级漏洞分析屏",
            "攻击者可控输入/前提",
            "缺失 guard / validation",
            "show_real_world_context",
            "真实利用与影响边界屏",
            "已验证影响边界",
            "漏洞已确认",
        ):
            if expected not in standard_script_text:
                raise SystemExit(f"FAILED: generated recording script is missing target identity marker: {expected}")
        if re.search(
            r"READY_[A-Z0-9_]+\s*=\s*['\"]?\$?\{?(?:REVIEWER_PAUSE_SHORT|REVIEWER_PAUSE_LONG|PAUSE_SHORT|PAUSE_LONG)\}?",
            standard_script_text,
        ):
            raise SystemExit("FAILED: generated recording script derives functional readiness waits from reviewer pause variables")
        quick_case_match = re.search(r"quick\)\n(?P<body>.*?)\n\s*;;", standard_script_text, re.DOTALL)
        if not quick_case_match:
            raise SystemExit("FAILED: generated quick mode case is missing from recording script")
        quick_case_body = quick_case_match.group("body")
        if "READY_WAIT_SECONDS" in quick_case_body or "READY_RETRY_COUNT" in quick_case_body:
            raise SystemExit("FAILED: quick mode must not rewrite functional readiness wait variables")
        if "find \"$ATTACH_DIR\" -maxdepth 4 -type f" not in standard_script_text:
            raise SystemExit("FAILED: generated recording script must discover Compose files under nested attachments/")
        standard_evidence = json.loads((standard_bundle / "verification-evidence.json").read_text(encoding="utf-8"))
        if "attachments/evidence/replay-output.log" not in standard_evidence.get("evidence_files", []):
            raise SystemExit("FAILED: generated verification evidence must register replay-output.log")
        if not (standard_bundle / "attachments/evidence/replay-output.log").exists():
            raise SystemExit("FAILED: generated confirmed bundle must include replay-output.log")

        def copy_standard_bundle(suffix: str) -> Path:
            copied = standard_bundle.parent / f"{standard_bundle.name}_{suffix}"
            if copied.exists():
                shutil.rmtree(copied)
            shutil.copytree(standard_bundle, copied)
            return copied

        def write_replay_manifest(bundle: Path, *, with_provenance: bool = True, source_kind: str = "copied_successful_transcript") -> None:
            payload = {
                "schema_version": 1,
                "validation_status": "passed",
                "promote_status": "promoted",
                "replay_logs": [
                    {
                        "path": "attachments/evidence/replay-output.log",
                        "source_kind": source_kind,
                        "trust_classification": "trusted_transcript",
                        "sha256": "0" * 64,
                        "notes": "Wrapper did not execute replay; transcript was validated from bundled evidence.",
                    }
                ],
            }
            if with_provenance:
                payload["replay_logs"][0]["source_path"] = "confirmed/.staging/.inputs/selftest.renderer-input.json"
                payload["replay_logs"][0]["provenance"] = "Copied from the selected renderer input evidence."
            (bundle / "bundle-build-manifest.json").write_text(
                json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )

        bad_placeholder_replay_log = copy_standard_bundle("placeholder_replay_log")
        (bad_placeholder_replay_log / "attachments/evidence/replay-output.log").write_text(
            "Zhulong reviewer replay log placeholder.\n"
            "Run the bundle-root replay script to refresh this file with live reviewer output.\n"
            "Replay contract direct-impact marker: DIRECT_IMPACT_CONFIRMED\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_replay_log),
            "--language",
            "zh-CN",
        ], plugin_root, "placeholder-only")

        good_copied_replay_log = copy_standard_bundle("copied_replay_with_provenance")
        write_replay_manifest(good_copied_replay_log, with_provenance=True)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_copied_replay_log),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_copied_replay_log)

        replay_log_untrusted_cases = {
            "empty": "",
            "marker_only": "DIRECT_IMPACT_CONFIRMED\n",
            "heading_marker_only": "Zhulong reviewer replay log\nGenerated at: 2026-06-16T00:00:00Z\nDIRECT_IMPACT_CONFIRMED\n",
            "explanatory_only": (
                "This file explains that the reviewer should run the bundle-root replay script.\n"
                "When the command succeeds it should print DIRECT_IMPACT_CONFIRMED.\n"
            ),
            "marker_without_transcript_signals": (
                "Zhulong reviewer replay log\n"
                "DIRECT_IMPACT_CONFIRMED\n"
                "The exploit should have succeeded here, but no command or raw output is included.\n"
            ),
        }
        for suffix, log_text in replay_log_untrusted_cases.items():
            bad_replay_log = copy_standard_bundle(f"untrusted_replay_log_{suffix}")
            (bad_replay_log / "attachments/evidence/replay-output.log").write_text(log_text, encoding="utf-8")
            run_expect_fail([
                sys.executable,
                str(plugin_root / "scripts/validate_report_bundle.py"),
                "--bundle-dir",
                str(bad_replay_log),
                "--language",
                "zh-CN",
            ], plugin_root, "placeholder-only or marker-only")
            shutil.rmtree(bad_replay_log)

        bad_copied_without_provenance = copy_standard_bundle("copied_replay_without_provenance")
        write_replay_manifest(bad_copied_without_provenance, with_provenance=False)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_copied_without_provenance),
            "--language",
            "zh-CN",
        ], plugin_root, "lacks copied transcript provenance")
        shutil.rmtree(bad_copied_without_provenance)

        good_ssrf_callback_bounded = copy_standard_bundle("ssrf_callback_bounded")
        ssrf_callback_supplement = next(good_ssrf_callback_bounded.glob("*_补充复现说明.md"))
        ssrf_callback_supplement.write_text(
            ssrf_callback_supplement.read_text(encoding="utf-8")
            + "\n\nSSRF 补充证据：监听器收到请求，METADATA HIT，request received；"
            "本包仅声称服务端出站请求可达和 callback observed，不声称响应内容、配置、凭据或敏感数据泄露。\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_ssrf_callback_bounded),
            "--language",
            "zh-CN",
        ], plugin_root)

        bad_ssrf_callback_overclaim = copy_standard_bundle("ssrf_callback_overclaim")
        ssrf_overclaim_supplement = next(bad_ssrf_callback_overclaim.glob("*_补充复现说明.md"))
        ssrf_overclaim_supplement.write_text(
            ssrf_overclaim_supplement.read_text(encoding="utf-8")
            + "\n\nSSRF 补充证据：监听器收到请求，METADATA HIT，request received；"
            "报告同时声称响应内容外显、配置泄露和凭据泄露，但没有响应内容 oracle token。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_ssrf_callback_overclaim),
            "--language",
            "zh-CN",
        ], plugin_root, "SSRF impact overclaim")

        good_ssrf_response_exposure = copy_standard_bundle("ssrf_response_exposure")
        ssrf_good_supplement = next(good_ssrf_response_exposure.glob("*_补充复现说明.md"))
        ssrf_good_supplement.write_text(
            ssrf_good_supplement.read_text(encoding="utf-8")
            + "\n\nSSRF 直接危害证据：内部 metadata response content 进入目标输出，"
            "INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED，并由 DIRECT_IMPACT_CONFIRMED replay 标记界定。\n",
            encoding="utf-8",
        )
        write_live_replay_log(
            good_ssrf_response_exposure,
            extra="INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED: internal response content exposed in target output\n",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_ssrf_response_exposure),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_ssrf_response_exposure)

        bad_root_relative_artifact = copy_standard_bundle("root_relative_attachment_artifact")
        nested_bad_compose = bad_root_relative_artifact / "attachments/poc/docker-compose.root-path.yml"
        nested_bad_compose.parent.mkdir(parents=True, exist_ok=True)
        nested_bad_compose.write_text("services:\n  app:\n    image: alpine:3.20\n", encoding="utf-8")
        bad_root_script = bad_root_relative_artifact / "run-selftest-jwt-recording.sh"
        bad_root_script.write_text(
            bad_root_script.read_text(encoding="utf-8")
            + "\nprintf '%s\\n' 'docker compose -f docker-compose.root-path.yml up --abort-on-container-exit'\n"
            + "run_logged_command 'docker compose -f docker-compose.root-path.yml up --abort-on-container-exit'\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_root_relative_artifact),
            "--language",
            "zh-CN",
        ], plugin_root, "references root-relative Compose path")

        def mutate_bundle_finding(bundle: Path, mutator) -> None:
            findings_path = bundle / "findings.json"
            if not findings_path.exists():
                shared_findings_path = bundle.parent / "findings.json"
                source_findings_path = shared_findings_path if shared_findings_path.exists() else standard_fixture
                shared_data = json.loads(source_findings_path.read_text(encoding="utf-8"))
                if isinstance(shared_data, dict) and isinstance(shared_data.get("findings"), list):
                    candidates = [item for item in shared_data["findings"] if isinstance(item, dict)]
                    selected = next(
                        (
                            item for item in candidates
                            if str(item.get("slug") or "").strip() == standard_bundle.name
                            or Path(str(item.get("filename") or item.get("report_file") or "")).stem == standard_bundle.name
                            or "硬编码" in str(item.get("vulnerability_name") or item.get("vulnerability_name_zh") or "")
                        ),
                        candidates[0] if candidates else None,
                    )
                elif isinstance(shared_data, dict):
                    selected = shared_data
                elif isinstance(shared_data, list) and shared_data and isinstance(shared_data[0], dict):
                    selected = shared_data[0]
                else:
                    raise SystemExit("FAILED: shared selftest findings.json has unexpected shape")
                if selected is None:
                    raise SystemExit("FAILED: shared selftest findings.json has no finding objects")
                findings_path.write_text(json.dumps(selected, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            data = json.loads(findings_path.read_text(encoding="utf-8"))
            if isinstance(data, dict) and isinstance(data.get("findings"), list):
                target = data["findings"][0]
            elif isinstance(data, dict):
                target = data
            else:
                raise SystemExit("FAILED: selftest bundle findings.json has unexpected shape")
            mutator(target)
            findings_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

        neutral_pr_bundle = copy_standard_bundle("prl_neutral_title_pass")
        mutate_bundle_finding(
            neutral_pr_bundle,
            lambda finding: finding.setdefault("cvss", {}).update({"vector": "CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:C/C:H/I:H/A:L"}),
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(neutral_pr_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(neutral_pr_bundle)

        bad_unauth_title = copy_standard_bundle("unauthenticated_title_prl")
        mutate_bundle_finding(
            bad_unauth_title,
            lambda finding: finding.setdefault("cvss", {}).update({"vector": "CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:C/C:H/I:H/A:L"}),
        )
        rewrite_docx_paragraphs(
            next(bad_unauth_title.glob("*.docx")),
            lambda text: (
                "gothinkster/node-express-realworld-example-app Unauthenticated 硬编码 JWT 密钥导致身份认证绕过 严重漏洞报告"
                if text.startswith("gothinkster/node-express-realworld-example-app")
                else text
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_unauth_title),
            "--language",
            "zh-CN",
        ], plugin_root, "title/CVSS/auth consistency failure")
        shutil.rmtree(bad_unauth_title)

        bad_zero_step = copy_standard_bundle("recording_zero_step")
        zero_script = bad_zero_step / "run-selftest-jwt-recording.sh"
        zero_script.write_text(
            zero_script.read_text(encoding="utf-8").replace("announce_step '代码'", "announce_step '0/2'"),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_zero_step),
            "--language",
            "zh-CN",
        ], plugin_root, "must not use 0/N step labels")
        shutil.rmtree(bad_zero_step)

        bad_unconditional_poc = copy_standard_bundle("unconditional_poc_confirmation")
        bad_confirm_script = bad_unconditional_poc / "attachments/unconditional-confirm.sh"
        bad_confirm_script.write_text(
            "#!/bin/sh\nset -eu\necho \"VULNERABILITY CONFIRMED\"\n",
            encoding="utf-8",
        )
        bad_confirm_script.chmod(0o755)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_unconditional_poc),
            "--language",
            "zh-CN",
        ], plugin_root, "without a nearby or preceding concrete success oracle")
        shutil.rmtree(bad_unconditional_poc)

        good_conditional_poc = copy_standard_bundle("conditional_poc_confirmation")
        good_confirm_script = good_conditional_poc / "attachments/conditional-confirm.sh"
        good_confirm_script.write_text(
            "#!/bin/sh\n"
            "set -eu\n"
            "RESPONSE='{\"ok\":true,\"status_code\":200}'\n"
            "if printf '%s' \"$RESPONSE\" | grep -q '\"ok\":true'; then\n"
            "  echo \"VULNERABILITY CONFIRMED\"\n"
            "else\n"
            "  echo \"$RESPONSE\"\n"
            "  exit 1\n"
            "fi\n",
            encoding="utf-8",
        )
        good_confirm_script.chmod(0o755)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_conditional_poc),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_conditional_poc)

        bad_root_syntax = copy_standard_bundle("root_script_syntax_error")
        bad_root_syntax_script = bad_root_syntax / "run-selftest-jwt-recording.sh"
        bad_root_syntax_script.write_text(
            bad_root_syntax_script.read_text(encoding="utf-8") + "\nif then\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_root_syntax),
            "--language",
            "zh-CN",
        ], plugin_root, "shell syntax check failed")
        shutil.rmtree(bad_root_syntax)

        bad_grep_echo_oracle = copy_standard_bundle("fail_open_grep_echo_oracle")
        bad_grep_echo_script = bad_grep_echo_oracle / "run-selftest-jwt-recording.sh"
        bad_grep_echo_script.write_text(
            bad_grep_echo_script.read_text(encoding="utf-8")
            + "\nprintf 'missing\\n' | grep --color=always '认证绕过成功' || echo '未检测到成功判据'\n"
            + "echo -e \"${G}═══ 漏洞已确认：gothinkster/node-express-realworld-example-app 硬编码 JWT 密钥导致身份认证绕过 ═══${N}\"\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_grep_echo_oracle),
            "--language",
            "zh-CN",
        ], plugin_root, "softens a success-oracle failure")
        shutil.rmtree(bad_grep_echo_oracle)

        bad_grep_true_oracle = copy_standard_bundle("fail_open_grep_true_oracle")
        bad_grep_true_script = bad_grep_true_oracle / "run-selftest-jwt-recording.sh"
        bad_grep_true_script.write_text(
            bad_grep_true_script.read_text(encoding="utf-8")
            + "\nprintf 'missing\\n' | grep -q '认证绕过成功' || true\n"
            + "printf '%s\\n' \"${G}VULNERABILITY CONFIRMED: gothinkster/node-express-realworld-example-app hardcoded JWT secret${N}\"\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_grep_true_oracle),
            "--language",
            "zh-CN",
        ], plugin_root, "softens a success-oracle failure")
        shutil.rmtree(bad_grep_true_oracle)

        bad_wrapper_confirmation = copy_standard_bundle("fail_open_wrapper_confirmation")
        bad_wrapper_script = bad_wrapper_confirmation / "run-selftest-jwt-recording.sh"
        bad_wrapper_script.write_text(
            bad_wrapper_script.read_text(encoding="utf-8")
            + "\nprintf 'missing\\n' | grep -q '认证绕过成功' || echo '未检测到成功判据'\n"
            + "highlight_success \"${G}═══ 漏洞已确认：gothinkster/node-express-realworld-example-app 硬编码 JWT 密钥导致身份认证绕过 ═══${N}\"\n"
            + "print_banner \"${G}VULNERABILITY CONFIRMED: hardcoded JWT secret${N}\"\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_wrapper_confirmation),
            "--language",
            "zh-CN",
        ], plugin_root, "softens a success-oracle failure")
        shutil.rmtree(bad_wrapper_confirmation)

        good_fail_closed_oracle = copy_standard_bundle("fail_closed_oracle")
        good_fail_closed_script = good_fail_closed_oracle / "run-selftest-jwt-recording.sh"
        good_fail_closed_script.write_text(
            good_fail_closed_script.read_text(encoding="utf-8")
            + "\nif ! printf '认证绕过成功\\n' | grep -q '认证绕过成功'; then\n"
            + "  echo '未检测到成功判据，不能确认漏洞。'\n"
            + "  exit 1\n"
            + "fi\n"
            + "echo '漏洞已确认：gothinkster/node-express-realworld-example-app 硬编码 JWT 密钥导致身份认证绕过'\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_fail_closed_oracle),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_fail_closed_oracle)

        bad_missing_compose_env = copy_standard_bundle("compose_missing_env_file")
        bad_missing_compose_env_file = bad_missing_compose_env / "attachments/docker-compose.zhulong.yml"
        bad_missing_compose_env_file.write_text(
            "services:\n"
            "  app:\n"
            "    image: alpine:3.20\n"
            "    env_file: .env\n"
            "    command: ['sh', '-c', 'sleep 1']\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_compose_env),
            "--language",
            "zh-CN",
        ], plugin_root, "env_file")
        shutil.rmtree(bad_missing_compose_env)

        bad_missing_compose_bind = copy_standard_bundle("compose_missing_bind_source")
        (bad_missing_compose_bind / "attachments/.env").write_text("SELFTEST=1\n", encoding="utf-8")
        (bad_missing_compose_bind / "attachments/docker-compose.zhulong.yml").write_text(
            "services:\n"
            "  app:\n"
            "    image: alpine:3.20\n"
            "    env_file:\n"
            "      - .env\n"
            "    volumes:\n"
            "      - ./missing-listener.py:/scripts/listener.py:ro\n"
            "    command: ['sh', '-c', 'sleep 1']\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_compose_bind),
            "--language",
            "zh-CN",
        ], plugin_root, "volume source")
        shutil.rmtree(bad_missing_compose_bind)

        good_compose_bundle = copy_standard_bundle("compose_existing_env_bind_and_named_volume")
        (good_compose_bundle / "attachments/.env").write_text("SELFTEST=1\n", encoding="utf-8")
        (good_compose_bundle / "attachments/listen.py").write_text("print('ready')\n", encoding="utf-8")
        (good_compose_bundle / "attachments/docker-compose.zhulong.yml").write_text(
            "services:\n"
            "  app:\n"
            "    image: alpine:3.20\n"
            "    env_file:\n"
            "      - .env\n"
            "    volumes:\n"
            "      - ./listen.py:/scripts/listen.py:ro\n"
            "      - named-cache:/cache\n"
            "      - type: bind\n"
            "        source: ./listen.py\n"
            "        target: /scripts/listen-copy.py\n"
            "        read_only: true\n"
            "    command: ['sh', '-c', 'sleep 1']\n"
            "volumes:\n"
            "  named-cache: {}\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_compose_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_compose_bundle)

        bad_english_docx = copy_standard_bundle("untranslated_english_docx")
        rewrite_docx_paragraphs(
            next(bad_english_docx.glob("*.docx")),
            lambda text: (
                "The vulnerable endpoint accepts attacker controlled input and forwards it to a sensitive server side operation. This vulnerability allows an attacker to trigger a security impact through the default configuration."
                if text.startswith("默认配置缺失 JWT_SECRET")
                else text
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_english_docx),
            "--language",
            "zh-CN",
        ], plugin_root, "long English natural-language paragraph")
        shutil.rmtree(bad_english_docx)

        bad_english_supplement = copy_standard_bundle("untranslated_english_supplement")
        bad_supplement_path = next(bad_english_supplement.glob("*_补充复现说明.md"))
        bad_supplement_path.write_text(
            bad_supplement_path.read_text(encoding="utf-8")
            + "\nThe vulnerable endpoint accepts attacker controlled input and forwards it to a sensitive server side operation. This vulnerability allows an attacker to trigger a security impact through the default configuration.\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_english_supplement),
            "--language",
            "zh-CN",
        ], plugin_root, "long English natural-language paragraph")
        shutil.rmtree(bad_english_supplement)

        bad_variant_candidate_confirmed = copy_standard_bundle("variant_candidate_confirmed_text")
        bad_variant_supplement = next(bad_variant_candidate_confirmed.glob("*_补充复现说明.md"))
        bad_variant_supplement.write_text(
            bad_variant_supplement.read_text(encoding="utf-8")
            + "\n变体候选已确认：该条结论来自已确认种子的复现模式，直接复用相同结论即可。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_variant_candidate_confirmed),
            "--language",
            "zh-CN",
        ], plugin_root, "must not assert a variant/同类候选已确认")
        shutil.rmtree(bad_variant_candidate_confirmed)

        safe_variant_candidate_wording = copy_standard_bundle("safe_variant_candidate_wording")
        safe_variant_supplement = next(safe_variant_candidate_wording.glob("*_补充复现说明.md"))
        safe_variant_supplement.write_text(
            safe_variant_supplement.read_text(encoding="utf-8")
            + "\n安全说明：variant candidate requires Docker verification before confirmation，不能作为确认依据。\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(safe_variant_candidate_wording),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(safe_variant_candidate_wording)

        bad_variant_candidates_in_bundle = copy_standard_bundle("variant_candidates_jsonl_in_bundle")
        (bad_variant_candidates_in_bundle / "attachments/evidence/variant-candidates.jsonl").write_text(
            json.dumps(valid_variant_candidate_record(), ensure_ascii=False, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_variant_candidates_in_bundle),
            "--language",
            "zh-CN",
        ], plugin_root, "must not include variant-candidates.jsonl")
        shutil.rmtree(bad_variant_candidates_in_bundle)

        bad_variant_candidates_proof = copy_standard_bundle("variant_candidates_jsonl_proof")
        bad_variant_proof_supplement = next(bad_variant_candidates_proof.glob("*_补充复现说明.md"))
        bad_variant_proof_supplement.write_text(
            bad_variant_proof_supplement.read_text(encoding="utf-8")
            + "\n补充：variant-candidates.jsonl 证明该变体已确认，可作为确认依据。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_variant_candidates_proof),
            "--language",
            "zh-CN",
        ], plugin_root, "variant candidate output must not be used as confirmed evidence")
        shutil.rmtree(bad_variant_candidates_proof)

        bad_variant_seed_similarity_confirmed = copy_standard_bundle("variant_seed_similarity_confirmed")
        bad_variant_similarity_supplement = next(bad_variant_seed_similarity_confirmed.glob("*_补充复现说明.md"))
        bad_variant_similarity_supplement.write_text(
            bad_variant_similarity_supplement.read_text(encoding="utf-8")
            + "\n审核说明：变体候选因为与种子相似已确认。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_variant_seed_similarity_confirmed),
            "--language",
            "zh-CN",
        ], plugin_root, "variant candidate output must not be used as confirmed evidence")
        shutil.rmtree(bad_variant_seed_similarity_confirmed)

        technical_english_ok = copy_standard_bundle("technical_english_ok")
        technical_supplement_path = next(technical_english_ok.glob("*_补充复现说明.md"))
        technical_supplement_path.write_text(
            technical_supplement_path.read_text(encoding="utf-8")
            + "\n```sh\ncurl -s http://localhost:3000/api/user -H 'Authorization: Bearer TOKEN'\n```\n"
            + "`{\"access_token\":\"TOKEN\",\"status_code\":200}`\n"
            + "`CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:C/C:H/I:H/A:L`\n"
            + "`VULNERABILITY CONFIRMED`\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(technical_english_ok),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(technical_english_ok)

        def quality_gate_bad_bundle(suffix: str, replacer) -> Path:
            bad_bundle = standard_bundle.parent / f"{standard_bundle.name}_{suffix}"
            shutil.copytree(standard_bundle, bad_bundle)
            rewrite_docx_paragraphs(next(bad_bundle.glob("*.docx")), replacer)
            return bad_bundle

        def remove_docx_real_world_section(docx_path: Path, *, language: str) -> None:
            stop_prefixes = ("位置：", "Location:")
            headings = (
                ("实际场景中的危害与利用方式",)
                if language == "zh-CN"
                else ("Real-World Exploitability",)
            )
            removing = False

            def replacer(text: str):
                nonlocal removing
                if any(heading in text for heading in headings):
                    removing = True
                    return None
                if removing and text.startswith(stop_prefixes):
                    removing = False
                    return text
                if removing:
                    return None
                return text

            rewrite_docx_paragraphs(docx_path, replacer)

        def remove_markdown_real_world_section(path: Path, *, language: str) -> None:
            text = path.read_text(encoding="utf-8")
            heading_fragment = "实际场景中的危害与利用方式" if language == "zh-CN" else "Practical Impact and Exploitation Path"
            kept: list[str] = []
            skipping = False
            for line in text.splitlines():
                if line.startswith("## ") and heading_fragment in line:
                    skipping = True
                    continue
                if skipping and line.startswith("## "):
                    skipping = False
                if not skipping:
                    kept.append(line)
            path.write_text("\n".join(kept) + "\n", encoding="utf-8")

        def weaken_markdown_success_evidence(path: Path, *, language: str) -> None:
            text = path.read_text(encoding="utf-8")
            evidence_heading = "关键成功证据" if language == "zh-CN" else "Key Success Evidence"
            replacement = "- 技术触发完成。" if language == "zh-CN" else "- Technical trigger completed."
            kept: list[str] = []
            in_evidence = False
            inserted = False
            for line in text.splitlines():
                if line.startswith("## ") and evidence_heading in line:
                    in_evidence = True
                    inserted = False
                    kept.append(line)
                    continue
                if in_evidence and line.startswith("## "):
                    if not inserted:
                        kept.append(replacement)
                    in_evidence = False
                if in_evidence:
                    if not inserted and line.strip():
                        kept.append(replacement)
                        inserted = True
                    continue
                kept.append(line)
            if in_evidence and not inserted:
                kept.append(replacement)
            path.write_text("\n".join(kept) + "\n", encoding="utf-8")

        def weaken_real_world_fallback_docx(docx_path: Path) -> None:
            protected_prefixes = (
                "漏洞描述",
                "影响版本",
                "漏洞危险性评估",
                "漏洞分析",
                "漏洞复现",
                "最终判定：",
                "攻击者条件",
                "服务端条件",
                "安全影响",
            )
            weak_markers = ("返回", "响应", "输出", "证据", "Docker", "验证", "确认", "成功", "记录")

            rewrite_docx_paragraphs(
                docx_path,
                lambda text: (
                    "攻击者条件：攻击者能够访问测试入口。"
                    if text.startswith("攻击者条件")
                    else "服务端条件：服务端运行测试组件。"
                    if text.startswith("服务端条件")
                    else "安全影响：存在机密性影响。"
                    if text.startswith("安全影响")
                    else None
                    if text.startswith("结果证据：") or text.startswith("实际结果：")
                    else "技术触发完成。"
                    if any(marker in text for marker in weak_markers) and not text.startswith(protected_prefixes)
                    else text
                ),
            )

        def weaken_en_real_world_fallback_docx(docx_path: Path) -> None:
            protected_prefixes = (
                "Vulnerability Description",
                "Affected Versions",
                "Risk Assessment",
                "Vulnerability Analysis",
                "Reproduction",
                "Final Verdict:",
                "Attacker Condition",
                "Server Condition",
                "Security Impact",
            )
            weak_markers = ("response", "output", "evidence", "docker", "verified", "confirmed", "success", "record")

            rewrite_docx_paragraphs(
                docx_path,
                lambda text: (
                    "Attacker Condition: attacker can reach the test entry point."
                    if text.startswith("Attacker Condition")
                    else "Server Condition: server runs the tested component."
                    if text.startswith("Server Condition")
                    else "Security Impact: confidentiality impact exists."
                    if text.startswith("Security Impact")
                    else None
                    if text.startswith("Evidence:") or text.startswith("Observed result:")
                    else "Technical trigger completed."
                    if any(marker in text.lower() for marker in weak_markers) and not text.startswith(protected_prefixes)
                    else text
                ),
            )

        def remove_docx_section(docx_path: Path, heading: str, stop_headings: set[str]) -> None:
            in_section = {"value": False}

            def replacer(text: str):
                if text == heading:
                    in_section["value"] = True
                    return None
                if in_section["value"] and text in stop_headings:
                    in_section["value"] = False
                    return text
                if in_section["value"]:
                    return None
                return text

            rewrite_docx_paragraphs(docx_path, replacer)

        def replace_docx_section_with_one_line(
            docx_path: Path,
            heading: str,
            stop_headings: set[str],
            replacement: str,
        ) -> None:
            state = {"in_section": False, "inserted": False}

            def replacer(text: str):
                if text == heading:
                    state["in_section"] = True
                    state["inserted"] = False
                    return text
                if state["in_section"] and text in stop_headings:
                    state["in_section"] = False
                    return text
                if state["in_section"]:
                    if not state["inserted"]:
                        state["inserted"] = True
                        return replacement
                    return None
                return text

            rewrite_docx_paragraphs(docx_path, replacer)

        bad_missing_real_world = copy_standard_bundle("missing_real_world_exploitability")
        remove_docx_real_world_section(next(bad_missing_real_world.glob("*.docx")), language="zh-CN")
        bad_missing_real_world_supplement = next(bad_missing_real_world.glob("*_补充复现说明.md"))
        remove_markdown_real_world_section(bad_missing_real_world_supplement, language="zh-CN")
        weaken_markdown_success_evidence(bad_missing_real_world_supplement, language="zh-CN")
        weaken_real_world_fallback_docx(next(bad_missing_real_world.glob("*.docx")))
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_real_world),
            "--language",
            "zh-CN",
        ], plugin_root, "VALIDATION FAILED")

        bad_en_missing_real_world = en_bundle.parent / f"{en_bundle.name}_missing_real_world_exploitability"
        if bad_en_missing_real_world.exists():
            shutil.rmtree(bad_en_missing_real_world)
        shutil.copytree(en_bundle, bad_en_missing_real_world)
        remove_docx_real_world_section(next(bad_en_missing_real_world.glob("*.docx")), language="en-US")
        bad_en_missing_real_world_supplement = next(bad_en_missing_real_world.glob("*_reproduction_note.md"))
        remove_markdown_real_world_section(bad_en_missing_real_world_supplement, language="en-US")
        weaken_markdown_success_evidence(bad_en_missing_real_world_supplement, language="en-US")
        weaken_en_real_world_fallback_docx(next(bad_en_missing_real_world.glob("*.docx")))
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_en_missing_real_world),
            "--language",
            "en-US",
        ], plugin_root, "VALIDATION FAILED")

        code_context_stop_headings = {"漏洞复现", "验证环境关键文件"}

        bad_missing_code_context = copy_standard_bundle("missing_code_context")
        remove_docx_section(next(bad_missing_code_context.glob("*.docx")), "关键代码上下文", code_context_stop_headings)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "missing or empty")

        bad_prose_code_context = copy_standard_bundle("prose_only_code_context")
        replace_docx_section_with_one_line(
            next(bad_prose_code_context.glob("*.docx")),
            "关键代码上下文",
            code_context_stop_headings,
            (
                "路径 src/app/routes/auth/auth.ts 行 18 到 26。攻击者可控输入会沿认证中间件传播到危险校验 sink，"
                "缺失 guard 是没有强制验证密钥配置；相邻校验不足以阻断公开密钥签名。Docker 已验证影响边界为认证绕过。"
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_prose_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "actual code-like snippet")

        bad_placeholder_code_context = copy_standard_bundle("placeholder_code_context")
        replace_docx_section_with_one_line(
            next(bad_placeholder_code_context.glob("*.docx")),
            "关键代码上下文",
            code_context_stop_headings,
            "代码上下文 1",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "placeholder-only code context")

        bad_en_placeholder_code_context = en_bundle.parent / f"{en_bundle.name}_placeholder_code_context"
        if bad_en_placeholder_code_context.exists():
            shutil.rmtree(bad_en_placeholder_code_context)
        shutil.copytree(en_bundle, bad_en_placeholder_code_context)
        replace_docx_section_with_one_line(
            next(bad_en_placeholder_code_context.glob("*.docx")),
            "Key Code Context",
            {"Reproduction", "Key Verification Environment Files"},
            "Key Code Context 1",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_en_placeholder_code_context),
            "--language",
            "en-US",
        ], plugin_root, "placeholder-only code context")

        bad_abbreviated_code_context = copy_standard_bundle("abbreviated_code_context")
        replace_docx_section_with_one_line(
            next(bad_abbreviated_code_context.glob("*.docx")),
            "关键代码上下文",
            code_context_stop_headings,
            (
                "1. src/app/routes/auth/auth.ts:18-26\n"
                "const secret = process.env.JWT_SECRET || 'superSecret';\n"
                "// ...\n"
                "// ... omitted validation logic ...\n"
                "app.use(jwt({ secret, algorithms: ['HS256'] }));\n"
                "攻击者可控输入从 Authorization 头传播到认证 sink。缺失 guard 是启动阶段没有强制校验 JWT_SECRET；"
                "相邻校验不足以阻断公开密钥签名。Docker 已验证影响边界为认证绕过。"
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_abbreviated_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "too abbreviated")

        bad_blue_code_context = copy_standard_bundle("blue_display_code_context")
        from docx import Document
        from docx.shared import Pt, RGBColor

        blue_docx_path = next(bad_blue_code_context.glob("*.docx"))
        blue_doc = Document(blue_docx_path)
        in_code_context = False
        for paragraph in blue_doc.paragraphs:
            text = paragraph.text.strip()
            if text == "关键代码上下文":
                in_code_context = True
                continue
            if in_code_context and text in code_context_stop_headings:
                break
            if in_code_context and ("const secret" in text or "app.use(jwt" in text):
                paragraph.style = "Intense Quote"
                for docx_run in paragraph.runs:
                    docx_run.font.size = Pt(14)
                    docx_run.font.color.rgb = RGBColor(0, 0, 255)
        blue_doc.save(blue_docx_path)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_blue_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "small readable code font")

        bad_placeholder_real_world = copy_standard_bundle("placeholder_real_world_exploitability")
        rewrite_docx_paragraphs(
            next(bad_placeholder_real_world.glob("*.docx")),
            lambda text: (
                "实际场景中的危害与利用方式"
                if "实际场景中的危害与利用方式" in text
                else "待补充"
                if (
                    text.startswith("实际使用场景")
                    or text.startswith("攻击者路径")
                    or text.startswith("触发调用链")
                    or text.startswith("直接危害证明")
                    or text.startswith("影响边界")
                    or text.startswith("服务端可达条件")
                    or text.startswith("影响外显通道")
                    or text.startswith("已验证影响边界")
                )
                else text
            ),
        )
        remove_markdown_real_world_section(next(bad_placeholder_real_world.glob("*_补充复现说明.md")), language="zh-CN")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_real_world),
            "--language",
            "zh-CN",
        ], plugin_root, "real-world exploitability section is too thin")

        bad_strong_boundary = copy_standard_bundle("strong_attacker_boundary_missing")
        remove_docx_real_world_section(next(bad_strong_boundary.glob("*.docx")), language="zh-CN")
        remove_markdown_real_world_section(next(bad_strong_boundary.glob("*_补充复现说明.md")), language="zh-CN")
        poc_boundary_script = bad_strong_boundary / "attachments/strong-boundary.sh"
        poc_boundary_script.write_text(
            "#!/bin/sh\n"
            "set -eu\n"
            "# PoC uses malicious JS and directly executes it to reach the component.\n"
            "printf 'ok {\"ok\":true}\\n' | grep -q 'ok'\n"
            "echo 'VULNERABILITY CONFIRMED'\n",
            encoding="utf-8",
        )
        poc_boundary_script.chmod(0o755)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_strong_boundary),
            "--language",
            "zh-CN",
        ], plugin_root, "strong-attacker-control PoC wording")

        good_strong_boundary = copy_standard_bundle("strong_attacker_boundary_explained")
        good_boundary_script = good_strong_boundary / "attachments/strong-boundary.sh"
        good_boundary_script.write_text(
            "#!/bin/sh\n"
            "set -eu\n"
            "# PoC mentions malicious JS, but the report explains the component boundary.\n"
            "printf 'ok {\"ok\":true}\\n' | grep -q 'ok'\n"
            "echo 'VULNERABILITY CONFIRMED'\n",
            encoding="utf-8",
        )
        good_boundary_script.chmod(0o755)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_strong_boundary),
            "--language",
            "zh-CN",
        ], plugin_root)

        bad_bundle_escape = copy_standard_bundle("bundle_escape_root_script")
        escape_script = bad_bundle_escape / "run-selftest-jwt-recording.sh"
        escape_script.write_text(
            escape_script.read_text(encoding="utf-8")
            + "\nREPO_ROOT=\"$(cd \"$SCRIPT_DIR/../../../../..\" && pwd)\"\n"
            + "printf '%s\\n' \"$REPO_ROOT\" >/dev/null\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_bundle_escape),
            "--language",
            "zh-CN",
        ], plugin_root, "depends on a parent path outside the confirmed bundle")

        bad_pkg_dependency = copy_standard_bundle("pkg_index_dependency")
        pkg_script = bad_pkg_dependency / "run-selftest-jwt-recording.sh"
        pkg_script.write_text(
            pkg_script.read_text(encoding="utf-8")
            + "\nprintf '%s\\n' 'unsafe external checkout path /pkg/index.js /pkg/security-research-YYYYMMDD-HHMMSS'\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_pkg_dependency),
            "--language",
            "zh-CN",
        ], plugin_root, "non-standalone path text")

        bad_workspace_marker = copy_standard_bundle("workspace_marker_leak")
        workspace_supplement = next(bad_workspace_marker.glob("*_补充复现说明.md"))
        workspace_supplement.write_text(
            workspace_supplement.read_text(encoding="utf-8")
            + "\n错误示例：材料仍提到 submitter-workspace 和 security-research-YYYYMMDD-HHMMSS。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_workspace_marker),
            "--language",
            "zh-CN",
        ], plugin_root, "non-standalone path text")

        bad_missing_target_identity = copy_standard_bundle("missing_target_identity")
        missing_target_script = bad_missing_target_identity / "run-selftest-jwt-recording.sh"
        missing_target_script.write_text(
            missing_target_script.read_text(encoding="utf-8")
            .replace("print_target_identity() {", "print_identity_removed() {")
            .replace("    print_target_identity\n    pause_step \"$PAUSE_SHORT\"\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_target_identity),
            "--language",
            "zh-CN",
        ], plugin_root, "target software/package")

        bad_legacy_target_identity = copy_standard_bundle("legacy_target_identity_labels")
        legacy_target_script = bad_legacy_target_identity / "run-selftest-jwt-recording.sh"
        legacy_target_script.write_text(
            legacy_target_script.read_text(encoding="utf-8")
            .replace("软件名称和版本号", "目标信息")
            .replace("测试软件名称", "目标软件")
            .replace("测试版本/分支", "版本号"),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_legacy_target_identity),
            "--language",
            "zh-CN",
        ], plugin_root, "reviewer-facing software/version identity line")

        bad_no_opening_identity_pause = copy_standard_bundle("no_opening_identity_pause")
        no_opening_pause_script = bad_no_opening_identity_pause / "run-selftest-jwt-recording.sh"
        no_opening_pause_script.write_text(
            no_opening_pause_script.read_text(encoding="utf-8")
            .replace("    print_target_identity\n    pause_step \"$PAUSE_SHORT\"\n", "    print_target_identity\n", 1),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_opening_identity_pause),
            "--language",
            "zh-CN",
        ], plugin_root, "opening tested software/version identity screen")

        bad_late_target_identity = copy_standard_bundle("late_target_identity")
        late_target_script = bad_late_target_identity / "run-selftest-jwt-recording.sh"
        late_target_script.write_text(
            late_target_script.read_text(encoding="utf-8")
            .replace("    print_target_identity\n", "", 1)
            .replace("    show_code_hint\n", "    show_code_hint\n    print_target_identity\n", 1),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_late_target_identity),
            "--language",
            "zh-CN",
        ], plugin_root, "before proof steps")

        bad_no_replay_code_context = copy_standard_bundle("replay_no_code_context_screen")
        no_replay_code_context_script = bad_no_replay_code_context / "run-selftest-jwt-recording.sh"
        no_replay_code_context_script.write_text(
            no_replay_code_context_script.read_text(encoding="utf-8")
            .replace("show_code_hint() {", "show_code_hint_removed() {")
            .replace("    show_code_hint\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_replay_code_context),
            "--language",
            "zh-CN",
        ], plugin_root, "code-context display")

        bad_no_replay_analysis = copy_standard_bundle("replay_no_analysis_screen")
        no_replay_analysis_script = bad_no_replay_analysis / "run-selftest-jwt-recording.sh"
        no_replay_analysis_script.write_text(
            no_replay_analysis_script.read_text(encoding="utf-8")
            .replace("show_vulnerability_analysis() {", "show_vulnerability_analysis_removed() {")
            .replace("    show_vulnerability_analysis\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_replay_analysis),
            "--language",
            "zh-CN",
        ], plugin_root, "code-level vulnerability analysis display")

        bad_no_replay_real_world = copy_standard_bundle("replay_no_real_world_screen")
        no_replay_real_world_script = bad_no_replay_real_world / "run-selftest-jwt-recording.sh"
        no_replay_real_world_script.write_text(
            no_replay_real_world_script.read_text(encoding="utf-8")
            .replace("show_real_world_context() {", "show_real_world_context_removed() {")
            .replace("    show_real_world_context\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_replay_real_world),
            "--language",
            "zh-CN",
        ], plugin_root, "real-world exploitability / impact-boundary display")

        bad_no_replay_final_summary = copy_standard_bundle("replay_no_final_summary_screen")
        no_replay_final_summary_script = bad_no_replay_final_summary / "run-selftest-jwt-recording.sh"
        no_replay_final_summary_script.write_text(
            no_replay_final_summary_script.read_text(encoding="utf-8")
            .replace("    show_evidence_summary\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_replay_final_summary),
            "--language",
            "zh-CN",
        ], plugin_root, "final evidence summary after proof commands")

        bad_unavailable_replay_context = copy_standard_bundle("replay_unavailable_fallback")
        unavailable_script = bad_unavailable_replay_context / "run-selftest-jwt-recording.sh"
        unavailable_script.write_text(
            unavailable_script.read_text(encoding="utf-8").replace(
                "    show_vulnerability_analysis\n",
                "    printf '%s\\n' 'analysis_unavailable'\n    show_vulnerability_analysis\n",
                1,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_unavailable_replay_context),
            "--language",
            "zh-CN",
        ], plugin_root, "unavailable fallback text")

        bad_no_final_summary_pause = copy_standard_bundle("no_final_summary_pause")
        no_final_pause_script = bad_no_final_summary_pause / "run-selftest-jwt-recording.sh"
        no_final_pause_script.write_text(
            no_final_pause_script.read_text(encoding="utf-8")
            .replace("    show_evidence_summary\n    pause_step \"$PAUSE_LONG\"\n", "    show_evidence_summary\n"),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_final_summary_pause),
            "--language",
            "zh-CN",
        ], plugin_root, "final evidence summary")

        bad_undefined_root_helper = copy_standard_bundle("undefined_root_helper")
        undefined_helper_script = bad_undefined_root_helper / "run-selftest-jwt-recording.sh"
        undefined_helper_script.write_text(
            undefined_helper_script.read_text(encoding="utf-8")
            .replace("    show_code_hint\n", "    run_docker_poc\n    show_code_hint\n", 1),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_undefined_root_helper),
            "--language",
            "zh-CN",
        ], plugin_root, "not defined in the same root script")

        good_helper_closed = copy_standard_bundle("helper_closed")
        helper_closed_script = good_helper_closed / "run-selftest-jwt-recording.sh"
        helper_closed_script.write_text(
            helper_closed_script.read_text(encoding="utf-8")
            + "\nverify_closed_helper() {\n    return 0\n}\n\nrun_closed_helper() {\n    verify_closed_helper\n}\n\nrun_closed_helper\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_helper_closed),
            "--language",
            "zh-CN",
        ], plugin_root)
        shutil.rmtree(good_helper_closed)

        bad_missing_replay_log = copy_standard_bundle("missing_replay_log")
        missing_log_script = bad_missing_replay_log / "run-selftest-jwt-recording.sh"
        missing_log_script.write_text(
            missing_log_script.read_text(encoding="utf-8")
            .replace('REPLAY_LOG="$EVIDENCE_DIR/replay-output.log"', 'REPLAY_TEXT="$EVIDENCE_DIR/replay-output.txt"')
            .replace('> "$REPLAY_LOG"', '> "$REPLAY_TEXT"')
            .replace('>> "$REPLAY_LOG"', '>> "$REPLAY_TEXT"'),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_replay_log),
            "--language",
            "zh-CN",
        ], plugin_root, "must write reviewer replay output")

        bad_unregistered_replay_log = copy_standard_bundle("unregistered_replay_log")
        unregistered_evidence_path = bad_unregistered_replay_log / "verification-evidence.json"
        unregistered_data = json.loads(unregistered_evidence_path.read_text(encoding="utf-8"))
        unregistered_data["evidence_files"] = [
            item for item in unregistered_data.get("evidence_files", [])
            if item != "attachments/evidence/replay-output.log"
        ]
        unregistered_evidence_path.write_text(
            json.dumps(unregistered_data, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_unregistered_replay_log),
            "--language",
            "zh-CN",
        ], plugin_root, "must be registered")

        bad_final_without_marker_check = copy_standard_bundle("final_without_marker_check")
        final_without_marker_script = bad_final_without_marker_check / "run-selftest-jwt-recording.sh"
        final_without_marker_script.write_text(
            final_without_marker_script.read_text(encoding="utf-8")
            .replace("    verify_success_marker \"$SUCCESS_MARKER\"\n", ""),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_final_without_marker_check),
            "--language",
            "zh-CN",
        ], plugin_root, "programmatic success-marker check")

        bad_explanatory_marker_text = copy_standard_bundle("explanatory_marker_text")
        explanatory_marker_script = bad_explanatory_marker_text / "run-selftest-jwt-recording.sh"
        explanatory_marker_script.write_text(
            explanatory_marker_script.read_text(encoding="utf-8")
            .replace(
                "    verify_success_marker \"$SUCCESS_MARKER\"\n",
                "    printf '%s\\n' 'if output contains the success marker then the vulnerability is confirmed'\n",
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_explanatory_marker_text),
            "--language",
            "zh-CN",
        ], plugin_root, "programmatic success-marker check")

        bad_log_without_raw_output = copy_standard_bundle("replay_log_without_raw_output")
        log_without_raw_script = bad_log_without_raw_output / "run-selftest-jwt-recording.sh"
        log_without_raw_script.write_text(
            log_without_raw_script.read_text(encoding="utf-8")
            .replace(' > "$command_output" 2>&1', ' >/dev/null 2>/dev/null')
            .replace('        cat "$command_output" >> "$REPLAY_LOG"\n', ''),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_log_without_raw_output),
            "--language",
            "zh-CN",
        ], plugin_root, "raw command stdout/stderr")

        bad_missing_direct_impact = copy_standard_bundle("missing_direct_impact")
        missing_direct_script = bad_missing_direct_impact / "run-selftest-jwt-recording.sh"
        missing_direct_script.write_text(
            re.sub(
                r"\nrecord_direct_impact_marker\(\) \{\n(?:    .+\n)+\}\n",
                "\n",
                missing_direct_script.read_text(encoding="utf-8"),
            )
            .replace("DIRECT_IMPACT_MARKER='DIRECT_IMPACT_CONFIRMED'\n", "")
            .replace("    record_direct_impact_marker \"$DIRECT_IMPACT_MARKER\"\n", "")
            .replace("_CONFIRMED", "_CHECKED")
            .replace("认证绕过成功", "认证检查完成")
            .replace("会话伪造成功", "会话检查完成")
            .replace("direct impact is supported by the DIRECT_IMPACT_CONFIRMED-equivalent marker ", "direct impact is supported by the replay marker "),
            encoding="utf-8",
        )
        for text_path in [
            path for path in bad_missing_direct_impact.rglob("*")
            if path.is_file() and path.suffix in {".json", ".log", ".md", ".sh", ".txt"}
        ]:
            text_path.write_text(
                text_path.read_text(encoding="utf-8")
                .replace("DIRECT_IMPACT_CONFIRMED", "replay success marker")
                .replace("DIRECT_AVAILABILITY_IMPACT_CONFIRMED", "replay availability marker")
                .replace("_CONFIRMED", "_CHECKED")
                .replace("认证绕过成功", "认证检查完成")
                .replace("会话伪造成功", "会话检查完成"),
                encoding="utf-8",
            )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_direct_impact),
            "--language",
            "zh-CN",
        ], plugin_root, "direct-impact marker in replay evidence")

        bad_raw_structured_docx = copy_standard_bundle("raw_structured_docx")
        rewrite_docx_paragraphs(
            next(bad_raw_structured_docx.glob("*.docx")),
            lambda text: (
                "{'oracle_token': '认证绕过成功', 'impact': ['auth bypass'], 'confirmed': True}"
                if text.startswith("成功判据：")
                else text
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_raw_structured_docx),
            "--language",
            "zh-CN",
        ], plugin_root, "raw serialized structured data")

        bad_mutable_runtime_identity = copy_standard_bundle("mutable_runtime_identity")
        mutable_supplement = next(bad_mutable_runtime_identity.glob("*_补充复现说明.md"))
        mutable_supplement.write_text(
            mutable_supplement.read_text(encoding="utf-8")
            + "\n测试版本/分支：current version\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_mutable_runtime_identity),
            "--language",
            "zh-CN",
        ], plugin_root, "mutable runtime/version identity")

        bad_direct_marker_mismatch = copy_standard_bundle("direct_marker_mismatch")
        marker_evidence_path = bad_direct_marker_mismatch / "verification-evidence.json"
        marker_evidence = json.loads(marker_evidence_path.read_text(encoding="utf-8"))
        marker_evidence["direct_impact_marker"] = "DIRECT_AVAILABILITY_IMPACT_CONFIRMED"
        marker_evidence_path.write_text(
            json.dumps(marker_evidence, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_direct_marker_mismatch),
            "--language",
            "zh-CN",
        ], plugin_root, "direct_impact_marker does not match")

        bad_replay_log_marker_missing = copy_standard_bundle("replay_log_marker_missing")
        replay_log_path = bad_replay_log_marker_missing / "attachments/evidence/replay-output.log"
        replay_log_path.write_text(
            "Zhulong reviewer replay log\n"
            "Generated at: 2026-06-16T00:00:00Z\n"
            "COMMAND: docker compose -f attachments/poc/docker-compose.selftest.yml up --abort-on-container-exit\n"
            "RAW OUTPUT: deterministic selftest replay completed without the direct-impact token\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_replay_log_marker_missing),
            "--language",
            "zh-CN",
        ], plugin_root, "does not contain the deterministic direct-impact marker")

        def build_batch_confirmed_dir(name: str) -> tuple[Path, Path, Path]:
            batch_workspace = workspace / name
            if batch_workspace.exists():
                shutil.rmtree(batch_workspace)
            batch_confirmed = batch_workspace / "confirmed"
            batch_confirmed.mkdir(parents=True)
            batch_config = json.loads((workspace / "asr-config.json").read_text(encoding="utf-8"))
            batch_config["workspace_root"] = batch_workspace.name
            batch_config["confirmed_output_dir"] = f"{batch_workspace.name}/confirmed"
            (batch_workspace / "asr-config.json").write_text(
                json.dumps(batch_config, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            first = batch_confirmed / standard_bundle.name
            second = batch_confirmed / f"{standard_bundle.name}_sibling"
            shutil.copytree(standard_bundle, first)
            shutil.copytree(standard_bundle, second)
            write_live_replay_logs(first, second)
            return batch_confirmed, first, second

        batch_valid, _batch_first, _batch_second = build_batch_confirmed_dir("issue15-batch-valid")
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_all_report_bundles.py"),
            "--confirmed-dir",
            str(batch_valid),
            "--language",
            "zh-CN",
        ], plugin_root)

        batch_replay_regression, _batch_ok, batch_bad_replay = build_batch_confirmed_dir("issue15-batch-replay-regression")
        batch_bad_replay_script = batch_bad_replay / "run-selftest-jwt-recording.sh"
        batch_bad_replay_script.write_text(
            batch_bad_replay_script.read_text(encoding="utf-8")
            .replace("    show_evidence_summary\n    pause_step \"$PAUSE_LONG\"\n", "    show_evidence_summary\n", 1),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_all_report_bundles.py"),
            "--confirmed-dir",
            str(batch_replay_regression),
            "--language",
            "zh-CN",
        ], plugin_root, "partial confirmed bundle or validation failure detected")

        batch_consistency_regression, batch_bad_consistency, _batch_ok_2 = build_batch_confirmed_dir("issue15-batch-consistency-regression")
        batch_evidence_path = batch_bad_consistency / "verification-evidence.json"
        batch_evidence = json.loads(batch_evidence_path.read_text(encoding="utf-8"))
        batch_evidence["direct_impact_marker"] = "DIRECT_AVAILABILITY_IMPACT_CONFIRMED"
        batch_evidence_path.write_text(json.dumps(batch_evidence, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_all_report_bundles.py"),
            "--confirmed-dir",
            str(batch_consistency_regression),
            "--language",
            "zh-CN",
        ], plugin_root, "partial confirmed bundle or validation failure detected")

        bad_unrelated_readiness = copy_standard_bundle("unrelated_readiness")
        unrelated_readiness_script = bad_unrelated_readiness / "run-selftest-jwt-recording.sh"
        unrelated_readiness_script.write_text(
            unrelated_readiness_script.read_text(encoding="utf-8").replace(
                "    show_real_world_context\n",
                "    show_real_world_context\n    curl -fsS http://unrelated.invalid/health >/dev/null\n",
                1,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_unrelated_readiness),
            "--language",
            "zh-CN",
        ], plugin_root, "not used by proof commands")

        bad_readiness_pause_reuse = copy_standard_bundle("readiness_pause_reuse")
        readiness_pause_script = bad_readiness_pause_reuse / "run-selftest-jwt-recording.sh"
        readiness_pause_script.write_text(
            readiness_pause_script.read_text(encoding="utf-8")
            + "\nwait_for_service_readiness() {\n"
              "    retry_count=0\n"
              "    while [ \"$retry_count\" -lt 3 ]; do\n"
              "        sleep \"$PAUSE_SHORT\"\n"
              "        retry_count=$((retry_count + 1))\n"
              "    done\n"
              "}\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_readiness_pause_reuse),
            "--language",
            "zh-CN",
        ], plugin_root, "readiness/startup/retry/backoff")

        bad_absolute_evidence_log_output = copy_standard_bundle("absolute_evidence_log_output")
        absolute_evidence_script = bad_absolute_evidence_log_output / "run-selftest-jwt-recording.sh"
        absolute_evidence_script.write_text(
            absolute_evidence_script.read_text(encoding="utf-8").replace(
                '"$REPLAY_LOG_REL" >&2',
                '"$REPLAY_LOG" >&2',
                1,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_absolute_evidence_log_output),
            "--language",
            "zh-CN",
        ], plugin_root, "bundle-relative replay log path")

        bad_displayed_command_only = copy_standard_bundle("displayed_command_only")
        displayed_only_script = bad_displayed_command_only / "run-selftest-jwt-recording.sh"
        displayed_only_script.write_text(
            re.sub(
                r"^(\s*)run_logged_command\s+(.+)$",
                r"\1printf '%s\n' \2",
                displayed_only_script.read_text(encoding="utf-8"),
                flags=re.MULTILINE,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_displayed_command_only),
            "--language",
            "zh-CN",
        ], plugin_root, "without an actual bundle-local execution path")

        bad_missing_helper_reference = copy_standard_bundle("missing_helper_reference")
        missing_helper_note = next(bad_missing_helper_reference.glob("*_补充复现说明.md"))
        missing_helper_note.write_text(
            missing_helper_note.read_text(encoding="utf-8")
            + "\n\n补充检查：如需复核，请执行 `bash ./missing-helper.sh`。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_helper_reference),
            "--language",
            "zh-CN",
        ], plugin_root, "missing local helper")

        bad_pause_overwrite = copy_standard_bundle("quick_pause_overwrite")
        pause_script = bad_pause_overwrite / "run-selftest-jwt-recording.sh"
        pause_script.write_text(
            pause_script.read_text(encoding="utf-8")
            .replace('PAUSE_SHORT="${REVIEWER_PAUSE_SHORT:-1}"', "PAUSE_SHORT=0")
            .replace('PAUSE_LONG="${REVIEWER_PAUSE_LONG:-2}"', "PAUSE_LONG=0"),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_pause_overwrite),
            "--language",
            "zh-CN",
        ], plugin_root, "quick mode overwrites reviewer pause settings")

        bad_hardcoded_pause = copy_standard_bundle("hardcoded_pause")
        hardcoded_pause_script = bad_hardcoded_pause / "run-selftest-jwt-recording.sh"
        hardcoded_pause_script.write_text(
            hardcoded_pause_script.read_text(encoding="utf-8")
            .replace('pause_step "$PAUSE_SHORT"', "pause_step 1", 1),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_hardcoded_pause),
            "--language",
            "zh-CN",
        ], plugin_root, "fixed reviewer pause")

        bad_missing_pre_command_pause = copy_standard_bundle("missing_pre_command_pause")
        pre_command_pause_script = bad_missing_pre_command_pause / "run-selftest-jwt-recording.sh"
        pre_command_pause_script.write_text(
            pre_command_pause_script.read_text(encoding="utf-8").replace(
                "    show_real_world_context\n",
                "    show_real_world_context\n    run_logged_command 'python3 attachments/poc/jwt-forge-poc.py'\n",
                1,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_pre_command_pause),
            "--language",
            "zh-CN",
        ], plugin_root, "pause before proof command execution transitions")

        bad_missing_post_command_pause = copy_standard_bundle("missing_post_command_pause")
        post_command_pause_script = bad_missing_post_command_pause / "run-selftest-jwt-recording.sh"
        post_command_pause_script.write_text(
            post_command_pause_script.read_text(encoding="utf-8").replace(
                "    show_real_world_context\n",
                "    show_real_world_context\n    pause_step \"$PAUSE_SHORT\"\n"
                "    run_logged_command 'python3 attachments/poc/jwt-forge-poc.py'\n"
                "    verify_success_marker \"$SUCCESS_MARKER\"\n",
                1,
            ),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_post_command_pause),
            "--language",
            "zh-CN",
        ], plugin_root, "pause after proof command output transitions")

        bad_recursive_replay = copy_standard_bundle("recursive_replay")
        recursive_script = bad_recursive_replay / "run-selftest-jwt-recording.sh"
        recursive_script.write_text(
            recursive_script.read_text(encoding="utf-8")
            + "\n./$(basename \"$0\") quick docker\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_recursive_replay),
            "--language",
            "zh-CN",
        ], plugin_root, "recursively invoke itself")

        bad_exec_recursive_replay = copy_standard_bundle("exec_recursive_replay")
        exec_recursive_script = bad_exec_recursive_replay / "run-selftest-jwt-recording.sh"
        exec_recursive_script.write_text(
            exec_recursive_script.read_text(encoding="utf-8")
            + "\nexec \"$0\" quick docker\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_exec_recursive_replay),
            "--language",
            "zh-CN",
        ], plugin_root, "recursively invoke itself")

        bad_time_exact = copy_standard_bundle("stale_exact_timing_summary")
        time_supplement = next(bad_time_exact.glob("*_补充复现说明.md"))
        time_supplement.write_text(
            time_supplement.read_text(encoding="utf-8")
            + "\n补充：可用性 proof 的稳定结论是触发耗时 1.37 seconds。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_time_exact),
            "--language",
            "zh-CN",
        ], plugin_root, "stale exact timings")

        good_time_range = copy_standard_bundle("time_range_summary")
        time_range_supplement = next(good_time_range.glob("*_补充复现说明.md"))
        time_range_supplement.write_text(
            time_range_supplement.read_text(encoding="utf-8")
            + "\n补充：可用性 proof 使用至少 1 秒阈值描述，并要求审核者查看最新日志中的精确数值。\n",
            encoding="utf-8",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_time_range),
            "--language",
            "zh-CN",
        ], plugin_root)

        good_nested_parent_path = copy_standard_bundle("nested_parent_path_inside_bundle")
        nested_script = good_nested_parent_path / "attachments/nested/inside-bundle.sh"
        nested_script.parent.mkdir(parents=True, exist_ok=True)
        nested_script.write_text(
            "#!/bin/sh\n"
            "set -eu\n"
            "SCRIPT_DIR=\"$(CDPATH= cd -- \"$(dirname -- \"$0\")\" && pwd)\"\n"
            "BUNDLE_LOCAL=\"$(cd \"$SCRIPT_DIR/..\" && pwd)\"\n"
            "test -d \"$BUNDLE_LOCAL\"\n",
            encoding="utf-8",
        )
        nested_script.chmod(0o755)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_nested_parent_path),
            "--language",
            "zh-CN",
        ], plugin_root)

        npm_warning_bundle = copy_standard_bundle("package_manager_install_warning")
        npm_script = npm_warning_bundle / "attachments/npm-install-warning.sh"
        npm_script.write_text("#!/bin/sh\nset -eu\nnpm install left-pad\n", encoding="utf-8")
        npm_script.chmod(0o755)
        npm_warning_output = run_capture([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(npm_warning_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        if "package manager install command" not in npm_warning_output:
            raise SystemExit("FAILED: package manager install warning fixture did not emit a warning")

        poc_label_warning_bundle = copy_standard_bundle("poc_label_warning")
        poc_label_supplement = next(poc_label_warning_bundle.glob("*_补充复现说明.md"))
        poc_label_supplement.write_text(
            poc_label_supplement.read_text(encoding="utf-8")
            + "\n补充：PoC-4 展示了最高编号复现路径，但根录屏脚本尚未覆盖该标签。\n",
            encoding="utf-8",
        )
        poc_label_output = run_capture([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(poc_label_warning_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        if "root recording script appears to miss the highest PoC label" not in poc_label_output:
            raise SystemExit("FAILED: PoC label drift fixture did not emit a warning")

        stale_video_bundle = copy_standard_bundle("stale_video_warning")
        stale_video = stale_video_bundle / "复现视频.mp4"
        stale_video.write_bytes(b"placeholder video bytes")
        os.utime(stale_video, (1, 1))
        stale_output = run_capture([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(stale_video_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        if "recording appears older than current reproduction script or report material" not in stale_output:
            raise SystemExit("FAILED: stale recording video fixture did not emit a warning")

        def standard_reviewer_paths(bundle: Path) -> tuple[str, str]:
            poc_matches = sorted((bundle / "attachments").rglob("jwt-forge-poc.py"))
            evidence_matches = sorted((bundle / "attachments").rglob("forged-token-response.json"))
            if not poc_matches or not evidence_matches:
                raise SystemExit("FAILED: standard bundle is missing expected reviewer evidence selftest attachments")
            return (
                poc_matches[0].relative_to(bundle).as_posix(),
                evidence_matches[0].relative_to(bundle).as_posix(),
            )

        def write_useful_reviewer_addendum(bundle: Path, *, extra: str = "") -> None:
            _poc_rel, evidence_rel = standard_reviewer_paths(bundle)
            (bundle / "reviewer-evidence-and-impact.md").write_text(
                "# 审核证据与影响说明\n\n"
                "## 攻击者能力与边界\n\n"
                "攻击者需要能够向已确认的 Docker 复现入口提交伪造 JWT；服务端条件是默认密钥配置生效。"
                "本包只声称 Docker 成功判据证明的认证绕过影响，不声称未验证的主机执行或容器逃逸。\n\n"
                "## 审核方最短复现\n\n"
                "在 bundle 根目录运行 `REVIEWER_PAUSE_SHORT=0 REVIEWER_PAUSE_LONG=0 ./run-selftest-jwt-recording.sh quick docker`。\n\n"
                "## 成功判据与证据映射\n\n"
                "- 成功判据：`认证绕过成功`\n"
                f"- 证据文件：`{evidence_rel}`\n\n"
                "## 已验证影响\n\n"
                "已验证影响是完整性与认证边界绕过，审核材料中的证据和 replay command 均为 bundle-local。\n"
                + (f"\n{extra}\n" if extra else ""),
                encoding="utf-8",
            )

        def write_standard_reviewer_index(
            bundle: Path,
            *,
            oracle: str = "认证绕过成功",
            artifact_paths: list[str] | None = None,
            command: str = "REVIEWER_PAUSE_SHORT=0 REVIEWER_PAUSE_LONG=0 ./run-selftest-jwt-recording.sh quick docker",
            extra: dict | None = None,
        ) -> None:
            poc_rel, evidence_rel = standard_reviewer_paths(bundle)
            data = {
                "schema_version": 1,
                "finding_slug": bundle.name,
                "bundle_root_command": command,
                "poc_files": [poc_rel],
                "evidence_outputs": artifact_paths if artifact_paths is not None else [evidence_rel],
                "success_oracles": [oracle],
                "real_world_exploitability_summary": "攻击者控制伪造 JWT，默认密钥配置让认证绕过在 Docker 中可达。",
                "boundaries": ["不声称容器逃逸或宿主机执行。"],
            }
            if extra:
                data.update(extra)
            index_path = bundle / "attachments/reviewer-evidence-index.json"
            index_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

        good_reviewer_index = copy_standard_bundle("reviewer_index_valid")
        write_useful_reviewer_addendum(good_reviewer_index)
        write_standard_reviewer_index(good_reviewer_index)
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_reviewer_index),
            "--language",
            "zh-CN",
        ], plugin_root)
        all_errors_valid = run_report_bundle_all_errors_json(plugin_root, good_reviewer_index, language="zh-CN")
        if all_errors_valid.get("valid") is not True or all_errors_valid.get("issues") != []:
            raise SystemExit(f"FAILED: valid bundle --all-errors --json did not return valid=true with no issues:\n{all_errors_valid}")
        all_errors_output_path = good_reviewer_index / "bundle-validation-errors.json"
        all_errors_output = run_report_bundle_all_errors_json(
            plugin_root,
            good_reviewer_index,
            language="zh-CN",
            output_errors=all_errors_output_path,
        )
        if all_errors_output.get("valid") is not True or all_errors_output.get("issues") != []:
            raise SystemExit("FAILED: --output-errors JSON did not preserve valid=true empty issue list")
        all_errors_output_path.unlink()

        bad_all_errors_bundle = copy_standard_bundle("all_errors_malformed")
        write_useful_reviewer_addendum(
            bad_all_errors_bundle,
            extra=(
                "## malformed all-errors fixture\n\n"
                "本包故意使用 minimal fixture replay 但没有说明来源依据。"
                "同时故意写入 SSRF response content exposes credentials to target output 的过强声明，用于验证 all-errors collector。\n"
            ),
        )
        write_standard_reviewer_index(
            bad_all_errors_bundle,
            artifact_paths=["attachments/evidence/missing-reviewer-artifact.log"],
        )
        bad_all_errors_evidence_path = bad_all_errors_bundle / "verification-evidence.json"
        bad_all_errors_evidence = json.loads(bad_all_errors_evidence_path.read_text(encoding="utf-8"))
        bad_all_errors_evidence["direct_impact_marker"] = "DIRECT_AVAILABILITY_IMPACT_CONFIRMED"
        bad_all_errors_evidence_path.write_text(
            json.dumps(bad_all_errors_evidence, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        bad_all_errors_log = bad_all_errors_bundle / "attachments/evidence/replay-output.log"
        bad_all_errors_log.write_text(
            "Zhulong reviewer replay log placeholder.\n"
            "Run the bundle-root replay script to refresh this file with live reviewer output.\n",
            encoding="utf-8",
        )
        bad_all_errors_script = bad_all_errors_bundle / "run-selftest-jwt-recording.sh"
        bad_all_errors_script.write_text(
            bad_all_errors_script.read_text(encoding="utf-8")
            .replace('REPLAY_LOG="$EVIDENCE_DIR/replay-output.log"', 'REPLAY_LOG="$EVIDENCE_DIR/unregistered-output.log"')
            .replace(
                "    show_real_world_context\n",
                "    show_real_world_context\n    run_logged_command 'python3 attachments/poc/jwt-forge-poc.py'\n",
                1,
            )
            .replace("    verify_success_marker \"$SUCCESS_MARKER\"\n", ""),
            encoding="utf-8",
        )
        replace_docx_section_with_one_line(
            next(bad_all_errors_bundle.glob("*.docx")),
            "关键代码上下文",
            code_context_stop_headings,
            "代码上下文 1",
        )
        bad_all_errors_supplement = next(bad_all_errors_bundle.glob("*_补充复现说明.md"))
        bad_all_errors_supplement.write_text(
            bad_all_errors_supplement.read_text(encoding="utf-8")
            + "\n补充检查：如需复核，请执行 `bash ./missing-helper.sh`。\n"
            + "\n调试泄漏：{'oracle_token': '认证绕过成功', 'impact': ['auth bypass'], 'confirmed': True}\n",
            encoding="utf-8",
        )
        bad_all_errors_payload = run_report_bundle_all_errors_json(
            plugin_root,
            bad_all_errors_bundle,
            language="zh-CN",
            expected_returncode=1,
        )
        if bad_all_errors_payload.get("valid") is not False:
            raise SystemExit("FAILED: malformed --all-errors payload must return valid=false")
        if len(bad_all_errors_payload.get("issues", [])) < 5:
            raise SystemExit(f"FAILED: malformed --all-errors payload returned fewer than 5 issues:\n{bad_all_errors_payload}")
        require_report_issue(bad_all_errors_payload, "REVIEWER_INDEX_ARTIFACT_MISSING")
        require_report_issue(bad_all_errors_payload, "REPLAY_LOG_UNREGISTERED")
        require_report_issue(bad_all_errors_payload, "REPLAY_LOG_PLACEHOLDER")
        require_report_issue(bad_all_errors_payload, "DIRECT_IMPACT_MARKER_MISMATCH")
        require_report_issue(bad_all_errors_payload, "REPLAY_LOG_MARKER_MISSING")
        require_report_issue(bad_all_errors_payload, "FIXTURE_PROVENANCE_MISSING")
        require_report_issue(bad_all_errors_payload, "SSRF_IMPACT_OVERCLAIM")
        require_report_issue(bad_all_errors_payload, "CODE_CONTEXT_MINIMUM_QUALITY")
        require_report_issue(bad_all_errors_payload, "REPLAY_HELPER_PAUSE_CONTRACT")
        if not (
            any(issue.get("code") == "RAW_STRUCTURED_OBJECT_LEAK" for issue in bad_all_errors_payload.get("issues", []))
            or any(issue.get("code") == "LOCAL_HELPER_MISSING" for issue in bad_all_errors_payload.get("issues", []))
        ):
            raise SystemExit("FAILED: malformed all-errors payload must include raw structured leakage or missing local helper")
        bad_all_errors_output_path = bad_all_errors_bundle / "bundle-validation-errors.json"
        bad_all_errors_output_payload = run_report_bundle_all_errors_json(
            plugin_root,
            bad_all_errors_bundle,
            language="zh-CN",
            expected_returncode=1,
            output_errors=bad_all_errors_output_path,
        )
        require_report_issue(bad_all_errors_output_payload, "REVIEWER_INDEX_ARTIFACT_MISSING")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_all_errors_bundle),
            "--language",
            "zh-CN",
        ], plugin_root, "VALIDATION FAILED")

        bad_index_variant_candidates_proof = copy_standard_bundle("reviewer_index_variant_candidates_proof")
        write_useful_reviewer_addendum(bad_index_variant_candidates_proof)
        write_standard_reviewer_index(
            bad_index_variant_candidates_proof,
            extra={"confirmation_basis": "variant-candidates.jsonl proves this variant is confirmed"},
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_index_variant_candidates_proof),
            "--language",
            "zh-CN",
        ], plugin_root, "variant candidate output must not be used as confirmed evidence")
        shutil.rmtree(bad_index_variant_candidates_proof)

        bad_index_missing_artifact = copy_standard_bundle("reviewer_index_missing_artifact")
        write_useful_reviewer_addendum(bad_index_missing_artifact)
        write_standard_reviewer_index(
            bad_index_missing_artifact,
            artifact_paths=["attachments/evidence/does-not-exist.log"],
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_index_missing_artifact),
            "--language",
            "zh-CN",
        ], plugin_root, "referenced artifact path does not exist")

        bad_index_outside_path = copy_standard_bundle("reviewer_index_outside_path")
        write_useful_reviewer_addendum(bad_index_outside_path)
        write_standard_reviewer_index(bad_index_outside_path, artifact_paths=["../outside.log"])
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_index_outside_path),
            "--language",
            "zh-CN",
        ], plugin_root, "must stay inside the bundle")

        bad_index_local_path = copy_standard_bundle("reviewer_index_local_path")
        write_useful_reviewer_addendum(bad_index_local_path)
        submitter_local_path = "/" + "Users/" + "localuser/tmp/evidence.log"
        write_standard_reviewer_index(bad_index_local_path, artifact_paths=[submitter_local_path])
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_index_local_path),
            "--language",
            "zh-CN",
        ], plugin_root, "absolute/operator-local")

        bad_index_missing_oracle_source = copy_standard_bundle("reviewer_index_missing_oracle_source")
        write_useful_reviewer_addendum(bad_index_missing_oracle_source)
        write_standard_reviewer_index(bad_index_missing_oracle_source, oracle="NEVER_OBSERVED_REVIEWER_ORACLE")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_index_missing_oracle_source),
            "--language",
            "zh-CN",
        ], plugin_root, "success oracle token is not present")

        bad_placeholder_addendum = copy_standard_bundle("reviewer_addendum_placeholder")
        (bad_placeholder_addendum / "reviewer-evidence-and-impact.md").write_text(
            "# 审核证据\n\nTODO\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_addendum),
            "--language",
            "zh-CN",
        ], plugin_root, "placeholder-only")

        bad_fixture_without_provenance = copy_standard_bundle("fixture_without_provenance")
        fixture_supplement = next(bad_fixture_without_provenance.glob("*_补充复现说明.md"))
        fixture_supplement.write_text(
            fixture_supplement.read_text(encoding="utf-8")
            + "\n补充：本包使用 minimal fixture replay。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_fixture_without_provenance),
            "--language",
            "zh-CN",
        ], plugin_root, "source-grounded provenance")

        good_fixture_provenance = copy_standard_bundle("fixture_with_provenance")
        write_useful_reviewer_addendum(
            good_fixture_provenance,
            extra=(
                "## fixture provenance\n\n"
                "本包使用最小 fixture，但它保留原始源码中的危险模式和 Docker 复现边界；"
                "该 fixture 足以复现认证绕过边界。未验证、不声称容器逃逸、宿主机 RCE 或更强影响。"
            ),
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_fixture_provenance),
            "--language",
            "zh-CN",
        ], plugin_root)

        bad_library_boundary = copy_standard_bundle("library_boundary_missing")
        library_supplement = next(bad_library_boundary.glob("*_补充复现说明.md"))
        library_supplement.write_text(
            library_supplement.read_text(encoding="utf-8")
            + "\n补充：这是一个库漏洞，public API 接收攻击者可控 key。\n",
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_library_boundary),
            "--language",
            "zh-CN",
        ], plugin_root, "consumer application boundary")

        good_q_style_boundary = copy_standard_bundle("q_style_library_boundary")
        write_useful_reviewer_addendum(
            good_q_style_boundary,
            extra=(
                "## library consumer boundary\n\n"
                "这是 q-style library/package 漏洞：public API `Q.set` 接收攻击者可控 key/name 参数。"
                "消费方或上层应用需要把用户字段名桥接到该 API，库本身不提供网络入口。"
                "单步路径是目标对象局部 prototype chain hijack；全局 `Object.prototype` 污染只在额外 consumer pattern 中成立，"
                "本包不声称无上层应用桥接即可远程触发。"
            ),
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_q_style_boundary),
            "--language",
            "zh-CN",
        ], plugin_root)

        good_speedtest_style_fixture = copy_standard_bundle("speedtest_style_fixture")
        speedtest_evidence_dir = good_speedtest_style_fixture / "attachments/evidence"
        speedtest_evidence_dir.mkdir(parents=True, exist_ok=True)
        (speedtest_evidence_dir / "webshell-oracle.txt").write_text(
            "WEBSHELL_CONFIRMED\nuid=82(www-data)\n",
            encoding="utf-8",
        )
        write_useful_reviewer_addendum(
            good_speedtest_style_fixture,
            extra=(
                "## Speedtest-style source-grounded fixture\n\n"
                "本包使用最小 fixture，保留原始源码中的 sed/entrypoint 危险模式；fixture 足以复现代码注入边界。"
                "成功判据包括 `WEBSHELL_CONFIRMED` 与 `uid=82(www-data)`，对应 `attachments/evidence/webshell-oracle.txt`。"
                "未验证、不声称容器逃逸、宿主机执行或匿名公开入口。"
            ),
        )
        write_standard_reviewer_index(
            good_speedtest_style_fixture,
            oracle="WEBSHELL_CONFIRMED",
            artifact_paths=["attachments/evidence/webshell-oracle.txt"],
            extra={"fixture_files": ["attachments/evidence/webshell-oracle.txt"]},
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(good_speedtest_style_fixture),
            "--language",
            "zh-CN",
        ], plugin_root)

        bad_severity_mismatch = copy_standard_bundle("severity_body_mismatch")
        mutate_bundle_finding(bad_severity_mismatch, lambda _finding: None)
        from docx import Document
        severity_docx_path = next(bad_severity_mismatch.glob("*.docx"))
        severity_doc = Document(severity_docx_path)
        severity_doc.add_paragraph("等级判定：中危")
        severity_doc.save(severity_docx_path)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_severity_mismatch),
            "--language",
            "zh-CN",
        ], plugin_root, "report body severity does not match")

        bad_claim_overreach = copy_standard_bundle("webshell_claim_without_oracle")
        write_useful_reviewer_addendum(
            bad_claim_overreach,
            extra=(
                "## 过强声明夹具\n\n"
                "本段故意声称已验证 HTTP webshell 命令执行，但成功判据仍只有 `认证绕过成功`，用于确认 validator 会拒绝 claim/oracle 不一致。"
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_claim_overreach),
            "--language",
            "zh-CN",
        ], plugin_root, "webshell or HTTP command execution")

        bad_missing_attacker = quality_gate_bad_bundle(
            "missing_attacker_condition",
            lambda text: None if text.startswith("攻击者条件") or text.startswith("入口/可控输入") else text,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_attacker),
            "--language",
            "zh-CN",
        ], plugin_root, "missing 攻击者条件")

        bad_missing_server = quality_gate_bad_bundle(
            "missing_server_condition",
            lambda text: None if text.startswith("服务端条件") else text,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_server),
            "--language",
            "zh-CN",
        ], plugin_root, "missing 服务端条件")

        bad_missing_impact = quality_gate_bad_bundle(
            "missing_security_impact",
            lambda text: None if text.startswith("安全影响") else text,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_impact),
            "--language",
            "zh-CN",
        ], plugin_root, "missing 安全影响")

        bad_placeholder_attacker = quality_gate_bad_bundle(
            "placeholder_attacker_condition",
            lambda text: "攻击者条件：待补充" if text.startswith("攻击者条件") else text,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_attacker),
            "--language",
            "zh-CN",
        ], plugin_root, "placeholder-only")

        bad_weak_impact = quality_gate_bad_bundle(
            "weak_security_impact",
            lambda text: "安全影响：该问题很危险。" if text.startswith("安全影响") else text,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_weak_impact),
            "--language",
            "zh-CN",
        ], plugin_root, "must mention a concrete CIA impact")
        for bad_quality_bundle in (
            bad_missing_attacker,
            bad_missing_server,
            bad_missing_impact,
            bad_placeholder_attacker,
            bad_weak_impact,
            bad_placeholder_replay_log,
            good_ssrf_callback_bounded,
            bad_ssrf_callback_overclaim,
            bad_root_relative_artifact,
            bad_missing_real_world,
            bad_missing_code_context,
            bad_prose_code_context,
            bad_placeholder_code_context,
            bad_en_placeholder_code_context,
            bad_abbreviated_code_context,
            bad_blue_code_context,
            bad_placeholder_real_world,
            bad_strong_boundary,
            good_strong_boundary,
            bad_en_missing_real_world,
            bad_bundle_escape,
            bad_pkg_dependency,
            bad_workspace_marker,
            bad_missing_target_identity,
            bad_legacy_target_identity,
            bad_no_opening_identity_pause,
            bad_late_target_identity,
            bad_no_replay_code_context,
            bad_no_replay_analysis,
            bad_no_replay_real_world,
            bad_no_replay_final_summary,
            bad_unavailable_replay_context,
            bad_no_final_summary_pause,
            bad_undefined_root_helper,
            bad_missing_replay_log,
            bad_unregistered_replay_log,
            bad_final_without_marker_check,
            bad_explanatory_marker_text,
            bad_log_without_raw_output,
            bad_missing_direct_impact,
            bad_raw_structured_docx,
            bad_mutable_runtime_identity,
            bad_direct_marker_mismatch,
            bad_replay_log_marker_missing,
            bad_unrelated_readiness,
            bad_readiness_pause_reuse,
            bad_absolute_evidence_log_output,
            bad_displayed_command_only,
            bad_missing_helper_reference,
            bad_pause_overwrite,
            bad_hardcoded_pause,
            bad_missing_pre_command_pause,
            bad_missing_post_command_pause,
            bad_recursive_replay,
            bad_exec_recursive_replay,
            bad_time_exact,
            good_time_range,
            good_nested_parent_path,
            npm_warning_bundle,
            poc_label_warning_bundle,
            stale_video_bundle,
            good_reviewer_index,
            bad_all_errors_bundle,
            bad_index_missing_artifact,
            bad_index_outside_path,
            bad_index_local_path,
            bad_index_missing_oracle_source,
            bad_placeholder_addendum,
            bad_fixture_without_provenance,
            good_fixture_provenance,
            bad_library_boundary,
            good_q_style_boundary,
            good_speedtest_style_fixture,
            bad_severity_mismatch,
            bad_claim_overreach,
        ):
            shutil.rmtree(bad_quality_bundle)
        legacy_marker_fixture = workspace / "legacy-english-analysis-markers-finding.json"
        legacy_marker_data = json.loads(standard_fixture.read_text(encoding="utf-8"))
        legacy_marker_data["vulnerability_id"] = "SELFTEST-LEGACY-MARKERS"
        legacy_marker_data["vulnerability_name"] = "导入URL服务端请求伪造"
        legacy_marker_data["severity"] = "高危"
        legacy_marker_data["severity_cn"] = "高危"
        legacy_marker_data["cvss"] = {
            "vector": "CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:U/C:H/I:N/A:N",
            "score": "7.5",
            "severity": "高危",
            "rationale": ["评估依据：服务端可被诱导访问内网资源，形成 SSRF 信息泄露风险。"],
        }
        legacy_marker_data["analysis"] = [
            "Location: api/src/services/files.ts importOne() accepts a user-supplied URL.",
            "Entry / controllable input: an authenticated attacker submits an import URL.",
            "Dangerous operation: axios.get() performs a server-side fetch.",
            "Trigger path: URL import -> server fetch -> attacker-controlled internal destination.",
            "Root cause: URL import lacks complete private-network deny-list validation.",
            "Why existing checks fail: the current deny list is incomplete for common internal ranges.",
        ]
        legacy_marker_data["security_impact"] = (
            "Docker replay confirms SSRF internal response content is exposed in target output; "
            "INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED bounds the direct impact."
        )
        legacy_marker_data["real_world_exploitability"] = [
            "实际使用场景：导入 URL 服务处理用户提交的远程地址，服务端默认可访问内部 metadata service。",
            "攻击者路径：认证攻击者控制导入 URL，使服务端请求内部服务。",
            "触发调用链：用户 URL 进入 importOne() 后到达 axios.get() 请求 sink。",
            "直接危害证明：内部响应内容进入目标输出，INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED 和 DIRECT_IMPACT_CONFIRMED 标记证明有界直接危害。",
            "影响边界：仅声称 Docker PoC 证明的内部响应外显，不声称未验证的更强主机影响。",
        ]
        legacy_marker_data["reproduction"][0]["results"].append(
            "结果证据：INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED 显示内部响应内容进入目标输出。"
        )
        legacy_marker_fixture.write_text(json.dumps(legacy_marker_data, ensure_ascii=False, indent=2), encoding="utf-8")
        run([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(legacy_marker_fixture),
            "--output-dir",
            str(workspace / "confirmed"),
            "--language",
            "zh-CN",
        ], plugin_root)
        legacy_marker_bundle = next(
            (
                path for path in (workspace / "confirmed").iterdir()
                if path.is_dir() and "导入URL" in path.name
            ),
            None,
        )
        if legacy_marker_bundle is None:
            raise SystemExit("FAILED: legacy marker fixture did not render a confirmed bundle")
        write_live_replay_log(
            legacy_marker_bundle,
            extra="INTERNAL_RESPONSE_EXFILTRATED_CONFIRMED: internal response content exposed in target output\n",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(legacy_marker_bundle),
            "--language",
            "zh-CN",
        ], plugin_root)
        legacy_marker_lines = "\n".join(docx_text(next(legacy_marker_bundle.glob("*.docx"))))
        if "Location:" in legacy_marker_lines or "Entry / controllable input:" in legacy_marker_lines:
            raise SystemExit("FAILED: renderer did not localize legacy English analysis markers for zh-CN output")
        if "位置：" not in legacy_marker_lines or "入口/可控输入：" not in legacy_marker_lines:
            raise SystemExit("FAILED: localized zh-CN analysis markers are missing from legacy marker fixture")
        missing_name_fixture = workspace / "missing-vulnerability-name-finding.json"
        missing_name_data = json.loads(standard_fixture.read_text(encoding="utf-8"))
        missing_name_data.pop("vulnerability_name", None)
        missing_name_data.pop("vulnerability_name_en", None)
        missing_name_data["title_zh"] = "gothinkster/node-express-realworld-example-app 默认配置下硬编码 JWT 密钥导致身份认证绕过并允许攻击者伪造任意用户 token 的完整漏洞报告标题"
        missing_name_fixture.write_text(json.dumps(missing_name_data, ensure_ascii=False, indent=2), encoding="utf-8")
        run_expect_fail([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(missing_name_fixture),
            "--output-dir",
            str(workspace / "confirmed"),
            "--language",
            "zh-CN",
        ], plugin_root, "must include vulnerability_name")

        bad_runtime_scope = zh_bundle.parent / f"{standard_bundle.name}_runtime_scope_overclaim"
        shutil.copytree(standard_bundle, bad_runtime_scope)
        runtime_scope_findings_path = bad_runtime_scope / "findings.json"
        if not runtime_scope_findings_path.exists():
            shutil.copy2(standard_fixture, runtime_scope_findings_path)
        runtime_scope_data = json.loads(runtime_scope_findings_path.read_text(encoding="utf-8"))
        runtime_scope_data["source_runtime_match"] = False
        runtime_scope_data.setdefault("impact", {})["affected_versions"] = "v2.9.1（Docker 验证版本），可能影响所有版本"
        runtime_scope_findings_path.write_text(
            json.dumps(runtime_scope_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_runtime_scope),
            "--language",
            "zh-CN",
        ], plugin_root, "source/runtime mismatch detected")

        nested_warning_bundle = zh_bundle.parent / f"{standard_bundle.name}_nested_attachment_warning"
        shutil.copytree(standard_bundle, nested_warning_bundle)
        nested_dir = nested_warning_bundle / "attachments/security-research-20260502-123456/evidence"
        nested_dir.mkdir(parents=True, exist_ok=True)
        nested_file = nested_dir / "forged-token-response.json"
        nested_file.write_text('{"ok":true,"nested":true}\n', encoding="utf-8")
        nested_evidence_data = json.loads((nested_warning_bundle / "verification-evidence.json").read_text(encoding="utf-8"))
        nested_evidence_data["evidence_files"].append("attachments/security-research-20260502-123456/evidence/forged-token-response.json")
        (nested_warning_bundle / "verification-evidence.json").write_text(
            json.dumps(nested_evidence_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        nested_proc = subprocess.run([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(nested_warning_bundle),
            "--language",
            "zh-CN",
        ], cwd=plugin_root, capture_output=True, text=True)
        nested_output = (nested_proc.stdout or "") + (nested_proc.stderr or "")
        if nested_proc.returncode == 0:
            raise SystemExit("FAILED: nested timestamped workspace attachment fixture unexpectedly validated")
        if "non-standalone path text" not in nested_output:
            raise SystemExit(
                "FAILED: nested timestamped workspace attachment fixture did not fail on the new path-redaction gate\n"
                + nested_output
            )
        shutil.rmtree(bad_runtime_scope)
        shutil.rmtree(nested_warning_bundle)
        events_before_bundle_validation = [
            json.loads(line)
            for line in (workspace / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        def normalized_event_name(event: dict[str, object]) -> str:
            return str(event.get("event_name") or event.get("event") or "")

        def normalized_event_details(event: dict[str, object]) -> dict[str, object]:
            details = event.get("details")
            if not isinstance(details, dict):
                return {}
            metadata = details.get("metadata")
            if not isinstance(metadata, list):
                return details
            result: dict[str, object] = {"summary": details.get("summary")}
            for item in metadata:
                if isinstance(item, dict) and isinstance(item.get("key"), str):
                    result[item["key"]] = item.get("value")
            return result

        bundle_validated_before = sum(
            1 for event in events_before_bundle_validation
            if normalized_event_name(event) == "bundle_validated"
        )
        for event_name, stage in [
            ("selftest_candidates_started", "candidate_generation"),
            ("selftest_triage_started", "triage"),
            ("selftest_verification_started", "verification"),
        ]:
            run([
                sys.executable,
                str(workspace / "bin/write-audit-event.py"),
                "--workspace-dir", str(workspace),
                "--event", event_name,
                "--stage", stage,
                "--status", "running",
                "--transition-kind", "advance",
                "--message", "Advance the bounded selftest workflow before bundle validation.",
                "--accept-current-revision",
            ], plugin_root)
        for _ in range(2):
            run([
                sys.executable,
                str(plugin_root / "scripts/validate_report_bundle.py"),
                "--bundle-dir",
                str(en_bundle),
                "--language",
                "en-US",
                "--write-audit-event",
            ], plugin_root)
        shutil.copy2(
            plugin_root / "assets/examples/confirmed-findings.example.json",
            workspace / "confirmed/findings.example.json",
        )
        shutil.copy2(
            plugin_root / "assets/confirmed-vuln-report-template.docx",
            workspace / "confirmed/confirmed-vuln-report-template.docx",
        )
        (workspace / "confirmed/.DS_Store").write_text("", encoding="utf-8")
        run([
            sys.executable,
            str(workspace / "bin/validate-all-report-bundles.py"),
            "--confirmed-dir",
            str(workspace / "confirmed"),
        ], plugin_root)
        partial_confirmed = workspace / "confirmed/C99-partial-confirmed"
        partial_confirmed.mkdir()
        shutil.copy2(zh_bundle / "verification-evidence.json", partial_confirmed / "verification-evidence.json")
        (partial_confirmed / "findings.json").write_text(
            json.dumps({"findings": [{"slug": "c99-partial-confirmed"}]}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(workspace / "bin/validate-all-report-bundles.py"),
            "--confirmed-dir",
            str(workspace / "confirmed"),
        ], plugin_root, "partial confirmed bundle")
        partial_json_proc = subprocess.run([
            sys.executable,
            str(workspace / "bin/validate-all-report-bundles.py"),
            "--confirmed-dir",
            str(workspace / "confirmed"),
            "--json",
        ], cwd=plugin_root, capture_output=True, text=True)
        if partial_json_proc.returncode == 0:
            raise SystemExit("FAILED: validate-all-report-bundles.py --json unexpectedly passed with a partial bundle")
        partial_json = json.loads(partial_json_proc.stdout)
        partial_entries = [
            item for item in partial_json.get("results", [])
            if item.get("name") == "C99-partial-confirmed"
        ]
        if not partial_entries or partial_entries[0].get("classification") != "partial_confirmed_bundle":
            raise SystemExit("FAILED: validate-all-report-bundles.py --json did not classify the partial bundle")
        helper_classes = {
            item.get("name"): item.get("classification")
            for item in partial_json.get("results", [])
            if item.get("name") in {
                "findings.example.json",
                "confirmed-vuln-report-template.docx",
                ".DS_Store",
            }
        }
        expected_helpers = {
            "findings.example.json",
            "confirmed-vuln-report-template.docx",
            ".DS_Store",
        }
        if set(helper_classes) != expected_helpers or set(helper_classes.values()) != {"ignored_helper_file"}:
            raise SystemExit("FAILED: validate-all-report-bundles.py --json did not ignore confirmed/ helper files")
        events_after_bundle_validation = [
            json.loads(line)
            for line in (workspace / "audit-events.jsonl").read_text(encoding="utf-8").splitlines()
            if line.strip()
        ]
        bundle_validated_events = [
            event for event in events_after_bundle_validation
            if normalized_event_name(event) == "bundle_validated"
        ]
        if len(bundle_validated_events) != bundle_validated_before + 2:
            raise SystemExit("FAILED: validate_report_bundle.py did not append both bundle_validated events")
        bundle_event = bundle_validated_events[-1]
        if bundle_event.get("stage") != "packaging" or bundle_event.get("to_status", bundle_event.get("status")) != "running":
            raise SystemExit("FAILED: bundle_validated event must use packaging/running without completing the audit")
        if bundle_event.get("transition_kind") != "observe":
            raise SystemExit("FAILED: repeated bundle validation must observe the existing packaging stage")
        details = normalized_event_details(bundle_event)
        if details.get("bundle") != f"confirmed/{en_bundle.name}":
            raise SystemExit("FAILED: bundle_validated event must store a workspace-relative bundle path")
        if details.get("verification_status") != "confirmed_in_docker":
            raise SystemExit("FAILED: bundle_validated event must preserve confirmed_in_docker evidence status")
        stage_status = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
        if stage_status.get("last_event_name", stage_status.get("last_event")) != "bundle_validated" or stage_status.get("status") != "running":
            raise SystemExit("FAILED: bundle validation must update state without marking the audit completed")

        bad_missing_verification = zh_bundle.parent / f"{zh_bundle.name}_missing_verification_evidence"
        shutil.copytree(zh_bundle, bad_missing_verification)
        (bad_missing_verification / "verification-evidence.json").unlink()
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_verification),
            "--language",
            "zh-CN",
        ], plugin_root, "confirmed bundle must include verification-evidence.json")

        bad_high_confidence = zh_bundle.parent / f"{zh_bundle.name}_high_confidence_status"
        shutil.copytree(zh_bundle, bad_high_confidence)
        high_confidence_data = json.loads((bad_high_confidence / "verification-evidence.json").read_text(encoding="utf-8"))
        high_confidence_data["verification_status"] = "high_confidence_unverified_due_to_sandbox_limitation"
        (bad_high_confidence / "verification-evidence.json").write_text(
            json.dumps(high_confidence_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_high_confidence),
            "--language",
            "zh-CN",
        ], plugin_root, "verification_status=confirmed_in_docker")

        bad_missing_evidence_file = zh_bundle.parent / f"{zh_bundle.name}_missing_evidence_file"
        shutil.copytree(zh_bundle, bad_missing_evidence_file)
        missing_evidence_data = json.loads((bad_missing_evidence_file / "verification-evidence.json").read_text(encoding="utf-8"))
        missing_evidence_data["evidence_files"] = ["attachments/evidence/missing.log"]
        (bad_missing_evidence_file / "verification-evidence.json").write_text(
            json.dumps(missing_evidence_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_evidence_file),
            "--language",
            "zh-CN",
        ], plugin_root, "does not exist inside bundle")

        bad_absolute_evidence = zh_bundle.parent / f"{zh_bundle.name}_absolute_evidence_path"
        shutil.copytree(zh_bundle, bad_absolute_evidence)
        absolute_evidence_data = json.loads((bad_absolute_evidence / "verification-evidence.json").read_text(encoding="utf-8"))
        absolute_evidence_data["evidence_files"] = [str((bad_absolute_evidence / "attachments/poc/path_traversal.py").resolve())]
        (bad_absolute_evidence / "verification-evidence.json").write_text(
            json.dumps(absolute_evidence_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_absolute_evidence),
            "--language",
            "zh-CN",
        ], plugin_root, "must be bundle-relative")

        bad_escape_evidence = zh_bundle.parent / f"{zh_bundle.name}_escape_evidence_path"
        shutil.copytree(zh_bundle, bad_escape_evidence)
        escape_evidence_data = json.loads((bad_escape_evidence / "verification-evidence.json").read_text(encoding="utf-8"))
        escape_evidence_data["evidence_files"] = ["../outside.log"]
        (bad_escape_evidence / "verification-evidence.json").write_text(
            json.dumps(escape_evidence_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_escape_evidence),
            "--language",
            "zh-CN",
        ], plugin_root, "must not escape the bundle with '..'")

        bad_empty_poc = zh_bundle.parent / f"{zh_bundle.name}_empty_poc_path"
        shutil.copytree(zh_bundle, bad_empty_poc)
        empty_poc_data = json.loads((bad_empty_poc / "verification-evidence.json").read_text(encoding="utf-8"))
        empty_poc_data["poc_path"] = ""
        (bad_empty_poc / "verification-evidence.json").write_text(
            json.dumps(empty_poc_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_empty_poc),
            "--language",
            "zh-CN",
        ], plugin_root, "verification-evidence.json poc_path must not be empty")

        bad_absolute_poc = zh_bundle.parent / f"{zh_bundle.name}_absolute_poc_path"
        shutil.copytree(zh_bundle, bad_absolute_poc)
        absolute_poc_data = json.loads((bad_absolute_poc / "verification-evidence.json").read_text(encoding="utf-8"))
        absolute_poc_data["poc_path"] = str((bad_absolute_poc / "attachments/poc/path_traversal.py").resolve())
        (bad_absolute_poc / "verification-evidence.json").write_text(
            json.dumps(absolute_poc_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_absolute_poc),
            "--language",
            "zh-CN",
        ], plugin_root, "must be bundle-relative")

        bad_escape_poc = zh_bundle.parent / f"{zh_bundle.name}_escape_poc_path"
        shutil.copytree(zh_bundle, bad_escape_poc)
        escape_poc_data = json.loads((bad_escape_poc / "verification-evidence.json").read_text(encoding="utf-8"))
        escape_poc_data["poc_path"] = "../outside-poc.py"
        (bad_escape_poc / "verification-evidence.json").write_text(
            json.dumps(escape_poc_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_escape_poc),
            "--language",
            "zh-CN",
        ], plugin_root, "must not escape the bundle with '..'")

        bad_symlink_escape = zh_bundle.parent / f"{zh_bundle.name}_symlink_escape"
        shutil.copytree(zh_bundle, bad_symlink_escape)
        outside_file = zh_bundle.parent / "outside-symlink-target.log"
        outside_file.write_text("outside evidence\n", encoding="utf-8")
        symlink_path = bad_symlink_escape / "attachments/evidence/outside-link.log"
        symlink_path.parent.mkdir(parents=True, exist_ok=True)
        symlink_path.symlink_to(outside_file)
        symlink_data = json.loads((bad_symlink_escape / "verification-evidence.json").read_text(encoding="utf-8"))
        symlink_data["evidence_files"] = ["attachments/evidence/outside-link.log"]
        (bad_symlink_escape / "verification-evidence.json").write_text(
            json.dumps(symlink_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_symlink_escape),
            "--language",
            "zh-CN",
        ], plugin_root, "escapes the bundle root")

        bad_docker_required = zh_bundle.parent / f"{zh_bundle.name}_docker_required_false"
        shutil.copytree(zh_bundle, bad_docker_required)
        docker_required_data = json.loads((bad_docker_required / "verification-evidence.json").read_text(encoding="utf-8"))
        docker_required_data["docker_required"] = False
        (bad_docker_required / "verification-evidence.json").write_text(
            json.dumps(docker_required_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_docker_required),
            "--language",
            "zh-CN",
        ], plugin_root, "docker_required must be true")

        bad_no_escalation = zh_bundle.parent / f"{zh_bundle.name}_no_severity_escalation"
        shutil.copytree(zh_bundle, bad_no_escalation)
        no_escalation_data = json.loads((bad_no_escalation / "verification-evidence.json").read_text(encoding="utf-8"))
        no_escalation_data["severity_escalation_attempted"] = False
        (bad_no_escalation / "verification-evidence.json").write_text(
            json.dumps(no_escalation_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_no_escalation),
            "--language",
            "zh-CN",
        ], plugin_root, "severity_escalation_attempted must be true")

        bad_empty_docker_command = zh_bundle.parent / f"{zh_bundle.name}_empty_docker_command"
        shutil.copytree(zh_bundle, bad_empty_docker_command)
        empty_command_data = json.loads((bad_empty_docker_command / "verification-evidence.json").read_text(encoding="utf-8"))
        empty_command_data["docker_command"] = ""
        (bad_empty_docker_command / "verification-evidence.json").write_text(
            json.dumps(empty_command_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_empty_docker_command),
            "--language",
            "zh-CN",
        ], plugin_root, "docker_command must not be empty")

        bad_placeholder_image = zh_bundle.parent / f"{zh_bundle.name}_placeholder_docker_image"
        shutil.copytree(zh_bundle, bad_placeholder_image)
        placeholder_image_data = json.loads((bad_placeholder_image / "verification-evidence.json").read_text(encoding="utf-8"))
        placeholder_image_data["docker_image"] = "project-specific Docker image or Docker Compose service"
        (bad_placeholder_image / "verification-evidence.json").write_text(
            json.dumps(placeholder_image_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_placeholder_image),
            "--language",
            "zh-CN",
        ], plugin_root, "must not use placeholder text")

        bad_missing_attachments = zh_bundle.parent / f"{zh_bundle.name}_missing_attachments"
        shutil.copytree(zh_bundle, bad_missing_attachments)
        shutil.rmtree(bad_missing_attachments / "attachments")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_missing_attachments),
            "--language",
            "zh-CN",
        ], plugin_root, "does not exist inside bundle")

        bad_multi_finding = zh_bundle.parent / f"{zh_bundle.name}_multi_finding"
        shutil.copytree(zh_bundle, bad_multi_finding)
        (bad_multi_finding / "findings.json").write_text(
            json.dumps({"findings": [{"slug": "one"}, {"slug": "two"}]}, ensure_ascii=False),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_multi_finding),
            "--language",
            "zh-CN",
        ], plugin_root, "per-bundle findings.json must describe exactly one confirmed vulnerability")

        bad_runtime_state = zh_bundle.parent / f"{zh_bundle.name}_runtime_state"
        shutil.copytree(zh_bundle, bad_runtime_state)
        (bad_runtime_state / ".omc" / "state").mkdir(parents=True)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_report_bundle.py"),
            "--bundle-dir",
            str(bad_runtime_state),
            "--language",
            "zh-CN",
        ], plugin_root, "final confirmed bundle must not contain runtime or source-control directory")

        # --- Finalization gate tests ---
        if not (workspace / "bin/finalize-audit-workspace.py").exists():
            raise SystemExit("FAILED: bootstrapped workspace is missing finalize-audit-workspace.py")

        def install_fake_docker_clean_helper(fixture: Path, *, force: bool = False) -> None:
            """Install the narrow deterministic cleanliness helper used by finalizer fixtures."""
            helper_dir = fixture / "bin"
            helper_dir.mkdir(parents=True, exist_ok=True)
            helper = helper_dir / "manage-docker-resources.py"
            if force or not helper.exists():
                helper.write_text(
                    "#!/usr/bin/env python3\n"
                    "import json\n"
                    "import sys\n"
                    "from datetime import datetime, timezone\n"
                    "from pathlib import Path\n"
                    "args = sys.argv[1:]\n"
                    "workspace = Path(args[args.index('--workspace-dir') + 1])\n"
                    "strict = '--strict' in args\n"
                    "status = workspace / 'docker/docker-cleanliness-status.json'\n"
                    "status.parent.mkdir(parents=True, exist_ok=True)\n"
                    "status.write_text(json.dumps({'schema_version': 1, 'checked_at': datetime.now(timezone.utc).isoformat(timespec='seconds').replace('+00:00', 'Z'), 'workspace': workspace.name, 'clean': True, 'strict': strict, 'counts': {}, 'note': 'selftest strict Docker cleanliness fixture'}, sort_keys=True) + '\\n', encoding='utf-8')\n",
                    encoding="utf-8",
                )
                helper.chmod(0o755)
            baseline = fixture / "docker/docker-resource-baseline.json"
            if not baseline.exists():
                baseline.parent.mkdir(parents=True, exist_ok=True)
                baseline.write_text(
                    json.dumps(
                        {
                            "schema_version": 1,
                            "captured_at": "2026-05-06T00:00:00Z",
                            "docker_available": True,
                            "images": [],
                            "volumes": [],
                            "networks": [],
                            "containers": [],
                            "build_cache": [],
                        },
                        sort_keys=True,
                    )
                    + "\n",
                    encoding="utf-8",
                )

        def install_zero_candidate_recon(fixture: Path) -> None:
            """Provide a production-valid structured zero-candidate proof."""
            recon_fixture = plugin_root / "assets/fixtures/recon-result/service"
            shutil.copytree(recon_fixture / "repo", repo_dir, dirs_exist_ok=True)
            shutil.copytree(recon_fixture / "workspace", fixture, dirs_exist_ok=True)
            shutil.copy2(
                fixture / "cases/complete-service.json",
                fixture / "recon-result.json",
            )

        def install_source_bound_confirmed_chain(fixture: Path) -> str:
            """Build one deterministic bundle whose verifier path is workspace-bound."""
            slug = build_wrapper_source_finding(
                plugin_root,
                repo_dir,
                fixture,
                slug="p9-source-bound-demo-app_Path_Traversal_高危漏洞报告",
            )
            contract_path = build_wrapper_contract(fixture, slug)
            contract = json.loads(contract_path.read_text(encoding="utf-8"))
            binding = contract["source_binding"]
            binding["materials"]["verifier_verdict"] = "verifier/CAND-0001/verifier-verdict.json"
            write_json_fixture(contract_path, contract)
            old_verdict = fixture / "verifier/verifier-verdict.json"
            if old_verdict.exists():
                old_verdict.unlink()
            tested_ref = str(binding["tested_ref"])
            target_ref = {
                "target_config": str(binding["materials"]["target_config"]),
                "tested_ref": tested_ref,
            }
            candidate = valid_candidate_contract(
                {
                    "target_ref": target_ref,
                    "evidence": {
                        "static_locations": [
                            {
                                "path": "src/importer.py",
                                "start_line": 1,
                                "end_line": 2,
                                "reason": "The source-bound fixture reaches the reviewed sink.",
                            }
                        ],
                        "dynamic_evidence": [],
                    },
                }
            )
            candidate_path = fixture / "candidates/CAND-0001/candidate.json"
            candidate_path.parent.mkdir(parents=True, exist_ok=True)
            write_json_fixture(candidate_path, candidate)
            verdict = valid_verifier_verdict(
                {
                    "candidate_id": "CAND-0001",
                    "target_ref": target_ref,
                }
            )
            verdict_path = fixture / "verifier/CAND-0001/verifier-verdict.json"
            verdict_path.parent.mkdir(parents=True, exist_ok=True)
            write_json_fixture(verdict_path, verdict)
            run(
                [
                    sys.executable,
                    str(plugin_root / "scripts/build_confirmed_bundle.py"),
                    "--workspace-dir",
                    str(fixture),
                    "--repo-root",
                    str(repo_dir),
                    "--contract",
                    str(contract_path),
                    "--language",
                    "zh-CN",
                ],
                plugin_root,
            )
            run(
                [
                    sys.executable,
                    str(plugin_root / "scripts/audit_disposition.py"),
                    "--workspace-dir",
                    str(fixture),
                    "--candidate",
                    str(candidate_path.relative_to(fixture)),
                    "--verdict",
                    str(verdict_path.relative_to(fixture)),
                    "--update-from-verdict",
                    "--write",
                ],
                plugin_root,
            )
            return slug

        real_docker_clean_helper = workspace / "bin/manage-docker-resources.py"
        real_docker_clean_helper_bytes = real_docker_clean_helper.read_bytes()
        real_docker_clean_helper_mode = real_docker_clean_helper.stat().st_mode & 0o777
        install_fake_docker_clean_helper(workspace, force=True)
        install_zero_candidate_recon(workspace)

        def seed_r2_finalization_path(fixture: Path, label: str) -> None:
            """Give minimal finalization fixtures the normal R2 workflow prefix.

            The finalizer is a terminal producer, not a bootstrapper.  A fresh
            R2 workspace must therefore enter intake through start and reach
            verification through the explicit forward graph before the
            finalizer records its finalization advance.
            """
            writer = plugin_root / "scripts" / "write_audit_event.py"
            for event_name, stage, transition_kind in (
                (f"{label}_intake_started", "intake", "start"),
                (f"{label}_recon_started", "recon", "advance"),
                (f"{label}_candidates_started", "candidate_generation", "advance"),
                (f"{label}_triage_started", "triage", "advance"),
                (f"{label}_verification_started", "verification", "advance"),
            ):
                run([
                    sys.executable,
                    str(writer),
                    "--workspace-dir", str(fixture),
                    "--event", event_name,
                    "--stage", stage,
                    "--status", "running",
                    "--transition-kind", transition_kind,
                    "--message", "Seed a legal R2 workflow prefix for the bounded finalization fixture.",
                    "--accept-current-revision",
                ], plugin_root)
            install_fake_docker_clean_helper(fixture)

        def disposition_item(
            *,
            item_id: str,
            state: str,
            source_type: str,
            docker_status: str,
            reason_code: str,
            confirmed_bundle_path: str = "",
            docker_applicable: bool = True,
            title: str = "selftest disposition item",
            materiality_rationale: str = "selftest material item",
        ) -> dict:
            return {
                "id": item_id,
                "title": title,
                "state": state,
                "source_type": source_type,
                "docker_applicable": docker_applicable,
                "docker_status": docker_status,
                "reason_code": reason_code,
                "confirmed_bundle_path": confirmed_bundle_path,
                "materiality_rationale": materiality_rationale,
            }

        def make_disposition_fixture(
            name: str,
            *,
            items: list[dict],
            copy_valid_bundle: bool,
        ) -> Path:
            fixture = repo_dir / name
            if fixture.exists():
                shutil.rmtree(fixture)
            (fixture / "confirmed").mkdir(parents=True, exist_ok=True)
            (fixture / "docker").mkdir(parents=True, exist_ok=True)
            (fixture / "asr-config.json").write_text(
                json.dumps({
                    "workspace_root": fixture.name,
                    "workspace_created_at": "2026-05-06T00:00:00Z",
                    "confirmed_output_dir": f"{fixture.name}/confirmed",
                }, indent=2),
                encoding="utf-8",
            )
            for filename, heading in (
                ("candidate-findings.md", "# Candidate Findings\n\n"),
                ("false-positives.md", "# False Positives and Non-Security Defects\n\n"),
                ("unverified-leads.md", "# Unverified Leads\n\n"),
                ("attack-surface.md", "# Attack Surface Handoff\n\n"),
            ):
                (fixture / filename).write_text(heading, encoding="utf-8")
            if copy_valid_bundle:
                shutil.copytree(zh_bundle, fixture / "confirmed" / zh_bundle.name)
            (fixture / "audit-disposition.json").write_text(
                json.dumps(
                    {
                        "schema_version": 1,
                        "generated_at": "2026-05-06T00:00:01Z",
                        "workspace": fixture.name,
                        "items": items,
                    },
                    ensure_ascii=False,
                    indent=2,
                    sort_keys=True,
                ) + "\n",
                encoding="utf-8",
            )
            return fixture

        valid_bundle_rel = f"confirmed/{zh_bundle.name}"
        valid_disposition_workspace = make_disposition_fixture(
            "security-research-disposition-valid",
            items=[
                disposition_item(
                    item_id="confirmed:selftest",
                    state="confirmed",
                    source_type="hybrid",
                    docker_status="reproduced",
                    reason_code="docker_reproduced",
                    confirmed_bundle_path=valid_bundle_rel,
                    materiality_rationale="Docker reproduction succeeded and bundle validation passes.",
                )
            ],
            copy_valid_bundle=True,
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/audit_disposition.py"),
            "--workspace-dir",
            str(valid_disposition_workspace),
            "--result",
            "completed_with_confirmed_bundles",
        ], plugin_root)
        for source_type, reason_code in (
            ("scanner", "scanner_only"),
            ("dependency", "dependency_only"),
            ("static", "static_only"),
            ("llm", "llm_only"),
        ):
            source_only_workspace = make_disposition_fixture(
                f"security-research-disposition-{source_type}",
                items=[
                    disposition_item(
                        item_id=f"confirmed:{source_type}",
                        state="confirmed",
                        source_type=source_type,
                        docker_status="reproduced",
                        reason_code=reason_code,
                        confirmed_bundle_path=valid_bundle_rel,
                    )
                ],
                copy_valid_bundle=True,
            )
            run_expect_fail([
                sys.executable,
                str(plugin_root / "scripts/audit_disposition.py"),
                "--workspace-dir",
                str(source_only_workspace),
                "--result",
                "completed_with_confirmed_bundles",
            ], plugin_root, f"source_type={source_type} cannot be confirmed")
        not_reproduced_workspace = make_disposition_fixture(
            "security-research-disposition-not-reproduced",
            items=[
                disposition_item(
                    item_id="confirmed:not-reproduced",
                    state="confirmed",
                    source_type="hybrid",
                    docker_status="failed",
                    reason_code="not_reproducible",
                    confirmed_bundle_path=valid_bundle_rel,
                )
            ],
            copy_valid_bundle=True,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/audit_disposition.py"),
            "--workspace-dir",
            str(not_reproduced_workspace),
            "--result",
            "completed_with_confirmed_bundles",
        ], plugin_root, "requires docker_status=reproduced")
        non_confirmed_points_to_confirmed = make_disposition_fixture(
            "security-research-disposition-non-confirmed-path",
            items=[
                disposition_item(
                    item_id="candidate:bad-path",
                    state="candidate",
                    source_type="manual",
                    docker_status="not_started",
                    reason_code="insufficient_evidence",
                    confirmed_bundle_path=valid_bundle_rel,
                )
            ],
            copy_valid_bundle=True,
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/audit_disposition.py"),
            "--workspace-dir",
            str(non_confirmed_points_to_confirmed),
            "--result",
            "completed_with_confirmed_bundles",
        ], plugin_root, "non-confirmed items must not point into confirmed/")
        for docker_status, reason_code in (
            ("blocked", "blocked_by_docker"),
            ("timed_out", "timed_out"),
            ("dirty_state", "dirty_docker"),
        ):
            blocking_workspace = make_disposition_fixture(
                f"security-research-disposition-{docker_status}",
                items=[
                    disposition_item(
                        item_id=f"candidate:{docker_status}",
                        state="candidate",
                        source_type="runtime",
                        docker_status=docker_status,
                        reason_code=reason_code,
                    )
                ],
                copy_valid_bundle=False,
            )
            run_expect_fail([
                sys.executable,
                str(plugin_root / "scripts/audit_disposition.py"),
                "--workspace-dir",
                str(blocking_workspace),
                "--result",
                "completed_no_confirmed_findings",
            ], plugin_root, "blocks completed_no_confirmed_findings")

        def write_integrity_fixture(
            fixture_workspace: Path,
            *,
            events: list[dict],
            status: dict,
            docker_status: dict,
            ledger: dict | None = None,
        ) -> None:
            fixture_workspace.mkdir(parents=True, exist_ok=True)
            (fixture_workspace / "docker").mkdir(parents=True, exist_ok=True)
            (fixture_workspace / "asr-config.json").write_text(
                json.dumps({
                    "workspace_root": fixture_workspace.name,
                    "workspace_created_at": "2026-05-06T00:00:00Z",
                    "confirmed_output_dir": f"{fixture_workspace.name}/confirmed",
                }, indent=2),
                encoding="utf-8",
            )
            canonical_status = {
                "schema_version": 1,
                "plugin": "zhulong",
                "plugin_version": "0.4.0-selftest",
                "last_event_at": "2026-05-06T00:00:01Z",
                "blocker": None,
                "resume_step": None,
                "workspace": fixture_workspace.name,
                "target_repo": ".",
                **status,
            }
            (fixture_workspace / "stage-status.json").write_text(
                json.dumps(canonical_status, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            canonical_docker_status = {
                "schema_version": 1,
                "checked_at": "2026-05-06T00:00:01Z",
                "workspace": fixture_workspace.name,
                "clean": bool(docker_status.get("clean")),
                "strict": bool(docker_status.get("strict")),
                "counts": docker_status.get("counts") if isinstance(docker_status.get("counts"), dict) else {},
                "note": str(docker_status.get("note") or "selftest strict Docker cleanliness fixture"),
            }
            (fixture_workspace / "docker/docker-cleanliness-status.json").write_text(
                json.dumps(canonical_docker_status, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            docker_digest = "sha256:" + hashlib.sha256((fixture_workspace / "docker/docker-cleanliness-status.json").read_bytes()).hexdigest()
            normalized_events = []
            for event in events:
                event = dict(event)
                if event.get("event") == "finalization_succeeded":
                    details = dict(event.get("details") or {})
                    details.update({
                        "docker_clean": True,
                        "docker_clean_strict": True,
                        "docker_cleanliness_path": "docker/docker-cleanliness-status.json",
                        "docker_cleanliness_sha256": docker_digest,
                        "docker_cleanliness_checked_at": canonical_docker_status["checked_at"],
                        "docker_cleanliness_workspace": fixture_workspace.name,
                    })
                    event["details"] = details
                normalized_events.append(event)
            (fixture_workspace / "audit-events.jsonl").write_text(
                "".join(json.dumps(event, ensure_ascii=False, sort_keys=True) + "\n" for event in normalized_events),
                encoding="utf-8",
            )
            if ledger is not None:
                (fixture_workspace / "audit-disposition.json").write_text(
                    json.dumps(ledger, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                    encoding="utf-8",
                )

        bad_integrity_workspace = repo_dir / "security-research-integrity-bad"
        write_integrity_fixture(
            bad_integrity_workspace,
            events=[
                {
                    "ts": "2026-05-06T00:00:01Z",
                    "event": "finalization_failed",
                    "stage": "finalization",
                    "status": "failed",
                    "message": "Completion gate failed.",
                    "details": {"expected_result": "completed_no_confirmed_findings"},
                }
            ],
            status={
                "stage": "completed",
                "status": "completed",
                "result": "completed_no_confirmed_findings",
                "completed_at": "2026-05-06T00:00:02Z",
            },
            docker_status={
                "schema_version": 1,
                "clean": False,
                "strict": True,
                "workspace": bad_integrity_workspace.name,
            },
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/assert_finalized_workspace.py"),
            "--workspace-dir",
            str(bad_integrity_workspace),
        ], plugin_root, "latest finalization event is finalization_failed")
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(bad_integrity_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            bad_integrity_workspace / "handoff-summary.md",
            "Finalization integrity: `blocked`",
            "handoff blocks manually completed failed finalization",
        )
        require_text(
            bad_integrity_workspace / "handoff-summary.md",
            "Completion gate passed: `false`",
            "handoff does not claim completion gate passed after failed finalization",
        )

        good_integrity_workspace = repo_dir / "security-research-integrity-good"
        recon_fixture = plugin_root / "assets/fixtures/recon-result/service"
        shutil.copytree(recon_fixture / "repo", repo_dir, dirs_exist_ok=True)
        shutil.copytree(recon_fixture / "workspace", good_integrity_workspace, dirs_exist_ok=True)
        write_integrity_fixture(
            good_integrity_workspace,
            events=[
                {
                    "ts": "2026-05-06T00:00:01Z",
                    "event": "finalization_succeeded",
                    "stage": "completed",
                    "status": "ok",
                    "message": "Audit finalized.",
                    "details": {
                        "result": "completed_no_confirmed_findings",
                        "docker_clean": True,
                        "validated_bundles": 0,
                    },
                }
            ],
            status={
                "stage": "completed",
                "status": "completed",
                "result": "completed_no_confirmed_findings",
            },
            docker_status={
                "schema_version": 1,
                "clean": True,
                "strict": True,
                "workspace": good_integrity_workspace.name,
            },
            ledger={
                "schema_version": 1,
                "generated_at": "2026-05-06T00:00:01Z",
                "workspace": good_integrity_workspace.name,
                "items": [],
            },
        )
        shutil.copy2(
            good_integrity_workspace / "cases/complete-service.json",
            good_integrity_workspace / "recon-result.json",
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/assert_finalized_workspace.py"),
            "--workspace-dir",
            str(good_integrity_workspace),
        ], plugin_root)
        integrity_json = json.loads(run_capture([
            sys.executable,
            str(plugin_root / "scripts/assert_finalized_workspace.py"),
            "--workspace-dir",
            str(good_integrity_workspace),
            "--json",
        ], plugin_root))
        if integrity_json.get("ok") is not True:
            raise SystemExit("FAILED: finalization integrity JSON did not pass for valid completion fixture")

        def write_issue19_workspace(
            name: str,
            *,
            result: str = "",
            status_stage: str = "verification",
            status_value: str = "running",
            handoff: str = "",
            copy_bundle: bool = False,
            partial_bundle: bool = False,
            docker_evidence: bool = False,
        ) -> Path:
            fixture = repo_dir / name
            if fixture.exists():
                shutil.rmtree(fixture)
            (fixture / "confirmed").mkdir(parents=True, exist_ok=True)
            (fixture / "docker").mkdir(parents=True, exist_ok=True)
            (fixture / "asr-config.json").write_text(
                json.dumps({
                    "workspace_root": fixture.name,
                    "workspace_created_at": "2026-05-06T00:00:00Z",
                    "confirmed_output_dir": f"{fixture.name}/confirmed",
                }, indent=2),
                encoding="utf-8",
            )
            status_doc = {
                "schema_version": 1,
                "plugin": "zhulong",
                "plugin_version": "selftest",
                "stage": status_stage,
                "status": status_value,
                "last_event_at": "2026-05-06T00:00:00Z",
                "blocker": "selftest blocker" if status_value in {"blocked", "paused"} else None,
                "resume_step": "resume selftest" if status_value in {"blocked", "paused"} else None,
                "workspace": fixture.name,
                "target_repo": str(repo_dir),
            }
            if result:
                status_doc["result"] = result
            (fixture / "stage-status.json").write_text(
                json.dumps(status_doc, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            (fixture / "audit-events.jsonl").write_text(
                json.dumps({
                    "ts": "2026-05-06T00:00:00Z",
                    "event": "selftest_fixture",
                    "stage": status_stage,
                    "status": status_value,
                    "message": "Issue 19 handoff consistency fixture.",
                    "details": {},
                }, ensure_ascii=False, sort_keys=True) + "\n",
                encoding="utf-8",
            )
            for filename, heading in (
                ("candidate-findings.md", "# Candidate Findings\n\n"),
                ("false-positives.md", "# False Positives\n\n"),
                ("unverified-leads.md", "# Unverified Leads\n\n"),
                ("attack-surface.md", "# Attack Surface Handoff\n\n"),
            ):
                (fixture / filename).write_text(heading, encoding="utf-8")
            if handoff:
                (fixture / "handoff-summary.md").write_text(handoff, encoding="utf-8")
            if copy_bundle:
                shutil.copytree(zh_bundle, fixture / "confirmed" / zh_bundle.name)
            if partial_bundle:
                partial = fixture / "confirmed" / "partial-docker-evidence"
                partial.mkdir()
                (partial / "verification-evidence.json").write_text(
                    json.dumps({"verification_status": "confirmed_in_docker"}, indent=2),
                    encoding="utf-8",
                )
            if docker_evidence:
                evidence_dir = fixture / "evidence" / "docker-evidence" / "case-1"
                evidence_dir.mkdir(parents=True, exist_ok=True)
                (evidence_dir / "verification-evidence.json").write_text(
                    json.dumps({"verification_status": "confirmed_in_docker"}, indent=2),
                    encoding="utf-8",
                )
            install_fake_docker_clean_helper(fixture)
            return fixture

        contradictory_handoff_workspace = write_issue19_workspace(
            "security-research-issue19-contradictory-handoff",
            handoff=(
                "# Handoff Summary\n\n"
                "- Confirmed bundles: 1\n"
                "- Formal seeded variant discovery completed and ready.\n"
                "- Docker evidence directory is a completed confirmed bundle.\n"
            ),
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir",
            str(contradictory_handoff_workspace),
            "--repo-root",
            str(repo_dir),
            "--skip-latest-check",
        ], plugin_root, "handoff/status consistency failed")
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(contradictory_handoff_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            contradictory_handoff_workspace / "handoff-summary.md",
            "Confirmed bundles: 0",
            "Issue 19 renderer rewrites zero-bundle count",
        )
        require_text(
            contradictory_handoff_workspace / "handoff-summary.md",
            "Status: `not_applicable_no_validated_confirmed_bundle`",
            "Issue 19 renderer makes formal variant not applicable",
        )
        forbid_text(
            contradictory_handoff_workspace / "handoff-summary.md",
            "Formal seeded variant discovery completed and ready",
            "Issue 19 renderer removes stale formal variant completion claim",
        )

        status_claim_workspace = write_issue19_workspace(
            "security-research-issue19-status-claim",
            result="completed_with_confirmed_bundles",
            status_stage="completed",
            status_value="completed",
            handoff="# Handoff Summary\n\n- State: completed_with_confirmed_bundles\n",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir",
            str(status_claim_workspace),
            "--repo-root",
            str(repo_dir),
            "--skip-latest-check",
        ], plugin_root, "validated_confirmed_bundle_count=0")

        docker_evidence_workspace = write_issue19_workspace(
            "security-research-issue19-docker-evidence-only",
            docker_evidence=True,
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(docker_evidence_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            docker_evidence_workspace / "handoff-summary.md",
            "State: `docker_evidence_collected_but_no_bundle`",
            "Issue 19 Docker evidence-only state",
        )
        require_text(
            docker_evidence_workspace / "handoff-summary.md",
            "Confirmed bundles: 0",
            "Issue 19 Docker evidence-only does not count as bundle",
        )

        partial_issue19_workspace = write_issue19_workspace(
            "security-research-issue19-partial-confirmed",
            partial_bundle=True,
        )
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(partial_issue19_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            partial_issue19_workspace / "handoff-summary.md",
            "Confirmed bundles: 0",
            "Issue 19 partial bundle is not validated",
        )
        require_text(
            partial_issue19_workspace / "handoff-summary.md",
            "Invalid or partial bundle directories: 1",
            "Issue 19 partial bundle count",
        )

        code_level_only_workspace = write_issue19_workspace(
            "security-research-issue19-code-level-only",
            handoff="# Handoff Summary\n\n- Code-level evidence reproduced; bundle-ready for confirmed output.\n",
        )
        (code_level_only_workspace / "evidence" / "code-level").mkdir(parents=True, exist_ok=True)
        (code_level_only_workspace / "evidence" / "code-level" / "verification-evidence.json").write_text(
            json.dumps({"verification_status": "code_level_reproduced"}, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/validate_workspace_state.py"),
            "--workspace-dir",
            str(code_level_only_workspace),
            "--repo-root",
            str(repo_dir),
            "--skip-latest-check",
        ], plugin_root, "bundle readiness")
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(code_level_only_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            code_level_only_workspace / "handoff-summary.md",
            "Confirmed bundles: 0",
            "Issue 19 code-level-only handoff remains no bundle",
        )

        valid_variant_workspace = write_issue19_workspace(
            "security-research-issue19-valid-variant",
            copy_bundle=True,
        )
        write_finalization_variant_artifacts(valid_variant_workspace)
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(valid_variant_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            valid_variant_workspace / "handoff-summary.md",
            "Confirmed bundles: 1",
            "Issue 19 valid bundle counted",
        )
        require_text(
            valid_variant_workspace / "handoff-summary.md",
            "Status: `completed`",
            "Issue 19 valid formal variant analysis can report completed",
        )

        invalid_variant_workspace = write_issue19_workspace(
            "security-research-issue19-invalid-variant",
            copy_bundle=True,
        )
        invalid_variant_dir = invalid_variant_workspace / "evidence" / "variant-analysis"
        invalid_variant_dir.mkdir(parents=True, exist_ok=True)
        write_variant_seed_card(
            invalid_variant_dir / "seeds.jsonl",
            {"confirmed_bundle_path": "evidence/docker-evidence/case-1"},
        )
        write_variant_candidates_jsonl(
            invalid_variant_dir / "variant-candidates.jsonl",
            [valid_variant_candidate_record()],
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(invalid_variant_workspace),
            "--result",
            "completed_with_confirmed_bundles",
        ], plugin_root, "variant seed confirmed_bundle_path must not point to candidate/manual/evidence-only material",
           extra_env={})

        blocked_finalization_workspace = repo_dir / "security-research-blocked-verification"
        blocked_finalization_workspace.mkdir(parents=True, exist_ok=True)
        (blocked_finalization_workspace / "confirmed").mkdir()
        (blocked_finalization_workspace / "docker").mkdir()
        (blocked_finalization_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": blocked_finalization_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{blocked_finalization_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        (blocked_finalization_workspace / "candidate-findings.md").write_text(
            "# Candidate Findings\n\n"
            "| Candidate ID | Suspected Weakness | Evidence So Far | Source-to-Sink Hypothesis | Docker Verification Plan | Status |\n"
            "| --- | --- | --- | --- | --- | --- |\n"
            "| C1 | SSRF | curl_exec sink | webhook url -> curl_exec | start runtime and test internal service | BLOCKED (Docker rate limit) |\n",
            encoding="utf-8",
        )
        (blocked_finalization_workspace / "unverified-leads.md").write_text("# Unverified Leads\n\n", encoding="utf-8")
        (blocked_finalization_workspace / "attack-surface.md").write_text(
            "# Attack Surface Handoff\n\n## Docker Verification Status\n\n"
            "- Running service target: BLOCKED - Docker Hub rate limit, no cached images.\n",
            encoding="utf-8",
        )
        (blocked_finalization_workspace / "docker/docker-cleanliness-status.json").write_text(
            json.dumps({"schema_version": 1, "clean": True, "strict": True}, indent=2),
            encoding="utf-8",
        )
        seed_r2_finalization_path(blocked_finalization_workspace, "blocked_finalization")
        blocked_proc = subprocess.run([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(blocked_finalization_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], cwd=plugin_root, capture_output=True, text=True, env=os.environ.copy())
        blocked_output = (blocked_proc.stdout or "") + (blocked_proc.stderr or "")
        if blocked_proc.returncode == 0:
            raise SystemExit("FAILED: blocked verification finalized as completed_no_confirmed_findings")
        for expected in (
            "Blocked Docker/runtime verification prevents completed_no_confirmed_findings",
            "blocked_verification",
            "docker login",
            "rerun Docker verification",
        ):
            if expected not in blocked_output:
                raise SystemExit(f"FAILED: blocked finalization output missing: {expected}\n{blocked_output}")
        blocked_events = (blocked_finalization_workspace / "audit-events.jsonl").read_text(encoding="utf-8")
        if "finalization_succeeded" in blocked_events:
            raise SystemExit("FAILED: blocked verification wrote finalization_succeeded")
        blocked_status = json.loads((blocked_finalization_workspace / "stage-status.json").read_text(encoding="utf-8"))
        if blocked_status.get("status") != "blocked" or blocked_status.get("blocker") != "blocked_verification":
            raise SystemExit("FAILED: blocked verification did not leave stage-status.json in blocked state")
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/assert_finalized_workspace.py"),
            "--workspace-dir",
            str(blocked_finalization_workspace),
        ], plugin_root, "stage-status.json does not declare a completed workspace")
        run([
            sys.executable,
            str(plugin_root / "scripts/render_handoff_summary.py"),
            "--workspace-dir",
            str(blocked_finalization_workspace),
            "--repo-root",
            str(repo_dir),
        ], plugin_root)
        require_text(
            blocked_finalization_workspace / "handoff-summary.md",
            "Blocked verification: `blocked_verification`",
            "handoff surfaces blocked verification",
        )

        stale_blocker_workspace = repo_dir / "security-research-stale-blocker"
        stale_blocker_workspace.mkdir(parents=True, exist_ok=True)
        (stale_blocker_workspace / "confirmed").mkdir()
        (stale_blocker_workspace / "docker").mkdir()
        (stale_blocker_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": stale_blocker_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{stale_blocker_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        (stale_blocker_workspace / "candidate-findings.md").write_text("# Candidate Findings\n\n", encoding="utf-8")
        (stale_blocker_workspace / "unverified-leads.md").write_text("# Unverified Leads\n\n", encoding="utf-8")
        (stale_blocker_workspace / "attack-surface.md").write_text(
            "# Attack Surface Handoff\n\n## Docker Verification Status\n\n"
            "- Docker gate: ready\n"
            "- Running service target: NOT STARTED (images being pulled)\n"
            "- Still blocked or missing: Image pull required\n",
            encoding="utf-8",
        )
        seed_r2_finalization_path(stale_blocker_workspace, "stale_blocker")
        stale_proc = subprocess.run([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(stale_blocker_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], cwd=plugin_root, capture_output=True, text=True, env=os.environ.copy())
        stale_output = (stale_proc.stdout or "") + (stale_proc.stderr or "")
        if stale_proc.returncode == 0:
            raise SystemExit("FAILED: stale attack-surface Docker blocker finalized as no-confirmed")
        for expected in ("images being pulled", "Blocked Docker/runtime verification prevents completed_no_confirmed_findings"):
            if expected not in stale_output:
                raise SystemExit(f"FAILED: stale blocker finalization output missing: {expected}\n{stale_output}")

        high_confidence_blocked_workspace = repo_dir / "security-research-high-confidence-blocked"
        high_confidence_blocked_workspace.mkdir(parents=True, exist_ok=True)
        (high_confidence_blocked_workspace / "confirmed").mkdir()
        (high_confidence_blocked_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": high_confidence_blocked_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{high_confidence_blocked_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        (high_confidence_blocked_workspace / "candidate-findings.md").write_text("# Candidate Findings\n\n", encoding="utf-8")
        (high_confidence_blocked_workspace / "attack-surface.md").write_text("# Attack Surface Handoff\n\n", encoding="utf-8")
        (high_confidence_blocked_workspace / "unverified-leads.md").write_text(
            "# Unverified Leads\n\n"
            "| Lead ID | Suspected Weakness | Evidence So Far | Missing Evidence | Docker Confirmation Status | Safe Resume Step | High-Confidence-Unverified? |\n"
            "| --- | --- | --- | --- | --- | --- | --- |\n"
            "| U1 | Optional Kafka TLS | rejectUnauthorized:false | deployment materiality | blocked_no_docker | configure Kafka and rerun Docker verification | Yes |\n",
            encoding="utf-8",
        )
        seed_r2_finalization_path(high_confidence_blocked_workspace, "high_confidence_blocked")
        high_conf_blocked_proc = subprocess.run([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(high_confidence_blocked_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], cwd=plugin_root, capture_output=True, text=True, env=os.environ.copy())
        high_conf_output = (high_conf_blocked_proc.stdout or "") + (high_conf_blocked_proc.stderr or "")
        if high_conf_blocked_proc.returncode == 0:
            raise SystemExit("FAILED: high-confidence blocked lead without materiality finalized as no-confirmed")
        if "Material blocker?" not in high_conf_output:
            raise SystemExit("FAILED: high-confidence blocked lead failure did not request materiality rationale")

        high_confidence_safe_workspace = repo_dir / "security-research-high-confidence-safe"
        high_confidence_safe_workspace.mkdir(parents=True, exist_ok=True)
        (high_confidence_safe_workspace / "confirmed").mkdir()
        (high_confidence_safe_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": high_confidence_safe_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{high_confidence_safe_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        install_zero_candidate_recon(high_confidence_safe_workspace)
        (high_confidence_safe_workspace / "candidate-findings.md").write_text("# Candidate Findings\n\n", encoding="utf-8")
        (high_confidence_safe_workspace / "unverified-leads.md").write_text(
            "# Unverified Leads\n\n"
            "| Lead ID | Suspected Weakness | Evidence So Far | Missing Evidence | Docker Confirmation Status | Safe Resume Step | High-Confidence-Unverified? | Material blocker? | Default runtime scope? | Why completion is still safe? |\n"
            "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |\n"
            "| U1 | Optional Kafka TLS | rejectUnauthorized:false | deployment materiality | blocked_no_docker | configure Kafka and rerun Docker verification | Yes | No | optional integration | Kafka is disabled in default runtime and this is a non-material optional integration follow-up. |\n",
            encoding="utf-8",
        )
        # The Markdown lead remains visible for handoff, but completion may
        # proceed only after the disposition ledger explicitly records the
        # non-material lead as false_positive.
        write_json_fixture(
            high_confidence_safe_workspace / "audit-disposition.json",
            {
                "schema_version": 1,
                "generated_at": "2026-05-06T00:00:01Z",
                "workspace": high_confidence_safe_workspace.name,
                "candidate_dispositions": [],
                "items": [
                    {
                        "id": "unverified-leads:u1",
                        "title": "Optional Kafka TLS",
                        "state": "false_positive",
                        "source_type": "runtime",
                        "docker_applicable": False,
                        "docker_status": "not_applicable",
                        "reason_code": "safe_config",
                        "confirmed_bundle_path": "",
                        "materiality_rationale": "Non-material optional integration; safe configuration disposition recorded explicitly.",
                    }
                ],
            },
        )
        seed_r2_finalization_path(high_confidence_safe_workspace, "high_confidence_safe")
        run_with_env([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(high_confidence_safe_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], plugin_root, {})
        require_text(
            high_confidence_safe_workspace / "SUMMARY.md",
            "completed_no_confirmed_findings",
            "finalization creates stable workspace SUMMARY.md",
        )

        legacy_clean_workspace = repo_dir / "security-research-legacy-clean-no-ledger"
        legacy_clean_workspace.mkdir(parents=True, exist_ok=True)
        (legacy_clean_workspace / "confirmed").mkdir()
        (legacy_clean_workspace / "docker").mkdir()
        (legacy_clean_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": legacy_clean_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{legacy_clean_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        install_zero_candidate_recon(legacy_clean_workspace)
        (legacy_clean_workspace / "docker/docker-cleanliness-status.json").write_text(
            json.dumps({"schema_version": 1, "clean": True, "strict": True}, indent=2),
            encoding="utf-8",
        )
        seed_r2_finalization_path(legacy_clean_workspace, "legacy_clean")
        if (legacy_clean_workspace / "audit-disposition.json").exists():
            raise SystemExit("FAILED: legacy clean fixture unexpectedly started with audit-disposition.json")
        run_with_env([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(legacy_clean_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], plugin_root, {})
        if not (legacy_clean_workspace / "audit-disposition.json").exists():
            raise SystemExit("FAILED: legacy clean no-confirmed finalization did not write audit-disposition.json")
        run([
            sys.executable,
            str(plugin_root / "scripts/assert_finalized_workspace.py"),
            "--workspace-dir",
            str(legacy_clean_workspace),
        ], plugin_root)

        stale_docker_status_workspace = repo_dir / "security-research-stale-docker-status"
        stale_docker_status_workspace.mkdir(parents=True, exist_ok=True)
        (stale_docker_status_workspace / "confirmed").mkdir()
        (stale_docker_status_workspace / "docker").mkdir()
        (stale_docker_status_workspace / "bin").mkdir()
        (stale_docker_status_workspace / "asr-config.json").write_text(
            json.dumps({
                "workspace_root": stale_docker_status_workspace.name,
                "workspace_created_at": "2026-05-06T00:00:00Z",
                "confirmed_output_dir": f"{stale_docker_status_workspace.name}/confirmed",
            }, indent=2),
            encoding="utf-8",
        )
        for filename, heading in (
            ("candidate-findings.md", "# Candidate Findings\n\n"),
            ("unverified-leads.md", "# Unverified Leads\n\n"),
            ("false-positives.md", "# False Positives\n\n"),
            ("attack-surface.md", "# Attack Surface Handoff\n\n"),
        ):
            (stale_docker_status_workspace / filename).write_text(heading, encoding="utf-8")
        install_zero_candidate_recon(stale_docker_status_workspace)
        seed_r2_finalization_path(stale_docker_status_workspace, "stale_docker_status")
        (stale_docker_status_workspace / "docker/docker-resource-baseline.json").write_text(
            json.dumps({
                "schema_version": 1,
                "captured_at": "2026-05-06T00:00:00Z",
                "docker_available": True,
                "images": [],
                "volumes": [],
                "networks": [],
                "containers": [],
                "build_cache": [],
            }, indent=2),
            encoding="utf-8",
        )
        stale_status_path = stale_docker_status_workspace / "docker/docker-cleanliness-status.json"
        stale_status_path.write_text(
            json.dumps({
                "schema_version": 1,
                "checked_at": "2026-05-06T00:00:01Z",
                "workspace": stale_docker_status_workspace.name,
                "clean": True,
                "strict": True,
                "counts": {},
                "note": "stale selftest fixture",
            }, indent=2),
            encoding="utf-8",
        )
        fake_manage = stale_docker_status_workspace / "bin/manage-docker-resources.py"
        fake_manage.write_text(
            "#!/usr/bin/env python3\n"
            "import sys\n"
            "print('simulated Docker helper failure without status refresh', file=sys.stderr)\n"
            "raise SystemExit(1)\n",
            encoding="utf-8",
        )
        fake_manage.chmod(0o755)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir",
            str(stale_docker_status_workspace),
            "--result",
            "completed_no_confirmed_findings",
        ], plugin_root, "Docker cleanliness check failed")

        # Test 1: Finalization with valid bundles succeeds
        # Remove partial/bad bundles first so only valid ones remain
        for bad in (
            partial_confirmed,
            bad_missing_verification,
            bad_high_confidence,
            bad_missing_evidence_file,
            bad_absolute_evidence,
            bad_escape_evidence,
            bad_empty_poc,
            bad_absolute_poc,
            bad_escape_poc,
            bad_symlink_escape,
            bad_docker_required,
            bad_no_escalation,
            bad_empty_docker_command,
            bad_placeholder_image,
            bad_missing_attachments,
            bad_multi_finding,
            bad_runtime_state,
            bad_runtime_scope,
            nested_warning_bundle,
        ):
            if bad.exists():
                shutil.rmtree(bad)
        # The preceding bundle-validator mutations intentionally leave several
        # invalid and unrelated valid bundles in confirmed/.  Remove all
        # visible bundle directories before building the one authority-bound
        # bundle used by the finalization chain; build_confirmed_bundle.py
        # validates the complete confirmed/ directory before promotion.
        for confirmed_entry in (workspace / "confirmed").iterdir():
            if confirmed_entry.is_dir() and not confirmed_entry.name.startswith("."):
                shutil.rmtree(confirmed_entry)
        source_bound_slug = install_source_bound_confirmed_chain(workspace)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--language", "auto",
            "--result", "completed_with_confirmed_bundles",
        ], plugin_root, "Seeded variant discovery gate",
           extra_env={})
        write_finalization_variant_artifacts(workspace)
        run_with_env([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--language", "auto",
            "--result", "completed_with_confirmed_bundles",
        ], plugin_root, {})
        finalized_status = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
        if finalized_status.get("status") != "completed":
            raise SystemExit("FAILED: finalization did not set stage-status.json status to completed")
        expected_final_stage = "finalization" if finalized_status.get("schema_version") == 2 else "completed"
        if finalized_status.get("stage") != expected_final_stage:
            raise SystemExit("FAILED: finalization did not set the protocol-canonical completed stage")
        if finalized_status.get("blocker") is not None:
            raise SystemExit("FAILED: finalization did not clear blocker in stage-status.json")
        if finalized_status.get("resume_step") is not None:
            raise SystemExit("FAILED: finalization did not clear resume_step in stage-status.json")
        finalized_handoff = (workspace / "handoff-summary.md").read_text(encoding="utf-8")
        if "running" in finalized_handoff.split("Status:")[1].split("\n")[0] if "Status:" in finalized_handoff else "":
            raise SystemExit("FAILED: finalized handoff-summary.md still reports running status")
        finalized_events = (workspace / "audit-events.jsonl").read_text(encoding="utf-8")
        if "finalization_succeeded" not in finalized_events:
            raise SystemExit("FAILED: finalization did not write finalization_succeeded event")
        if "finalization_started" not in finalized_events:
            raise SystemExit("FAILED: finalization did not write finalization_started event")
        if "bundle_validation_outcome" not in finalized_events:
            raise SystemExit("FAILED: finalization did not write bundle_validation_outcome event")
        if "audit_disposition_outcome" not in finalized_events:
            raise SystemExit("FAILED: finalization did not write audit_disposition_outcome event")
        finalized_ledger = json.loads((workspace / "audit-disposition.json").read_text(encoding="utf-8"))
        if not finalized_ledger.get("items"):
            raise SystemExit("FAILED: finalization did not write audit-disposition.json items for confirmed bundles")
        run([
            sys.executable,
            str(workspace / "bin/assert-finalized-workspace.py"),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)
        require_text(
            workspace / "SUMMARY.md",
            "completed_with_confirmed_bundles",
            "bundle finalization writes stable workspace SUMMARY.md",
        )
        # Test 2: Finalization with no confirmed bundles succeeds under completed_no_confirmed_findings
        def reset_finalization_test_to_verification() -> None:
            current = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
            writer = workspace / "bin/write-audit-event.py"
            common = [
                sys.executable, str(writer), "--workspace-dir", str(workspace),
                "--accept-current-revision", "--subject", "run:finalization-selftest",
                "--evidence-ref", "handoff-summary.md",
                "--next-action-json", json.dumps({
                    "action_id": "resume-finalization-selftest",
                    "action_type": "review",
                    "subject_ids": ["run:finalization-selftest"],
                    "summary": "Review the finalization test transition before continuing.",
                    "evidence_refs": ["handoff-summary.md"],
                }, sort_keys=True),
                "--details-json", json.dumps({
                    "summary": "Reset finalization selftest workflow.",
                    "reason_detail": "The selftest starts a new bounded finalization branch after validating prior output.",
                }, sort_keys=True),
            ]
            if current.get("stage") == "finalization" and current.get("status") == "completed":
                run(common + [
                    "--event", "selftest_finalization_reopened",
                    "--stage", "finalization",
                    "--status", "running",
                    "--transition-kind", "reopen",
                    "--reason-code", "validation_failed",
                    "--message", "Reopen the completed finalization test branch.",
                ], plugin_root)
                current = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
            if current.get("stage") == "finalization" and current.get("status") == "running":
                run(common + [
                    "--event", "selftest_returned_to_verification",
                    "--stage", "verification",
                    "--status", "running",
                    "--transition-kind", "return",
                    "--reason-code", "validation_failed",
                    "--message", "Return finalization selftest work to verification.",
                ], plugin_root)
                return
            if current.get("stage") != "verification" or current.get("status") != "running":
                raise SystemExit("FAILED: finalization selftest could not reset to verification/running")

        reset_finalization_test_to_verification()
        # Remove all bundle dirs to simulate no-finding workspace
        for entry in (workspace / "confirmed").iterdir():
            if entry.is_dir() and not entry.name.startswith("."):
                shutil.rmtree(entry)
        for authority_dir in (workspace / "candidates", workspace / "verifier"):
            if authority_dir.exists():
                shutil.rmtree(authority_dir)
        disposition_path = workspace / "audit-disposition.json"
        if disposition_path.exists():
            disposition_path.unlink()
        run_with_env([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_no_confirmed_findings",
        ], plugin_root, {})
        no_finding_status = json.loads((workspace / "stage-status.json").read_text(encoding="utf-8"))
        if no_finding_status.get("status") != "completed":
            raise SystemExit("FAILED: no-finding finalization did not set status to completed")
        expected_no_finding_stage = "finalization" if no_finding_status.get("schema_version") == 2 else "completed"
        if no_finding_status.get("stage") != expected_no_finding_stage:
            raise SystemExit("FAILED: no-finding finalization did not set the protocol-canonical completed stage")
        no_finding_handoff = (workspace / "handoff-summary.md").read_text(encoding="utf-8")
        if "No confirmed vulnerabilities" not in no_finding_handoff:
            raise SystemExit("FAILED: no-finding handoff does not show 'No confirmed vulnerabilities'")
        if "initial_probing" in no_finding_handoff:
            raise SystemExit("FAILED: no-finding handoff still shows stale initial_probing")
        if "running" in no_finding_handoff.split("Status:")[1].split("\n")[0] if "Status:" in no_finding_handoff else "":
            raise SystemExit("FAILED: no-finding handoff still reports running status")
        require_text(
            workspace / "SUMMARY.md",
            "completed_no_confirmed_findings",
            "no-finding finalization updates generated SUMMARY.md",
        )
        no_finding_ledger = json.loads((workspace / "audit-disposition.json").read_text(encoding="utf-8"))
        if any(item.get("state") == "confirmed" for item in no_finding_ledger.get("items", [])):
            raise SystemExit("FAILED: refreshed no-finding audit-disposition.json kept stale confirmed items")
        run([
            sys.executable,
            str(workspace / "bin/assert-finalized-workspace.py"),
            "--workspace-dir",
            str(workspace),
        ], plugin_root)

        # Test 3: Finalization fails when partial confirmed bundles exist
        reset_finalization_test_to_verification()
        partial_for_gate = workspace / "confirmed/C99-partial-gate-test"
        partial_for_gate.mkdir(parents=True, exist_ok=True)
        (partial_for_gate / "verification-evidence.json").write_text(
            json.dumps({"verification_status": "confirmed_in_docker"}, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_with_confirmed_bundles",
        ], plugin_root, "partial confirmed bundle",
           extra_env={})

        # Test 4: Finalization fails when result=completed_with_confirmed_bundles but zero bundles validate
        shutil.rmtree(partial_for_gate)
        reset_finalization_test_to_verification()
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_with_confirmed_bundles",
        ], plugin_root, "requires at least one validated confirmed bundle",
           extra_env={})

        # Test 5: Finalization fails when Docker strict cleanliness fails
        # Re-render a valid bundle for this test
        real_docker_clean_helper.write_bytes(real_docker_clean_helper_bytes)
        real_docker_clean_helper.chmod(real_docker_clean_helper_mode)
        run([
            sys.executable,
            str(workspace / "bin/render-confirmed-vuln-docx.py"),
            "--input",
            str(plugin_root / "assets/examples/confirmed-findings.example.json"),
            "--output-dir", str(workspace / "confirmed"),
            "--language", "zh-CN",
        ], plugin_root)
        for rendered_bundle in sorted(
            path for path in (workspace / "confirmed").iterdir()
            if path.is_dir() and not path.name.startswith(".")
        ):
            write_live_replay_logs(rendered_bundle)
        write_finalization_variant_artifacts(workspace)
        reset_finalization_test_to_verification()
        # Use a fake baseline that will make verify-clean fail
        fake_docker_baseline = workspace / "docker" / "docker-resource-baseline.json"
        fake_docker_baseline.write_text(
            json.dumps({
                "schema_version": 1, "captured_at": "2026-04-30T00:00:00Z",
                "docker_available": True,
                "images": [], "volumes": [], "networks": [], "containers": [], "build_cache": [],
            }, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(workspace / "bin/manage-docker-resources.py"),
            "--workspace-dir",
            str(workspace),
            "--capture-baseline",
        ], plugin_root, "Refusing to overwrite existing Docker resource baseline")
        real_docker_clean_helper.write_text(
            "#!/usr/bin/env python3\n"
            "import sys\n"
            "print('deterministic selftest Docker helper failure', file=sys.stderr)\n"
            "raise SystemExit(1)\n",
            encoding="utf-8",
        )
        real_docker_clean_helper.chmod(0o755)
        # The real Docker environment likely has resources, so verify-clean --strict should fail
        # But if Docker is clean, this test would pass incorrectly. Use a fixture instead.
        fake_current = workspace / "docker" / "current-finalize-fixture.json"
        fake_current.write_text(
            json.dumps({
                "schema_version": 1, "captured_at": "2026-04-30T00:01:00Z",
                "docker_available": True,
                "images": [{"id": "sha256:leftover", "repository": "leftover", "tag": "latest",
                            "labels": {"org.zhulong.managed": "true",
                                       "org.zhulong.workspace": workspace.name}}],
                "volumes": [], "networks": [], "containers": [], "build_cache": [],
            }, indent=2),
            encoding="utf-8",
        )
        # We can't easily inject the current-file into the finalization gate's Docker check,
        # so test the Docker failure path by removing the baseline entirely
        real_baseline = workspace / "docker" / "docker-resource-baseline.json"
        backup_baseline = workspace / "docker" / "docker-resource-baseline.json.bak"
        real_baseline.rename(backup_baseline)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_with_confirmed_bundles",
        ], plugin_root, "completion authority chain")
        backup_baseline.rename(real_baseline)

        # Test 5a: completed_no_confirmed_findings fails when partial confirmed bundles exist
        reset_finalization_test_to_verification()
        # Remove the valid bundle rendered for Test 5
        for entry in (workspace / "confirmed").iterdir():
            if entry.is_dir() and not entry.name.startswith("."):
                shutil.rmtree(entry)
        partial_no_finding = workspace / "confirmed/C98-partial-no-finding"
        partial_no_finding.mkdir(parents=True, exist_ok=True)
        (partial_no_finding / "verification-evidence.json").write_text(
            json.dumps({"verification_status": "confirmed_in_docker"}, indent=2),
            encoding="utf-8",
        )
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_no_confirmed_findings",
        ], plugin_root, "partial confirmed bundle",
           extra_env={})
        shutil.rmtree(partial_no_finding)

        # Test 5b: completed_no_confirmed_findings fails when Docker cleanliness fails
        reset_finalization_test_to_verification()
        real_baseline2 = workspace / "docker" / "docker-resource-baseline.json"
        backup_baseline2 = workspace / "docker" / "docker-resource-baseline.json.bak2"
        real_baseline2.rename(backup_baseline2)
        run_expect_fail([
            sys.executable,
            str(plugin_root / "scripts/finalize_audit_workspace.py"),
            "--workspace-dir", str(workspace),
            "--result", "completed_no_confirmed_findings",
        ], plugin_root, "Docker cleanliness check failed")
        backup_baseline2.rename(real_baseline2)

        # Test 5c: missing audit-event writer must be visible, not silent.
        reset_finalization_test_to_verification()
        isolated_finalizer_dir = Path(tempdir) / "isolated-finalizer"
        isolated_finalizer_dir.mkdir(parents=True, exist_ok=True)
        isolated_finalizer = isolated_finalizer_dir / "finalize_audit_workspace.py"
        shutil.copy2(plugin_root / "scripts/finalize_audit_workspace.py", isolated_finalizer)
        shutil.copy2(plugin_root / "scripts/audit_state_io.py", isolated_finalizer_dir / "audit_state_io.py")
        shutil.copy2(plugin_root / "scripts/audit_text_safety.py", isolated_finalizer_dir / "audit_text_safety.py")
        shutil.copy2(plugin_root / "scripts/audit_transition_policy.py", isolated_finalizer_dir / "audit_transition_policy.py")
        shutil.copy2(plugin_root / "scripts/validate_audit_protocol.py", isolated_finalizer_dir / "validate_audit_protocol.py")
        shutil.copy2(plugin_root / "scripts/workspace_state.py", isolated_finalizer_dir / "workspace_state.py")
        shutil.copy2(plugin_root / "scripts/blocked_verification.py", isolated_finalizer_dir / "blocked_verification.py")
        shutil.copy2(plugin_root / "scripts/validate_candidate.py", isolated_finalizer_dir / "validate_candidate.py")
        shutil.copy2(plugin_root / "scripts/candidate_identity.py", isolated_finalizer_dir / "candidate_identity.py")
        shutil.copy2(plugin_root / "scripts/validate_verifier_verdict.py", isolated_finalizer_dir / "validate_verifier_verdict.py")
        shutil.copy2(plugin_root / "scripts/audit_disposition.py", isolated_finalizer_dir / "audit_disposition.py")
        workspace_writer = workspace / "bin/write-audit-event.py"
        hidden_workspace_writer = workspace / "bin/write-audit-event.py.hidden-for-selftest"
        workspace_writer.rename(hidden_workspace_writer)
        try:
            proc = subprocess.run([
                sys.executable,
                str(isolated_finalizer),
                "--workspace-dir", str(workspace),
                "--result", "completed_no_confirmed_findings",
            ], cwd=plugin_root, capture_output=True, text=True,
                env=os.environ.copy())
        finally:
            hidden_workspace_writer.rename(workspace_writer)
        if proc.returncode == 0:
            raise SystemExit("FAILED: finalization silently continued when its state-changing writer was missing")
        if "FINALIZATION FAILED: audit event writer not found" not in proc.stderr:
            raise SystemExit("FAILED: missing finalization writer did not produce a stable fatal error")

        claude_home = Path(tempdir) / "claude-home"
        claude_home.mkdir(parents=True, exist_ok=True)
        default_sync_output = run_capture([
            "bash",
            str(plugin_root / "scripts/sync_to_claude_skill.sh"),
            "--claude-skills-dir",
            str(claude_home / "skills"),
            "--keep-backups",
            "2",
        ], plugin_root)
        if "Prompt template sync:" not in default_sync_output or "skipped" not in default_sync_output:
            raise SystemExit("FAILED: sync_to_claude_skill.sh should skip external prompt template sync by default")
        if "Prompt template synced from canonical source:" in default_sync_output:
            raise SystemExit("FAILED: sync_to_claude_skill.sh wrote an external prompt template by default")
        prompt_template_output = Path(tempdir) / "prompt-template.md"
        run([
            "bash",
            str(plugin_root / "scripts/sync_to_claude_skill.sh"),
            "--claude-skills-dir",
            str(claude_home / "skills"),
            "--keep-backups",
            "2",
            "--prompt-template-output",
            str(prompt_template_output),
        ], plugin_root)
        if prompt_template_output.read_text(encoding="utf-8") != (
            plugin_root / "assets/references/claude-code-invocation-template.md"
        ).read_text(encoding="utf-8"):
            raise SystemExit("FAILED: sync_to_claude_skill.sh did not honor --prompt-template-output")
        installed_skill = claude_home / "skills" / "zhulong"
        require_installed_package_hygiene(installed_skill, "Claude installed skill")
        if not (installed_skill / "SKILL.md").exists():
            raise SystemExit("FAILED: Claude skill sync did not create SKILL.md")
        if (installed_skill / "AGENTS.md").exists():
            raise SystemExit("FAILED: Claude skill sync copied repo-root AGENTS.md into installed skill")
        if not (installed_skill / "scripts/check_security_tooling.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy scripts")
        if not (installed_skill / "scripts/check_docker_gate.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy check_docker_gate.sh")
        if not (installed_skill / "scripts/run_initial_probes.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy run_initial_probes.sh")
        if not (installed_skill / "scripts/run_verification_case.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy run_verification_case.sh")
        if not (installed_skill / "scripts/check_sandbox_preflight.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy check_sandbox_preflight.py")
        if not (installed_skill / "scripts/evidence_io.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy evidence_io.py")
        if not (installed_skill / "scripts/render_handoff_summary.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy render_handoff_summary.py")
        if not (installed_skill / "scripts/workspace_state.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy workspace_state.py")
        if not (installed_skill / "scripts/asr_start.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy asr_start.sh")
        if not (installed_skill / "scripts/resolve_skill_root.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy resolve_skill_root.sh")
        if not (installed_skill / "scripts/zhulong_audit.sh").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy zhulong_audit.sh")
        if not (installed_skill / "scripts/write_audit_event.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy write_audit_event.py")
        if not (installed_skill / "scripts/audit_text_safety.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy audit_text_safety.py")
        if not (installed_skill / "scripts/audit_transition_policy.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy audit_transition_policy.py")
        if not (installed_skill / "scripts/validate_workspace_state.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy validate_workspace_state.py")
        if not (installed_skill / "scripts/find_variant_candidates.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy find_variant_candidates.py")
        if not (installed_skill / "scripts/assert_finalized_workspace.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy assert_finalized_workspace.py")
        if not (installed_skill / "scripts/audit_disposition.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy audit_disposition.py")
        if not (installed_skill / "scripts/blocked_verification.py").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy blocked_verification.py")
        if not (installed_skill / "assets/tool-registry.json").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy assets")
        if not (installed_skill / "assets/references/python-library-audit-playbook.md").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy Python Library playbook")
        if not (installed_skill / "assets/references/php-swoole-audit-playbook.md").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy PHP/Swoole playbook")
        if not (installed_skill / "assets/references/docker-registry-fallbacks.example.json").exists():
            raise SystemExit("FAILED: Claude skill sync did not copy registry fallback example")
        run([
            sys.executable,
            str(installed_skill / "scripts/selftest_plugin.py"),
        ], installed_skill)
        backups = sorted((claude_home / "skills" / ".zhulong-backups").glob("zhulong.backup.*"))
        if len(backups) > 2:
            raise SystemExit("FAILED: sync_to_claude_skill.sh did not enforce backup retention")
        top_level_backups = sorted((claude_home / "skills").glob("zhulong.backup.*"))
        if top_level_backups:
            raise SystemExit("FAILED: sync_to_claude_skill.sh left loadable backups at skills root")

        codex_home = Path(tempdir) / "codex-home"
        codex_skills_dir = codex_home / "skills"
        codex_home.mkdir(parents=True, exist_ok=True)
        codex_sync_output = run_capture([
            "bash",
            str(plugin_root / "scripts/sync_to_codex_skill.sh"),
            "--codex-skills-dir",
            str(codex_skills_dir),
            "--keep-backups",
            "2",
        ], plugin_root)
        if "Codex skill synced successfully." not in codex_sync_output:
            raise SystemExit("FAILED: sync_to_codex_skill.sh did not report success")
        if "Installed skill directory:" not in codex_sync_output or str(codex_skills_dir / "zhulong") not in codex_sync_output:
            raise SystemExit("FAILED: sync_to_codex_skill.sh did not report installed path")
        codex_installed_skill = codex_skills_dir / "zhulong"
        require_installed_package_hygiene(codex_installed_skill, "Codex installed skill")
        if (codex_installed_skill / "AGENTS.md").exists():
            raise SystemExit("FAILED: Codex skill sync copied repo-root AGENTS.md into installed skill")
        for rel in INSTALLED_SKILL_REQUIRED_FILES:
            if not (codex_installed_skill / rel).exists():
                raise SystemExit(f"FAILED: Codex skill sync did not copy required installed file: {rel}")
        for rel in FORBIDDEN_INSTALLED_TOP_LEVEL:
            if (codex_installed_skill / rel).exists():
                raise SystemExit(f"FAILED: Codex skill sync copied forbidden top-level material: {rel}")
        require_text(
            codex_installed_skill / "docs/CODEX_SKILL_ADAPTATION.md",
            "Codex reads user skills from `$HOME/.agents/skills`",
            "installed Codex skill adaptation docs",
        )
        require_text(
            codex_installed_skill / "README.plugin-package.md",
            "Zhulong",
            "installed Codex package README",
        )
        require_text(
            codex_installed_skill / "INSTALL.plugin-package.md",
            "Install",
            "installed Codex package install notes",
        )
        for _ in range(3):
            run([
                "bash",
                str(plugin_root / "scripts/sync_to_codex_skill.sh"),
                "--codex-skills-dir",
                str(codex_skills_dir),
                "--keep-backups",
                "2",
            ], plugin_root)
        codex_backups = sorted((codex_skills_dir / ".zhulong-backups").glob("zhulong.backup.*"))
        if len(codex_backups) != 2:
            raise SystemExit("FAILED: sync_to_codex_skill.sh did not enforce backup retention")
        if not all(backup.parent.name.startswith(".") for backup in codex_backups):
            raise SystemExit("FAILED: sync_to_codex_skill.sh backup directories are not hidden")
        codex_top_level_backups = sorted(codex_skills_dir.glob("zhulong.backup.*"))
        if codex_top_level_backups:
            raise SystemExit("FAILED: sync_to_codex_skill.sh left loadable backups at skills root")
        run([
            sys.executable,
            str(codex_installed_skill / "scripts/selftest_plugin.py"),
        ], codex_installed_skill)

    print(f"SELFTEST PASSED: {plugin_root}")


if __name__ == "__main__":
    main()
